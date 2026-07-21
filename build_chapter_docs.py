"""
Builds one Word doc per chapter from raw extracted text.
Cleans boilerplate, strips anatomy/physiology figure captions,
keeps clinical content (tables, criteria, management, dosing).
Structures each doc with consistent clinical sections.
"""
import sys, os, re, traceback
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(line_buffering=True)

IN_DIR  = r"C:\Users\JasonOng\Desktop\local docs\personal\rach"
OUT_DIR = r"C:\Users\JasonOng\Desktop\local docs\personal\rach"
LOG     = r"c:\Users\JasonOng\AST_Agent\docbuild_log.txt"

def log(msg):
    print(msg, flush=True)
    with open(LOG, "a", encoding="utf-8") as f:
        f.write(msg + "\n")

CHAPTERS = [
    ("Ch100_Anatomy_SmallLargeIntestine.txt",         "Ch 100: Small & Large Intestine — Anatomy and Developmental Anomalies"),
    ("Ch101_SmallIntestinal_Motor_Sensory.txt",        "Ch 101: Small Intestinal Motor and Sensory Function"),
    ("Ch102_Colonic_Motor_Sensory.txt",                "Ch 102: Colonic Motor and Sensory Function"),
    ("Ch103_Electrolyte_Absorption_Secretion.txt",     "Ch 103: Intestinal Electrolyte Absorption and Secretion"),
    ("Ch104_Digestion_Absorption_Macro.txt",           "Ch 104: Digestion and Absorption of Carbohydrate, Protein, and Fat"),
    ("Ch105_Micronutrients_Absorption.txt",            "Ch 105: Digestion and Absorption of Micronutrients"),
    ("Ch106_Maldigestion_Malabsorption.txt",           "Ch 106: Maldigestion and Malabsorption"),
    ("Ch107_Small_Intestinal_Bacterial_Overgrowth.txt","Ch 107: Small Intestinal Bacterial Overgrowth (SIBO)"),
    ("Ch108_Short_Bowel_Syndrome.txt",                 "Ch 108: Short Bowel Syndrome"),
    ("Ch109_Celiac_Disease.txt",                       "Ch 109: Celiac Disease"),
    ("Ch110_Tropical_Diarrhea_Malabsorption.txt",      "Ch 110: Tropical Diarrhea and Malabsorption"),
    ("Ch111_Whipple_Disease.txt",                      "Ch 111: Whipple Disease"),
    ("Ch112_Infectious_Enteritis_Proctocolitis.txt",   "Ch 112: Infectious Enteritis and Proctocolitis"),
    ("Ch113_Food_Poisoning.txt",                       "Ch 113: Food Poisoning"),
    ("Ch114_Cdiff_AAD.txt",                            "Ch 114: C. difficile Infection and Antibiotic-Associated Diarrhea"),
    ("Ch115_Intestinal_Protozoa.txt",                  "Ch 115: Intestinal Protozoa"),
    ("Ch116_Intestinal_Worms.txt",                     "Ch 116: Intestinal Worms"),
    ("Ch117_IBD_Epidemiology_Diagnosis.txt",           "Ch 117: IBD — Epidemiology, Pathogenesis, and Diagnosis"),
    ("Ch118_IBD_Management.txt",                       "Ch 118: IBD — Management"),
    ("Ch119_Ileostomies_Colostomies.txt",              "Ch 119: Ileostomies, Colostomies, and Anastomoses"),
    ("Ch120_Intestinal_Ischemia.txt",                  "Ch 120: Intestinal Ischemia"),
    ("Ch121_Intestinal_Ulcerations.txt",               "Ch 121: Intestinal Ulcerations"),
    ("Ch122_Appendicitis.txt",                         "Ch 122: Appendicitis"),
    ("Ch123_Diverticular_Disease.txt",                 "Ch 123: Diverticular Disease of the Colon"),
    ("Ch124_IBS.txt",                                  "Ch 124: Irritable Bowel Syndrome (IBS)"),
    ("Ch125_Intestinal_Obstruction.txt",               "Ch 125: Intestinal Obstruction"),
    ("Ch126_Ileus_PseudoObstruction.txt",              "Ch 126: Ileus and Pseudo-Obstruction Syndromes"),
    ("Ch127_Small_Bowel_Tumors.txt",                   "Ch 127: Tumors of the Small Intestine"),
    ("Ch128_Colonic_Polyps_Polyposis.txt",             "Ch 128: Colonic Polyps and Polyposis Syndromes"),
    ("Ch129_Colorectal_Cancer.txt",                    "Ch 129: Colorectal Cancer"),
    ("Ch130_Other_Colon_Diseases.txt",                 "Ch 130: Other Diseases of the Colon"),
    ("Ch131_Anal_Diseases.txt",                        "Ch 131: Anal Diseases"),
]

# Sections that are primarily anatomy/physiology (skip figure captions for these)
SKIP_FIGURE_KEYWORDS = [
    'pathogenesis', 'mechanism', 'embryology', 'anatomy', 'histology',
    'structure', 'pathway', 'signaling', 'receptor', 'gene', 'chromosome',
    'molecular', 'cellular', 'ultrastructure', 'intestinal epithelium',
]

def is_skip_figure(caption):
    low = caption.lower()
    return any(kw in low for kw in SKIP_FIGURE_KEYWORDS)

def clean_raw(text):
    # Remove full navigation sidebar block (from "Elsevier" down to "Book Page Loaded")
    text = re.sub(
        r'Elsevier\s*\nSkip to main content.*?Book Page Loaded\s*\n',
        '\n', text, flags=re.DOTALL
    )
    # Remove page markers from extraction script
    text = re.sub(r'CHAPTER:[^\n]*\n', '', text)
    text = re.sub(r'Pages extracted:[^\n]*\n', '', text)
    text = re.sub(r'={3,}\n', '', text)
    text = re.sub(r'[—\-]{10,}\n', '', text)
    text = re.sub(r'Page \d+\s*\n', '', text)
    text = re.sub(r'--- Reader page[^\n]*---\n', '', text)

    # Remove figure captions for anatomy/physiology figures; keep clinical figures
    def handle_fig(m):
        caption = m.group(0)
        if is_skip_figure(caption):
            return ''
        # Keep clinical figures but shorten them
        first_sentence = caption.split('.')[0] + '.' if '.' in caption else caption[:120]
        return first_sentence.strip() + '\n'
    text = re.sub(r'FIG\.\s*\d+[^\n]{0,600}\n', handle_fig, text)
    text = re.sub(r'Follow for extended description\s*\n', '', text)

    # Remove navigation artifacts
    for pat in [
        r'Open/Close Margin\s*\n', r'Bookmark page\s*\n',
        r'Back to Page[^\n]*\n', r'Go to Page\s*\n',
        r'/ \d+\s*\n', r'Previous\s*\n', r'Next\s*\n',
        r'More Options\s*\n', r'Reader Preferences\s*\n',
        r'Search across book\s*\n', r'More book options\s*\n',
        r'Highlights, Notes, Bookmarks[^\n]*\n',
        r'Table of Contents\s*\n', r'Go to First Page\s*\n',
        r'Close\s*\n', r'Skip to [^\n]+\n', r'Back\s*\n',
    ]:
        text = re.sub(pat, '', text)

    # Collapse excess blank lines
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()

def split_into_blocks(text):
    """Split cleaned text into labelled blocks: TABLE, HEADING, BODY."""
    blocks = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Detect table heading (e.g. "Table 114.3")
        if re.match(r'^Table\s+\d+[\.\d]*\s*$', stripped):
            title = stripped
            i += 1
            # Skip blank line
            while i < len(lines) and not lines[i].strip():
                i += 1
            # Gather table title line
            table_title = ''
            if i < len(lines) and lines[i].strip() and '\t' not in lines[i] and len(lines[i].strip()) < 120:
                table_title = lines[i].strip()
                i += 1
            # Collect table rows (tab-separated or multi-column)
            table_lines = []
            while i < len(lines):
                l = lines[i]
                if not l.strip():
                    i += 1
                    break
                table_lines.append(l.rstrip())
                i += 1
            blocks.append(('TABLE', title, table_title, table_lines))
            continue

        # Section heading heuristic: short line, no period, often all-title-case or uppercase
        if (len(stripped) < 80 and not stripped.endswith(',')
                and not re.search(r'\d{1,3}\.\d', stripped)
                and stripped.isupper()):
            blocks.append(('HEADING', stripped, None, None))
            i += 1
            continue

        # Regular body text
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip():
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(('BODY', ' '.join(para_lines), None, None))

    return blocks

# ── Word helpers ────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(14)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(12)
    return p

def add_body_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def add_bullet_para(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_table_block(doc, table_ref, table_title, table_lines):
    # Table reference heading
    p = doc.add_paragraph()
    run = p.add_run(f"{table_ref}" + (f" — {table_title}" if table_title else ""))
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if not table_lines:
        return

    # Parse rows — try tab-separated first, then whitespace-split with 2+ spaces
    parsed_rows = []
    for line in table_lines:
        if '\t' in line:
            parsed_rows.append([c.strip() for c in line.split('\t')])
        else:
            parts = re.split(r'\s{2,}', line.strip())
            parsed_rows.append(parts if len(parts) > 1 else [line.strip()])

    if not parsed_rows:
        return

    max_cols = max(len(r) for r in parsed_rows)
    if max_cols < 2:
        for r in parsed_rows:
            add_bullet_para(doc, r[0])
        return

    t = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    t.style = 'Table Grid'
    for ri, row in enumerate(parsed_rows):
        for ci in range(max_cols):
            cell = t.cell(ri, ci)
            cell_text = row[ci] if ci < len(row) else ''
            cell.text = cell_text
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            if ri == 0:
                set_cell_bg(cell, 'BDD7EE')
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

CLINICAL_SECTION_KEYWORDS = {
    'CLINICAL FEATURES': ['clinical feature', 'sign', 'symptom', 'presentation', 'manifestation'],
    'DIAGNOSIS': ['diagnosis', 'diagnostic', 'criteria', 'test', 'investigation', 'workup', 'laboratory', 'imaging', 'endoscop'],
    'MANAGEMENT': ['management', 'treatment', 'therapy', 'therapeutic', 'drug', 'medic', 'surgery', 'surgical', 'dose', 'mg'],
    'COMPLICATIONS': ['complication', 'adverse', 'risk', 'prognosis', 'outcome', 'mortality', 'morbidity'],
    'EPIDEMIOLOGY': ['epidemiol', 'incidence', 'prevalence', 'age', 'sex', 'gender', 'race', 'geographic'],
}

def classify_section(heading):
    low = heading.lower()
    for label, keywords in CLINICAL_SECTION_KEYWORDS.items():
        if any(kw in low for kw in keywords):
            return label
    return None

def build_doc(title, raw_text, out_path):
    cleaned = clean_raw(raw_text)
    blocks  = split_into_blocks(cleaned)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ──
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(18)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Sleisenger & Fordtran's GI and Liver Disease  |  "
        "Clinical Reference — Gastro Residency (Singapore)"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    current_section = None

    for block in blocks:
        kind = block[0]

        if kind == 'TABLE':
            _, ref, t_title, t_lines = block
            add_table_block(doc, ref, t_title, t_lines)

        elif kind == 'HEADING':
            heading_text = block[1]
            classified = classify_section(heading_text)
            if classified and classified != current_section:
                current_section = classified
                add_h1(doc, classified)
            add_h2(doc, heading_text.title())

        elif kind == 'BODY':
            body_text = block[1]
            # Skip if this is a repeat of navigation boilerplate
            if any(nav in body_text for nav in ['Skip to main content', 'Book Page Loaded', 'Table of Contents sections']):
                continue
            # Detect inline bullet-style lines
            if body_text.startswith(('•', '-', '*')):
                add_bullet_para(doc, body_text.lstrip('•-* '))
            else:
                add_body_para(doc, body_text)

    doc.save(out_path)

def main():
    with open(LOG, "w", encoding="utf-8") as f:
        f.write("=== build_chapter_docs.py ===\n")
    log(f"Building {len(CHAPTERS)} Word docs -> {OUT_DIR}\n")

    done, failed = 0, []

    for idx, (filename, title) in enumerate(CHAPTERS):
        in_path  = os.path.join(IN_DIR, filename)
        out_path = os.path.join(OUT_DIR, filename.replace('.txt', '.docx'))

        log(f"[{idx+1:02d}/{len(CHAPTERS)}] {title}")
        try:
            with open(in_path, 'r', encoding='utf-8') as f:
                raw = f.read()
            build_doc(title, raw, out_path)
            log(f"  -> Saved: {out_path}")
            done += 1
        except Exception as e:
            log(f"  -> FAILED: {e}")
            log(traceback.format_exc())
            failed.append(filename)

    log(f"\n=== Done: {done}/{len(CHAPTERS)} ===")
    if failed:
        log("Failed: " + ', '.join(failed))

main()
