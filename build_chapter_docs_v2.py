"""
build_chapter_docs_v2.py
Generates 32 Word documents for Sleisenger & Fordtran's GI and Liver Disease
Chapters 100–131 (Part XI: Small and Large Intestine)
Clinical reference notes for Singapore gastroenterology residency (Rachel Teo)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_DIR = r"C:\Users\JasonOng\Desktop\local docs\personal\rach"

# ── Colour constants ──────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
MED_BLUE    = RGBColor(0x2E, 0x74, 0xB5)
PEARL_BG    = "DEEAF1"   # light blue shading for pearls
TABLE_HDR   = "BDD7EE"   # table header shading
SUBTITLE_GREY = RGBColor(0x80, 0x80, 0x80)
HR_COLOUR   = "B4C7E7"   # horizontal rule colour

SUBTITLE_TEXT = "Sleisenger & Fordtran's GI and Liver Disease | Clinical Reference — Singapore GI Residency"


# ── Formatting helpers ────────────────────────────────────────────────────────

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def shade_paragraph(para, hex_colour):
    """Apply solid background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    pPr.append(shd)


def add_bottom_border(para, hex_colour="B4C7E7"):
    """Add a light-blue bottom border to a paragraph (acts as HR)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size  = Pt(16)
    run.font.color.rgb = DARK_BLUE
    run.font.bold  = True
    para.paragraph_format.space_after = Pt(2)
    return para


def add_subtitle(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(SUBTITLE_TEXT)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = SUBTITLE_GREY
    para.paragraph_format.space_after = Pt(10)
    return para


def add_h1(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(3)
    add_bottom_border(para, HR_COLOUR)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = MED_BLUE
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(2)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    para.paragraph_format.space_after = Pt(2)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    para.paragraph_format.space_after = Pt(1)
    return para


def add_pearl(doc, text):
    """Shaded light-blue pearl bullet."""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(2)
    shade_paragraph(para, PEARL_BG)
    return para


def add_pearls_section(doc, pearls):
    add_h1(doc, "Key Clinical Pearls")
    for p in pearls:
        add_pearl(doc, p)


def add_table(doc, headers, rows):
    """Add a styled table with shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
                run.font.bold = True
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  TABLE_HDR)
        tcPr.append(shd)
    # data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            cells[c_idx].text = cell_text
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()  # spacing after table


def build_doc(chapter_data):
    doc = Document()
    set_margins(doc)
    add_title(doc, chapter_data['title'])
    add_subtitle(doc)
    add_pearls_section(doc, chapter_data['pearls'])
    for section in chapter_data['sections']:
        add_h2(doc, section['heading'])
        for item in section['content']:
            if item['type'] == 'bullet':
                add_bullet(doc, item['text'], item.get('level', 0))
            elif item['type'] == 'body':
                add_body(doc, item['text'])
            elif item['type'] == 'table':
                add_table(doc, item['headers'], item['rows'])
            elif item['type'] == 'h2':
                add_h2(doc, item['text'])
    return doc


def b(text, level=0):
    return {'type': 'bullet', 'text': text, 'level': level}

def bd(text):
    return {'type': 'body', 'text': text}

def tbl(headers, rows):
    return {'type': 'table', 'headers': headers, 'rows': rows}

def h2(text):
    return {'type': 'h2', 'text': text}


# ── Chapter data ──────────────────────────────────────────────────────────────

CHAPTERS = []

# ── Ch100 ──
CHAPTERS.append({
    'filename': 'Ch100_Anatomy_SmallLargeIntestine.docx',
    'title': 'Ch100 — Anatomy & Developmental Anomalies of the Small and Large Intestine',
    'pearls': [
        "Meckel's diverticulum — rule of 2s: 2% population, 2 feet from ileocaecal valve, 2 inches long, 2:1 M:F, presents in first 2 years of life.",
        "Hirschsprung disease = aganglionosis (absent Auerbach & Meissner plexuses); hallmark is failure to pass meconium within 48 h of birth.",
        "Malrotation risk: midgut volvulus — a surgical emergency presenting with bilious vomiting in a neonate.",
        "Intestinal atresia is the most common cause of neonatal bowel obstruction.",
    ],
    'sections': [
        {'heading': "Meckel's Diverticulum",
         'content': [
             b("True diverticulum; remnant of omphalomesenteric (vitelline) duct; antimesenteric border of ileum."),
             b("Complications: GI bleeding (ectopic gastric mucosa secretes acid → ileal ulceration), intestinal obstruction (band adhesion, intussusception, volvulus around Meckel's), Meckelitis (mimics appendicitis), perforation."),
             b("Diagnosis of bleeding: Technetium-99m pertechnetate scan detects ectopic gastric mucosa (sensitivity ~85% in children, lower in adults)."),
             b("Treatment: symptomatic Meckel's → surgical resection (open or laparoscopic); incidental finding in adults — conservative if no symptoms."),
         ]},
        {'heading': "Hirschsprung Disease",
         'content': [
             b("Absence of ganglion cells in Auerbach (myenteric) and Meissner (submucosal) plexuses — aganglionic segment fails to relax."),
             b("Short-segment (rectosigmoid, 75%) vs. long-segment (proximal to sigmoid) vs. total colonic aganglionosis."),
             b("Presentation: failure to pass meconium in 48 h; abdominal distension; ribbon stools; older children with chronic constipation and growth failure."),
             b("Diagnosis: suction rectal biopsy — absence of ganglion cells + acetylcholinesterase staining of hypertrophied nerve trunks. Contrast enema shows transition zone."),
             b("Treatment: surgical resection of aganglionic segment with pull-through of ganglionic bowel to anus (Swenson, Soave, Duhamel procedures)."),
             b("Complication: Hirschsprung-associated enterocolitis (HAEC) — life-threatening; rectal irrigation + antibiotics + urgent surgery."),
         ]},
        {'heading': "Malrotation",
         'content': [
             b("Normal rotation of midgut is 270° counterclockwise; malrotation leaves DJ flexure to the right of midline → narrow mesenteric base → predisposes to midgut volvulus."),
             b("Ladd's bands: peritoneal bands from abnormally positioned caecum crossing duodenum → duodenal obstruction."),
             b("Presentation: bilious vomiting in neonate; malrotation can present at any age (10–15% in adults) with intermittent obstruction."),
             b("Diagnosis: upper GI contrast series — DJ flexure positioned to the right or inferiorly. CT for adults."),
             b("Treatment: Ladd's procedure — volvulus reduction (counter-clockwise), division of Ladd's bands, appendicectomy, broadening of mesenteric base. Surgical emergency if volvulus present."),
         ]},
        {'heading': "Intestinal Duplications",
         'content': [
             b("Cystic or tubular structures adjacent to native bowel; most common in ileum."),
             b("Share blood supply and often muscular wall with adjacent bowel."),
             b("Complications: obstruction, bleeding (if ectopic gastric mucosa), intussusception, infection."),
             b("Treatment: surgical excision."),
         ]},
    ],
})

# ── Ch101 ──
CHAPTERS.append({
    'filename': 'Ch101_SmallIntestinalMotorSensory.docx',
    'title': 'Ch101 — Small Intestinal Motor and Sensory Function',
    'pearls': [
        "Migrating motor complex (MMC) cycles every 90–120 min in the fasting state; Phase III 'housekeeping wave' sweeps bacteria aborally — disruption leads to SIBO.",
        "Small bowel transit time is normally 2–4 hours.",
        "Visceral hypersensitivity (lowered pain threshold to luminal distension) is the central mechanism in IBS.",
        "90% of the body's serotonin (5-HT) resides in enterochromaffin cells of the gut — critical neurotransmitter for motility.",
    ],
    'sections': [
        {'heading': "Fasting and Fed Motor Patterns",
         'content': [
             b("MMC phases: Phase I (motor quiescence) → Phase II (irregular contractions) → Phase III (regular high-amplitude contractions, 'housekeeper') → Phase I again."),
             b("Fed pattern: loss of MMC; sustained irregular mixing contractions driven by nutrients + CCK, GLP-1, peptide YY."),
             b("Disruption of MMC → bacterial stasis → SIBO (common in DM autonomic neuropathy, scleroderma, post-surgical blind loops)."),
         ]},
        {'heading': "Visceral Hypersensitivity",
         'content': [
             b("Afferent nerves sensitised at the level of the gut (peripheral sensitisation) and spinal cord/brain (central sensitisation)."),
             b("Results in pain at lower luminal pressures/volumes than normal."),
             b("Key mechanism in IBS, functional dyspepsia, and other functional GI disorders."),
             b("Modulated by psychological state — explains benefit of antidepressants and psychological therapies."),
         ]},
        {'heading': "Clinical Approach to Small Bowel Dysmotility",
         'content': [
             b("Symptoms: bloating, nausea, early satiety, abdominal distension, diarrhoea or constipation, SIBO."),
             b("Investigations: gastric emptying scintigraphy (solid meal); antroduodenal manometry (gold standard for motility disorder — identifies neuropathic vs myopathic pattern); hydrogen breath test for SIBO; small bowel transit with barium/scintigraphy."),
             b("Management: dietary modification (small frequent meals, low-fat, low-fibre for gastroparesis); prokinetics — metoclopramide 10mg TDS AC, domperidone 10mg TDS AC, erythromycin 250mg TDS (motilin agonist); treat SIBO if present; nutritional support for severe cases."),
         ]},
    ],
})

# ── Ch102 ──
CHAPTERS.append({
    'filename': 'Ch102_ColonicMotorSensory.docx',
    'title': 'Ch102 — Colonic Motor and Sensory Function',
    'pearls': [
        "Normal colonic transit time: 24–48 hours. Slow transit constipation (>72 h) vs outlet obstruction — different management pathways.",
        "Mass movements (1–3/day) propel stool from right colon to sigmoid — absent or infrequent in slow transit constipation.",
        "Chagas disease (Trypanosoma cruzi) destroys myenteric plexus → acquired megacolon; endemic in Latin America.",
        "Hirschsprung disease in adults is rare but exists — diagnosed by rectal biopsy showing aganglionosis.",
    ],
    'sections': [
        {'heading': "Constipation Types and Investigation",
         'content': [
             b("Slow transit constipation: infrequent urge, no straining; colonic transit study with radiopaque markers — retention throughout colon."),
             b("Outlet obstruction / dyssynergia (anismus): straining, incomplete evacuation, feeling of blockage; defaecography shows paradoxical puborectalis contraction; anorectal manometry shows failure of sphincter relaxation on straining."),
             b("Normal transit with IBS-C: normal transit time, predominant symptom is pain relieved by defaecation."),
             tbl(["Test", "What it measures", "Use for"],
                 [["Colonic transit (radiopaque markers)", "Global transit time (normal <72h for >80% of markers to pass)", "Slow transit constipation"],
                  ["Defaecography / MRI proctography", "Pelvic floor dynamics, rectal prolapse, rectocele", "Outlet obstruction, SRUS"],
                  ["Anorectal manometry", "Sphincter pressures, RAIR, rectal sensation", "Hirschsprung, dyssynergia, incontinence"],
                  ["Balloon expulsion test", "Ability to expel 50mL water-filled balloon", "Screening for outlet obstruction"]]),
         ]},
        {'heading': "Diarrhoea Mechanisms",
         'content': [
             b("Secretory: active ion secretion (Cl-) or inhibited Na absorption; large-volume watery diarrhoea; persists with fasting; osmotic gap <50 mOsm/kg."),
             b("Osmotic: poorly absorbable solutes retain water; stops with fasting; osmotic gap >125 mOsm/kg."),
             b("Motility: rapid transit reduces contact time; typically not large volume."),
             b("Inflammatory/exudative: mucosal damage → blood, mucus, protein loss; stool WBC positive."),
         ]},
        {'heading': "Faecal Incontinence",
         'content': [
             b("Causes: internal anal sphincter (IAS) defect (obstetric, surgical), external anal sphincter (EAS) defect, pudendal neuropathy (obstetric, DM), rectal hyposensitivity, diarrhoea, overflow (severe constipation)."),
             b("Investigations: endoanal ultrasound (EAUS) — gold standard for sphincter anatomy; anorectal manometry — resting/squeeze pressures; pudendal nerve terminal motor latency; MRI pelvis for complex cases."),
             b("Treatment: dietary modification, loperamide to firm stool; biofeedback therapy — strengthens EAS, improves rectal sensation (moderate evidence); overlapping sphincteroplasty for anterior defect; sacral nerve stimulation (SNS) — effective even without sphincter defect — improves continence by modulating sacral afferents; colostomy as last resort."),
         ]},
    ],
})

# ── Ch103 ──
CHAPTERS.append({
    'filename': 'Ch103_ElectrolyteAbsorptionSecretion.docx',
    'title': 'Ch103 — Intestinal Electrolyte Absorption and Secretion',
    'pearls': [
        "Stool osmotic gap = 290 − 2 × (stool Na + stool K): >125 mOsm/kg = osmotic diarrhoea; <50 mOsm/kg = secretory diarrhoea.",
        "Secretory diarrhoea persists with fasting (large volume, watery); osmotic diarrhoea stops with fasting.",
        "Cholera toxin is the prototype of secretory diarrhoea — cAMP-mediated Cl⁻ secretion via CFTR.",
        "ORS works because Na-glucose co-transporter (SGLT1) is intact even in cholera — glucose drives Na and water absorption.",
    ],
    'sections': [
        {'heading': "Stool Osmotic Gap — Practical Use",
         'content': [
             b("Formula: 290 − 2 × (stool [Na] + stool [K]). Stool osmolality assumed 290 mOsm/kg."),
             b(">125: osmotic — caused by unabsorbed solutes (Mg, sorbitol, lactulose, lactose in lactase deficiency)."),
             b("<50: secretory — caused by active secretion or blocked absorption (toxins, neuroendocrine tumours, bile acid malabsorption, congenital Cl diarrhoea)."),
             b("50–125: indeterminate — may be mixed."),
         ]},
        {'heading': "Causes of Secretory Diarrhoea",
         'content': [
             b("VIPoma (Verner-Morrison): watery diarrhoea, hypokalaemia, achlorhydria (WDHA/pancreatic cholera); VIP measured in plasma; octreotide treatment."),
             b("Carcinoid / NET: serotonin-mediated; flushing + diarrhoea; urine 5-HIAA elevated."),
             b("Zollinger-Ellison syndrome: gastrin hypersecretion → acid inactivates pancreatic lipase → secretory + osmotic component."),
             b("Microscopic colitis: collagenous or lymphocytic; watery diarrhoea in older women."),
             b("Bile acid malabsorption (BAM): ileal resection/disease → bile acids reach colon → secretory diarrhoea; SeHCAT scan diagnoses; cholestyramine or colesevelam treatment."),
             b("Congenital chloride diarrhoea: rare; Cl-HCO3 exchanger defect; high stool Cl-, metabolic alkalosis."),
         ]},
        {'heading': "Causes of Osmotic Diarrhoea",
         'content': [
             b("Lactose intolerance: secondary lactase deficiency (after gastroenteritis, coeliac, Crohn's); hydrogen breath test with lactose load."),
             b("Sorbitol, fructose: sugar-free gums, fruit juice, high-fructose corn syrup."),
             b("Lactulose, polyethylene glycol, Mg-containing antacids — iatrogenic osmotic diarrhoea."),
         ]},
        {'heading': "Oral Rehydration Therapy (ORT)",
         'content': [
             b("WHO ORS (2003): Na 75 mmol/L, glucose 75 mmol/L, K 20 mmol/L, citrate 10 mmol/L, Cl 65 mmol/L; osmolality 245 mOsm/L (reduced from original 311)."),
             b("Glucose enhances Na absorption via SGLT1 co-transporter — remains functional even in secretory diarrhoea."),
             b("Reduced-osmolality ORS reduces stool output and duration vs original ORS; suitable for all ages."),
             b("Rice-based ORS: provides small glucose polymers — equivalent or superior efficacy in cholera."),
             b("IV rehydration for severe dehydration, ileus, or shock; transition to ORT as soon as tolerated."),
         ]},
    ],
})

# ── Ch104 ──
CHAPTERS.append({
    'filename': 'Ch104_DigestionAbsorptionMacronutrients.docx',
    'title': 'Ch104 — Digestion and Absorption of Carbohydrates, Proteins, and Fat',
    'pearls': [
        "Fat malabsorption (steatorrhoea) is clinically the most significant — leads to fat-soluble vitamin deficiency (A, D, E, K) and weight loss.",
        "Pancreatic lipase requires bile salts to form mixed micelles for fat absorption.",
        "Sudan III stain of stool for fat (qualitative); 72-hour faecal fat collection >7g/day = steatorrhoea (gold standard, quantitative).",
        "D-xylose test distinguishes mucosal disease (low D-xylose absorption) from luminal/pancreatic disease (normal D-xylose).",
        "Bile acid malabsorption causes watery secretory diarrhoea — especially after ileal resection or disease.",
    ],
    'sections': [
        {'heading': "Fat Digestion and Absorption",
         'content': [
             b("Steps: lingual/gastric lipase (minimal) → pancreatic lipase + colipase (duodenum, major step) → bile salt mixed micelles solubilise fatty acids + monoglycerides → enterocyte uptake → re-esterified to triglycerides → packaged into chylomicrons → lymphatics (thoracic duct) → systemic circulation."),
             b("Disruption causes: pancreatic exocrine insufficiency (insufficient lipase), bile acid deficiency (cholestatic liver disease, ileal resection, bacterial deconjugation in SIBO), mucosal disease (coeliac, Crohn's), lymphatic obstruction (Whipple's disease, lymphangiectasia)."),
             b("Clinical: steatorrhoea (pale, bulky, floating, offensive stools), fat-soluble vitamin deficiency, weight loss, hypoalbuminaemia."),
         ]},
        {'heading': "Carbohydrate Absorption",
         'content': [
             b("Salivary + pancreatic amylase → oligosaccharides → brush border disaccharidases (lactase, sucrase-isomaltase, maltase) → monosaccharides → absorbed via SGLT1 (glucose, galactose) and GLUT5 (fructose)."),
             b("Lactase is the most vulnerable disaccharidase — congenital alactasia (rare) or secondary acquired lactase deficiency after mucosal injury."),
             b("Unabsorbed carbohydrates → colonic fermentation → bloating, flatulence, osmotic diarrhoea (high osmotic gap)."),
         ]},
        {'heading': "Protein Digestion",
         'content': [
             b("Pepsin (gastric, pH-dependent) → pancreatic proteases (trypsin, chymotrypsin, elastase — activated by enterokinase) → amino acids and small peptides → enterocyte absorption."),
             b("Isolated protein malabsorption is rare clinically; usually occurs alongside fat malabsorption."),
             b("Protein-losing enteropathy: mucosal disease or lymphatic obstruction → faecal protein loss → hypoalbuminaemia, oedema. Alpha-1-antitrypsin faecal clearance diagnoses."),
         ]},
        {'heading': "Investigating Malabsorption",
         'content': [
             tbl(["Test", "Purpose", "Interpretation"],
                 [["72h faecal fat", "Quantify steatorrhoea", ">7g/day = abnormal (gold standard)"],
                  ["Faecal elastase-1", "Pancreatic exocrine function", "<100 µg/g = severe, 100-200 = moderate PEI"],
                  ["D-xylose test", "Small bowel mucosal absorption", "Low urine xylose (<4g/5h) = mucosal disease"],
                  ["Sudan III stool stain", "Qualitative fat malabsorption", "Positive = steatorrhoea screen"],
                  ["H2 breath test (lactose/glucose)", "Lactose intolerance / SIBO", "Early H2 rise = SIBO; late rise = carbohydrate malabsorption"]]),
         ]},
    ],
})

# ── Ch105 ──
CHAPTERS.append({
    'filename': 'Ch105_DigestionAbsorptionMicronutrients.docx',
    'title': 'Ch105 — Digestion and Absorption of Micronutrients',
    'pearls': [
        "Vitamin B12 absorption requires intrinsic factor (IF) secreted by gastric parietal cells — deficiency from pernicious anaemia, total gastrectomy, ileal disease/resection.",
        "Iron is absorbed as Fe²⁺ in the duodenum — deficiency anaemia is the most common nutritional anaemia globally and the most common GI indication for endoscopy.",
        "Zinc deficiency → acrodermatitis enteropathica: perioral and perianal rash, alopecia, diarrhoea — seen in IBD, post-ileostomy.",
        "Vitamin D deficiency is common in Singapore despite tropical sun exposure (indoor lifestyle, sunscreen use).",
    ],
    'sections': [
        {'heading': "Fat-Soluble Vitamins (A, D, E, K)",
         'content': [
             b("All require bile salts and intact fat absorption; deficiency in cholestatic liver disease, pancreatic exocrine insufficiency, short bowel syndrome."),
             b("Vitamin A: night blindness (early), xerophthalmia, keratomalacia; replace with oral vitamin A supplementation."),
             b("Vitamin D: osteomalacia (adults), rickets (children), hypocalcaemia, proximal myopathy; replace with cholecalciferol 800–2000 IU/day or higher loading doses."),
             b("Vitamin E: peripheral neuropathy (spinocerebellar), haemolytic anaemia, retinopathy in abetalipoproteinaemia."),
             b("Vitamin K: coagulopathy — prolonged PT corrected by IV/SC vitamin K (not by FFP alone); differentiates from liver synthetic failure (PT not corrected by vitamin K)."),
         ]},
        {'heading': "Water-Soluble Vitamins",
         'content': [
             b("Vitamin B12 (cobalamin): deficiency → megaloblastic anaemia, subacute combined degeneration of spinal cord (posterior and lateral column demyelination); causes: pernicious anaemia, gastrectomy, ileal resection, Crohn's ileitis, SIBO (bacteria consume B12), fish tapeworm; diagnosis: low serum B12, elevated MMA and homocysteine; treatment: IM hydroxocobalamin 1mg on alternate days × 6 doses then monthly (or high-dose oral 1mg/day if no IF deficiency)."),
             b("Folate: deficiency → megaloblastic anaemia (without neurological features); causes: poor intake, pregnancy, malabsorption, methotrexate; treat with folic acid 5mg/day."),
             b("Thiamine (B1): Wernicke's encephalopathy (confusion, ophthalmoplegia, ataxia) and Korsakoff syndrome; give thiamine 200–500mg IV TDS before any glucose in at-risk patients."),
             b("Vitamin C (ascorbic acid): scurvy — perifollicular haemorrhage, corkscrew hairs, bleeding gums, poor wound healing; treat with ascorbic acid 500mg TDS."),
         ]},
        {'heading': "Clinically Important Minerals",
         'content': [
             b("Iron: non-haem iron reduced to Fe²⁺ by gastric acid and duodenal ferric reductase; absorbed in duodenum/proximal jejunum; iron deficiency anaemia is universal indication for upper + lower GI endoscopy to exclude occult bleeding, especially in men and post-menopausal women."),
             b("Zinc: essential for wound healing, immune function, gut integrity; deficient in chronic diarrhoea, IBD, Crohn's, post-ileostomy; supplement with zinc gluconate/sulphate."),
             b("Magnesium: hypomagnesaemia with chronic PPI use (impairs TRPM6 channel), short bowel syndrome, prolonged diarrhoea; can cause hypokalaemia and hypocalcaemia refractory to replacement unless Mg corrected first."),
             b("Calcium: absorption enhanced by vitamin D; depletion in fat malabsorption, short bowel (especially colon-in-continuity with hyperoxaluria → oxalate renal stones)."),
         ]},
    ],
})

# ── Ch106 ──
CHAPTERS.append({
    'filename': 'Ch106_MaldigestionMalabsorption.docx',
    'title': 'Ch106 — Maldigestion and Malabsorption',
    'pearls': [
        "Cardinal presentation: diarrhoea, weight loss, bloating, steatorrhoea — think malabsorption.",
        "Test stepwise: exclude coeliac first (most common mucosal cause); then pancreatic exocrine insufficiency; then SIBO.",
        "Small bowel biopsy is gold standard for mucosal disease.",
        "Faecal elastase <100 µg/g = severe pancreatic exocrine insufficiency (PEI); <200 µg/g = moderate.",
    ],
    'sections': [
        {'heading': "Causes by Site of Defect",
         'content': [
             b("Luminal phase defects: Pancreatic exocrine insufficiency (chronic pancreatitis, pancreatic cancer, cystic fibrosis, post-pancreatectomy) — insufficient lipase/protease; Bile acid deficiency (obstructive jaundice, ileal resection >100cm, SIBO with bile acid deconjugation, cholestyramine overuse)."),
             b("Mucosal phase defects: Coeliac disease (villous atrophy, proximal small bowel); Crohn's disease (patchy, any segment); tropical sprue (diffuse, responds to antibiotics); Whipple's disease (PAS-positive macrophages); radiation enteritis; eosinophilic gastroenteritis."),
             b("Post-mucosal / lymphatic defects: Intestinal lymphangiectasia (primary or secondary to lymphoma, Whipple's, constrictive pericarditis) → fat malabsorption + protein-losing enteropathy + lymphopenia."),
         ]},
        {'heading': "Investigation Algorithm",
         'content': [
             b("Step 1 — Serology: anti-tTG IgA + total IgA (for coeliac); FBC, CRP, albumin, LFTs."),
             b("Step 2 — Stool tests: faecal elastase (PEI); Sudan stain (qualitative fat); stool culture/C&S/ova & parasites (exclude infective cause)."),
             b("Step 3 — Imaging: CT enterography (mucosal disease, Crohn's, lymph nodes, pancreas); MRI enterography (preferred in young patients — no radiation)."),
             b("Step 4 — Endoscopy + biopsy: upper GI endoscopy with duodenal biopsy (4 from D2, 2 from D1) for coeliac, Whipple's; colonoscopy + terminal ileal biopsy for Crohn's."),
             b("Step 5 — Functional tests: hydrogen breath test for SIBO (glucose or lactulose); SeHCAT scan for bile acid malabsorption; D-xylose absorption test."),
         ]},
        {'heading': "Nutritional Consequences and Replacement",
         'content': [
             b("Fat-soluble vitamins (ADEK): supplementation essential — water-miscible forms preferred in severe fat malabsorption."),
             b("Vitamin B12: IM hydroxocobalamin if ileal disease/resection — oral absorption unreliable without terminal ileum."),
             b("Iron: oral ferrous sulphate first-line; IV iron (ferric carboxymaltose) if oral not tolerated or malabsorption persistent."),
             b("Calcium + vitamin D: essential in any prolonged malabsorption syndrome; dual energy X-ray absorptiometry (DEXA) for baseline bone density."),
             b("Zinc and magnesium: supplement in IBD, SBS, chronic diarrhoea."),
             b("Pancreatic enzyme replacement therapy (PERT): Creon (pancreatin) with all fatty meals — start at 25,000–50,000 IU lipase/meal; titrate up to 75,000–100,000 IU/meal if inadequate response."),
         ]},
    ],
})

# ── Ch107 ──
CHAPTERS.append({
    'filename': 'Ch107_SIBO.docx',
    'title': 'Ch107 — Small Intestinal Bacterial Overgrowth (SIBO)',
    'pearls': [
        "SIBO = >10⁵ CFU/mL of bacteria in proximal small bowel aspirate.",
        "Classic predisposing factors: prior GI surgery (Billroth II, blind loops), motility disorders (DM autonomic neuropathy, scleroderma), anatomic abnormalities (strictures, diverticula), achlorhydria/PPI use, advanced age.",
        "Glucose hydrogen breath test: early H₂ rise (<90 min) = SIBO; sensitivity 62%, specificity 83%.",
        "Gold standard: jejunal aspirate culture >10⁵ CFU/mL — cumbersome but definitive.",
        "Treatment of choice: rifaximin 550mg TDS × 14 days (minimal systemic absorption, no significant dysbiosis).",
    ],
    'sections': [
        {'heading': "Clinical Features",
         'content': [
             b("Bloating (especially postprandial), flatulence, abdominal cramps, diarrhoea — often watery or steatorrhoeic."),
             b("Vitamin B12 deficiency: bacteria in small bowel consume cobalamin before it can be absorbed; paradoxically folate levels may be normal or elevated (bacteria synthesise folate)."),
             b("Fat-soluble vitamin deficiency (A, D, E, K) from bile acid deconjugation by bacteria → impaired micelle formation → steatorrhoea."),
             b("Weight loss, protein malnutrition in severe or prolonged SIBO."),
         ]},
        {'heading': "Diagnosis",
         'content': [
             b("Glucose H₂ breath test: 75g glucose oral → measure breath H₂ at baseline and every 15–20 min for 2–3h; early rise ≥20 ppm above baseline <90 min = SIBO."),
             b("Lactulose breath test: less specific (colonic fermentation of lactulose can cause false positives); not preferred."),
             b("Jejunal aspirate culture: aspirate during upper GI endoscopy; quantitative culture >10⁵ CFU/mL confirms SIBO; gold standard but invasive."),
             b("Empirical treatment trial is acceptable in patients with typical clinical features and strong predisposing factors."),
         ]},
        {'heading': "Treatment",
         'content': [
             b("First-line: rifaximin 550mg TDS × 14 days — non-absorbable antibiotic targeting gut flora; minimal systemic side-effects; low resistance rates."),
             b("Alternatives: co-amoxiclav 625mg TDS × 7–10d; metronidazole 400mg TDS + ciprofloxacin 500mg BD × 7d; doxycycline 100mg BD × 7d."),
             b("Treat the underlying cause: correct anatomical problem if surgically feasible; optimise DM/scleroderma management; review PPI necessity."),
             b("Recurrent SIBO: rotating antibiotic courses (1–2 weeks every 4–8 weeks); prokinetics as adjunct (erythromycin 250mg TDS AC, metoclopramide 10mg TDS) to restore MMC."),
             b("B12 replacement: IM hydroxocobalamin; fat-soluble vitamins as needed."),
         ]},
        {'heading': "Special Considerations",
         'content': [
             b("Post-Billroth II/Roux-en-Y: afferent limb stasis → high SIBO risk; high recurrence after treatment."),
             b("Diabetic gastroparesis: prokinetics reduce SIBO recurrence by restoring MMC."),
             b("Scleroderma: systemic disease impairs motility at multiple levels — difficult to eradicate SIBO; rotating antibiotics standard of care."),
             b("PPI-associated SIBO: evidence for causality is moderate; weigh risks vs benefits of PPI continuation."),
         ]},
    ],
})

# ── Ch108 ──
CHAPTERS.append({
    'filename': 'Ch108_ShortBowelSyndrome.docx',
    'title': 'Ch108 — Short Bowel Syndrome (SBS)',
    'pearls': [
        "SBS = <200 cm functional small bowel remaining after resection (some use <150 cm).",
        "Ileocaecal valve (ICV) preservation critically reduces diarrhoea and SIBO risk — document ICV status in every SBS patient.",
        "End-jejunostomy with <100 cm remaining almost always requires long-term parenteral nutrition (PN).",
        "Teduglutide (GLP-2 analogue, Revestive) is approved for PN-dependent SBS — promotes intestinal adaptation and reduces PN requirements.",
        "Ileum adapts better than jejunum — ileal resection causes more nutritional sequelae.",
    ],
    'sections': [
        {'heading': "Phases of SBS",
         'content': [
             b("Acute phase (0–3 months): massive secretory diarrhoea, high fluid/electrolyte losses (Na, K, Mg); TPN is mandatory; fluid and electrolyte management critical; gastric acid hypersecretion → add high-dose PPI."),
             b("Adaptation phase (3 months – 2 years): bowel hypertrophy, villous enlargement, crypt deepening, slowing of transit; gradual introduction of enteral nutrition stimulates adaptation; GLP-2 released by remaining ileum/colon drives adaptation."),
             b("Maintenance/stabilisation: long-term management — some patients achieve oral autonomy; others remain PN-dependent; depends on bowel remnant length and anatomy."),
         ]},
        {'heading': "Nutritional Management",
         'content': [
             b("TPN → parenteral supplementation → enteral feeds → oral diet — stepwise transition as tolerated."),
             b("End-jejunostomy: high Na and water losses; restrict hypotonic fluids (water, tea, coffee worsen output by triggering secretion); use St Mark's solution (Na 90 mmol/L, glucose 20g/L — sipped through the day); high complex carbohydrate, low fat, no lactose diet."),
             b("Colon-in-continuity: colon salvages additional calories via fermentation of unabsorbed CHO; high complex carbohydrate diet beneficial; low-oxalate diet essential to prevent calcium oxalate renal stones (fat malabsorption → increased free oxalate absorption by colon)."),
             b("Micronutrient monitoring: B12 (IM injections), fat-soluble vitamins, zinc, Mg, iron — monitor 3–6 monthly."),
         ]},
        {'heading': "Pharmacotherapy",
         'content': [
             b("Loperamide 4mg QDS (30 min before meals) + 2mg after each loose stool — reduces intestinal transit and secretion; first-line antidiarrhoeal."),
             b("Codeine phosphate 30mg QDS — stronger antidiarrhoeal; be aware of constipation risk."),
             b("Octreotide: reduces intestinal secretions but impairs intestinal adaptation — use short-term in acute phase if secretory losses uncontrolled; avoid long-term in adaptation phase."),
             b("PPIs: essential in early SBS — gastric acid hypersecretion inactivates pancreatic enzymes and causes peptic ulceration; use IV omeprazole/pantoprazole in acute phase then oral."),
             b("Cholestyramine: for bile acid diarrhoea in patients with >100cm ileum resected but colon intact — do not use in end-jejunostomy patients (worsens fat malabsorption)."),
         ]},
        {'heading': "Teduglutide (Revestive)",
         'content': [
             b("GLP-2 analogue: 0.05 mg/kg SC once daily."),
             b("Promotes intestinal mucosal growth, increases villous height, reduces gastric emptying."),
             b("Clinical trials: reduces weekly PN volume by ~4.4L/week vs placebo; 27% achieve ≥1 day/week PN reduction."),
             b("Monitor: intestinal polyp surveillance (colonoscopy before starting, then every 5 years); contraindicated in active GI malignancy or within 5 years of GI cancer."),
             b("Expensive — requires specialist prescribing; PBS-subsidised criteria vary by country."),
         ]},
    ],
})

print("Chapters 100–108 data loaded.")

# ── Ch109 ──
CHAPTERS.append({
    'filename': 'Ch109_CoeliacDisease.docx',
    'title': 'Ch109 — Coeliac Disease',
    'pearls': [
        "HLA-DQ2 (90%) and HLA-DQ8 (10%) are necessary but not sufficient for coeliac disease; negative HLA virtually excludes the diagnosis.",
        "Classic presentation (malabsorption + diarrhoea) is now a minority — most present atypically: iron deficiency anaemia, osteoporosis, elevated LFTs, neurological symptoms, infertility.",
        "Serology: anti-tTG IgA is the preferred first test; always check total IgA to exclude IgA deficiency (prevalence ~1 in 400).",
        "Do NOT start a gluten-free diet before diagnostic biopsy — histology will normalise.",
        "Complications: refractory coeliac disease (RCD), enteropathy-associated T-cell lymphoma (EATL — rare but life-threatening), small bowel adenocarcinoma.",
    ],
    'sections': [
        {'heading': "Diagnosis",
         'content': [
             b("Serology: anti-tTG IgA (sensitivity 95%, specificity 95%); anti-endomysial IgA (highly specific, operator-dependent); anti-deamidated gliadin peptide (DGP) IgG — use if IgA-deficient."),
             b("Endoscopy: duodenal biopsies essential — 4 biopsies from D2 + 2 from D1 (Marsh staging); scalloped folds, mosaic pattern, loss of Kerckring folds are endoscopic clues."),
             tbl(["Marsh Grade", "Histology", "Significance"],
                 [["Marsh 1", "Intraepithelial lymphocytes (IEL) >25/100 epithelial cells; normal villi", "Compatible, not diagnostic alone"],
                  ["Marsh 2", "IEL + crypt hyperplasia", "Compatible"],
                  ["Marsh 3a/b/c", "IEL + crypts + partial/subtotal/total villous atrophy", "Diagnostic (Marsh 3 = confirmatory)"]]),
             b("Capsule endoscopy: role in suspected RCD and evaluating small bowel extent."),
             b("HLA typing: DQ2/DQ8 negative = essentially excludes coeliac. Useful to reassure patients already on GFD before biopsy."),
         ]},
        {'heading': "Gluten-Free Diet (GFD)",
         'content': [
             b("Avoid: wheat, rye, barley. Oats are controversial (contain avenin — tolerated by most but pure oats required)."),
             b("Cross-contamination risk: shared cooking surfaces, fryers, sauces containing gluten — patient education essential."),
             b("Dietician referral mandatory: GFD adherence is complex and socially challenging; dedicated coeliac dietician improves outcomes."),
             b("Monitoring response: repeat anti-tTG at 6–12 months post-GFD; should fall significantly. Histological healing may lag serology by 1–2 years."),
             b("Hidden gluten: medications, supplements, processed foods, soy sauce, malt vinegar — check labels."),
             b("Annual DEXA scan for bone density; consider calcium + vitamin D supplementation."),
         ]},
        {'heading': "Refractory Coeliac Disease (RCD)",
         'content': [
             b("Definition: persistent/recurrent malabsorption symptoms and villous atrophy on GFD for >12 months despite strict adherence."),
             b("RCD Type I: normal IEL phenotype; better prognosis; treat with steroids (budesonide or prednisolone) ± azathioprine."),
             b("RCD Type II: aberrant clonal IEL population (CD3+, CD8−, surface CD3−) — high risk of EATL (>50% at 5 years); poor prognosis; treat with cladribine + stem cell transplant in selected cases."),
             b("Investigate RCD: CT/MRI enterography (lymph nodes, cavitating lymph node syndrome, complications); capsule endoscopy for ulcers; push enteroscopy for biopsy."),
             b("Complications to exclude before diagnosing RCD: dietary non-compliance (most common), SIBO, pancreatic exocrine insufficiency, microscopic colitis, lymphoma."),
         ]},
        {'heading': "Singapore Notes",
         'content': [
             b("Coeliac disease increasingly diagnosed in Singapore — prevalence ~0.5–1% based on serology studies; historically underdiagnosed."),
             b("GFD difficult to implement in Singapore: most local hawker food (noodles, dumplings, soy sauce) contains gluten; dedicated GF options limited but improving."),
             b("Annual bone density monitoring recommended — particularly relevant as baseline vitamin D deficiency is common."),
             b("Consider coeliac screening in first-degree relatives (10-fold increased risk), Type 1 DM, Down syndrome, Turner syndrome, autoimmune thyroid disease."),
         ]},
    ],
})

# ── Ch110 ──
CHAPTERS.append({
    'filename': 'Ch110_TropicalDiarrhoeaMalabsorption.docx',
    'title': 'Ch110 — Tropical Diarrhoea and Malabsorption',
    'pearls': [
        "Tropical sprue = chronic malabsorption in tropical residents following an infective trigger — affects the entire small bowel (unlike coeliac which is proximal-predominant).",
        "Tropical sprue responds to tetracycline + folic acid — distinguishing feature from coeliac (responds to GFD).",
        "Environmental enteropathy = subclinical mucosal damage in children in tropical low-resource settings — contributes to stunting and impaired vaccine responses.",
        "In Singapore/SE Asia returnees with malabsorption: always exclude Giardia, SIBO, coeliac, strongyloides, and tropical sprue.",
    ],
    'sections': [
        {'heading': "Tropical Sprue",
         'content': [
             b("Epidemiology: residents or long-term visitors to tropics (SE Asia, India, Caribbean); after infective trigger (often enterotoxigenic E. coli or other organisms)."),
             b("Clinical: chronic diarrhoea, steatorrhoea, weight loss, folate deficiency (early — proximal bowel involved), B12 deficiency (late — ileal involvement), megaloblastic anaemia."),
             b("Biopsy: partial villous atrophy (less severe than coeliac Marsh 3); affects entire small bowel; IEL present."),
             b("Treatment: tetracycline 250mg QDS × 3–6 months + folic acid 5mg/day; response typically within 1–2 weeks; treat until remission (usually 3–6 months)."),
             b("Alternative to tetracycline: co-trimoxazole 960mg BD × 3–6 months."),
         ]},
        {'heading': "Brainerd Diarrhoea",
         'content': [
             b("Epidemic secretory diarrhoea — waterborne or foodborne; outbreak-associated."),
             b("Large-volume watery diarrhoea, persists for months; no antibiotic response."),
             b("Self-limiting over 12–24 months; aetiology unidentified."),
         ]},
        {'heading': "Environmental Enteropathy",
         'content': [
             b("Subclinical villous blunting + increased intestinal permeability in children in developing-country settings."),
             b("Impairs oral vaccine responses (polio, rotavirus); contributes to growth stunting."),
             b("Nutrition and WASH (water, sanitation, hygiene) interventions are key."),
         ]},
        {'heading': "Differential Diagnosis: Malabsorption in Singapore/Tropics",
         'content': [
             tbl(["Diagnosis", "Key Features", "Treatment"],
                 [["Giardia lamblia", "Bloating, intermittent diarrhoea, stool antigen positive", "Metronidazole 400mg TDS × 5–7d OR tinidazole 2g stat"],
                  ["SIBO", "Prior surgery/motility disorder, bloating, breath test positive", "Rifaximin 550mg TDS × 14d"],
                  ["Coeliac disease", "anti-tTG IgA positive, Marsh 3 biopsy", "Strict GFD"],
                  ["Tropical sprue", "Tropical resident, responds to tetracycline + folate", "Tetracycline 250mg QDS × 3–6m + folic acid"],
                  ["Strongyloides", "Eosinophilia, larva currens, from tropics", "Ivermectin 200 µg/kg OD × 2d"],
                  ["Pancreatic exocrine insufficiency", "Faecal elastase <100, steatorrhoea, chronic pancreatitis", "Pancreatin (Creon) with meals"]]),
         ]},
    ],
})

# ── Ch111 ──
CHAPTERS.append({
    'filename': 'Ch111_WhippleDisease.docx',
    'title': "Ch111 — Whipple's Disease",
    'pearls': [
        "Caused by Tropheryma whipplei — a gram-positive actinomycete; rare but a classic board question.",
        "Classic triad: diarrhoea + weight loss + migratory polyarthralgia (arthralgia often precedes GI symptoms by years).",
        "CNS involvement (dementia, ophthalmoplegia, myoclonus) = poor prognosis — always ask about neurological symptoms.",
        "Diagnosis: PAS-positive macrophages on small bowel biopsy + confirmatory PCR.",
        "Treatment: IV ceftriaxone 2g OD × 2 weeks → oral co-trimoxazole (TMP-SMX) 960mg BD × 1–2 years.",
    ],
    'sections': [
        {'heading': "Clinical Features",
         'content': [
             b("GI: steatorrhoea, watery diarrhoea, malabsorption, abdominal pain, hepatosplenomegaly."),
             b("Joint: migratory, non-destructive polyarthralgia or polyarthritis (seronegative); often precedes GI diagnosis by 5–10 years; may be misdiagnosed as seronegative RA."),
             b("CNS: cognitive decline/dementia, supranuclear gaze palsy (inability to look downward), oculomasticatory myorhythmia (pathognomonic — rhythmic eye movements + jaw contractions), cerebellar ataxia; very poor prognosis."),
             b("Cardiac: culture-negative endocarditis; pericarditis."),
             b("Skin: hyperpigmentation (sun-exposed areas)."),
             b("Systemic: lymphadenopathy (mesenteric LAD on CT), fever, anaemia."),
         ]},
        {'heading': "Diagnosis",
         'content': [
             b("Upper GI endoscopy: shaggy pale yellow/cream-coloured mucosa with eroded red patches in duodenum."),
             b("Biopsy (D2/D3): lamina propria packed with foamy macrophages positive on PAS stain (magenta granules); distorted villi; dilated lacteals."),
             b("PCR for T. whipplei: confirm on biopsy tissue; also perform on CSF if CNS involvement suspected."),
             b("CT abdomen: low-attenuation mesenteric/retroperitoneal lymph nodes (due to lipid deposition) — characteristic finding."),
             b("Culture: possible but extremely slow; not routine."),
         ]},
        {'heading': "Treatment",
         'content': [
             b("Induction: IV ceftriaxone 2g OD × 14 days — ensures CNS penetration (TMP-SMX has poor CNS penetration for induction)."),
             b("Maintenance: oral TMP-SMX 960mg BD × minimum 1 year (many recommend 2 years to prevent CNS relapse)."),
             b("CNS Whipple's: treat for minimum 2 years; monitor with repeat CSF PCR; longer treatment may be needed."),
             b("Sulfa allergy alternative: doxycycline 100mg BD + hydroxychloroquine 200mg TDS (intracellular alkalinisation enhances doxycycline efficacy)."),
             b("Relapse: re-treat with IV ceftriaxone; monitor with repeat biopsy and PCR at 6–12 months."),
             b("Immune reconstitution inflammatory syndrome (IRIS): can occur after treatment initiation — masquerades as worsening."),
         ]},
    ],
})

# ── Ch112 ──
CHAPTERS.append({
    'filename': 'Ch112_InfectiousEnteritisProctocolitis.docx',
    'title': 'Ch112 — Infectious Enteritis and Proctocolitis',
    'pearls': [
        "Most acute diarrhoea is self-limiting — antibiotics are NOT routinely indicated.",
        "Bloody diarrhoea + fever = inflammatory diarrhoea → stool culture + consider antibiotics.",
        "EHEC O157:H7: avoid antibiotics (risk of HUS from toxin release) and antiperistaltics.",
        "Travellers' diarrhoea empiric Rx: azithromycin 1g stat or 500mg OD × 3 days — preferred over fluoroquinolones in SE Asia due to high quinolone resistance.",
        "Always screen STI-related proctocolitis broadly in MSM patients.",
    ],
    'sections': [
        {'heading': "Approach to Acute Diarrhoea",
         'content': [
             b("Watery, non-bloody, afebrile, <3 days: likely viral (norovirus, rotavirus) or ETEC → supportive (ORS, diet). No antibiotics."),
             b("Bloody diarrhoea or fever or systemic toxicity: inflammatory → stool culture + C. diff PCR + consider antibiotics after specimen."),
             b("Immunocompromised (HIV, transplant, IBD on biologics): lower threshold for investigation and treatment; consider CMV, Cryptosporidium, Microsporidium."),
             b("Admission criteria: severe dehydration, bloody diarrhoea + systemic toxicity, suspected HUS, elderly/frail, unable to maintain oral intake."),
         ]},
        {'heading': "Common Pathogens — Singapore Context",
         'content': [
             tbl(["Pathogen", "Key Features", "Treatment"],
                 [["Campylobacter", "Most common bacterial diarrhoea SG; bloody diarrhoea; reactive arthritis, Guillain-Barré post-infection", "Azithromycin 500mg OD × 3d (quinolone resistance ~50% in SG)"],
                  ["Salmonella non-typhoidal", "Poultry/eggs; usually self-limiting; bacteraemia risk in elderly/immunocompromised", "Self-limiting; ciprofloxacin or azithromycin if severe/bacteraemic"],
                  ["Shigella", "Dysentery; low inoculum dose; outbreaks in institutions", "Azithromycin 500mg OD × 3d (avoid quinolones — resistance)"],
                  ["EHEC O157:H7", "Undercooked beef; bloody diarrhoea; NO antibiotics (↑HUS risk)", "Supportive only; monitor renal function for HUS"],
                  ["Vibrio cholerae", "Watery diarrhoea → rapid dehydration; rice-water stools", "ORS (mainstay) ± doxycycline 300mg stat (reduces duration)"],
                  ["Yersinia enterocolitica", "RIF pain mimics appendicitis; reactive arthritis; cold enrichment required for culture", "Usually self-limiting; ciprofloxacin if severe"]]),
         ]},
        {'heading': "Proctocolitis (STI-Related)",
         'content': [
             b("Relevant in MSM patients — always screen broadly for multiple STIs simultaneously."),
             b("Neisseria gonorrhoeae: purulent rectal discharge, tenesmus; ceftriaxone 500mg IM stat (1g IM if weight >150kg)."),
             b("Chlamydia trachomatis LGV (serovars L1–L3): severe bloody proctitis, perirectal fistulae, inguinal lymphadenopathy; treat with doxycycline 100mg BD × 21 days."),
             b("Chlamydia non-LGV: milder proctitis; doxycycline 100mg BD × 7 days."),
             b("HSV proctitis: vesicles, severe anorectal pain; aciclovir 400mg TDS × 10d or valaciclovir 1g BD × 10d."),
             b("Syphilis (proctitis/anal chancre): benzathine penicillin 2.4MU IM stat (primary/secondary); refer sexual health."),
             b("Screen: HIV, hepatitis B/C, gonorrhoea, chlamydia (including LGV PCR), syphilis serology — all simultaneously in symptomatic MSM."),
         ]},
    ],
})

# ── Ch113 ──
CHAPTERS.append({
    'filename': 'Ch113_FoodPoisoning.docx',
    'title': 'Ch113 — Food Poisoning',
    'pearls': [
        "Preformed toxin = rapid onset <6h: Staph aureus, Bacillus cereus (emetic form), Clostridium perfringens (reheated meat, 8–16h).",
        "Botulism: descending flaccid paralysis, diplopia, dysphagia — canned foods, honey (infants); treat with trivalent antitoxin + ICU.",
        "Scombroid: histamine release from spoiled fish (tuna, mackerel) — flushing, urticaria, headache within minutes; antihistamines.",
        "Ciguatera: reef fish toxin (barracuda, red snapper) — paresthesias, temperature reversal; no specific treatment; may last months.",
        "Nearly all food poisoning is managed supportively — diagnose the source, notify public health if outbreak.",
    ],
    'sections': [
        {'heading': "Rapid-Onset Toxin-Mediated (<6 hours)",
         'content': [
             b("Staphylococcus aureus: preformed heat-stable enterotoxin; high-risk foods: custards, cream pastries, rice salads, processed meats; onset 1–6h; severe vomiting ± diarrhoea; self-limiting 24h; no antibiotics."),
             b("Bacillus cereus (emetic): heat-stable cereulide toxin; fried rice/noodles left at room temperature; onset 1–5h; predominantly vomiting; self-limiting."),
             b("Clostridium perfringens: heat-labile spore-forming toxin; reheated meat and poultry; onset 8–16h; watery diarrhoea, cramps; self-limiting 24–48h."),
             b("Bacillus cereus (diarrhoeal): heat-labile enterotoxin; meats, vegetables, sauces; onset 8–16h; diarrhoea predominant; self-limiting."),
         ]},
        {'heading': "Botulism",
         'content': [
             b("Caused by Clostridium botulinum neurotoxin (types A, B, E commonest)."),
             b("Sources: home-canned foods, preserved meats, smoked fish; infant botulism from honey."),
             b("Clinical: descending flaccid paralysis — diplopia → dysarthria → dysphagia → limb weakness → respiratory failure; autonomic features (dry mouth, constipation, urinary retention, dilated pupils); NO fever."),
             b("Diagnosis: mouse bioassay of stool/food; toxin ELISA; electromyography (incremental response to repetitive stimulation)."),
             b("Treatment: trivalent (ABE) antitoxin ASAP (reduces duration, does not reverse established paralysis); ICU + mechanical ventilation if respiratory failure; notifiable disease."),
         ]},
        {'heading': "Marine Toxins",
         'content': [
             b("Scombroid poisoning: tuna, mackerel, mahi-mahi inadequately refrigerated → bacterial histidine decarboxylase converts histidine → histamine; onset within minutes of eating; flushing, urticaria, headache, palpitations, diarrhoea; treatment: antihistamines (chlorphenamine 4mg or cetirizine 10mg); usually self-limiting."),
             b("Ciguatera fish poisoning: ciguatoxin from dinoflagellates (Gambierdiscus) bioaccumulates in reef fish (barracuda, red snapper, grouper, amberjack); onset 3–5h; GI then neurological; temperature reversal (cold feels hot) is pathognomonic; arthralgia, bradycardia; may last months–years; no specific antidote; avoid alcohol, nuts, fish (may worsen symptoms); symptomatic treatment."),
             b("Paralytic shellfish poisoning: saxitoxin; mussels, clams, scallops during algal bloom; ascending paralysis; respiratory failure possible; supportive ± mechanical ventilation."),
             b("Amnesic shellfish poisoning: domoic acid; mussels; short-term memory loss, confusion, seizures; supportive."),
         ]},
    ],
})

# ── Ch114 ──
CHAPTERS.append({
    'filename': 'Ch114_CDiff_AntibioticDiarrhoea.docx',
    'title': 'Ch114 — Clostridioides difficile Infection and Antibiotic-Associated Diarrhoea',
    'pearls': [
        "CDI = most common healthcare-associated infection; always consider in patients with diarrhoea on or within 8 weeks of antibiotics.",
        "Highest-risk antibiotics: fluoroquinolones, clindamycin, cephalosporins, amoxicillin-clavulanate.",
        "Test only symptomatic patients (≥3 unformed stools/24h) — do NOT test formed stool or asymptomatic carriers.",
        "Fidaxomicin preferred over vancomycin for first recurrence — lower recurrence rate.",
        "FMT (faecal microbiota transplantation) is >90% effective for multiply recurrent CDI.",
    ],
    'sections': [
        {'heading': "Diagnosis",
         'content': [
             b("Two-step algorithm: GDH antigen EIA (screening — high sensitivity) + toxin A/B EIA (confirmation — high specificity)."),
             b("NAAT/PCR: highly sensitive; detects toxin genes, not active toxin — may detect asymptomatic colonisers; use only in symptomatic patients."),
             b("Sigmoidoscopy/colonoscopy: pseudomembranes (yellow-white plaques) in fulminant CDI; not routine for diagnosis."),
             b("Do NOT perform test-of-cure after treatment — patients may remain PCR-positive after clinical resolution."),
         ]},
        {'heading': "Treatment by Severity",
         'content': [
             tbl(["Severity", "Criteria", "Treatment"],
                 [["Non-severe", "WBC ≤15×10⁹/L AND Cr <1.5×ULN", "Fidaxomicin 200mg BD × 10d (preferred) OR vancomycin 125mg QDS PO × 10d"],
                  ["Severe", "WBC >15×10⁹/L OR Cr ≥1.5×ULN", "Fidaxomicin 200mg BD × 10d OR vancomycin 125mg QDS × 10d"],
                  ["Fulminant", "Hypotension, ileus, toxic megacolon, peritonitis", "Vancomycin 500mg QDS PO/NG + IV metronidazole 500mg TDS; rectal vancomycin if ileus; urgent surgical consult"],
                  ["1st recurrence", "CDI recurrence within 8 weeks", "Fidaxomicin 200mg BD × 10d (preferred) OR vancomycin pulsed/tapered"],
                  ["2nd+ recurrence", "Multiple recurrences", "FMT (capsule, colonoscopic, or enema — >90% efficacy); bezlotoxumab 10mg/kg IV to prevent future recurrence"]]),
         ]},
        {'heading': "Prevention",
         'content': [
             b("Hand hygiene: soap and water essential — alcohol gel DOES NOT kill C. diff spores."),
             b("Contact precautions: gloves + gown until diarrhoea resolves; single-side room preferred."),
             b("Environmental decontamination: sodium hypochlorite (bleach) 1:10 solution — effective sporicidal agent."),
             b("Antimicrobial stewardship: reduce unnecessary antibiotic use, especially high-risk antibiotics."),
             b("Probiotics: Saccharomyces boulardii or Lactobacillus rhamnosus — modest reduction in AAD and CDI when co-prescribed with high-risk antibiotics; evidence of quality moderate."),
         ]},
    ],
})

# ── Ch115 ──
CHAPTERS.append({
    'filename': 'Ch115_IntestinalProtozoa.docx',
    'title': 'Ch115 — Intestinal Protozoa',
    'pearls': [
        "Giardia lamblia is the most common intestinal parasite worldwide — stool antigen or PCR for diagnosis.",
        "Entamoeba histolytica (pathogenic) vs E. dispar (non-pathogenic) are morphologically identical — distinguish by antigen test or PCR.",
        "Cryptosporidium causes self-limiting diarrhoea in immunocompetent hosts but can be life-threatening in HIV/immunocompromised — ART is key.",
        "Amoebic liver abscess: single, right lobe, fever + right upper quadrant pain in a traveller; serology positive; treat with metronidazole ± aspiration.",
    ],
    'sections': [
        {'heading': "Giardia lamblia",
         'content': [
             b("Transmission: contaminated water (hikers, travellers, day-care centres), person-to-person faeco-oral."),
             b("Clinical: watery diarrhoea, bloating, flatulence, malabsorption, weight loss; may become chronic; asymptomatic carriage possible."),
             b("Diagnosis: stool antigen ELISA or PCR (preferred over microscopy for trophozoites/cysts); duodenal aspirate or string test if stool tests negative."),
             b("Treatment: metronidazole 400mg TDS × 5–7 days OR tinidazole 2g single dose (better compliance, similar efficacy); alternative: albendazole 400mg OD × 5 days."),
             b("Treat asymptomatic carriers in low-endemicity settings (Singapore); in endemic settings, may observe."),
             b("Refractory giardiasis: combination metronidazole + albendazole or quinacrine; check for immunodeficiency (IgA, IgG deficiency)."),
         ]},
        {'heading': "Entamoeba histolytica",
         'content': [
             b("Transmission: faeco-oral, contaminated water/food; global distribution; higher risk in travellers to tropics."),
             b("Intestinal amoebiasis: amoebic colitis — bloody diarrhoea, cramps, tenesmus; colonoscopy shows flask-shaped ulcers (undermined edges), haemorrhagic mucosa; amoeboma (inflammatory mass)."),
             b("Amoebic liver abscess (ALA): most common extraintestinal form; single large abscess, right lobe; fever, RUQ pain, tender hepatomegaly; serology (IHA, ELISA) positive in >90%; stool examination often negative; US/CT confirms."),
             b("Diagnosis: stool antigen/PCR for intestinal amoebiasis; serology for ALA; biopsy rarely needed."),
             b("Treatment: metronidazole 800mg TDS × 10 days (or tinidazole 2g OD × 5 days) FOLLOWED by luminal agent — diloxanide furoate 500mg TDS × 10 days OR paromomycin 500mg TDS × 7 days — to eradicate cysts in colon."),
             b("ALA drainage: indicated if large (>5cm), left lobe (risk of pericardial rupture), not responding to metronidazole after 3–5 days, or diagnostic uncertainty. Most respond to metronidazole alone."),
         ]},
        {'heading': "Cryptosporidium",
         'content': [
             b("Oocysts resist standard chlorination; waterborne outbreaks in municipal water supplies."),
             b("Immunocompetent: self-limiting profuse watery diarrhoea 1–2 weeks; supportive management."),
             b("HIV/immunocompromised: chronic severe diarrhoea, malnutrition, cholangiopathy; life-threatening."),
             b("Diagnosis: modified Ziehl-Neelsen stool stain; stool antigen (DFA or ELISA); PCR most sensitive."),
             b("Treatment: immunocompetent — supportive; HIV — nitazoxanide 500mg BD × 3 days + ART (immune reconstitution is key); no highly effective drug in severely immunocompromised without immune recovery."),
         ]},
        {'heading': "Other Intestinal Protozoa",
         'content': [
             tbl(["Organism", "Key Feature", "Treatment"],
                 [["Cyclospora cayetanensis", "Imported fruit/herbs; prolonged diarrhoea; acid-fast oocysts", "TMP-SMX 960mg BD × 7–10 days"],
                  ["Cystoisospora belli", "HIV/immunocompromised; acid-fast oocysts larger than Cryptosporidium", "TMP-SMX 960mg QDS × 10d then BD × 3 weeks; secondary prophylaxis in HIV"],
                  ["Microsporidium", "HIV; chronic diarrhoea; modified trichrome stain; small spores", "Albendazole 400mg BD (Encephalitozoon); fumagillin for Enterocytozoon bieneusi (not widely available)"],
                  ["Blastocystis hominis", "Controversial pathogenicity; common in Singapore; most asymptomatic", "Treat if symptomatic after exclusion of other causes: metronidazole or co-trimoxazole"]]),
         ]},
    ],
})

# ── Ch116 ──
CHAPTERS.append({
    'filename': 'Ch116_IntestinalWorms.docx',
    'title': 'Ch116 — Intestinal Worms (Helminths)',
    'pearls': [
        "Strongyloides stercoralis is the most dangerous intestinal helminth — hyperinfection syndrome in immunocompromised patients (steroids, organ transplant, HTLV-1) can be fatal.",
        "Screen ALL patients from endemic areas (SE Asia including Singapore) for strongyloides before starting immunosuppression — serology (ELISA) is the most practical screen.",
        "Eosinophilia + travel history to tropics = helminthic infection until proven otherwise.",
        "Liver flukes (Clonorchis/Opisthorchis) are risk factors for cholangiocarcinoma — common in SE Asian populations who eat raw/undercooked freshwater fish.",
    ],
    'sections': [
        {'heading': "Intestinal Nematodes (Roundworms)",
         'content': [
             b("Ascaris lumbricoides: largest intestinal roundworm (15–35cm); massive infection → intestinal obstruction; biliary ascariasis (worm migrates into bile duct) → biliary colic, cholangitis, pancreatitis; Löffler syndrome (pulmonary eosinophilia) during larval migration; treatment: albendazole 400mg stat OR mebendazole 500mg stat."),
             b("Hookworm (Ancylostoma duodenale, Necator americanus): skin penetration (ground itch); iron deficiency anaemia from intestinal blood loss; protein-losing enteropathy; treatment: albendazole 400mg OD × 3 days."),
             b("Strongyloides stercoralis: unique autoinfection cycle — larvae can complete life cycle within host; low-grade chronic infection in immunocompetent (perianal pruritus, larva currens, mild diarrhoea, eosinophilia); hyperinfection in immunocompromised: massive larval dissemination, carry gut bacteria → gram-negative septicaemia, meningitis, multi-organ failure; diagnosis: stool O&P (low sensitivity — single stool exam <30%); serology ELISA (sensitivity ~85%); duodenal aspirate/string test for confirmation; treatment: ivermectin 200 µg/kg OD × 2 days (preferred — superior efficacy); albendazole 400mg BD × 7 days (less effective); hyperinfection: ivermectin daily until parasite-free (minimum 2 weeks)."),
             b("Enterobius vermicularis (Pinworm): most common helminth in children; perianal itch worse at night (female worm deposits eggs); diagnosis: scotch tape test applied perianally first thing in morning; treatment: mebendazole 100mg stat + repeat in 2 weeks OR albendazole 400mg stat + repeat 2 weeks; treat all household members."),
             b("Trichuris trichiura (Whipworm): right colon; heavy infection → Trichuris dysentery syndrome (bloody diarrhoea, rectal prolapse, anaemia in children); treatment: mebendazole 500mg stat OR albendazole 400mg OD × 3 days."),
         ]},
        {'heading': "Cestodes (Tapeworms)",
         'content': [
             b("Taenia saginata (beef tapeworm): motile proglottids passed in stool; diagnosis: proglottid identification; treatment: praziquantel 10mg/kg stat."),
             b("Taenia solium (pork tapeworm): similar GI presentation; DANGER — if T. solium eggs are ingested → cysticercosis (larval cysts in brain → neurocysticercosis, epilepsy, hydrocephalus); treatment of intestinal T. solium: praziquantel 10mg/kg stat; neurocysticercosis: albendazole 7.5mg/kg BD + steroids + anti-epileptics ± neurosurgery."),
             b("Hymenolepis nana (dwarf tapeworm): most common tapeworm in Singapore; self-infection possible; praziquantel 25mg/kg stat."),
             b("Diphyllobothrium latum (fish tapeworm): raw freshwater fish; B12 deficiency (tapeworm preferentially absorbs B12); praziquantel 10mg/kg stat."),
         ]},
        {'heading': "Trematodes (Flukes)",
         'content': [
             tbl(["Fluke", "Key Feature", "Treatment"],
                 [["Clonorchis sinensis / Opisthorchis viverrini", "Raw freshwater fish; SE Asia; risk factor for cholangiocarcinoma; biliary obstruction", "Praziquantel 25mg/kg TDS × 2 days"],
                  ["Fasciola hepatica", "Watercress; tender hepatomegaly, eosinophilia; serpiginous liver tracks on CT", "Triclabendazole 10mg/kg single dose"],
                  ["Schistosoma mansoni/japonicum", "Fresh water contact; portal hypertension, hepatosplenomegaly; haematuria (S. haematobium)", "Praziquantel 40mg/kg in 2 divided doses same day"]]),
         ]},
    ],
})

# ── Ch117 ──
CHAPTERS.append({
    'filename': 'Ch117_IBD_Epidemiology_Pathogenesis_Diagnosis.docx',
    'title': 'Ch117 — Inflammatory Bowel Disease: Epidemiology, Pathogenesis, and Diagnosis',
    'pearls': [
        "IBD incidence is rising rapidly in Singapore and Asia — particularly among young urbanised populations.",
        "Crohn's disease: transmural, skip lesions, any segment mouth-to-anus; UC: mucosal, continuous from rectum proximally.",
        "Key discriminating features for CD: perianal disease, granulomas, cobblestoning, rectal sparing.",
        "Faecal calprotectin >250 µg/g indicates active mucosal inflammation — useful for monitoring remission and predicting relapse.",
        "Always exclude infection (stool culture, C. diff, CMV) before starting immunosuppression.",
    ],
    'sections': [
        {'heading': "Montreal Classification",
         'content': [
             tbl(["CD — Location (L)", "CD — Behaviour (B)", "CD — Age (A)"],
                 [["L1: Ileal", "B1: Non-stricturing, non-penetrating", "A1: <17 years"],
                  ["L2: Colonic", "B2: Stricturing", "A2: 17–40 years"],
                  ["L3: Ileocolonic", "B3: Penetrating (fistulising)", "A3: >40 years"],
                  ["L4: Upper GI", "+p: Perianal disease modifier", ""]]),
             tbl(["UC — Extent (E)", "UC — Severity (S)", ""],
                 [["E1: Proctitis (distal to rectosigmoid)", "S0: Clinical remission", ""],
                  ["E2: Left-sided (to splenic flexure)", "S1: Mild", ""],
                  ["E3: Pancolitis (proximal to splenic flexure)", "S2: Moderate", ""],
                  ["", "S3: Severe", ""]]),
         ]},
        {'heading': "Diagnosis of Crohn's Disease",
         'content': [
             b("Ileocolonoscopy + biopsy: essential first investigation; skip lesions, aphthous ulcers, cobblestoning, linear/serpiginous ulcers, rectal sparing."),
             b("Histology: patchy/focal chronic active inflammation; non-caseating epithelioid granulomas (present in 30–40% of biopsies — not required for diagnosis)."),
             b("CT/MRI enterography: small bowel disease extent, mural thickening, enhancement, mesenteric fat stranding, lymph nodes, fistulae, abscesses."),
             b("MRI pelvis: mandatory in perianal CD for fistula/abscess mapping (Parks classification)."),
             b("Capsule endoscopy: small bowel mucosal disease not seen on cross-sectional imaging; contraindicated in stricturing CD until patency capsule confirms no stricture."),
         ]},
        {'heading': "Diagnosis of Ulcerative Colitis",
         'content': [
             b("Colonoscopy + biopsy: continuous inflammation from rectum; diffuse mucosal erythema, granularity, friability, loss of vascular pattern, spontaneous bleeding; pseudopolyps in chronic disease."),
             b("Rectal sparing: suggests Crohn's, drug-induced colitis, or backwash ileitis (long-standing extensive UC)."),
             b("Histology: diffuse plasmacytosis, cryptitis, crypt abscesses, crypt architectural distortion; absence of granulomas."),
             b("Endoscopic severity: Mayo score or UCEIS (UC Endoscopic Index of Severity); MES 0–3."),
         ]},
        {'heading': "Differential Diagnosis",
         'content': [
             b("Infective colitis: always exclude before starting immunosuppression — stool culture (Campylobacter, Salmonella, Shigella, EHEC), C. diff PCR, CMV (in severely ill UC patients), amoeba."),
             b("Ischaemic colitis: watershed areas, elderly, precipitant (hypotension, aortic surgery)."),
             b("Microscopic colitis: normal endoscopy — diagnosis by biopsy."),
             b("Behçet's disease: deep ileal/caecal ulcers + aphthous oral ulcers + genital ulcers + uveitis."),
             b("NSAID enteropathy: history of NSAID use; characteristic diaphragm strictures."),
             b("Radiation colitis: history of pelvic radiotherapy."),
         ]},
        {'heading': "Biomarkers",
         'content': [
             b("CRP: correlates with CD disease activity (less reliable in UC mucosal inflammation without systemic response); normal CRP does not exclude active CD in ~30%."),
             b("Faecal calprotectin (FCP): sensitive marker of mucosal inflammation; >250 µg/g = active inflammation; <50 = remission; useful for monitoring response to treatment and predicting relapse; not specific (elevated in infection, NSAID enteropathy, bowel cancer)."),
             b("Faecal lactoferrin: similar utility to FCP."),
             b("Serology: p-ANCA (UC associated), ASCA (CD associated) — limited diagnostic utility alone; not recommended as primary diagnostic test."),
         ]},
    ],
})

# ── Ch118 ──
CHAPTERS.append({
    'filename': 'Ch118_IBD_Management.docx',
    'title': 'Ch118 — Inflammatory Bowel Disease: Management',
    'pearls': [
        "Treat-to-target (T2T) strategy: target mucosal healing (not just symptomatic remission) — associated with improved long-term outcomes.",
        "Steroids for induction only — NEVER for maintenance of IBD remission.",
        "Before starting thiopurines (azathioprine/6-MP): check TPMT and NUDT15 genotype — NUDT15 variant is more prevalent in Asians and causes severe leucopenia.",
        "Screen for latent TB (Quantiferon Gold), hepatitis B (HBsAg, anti-HBc, anti-HBs), and varicella immunity before starting biologics.",
        "IBD and pregnancy: mesalazine is safe throughout; anti-TNF safe in 1st/2nd trimester; methotrexate absolutely contraindicated.",
    ],
    'sections': [
        {'heading': "UC Management",
         'content': [
             b("Mild-moderate proctitis (E1): mesalazine suppository 1g OD (first-line — topical superior to oral for proctitis alone); add oral mesalazine 2–4.8g/day if not adequate."),
             b("Mild-moderate left-sided/extensive (E2/E3): oral mesalazine 2–4.8g/day + topical mesalazine enema (combined oral + topical is superior to either alone); add oral prednisolone 40mg/day tapering over 8 weeks if no response within 2 weeks."),
             b("Severe UC (inpatient admission): IV hydrocortisone 400mg/day (100mg QDS); thromboprophylaxis with LMWH; nil by mouth/light diet; daily monitoring of stool chart, CRP, FBC, electrolytes; flexible sigmoidoscopy within 24h to exclude CMV and confirm diagnosis."),
             b("Severe UC — Oxford criteria for rescue therapy decision at Day 3: stool frequency >8/day OR CRP >45 mg/L + stool frequency 3–8/day → predict 85% colectomy risk → consider rescue therapy."),
             b("Rescue therapy: infliximab 5mg/kg IV OR ciclosporin 2mg/kg/day IV (bridge to thiopurine); discuss colectomy simultaneously with colorectal surgery; if rescue fails → colectomy."),
             b("Maintenance of UC remission: mesalazine for E2/E3; azathioprine 2–2.5mg/kg/day (steroid-dependent); anti-TNF ± azathioprine for moderate-severe; vedolizumab or ustekinumab for biologic-refractory."),
         ]},
        {'heading': "CD Management",
         'content': [
             b("Mild-moderate ileocolonic CD: budesonide MMX 9mg/day × 8–12 weeks (ileocolonic-release; minimal systemic exposure); avoid systemic steroids as first-line if budesonide adequate."),
             b("Moderate-severe CD: prednisolone 40mg/day tapering; early immunomodulator introduction (azathioprine 2–2.5mg/kg/day or 6-MP 1–1.5mg/kg/day) for steroid-dependent/recurrent disease."),
             b("High-risk CD (young age, perianal disease, penetrating behaviour, early steroid dependence): early biologic therapy — anti-TNF ± azathioprine combination."),
             b("Perianal CD: combined surgical drainage (EUA + seton placement) + infliximab maintenance; MRI pelvis staging before surgery; adalimumab or vedolizumab alternatives."),
         ]},
        {'heading': "Biologics",
         'content': [
             tbl(["Biologic", "Mechanism", "Induction", "Maintenance", "Notes"],
                 [["Infliximab", "Anti-TNFα (IV)", "5mg/kg at 0,2,6 wk", "5mg/kg 8-weekly", "Combination with AZA reduces immunogenicity; check drug levels + anti-drug Ab at loss of response"],
                  ["Adalimumab", "Anti-TNFα (SC)", "160/80/40mg", "40mg 2-weekly", "Convenient SC; similar efficacy to IFX"],
                  ["Vedolizumab", "Anti-α4β7 (gut-selective, IV/SC)", "300mg at 0,2,6 wk", "300mg 8-weekly", "Slower onset; safer infection/malignancy profile; preferred in elderly/prior malignancy"],
                  ["Ustekinumab", "Anti-IL12/23 p40 (IV then SC)", "IV weight-based (~6mg/kg)", "SC 90mg 8–12 wkly", "Good safety profile; increasingly first-line in CD"],
                  ["Tofacitinib", "JAK inhibitor (oral)", "10mg BD × 8 weeks", "5mg BD", "UC only; rapid onset; caution: thromboembolism/MACE in older patients; avoid in high CV risk"]]),
         ]},
        {'heading': "Monitoring on Immunosuppressives",
         'content': [
             b("Azathioprine/6-MP: FBC monthly × 3 months then 3-monthly; LFTs 3-monthly; TPMT/NUDT15 before starting; thiopurine metabolite levels (6-TGN, 6-MMP) if non-response or toxicity."),
             b("Anti-TNF agents: drug trough levels + anti-drug antibodies at loss of response to guide dose optimisation vs switch."),
             b("Annual screening on biologics: Quantiferon Gold (latent TB); hepatitis B serology; skin check for non-melanoma skin cancer."),
         ]},
        {'heading': "Surgical Indications and Cancer Surveillance",
         'content': [
             b("UC surgery: medically refractory disease, dysplasia (HGD or LGD multifocal), colorectal cancer, toxic megacolon. IPAA (ileal pouch-anal anastomosis) is restorative procedure of choice in elective setting."),
             b("CD surgery: intestinal obstruction, fistula/abscess not responding to medical therapy, cancer, growth failure in children. Conservative resection principle — preserve bowel length."),
             b("CRC surveillance: colonoscopy starting 8–10 years after diagnosis of extensive colitis (E2/E3 UC; colonic CD); frequency 1–3 yearly based on risk factors (low risk: every 3 years; high risk: every year — active inflammation, PSC, prior dysplasia, foreshortened colon, family history CRC)."),
         ]},
    ],
})

print("Chapters 109–118 data loaded.")

# ── Ch119 ──
CHAPTERS.append({
    'filename': 'Ch119_Ileostomies_Colostomies_Anastomoses.docx',
    'title': 'Ch119 — Ileostomies, Colostomies, and Intestinal Anastomoses',
    'pearls': [
        "End ileostomy outputs 800–1000 mL/day — Na and Mg depletion are the critical electrolyte complications.",
        "High-output ileostomy (>2L/day): restrict hypotonic fluids (water, tea, coffee paradoxically increase output), use St Mark's oral rehydration solution.",
        "Loop ileostomy is used to protect a distal anastomosis — typically reversed within 8–12 weeks.",
        "Parastomal hernia is the most common late stoma complication.",
        "Anastomotic leak presents Day 3–5 post-op with fever, raised CRP, peritonism — CT scan urgently.",
    ],
    'sections': [
        {'heading': "Types of Stomas",
         'content': [
             b("End colostomy: permanent — typically after abdominoperineal resection (APR) for low rectal cancer; left iliac fossa; formed stool output."),
             b("Loop colostomy: temporary decompression for obstructing left-sided cancer or colonic injury; two limbs brought to surface."),
             b("End ileostomy: after total proctocolectomy (e.g., medically refractory UC) or Hartmann's reversal not feasible; right iliac fossa; 800–1000mL liquid output/day."),
             b("Loop ileostomy: to defunctiontection a distal anastomosis (low anterior resection, IPAA, ileal pouch); reversed when anastomosis healed (8–12 weeks); water-soluble contrast enema before reversal confirms anastomotic integrity."),
         ]},
        {'heading': "Stoma Complications",
         'content': [
             tbl(["Complication", "Timing", "Management"],
                 [["Bleeding", "Early", "Direct pressure; rarely re-exploration"],
                  ["Ischaemia/Necrosis", "Early (24–48h)", "Assess viability with glass tube and torch; re-surgery if necrosis below fascia"],
                  ["Retraction", "Early–late", "Convex appliance; surgical revision if severe"],
                  ["High output", "Early–late", "Loperamide, codeine, St Mark's solution, dietary modification (see below)"],
                  ["Parastomal hernia", "Late (most common)", "Mesh repair or relocation if symptomatic"],
                  ["Prolapse", "Late", "Manual reduction; surgical revision"],
                  ["Stenosis", "Late", "Digital dilatation; surgical revision if severe"],
                  ["Skin irritation/dermatitis", "Any time", "Stoma nurse; correct appliance fit; barrier creams"]]),
         ]},
        {'heading': "High-Output Ileostomy Management",
         'content': [
             b("Definition: output >1500 mL/24h (some use >2000 mL/24h)."),
             b("Causes: early post-operative, short bowel syndrome, SIBO, CDI, bowel inflammation, dietary indiscretion."),
             b("Restrict hypotonic fluids: water, juice, carbonated drinks, tea — these generate a net sodium secretion and WORSEN output."),
             b("St Mark's oral rehydration solution: Na 90 mmol/L, glucose 20g/L — sip throughout the day (200 mL/h); replaces sodium and water."),
             b("Pharmacotherapy (in order): loperamide 4mg QDS (30 min before meals) + 2mg after each loose episode; codeine phosphate 30mg QDS; PPI to reduce gastric hypersecretion; octreotide 100 µg SC TDS if refractory."),
             b("Electrolyte replacement: Mg deficiency — oral magnesium glycerophosphate (poorly absorbed, less diarrhoea); IV Mg sulphate if severe; Na depletion — St Mark's solution + dietary salt."),
         ]},
        {'heading': "Anastomotic Complications",
         'content': [
             b("Anastomotic leak: fever + raised CRP Day 3–5 + peritonism ± faeculent drain output; CT abdomen with IV contrast to characterise; small contained leak → IV antibiotics + radiological drainage; generalised peritonitis → emergency re-laparotomy + washout ± defunctioning stoma."),
             b("Anastomotic stricture: late complication; symptoms of incomplete obstruction; endoscopic balloon dilatation for short fibrotic strictures (3–5 serial dilatations); surgical resection and re-anastomosis for long or refractory strictures."),
         ]},
    ],
})

# ── Ch120 ──
CHAPTERS.append({
    'filename': 'Ch120_IntestinalIschaemia.docx',
    'title': 'Ch120 — Intestinal Ischaemia',
    'pearls': [
        "Acute mesenteric ischaemia (AMI): surgical emergency — 'pain out of proportion to physical examination'; mortality >50% even with treatment.",
        "SMA embolism (50% of AMI) — cardiac source: AF, post-MI mural thrombus.",
        "CT angiography is the investigation of choice for AMI — immediate IV heparin while awaiting imaging.",
        "Ischaemic colitis is the most common form — typically affects watershed areas (splenic flexure, rectosigmoid junction).",
        "NOMI (non-occlusive mesenteric ischaemia) occurs in critically ill patients with low-flow states — vasopressors are a precipitant.",
    ],
    'sections': [
        {'heading': "Acute Mesenteric Ischaemia (AMI)",
         'content': [
             tbl(["Type", "% AMI", "Cause", "Key Feature"],
                 [["SMA embolism", "50%", "AF, post-MI, endocarditis, aortic atheroma", "Sudden onset; no preceding angina; source identifiable"],
                  ["SMA thrombosis", "15–25%", "Atherosclerosis at SMA origin", "Chronic post-prandial pain (intestinal angina) precedes acute event"],
                  ["Mesenteric venous thrombosis (MVT)", "5–15%", "Hypercoagulable states, portal hypertension, OCP", "Slower onset; insidious presentation; younger patients"],
                  ["NOMI", "20–30%", "Cardiogenic shock, vasopressors, post-cardiac surgery", "ICU patient; no occlusion on CT — diffuse vasoconstriction"]]),
             b("Clinical: sudden severe abdominal pain (periumbilical/diffuse), nausea, vomiting, forceful defaecation (gut emptying) initially; 'pain out of proportion' — peritonism is a LATE sign indicating infarction."),
             b("Investigations: CT angiography — defines cause (embolus vs thrombosis vs venous), bowel viability (pneumatosis, portal venous gas = poor prognostic signs), extent. Immediate IV heparin once clinical suspicion raised (do not wait for imaging if strong suspicion)."),
             b("Management: SMA embolism — catheter-directed thrombolysis or surgical embolectomy + second-look laparotomy; SMA thrombosis — surgical or endovascular revascularisation; MVT — anticoagulation (LMWH then warfarin × 6 months minimum); NOMI — treat underlying shock, intra-arterial papaverine via SMA catheter."),
         ]},
        {'heading': "Chronic Mesenteric Ischaemia",
         'content': [
             b("Post-prandial abdominal pain ('intestinal angina') 15–30 min after eating → weight loss and sitophobia (fear of eating)."),
             b("Usually requires ≥2 of 3 mesenteric vessels to be significantly stenosed."),
             b("Diagnosis: CT angiography or MRA (non-invasive); Doppler US for screening."),
             b("Treatment: endovascular revascularisation (stenting) preferred over open surgery in most centres; optimise cardiovascular risk factors; antiplatelet therapy."),
         ]},
        {'heading': "Ischaemic Colitis",
         'content': [
             b("Most common form of intestinal ischaemia; usually self-limiting."),
             b("Affects watershed areas: splenic flexure and rectosigmoid junction — these are areas of limited collateral supply."),
             b("Precipitants: hypotension (any cause), aortic surgery (ligation of IMA), constipation, cocaine, vasoconstrictors, polycythaemia."),
             b("Clinical: sudden onset crampy left-sided abdominal pain + passage of bright red blood per rectum (often within 24h of pain onset); no preceding systemic illness."),
             b("Colonoscopy: segmental oedema, haemorrhage, erosions, ulceration with ABRUPT BORDERS at watershed (pathognomonic); do not over-insufflate — risk of perforation."),
             b("CT: colonic wall thickening ± thumbprinting at watershed areas."),
             b("Management: most resolve with bowel rest, IV fluids, antibiotics (metronidazole + ciprofloxacin) for translocation risk; 10–20% progress to transmural ischaemia requiring colectomy; stricture may develop 4–6 weeks later."),
         ]},
    ],
})

# ── Ch121 ──
CHAPTERS.append({
    'filename': 'Ch121_IntestinalUlcerations_NonIBD.docx',
    'title': 'Ch121 — Intestinal Ulcerations and Strictures (Non-IBD)',
    'pearls': [
        "NSAID-induced small bowel injury is common and underdiagnosed — capsule endoscopy reveals lesions in 30–70% of chronic NSAID users.",
        "Behçet's disease: ileocaecal ulcers + oral aphthous ulcers + genital ulcers + uveitis — requires multidisciplinary management.",
        "Solitary rectal ulcer syndrome (SRUS): related to internal rectal prolapse and straining — often misdiagnosed as IBD or cancer.",
        "Diaphragm strictures of the small bowel are pathognomonic of NSAID enteropathy.",
    ],
    'sections': [
        {'heading': "NSAID Enteropathy",
         'content': [
             b("Prevalence: capsule endoscopy shows small bowel lesions in 30–70% of chronic NSAID users."),
             b("Mechanism: direct mucosal injury (uncouples mitochondrial oxidative phosphorylation) + disruption of enterohepatic bile acid circulation → increased mucosal permeability."),
             b("Clinical: iron deficiency anaemia, hypoalbuminaemia (protein-losing enteropathy), occult GI blood loss, abdominal pain; may be asymptomatic."),
             b("Characteristic lesions: circumferential 'diaphragm' strictures of small bowel (pathognomonic — not caused by any other drug or disease); may cause intermittent SBO."),
             b("Diagnosis: capsule endoscopy (patency capsule first if stricture suspected); CT/MRI enterography for strictures."),
             b("Management: stop NSAID (lesions often heal within weeks); if NSAID necessary — lowest dose COX-2 selective agent + PPI; misoprostol 200 µg QDS may reduce NSAID enteropathy; strictures may require surgical resection."),
         ]},
        {'heading': "Behçet's Disease — GI Involvement",
         'content': [
             b("Ileocaecal region is most commonly affected (85% of GI Behçet's); deep, large, punched-out ulcers; may perforate or fistulise."),
             b("Diagnosis: International Study Group criteria — oral ulcers + ≥2 of: genital ulcers, eye lesions, skin lesions, positive pathergy test."),
             b("GI complications: perforation, haemorrhage, fistula, stricture."),
             b("Treatment: 5-ASA (mesalazine) for mild; prednisolone for moderate; azathioprine for steroid-dependent; infliximab for refractory or severe/perforating disease; thalidomide (aphthous ulcers); co-management with ophthalmology and rheumatology."),
         ]},
        {'heading': "Solitary Rectal Ulcer Syndrome (SRUS)",
         'content': [
             b("Misnomer — multiple ulcers common; may appear as polypoid lesion mimicking cancer."),
             b("Location: anterior rectal wall, 3–10 cm from anal verge."),
             b("Pathogenesis: internal rectal prolapse + repetitive straining → ischaemia and mechanical trauma of anterior rectal wall."),
             b("Histology (pathognomonic): fibromuscular obliteration of lamina propria + distorted glandular architecture + muscularis mucosae thickening — distinguishes from IBD and malignancy."),
             b("Clinical: rectal bleeding (small volumes), mucus discharge, incomplete evacuation, tenesmus, perineal pain."),
             b("Treatment: biofeedback (first-line — addresses paradoxical sphincter contraction and straining behaviour); high-fibre diet + stool softeners; avoid manual digitation; topical sucralfate enemas; surgery (rectopexy) for full-thickness rectal prolapse with SRUS."),
         ]},
        {'heading': "Other Drug-Induced Intestinal Ulceration",
         'content': [
             b("Enteric-coated potassium chloride: small bowel ulcers and strictures."),
             b("Mycophenolate mofetil (in transplant patients): colonic and small bowel ulcers, diarrhoea — may mimic IBD or GVHD."),
             b("Checkpoint inhibitors (pembrolizumab, nivolumab, ipilimumab): immune-mediated enterocolitis; diarrhoea ± blood; treat with steroids (prednisolone 1mg/kg/day) ± infliximab; hold immunotherapy."),
             b("Iron tablets: direct mucosal toxicity — rare small bowel ulceration."),
         ]},
    ],
})

# ── Ch122 ──
CHAPTERS.append({
    'filename': 'Ch122_Appendicitis.docx',
    'title': 'Ch122 — Appendicitis',
    'pearls': [
        "Most common surgical emergency worldwide; peak incidence in 10–20s but any age.",
        "Alvarado score (max 10): ≥7 = high probability → proceed to surgery; 5–6 = equivocal → CT scan; <5 = low probability.",
        "CT abdomen+pelvis (IV contrast): sensitivity 94%, specificity 95% — gold standard for adults.",
        "MRI preferred over CT in pregnant women and children (no ionising radiation).",
        "Antibiotic-only treatment viable for uncomplicated appendicitis (~70% success at 1 year) but requires counselling on recurrence (~39% at 5 years).",
    ],
    'sections': [
        {'heading': "Clinical Features",
         'content': [
             b("Classic presentation: periumbilical pain migrating to right iliac fossa over 4–6h, anorexia, low-grade fever, nausea and vomiting."),
             b("Examination: RIF tenderness (McBurney's point), Rovsing's sign (LIF palpation causes RIF pain), psoas sign (pain with right hip extension — retrocaecal appendix), obturator sign (pain with right hip internal rotation — pelvic appendix)."),
             b("High or retrocaecal appendix: may present with flank pain, no RIF tenderness — atypical."),
             b("Pelvic appendix: may cause dysuria, suprapubic pain, diarrhoea; in young women, always consider gynaecological causes (ectopic pregnancy, ovarian cyst/torsion, PID)."),
             b("Elderly: atypical presentation, fever often absent — higher perforation rate at presentation."),
         ]},
        {'heading': "Alvarado Score (MANTRELS)",
         'content': [
             tbl(["Feature", "Points"],
                 [["Migration of pain to RIF", "1"],
                  ["Anorexia", "1"],
                  ["Nausea/Vomiting", "1"],
                  ["RIF tenderness", "2"],
                  ["Rebound tenderness", "1"],
                  ["Elevated temperature (>37.3°C)", "1"],
                  ["Leukocytosis (WBC >10×10⁹/L)", "2"],
                  ["Shift to left (neutrophilia)", "1"],
                  ["TOTAL", "10"]]),
             b("Score ≥7: high probability → proceed directly to surgery. Score 5–6: equivocal → CT scan. Score <5: low probability → observe, consider alternative diagnoses."),
         ]},
        {'heading': "Management",
         'content': [
             b("Uncomplicated appendicitis: laparoscopic appendicectomy — standard of care; day-case possible; lower wound infection and shorter recovery vs open."),
             b("Antibiotic-only treatment (non-operative): IV co-amoxiclav or ertapenem → oral amoxicillin-clavulanate × 10 days total; success ~70% at 1 year; recurrence ~39% at 5 years; discuss fully with patient; not appropriate if CRP >70 or faecolith present (higher failure risk)."),
             b("Appendix mass / phlegmon: conservative management with IV antibiotics (piperacillin-tazobactam) ± CT-guided drainage of abscess; interval appendicectomy at 6–8 weeks (perform colonoscopy first in adults >40 years to exclude underlying caecal tumour)."),
             b("Perforated appendicitis with generalised peritonitis: emergency laparoscopic appendicectomy + washout; IV antibiotics; if haemodynamically unstable → open surgery."),
         ]},
    ],
})

# ── Ch123 ──
CHAPTERS.append({
    'filename': 'Ch123_DiverticularDisease.docx',
    'title': 'Ch123 — Diverticular Disease of the Colon',
    'pearls': [
        "Diverticulosis in ~50% of Western population >60 years; most asymptomatic.",
        "Left-sided diverticulitis predominates in Western patients; RIGHT-sided diverticulitis is common in Asian (including Singapore) patients and can mimic appendicitis.",
        "CT abdomen+pelvis with IV contrast: sensitivity 97%, specificity 100% — gold standard for acute diverticulitis.",
        "Uncomplicated diverticulitis: emerging evidence that antibiotics may not be necessary in immunocompetent patients with mild disease.",
        "Diverticular bleeding: most common cause of acute lower GI bleeding — 70–80% stops spontaneously.",
    ],
    'sections': [
        {'heading': "Hinchey Classification",
         'content': [
             tbl(["Stage", "Description", "Management"],
                 [["Stage Ia", "Pericolic abscess (<4cm) / phlegmon", "IV antibiotics; usually resolves"],
                  ["Stage Ib", "Mesenteric abscess", "IV antibiotics; consider percutaneous drainage"],
                  ["Stage II", "Pelvic or distant abscess (>4cm)", "CT-guided percutaneous drainage + IV antibiotics; Hartmann's if fails"],
                  ["Stage III", "Purulent peritonitis", "Emergency laparotomy; Hartmann's or primary anastomosis + loop ileostomy"],
                  ["Stage IV", "Faecal peritonitis", "Emergency laparotomy; Hartmann's procedure"]]),
         ]},
        {'heading': "Acute Diverticulitis — Diagnosis and Management",
         'content': [
             b("CT abdomen+pelvis (IV contrast): peridiverticular fat stranding, bowel wall thickening, abscess, free air/fluid."),
             b("Colonoscopy contraindicated during acute attack — risk of perforation; perform at 6–8 weeks post-recovery to exclude CRC (particularly important in first episode)."),
             b("Uncomplicated (Stage Ia/mild): oral amoxicillin-clavulanate 625mg TDS × 7–10 days; emerging evidence that antibiotics may be withheld in immunocompetent patients with mild CT-confirmed uncomplicated diverticulitis; liquid diet → low-residue diet; outpatient if tolerating orals."),
             b("Stage Ib–II: IV piperacillin-tazobactam 4.5g TDS (or cefuroxime 1.5g TDS + metronidazole 500mg TDS) + CT-guided drainage if abscess >4cm."),
             b("Elective surgery (sigmoid colectomy): after recurrent episodes (≥2 attacks), immunocompromised, young age (<50), inability to exclude malignancy; laparoscopic approach preferred; discuss colostomy vs primary anastomosis risk with patient."),
         ]},
        {'heading': "Right-Sided Diverticulitis (Singapore Context)",
         'content': [
             b("Common in younger Asian patients (30–50s); right-sided single or multiple caecal/ascending colon diverticula."),
             b("Mimics appendicitis — CT differentiates (diverticulitis: fat stranding centred on diverticulum, not appendix; appendix appears normal)."),
             b("Management: primarily medical (IV antibiotics); avoid unnecessary surgery; overall good prognosis; recurrence less common than left-sided."),
         ]},
        {'heading': "Diverticular Bleeding",
         'content': [
             b("Most common cause of acute lower GI haemorrhage; typically painless, sudden onset, massive bright red rectal bleeding in older patients."),
             b("70–80% stops spontaneously; 25% rebleed."),
             b("Resuscitate: IV access, FBC, group & crossmatch, coagulation; reverse anticoagulation if possible."),
             b("Investigation: CT angiography (CTA) if haemodynamically significant — identifies active bleeding for IR-guided embolisation; colonoscopy once stabilised (within 24h after preparation if haemostasis achieved) — allows endoscopic haemostasis."),
             b("Treatment: endoscopic (clips, epinephrine + clip) or IR embolisation; surgery (segmental colectomy) if ongoing haemorrhage uncontrolled endoscopically/radiologically."),
         ]},
    ],
})

# ── Ch124 ──
CHAPTERS.append({
    'filename': 'Ch124_IBS.docx',
    'title': 'Ch124 — Irritable Bowel Syndrome (IBS)',
    'pearls': [
        "Rome IV criteria: recurrent abdominal pain ≥1 day/week for ≥3 months (onset >6 months ago) with ≥2 of: related to defaecation, change in stool frequency, change in stool form.",
        "Prevalence in Singapore ~10–15%; predominantly younger age; significant impact on quality of life.",
        "Faecal calprotectin useful to distinguish IBS from organic IBD (FCP <50 µg/g in IBS).",
        "Low FODMAP diet effective in 50–70% of IBS patients — requires dietician supervision; local food culture makes implementation challenging.",
        "Biopsychosocial model: gut-brain axis dysfunction + visceral hypersensitivity + microbiome changes all contribute.",
    ],
    'sections': [
        {'heading': "Rome IV Criteria and Red Flags",
         'content': [
             b("IBS subtypes: IBS-D (loose/watery stools ≥25% BMs), IBS-C (hard/lumpy ≥25% BMs), IBS-M (mixed), IBS-U (unclassified)."),
             b("Red flags requiring investigation: age >50 years (new symptom onset), rectal bleeding, unintentional weight loss, nocturnal diarrhoea, iron deficiency anaemia, palpable abdominal/rectal mass, family history CRC/IBD, fever."),
             b("Baseline investigations: FBC, CRP, faecal calprotectin, coeliac serology (anti-tTG IgA); stool MCS if recent antibiotic use or diarrhoea-predominant; colonoscopy if red flags present or age >50."),
         ]},
        {'heading': "Management — Stepped Care",
         'content': [
             b("Lifestyle: regular meal times (3 meals/day); avoid large meals; reduce insoluble fibre (bran — worsens symptoms); reduce gas-producing foods; limit alcohol, caffeine, carbonated drinks; regular exercise (improves GI transit and psychological wellbeing)."),
             b("Low FODMAP diet (Fermentable Oligosaccharides, Disaccharides, Monosaccharides and Polyols): eliminate high-FODMAP foods × 4–8 weeks → systematic reintroduction; supervised by GI dietician; effective in 50–70%; not a permanent diet — identify individual triggers."),
             b("Singapore context: many local foods are high FODMAP (durian, watermelon, lontong, fried rice with onion/garlic, wheat-based noodles); patient education on practical substitutions is essential."),
         ]},
        {'heading': "Drug Treatment",
         'content': [
             tbl(["Subtype", "Drug", "Dose", "Mechanism"],
                 [["IBS-D", "Loperamide", "2mg PRN", "Opioid receptor — reduces transit (symptom relief for diarrhoea; does NOT help pain)"],
                  ["IBS-D", "Rifaximin", "550mg TDS × 14 days", "Non-absorbable antibiotic; modulates gut microbiome; not subsidised in Singapore"],
                  ["IBS-C", "Ispaghula husk (Fybogel)", "1 sachet BD", "Soluble fibre — bulks and softens stool"],
                  ["IBS-C", "Macrogol (PEG)", "1–3 sachets/day", "Osmotic laxative"],
                  ["IBS-C", "Linaclotide", "290µg OD (30min before breakfast)", "GC-C agonist — secretes fluid + reduces visceral pain; approved in Singapore"],
                  ["IBS-C", "Prucalopride", "2mg OD", "5-HT4 agonist — colonic prokinetic"],
                  ["All (pain/spasm)", "Mebeverine", "135mg TDS AC", "Antispasmodic; smooth muscle relaxant"],
                  ["All (pain/spasm)", "Hyoscine butylbromide", "10mg TDS", "Anticholinergic antispasmodic"],
                  ["All", "Amitriptyline (low-dose)", "10–50mg nocte", "Neuromodulator; best for IBS-D/pain; constipating side-effect useful in IBS-D"],
                  ["All (anxiety/IBS-C)", "SSRIs", "Standard doses", "For psychological comorbidity and IBS-C (accelerate transit)"]]),
         ]},
        {'heading': "Psychological Therapies",
         'content': [
             b("Cognitive behavioural therapy (CBT): evidence-based for global IBS symptom reduction; particularly effective for severe/refractory IBS."),
             b("Gut-directed hypnotherapy: strong evidence; 12-session course; reduces visceral hypersensitivity."),
             b("Mindfulness-based stress reduction: emerging evidence."),
             b("Consider psychology/psychiatry referral for: moderate-severe IBS not responding to medical treatment, significant anxiety/depression, catastrophising behaviour, high healthcare utilisation."),
         ]},
    ],
})

# ── Ch125 ──
CHAPTERS.append({
    'filename': 'Ch125_IntestinalObstruction.docx',
    'title': 'Ch125 — Intestinal Obstruction',
    'pearls': [
        "SBO most common cause = adhesions (60–70%) from prior abdominal/pelvic surgery.",
        "LBO most common cause = colorectal cancer (60%).",
        "Closed-loop obstruction = surgical emergency — risk of rapid strangulation and perforation.",
        "Gastrografin (water-soluble contrast) challenge has both diagnostic AND therapeutic benefit in adhesional SBO — passage of contrast to colon within 24h predicts non-operative resolution.",
        "Caecal diameter >12cm on CT → risk of perforation; urgent decompression needed.",
    ],
    'sections': [
        {'heading': "Small Bowel Obstruction (SBO)",
         'content': [
             b("Causes: adhesions 60–70% (prior surgery), hernia 10–15% (inguinal, femoral, incisional — always examine groin/abdominal wall), malignancy 10%, Crohn's disease stricture, volvulus, gallstone ileus (air in biliary tree on AXR + ectopic gallstone — Rigler's triad)."),
             b("Clinical: colicky abdominal pain, vomiting (early and profuse in high SBO; late in low SBO), abdominal distension (more pronounced in low SBO), obstipation; tinkling bowel sounds (early) → silence (late = ileus or strangulation)."),
             b("Investigations: AXR — dilated small bowel loops >3cm, central position, valvulae conniventes (plicae circulares) visible, paucity of colonic gas; CT abdomen+pelvis — transition point, cause, strangulation signs (thickened wall, reduced/absent mucosal enhancement, mesenteric fat stranding, portal venous gas, pneumatosis, ascites)."),
             b("Strangulation features (surgical emergency): fever, leucocytosis, severe unremitting pain (no longer colicky), peritonism, pneumatosis on CT."),
         ]},
        {'heading': "SBO Management",
         'content': [
             b("Resuscitation: IV access, IV fluids (Normal saline or Hartmann's), urinary catheter for fluid balance, NBM, NGT for decompression."),
             b("Conservative management: 75–80% of adhesional SBO resolves non-operatively; NGT decompression, IV fluids, serial abdominal examination."),
             b("Gastrografin challenge: 100 mL water-soluble contrast via NGT or PO; AXR at 4h and 24h; contrast reaching colon = high probability of non-operative resolution; failure to progress → surgery; Gastrografin also has therapeutic effect (osmotic agent reduces bowel oedema)."),
             b("Surgical indications: complete obstruction, closed-loop obstruction, strangulation, failure of conservative management at 24–48h, hernia causing obstruction (urgent repair)."),
             b("Surgery: laparoscopic adhesiolysis preferred (lower morbidity, shorter recovery); bowel resection if non-viable bowel."),
         ]},
        {'heading': "Large Bowel Obstruction (LBO)",
         'content': [
             b("Causes: CRC 60%, diverticular stricture 10–15%, sigmoid/caecal volvulus 10%, extrinsic compression (ovarian cancer, pelvic recurrence), pseudo-obstruction (Ogilvie's)."),
             b("Clinical: absolute constipation, progressive abdominal distension, nausea, late vomiting (may be faeculent); caecal diameter >12 cm = risk of ischaemia and perforation."),
             b("Sigmoid volvulus: AXR 'coffee-bean' sign (inverted U of dilated sigmoid); CT confirms; treatment: emergency flexible sigmoidoscopy for detorsion + flatus tube placement (success ~70–90%); definitive treatment is elective sigmoid colectomy (recurrence rate high without resection)."),
             b("Caecal volvulus: AXR — dilated loop in left upper quadrant; CT confirms; treatment: usually right hemicolectomy (detorsion alone has high recurrence; colonoscopic detorsion rarely successful)."),
             b("CRC causing obstruction: options — Hartmann's (resection + end colostomy) or resection + primary anastomosis + defunctioning loop ileostomy; self-expanding metallic stent (SEMS) as bridge to elective surgery for resectable disease; SEMS palliative for metastatic/unresectable disease (risk: perforation 1–4%, stent migration 5–10%)."),
         ]},
    ],
})

# ── Ch126 ──
CHAPTERS.append({
    'filename': 'Ch126_Ileus_PseudoObstruction.docx',
    'title': 'Ch126 — Ileus and Pseudo-Obstruction Syndromes',
    'pearls': [
        "Post-operative ileus is normal for 2–3 days after abdominal surgery; investigate electrolytes (K⁺, Mg²⁺, Na⁺) and medications (opioids, anticholinergics) if prolonged.",
        "Acute colonic pseudo-obstruction (ACPO/Ogilvie's): massive colonic dilatation WITHOUT mechanical obstruction — CT to confirm before treatment.",
        "Neostigmine 2mg IV is first-line treatment for ACPO — 80–90% success; monitor for bradycardia and bronchospasm; have atropine at bedside.",
        "Caecal diameter >12 cm on AXR/CT = risk of perforation — urgent decompression.",
    ],
    'sections': [
        {'heading': "Post-Operative Ileus",
         'content': [
             b("Normal physiology: small bowel recovers motility within 24h; stomach 24–48h; colon 3–5 days after bowel surgery."),
             b("Prolonged ileus (>5 days): exclude mechanical obstruction, anastomotic leak, intra-abdominal abscess, electrolyte disturbance (K⁺ <3.5, Mg²⁺ <0.7, Na⁺ imbalance), medications (opioids most common, anticholinergics, calcium channel blockers)."),
             b("Prevention: enhanced recovery after surgery (ERAS) protocols: early mobilisation, early enteral nutrition, chewing gum (stimulates motilin release), minimise opioids (multimodal analgesia — NSAID/paracetamol + regional blocks), minimise nasogastric tubes."),
             b("Alvimopan: peripheral µ-opioid receptor antagonist; reduces ileus after bowel resection when opioids used; not widely available; monitor for cardiovascular events."),
         ]},
        {'heading': "Acute Colonic Pseudo-Obstruction (ACPO / Ogilvie's Syndrome)",
         'content': [
             b("Definition: acute massive colonic dilatation without mechanical cause — imbalance between sympathetic (inhibitory) and parasympathetic (excitatory) colonic innervation."),
             b("Risk factors: elderly hospitalised patients; post-surgery (especially orthopaedic — hip/knee); post-trauma; sepsis; cardiac/respiratory illness; opioids; electrolyte disturbance."),
             b("Diagnosis: CT abdomen to exclude mechanical obstruction (transition point absent in ACPO); caecal diameter >12 cm on AXR/CT = high perforation risk → urgent intervention."),
             b("Management: correct electrolytes; stop anticholinergics and opioids; ambulation; NG tube decompression; nil by mouth."),
             b("Neostigmine: 2mg IV over 3–5 minutes; mechanism: acetylcholinesterase inhibitor → increases ACh → stimulates colonic motility; success 80–90%; contraindications: bradycardia, bronchospasm, bowel obstruction, renal failure; have atropine 0.6mg IV ready; monitor ECG and vital signs for 30 min; may repeat once."),
             b("Colonoscopic decompression: if neostigmine fails or contraindicated; flatus tube placement; success ~70–80%; recurrence common."),
             b("Surgery (caecostomy or colostomy): only if all medical measures fail or perforation suspected."),
         ]},
        {'heading': "Chronic Intestinal Pseudo-Obstruction (CIPO)",
         'content': [
             b("Causes: connective tissue disease (scleroderma — most common), amyloidosis, mitochondrial cytopathy, idiopathic (enteric neuromuscular disease), paraneoplastic (anti-Hu antibodies — SCLC)."),
             b("Clinical: recurrent/chronic abdominal distension, nausea, early satiety, constipation or diarrhoea; significantly impairs quality of life."),
             b("Management: nutritional support (TPN or enteral feeding); prokinetics (erythromycin, metoclopramide, domperidone, prucalopride); antibiotics for SIBO (common in CIPO); neostigmine for acute-on-chronic flares; octreotide may paradoxically help in scleroderma (suppresses postprandial motility, allows MMC to return); rarely surgery."),
         ]},
    ],
})

print("Chapters 119–126 data loaded.")

# ── Ch127 ──
CHAPTERS.append({
    'filename': 'Ch127_SmallIntestineTumors.docx',
    'title': 'Ch127 — Tumours of the Small Intestine',
    'pearls': [
        "Small bowel tumours are rare but 50% of symptomatic ones are malignant.",
        "Most common malignancies: adenocarcinoma (duodenum), NET/carcinoid (ileum), GIST (jejunum/ileum), lymphoma.",
        "Carcinoid syndrome (flushing, diarrhoea, bronchospasm, right-sided valve disease) occurs when liver metastases allow serotonin to bypass hepatic metabolism.",
        "GIST: c-KIT (CD117) positive — imatinib 400mg OD for metastatic/unresectable disease.",
        "Diagnosis is often delayed due to non-specific symptoms — capsule endoscopy and CT/MRI enterography are key.",
    ],
    'sections': [
        {'heading': "Benign Tumours",
         'content': [
             b("Adenoma: duodenal adenomas associated with FAP (periampullary location) — annual duodenoscopy in FAP; endoscopic polypectomy for <3cm lesions; surgical excision for large/ampullary."),
             b("Lipoma, leiomyoma, haemangioma: usually asymptomatic; may cause intussusception (adult intussusception = malignant lead point until proven otherwise → surgery)."),
         ]},
        {'heading': "Small Bowel Adenocarcinoma",
         'content': [
             b("Most common malignant small bowel tumour; duodenum (D2–D3) > jejunum > ileum."),
             b("Risk factors: coeliac disease, Crohn's disease (ileal), FAP (periampullary), Lynch syndrome, Peutz-Jeghers syndrome."),
             b("Symptoms: obstruction, occult GI bleeding, iron deficiency anaemia, weight loss, jaundice (periampullary)."),
             b("Diagnosis: CT/MRI enterography; ERCP for periampullary; upper GI endoscopy (duodenal); capsule endoscopy (jejunoileal)."),
             b("Treatment: surgical resection (Whipple's for periampullary/duodenal; segmental resection for jejunoileal); adjuvant FOLFOX × 6 months for Stage III; poor prognosis if metastatic."),
         ]},
        {'heading': "Neuroendocrine Tumours (NET) / Carcinoid",
         'content': [
             b("Most common in ileum (terminal ileum); also appendix, rectum."),
             b("Serotonin-secreting; carcinoid syndrome (diarrhoea, flushing, bronchospasm, right-sided valve disease — tricuspid regurgitation) only when liver metastases present (serotonin bypasses hepatic metabolism)."),
             b("Diagnosis: serum chromogranin A (elevated); urine 5-HIAA (elevated in carcinoid syndrome); CT/MRI — primary + liver metastases; ⁶⁸Ga-DOTATATE PET-CT (superior somatostatin receptor imaging)."),
             b("Treatment: surgical resection for localised disease; somatostatin analogues (octreotide LAR 30mg IM monthly or lanreotide autogel 120mg SC monthly) for symptom control and antiproliferative effect (PROMID trial, CLARINET trial); PRRT (Lutetium-177 DOTATATE) for somatostatin receptor positive metastatic NETs (NETTER-1 trial); everolimus for progressive disease."),
         ]},
        {'heading': "Gastrointestinal Stromal Tumour (GIST)",
         'content': [
             b("c-KIT (CD117) positive in 95%; PDGFRA mutation in 5%; rare c-KIT and PDGFRA wild-type (NF1, SDH-deficient, BRAF mutated)."),
             b("Most common mesenchymal tumour of GI tract; small bowel GISTs often present with GI bleeding (hypervascular tumour bleeds into lumen)."),
             b("CT: hypervascular mass, may be exophytic; PET-CT useful for staging."),
             b("Treatment: surgical resection (R0) for localised; imatinib 400mg OD for metastatic/unresectable or adjuvant in high-risk (size >5cm, high mitotic rate); sunitinib for imatinib-refractory; regorafenib third-line; mutational analysis (exon 9 vs 11 of c-KIT) guides imatinib dose."),
         ]},
        {'heading': "Primary Small Bowel Lymphoma",
         'content': [
             b("Jejunum most commonly involved; increased risk in immunocompromised (HIV, immunosuppression after transplant), coeliac disease (EATL — enteropathy-associated T-cell lymphoma)."),
             b("Types: DLBCL, MALT, Burkitt's, EATL (poor prognosis)."),
             b("MALT lymphoma of stomach — treat H. pylori first (antibiotic eradication alone may cause regression in localised H. pylori-associated MALT)."),
             b("DLBCL: R-CHOP (rituximab + cyclophosphamide + doxorubicin + vincristine + prednisolone); EATL: poor prognosis — autologous stem cell transplant in first remission considered; treatment usually includes surgery + R-CHOP-like regimen."),
         ]},
    ],
})

# ── Ch128 ──
CHAPTERS.append({
    'filename': 'Ch128_ColonicPolypsPolyposisSyndromes.docx',
    'title': 'Ch128 — Colonic Polyps and Polyposis Syndromes',
    'pearls': [
        "Adenomatous polyps are the major CRC precursors — risk of advanced neoplasia: size >10mm, villous histology, high-grade dysplasia, ≥3 polyps.",
        "SSA (sessile serrated adenoma/lesion): flat, pale, right-sided; BRAF V600E mutation; serrated pathway to CRC — associated with MSI-H tumours.",
        "FAP: APC gene mutation; hundreds to thousands of adenomas; CRC risk near 100% untreated — prophylactic colectomy by age 20.",
        "Lynch syndrome: MMR gene mutations (MLH1, MSH2, MSH6, PMS2); 50–80% lifetime CRC risk; universal MMR/MSI testing of ALL CRC specimens recommended.",
        "Aspirin 600mg OD reduces CRC risk in Lynch syndrome (CAPP2 trial).",
    ],
    'sections': [
        {'heading': "Adenoma Surveillance Intervals",
         'content': [
             tbl(["Finding at index colonoscopy", "Next surveillance colonoscopy"],
                 [["1–2 tubular adenomas <10mm, LGD", "5 years"],
                  ["3–4 adenomas <10mm", "3 years"],
                  ["≥5 adenomas OR ≥3 adenomas with 1 ≥10mm", "1 year"],
                  ["Any adenoma ≥10mm", "3 years"],
                  ["Villous or tubulovillous adenoma", "3 years"],
                  ["High-grade dysplasia (HGD) in adenoma", "1 year (after confirmation of complete resection)"],
                  ["Piecemeal resection of large polyp", "3–6 months (to confirm complete resection)"]]),
         ]},
        {'heading': "Sessile Serrated Lesion (SSL) / Sessile Serrated Adenoma (SSA)",
         'content': [
             b("Flat, pale, right-sided; difficult to detect — often missed on colonoscopy; pale rim of bubbles at edge of lesion."),
             b("BRAF V600E mutation; epigenetic silencing via CpG island methylator phenotype (CIMP) → MLH1 silencing → MSI-H → CRC (serrated pathway, ~20% of CRCs)."),
             b("SSL <10mm without dysplasia: surveillance at 3 years; SSL ≥10mm or with dysplasia: 1 year (same as advanced adenoma)."),
             b("Traditional serrated adenoma (TSA): rare; left-sided; pinecone appearance; KRAS or BRAF mutation; treat as adenoma."),
         ]},
        {'heading': "FAP (Familial Adenomatous Polyposis)",
         'content': [
             b("APC gene mutation (chromosome 5q21); autosomal dominant; >100 adenomas (classical) or 20–100 (attenuated FAP)."),
             b("CRC risk approaches 100% by 4th decade if untreated; onset of adenomas typically in teenage years."),
             b("Extracolonic manifestations: CHRPE (congenital hypertrophy of retinal pigment epithelium), desmoid tumours (fibromatosis — can be life-threatening, esp. after surgery), periampullary/duodenal adenomas (second most common cause of death in FAP), papillary thyroid cancer, epidermoid cysts, mandibular osteomas."),
             b("Surveillance: annual flexible sigmoidoscopy from age 12–14 until polyposis documented → plan colectomy; duodenoscopy for periampullary polyps from age 25–30."),
             b("Surgery: total colectomy + ileorectal anastomosis (IRA) if rectal polyps manageable; proctocolectomy + IPAA if severe rectal polyposis; laparoscopic preferred."),
             b("Chemoprevention: sulindac (COX-2 inhibitor) or celecoxib — reduces polyp burden but does NOT replace surveillance and surgery."),
         ]},
        {'heading': "Lynch Syndrome (HNPCC)",
         'content': [
             b("Autosomal dominant mutations in MMR genes: MLH1, MSH2 (most common), MSH6, PMS2; EPCAM deletions (silences MSH2)."),
             b("Lifetime CRC risk 50–80% (MLH1/MSH2); 20–30% (MSH6/PMS2); polyps often synchronous, right-sided, MSI-H."),
             b("Extracolonic cancers: endometrial (40–60%), ovarian, gastric, urological (upper tract urothelial), hepatobiliary, small bowel."),
             b("Identification: Amsterdam II criteria; revised Bethesda guidelines; universal IHC/MSI testing of all CRC (recommended by Singapore Cancer Network)."),
             b("Surveillance: colonoscopy every 1–2 years from age 20–25 (or 5 years before youngest CRC in family); annual gynaecological surveillance from age 30–35 (transvaginal USS + endometrial biopsy); consider URS surveillance."),
             b("Chemoprevention: aspirin 600mg OD significantly reduces CRC incidence in Lynch syndrome (CAPP2 trial — 63% risk reduction after 2 years use)."),
         ]},
        {'heading': "Other Polyposis Syndromes",
         'content': [
             tbl(["Syndrome", "Gene", "Polyp Type", "Key Features", "Cancer Risk"],
                 [["Peutz-Jeghers (PJS)", "STK11/LKB1", "Hamartoma", "Mucocutaneous melanin pigmentation (lips, buccal mucosa, fingers); intussusception", "CRC, breast, ovarian, pancreatic, gastric"],
                  ["Juvenile Polyposis (JPS)", "SMAD4 / BMPR1A", "Hamartoma", "SMAD4 mutation → hereditary haemorrhagic telangiectasia; massive GI bleeding risk", "CRC, gastric"],
                  ["PTEN hamartoma (Cowden)", "PTEN", "Hamartoma + adenoma", "Macrocephaly, facial trichilemmomas, Lhermitte-Duclos disease", "Breast, thyroid, endometrial"]]),
         ]},
    ],
})

# ── Ch129 ──
CHAPTERS.append({
    'filename': 'Ch129_ColorectalCancer.docx',
    'title': 'Ch129 — Colorectal Cancer (CRC)',
    'pearls': [
        "#1 cancer in Singapore men; #2 overall — incidence rising rapidly in Asian populations.",
        "40% present at Stage III/IV — late presentation underscores importance of screening.",
        "MSI-H/dMMR CRC: responds to pembrolizumab immunotherapy — test ALL metastatic CRC for MMR/MSI status.",
        "RAS/RAF wild-type only: eligible for anti-EGFR therapy (cetuximab, panitumumab).",
        "Rectal cancer: MRI pelvis for local staging, total mesorectal excision (TME) is surgical standard, neoadjuvant CRT for T3/T4 or N+.",
    ],
    'sections': [
        {'heading': "Risk Factors and Clinical Features",
         'content': [
             b("Major risk factors: advancing age, family history CRC, Lynch syndrome/FAP, long-standing extensive IBD, adenomatous polyps, obesity, physical inactivity, red/processed meat consumption, smoking, alcohol, type 2 diabetes."),
             b("Right colon: vague symptoms, iron deficiency anaemia (occult bleeding), palpable mass (may be large), weight loss — often late diagnosis."),
             b("Left colon/sigmoid: change in bowel habit (narrowing of stool calibre, alternating constipation/diarrhoea), bright red or dark PR blood, obstructive symptoms."),
             b("Rectum: tenesmus, mucus PR, bright red bleeding, sensation of incomplete emptying, urgency."),
             b("New rectal bleeding or change in bowel habit at any age >50 requires colonoscopy; urgent referral (2-week wait pathway equivalent in Singapore)."),
         ]},
        {'heading': "Diagnosis and Staging",
         'content': [
             b("Colonoscopy + biopsy: gold standard for diagnosis; complete colonoscopy to exclude synchronous cancers (3–5% incidence)."),
             b("CT chest/abdomen/pelvis: staging for distant metastases (lung, liver, peritoneum); define T and N stage."),
             b("MRI pelvis: T and N staging for RECTAL cancer — circumferential resection margin (CRM), extramural vascular invasion (EMVI), mesorectal fascia involvement."),
             b("PET-CT: before resection of isolated distant metastases (liver, lung) to exclude occult disease."),
             tbl(["Stage", "TNM", "5-year Survival"],
                 [["Stage I", "T1–2, N0, M0", ">90%"],
                  ["Stage II", "T3–4, N0, M0", "70–85%"],
                  ["Stage III", "Any T, N1–2, M0", "40–80%"],
                  ["Stage IV", "Any T, Any N, M1", "10–15% (improving with targeted therapy)"]]),
         ]},
        {'heading': "Management by Stage",
         'content': [
             b("Stage I: surgical resection alone (laparoscopic preferred — equivalent oncological outcomes, faster recovery)."),
             b("Stage II high-risk (T4, perforation, obstruction, <12 lymph nodes retrieved, poor differentiation, LVI, PNI, MSS): adjuvant FOLFOX × 6 months; Stage II MSI-H has good prognosis without chemotherapy."),
             b("Stage III: adjuvant FOLFOX × 6 months or CAPOX × 3 months (equally effective for Stage III — IDEA trial); start within 8 weeks of surgery."),
             b("Stage IV (metastatic): palliative chemotherapy; FOLFOX ± bevacizumab or FOLFIRI ± bevacizumab as first-line for RAS-mutant; FOLFOX/FOLFIRI ± cetuximab or panitumumab for RAS/BRAF wild-type (left-sided primary preferred for anti-EGFR); FOLFOXIRI + bevacizumab for fit patients (high response rate for conversion); pembrolizumab first-line for MSI-H/dMMR metastatic CRC (KEYNOTE-177)."),
             b("Oligometastatic disease (resectable liver/lung mets): resection with curative intent; ablation (RFA, microwave); CRS + HIPEC for selected peritoneal metastases."),
         ]},
        {'heading': "Rectal Cancer Specific",
         'content': [
             b("MRI pelvis staging: T stage, N stage, CRM (margin <1mm = involved), EMVI, mesorectal fascia — guides neoadjuvant therapy decision."),
             b("Total mesorectal excision (TME): surgical gold standard — en bloc resection of rectum within the mesorectal envelope; reduces local recurrence from ~30% to <10%."),
             b("Neoadjuvant therapy for T3–T4 or N+ rectal cancer: long-course CRT (50Gy + capecitabine × 5–6 weeks → surgery 8–12 weeks later) OR short-course RT (25Gy × 5 days → surgery 1 week later or 8–12 weeks if good response sought)."),
             b("Total neoadjuvant therapy (TNT): induction FOLFOX/CAPOX before CRT (or after) — increases complete response rates; allows watch-and-wait in clinical complete responders."),
             b("Watch-and-wait (W&W): clinical complete response after CRT → organ-preservation (defer surgery); strict surveillance with DRE + MRI pelvis + endoscopy; regrowth salvaged with surgery."),
             b("Low anterior resection vs APR: distance from anal verge determines; TME with sphincter-preserving LAR if ≥1cm from sphincter complex; APR with permanent colostomy for low tumours involving sphincter complex."),
         ]},
        {'heading': "Post-Resection Surveillance and Screening",
         'content': [
             b("CEA: 3-monthly × 3 years then 6-monthly × 2 years; CT chest/abdomen/pelvis: 6-monthly × 3 years then annually × 2 years."),
             b("Colonoscopy: 1 year post-resection; if clear → 3-yearly surveillance."),
             b("Singapore MOH CRC Screening: FIT (faecal immunochemical test) biennially for ages 50–75 (subsidised under Screen for Life); positive FIT → colonoscopy within 6 weeks; colonoscopy recommended for high-risk individuals (family history, Lynch, IBD, polyposis) at younger ages and shorter intervals."),
         ]},
    ],
})

# ── Ch130 ──
CHAPTERS.append({
    'filename': 'Ch130_OtherDiseasesOfColon.docx',
    'title': 'Ch130 — Other Diseases of the Colon',
    'pearls': [
        "Microscopic colitis: normal colonoscopy but biopsy shows inflammation — common cause of watery diarrhoea in middle-aged/elderly women; budesonide 9mg/day is most effective treatment.",
        "NSAIDs (especially diclofenac), SSRIs (sertraline), PPIs (lansoprazole), and statins are associated with microscopic colitis.",
        "Sigmoid volvulus: 'coffee-bean' sign on AXR; colonoscopic detorsion is first-line — recurrence high without elective resection.",
        "Melanosis coli: benign brown discolouration from anthraquinone laxatives (senna) — reversible; adenomas appear pale against this background.",
        "Pneumatosis coli: gas cysts in bowel wall — exclude ischaemia; benign form treated with high-flow oxygen.",
    ],
    'sections': [
        {'heading': "Microscopic Colitis (MC)",
         'content': [
             b("Subtypes: collagenous colitis (CC — subepithelial collagen band >10 µm on trichrome stain) and lymphocytic colitis (LC — ≥20 intraepithelial lymphocytes per 100 epithelial cells)."),
             b("Clinical: watery, non-bloody diarrhoea (5–10 stools/day including nocturnal); abdominal cramps; normal colonoscopy and CT; weight loss uncommon."),
             b("Associated conditions: autoimmune thyroid disease, coeliac disease, rheumatoid arthritis."),
             b("Drug triggers: NSAIDs (most common), SSRIs (especially sertraline, paroxetine), PPIs (especially lansoprazole — not all PPIs equally culpable), statins, ranitidine (withdrawn from market but historical), beta-blockers."),
             b("Biopsy: take multiple biopsies from entire colon (right, transverse, left) — patchy; do not rely on rectal biopsy alone."),
             tbl(["Treatment", "Indication", "Dose"],
                 [["Stop offending drug", "First step always", "—"],
                  ["Budesonide (preferred)", "Moderate-severe or drug not identified", "9mg/day × 6–8 weeks; taper; maintenance 6mg/day if relapsing"],
                  ["Bismuth subsalicylate", "Mild cases", "525mg TDS × 8 weeks"],
                  ["Cholestyramine", "Mild; bile acid component", "4g TDS"],
                  ["Mesalazine", "Moderate; alternative", "2–3g/day"],
                  ["Azathioprine / budesonide", "Steroid-dependent refractory", "Combination"]]),
         ]},
        {'heading': "Colonic Volvulus",
         'content': [
             b("Sigmoid volvulus (80%): elderly, institutionalised, chronic constipation, neuropsychiatric medications; AXR — 'coffee-bean' sign (inverted U, no haustration, pointing to right upper quadrant); CT confirms."),
             b("Treatment: emergency flexible sigmoidoscopy for detorsion + flatus tube 24–48h (success rate 70–90%); assess mucosal viability during procedure; necrotic mucosa → emergency Hartmann's; elective sigmoid colectomy within same admission (high recurrence without resection)."),
             b("Caecal volvulus: younger patients; AXR — dilated caecum in left upper quadrant; CT confirms; treatment: usually right hemicolectomy (colonoscopic detorsion has low success rate)."),
         ]},
        {'heading': "Melanosis Coli",
         'content': [
             b("Benign brown/black discolouration of colonic mucosa from lipofuscin accumulation in macrophages — result of anthraquinone laxative use (senna, cascara)."),
             b("Clinically significant: adenomatous polyps appear pale/white against the brown background — may be easier to identify."),
             b("Reversible with cessation of anthraquinone laxatives (over months)."),
             b("No malignant potential; no treatment required."),
         ]},
        {'heading': "Pneumatosis Coli",
         'content': [
             b("Gas cysts in the bowel wall; can affect any segment; right colon most common."),
             b("Benign (chronic) form: asymptomatic or mild symptoms (mucus, diarrhoea, rectal bleeding); may be incidental finding on CT or colonoscopy."),
             b("Life-threatening causes: bowel ischaemia, necrotising enterocolitis (NEC), severe infection — always exclude clinically."),
             b("CT: gas in bowel wall (linear or cystic); portal venous gas = associated with ischaemia → emergency laparotomy."),
             b("Benign form treatment: high-flow oxygen therapy (100% O₂ via NRB mask for 3–5 days) causes resolution by diluting intraluminal nitrogen; hyperbaric oxygen for refractory cases; surgery not required."),
         ]},
        {'heading': "Solitary Rectal Ulcer Syndrome (SRUS)",
         'content': [
             b("See Ch121 for detailed management. Summary: biofeedback first-line; topical sucralfate; fibre; avoid straining; rectopexy for full-thickness prolapse."),
         ]},
    ],
})

# ── Ch131 ──
CHAPTERS.append({
    'filename': 'Ch131_AnalDiseases.docx',
    'title': 'Ch131 — Anal Diseases',
    'pearls': [
        "Haemorrhoids graded I–IV; rubber band ligation (RBL) is the most effective non-surgical treatment for Grade I–III.",
        "Anal fissure: most are in posterior midline — LATERAL position suggests Crohn's, TB, malignancy, or STI.",
        "Anorectal abscess: drain urgently — do NOT treat with antibiotics alone (will not resolve without drainage).",
        "Anal SCC: HPV-related; Nigro protocol (5-FU + mitomycin C + radiotherapy) — 80% complete response; surgery (APR) reserved for residual/recurrent disease.",
        "Digital rectal examination (DRE) and anoscopy are essential first-line assessments for all anal complaints.",
    ],
    'sections': [
        {'heading': "Haemorrhoids",
         'content': [
             tbl(["Grade", "Features", "Treatment"],
                 [["Grade I", "Bleed on defaecation; no prolapse", "High-fibre diet, topical creams, RBL"],
                  ["Grade II", "Prolapse during defaecation; spontaneously reduce", "RBL (most effective)"],
                  ["Grade III", "Prolapse; require manual reduction", "RBL or haemorrhoidectomy"],
                  ["Grade IV", "Permanently prolapsed; cannot be reduced", "Haemorrhoidectomy (excisional or stapled)"]]),
             b("Internal haemorrhoids: above dentate line; columnar epithelium → painless; bleed bright red."),
             b("External haemorrhoids: below dentate line; squamous epithelium → painful when thrombosed."),
             b("Thrombosed external haemorrhoid: severe acute pain, tense blue perianal lump; incision and evacuation within 72h (most effective) — after 72h, pain usually beginning to resolve and conservative management preferred (sitz baths, analgesia, stool softeners)."),
             b("RBL: applied to haemorrhoid pedicle 1cm above dentate line; 2–3 bands per session; repeat every 4–6 weeks; success 80–90% for Grade I–III; complications: pain (mild), haemorrhage, rare pelvic sepsis (urgent if fever, urinary retention)."),
         ]},
        {'heading': "Anal Fissure",
         'content': [
             b("Acute (<6 weeks): fresh tear; smooth edges; posterior midline (90%) or anterior midline (10%)."),
             b("Chronic (>6 weeks): indurated raised edges, sentinel skin tag distally, hypertrophied anal papilla proximally, exposed white internal anal sphincter fibres at base."),
             b("Lateral fissure: investigate for Crohn's disease, TB, syphilis, HIV, anal cancer."),
             b("Pathophysiology: high internal anal sphincter (IAS) tone → reduced blood supply → ischaemia → poor healing."),
             tbl(["Treatment", "Details", "Success Rate"],
                 [["Fibre + stool softeners", "First step; bulks stool, reduces trauma", "~40% acute fissures heal"],
                  ["GTN 0.4% cream BD (6 weeks)", "Chemical sphincterotomy; headache side effect; apply perianally", "50–60%; high recurrence"],
                  ["Diltiazem 2% cream BD (8 weeks)", "Preferred over GTN — fewer headaches; similar efficacy", "50–65%; lower recurrence"],
                  ["Botulinum toxin injection", "20 IU into each side of IAS; day procedure; temporary", "~60–80%; ~30% recurrence"],
                  ["Lateral internal sphincterotomy (LIS)", "Most effective; surgical; risk of incontinence (1–5% permanent)", "90–95% success"]]),
         ]},
        {'heading': "Anorectal Abscess and Fistula-in-Ano",
         'content': [
             b("Cryptoglandular origin (90%): infection of anal glands in intersphincteric space spreads to form abscess."),
             b("Abscess types: perianal (most common — subcutaneous), ischiorectal, intersphincteric, supralevator; horseshoe abscess spans both sides."),
             b("Treatment: urgent incision and drainage (I&D) under anaesthesia; antibiotics not a substitute for drainage but added for immunocompromised, systemic sepsis, or extensive cellulitis; examine for fistula at EUA."),
             b("Fistula-in-ano: Parks classification — intersphincteric (most common, 70%), transsphincteric (25%), suprasphincteric, extrasphincteric (rare)."),
             tbl(["Fistula Type", "Sphincter Involvement", "Treatment"],
                 [["Intersphincteric", "Crosses IAS only", "Fistulotomy (low risk of incontinence)"],
                  ["Low transsphincteric (<30% EAS)", "Crosses lower EAS", "Fistulotomy"],
                  ["High transsphincteric (>30% EAS)", "Crosses upper EAS", "Seton (loose — draining; cutting — staged sphincterotomy); LIFT; VAAFT; fibrin glue"],
                  ["Suprasphincteric / extrasphincteric", "Over/outside puborectalis", "Seton; complex repair; diversion"]]),
             b("Crohn's perianal fistula: combined infliximab (induction + maintenance) + surgical seton drainage; MRI pelvis for staging; local repair futile without disease quiescence; fistula plug, LIFT, video-assisted anal fistula treatment (VAAFT) in selected cases."),
         ]},
        {'heading': "Anal Cancer",
         'content': [
             b("Squamous cell carcinoma (SCC) 90%; HPV 16 and 18 are the main oncogenic drivers."),
             b("Risk factors: HIV (40× increased risk), MSM, immunosuppression (transplant), prior anogenital warts/HPV disease, history of cervical/vulval/vaginal cancer, smoking."),
             b("Clinical: bleeding per rectum, anal pain, perianal mass, change in bowel habit, incontinence; often initially misdiagnosed as haemorrhoid."),
             b("Biopsy required: DRE + anoscopy; biopsy any suspicious lesion."),
             b("Staging: MRI pelvis (T and N staging); CT chest/abdomen (distant metastases); inguinal LN palpation."),
             b("Nigro protocol (definitive chemoradiotherapy): 5-FU + mitomycin C (or cisplatin) + concurrent pelvic radiotherapy (45–59.4 Gy); 80% complete response; preserves anal sphincter; do NOT perform upfront surgery."),
             b("Surgery (APR): reserved for residual disease at 8–12 weeks post-CRT (biopsy-confirmed) or late recurrence; results in permanent colostomy."),
             b("HPV vaccination: Gardasil 9 prevents HPV 16/18 → effective prevention of anal and cervical cancers; recommended for MSM up to age 45 in Singapore if not previously vaccinated."),
         ]},
        {'heading': "Pilonidal Sinus",
         'content': [
             b("Midline natal cleft; predominantly young men; chronic hair follicle infection + foreign body reaction."),
             b("Acute abscess: I&D; do NOT perform definitive excision in acute setting."),
             b("Chronic/recurrent pilonidal disease: excision with off-midline closure (Karydakis flap or Bascom procedure) — significantly lower recurrence than midline closure."),
             b("Lay-open and granulation: simple pilonidal pit with minimal tracks; good for limited disease."),
         ]},
        {'heading': "Pruritus Ani",
         'content': [
             b("Causes: anorectal pathology (haemorrhoids, fissure, faecal incontinence with soiling, prolapse), skin disease (psoriasis, lichen sclerosus, contact dermatitis), infection (Candida albicans, Enterobius vermicularis), dietary triggers (caffeine, citrus, tomato, spices, beer), idiopathic (commonest after secondary causes excluded)."),
             b("Treatment: address underlying cause first; strict perianal hygiene (wash after defaecation, pat dry, avoid vigorous wiping); avoid over-washing with soap (disrupts skin barrier); short-course topical hydrocortisone 1% cream BD × 2–4 weeks (only if dermatitis/lichen present — avoid chronic use → skin atrophy); treating Candida: clotrimazole 1% cream; barrier cream (zinc oxide); dietary diary to identify triggers."),
         ]},
    ],
})

print("Chapters 127–131 data loaded.")

# ── Main execution ────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    total = len(CHAPTERS)
    print(f"\nGenerating {total} chapter documents...")
    print(f"Output directory: {OUT_DIR}\n")

    for i, chapter in enumerate(CHAPTERS, 1):
        doc = build_doc(chapter)
        out_path = os.path.join(OUT_DIR, chapter['filename'])
        doc.save(out_path)
        print(f"[{i:02d}/{total}] Saved: {chapter['filename']}")

    print(f"\nAll {total} documents generated successfully.")


if __name__ == '__main__':
    main()
