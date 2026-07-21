"""
HyESys H100 vs H125 ESS & PCS Spec Comparison Builder
Saves to: C:/Users/JasonOng/OneDrive - Advancer Global Ltd/AST BD/HyESys Dept/3. Hardware (PCS.BATT)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = (r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd\AST BD"
          r"\HyESys Dept\3. Hardware (PCS.BATT)"
          r"\HyESys H100 vs H125 ESS-PCS Spec Comparison.docx")

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1A, 0x28, 0x58)   # AST navy
GOLD    = RGBColor(0xF0, 0xA5, 0x00)   # AST gold
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LGREY   = RGBColor(0xF2, 0xF4, 0xF8)   # light row
MGREY   = RGBColor(0xD9, 0xDD, 0xE8)   # section header row
DKGREY  = RGBColor(0x3A, 0x4A, 0x6A)   # sub-header row (dark blue-grey)
GREEN   = RGBColor(0x00, 0x70, 0x50)
RED     = RGBColor(0xC0, 0x20, 0x20)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '%02X%02X%02X' % (rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_fmt(para, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0, italic=False):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    if level == 1:
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(16)
        # underline with gold
        r2 = p.add_run()
        r2.font.size = Pt(4)
    elif level == 2:
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DKGREY

def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Note: {text}")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

# ── Table builder ──────────────────────────────────────────────────────────────
COLD_WIDTH  = Cm(6.0)
COL1_WIDTH  = Cm(7.0)
COL2_WIDTH  = Cm(7.2)
COL3_WIDTH  = Cm(7.2)

def make_table(doc, rows_data):
    """
    rows_data: list of dicts
      { 'type': 'header'|'section'|'row'|'note',
        'cols': [str, str, str],     # for 3-col rows
        'text': str,                 # for section/note
        'highlight': [None|'green'|'red'|'yellow', ...] }
    """
    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    for rd in rows_data:
        rtype = rd.get('type', 'row')

        if rtype == 'section':
            row = tbl.add_row()
            row.height = Cm(0.65)
            for i, cell in enumerate(row.cells):
                set_cell_bg(cell, MGREY)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            p = merged.paragraphs[0]
            r = p.add_run(rd['text'])
            r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
            p.paragraph_format.left_indent = Pt(4)
            continue

        if rtype == 'header':
            row = tbl.add_row()
            row.height = Cm(0.75)
            cols_text = rd.get('cols', ['Parameter', 'H100 — v1.0', 'H125 — v2.0'])
            widths = [COLD_WIDTH, COL2_WIDTH, COL3_WIDTH]
            for i, (cell, txt, w) in enumerate(zip(row.cells, cols_text, widths)):
                cell.width = w
                set_cell_bg(cell, NAVY)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                r = p.add_run(txt)
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if rtype == 'row':
            row = tbl.add_row()
            row.height = Cm(0.58)
            cols_text = rd.get('cols', ['', '', ''])
            highlight  = rd.get('highlight', [None, None, None])
            alt_bg     = rd.get('alt', False)
            widths     = [COLD_WIDTH, COL2_WIDTH, COL3_WIDTH]

            for i, (cell, txt, hl, w) in enumerate(zip(row.cells, cols_text, highlight, widths)):
                cell.width = w
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # background
                if hl == 'green':
                    set_cell_bg(cell, RGBColor(0xE6, 0xF7, 0xEF))
                elif hl == 'red':
                    set_cell_bg(cell, RGBColor(0xFD, 0xEC, 0xEC))
                elif hl == 'yellow':
                    set_cell_bg(cell, RGBColor(0xFF, 0xF8, 0xE0))
                elif alt_bg and i > 0:
                    set_cell_bg(cell, LGREY)
                elif i == 0:
                    set_cell_bg(cell, RGBColor(0xF6, 0xF7, 0xFA))

                p = cell.paragraphs[0]
                p.paragraph_format.left_indent = Pt(4)
                r = p.add_run(txt)
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True; r.font.color.rgb = RGBColor(0x30, 0x40, 0x60)
                else:
                    r.font.color.rgb = RGBColor(0x20, 0x20, 0x30)
            continue

    return tbl

# ── Document ───────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.0)

# ── Cover header ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("ADVANCER SMART TECHNOLOGY PTE LTD")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HyESys Department  ·  Hardware Specification Comparison")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x60, 0x60, 0x80)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run("HyESys H100 (v1.0)  vs  H125 (v2.0)")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ESS System  &  PCS Converter — Side-by-Side Specification Comparison")
r.font.size = Pt(12); r.font.color.rgb = DKGREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Prepared by Jason Ong  ·  June 2026  ·  For Internal Reference")
r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x90)

doc.add_paragraph()

# ── Source documents box ───────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("Source Documents")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY

sources = [
    ("H100 ESS System Manual",  "FZKJ-100KW206KWH-B1-F280  储能系统操作与运维指导手册",   "Beijing Fangzhi Technology",        "Jan 2024"),
    ("H125 ESS System Spec",    "FZKJ-125kW/236kWh-B1-F280  储能系统一体柜产品规格书",     "Zhejiang Fangzhi Juneng Technology", "Jan 2025"),
    ("H100 PCS Spec",           "INPPCS-100/0.4-W-14-C1-OS  产品规格书",                   "Beijing INP Electric",               "—"),
    ("H125 PCS Spec",           "SP125HCPS Energy Storage Converter Product Specification", "Sino Soar Hybrid (Shenzhen)",        "Jul 2025 v1.22"),
]
for label, title, mfr, date in sources:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f"{label}: ")
    r.bold = True; r.font.size = Pt(8.5)
    r = p.add_run(f"{title}  [{mfr}, {date}]")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x40, 0x40, 0x60)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — ESS SYSTEM COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 1 — ESS SYSTEM COMPARISON", 1)
add_note(doc, "H100 specs from the grid-connected cabinet (并网柜) section of the Operations Manual. "
              "H125 specs from the Product Specification Sheet.")

rows = [
    {'type': 'header', 'cols': ['Parameter', 'H100 System  (v1.0)', 'H125 System  (v2.0)']},
    {'type': 'section', 'text': '1.1  System Identity'},
    {'type': 'row', 'cols': ['Model Number', 'FZKJ-100KW206KWH-B1-F280', 'FZKJ-125kW/236kWh-B1-F280']},
    {'type': 'row', 'cols': ['Cabinet Type', 'Integrated ESS (grid-connected + off-grid cabinets)', 'Single integrated ESS cabinet'], 'alt': True},
    {'type': 'row', 'cols': ['Manufacturer', 'Beijing Fangzhi Technology Co., Ltd', 'Zhejiang Fangzhi Juneng Technology Co., Ltd']},
    {'type': 'row', 'cols': ['Document Date', 'January 2024', 'January 2025 (V1)'], 'alt': True},
    {'type': 'row', 'cols': ['Application', 'Grid-connected commercial/industrial ESS\n(Singapore project, 200 kWh)', 'User-side small commercial/industrial ESS\n& distribution grid management']},

    {'type': 'section', 'text': '1.2  Battery / DC Side'},
    {'type': 'row', 'cols': ['Cell Chemistry', 'Lithium Iron Phosphate (LFP)', 'Lithium Iron Phosphate (LFP)'], 'alt': True},
    {'type': 'row', 'cols': ['Cell Specification', 'Not specified in this document', '3.2 V / 280 Ah, prismatic aluminium']},
    {'type': 'row', 'cols': ['PACK Specification', 'Not specified (7 PACKs implied by H50 context)', '1P33S — 105.6 V / 280 Ah per PACK'], 'alt': True},
    {'type': 'row', 'cols': ['Battery String Config.', '1 string', '1P264S  (= 1P33S × 8 PACKs)']},
    {'type': 'row', 'cols': ['Number of PACKs', '7 (H50 system) / 22 (H125 system)', '8 PACKs per cabinet'], 'alt': True},
    {'type': 'row', 'cols': ['BMS Architecture', 'BMU → BCMU (2-level)', '2-level: 8 slave BMUs + 1 master BCM']},
    {'type': 'row', 'cols': ['Cell Balancing', 'Not specified', 'Passive balancing'], 'alt': True},
    {'type': 'row', 'cols': ['DC Nominal Voltage', '~745 V (observed in manual screenshots)', '844.8 V DC']},
    {'type': 'row', 'cols': ['DC Voltage Range', '600 – 900 V DC  (full load: 630 – 900 V)', '739 – 950 V  (cell: 2.8 – 3.6 V)'], 'alt': True,
     'highlight': [None, None, 'yellow']},
    {'type': 'row', 'cols': ['Max DC Current', '192 A', 'Not directly stated (PCS rated 200 A)']},
    {'type': 'row', 'cols': ['Rated Energy (System)', '206 kWh', '236.5 kWh  (@25°C, 0.5C, BOL)'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Cycle Life', '6,000 cycles  @25°C, 90% DOD, 80% EOL, 0.5C', 'Not specified in spec sheet']},

    {'type': 'section', 'text': '1.3  AC Side (Grid-Connected)'},
    {'type': 'row', 'cols': ['PCS Rated Power', '100 kW', '125 kW'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Rated AC Voltage', '400 Vac', '320 – 440 Vac  (rated 400 V)']},
    {'type': 'row', 'cols': ['AC Voltage Range', '372 – 428 Vac', '320 – 440 Vac  (broader range)'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Rated Frequency', '50 Hz', '50 Hz  (±5 Hz tolerance)']},
    {'type': 'row', 'cols': ['Frequency Range', '±0.5 Hz', '45 – 55 Hz'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Power Factor', '±0.99', 'Not stated (PCS supports -1 to +1)']},
    {'type': 'row', 'cols': ['Voltage THD', '≤5% (at rated power)', 'Not stated'], 'alt': True},
    {'type': 'row', 'cols': ['Current THD', '≤3% (at rated power)', '≤3% (≥30% load, from PCS spec)']},
    {'type': 'row', 'cols': ['3-Phase Voltage Imbalance', '≤2% rated value\n≤4% transient value', 'Not stated'], 'alt': True},
    {'type': 'row', 'cols': ['Max System Efficiency', '>98%', '≈89%  (system round-trip)\nPCS alone: 98.5% max'],
     'highlight': [None, 'green', 'yellow']},
    {'type': 'row', 'cols': ['Wiring', '3P3W + PE', '3P + N  (AC 320-440V)'], 'alt': True},

    {'type': 'section', 'text': '1.4  Physical & Environmental'},
    {'type': 'row', 'cols': ['Cabinet Dimensions\n(W × H × D)', '1,450 × 2,210 × 1,420 mm', '1,400 × 1,400 × 2,370 mm'],
     'highlight': [None, None, 'yellow']},
    {'type': 'row', 'cols': ['Cabinet Weight', '≈2,800 kg  (≈2.8 T)', '≈2,700 kg'], 'alt': True},
    {'type': 'row', 'cols': ['Cabinet Colour', 'RAL7035', 'Not specified']},
    {'type': 'row', 'cols': ['Protection Rating', 'IP54  (grid cabinet)\nIP45  (off-grid cabinet)', 'Meets outdoor use requirements'], 'alt': True,
     'highlight': [None, 'green', None]},
    {'type': 'row', 'cols': ['Cooling', 'Forced air cooling  +  Air conditioning', 'Air conditioning (battery)\n+ Forced air cooling (electrical)']},
    {'type': 'row', 'cols': ['Operating Temperature', '-20°C to +55°C', '-20°C to +50°C'], 'alt': True},
    {'type': 'row', 'cols': ['Storage Temperature', '-20°C to +70°C', '10°C to +35°C  (optimal storage)'],
     'highlight': [None, None, 'yellow']},
    {'type': 'row', 'cols': ['Operating Humidity', '0 – 95%', '0 – 95%  (optimal: dry)'], 'alt': True},
    {'type': 'row', 'cols': ['Altitude', '<2,000 m', '<2,000 m']},

    {'type': 'section', 'text': '1.5  Subsystems'},
    {'type': 'row', 'cols': ['Fire Suppression', 'Canister fire suppression system\n(electromagnetic valve, detection pipe)', 'Aerosol + water spray\nPre-installed external sprinkler piping'], 'alt': True},
    {'type': 'row', 'cols': ['Thermal Management', 'Air conditioning + forced ventilation', 'Air conditioning (battery compartment)\n+ forced air cooling (electrical)']},
    {'type': 'row', 'cols': ['Lighting', 'Built-in cabinet lighting', 'Not specified'], 'alt': True},
    {'type': 'row', 'cols': ['Sound/Thermal Insulation', 'Yes — insulation layer present', 'Not specified']},

    {'type': 'section', 'text': '1.6  Management & Communications'},
    {'type': 'row', 'cols': ['Local EMS (LEMS)', 'Yes — touchscreen LEMS\n(charge/discharge strategy, event log)', 'Yes — local management system + cloud platform'], 'alt': True},
    {'type': 'row', 'cols': ['Cloud Platform', 'Yes — via 4G / WTU module', 'Yes — 4G + cloud (FZKJ platform)']},
    {'type': 'row', 'cols': ['EMS Functions', 'Peak shaving, load following, anti-backflow\ncharge/discharge scheduling', 'Peak shaving, dynamic capacity expansion\nload following, anti-backflow, light-storage interaction\nmulti-cabinet parallel, OTA, SOH management'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Communication Protocols', 'RS485, CAN, Daisy-chain, Ethernet, UART, 4G', 'RS485, CAN, Ethernet, 4G']},
    {'type': 'row', 'cols': ['Metering', 'Built-in energy meter (DL/T 645 compliant)', 'AC energy meter + current transformer (2 sets)'], 'alt': True},
    {'type': 'row', 'cols': ['Other Sensors', 'Composite detector (CO/VOC/smoke)\nDoor switch, water sensor, temperature', 'Water sensor, door switch\nComposite sensor, temperature/humidity']},
    {'type': 'row', 'cols': ['Remote Monitoring', 'Yes', 'Yes'], 'alt': True},
]

make_table(doc, rows)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — PCS COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 2 — PCS (POWER CONVERSION SYSTEM) COMPARISON", 1)
add_note(doc, "H100 PCS: INPPCS-100/0.4-W-14-C1-OS by Beijing INP Electric Co., Ltd.  "
              "H125 PCS: SP125HCPS by Sino Soar Hybrid (Shenzhen) Technology Co., Ltd.  "
              "Both PCS units are rack-mount (insert frame) modules installed inside the ESS cabinet.")

rows2 = [
    {'type': 'header', 'cols': ['Parameter', 'H100 PCS  (INP / INPPCS-100)', 'H125 PCS  (Sino Soar / SP125HCPS)']},
    {'type': 'section', 'text': '2.1  Identity & Topology'},
    {'type': 'row', 'cols': ['Model', 'INPPCS-100/0.4-W-14-C1-OS', 'SP125HCPS']},
    {'type': 'row', 'cols': ['Manufacturer', 'Beijing INP Electric Co., Ltd\n(北京英博电气股份有限公司)', 'Sino Soar Hybrid (Shenzhen) Technology Co., Ltd\n(中腾微网)'], 'alt': True},
    {'type': 'row', 'cols': ['Rated AC Power', '100 kW', '125 kW'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Topology', 'Three-level topology (3-phase 4-arm bridge)', 'T-type three-level inverter'], 'alt': True},
    {'type': 'row', 'cols': ['Grid Isolation', 'Non-isolated (transformerless)', 'Not specified (transformerless implied)']},
    {'type': 'row', 'cols': ['Installation Type', 'Insert frame (rack-mount)', 'Insert frame (rack-mount)'], 'alt': True},
    {'type': 'row', 'cols': ['HMI', 'No built-in display\nRequires external 7" HMI screen', 'No built-in display\n(EMS screen optional)']},

    {'type': 'section', 'text': '2.2  DC Side (Battery Interface)'},
    {'type': 'row', 'cols': ['DC Channels', '1', '1']},
    {'type': 'row', 'cols': ['DC Voltage Range\n(operational)', 'DC 600 – 950 V\n(full rated output: 630 – 900 V)', 'DC 680 – 950 V\n(rated range: 680 – 900 V)'], 'alt': True,
     'highlight': [None, None, 'yellow']},
    {'type': 'row', 'cols': ['Minimum Battery Voltage', '600 V DC', '680 V DC'],
     'highlight': [None, 'green', 'yellow']},
    {'type': 'row', 'cols': ['Maximum Battery Voltage', '950 V DC', '950 V DC'], 'alt': True},
    {'type': 'row', 'cols': ['Rated Battery Voltage', '630 – 900 V  (for full power)', '680 – 900 V  (for full power)']},
    {'type': 'row', 'cols': ['Max DC Current', '192 A', '200 A'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Rated DC Power', '100 kW', '125 kW'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Voltage Regulation', '≤±2%', 'Not stated'], 'alt': True},
    {'type': 'row', 'cols': ['Current Regulation', '≤±5%', 'Not stated']},
    {'type': 'row', 'cols': ['BMS Communication', 'CAN', 'CAN or RS485  (selectable)'], 'alt': True,
     'highlight': [None, None, 'green']},

    {'type': 'section', 'text': '2.3  AC Side — Grid-Connected Mode'},
    {'type': 'row', 'cols': ['Rated Output Power', '100 kW', '125 kVA  (≈125 kW at PF=1)'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Rated Output Current', '145 A', '181 A'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Rated AC Voltage', 'AC 400 V', 'AC 400 V / 230 V  (3P/1P)']},
    {'type': 'row', 'cols': ['Grid Voltage Range', '380 V  (−20% to +15%)\n= 304 – 437 V', 'AC 400 V  (−20% to +15%)\n= 320 – 460 V'], 'alt': True},
    {'type': 'row', 'cols': ['Frequency Range', '50 Hz / 60 Hz  ±2.5 Hz', '50 Hz: 47 – 52 Hz\n60 Hz: 57 – 62 Hz']},
    {'type': 'row', 'cols': ['Current THD', '≤3%  (full load)', '<3%  (>30% load)'], 'alt': True},
    {'type': 'row', 'cols': ['Power Factor Control', '−0.99 to +0.99', '−1 to +1  (unity achievable)']},
    {'type': 'row', 'cols': ['DC Current Component', '≤0.5%', 'Not stated'], 'alt': True},
    {'type': 'row', 'cols': ['Overload Capacity', '1.1× continuous\n1.2× for 1 minute', 'Ie×1.1 to Ie×1.25 for 100 s\n>Ie×1.25 for 300 ms\n(150% instantaneous)'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Charge/Discharge Switching', '<100 ms', '<10 ms  (seamless grid/off-grid)'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Wiring Mode', '3-phase 4-wire', '3-phase 3-wire or 4-wire']},
    {'type': 'row', 'cols': ['AC Parallel Operation', 'Multiple units in parallel (AC side)', 'Up to 15 units parallel\n(grid-connected or off-grid)'], 'alt': True,
     'highlight': [None, None, 'green']},

    {'type': 'section', 'text': '2.4  AC Side — Off-Grid Mode'},
    {'type': 'row', 'cols': ['Off-Grid Rated Output', '100 kW / 100 kVA', '125 kVA\n(max 137.5 kVA = 110%)']},
    {'type': 'row', 'cols': ['Max Output Current\n(off-grid)', '160 A (implied by overload)', '200 A'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Off-Grid Output Voltage', 'AC 400 V  (±3%)', 'AC 400 V / 230 V  (±3%)']},
    {'type': 'row', 'cols': ['Output Voltage THD', '≤3%  (linear load)', '<3%  (resistive load)'], 'alt': True},
    {'type': 'row', 'cols': ['Unbalanced Load', '100%', '100%']},
    {'type': 'row', 'cols': ['Off-Grid Frequency', '50 Hz / 60 Hz', '50 Hz / 60 Hz'], 'alt': True},
    {'type': 'row', 'cols': ['Off-Grid Parallel', 'Not supported', 'Supported (up to 15 units)'],
     'highlight': [None, 'red', 'green']},
    {'type': 'row', 'cols': ['Black Start', 'Not specified', 'Supported'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Island Protection', 'Anti-islanding protection', 'Island protection + black start']},

    {'type': 'section', 'text': '2.5  Efficiency & Power Quality'},
    {'type': 'row', 'cols': ['Max Conversion Efficiency', '≥98%', '98.5%'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Standby Power Consumption', 'Not specified', '<15 W'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['No-Load Power Loss', 'Not specified', '<200 W'], 'alt': True,
     'highlight': [None, None, 'green']},

    {'type': 'section', 'text': '2.6  Physical Specifications'},
    {'type': 'row', 'cols': ['Dimensions  (W × L × H)', '480 × 720 × 260 mm', '440 × 720 × 200 mm'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Weight', '70 kg', '≤50 kg'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Cooling Method', 'Forced air cooling\n(rear intake → front outlet)', 'Intelligent variable-speed fan\n(front intake → rear outlet)'],
     'highlight': [None, None, 'yellow']},
    {'type': 'row', 'cols': ['Rated Airflow', '467 CFM  (13.2 m³/min)', '500 CFM  (14.1 m³/min)'], 'alt': True},
    {'type': 'row', 'cols': ['Protection Rating', 'IP20', 'IP20  (core control: IP5X)'],
     'highlight': [None, None, 'green']},

    {'type': 'section', 'text': '2.7  Environmental Conditions'},
    {'type': 'row', 'cols': ['Operating Temperature', '-20°C to +50°C\n(derate above +45°C)', '-30°C to +60°C\n(derate above +45°C)'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Relative Humidity', '≤95%  (no condensation)', '5% – 95%']},
    {'type': 'row', 'cols': ['Altitude', '>2,000 m: derate', '<4,000 m  (derate above 2,000 m)'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Noise Level', '≤75 dB', 'Not specified']},

    {'type': 'section', 'text': '2.8  Communications & Control'},
    {'type': 'row', 'cols': ['EMS Communication', 'Ethernet or RS485', 'RS485  (EMS port)'], 'alt': True},
    {'type': 'row', 'cols': ['BMS Communication', 'CAN', 'CAN or RS485  (selectable)'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Display Communication', 'Ethernet', 'Ethernet / RS485'], 'alt': True},
    {'type': 'row', 'cols': ['Parallel Communication', 'AC side parallel', 'Dedicated parallel bus\n(up to 15 units, CAN-based)'],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['I/O Dry Contacts', 'Emergency stop input via IO', 'DI: 2 channels\nDO: 2 channels'], 'alt': True,
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Emergency Stop', 'IO input from external switch', 'Dry contact input  (X1 / X1_com)']},

    {'type': 'section', 'text': '2.9  Certifications & Standards'},
    {'type': 'row', 'cols': ['Certifications', 'GB/T 34120, GB/T 34133, GB/T 36547\nNB/T 33014/33015/33016\nIEC-62477, IEC-61000', 'CE, IEC62477, IEC6100, EN50549'], 'alt': True},
    {'type': 'row', 'cols': ['Grid Standards', 'Chinese national grid standards\n(GB, NB, Q/GDW series)', 'International (IEC) + European (EN50549)'],
     'highlight': [None, None, 'green']},

    {'type': 'section', 'text': '2.10  Advanced Features'},
    {'type': 'row', 'cols': ['Grid-Forming (off-grid)', 'Limited  (no off-grid parallel)', 'Full  (black start, island, up to 15-unit parallel)'], 'alt': True,
     'highlight': [None, 'red', 'green']},
    {'type': 'row', 'cols': ['Oil/Generator Hybrid', 'Not specified', 'Supported  (oil-engine hybrid mode)',],
     'highlight': [None, None, 'green']},
    {'type': 'row', 'cols': ['Remote OTA Update', 'Via cloud platform', 'Yes  (EMS and inverter OTA)'], 'alt': True},
    {'type': 'row', 'cols': ['Active/Reactive Control', 'Active + reactive (P/Q) regulation', 'Active + reactive independent 3-phase control']},
    {'type': 'row', 'cols': ['Frequency Response', 'Frequency abnormality tolerance', 'High/low voltage ride-through\nFrequency ride-through'], 'alt': True,
     'highlight': [None, None, 'green']},
]

make_table(doc, rows2)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — KEY DIFFERENCES SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART 3 — KEY DIFFERENCES AT A GLANCE", 1)

p = doc.add_paragraph()
r = p.add_run("H125 (v2.0) upgrades vs H100 (v1.0)")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY

improvements = [
    ("Power",          "+25% rated AC power (125 kW vs 100 kW); +25% max DC current (200 A vs 192 A)"),
    ("Energy",         "+15% usable energy (236.5 kWh vs 206 kWh) with same-size cabinet footprint"),
    ("Battery",        "Cell-level spec defined (3.2V/280Ah LFP); H100 relied on legacy PACK spec"),
    ("PCS Efficiency", "98.5% max vs ≥98% — marginal improvement; standby loss defined (<15 W)"),
    ("PCS Weight",     "≤50 kg vs 70 kg — 29% lighter PCS module, same rack form factor"),
    ("Off-Grid",       "H125 PCS supports up to 15-unit parallel off-grid + black start; H100 PCS: not supported"),
    ("Grid Voltage",   "H125 PCS accepts wider DC window (680 V floor vs 630 V) — less discharge depth"),
    ("Temperature",    "H125 PCS rated to -30°C (vs -20°C); altitude up to 4,000 m (vs 3,000 m)"),
    ("Switching",      "H125 PCS: <10 ms seamless switching; H100 PCS: <100 ms — 10× improvement"),
    ("Standards",      "H125 PCS carries CE + EN50549 — suitable for Singapore/international deployment"),
    ("Parallel",       "H125 PCS: hardware parallel bus (up to 15 units); H100 PCS: AC-side parallel only"),
    ("Certifications", "H125: IEC + CE certified; H100: Chinese national standards only"),
]

for title, detail in improvements:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{title}: ")
    r.bold = True; r.font.size = Pt(9)
    r = p.add_run(detail)
    r.font.size = Pt(9)

p = doc.add_paragraph()
r = p.add_run("H100 (v1.0) advantages retained / areas where H125 is not better")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80, 0x30, 0x20)
p.paragraph_format.space_before = Pt(8)

notes = [
    ("Cycle Life",   "H100 explicitly rated 6,000 cycles; H125 spec sheet does not state cycle life"),
    ("DC Window",    "H100 PCS operates from 600 V DC (lower floor) — allows deeper battery discharge"),
    ("IP Rating",    "H100 cabinet is IP54 (outdoor-grade); H125 spec only states 'meets outdoor use'"),
    ("Noise",        "H100 PCS noise stated ≤75 dB; H125 PCS noise not specified"),
    ("System Eff.",  "H100 system efficiency stated >98% (likely PCS-only); H125 system ≈89% round-trip"),
]

for title, detail in notes:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{title}: ")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x30, 0x20)
    r = p.add_run(detail)
    r.font.size = Pt(9)

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Disclaimer: Specifications compiled directly from manufacturer documents as listed above. "
              "Values marked (v1.0) and (v2.0) refer to HyESys internal versioning, not manufacturer version numbers. "
              "Verify with manufacturer before procurement decisions.")
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
