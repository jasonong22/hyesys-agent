"""
Builds FEE2026_Mechanical_Equations_Reference.docx
108 equations across ME 101-106 / ME 201-206, each with legend + worked example.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DST = r'C:\Users\JasonOng\Desktop\local docs\personal\PE\FEE2026_Mechanical_Equations_Reference.docx'

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(0.95)
    sec.right_margin  = Inches(0.95)

C = {
    'title':      RGBColor(0x1A, 0x1A, 0x5E),
    'topic':      RGBColor(0x00, 0x3D, 0x7A),
    'subtopic':   RGBColor(0x00, 0x5C, 0xA8),
    'eqname':     RGBColor(0x12, 0x40, 0x12),
    'formula':    RGBColor(0x0D, 0x0D, 0x7A),
    'legend_hdr': RGBColor(0xFF, 0xFF, 0xFF),
    'body':       RGBColor(0x1A, 0x1A, 0x1A),
    'note':       RGBColor(0x55, 0x55, 0x55),
    'ex_hdr':     RGBColor(0x7B, 0x4B, 0x00),
    'ex_ans':     RGBColor(0x15, 0x57, 0x24),
    'defn':       RGBColor(0x1A, 0x1A, 0x60),
}
SHADING = {
    'topic':      'D0E4F5',
    'subtopic':   'EBF5FF',
    'eq':         'F5F5F5',
    'legend_hdr': '003D7A',
    'legend_row': 'F0F6FF',
    'legend_alt': 'FAFAFA',
    'ex_hdr':     'FEF3CD',
    'ex_body':    'FFFEF5',
    'ex_ans':     'D4EDDA',
    'defn':       'EAF0FF',
}

def _shd(hex_fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    return shd

def ap(text='', bold=False, italic=False, mono=False,
       color=None, shading=None, align=None, sz=10, before=40, after=40):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 6)
    p.paragraph_format.space_after  = Pt(after  / 6)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if shading:
        p._p.get_or_add_pPr().append(_shd(shading))
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if mono: r.font.name = 'Courier New'
        r.font.size = Pt(sz)
        if color:
            r.font.color.rgb = color if isinstance(color, RGBColor) \
                               else RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
    return p

def topic_hdr(code, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p._p.get_or_add_pPr().append(_shd(SHADING['topic']))
    r = p.add_run(f'  {code}  {title}')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = C['topic']

def subtopic_hdr(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p._p.get_or_add_pPr().append(_shd(SHADING['subtopic']))
    r = p.add_run(f'  {text}')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C['subtopic']

def ex_block(problem, steps, answer):
    ap('  Worked Example', bold=True, shading=SHADING['ex_hdr'],
       color=C['ex_hdr'], sz=8.5, before=3, after=1)
    ap(f'  Q: {problem}', italic=True, sz=9,
       shading=SHADING['ex_body'], color=C['body'], before=0, after=2)
    for step in steps:
        ap(f'     {step}', mono=True, sz=8.5,
           shading=SHADING['ex_body'], color=C['formula'], before=0, after=0)
    ap(f'  Ans: {answer}', bold=True, sz=9,
       shading=SHADING['ex_ans'], color=C['ex_ans'], before=2, after=8)

def eq_block(number, name, formula, legend_rows, note=None, definition=None, example=None):
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(5)
    p_name.paragraph_format.space_after  = Pt(1)
    rn = p_name.add_run(f'  [{number}]  {name}')
    rn.bold = True; rn.font.size = Pt(10); rn.font.color.rgb = C['eqname']

    p_form = doc.add_paragraph()
    p_form.paragraph_format.space_before = Pt(0)
    p_form.paragraph_format.space_after  = Pt(1)
    p_form.paragraph_format.left_indent  = Inches(0.35)
    p_form._p.get_or_add_pPr().append(_shd(SHADING['eq']))
    rf = p_form.add_run(f'  {formula}')
    rf.bold = True; rf.font.name = 'Courier New'
    rf.font.size = Pt(10.5); rf.font.color.rgb = C['formula']

    if note:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(0)
        p_note.paragraph_format.space_after  = Pt(1)
        p_note.paragraph_format.left_indent  = Inches(0.35)
        rno = p_note.add_run(f'  > {note}')
        rno.italic = True; rno.font.size = Pt(8.5); rno.font.color.rgb = C['note']

    if definition:
        p_def = doc.add_paragraph()
        p_def.paragraph_format.space_before = Pt(1)
        p_def.paragraph_format.space_after  = Pt(2)
        p_def.paragraph_format.left_indent  = Inches(0.35)
        p_def._p.get_or_add_pPr().append(_shd(SHADING['defn']))
        r_lbl = p_def.add_run('  Definition: ')
        r_lbl.bold = True; r_lbl.font.size = Pt(8.5)
        r_lbl.font.color.rgb = C['defn']
        r_body = p_def.add_run(definition)
        r_body.italic = True; r_body.font.size = Pt(8.5)
        r_body.font.color.rgb = C['defn']

    tbl = doc.add_table(rows=1 + len(legend_rows), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(0.85), Inches(4.10), Inches(1.20)]
    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    hdr_texts = ['Symbol', 'Description', 'SI Units']
    for i, cell in enumerate(tbl.rows[0].cells):
        cell._tc.get_or_add_tcPr().append(_shd(SHADING['legend_hdr']))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr_texts[i])
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = C['legend_hdr']

    for ri, (sym, desc, unit) in enumerate(legend_rows):
        row = tbl.rows[ri + 1]
        fill = SHADING['legend_row'] if ri % 2 == 0 else SHADING['legend_alt']
        for ci, (cell, val) in enumerate(zip(row.cells, [sym, desc, unit])):
            cell._tc.get_or_add_tcPr().append(_shd(fill))
            p = cell.paragraphs[0]
            if ci == 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True; run.font.name = 'Courier New'
                run.font.color.rgb = C['formula']
            else:
                run.font.color.rgb = C['body']

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if example:
        ex_block(example['problem'], example['steps'], example['answer'])


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT HEADER
# ══════════════════════════════════════════════════════════════════════════════
ap('PROFESSIONAL ENGINEERS EXAMINATION FEE 2026', bold=True, sz=15,
   color=C['title'], align='center', before=0, after=10)
ap('Mechanical Discipline — Complete Equation Reference with Worked Examples',
   bold=True, sz=12, color=C['subtopic'], align='center', before=0, after=6)
ap('Topics: ME 101-106 / ME 201-206  |  All equations web-verified  |  SI units throughout',
   sz=9, color=C['note'], align='center', before=0, after=20)
ap('-' * 105, mono=True, sz=7, color=C['note'], before=0, after=16)

# ══════════════════════════════════════════════════════════════════════════════
# ME 101/201  CONTROL AND INSTRUMENTATION
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 101 / 201', 'Control and Instrumentation')
subtopic_hdr('Transfer Functions & System Modelling')

eq_block(1, 'Standard Second-Order Transfer Function',
    'G(s) = wn^2 / (s^2 + 2*zeta*wn*s + wn^2)',
    [('G(s)', 'Transfer function (output/input in Laplace domain)', 'dimensionless'),
     ('s',    'Complex frequency variable (Laplace operator)', 'rad/s'),
     ('wn',   'Undamped natural frequency', 'rad/s'),
     ('zeta', 'Damping ratio', 'dimensionless')],
    definition="Describes how a second-order LTI system (e.g. spring-mass-damper, motor with inertia) maps an input to an output in the Laplace domain. Read off wn^2 from the s^0 coefficient and 2*zeta*wn from the s^1 coefficient of the denominator. Assumes no zeros in the numerator; the presence of zeros alters transient response formulae (eqs. 5-8).",
    example={
        'problem': 'A system has G(s) = 100/(s^2 + 6s + 100). Find wn and zeta.',
        'steps': [
            'wn^2 = 100  =>  wn = 10 rad/s',
            '2*zeta*wn = 6  =>  zeta = 6/(2*10) = 0.30',
        ],
        'answer': 'wn = 10 rad/s, zeta = 0.30 (underdamped)'})

eq_block(2, 'Closed-Loop Characteristic Equation',
    '1 + G(s)*H(s) = 0',
    [('G(s)', 'Forward-path (plant) transfer function', 'dimensionless'),
     ('H(s)', 'Feedback transfer function', 'dimensionless')],
    note='Poles of this equation determine stability. Roots must have negative real parts.',
    definition="Setting 1 + G(s)H(s) = 0 gives the closed-loop poles that govern stability and transient response. The system is BIBO stable if and only if every root lies in the left half s-plane (negative real part). Apply Routh-Hurwitz to the resulting polynomial for a quick stability check without solving for all roots.",
    example={
        'problem': 'G(s) = 10/(s+5), H(s) = 1 (unity feedback). Find closed-loop pole.',
        'steps': [
            '1 + 10/(s+5) = 0  =>  s+5+10 = 0  =>  s = -15',
        ],
        'answer': 'Closed-loop pole at s = -15 (stable, left half-plane)'})

eq_block(3, 'Damping Ratio (from characteristic polynomial as^2 + bs + c = 0)',
    'zeta = b / (2*sqrt(a*c))',
    [('zeta', 'Damping ratio', 'dimensionless'),
     ('a',    'Coefficient of s^2 term', 'varies'),
     ('b',    'Coefficient of s term', 'varies'),
     ('c',    'Constant (s^0) term', 'varies')],
    definition="Extracts the damping ratio directly from the three polynomial coefficients of a second-order characteristic equation. Also gives wn = sqrt(c/a). Requires a true second-order polynomial with no repeated poles or numerator zeros; adding zeros changes the transient shape but not the pole locations.",
    example={
        'problem': 'Characteristic polynomial: 2s^2 + 8s + 18 = 0. Find zeta.',
        'steps': [
            'a=2, b=8, c=18',
            'zeta = 8 / (2*sqrt(2*18)) = 8 / (2*sqrt(36)) = 8/12 = 0.667',
        ],
        'answer': 'zeta = 0.667 (underdamped)'})

subtopic_hdr('Transient Response (Underdamped, 0 < zeta < 1)')

eq_block(4, 'Damped Natural Frequency',
    'wd = wn * sqrt(1 - zeta^2)',
    [('wd',   'Damped natural frequency', 'rad/s'),
     ('wn',   'Undamped natural frequency', 'rad/s'),
     ('zeta', 'Damping ratio', 'dimensionless')],
    definition="The actual oscillation frequency of an underdamped system after a disturbance; always less than wn due to energy dissipation by damping. Only defined for underdamped systems (0 < zeta < 1); for zeta >= 1 the system does not oscillate. Used as the base frequency in all transient response formulae (eqs. 5-7).",
    example={
        'problem': 'wn = 10 rad/s, zeta = 0.30. Find wd.',
        'steps': [
            'wd = 10 * sqrt(1 - 0.09) = 10 * sqrt(0.91) = 10 * 0.954 = 9.54 rad/s',
        ],
        'answer': 'wd = 9.54 rad/s'})

eq_block(5, 'Rise Time (10-90%)',
    'tr = (pi - phi) / wd    where phi = arctan(sqrt(1-zeta^2) / zeta)',
    [('tr',   'Rise time', 's'),
     ('phi',  'Phase angle = arctan(sqrt(1-zeta^2)/zeta)', 'rad'),
     ('wd',   'Damped natural frequency', 'rad/s'),
     ('zeta', 'Damping ratio', 'dimensionless')],
    definition="Time for the step response to travel from 10% to 90% of its final value; measures response speed. Valid only for standard second-order underdamped systems (0 < zeta < 1) with no zeros. Faster rise (lower tr) requires higher wn or lower zeta, but lower zeta also increases overshoot — a fundamental speed-accuracy trade-off.",
    example={
        'problem': 'wn = 10 rad/s, zeta = 0.30. Find tr.',
        'steps': [
            'wd = 9.54 rad/s',
            'phi = arctan(0.954/0.30) = arctan(3.18) = 1.265 rad',
            'tr = (3.1416 - 1.265) / 9.54 = 1.877 / 9.54 = 0.197 s',
        ],
        'answer': 'tr = 0.197 s'})

eq_block(6, 'Peak Time',
    'tp = pi / wd',
    [('tp', 'Time to first peak overshoot', 's'),
     ('wd', 'Damped natural frequency', 'rad/s')],
    definition="Time at which the step response reaches its first (maximum) overshoot peak. Derived by differentiating the underdamped step response and finding the first zero crossing. Valid only for underdamped systems (0 < zeta < 1); critically and overdamped systems produce no peak.",
    example={
        'problem': 'wd = 9.54 rad/s. Find tp.',
        'steps': [
            'tp = pi / 9.54 = 3.1416 / 9.54 = 0.329 s',
        ],
        'answer': 'tp = 0.329 s'})

eq_block(7, 'Percent Overshoot',
    '%OS = exp(-pi*zeta / sqrt(1 - zeta^2)) * 100',
    [('%OS',  'Percent overshoot', '%'),
     ('zeta', 'Damping ratio', 'dimensionless')],
    note='Valid for standard 2nd-order systems without zeros.',
    definition="How much the step response exceeds its final value expressed as a percentage; depends only on zeta for a standard second-order system. Valid only for underdamped systems (zeta < 1); overdamped and critically damped systems produce zero overshoot. A system with zeta = 0.7 gives approximately 5% overshoot — a common design target.",
    example={
        'problem': 'zeta = 0.30. Find %OS.',
        'steps': [
            '%OS = exp(-pi*0.30 / sqrt(1-0.09)) * 100',
            '    = exp(-0.9425 / 0.954) * 100 = exp(-0.988) * 100 = 37.2%',
        ],
        'answer': '%OS = 37.2%'})

eq_block(8, 'Settling Time (2% criterion)',
    'ts = 4 / (zeta * wn)',
    [('ts',   'Settling time (response within +/-2% of final value)', 's'),
     ('zeta', 'Damping ratio', 'dimensionless'),
     ('wn',   'Undamped natural frequency', 'rad/s')],
    note='For 5% criterion use ts = 3/(zeta*wn).',
    definition="Approximate time for the step response to enter and permanently remain within +/-2% of its final value. The formula ts = 4/(zeta*wn) is an engineering approximation based on the exponential decay envelope of the damped oscillation; the exact value involves logarithmic terms. Valid for underdamped systems; becomes less accurate for very low or very high zeta.",
    example={
        'problem': 'wn = 10 rad/s, zeta = 0.30. Find ts (2% criterion).',
        'steps': [
            'ts = 4 / (0.30 * 10) = 4 / 3 = 1.33 s',
        ],
        'answer': 'ts = 1.33 s'})

subtopic_hdr('Steady-State Errors (Unity Feedback)')

eq_block(9, 'Position Error Constant (Type 0 system - step input)',
    'Kp = lim[s->0] G(s)     =>     e_ss = 1 / (1 + Kp)',
    [('Kp',   'Position error constant', 'dimensionless'),
     ('e_ss', 'Steady-state error', 'same units as input'),
     ('G(s)', 'Open-loop transfer function', 'dimensionless')],
    definition="Quantifies the steady-state error of a unity-feedback Type 0 system (no open-loop integrators) to a unit step input. A higher Kp reduces but never eliminates the error; only adding an integrator (making the system Type 1) gives zero steady-state step error. Apply only to stable closed-loop systems.",
    example={
        'problem': 'G(s) = 20/(s+4), unity feedback, unit step input. Find Kp and e_ss.',
        'steps': [
            'Kp = lim[s->0] 20/(s+4) = 20/4 = 5',
            'e_ss = 1/(1+5) = 1/6 = 0.167',
        ],
        'answer': 'Kp = 5, e_ss = 0.167 (16.7% of step magnitude)'})

eq_block(10, 'Velocity Error Constant (Type 1 system - ramp input)',
    'Kv = lim[s->0] s*G(s)     =>     e_ss = 1 / Kv',
    [('Kv',   'Velocity error constant', 's^-1'),
     ('e_ss', 'Steady-state error', 's'),
     ('G(s)', 'Open-loop transfer function', 'dimensionless')],
    definition="Quantifies the steady-state positional lag of a unity-feedback Type 1 system (exactly one open-loop integrator) tracking a constant-velocity ramp input. A Type 0 system cannot track a ramp (infinite error); a Type 2 system tracks a ramp with zero error. Higher Kv means tighter ramp tracking.",
    example={
        'problem': 'G(s) = 5/(s*(s+2)), unity feedback, ramp input. Find Kv and e_ss.',
        'steps': [
            'Kv = lim[s->0] s * 5/(s*(s+2)) = lim[s->0] 5/(s+2) = 5/2 = 2.5 s^-1',
            'e_ss = 1/Kv = 1/2.5 = 0.40',
        ],
        'answer': 'Kv = 2.5 s^-1, e_ss = 0.40'})

subtopic_hdr('PID Controller')

eq_block(11, 'PID Controller Transfer Function',
    'C(s) = Kp + Ki/s + Kd*s',
    [('C(s)', 'Controller transfer function', 'dimensionless'),
     ('Kp',   'Proportional gain', 'dimensionless'),
     ('Ki',   'Integral gain', 's^-1'),
     ('Kd',   'Derivative gain', 's'),
     ('s',    'Laplace operator', 'rad/s')],
    definition="Combines three control actions: P (proportional — responds to current error magnitude), I (integral — accumulates past error to eliminate steady-state offset, raises system type by 1), D (derivative — reacts to rate of change of error, adds damping but amplifies noise). The Ki/s term is an integrator; Kd*s is a differentiator. Gains are tuned for the required speed, stability margin, and disturbance rejection.",
    example={
        'problem': 'Kp=8, Ki=4, Kd=0.5. Write C(s) and find the combined transfer function.',
        'steps': [
            'C(s) = 8 + 4/s + 0.5s',
            'Combined: C(s) = (0.5s^2 + 8s + 4) / s',
        ],
        'answer': 'C(s) = (0.5s^2 + 8s + 4)/s'})

subtopic_hdr('Frequency Domain Stability Margins')

eq_block(12, 'Phase Margin',
    'PM = 180 deg + angle[G(j*wgc)]',
    [('PM',       'Phase margin (system stable if PM > 0 deg)', 'degrees'),
     ('wgc',      'Gain crossover frequency (where |G(jw)| = 1)', 'rad/s'),
     ('angle G',  'Phase angle of open-loop TF at wgc', 'degrees')],
    definition="Measures how much additional phase lag the system can tolerate before becoming unstable; read from the Bode phase plot at the gain crossover frequency wgc. A positive PM indicates stability; PM = 0 deg means the system is on the verge of instability. Typical design target: PM = 30-60 deg. A small PM (< 20 deg) produces a highly oscillatory closed-loop response.",
    example={
        'problem': 'At gain crossover frequency, G(jwgc) has phase angle -145 deg. Find PM.',
        'steps': [
            'PM = 180 + (-145) = 35 deg',
        ],
        'answer': 'PM = 35 deg (stable, PM > 0 deg)'})

eq_block(13, 'Gain Margin',
    'GM = 1 / |G(j*wpc)|     or     GM_dB = -20*log10|G(j*wpc)|',
    [('GM',       'Gain margin (stable if GM > 1)', 'dimensionless (or dB)'),
     ('wpc',      'Phase crossover frequency (where angle G = -180 deg)', 'rad/s'),
     ('|G(jw)|',  'Magnitude of open-loop TF at wpc', 'dimensionless')],
    definition="How much the open-loop gain can increase beyond its current value before the closed-loop system becomes unstable; read from the Bode magnitude plot at the phase crossover frequency wpc (where phase = -180 deg). Stable when GM > 1 (positive dB). A system with both PM > 0 and GM > 1 is guaranteed stable for negative unity feedback.",
    example={
        'problem': '|G(jwpc)| = 0.25 at phase crossover. Find GM in linear and dB.',
        'steps': [
            'GM = 1/0.25 = 4',
            'GM_dB = 20*log10(4) = 12.0 dB',
        ],
        'answer': 'GM = 4 (12.0 dB) — system is stable'})


# ══════════════════════════════════════════════════════════════════════════════
# ME 102/202  DYNAMICS AND VIBRATIONS
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 102 / 202', 'Dynamics and Vibrations')
subtopic_hdr('Kinematics - Particle (Constant Acceleration)')

eq_block(14, "Newton's Second Law - Linear Motion",
    'F = m * a',
    [('F', 'Net force', 'N'),
     ('m', 'Mass', 'kg'),
     ('a', 'Linear acceleration', 'm/s^2')],
    definition="States that the net force on a body equals its mass times its acceleration. F is the vector sum of ALL external forces; internal forces (e.g. between connected parts) cancel. Applies only in an inertial (non-accelerating) reference frame. For a non-inertial frame, add a pseudo-force (-m*a_frame) to account for the frame's own acceleration.",
    example={
        'problem': 'A 5 kg block experiences a net force of 30 N. Find acceleration.',
        'steps': ['a = F/m = 30/5 = 6 m/s^2'],
        'answer': 'a = 6 m/s^2'})

eq_block(15, 'SUVAT Equations of Motion (constant acceleration)',
    'v = u + at  |  s = ut + (1/2)*a*t^2  |  v^2 = u^2 + 2as',
    [('v', 'Final velocity', 'm/s'),
     ('u', 'Initial velocity', 'm/s'),
     ('a', 'Constant acceleration', 'm/s^2'),
     ('t', 'Time', 's'),
     ('s', 'Displacement', 'm')],
    definition="Three kinematic equations relating displacement, velocity, acceleration, and time under constant (uniform) acceleration. Strictly invalid when acceleration varies with time or position. Select the equation that contains the three known quantities and the one unknown. For variable acceleration, integrate a(t) directly: v = integral(a dt), s = integral(v dt).",
    example={
        'problem': 'Car starts from rest, accelerates at 3 m/s^2 for 8 s. Find v and s.',
        'steps': [
            'v = 0 + 3*8 = 24 m/s',
            's = 0 + (1/2)*3*8^2 = (1/2)*3*64 = 96 m',
        ],
        'answer': 'v = 24 m/s, s = 96 m'})

eq_block(16, 'Centripetal Acceleration',
    'ac = v^2/r = w^2*r',
    [('ac', 'Centripetal (radial) acceleration, directed toward centre', 'm/s^2'),
     ('v',  'Tangential velocity', 'm/s'),
     ('r',  'Radius of circular path', 'm'),
     ('w',  'Angular velocity', 'rad/s')],
    definition="Inward (radial) acceleration required to maintain circular motion at constant speed; acts toward the centre of curvature, never tangentially. The centripetal force F = m*ac is provided by the net inward force (tension, friction, or normal force). No work is done by centripetal force since it is always perpendicular to velocity.",
    example={
        'problem': 'A ball on a 1.2 m string rotates at w = 4 rad/s. Find centripetal acceleration.',
        'steps': ['ac = w^2 * r = 4^2 * 1.2 = 16 * 1.2 = 19.2 m/s^2'],
        'answer': 'ac = 19.2 m/s^2 (directed toward centre)'})

subtopic_hdr('Kinetics - Energy and Momentum')

eq_block(17, 'Work-Energy Theorem',
    'W_net = DeltaKE = (1/2)*m*v2^2 - (1/2)*m*v1^2',
    [('W_net', 'Net work done on the body', 'J'),
     ('KE',    'Kinetic energy', 'J'),
     ('m',     'Mass', 'kg'),
     ('v1',    'Initial velocity', 'm/s'),
     ('v2',    'Final velocity', 'm/s')],
    definition="Net work done by all forces on a body equals its change in kinetic energy; valid regardless of path. Conservative forces (gravity, spring) are conveniently handled via potential energy (W_cons = -DeltaPE). For a varying force, compute work as the integral of F*ds along the path. Applies to rigid bodies undergoing translation; for rotation use the rotational kinetic energy (1/2)*I*w^2.",
    example={
        'problem': 'A 2 kg object accelerates from 3 m/s to 7 m/s. Find net work done.',
        'steps': [
            'W_net = (1/2)*2*7^2 - (1/2)*2*3^2',
            '      = (1/2)*2*49 - (1/2)*2*9 = 49 - 9 = 40 J',
        ],
        'answer': 'W_net = 40 J'})

eq_block(18, 'Conservation of Linear Momentum',
    'm1*v1 + m2*v2 = m1*v1\' + m2*v2\'',
    [('m1, m2',    'Masses of bodies 1 and 2', 'kg'),
     ("v1, v2",    'Velocities before collision', 'm/s'),
     ("v1', v2'",  'Velocities after collision', 'm/s')],
    note='Valid when no external impulse acts on the system.',
    definition="Total linear momentum is conserved when no external impulse (force x time) acts on the system during the collision interval. Internal contact forces cancel in Newton's 3rd law pairs. Gravity and friction are external forces but are negligible during brief high-force impacts. Always apply simultaneously with the coefficient of restitution (eq. 19) to solve for both post-collision velocities.",
    example={
        'problem': '3 kg at 4 m/s hits stationary 1 kg. After: 1 kg moves at 6 m/s. Find v1\'.',
        'steps': [
            '3*4 + 1*0 = 3*v1\' + 1*6',
            '12 = 3*v1\' + 6  =>  v1\' = 2 m/s',
        ],
        'answer': '3 kg body continues at 2 m/s'})

eq_block(19, 'Coefficient of Restitution',
    "e = (v2' - v1') / (v1 - v2)",
    [('e',        'Coefficient of restitution  (0 = plastic, 1 = elastic)', 'dimensionless'),
     ("v1, v2",   'Velocities before impact', 'm/s'),
     ("v1', v2'", 'Velocities after impact', 'm/s')],
    definition="Ratio of relative separation speed to relative approach speed along the line of impact. e = 1: perfectly elastic (no kinetic energy lost). e = 0: perfectly plastic (bodies stick together, maximum energy loss). Real impacts: 0 < e < 1. Apply simultaneously with conservation of momentum (eq. 18) to form two equations and solve for the two unknown post-collision velocities.",
    example={
        'problem': "Same collision: v1=4, v2=0, v1'=2, v2'=6. Find e.",
        'steps': [
            "e = (v2' - v1') / (v1 - v2) = (6 - 2) / (4 - 0) = 4/4 = 1.0",
        ],
        'answer': 'e = 1.0 (perfectly elastic collision)'})

subtopic_hdr('Kinematics and Kinetics - Rotation')

eq_block(20, "Newton's Second Law - Rotational Motion",
    'M = I * alpha',
    [('M',     'Net moment (torque)', 'N*m'),
     ('I',     'Mass moment of inertia about rotation axis', 'kg*m^2'),
     ('alpha', 'Angular acceleration', 'rad/s^2')],
    definition="Rotational analogue of F = ma: net torque about an axis equals the mass moment of inertia about that same axis multiplied by angular acceleration. I and M must be referenced to the same axis. Use the parallel-axis theorem I = I_cm + m*d^2 to shift I from the centroidal axis to a parallel axis at distance d.",
    example={
        'problem': 'A flywheel (I = 0.5 kg*m^2) has net torque 4 N*m. Find angular acceleration.',
        'steps': ['alpha = M/I = 4/0.5 = 8 rad/s^2'],
        'answer': 'alpha = 8 rad/s^2'})

eq_block(21, 'Angular Momentum',
    'L = I * w     and     M = dL/dt',
    [('L', 'Angular momentum', 'kg*m^2/s'),
     ('I', 'Mass moment of inertia', 'kg*m^2'),
     ('w', 'Angular velocity', 'rad/s'),
     ('M', 'Applied moment (torque)', 'N*m')],
    definition="Angular momentum L = I*w is conserved when the net external torque is zero (e.g. a spinning body with no friction or applied torque). The relation M = dL/dt is the most general rotational form of Newton's 2nd law, valid even when I changes (e.g. figure skater drawing arms in — I decreases so w increases to keep L constant).",
    example={
        'problem': 'Flywheel I = 0.5 kg*m^2 spinning at w = 100 rad/s. Find L.',
        'steps': ['L = I*w = 0.5 * 100 = 50 kg*m^2/s'],
        'answer': 'L = 50 kg*m^2/s'})

subtopic_hdr('Mechanical Vibrations - SDOF Systems')

eq_block(22, 'Undamped Natural Frequency',
    'wn = sqrt(k/m)',
    [('wn', 'Undamped natural frequency', 'rad/s'),
     ('k',  'Spring stiffness', 'N/m'),
     ('m',  'Mass', 'kg')],
    definition="Frequency at which a SDOF spring-mass system oscillates freely with no energy dissipation. For an undamped system, resonance occurs when excitation frequency equals wn, producing theoretically infinite amplitude. Applies to any SDOF system reducible to an equivalent k and m. Convert to Hz by f = wn / (2*pi). For a pendulum: wn = sqrt(g/L).",
    example={
        'problem': 'Spring-mass system: k = 4000 N/m, m = 10 kg. Find wn and frequency in Hz.',
        'steps': [
            'wn = sqrt(4000/10) = sqrt(400) = 20 rad/s',
            'f = wn/(2*pi) = 20/6.283 = 3.18 Hz',
        ],
        'answer': 'wn = 20 rad/s (3.18 Hz)'})

eq_block(23, 'Equation of Motion - Undamped Free Vibration',
    'm*x_ddot + k*x = F(t)',
    [('m',    'Mass', 'kg'),
     ('x_ddot', 'Acceleration (second derivative of x)', 'm/s^2'),
     ('k',    'Stiffness', 'N/m'),
     ('x',    'Displacement from equilibrium', 'm'),
     ('F(t)', 'External forcing function', 'N')],
    definition="Governing differential equation for a spring-mass system with no damping. With F(t) = 0 (free vibration), the solution is pure sinusoidal motion at wn. With harmonic forcing F(t) = F0*sin(w*t), the steady-state amplitude is F0/k / |1-(w/wn)^2|, which becomes unbounded as w approaches wn (resonance). Valid for small oscillations (linear spring behaviour).",
    example={
        'problem': 'm = 2 kg, k = 800 N/m, no forcing. Write EOM and state wn.',
        'steps': [
            '2*x_ddot + 800*x = 0',
            'wn = sqrt(800/2) = sqrt(400) = 20 rad/s',
        ],
        'answer': '2*x_ddot + 800*x = 0; wn = 20 rad/s'})

eq_block(24, 'Equation of Motion - Damped Free Vibration',
    'm*x_ddot + c*x_dot + k*x = F(t)',
    [('m',    'Mass', 'kg'),
     ('c',    'Viscous damping coefficient', 'N*s/m'),
     ('x_dot', 'Velocity', 'm/s'),
     ('k',    'Stiffness', 'N/m'),
     ('x',    'Displacement', 'm'),
     ('F(t)', 'External force', 'N')],
    definition="Standard SDOF equation of motion with viscous (velocity-proportional) damping. Free response depends on zeta: underdamped (zeta < 1, oscillatory exponential decay), critically damped (zeta = 1, fastest non-oscillatory return), overdamped (zeta > 1, sluggish non-oscillatory return). Most real mechanical damping is approximated as viscous for analytical convenience.",
    example={
        'problem': 'm=2 kg, c=12 N*s/m, k=800 N/m. Write EOM and classify damping.',
        'steps': [
            '2*x_ddot + 12*x_dot + 800*x = 0',
            'cc = 2*sqrt(k*m) = 2*sqrt(1600) = 80 N*s/m',
            'zeta = c/cc = 12/80 = 0.15  =>  underdamped',
        ],
        'answer': 'EOM: 2*x_ddot + 12*x_dot + 800*x = 0; zeta = 0.15 (underdamped)'})

eq_block(25, 'Damping Ratio',
    'zeta = c / (2*sqrt(k*m)) = c / (2*m*wn)',
    [('zeta', 'Damping ratio', 'dimensionless'),
     ('c',    'Damping coefficient', 'N*s/m'),
     ('k',    'Stiffness', 'N/m'),
     ('m',    'Mass', 'kg'),
     ('wn',   'Natural frequency', 'rad/s')],
    definition="Dimensionless ratio of actual damping to critical damping. zeta < 1: underdamped (oscillatory decay). zeta = 1: critically damped (no oscillation, fastest recovery). zeta > 1: overdamped (sluggish). Can also be determined experimentally from the logarithmic decrement of a free decay record (eq. 28). Typical values: structural steel 0.01-0.05; mechanical systems 0.05-0.3.",
    example={
        'problem': 'm = 2 kg, c = 12 N*s/m, k = 800 N/m. Find zeta.',
        'steps': ['zeta = 12 / (2*sqrt(800*2)) = 12 / (2*40) = 12/80 = 0.15'],
        'answer': 'zeta = 0.15 (underdamped system)'})

eq_block(26, 'Critical Damping Coefficient',
    'cc = 2*sqrt(k*m) = 2*m*wn',
    [('cc', 'Critical damping coefficient (boundary: zeta = 1)', 'N*s/m'),
     ('k',  'Stiffness', 'N/m'),
     ('m',  'Mass', 'kg'),
     ('wn', 'Undamped natural frequency', 'rad/s')],
    definition="Minimum damping coefficient at which the system returns to equilibrium without oscillating; defines the boundary between underdamped (c < cc) and overdamped (c > cc) behaviour. Serves as the reference value for computing the dimensionless damping ratio zeta = c/cc. Increasing stiffness k or mass m raises cc.",
    example={
        'problem': 'm = 5 kg, k = 2000 N/m. Find critical damping coefficient.',
        'steps': ['cc = 2*sqrt(k*m) = 2*sqrt(2000*5) = 2*sqrt(10000) = 2*100 = 200 N*s/m'],
        'answer': 'cc = 200 N*s/m'})

eq_block(27, 'Resonance Frequency (Damped Forced Vibration)',
    'wr = wn * sqrt(1 - 2*zeta^2)',
    [('wr',   'Frequency of maximum amplitude response', 'rad/s'),
     ('wn',   'Undamped natural frequency', 'rad/s'),
     ('zeta', 'Damping ratio', 'dimensionless')],
    note='Valid only when zeta <= 1/sqrt(2) = 0.707.',
    definition="Excitation frequency that produces the maximum steady-state vibration amplitude in a damped system; always lower than wn. Criterion: only exists when zeta < 1/sqrt(2) = 0.707. For zeta >= 0.707, the amplitude-frequency curve has no peak and decreases monotonically from the static deflection — resonance effectively disappears. For very light damping (zeta -> 0), wr approaches wn.",
    example={
        'problem': 'wn = 20 rad/s, zeta = 0.15. Find wr.',
        'steps': [
            'wr = 20 * sqrt(1 - 2*0.15^2) = 20 * sqrt(1 - 0.045) = 20 * sqrt(0.955)',
            '   = 20 * 0.977 = 19.54 rad/s',
        ],
        'answer': 'wr = 19.54 rad/s'})

eq_block(28, 'Logarithmic Decrement',
    'delta = ln(xn / x(n+1)) = 2*pi*zeta / sqrt(1 - zeta^2)',
    [('delta',    'Logarithmic decrement', 'dimensionless'),
     ('xn',       'Amplitude of nth cycle', 'm'),
     ('x(n+1)',   'Amplitude of (n+1)th cycle', 'm'),
     ('zeta',     'Damping ratio', 'dimensionless')],
    definition="Natural log of the ratio of consecutive free-vibration peak amplitudes; used to measure zeta experimentally from a free decay record. For better accuracy average over N cycles: delta = (1/N)*ln(x1/x_{N+1}). Assumes constant viscous damping over the measured cycles and an underdamped system. A quick non-destructive method to characterise damping in structures and rotating machinery.",
    example={
        'problem': 'Consecutive amplitudes: x1 = 18 mm, x2 = 12 mm. Find zeta.',
        'steps': [
            'delta = ln(18/12) = ln(1.5) = 0.405',
            'zeta = delta / sqrt(4*pi^2 + delta^2) = 0.405 / sqrt(39.48 + 0.164) = 0.405/6.296 = 0.0644',
        ],
        'answer': 'zeta = 0.064'})


# ══════════════════════════════════════════════════════════════════════════════
# ME 103/203  FLUID MECHANICS
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 103 / 203', 'Fluid Mechanics')
subtopic_hdr('Fluid Properties')

eq_block(29, 'Density',
    'rho = m / V',
    [('rho', 'Mass density', 'kg/m^3'),
     ('m',   'Mass', 'kg'),
     ('V',   'Volume', 'm^3')],
    definition="Mass per unit volume; a fundamental fluid property appearing in nearly every flow equation. For liquids, rho is nearly constant (incompressible); for gases, rho varies strongly with pressure and temperature and is found from the ideal gas law: rho = P/(R*T). Water at 20 degC: 998 kg/m^3; air at 20 degC, 1 atm: 1.20 kg/m^3.",
    example={
        'problem': 'A 5-litre container holds 4.65 kg of oil. Find its density.',
        'steps': ['rho = m/V = 4.65 / 0.005 = 930 kg/m^3'],
        'answer': 'rho = 930 kg/m^3 (typical light crude oil)'})

eq_block(30, "Newton's Law of Viscosity",
    'tau = mu * (du/dy)',
    [('tau',   'Shear stress in fluid', 'Pa'),
     ('mu',    'Dynamic viscosity', 'Pa*s'),
     ('du/dy', 'Velocity gradient perpendicular to flow', 's^-1')],
    definition="Defines viscosity as the proportionality constant between shear stress and velocity gradient (shear rate); applies to Newtonian fluids only (water, air, most mineral oils). Non-Newtonian fluids (blood, paint, polymer melts, cornstarch) deviate — do not apply. For gases, mu increases with temperature; for liquids, mu decreases with temperature.",
    example={
        'problem': 'Oil (mu=0.1 Pa*s) between plates 5 mm apart; upper plate at 2 m/s, lower fixed. Find tau.',
        'steps': [
            'du/dy = 2 / 0.005 = 400 s^-1',
            'tau = 0.1 * 400 = 40 Pa',
        ],
        'answer': 'tau = 40 Pa'})

eq_block(31, 'Kinematic Viscosity',
    'nu = mu / rho',
    [('nu',  'Kinematic viscosity', 'm^2/s'),
     ('mu',  'Dynamic viscosity', 'Pa*s'),
     ('rho', 'Fluid density', 'kg/m^3')],
    definition="Dynamic viscosity normalised by density; appears naturally in the Reynolds number (Re = V*D/nu) and is directly measured by many viscometers (e.g. Ostwald). Typical values: water at 20 degC: 1.0e-6 m^2/s (1 cSt); air at 20 degC: 1.5e-5 m^2/s. For gases, nu increases with temperature even when rho decreases.",
    example={
        'problem': 'Oil: mu = 0.1 Pa*s, rho = 900 kg/m^3. Find nu.',
        'steps': ['nu = 0.1 / 900 = 1.11e-4 m^2/s = 111 cSt'],
        'answer': 'nu = 1.11e-4 m^2/s'})

subtopic_hdr('Hydrostatics')

eq_block(32, 'Hydrostatic Pressure',
    'P = P0 + rho*g*h',
    [('P',   'Absolute pressure at depth h', 'Pa'),
     ('P0',  'Pressure at free surface', 'Pa'),
     ('rho', 'Fluid density', 'kg/m^3'),
     ('g',   'Gravitational acceleration (9.81)', 'm/s^2'),
     ('h',   'Depth below free surface', 'm')],
    definition="Pressure in a static fluid increases linearly with depth due to the weight of fluid above. P0 at the free surface is usually atmospheric pressure (101.3 kPa). For gauge pressure (relative to atmospheric), P_gauge = rho*g*h. Applies strictly to stationary fluids; does not account for velocity-driven pressure changes (Bernoulli, eq. 36).",
    example={
        'problem': 'Find gauge pressure at 15 m depth in seawater (rho = 1025 kg/m^3).',
        'steps': ['P_gauge = rho*g*h = 1025 * 9.81 * 15 = 150,919 Pa = 151 kPa'],
        'answer': 'P_gauge = 151 kPa'})

eq_block(33, "Buoyancy Force - Archimedes' Principle",
    'Fb = rho_fluid * V_sub * g',
    [('Fb',        'Buoyancy (upward) force', 'N'),
     ('rho_fluid', 'Density of surrounding fluid', 'kg/m^3'),
     ('V_sub',     'Volume of body submerged in fluid', 'm^3'),
     ('g',         'Gravitational acceleration', 'm/s^2')],
    definition="Upward buoyancy force on a submerged or floating body equals the weight of fluid displaced. A body floats when Fb >= W (buoyancy >= weight). For partial submersion, V_sub is only the submerged portion; the waterplane area adjusts until equilibrium. Applies to any static fluid regardless of body shape.",
    example={
        'problem': 'A 0.2 m^3 object submerged in water. Find buoyancy force.',
        'steps': ['Fb = 1000 * 0.2 * 9.81 = 1962 N'],
        'answer': 'Fb = 1962 N (= 1.96 kN)'})

subtopic_hdr('Flow Equations')

eq_block(34, 'Continuity Equation - Incompressible Flow',
    'A1*V1 = A2*V2 = Q',
    [('A1, A2', 'Cross-sectional areas at sections 1 and 2', 'm^2'),
     ('V1, V2', 'Mean flow velocities at sections 1 and 2', 'm/s'),
     ('Q',      'Volumetric flow rate (constant)', 'm^3/s')],
    definition="Conservation of mass for steady, incompressible (constant-density) flow through a streamtube or pipe; volumetric flow rate Q is constant. Applies to liquids and low-speed gas flows (Ma < 0.3). A smaller cross-sectional area results in a higher velocity (and by Bernoulli, eq. 36, a lower pressure). Use A = pi*D^2/4 for circular pipes.",
    example={
        'problem': 'Water flows at V1=3 m/s in 200 mm pipe. Pipe reduces to 100 mm. Find V2.',
        'steps': [
            'A1 = pi*(0.2)^2/4 = 0.03142 m^2',
            'A2 = pi*(0.1)^2/4 = 0.007854 m^2',
            'V2 = A1*V1/A2 = 0.03142*3 / 0.007854 = 12 m/s',
        ],
        'answer': 'V2 = 12 m/s'})

eq_block(35, 'Continuity Equation - Compressible Flow',
    'rho1*A1*V1 = rho2*A2*V2 = m_dot',
    [('rho1, rho2', 'Fluid densities at sections 1 and 2', 'kg/m^3'),
     ('A1, A2',     'Cross-sectional areas', 'm^2'),
     ('V1, V2',     'Velocities', 'm/s'),
     ('m_dot',      'Mass flow rate (constant)', 'kg/s')],
    definition="Conservation of mass for steady compressible flow (variable density). Mass flow rate m_dot is conserved along the duct. Required for high-speed gas flows (Ma > 0.3) or processes with large pressure or temperature changes. Density is found from the ideal gas law at each section; combine with Bernoulli or isentropic relations for the complete solution.",
    example={
        'problem': 'Air: rho1=1.2 kg/m^3, V1=50 m/s, A1=0.1 m^2. At exit: rho2=0.9, A2=0.08 m^2. Find V2.',
        'steps': [
            'm_dot = rho1*A1*V1 = 1.2*0.1*50 = 6 kg/s',
            'V2 = m_dot/(rho2*A2) = 6/(0.9*0.08) = 83.3 m/s',
        ],
        'answer': 'V2 = 83.3 m/s'})

eq_block(36, "Bernoulli's Equation (steady, inviscid, incompressible)",
    'P1 + (1/2)*rho*V1^2 + rho*g*z1  =  P2 + (1/2)*rho*V2^2 + rho*g*z2',
    [('P',   'Static pressure', 'Pa'),
     ('rho', 'Fluid density', 'kg/m^3'),
     ('V',   'Flow velocity', 'm/s'),
     ('g',   'Gravitational acceleration', 'm/s^2'),
     ('z',   'Elevation above datum', 'm')],
    note='Head form (divide by rho*g): P/rho*g + V^2/2g + z = H (total head, m)',
    definition="Conservation of mechanical energy along a streamline. Four key assumptions: (1) steady flow, (2) inviscid (no viscous losses), (3) incompressible (constant rho), (4) applied along a single streamline. For real pipe systems with friction and fittings, add a head-loss term hL: H1 = H2 + hL. Never apply across a pump or turbine without adding the head gained or lost.",
    example={
        'problem': 'Pipe narrows 200mm->100mm, same elevation. P1=150 kPa, V1=3 m/s. Find P2 (water).',
        'steps': [
            'V2 = 12 m/s (from continuity, eq. 34)',
            'P2 = P1 + (1/2)*rho*(V1^2 - V2^2)',
            'P2 = 150000 + (1/2)*1000*(9 - 144) = 150000 - 67500 = 82,500 Pa',
        ],
        'answer': 'P2 = 82.5 kPa'})

eq_block(37, 'Momentum Equation - Steady 1-D Flow',
    'Sum(F) = rho*Q*(V2 - V1) = m_dot*(V2 - V1)',
    [('Sum(F)',  'Net external force on control volume', 'N'),
     ('rho',    'Fluid density', 'kg/m^3'),
     ('Q',      'Volumetric flow rate', 'm^3/s'),
     ('V1, V2', 'Velocities at inlet and outlet', 'm/s'),
     ('m_dot',  'Mass flow rate', 'kg/s')],
    definition="Newton's 2nd law applied to a fluid control volume with steady mass flow; net external force equals the rate of momentum leaving minus entering. Used to find forces on pipe bends, nozzles, vanes, and turbine blades where Bernoulli cannot directly give forces. Vectorial: apply as two separate equations in x and y directions. Pressure forces at inlet/outlet must be included in Sum(F).",
    example={
        'problem': 'Water m_dot=10 kg/s hits a flat plate at 5 m/s and is deflected 90 deg. Find reaction force.',
        'steps': [
            'Fx = m_dot*(0 - 5) = -50 N (x-direction)',
            'Fy = m_dot*(5 - 0) = +50 N (y-direction)',
            '|F| = sqrt(50^2 + 50^2) = 70.7 N',
        ],
        'answer': 'Reaction force = 70.7 N at 45 deg'})

subtopic_hdr('Dimensionless Numbers')

eq_block(38, 'Reynolds Number',
    'Re = rho*V*D/mu = V*D/nu',
    [('Re', 'Reynolds number (Re<2300 laminar; Re>4000 turbulent in pipe)', 'dimensionless'),
     ('rho', 'Fluid density', 'kg/m^3'),
     ('V',   'Characteristic velocity', 'm/s'),
     ('D',   'Characteristic length (pipe diameter)', 'm'),
     ('mu',  'Dynamic viscosity', 'Pa*s'),
     ('nu',  'Kinematic viscosity nu=mu/rho', 'm^2/s')],
    definition="Ratio of inertia forces to viscous forces; the primary parameter governing flow regime. For internal pipe flow: Re < 2300 laminar, 2300-4000 transitional, Re > 4000 turbulent (Munson et al., White). For external flow over a flat plate, transition is typically at Re ~5e5. For non-circular ducts, use the hydraulic diameter Dh = 4*A/P (A = cross-sectional area, P = wetted perimeter).",
    example={
        'problem': 'Water (rho=1000, mu=0.001 Pa*s) at 2 m/s in 50 mm pipe. Find Re.',
        'steps': ['Re = 1000*2*0.05 / 0.001 = 100,000'],
        'answer': 'Re = 100,000 => turbulent flow (Re >> 4000)'})

eq_block(39, 'Froude Number',
    'Fr = V / sqrt(g*L)',
    [('Fr', 'Froude number (Fr<1 subcritical; Fr>1 supercritical)', 'dimensionless'),
     ('V',  'Flow velocity', 'm/s'),
     ('g',  'Gravitational acceleration', 'm/s^2'),
     ('L',  'Characteristic length (e.g. hydraulic depth)', 'm')],
    definition="Ratio of inertia to gravitational forces; governs free-surface and open-channel flow. Fr < 1: subcritical (slow/deep), surface waves can propagate upstream. Fr > 1: supercritical (fast/shallow), disturbances cannot propagate upstream. Fr = 1: critical flow (transition). L is hydraulic depth (cross-sectional area / water surface width) for non-rectangular channels.",
    example={
        'problem': 'River: V = 1.5 m/s, hydraulic depth = 0.8 m. Find Fr.',
        'steps': ['Fr = 1.5 / sqrt(9.81*0.8) = 1.5 / 2.801 = 0.535'],
        'answer': 'Fr = 0.535 < 1 => subcritical (tranquil) flow'})

eq_block(40, 'Mach Number',
    'Ma = V / c     where  c = sqrt(gamma*R*T)',
    [('Ma',    'Mach number (Ma<1 subsonic; Ma>1 supersonic)', 'dimensionless'),
     ('V',     'Flow velocity', 'm/s'),
     ('c',     'Speed of sound in fluid', 'm/s'),
     ('gamma', 'Ratio of specific heats cp/cv (air = 1.4)', 'dimensionless'),
     ('R',     'Specific gas constant (air = 287)', 'J/(kg*K)'),
     ('T',     'Absolute temperature', 'K')],
    definition="Ratio of flow velocity to the local speed of sound; determines the importance of compressibility effects. Ma < 0.3: compressibility negligible (< 5% error using incompressible equations). 0.3 < Ma < 1: subsonic compressible. Ma = 1: sonic (choked flow at a nozzle throat). Ma > 1: supersonic (shock waves possible). Speed of sound depends on temperature only (not pressure) for an ideal gas.",
    example={
        'problem': 'Aircraft at V=340 m/s. Air T=288 K. Find Ma.',
        'steps': [
            'c = sqrt(1.4*287*288) = sqrt(115,834) = 340.3 m/s',
            'Ma = 340/340.3 = 0.999',
        ],
        'answer': 'Ma = 1.0 (transonic)'})

subtopic_hdr('Pipe Flow - Head Losses')

eq_block(41, 'Darcy-Weisbach Equation - Major (Friction) Loss',
    'hf = f * (L/D) * V^2/(2g)',
    [('hf', 'Friction head loss', 'm'),
     ('f',  'Darcy friction factor (dimensionless)', 'dimensionless'),
     ('L',  'Pipe length', 'm'),
     ('D',  'Internal pipe diameter', 'm'),
     ('V',  'Mean flow velocity', 'm/s'),
     ('g',  'Gravitational acceleration', 'm/s^2')],
    note='Fanning form (some texts): hf = 4f(L/D)(V^2/2g). Darcy f = 4 x Fanning f.',
    definition="Calculates head loss due to pipe wall friction for fully developed flow in a straight pipe. Valid for both laminar and turbulent regimes using the appropriate friction factor: f = 64/Re for laminar (eq. 42); Colebrook-White or Moody chart for turbulent (eq. 43). L/D must be sufficient for fully developed conditions (approximately > 20 for turbulent flow).",
    example={
        'problem': 'Water V=3 m/s, L=500 m, D=150 mm, f=0.015. Find hf.',
        'steps': [
            'hf = 0.015 * (500/0.15) * (3^2/(2*9.81))',
            'hf = 0.015 * 3333 * 0.459 = 22.9 m',
        ],
        'answer': 'hf = 22.9 m'})

eq_block(42, 'Laminar Friction Factor - Hagen-Poiseuille',
    'f = 64 / Re       (valid for Re < 2300)',
    [('f',  'Darcy friction factor', 'dimensionless'),
     ('Re', 'Reynolds number', 'dimensionless')],
    definition="Exact analytical result for the Darcy friction factor in fully developed laminar pipe flow, derived from the parabolic (Hagen-Poiseuille) velocity profile. Valid only for Re < 2300. Above this limit, the flow transitions to turbulent and this formula significantly underestimates friction losses. Note: f = 64/Re is the Darcy form; some texts use the Fanning form f = 16/Re.",
    example={
        'problem': 'Viscous oil flows in a pipe with Re = 1200. Find Darcy friction factor.',
        'steps': ['f = 64/Re = 64/1200 = 0.0533'],
        'answer': 'f = 0.053 (laminar flow, Re < 2300)'})

eq_block(43, 'Turbulent Friction Factor - Colebrook-White Equation',
    '1/sqrt(f) = -2*log10[ eps/(3.7*D) + 2.51/(Re*sqrt(f)) ]',
    [('f',   'Darcy friction factor', 'dimensionless'),
     ('eps', 'Absolute pipe roughness', 'm'),
     ('D',   'Pipe internal diameter', 'm'),
     ('Re',  'Reynolds number', 'dimensionless')],
    note='Implicit equation - requires iteration or Moody chart. Valid for turbulent flow (Re > 4000).',
    definition="Implicit equation for the Darcy friction factor in turbulent pipe flow (Re > 4000); must be solved iteratively (start with f = 0.02, substitute, repeat until convergence in 3-4 iterations). Combines smooth-wall and rough-wall behaviour on the Moody chart. The explicit Swamee-Jain approximation gives < 3% error: 1/sqrt(f) = -2*log10[eps/(3.7D) + 5.74/Re^0.9].",
    example={
        'problem': 'Steel pipe: eps=0.046 mm, D=100 mm, Re=100,000. Estimate f (1st iteration, f0=0.018).',
        'steps': [
            '1/sqrt(f) = -2*log10[0.046/(3.7*100) + 2.51/(100000*sqrt(0.018))]',
            '         = -2*log10[1.24e-4 + 1.87e-4] = -2*log10[3.11e-4] = 7.01',
            'f1 = 1/7.01^2 = 0.0203  (iterate to convergence)',
        ],
        'answer': 'f = 0.020 after convergence'})

eq_block(44, 'Minor (Local) Losses - Fittings, Entry, Exit',
    'hm = KL * V^2/(2g)',
    [('hm', 'Minor head loss', 'm'),
     ('KL', 'Loss coefficient (K_entry=0.5, K_exit=1.0)', 'dimensionless'),
     ('V',  'Velocity at the fitting', 'm/s'),
     ('g',  'Gravitational acceleration', 'm/s^2')],
    definition="Head loss at flow disturbances (fittings, valves, bends, entries, exits) due to flow separation and turbulent mixing. KL is empirical and obtained from tables. Common values: sharp-edged pipe entry 0.5, pipe exit (into reservoir) 1.0, fully open globe valve ~10, 90 deg elbow 0.3-1.5. In short pipe systems with many fittings, minor losses can dominate friction losses.",
    example={
        'problem': 'Sharp-edged pipe entry (KL=0.5), V=3 m/s. Find hm.',
        'steps': ['hm = 0.5 * 3^2/(2*9.81) = 0.5 * 9/19.62 = 0.229 m'],
        'answer': 'hm = 0.23 m'})

eq_block(45, 'Sudden Expansion Loss - Borda-Carnot',
    'h_exp = (V1 - V2)^2 / (2g)',
    [('h_exp',  'Head loss at sudden expansion', 'm'),
     ('V1',     'Upstream velocity (smaller pipe)', 'm/s'),
     ('V2',     'Downstream velocity (larger pipe)', 'm/s'),
     ('A1, A2', 'Cross-sectional areas upstream and downstream', 'm^2'),
     ('g',      'Gravitational acceleration', 'm/s^2')],
    definition="Head loss when flow abruptly expands into a larger-diameter pipe due to flow separation and turbulent dissipation in the re-attachment zone. Always a loss (energy destroyed, not converted). V2 must be found first from continuity (eq. 34). For a gradual expansion (diffuser), losses are much smaller. Sudden contraction has KL ~ 0.5 (minor loss form, eq. 44).",
    example={
        'problem': 'Flow at 4 m/s in 80 mm pipe suddenly expands to 150 mm. Find h_exp.',
        'steps': [
            'V2 = V1*(A1/A2) = 4*(80/150)^2 = 4*0.284 = 1.14 m/s',
            'h_exp = (4 - 1.14)^2 / (2*9.81) = (2.86)^2/19.62 = 8.18/19.62 = 0.417 m',
        ],
        'answer': 'h_exp = 0.42 m'})

subtopic_hdr('Fluid Machinery')

eq_block(46, 'Hydraulic Power',
    'P = rho*g*Q*H',
    [('P',   'Hydraulic power delivered to fluid', 'W'),
     ('rho', 'Fluid density', 'kg/m^3'),
     ('g',   'Gravitational acceleration', 'm/s^2'),
     ('Q',   'Volumetric flow rate', 'm^3/s'),
     ('H',   'Total head developed', 'm')],
    definition="Power transferred between a pump/turbine and the fluid. For a pump, H is the total head added (net useful head); for a turbine, H is the head extracted. Q must be in m^3/s and H in metres of fluid head. This is the ideal fluid power; actual shaft power differs by the pump or turbine efficiency (eq. 47).",
    example={
        'problem': 'Pump delivers Q=0.05 m^3/s of water at H=25 m total head. Find hydraulic power.',
        'steps': ['P = 1000*9.81*0.05*25 = 12,263 W'],
        'answer': 'P = 12.3 kW'})

eq_block(47, 'Pump Overall Efficiency',
    'eta = P_fluid / P_shaft = rho*g*Q*H / P_shaft',
    [('eta',      'Overall pump efficiency', 'dimensionless (0 to 1)'),
     ('P_fluid',  'Power delivered to fluid', 'W'),
     ('P_shaft',  'Shaft (brake) power input to pump', 'W'),
     ('rho',      'Fluid density', 'kg/m^3'),
     ('Q',        'Flow rate', 'm^3/s'),
     ('H',        'Total head', 'm')],
    definition="Ratio of useful hydraulic power delivered to the fluid to the mechanical shaft power consumed. Accounts for all internal losses: hydraulic (flow separation), mechanical (bearing and seal friction), and volumetric (internal leakage). Typical centrifugal pump: 60-85% at best efficiency point (BEP). Operating far from BEP reduces efficiency and can cause cavitation.",
    example={
        'problem': 'Hydraulic power = 12.3 kW, shaft power input = 15 kW. Find eta.',
        'steps': ['eta = 12263/15000 = 0.818'],
        'answer': 'eta = 81.8%'})


# ══════════════════════════════════════════════════════════════════════════════
# ME 104/204  MECHANICS AND MATERIALS
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 104 / 204', 'Mechanics and Materials')
subtopic_hdr('Stress, Strain and Elastic Constants')

eq_block(48, 'Direct (Normal) Stress',
    'sigma = F / A',
    [('sigma', 'Normal stress (positive = tension)', 'Pa'),
     ('F',     'Axial force (positive = tensile)', 'N'),
     ('A',     'Cross-sectional area', 'm^2')],
    definition="Force per unit area acting perpendicular (normal) to the cross-section; positive = tension, negative = compression. Assumes force is uniformly distributed across the section — valid away from load application points and geometric discontinuities (Saint-Venant's principle). Near holes, notches, or fillets, apply a stress concentration factor Kt: sigma_max = Kt * (F/A).",
    example={
        'problem': 'A 25 mm diameter steel rod carries 50 kN tensile load. Find sigma.',
        'steps': [
            'A = pi*(0.025)^2/4 = 4.909e-4 m^2',
            'sigma = 50000 / 4.909e-4 = 101.9 MPa',
        ],
        'answer': 'sigma = 102 MPa (tensile)'})

eq_block(49, 'Direct (Normal) Strain',
    'eps = deltaL / L0',
    [('eps',    'Normal strain', 'dimensionless (m/m)'),
     ('deltaL', 'Change in length', 'm'),
     ('L0',     'Original (gauge) length', 'm')],
    definition="Fractional elongation or contraction along the direction of loading; dimensionless. Positive = extension (tensile load), negative = shortening (compressive load). Assumes uniform deformation over the gauge length L0. For non-uniform deformation, strain is defined locally as eps = d(deltaL)/dL. Typical elastic strains in steel are on the order of 500-2000 microstrain (1e-4 to 2e-3).",
    example={
        'problem': 'A 500 mm rod elongates by 0.255 mm under load. Find eps.',
        'steps': ['eps = 0.255 / 500 = 5.10e-4 (510 microstrain)'],
        'answer': 'eps = 5.1e-4'})

eq_block(50, "Hooke's Law - Elastic Region",
    'sigma = E * eps',
    [('sigma', 'Normal stress', 'Pa'),
     ('E',     "Young's modulus (steel = 200 GPa)", 'Pa'),
     ('eps',   'Normal strain', 'dimensionless')],
    definition="Linear proportionality between stress and strain; the foundational equation of linear elasticity. Valid only below the proportional limit (linear portion of the stress-strain curve, below yield). E is a material constant: steel ~200 GPa, aluminium ~70 GPa, concrete ~30 GPa (compression only), rubber ~0.01-0.1 GPa. Above the yield stress, plastic deformation occurs and Hooke's law no longer applies.",
    example={
        'problem': 'Steel rod (E=200 GPa) under sigma=102 MPa. Find eps.',
        'steps': ['eps = sigma/E = 102e6 / 200e9 = 5.1e-4'],
        'answer': 'eps = 5.1e-4 (510 microstrain)'})

eq_block(51, 'Shear Stress and Shear Strain',
    'tau = Fs / A     and     gamma = tau / G',
    [('tau',   'Shear stress', 'Pa'),
     ('Fs',    'Shear force', 'N'),
     ('A',     'Area on which shear acts', 'm^2'),
     ('gamma', 'Shear strain (engineering)', 'dimensionless (rad)'),
     ('G',     'Shear modulus (steel = 79 GPa)', 'Pa')],
    definition="Shear stress acts tangential to the cross-section; shear strain is the angular distortion in radians. The shear form of Hooke's law (tau = G*gamma) holds in the elastic region below the shear yield stress. tau = Fs/A is the average shear stress; actual distribution is non-uniform (parabolic for rectangular sections — eq. 59). G relates to E and nu via eq. 53.",
    example={
        'problem': '20x20 mm section, 8 kN shear force, G=80 GPa. Find tau and gamma.',
        'steps': [
            'tau = 8000 / (0.020*0.020) = 20 MPa',
            'gamma = tau/G = 20e6 / 80e9 = 2.5e-4 rad',
        ],
        'answer': 'tau = 20 MPa, gamma = 2.5e-4 rad'})

eq_block(52, "Poisson's Ratio",
    'nu = -eps_lateral / eps_axial',
    [('nu',          "Poisson's ratio (steel = 0.30)", 'dimensionless'),
     ('eps_lateral', 'Transverse (lateral) strain', 'dimensionless'),
     ('eps_axial',   'Axial (longitudinal) strain', 'dimensionless')],
    definition="Dimensionless material property relating lateral contraction to axial elongation under uniaxial loading. The negative sign ensures nu is positive (lateral and axial strains are opposite in sign). Typical values: steel 0.28-0.30, aluminium 0.33, concrete 0.15-0.20, rubber ~0.5 (nearly incompressible), cork ~0 (no lateral deformation). Thermodynamic limits: -1 < nu < 0.5 for isotropic materials.",
    example={
        'problem': 'Steel (nu=0.30) under axial strain eps_axial = 5.1e-4. Find lateral strain.',
        'steps': ['eps_lateral = -nu * eps_axial = -0.30 * 5.1e-4 = -1.53e-4'],
        'answer': 'eps_lateral = -1.53e-4 (contraction)'})

eq_block(53, 'Relationship Between Elastic Constants',
    'G = E / [2*(1 + nu)]     and     K = E / [3*(1 - 2*nu)]',
    [('G',  'Shear modulus', 'Pa'),
     ('E',  "Young's modulus", 'Pa'),
     ('nu', "Poisson's ratio", 'dimensionless'),
     ('K',  'Bulk modulus', 'Pa')],
    definition="For a linear isotropic material, only two independent elastic constants fully characterise the behaviour; E, G, nu, and K are all interrelated. Knowing any two allows the others to be derived. Assumes isotropy (same properties in all directions); invalid for anisotropic materials such as composites, wood, or crystals. Physically, nu -> 0.5 gives K -> infinity (incompressible material).",
    example={
        'problem': 'Steel: E = 200 GPa, nu = 0.30. Find G.',
        'steps': ['G = 200e9 / (2*(1+0.30)) = 200e9 / 2.60 = 76.9 GPa'],
        'answer': 'G = 76.9 GPa (literature: ~79 GPa)'})

subtopic_hdr('Thermal Effects')

eq_block(54, 'Free Thermal Strain',
    'eps_T = alpha * DeltaT',
    [('eps_T',   'Thermal strain (free expansion)', 'dimensionless'),
     ('alpha',   'Coefficient of linear thermal expansion (steel = 12e-6)', '/degC'),
     ('DeltaT',  'Temperature change', 'degC (or K)')],
    definition="Dimensional change per unit length of an unconstrained member due to temperature change; no stress is generated if the member is free to expand or contract. alpha is a material property: steel ~12e-6/degC, aluminium ~23e-6/degC, concrete ~10e-6/degC. Temperature changes in degC and K are numerically equal. For 2D/3D, thermal strain acts equally in all directions for isotropic materials.",
    example={
        'problem': 'Steel rail (alpha=12e-6 /degC) heated by 40 degC. Find eps_T and extension of a 10 m rail.',
        'steps': [
            'eps_T = 12e-6 * 40 = 4.8e-4',
            'Extension = eps_T * L0 = 4.8e-4 * 10 = 4.8 mm',
        ],
        'answer': 'eps_T = 4.8e-4 (4.8 mm per 10 m rail)'})

eq_block(55, 'Constrained Thermal Stress',
    'sigma_T = E * alpha * DeltaT',
    [('sigma_T', 'Thermal stress (compressive if expansion prevented)', 'Pa'),
     ('E',       "Young's modulus", 'Pa'),
     ('alpha',   'Coefficient of thermal expansion', '/degC'),
     ('DeltaT',  'Temperature change', 'degC')],
    definition="Stress induced when thermal expansion is fully prevented by rigid supports. Compressive if heated and restrained; tensile if cooled and restrained. This is the fully-constrained (statically determinate) case. For partial constraint or statically indeterminate systems, a compatibility equation must be combined with equilibrium to find the actual stress.",
    example={
        'problem': 'Same steel rail rigidly fixed at both ends. DeltaT = 40 degC. Find sigma_T.',
        'steps': ['sigma_T = 200e9 * 12e-6 * 40 = 96 MPa (compressive)'],
        'answer': 'sigma_T = 96 MPa compressive'})

subtopic_hdr('Bending of Beams')

eq_block(56, 'Second Moment of Area - Rectangle (about centroidal axis)',
    'Ixx = b*h^3 / 12',
    [('Ixx', 'Second moment of area about horizontal centroidal axis', 'm^4'),
     ('b',   'Width of rectangle', 'm'),
     ('h',   'Height (depth) of rectangle', 'm')],
    definition="Area moment of inertia of a rectangle about its horizontal centroidal axis; quantifies a cross-section's resistance to bending. The formula I = bh^3/12 is about the centroidal axis parallel to b (depth h controls resistance). To shift to a parallel axis at distance d, use the parallel-axis theorem: I_new = I_centroid + A*d^2. For compound sections, sum individual contributions.",
    example={
        'problem': 'Beam cross-section: 60 mm wide x 120 mm deep. Find Ixx.',
        'steps': ['Ixx = 0.060 * (0.120)^3 / 12 = 0.060 * 1.728e-3 / 12 = 8.64e-6 m^4'],
        'answer': 'Ixx = 8.64e-6 m^4'})

eq_block(57, 'Second Moment of Area - Solid Circle',
    'I = pi*d^4 / 64 = pi*r^4 / 4',
    [('I', 'Second moment of area about diameter', 'm^4'),
     ('d', 'Diameter', 'm'),
     ('r', 'Radius', 'm')],
    definition="Area moment of inertia for a solid circular cross-section about a diameter; used in bending calculations. Note the distinction from torsion: I = pi*d^4/64 for bending; polar moment J = pi*d^4/32 = 2*I for torsion. For a hollow circle, subtract the inner contribution: I = pi*(do^4 - di^4)/64.",
    example={
        'problem': '80 mm diameter solid shaft. Find I.',
        'steps': ['I = pi*(0.080)^4 / 64 = pi*4.096e-6/64 = 2.01e-7 m^4'],
        'answer': 'I = 2.01e-7 m^4'})

eq_block(58, 'Flexure Formula (Bending Stress - Euler-Bernoulli)',
    'sigma = M*y / I     =>     sigma_max = M / Z',
    [('sigma',     'Bending stress at distance y from neutral axis', 'Pa'),
     ('M',         'Bending moment', 'N*m'),
     ('y',         'Distance from neutral axis to point of interest', 'm'),
     ('I',         'Second moment of area about neutral axis', 'm^4'),
     ('sigma_max', 'Maximum bending stress (at extreme fibre)', 'Pa'),
     ('Z',         'Section modulus Z = I/y_max', 'm^3')],
    definition="Bending stress is linearly distributed: zero at the neutral axis, maximum tension and compression at the extreme fibres. Assumes Euler-Bernoulli beam theory (plane sections remain plane after bending). Valid for slender beams (span/depth > ~10) and small deflections. Section modulus Z = I/y_max is convenient for directly calculating the maximum stress. Not valid near concentrated loads or supports.",
    example={
        'problem': 'Simply supported beam, M_max=12 kN*m, 60x120 mm rectangle. Find sigma_max.',
        'steps': [
            'I = 8.64e-6 m^4 (from eq. 56)',
            'y_max = 120/2 mm = 0.060 m',
            'sigma_max = 12000 * 0.060 / 8.64e-6 = 83.3 MPa',
        ],
        'answer': 'sigma_max = 83.3 MPa (at extreme fibres)'})

eq_block(59, 'Shear Stress in Beams',
    'tau = V*Q / (I*b)',
    [('tau', 'Horizontal shear stress at depth y', 'Pa'),
     ('V',   'Shear force at the cross-section', 'N'),
     ('Q',   'First moment of area of section above y about neutral axis', 'm^3'),
     ('I',   'Second moment of area of full cross-section', 'm^4'),
     ('b',   'Width of cross-section at depth y', 'm')],
    definition="Horizontal shear stress distribution in a beam due to transverse shear force. Q = A_above * y_bar_above (first moment of area of the section above the point of interest about the neutral axis). For a rectangle: parabolic distribution, zero at top/bottom surfaces, maximum at neutral axis = (3/2) * V/A. In I-beams, shear concentrates in the web.",
    example={
        'problem': '60x120 mm beam, V=20 kN. Find tau at the neutral axis.',
        'steps': [
            'Q_NA = (b * h/2) * (h/4) = 0.060 * 0.060 * 0.030 = 1.08e-4 m^3',
            'tau_NA = 20000*1.08e-4 / (8.64e-6*0.060) = 2.16 / 5.18e-7 = 4.17 MPa',
        ],
        'answer': 'tau_NA = 4.17 MPa'})

eq_block(60, 'Beam Deflection - Simply-Supported, Central Point Load',
    'delta_max = F*L^3 / (48*E*I)',
    [('delta_max', 'Maximum deflection at mid-span', 'm'),
     ('F',         'Point load at mid-span', 'N'),
     ('L',         'Span length', 'm'),
     ('E',         "Young's modulus", 'Pa'),
     ('I',         'Second moment of area', 'm^4')],
    definition="Maximum deflection at mid-span for a simply-supported beam with a concentrated central point load. The factor 48 comes from double integration of the bending moment diagram. Assumes Euler-Bernoulli beam (slender, small deflections, elastic). For a uniformly distributed load w over the full span: delta_max = 5*w*L^4 / (384*E*I).",
    example={
        'problem': '6 m span, central 20 kN load, 60x120 mm section (E=200 GPa). Find delta_max.',
        'steps': [
            'I = 8.64e-6 m^4',
            'delta_max = 20000*6^3 / (48*200e9*8.64e-6)',
            '          = 20000*216 / (82,944,000) = 4,320,000 / 82,944,000 = 0.0521 m',
        ],
        'answer': 'delta_max = 52.1 mm'})

eq_block(61, 'Beam Deflection - Cantilever, Tip Point Load',
    'delta_max = F*L^3 / (3*E*I)',
    [('delta_max', 'Maximum deflection at free end', 'm'),
     ('F',         'Point load at free end', 'N'),
     ('L',         'Cantilever length', 'm'),
     ('E',         "Young's modulus", 'Pa'),
     ('I',         'Second moment of area', 'm^4')],
    definition="Maximum deflection at the free tip of a cantilever beam under a concentrated end load. The denominator factor 3 (vs 48 for a simply-supported central load) means a cantilever deflects 16 times more than a simply-supported beam of identical L, F, and EI. For a cantilever with uniformly distributed load w: delta_max = w*L^4 / (8*E*I).",
    example={
        'problem': 'Cantilever 2 m long, 5 kN tip load, 60x120 mm section (E=200 GPa). Find delta_max.',
        'steps': [
            'delta_max = 5000*(2)^3 / (3*200e9*8.64e-6)',
            '          = 5000*8 / 5,184,000 = 40,000/5,184,000 = 0.00772 m',
        ],
        'answer': 'delta_max = 7.7 mm'})

subtopic_hdr('Torsion')

eq_block(62, 'Torsion Formula - Circular Shaft (Coulomb)',
    'tau = T*r / J     =>     tau_max = T*(d/2) / J',
    [('tau',     'Shear stress at radius r', 'Pa'),
     ('T',       'Applied torque', 'N*m'),
     ('r',       'Radial distance from shaft axis', 'm'),
     ('J',       'Polar second moment of area', 'm^4'),
     ('d',       'Shaft diameter', 'm')],
    definition="Shear stress varies linearly from zero at the shaft centre to maximum at the outer surface. Applies only to circular cross-sections (solid or hollow); non-circular sections require Saint-Venant's torsion theory. The formula is exact for uniform circular shafts in pure torsion within the elastic range. Stress distribution is independent of material (elastic or plastic, if elastic limit not exceeded).",
    example={
        'problem': 'Solid shaft d=40 mm, T=500 N*m. Find tau_max.',
        'steps': [
            'J = pi*(0.040)^4/32 = 2.513e-7 m^4',
            'tau_max = T*(d/2)/J = 500*0.020 / 2.513e-7 = 10/2.513e-7 = 39.8 MPa',
        ],
        'answer': 'tau_max = 39.8 MPa'})

eq_block(63, 'Polar Second Moment of Area - Solid Shaft',
    'J = pi*d^4 / 32',
    [('J', 'Polar second moment of area', 'm^4'),
     ('d', 'Shaft diameter', 'm')],
    definition="Geometric property of a solid circular cross-section controlling torsional stiffness; the torsional equivalent of the area second moment I for bending. J = 2*I for a circle since J = I_x + I_y = 2*(pi*d^4/64). Appears in both the torsion stress formula (tau = T*r/J) and the angle-of-twist formula (phi = T*L/(G*J)).",
    example={
        'problem': 'Find J for a 40 mm diameter solid shaft.',
        'steps': ['J = pi*(0.040)^4/32 = pi*2.56e-6/32 = 2.513e-7 m^4'],
        'answer': 'J = 2.513e-7 m^4'})

eq_block(64, 'Polar Second Moment of Area - Hollow Shaft',
    'J = pi*(do^4 - di^4) / 32',
    [('J',  'Polar second moment of area', 'm^4'),
     ('do', 'External (outer) diameter', 'm'),
     ('di', 'Internal (inner) diameter', 'm')],
    definition="Polar second moment for a hollow circular shaft obtained by subtracting the inner hole contribution from the solid shaft value. Hollow shafts are weight-efficient: material near the centre (low r) contributes little to J, so removing it saves mass with minimal torsional stiffness penalty. For the same J, a hollow shaft has a larger outer diameter and lower maximum stress than a solid shaft.",
    example={
        'problem': 'Hollow shaft: do=60 mm, di=40 mm. Find J.',
        'steps': [
            'J = pi*((0.060)^4 - (0.040)^4)/32',
            '  = pi*(1.296e-5 - 2.560e-6)/32 = pi*1.040e-5/32 = 1.021e-6 m^4',
        ],
        'answer': 'J = 1.021e-6 m^4'})

eq_block(65, 'Angle of Twist',
    'phi = T*L / (G*J)',
    [('phi', 'Angle of twist', 'rad'),
     ('T',   'Applied torque', 'N*m'),
     ('L',   'Shaft length', 'm'),
     ('G',   'Shear modulus', 'Pa'),
     ('J',   'Polar second moment of area', 'm^4')],
    definition="Total angular deformation along a uniform shaft under applied torque; the torsional equivalent of axial extension delta = P*L/(A*E). G*J is the torsional rigidity. For stepped shafts or shafts with torques applied at multiple points, sum contributions: phi_total = sum(T_i * L_i / (G_i * J_i)). Convert result from radians to degrees by multiplying by 180/pi.",
    example={
        'problem': '40 mm solid shaft (G=79 GPa), L=1 m, T=500 N*m. Find phi.',
        'steps': [
            'J = 2.513e-7 m^4',
            'phi = 500*1 / (79e9*2.513e-7) = 500/19,853 = 0.02519 rad = 1.44 deg',
        ],
        'answer': 'phi = 0.0252 rad (1.44 deg)'})

subtopic_hdr("Yield Criteria and Stress Transformations")

eq_block(66, "Principal Stresses - Mohr's Circle",
    'sigma1,2 = (sx+sy)/2 +/- sqrt[((sx-sy)/2)^2 + txy^2]',
    [('sigma1, sigma2', 'Major and minor principal stresses', 'Pa'),
     ('sx',            'Normal stress on x-face', 'Pa'),
     ('sy',            'Normal stress on y-face', 'Pa'),
     ('txy',           'Shear stress on x-face', 'Pa')],
    definition="At any point in a stressed body, there exist orientations (principal planes) on which shear stress is zero; the stresses on these planes are the principal stresses sigma1 (maximum) and sigma2 (minimum). Found geometrically as the centre plus or minus the radius of Mohr's circle. Principal stresses are needed as inputs for yield criteria (eqs. 68-69). For 3D problems, a third principal stress sigma3 also exists.",
    example={
        'problem': 'sx=80 MPa, sy=20 MPa, txy=30 MPa. Find principal stresses.',
        'steps': [
            'Centre C = (80+20)/2 = 50 MPa',
            'R = sqrt[((80-20)/2)^2 + 30^2] = sqrt[900+900] = sqrt(1800) = 42.4 MPa',
            'sigma1 = 50+42.4 = 92.4 MPa;  sigma2 = 50-42.4 = 7.6 MPa',
        ],
        'answer': 'sigma1 = 92.4 MPa, sigma2 = 7.6 MPa'})

eq_block(67, 'Maximum Shear Stress',
    'tau_max = sqrt[((sx-sy)/2)^2 + txy^2]  =  (sigma1 - sigma2)/2',
    [('tau_max',     'Maximum in-plane shear stress', 'Pa'),
     ('sx, sy',      'Normal stresses on x- and y-faces', 'Pa'),
     ('txy',         'Shear stress', 'Pa'),
     ('sigma1, sigma2', 'Principal stresses', 'Pa')],
    definition="Maximum in-plane shear stress equals the radius of Mohr's circle; acts on planes at 45 deg to the principal planes. For a full 3D stress state, also check the absolute maximum shear stress from all three principal stress pairs: tau_abs_max = max(|sigma1-sigma2|, |sigma2-sigma3|, |sigma1-sigma3|) / 2. This is the value used in the Tresca criterion (eq. 69).",
    example={
        'problem': 'Using results from eq. 66: sigma1=92.4 MPa, sigma2=7.6 MPa. Find tau_max.',
        'steps': ['tau_max = (sigma1-sigma2)/2 = (92.4-7.6)/2 = 84.8/2 = 42.4 MPa'],
        'answer': 'tau_max = 42.4 MPa'})

eq_block(68, 'Von Mises Yield Criterion (2-D)',
    'sigma_VM = sqrt(sigma1^2 - sigma1*sigma2 + sigma2^2) <= sigma_y',
    [('sigma_VM',      'Von Mises equivalent stress (distortion energy criterion)', 'Pa'),
     ('sigma1, sigma2', 'Principal stresses', 'Pa'),
     ('sigma_y',       'Uniaxial yield strength of material', 'Pa')],
    definition="Predicts yielding when the distortion strain energy density equals the uniaxial yield value; generally more accurate than Tresca for ductile metals. For pure shear (sigma1 = -sigma2 = tau), Von Mises predicts yield at tau = sigma_y/sqrt(3) = 0.577*sigma_y, which is 15.5% higher than Tresca's prediction of 0.5*sigma_y. Preferred in engineering design for steel and aluminium alloys.",
    example={
        'problem': 'sigma1=92.4 MPa, sigma2=7.6 MPa. Material sigma_y=250 MPa. Does it yield?',
        'steps': [
            'sigma_VM = sqrt(92.4^2 - 92.4*7.6 + 7.6^2)',
            '         = sqrt(8537.8 - 702.2 + 57.8) = sqrt(7893) = 88.8 MPa',
            '88.8 MPa < 250 MPa  =>  no yielding',
        ],
        'answer': 'sigma_VM = 88.8 MPa < 250 MPa => safe (FoS = 2.82)'})

eq_block(69, 'Tresca Yield Criterion',
    'tau_max = (sigma1 - sigma2)/2 <= sigma_y/2',
    [('tau_max',       'Maximum shear stress', 'Pa'),
     ('sigma1, sigma2', 'Principal stresses (sigma1 >= sigma2)', 'Pa'),
     ('sigma_y',       'Yield strength', 'Pa')],
    note='Tresca is more conservative (lower predicted failure load) than Von Mises.',
    definition="Predicts yielding when the maximum shear stress reaches half the uniaxial yield strength; also called the maximum shear stress criterion. More conservative than Von Mises (predicts failure at a 13.4% lower load in the worst case). Preferred in pressure vessel codes and safety-critical designs. For 3D stress states, evaluate all three principal stress pairs to find the true maximum shear stress.",
    example={
        'problem': 'tau_max=42.4 MPa (from eq. 67). sigma_y=250 MPa. Check Tresca criterion.',
        'steps': [
            'sigma_y/2 = 250/2 = 125 MPa',
            '42.4 MPa < 125 MPa  =>  no yielding',
        ],
        'answer': 'Safe. Tresca FoS = 125/42.4 = 2.95'})

subtopic_hdr('Pressure Vessels')

eq_block(70, 'Thin-Walled Cylinder - Hoop (Circumferential) Stress',
    'sigma_h = P*d / (2*t)',
    [('sigma_h', 'Hoop (circumferential) stress (= 2 x longitudinal)', 'Pa'),
     ('P',       'Internal gauge pressure', 'Pa'),
     ('d',       'Internal diameter', 'm'),
     ('t',       'Wall thickness', 'm')],
    note='Valid when d/t >= 20 (thin-wall assumption).',
    definition="Circumferential (hoop) stress in a thin-walled cylinder; the critical (larger) stress for design. Thin-wall criterion: d/t >= 20 (equivalently t/r_mean <= 0.1), where d is the internal diameter. Below this threshold, the stress varies significantly through the wall and Lame equations (eq. 72) must be used. Hoop stress is always twice the longitudinal stress in a closed-ended cylinder.",
    example={
        'problem': 'Pressure vessel: d=400 mm, t=10 mm, P=2 MPa. Find sigma_h.',
        'steps': ['sigma_h = 2e6*0.4 / (2*0.010) = 800,000/0.020 = 40 MPa'],
        'answer': 'sigma_h = 40 MPa (hoop)'})

eq_block(71, 'Thin-Walled Cylinder - Longitudinal (Axial) Stress',
    'sigma_L = P*d / (4*t)',
    [('sigma_L', 'Longitudinal (axial) stress (sigma_L = sigma_h/2)', 'Pa'),
     ('P',       'Internal gauge pressure', 'Pa'),
     ('d',       'Internal diameter', 'm'),
     ('t',       'Wall thickness', 'm')],
    definition="Axial stress in a closed-ended thin-walled cylinder due to internal pressure acting on the end caps. Always equal to half the hoop stress (eq. 70): sigma_L = sigma_h/2. Applies only to closed-ended cylinders; open-ended cylinders (no end caps) have zero longitudinal pressure-induced stress. Thin-wall criterion: d/t >= 20. For a spherical thin-walled vessel, stress is equal in all directions: sigma = P*d/(4*t).",
    example={
        'problem': 'Same vessel as eq. 70: d=400 mm, t=10 mm, P=2 MPa. Find sigma_L.',
        'steps': ['sigma_L = 2e6*0.4 / (4*0.010) = 800,000/0.040 = 20 MPa'],
        'answer': 'sigma_L = 20 MPa (= sigma_h/2)'})

eq_block(72, 'Thick-Walled Cylinder - Lame Equations',
    'sigma_theta = A + B/r^2     (hoop)     and     sigma_r = A - B/r^2     (radial)',
    [('sigma_theta', 'Hoop (tangential) stress at radius r', 'Pa'),
     ('sigma_r',     'Radial stress at radius r', 'Pa'),
     ('A, B',        'Lame constants from boundary conditions', 'Pa,  Pa*m^2'),
     ('r',           'Radial position', 'm')],
    definition="For thick-walled cylinders (d/t < 20), the thin-wall uniform-stress assumption is invalid; hoop stress varies with r and must be computed from Lame's equations. Constants A and B are determined from two pressure boundary conditions: sigma_r = -P_i at r = r_i and sigma_r = -P_o at r = r_o. Maximum hoop stress always occurs at the inner radius (r = r_i) and governs design.",
    example={
        'problem': 'Cylinder ri=50 mm, ro=100 mm, internal P=60 MPa, external P=0. Find sigma_theta at inner wall.',
        'steps': [
            'BC: sigma_r(ri) = A - B/ri^2 = -60 MPa;  sigma_r(ro) = A - B/ro^2 = 0',
            'From 2nd: A = B/ro^2 = B/0.01;  Sub: B/0.01 - B/0.0025 = -60  =>  -300B = -60',
            'B = 0.2 MPa*m^2;  A = 0.2/0.01 = 20 MPa',
            'sigma_theta(ri) = 20 + 0.2/0.0025 = 20 + 80 = 100 MPa',
        ],
        'answer': 'sigma_theta = 100 MPa at inner wall (maximum hoop stress)'})

subtopic_hdr('Buckling')

eq_block(73, 'Euler Critical Buckling Load',
    'P_cr = pi^2*E*I / Le^2',
    [('P_cr', 'Critical (Euler) buckling load', 'N'),
     ('E',    "Young's modulus", 'Pa'),
     ('I',    'Minimum second moment of area of cross-section', 'm^4'),
     ('Le',   'Effective length (Le = K*L, K depends on end conditions)', 'm')],
    note='End condition K: both-pinned=1.0; one-fixed-one-free=2.0; both-fixed=0.5; fixed-pinned=0.7.',
    definition="Axial load at which a slender column suddenly deflects laterally (elastic bifurcation buckling). Valid only for: (1) slender columns where Le/r_min is above the limit for the material (Le/r > ~120 for structural steel), (2) elastic behaviour (stress below proportional limit), (3) ideal geometry (no initial imperfections). Always use the minimum I (weakest axis) for I in the formula. Divide P_cr by an appropriate factor of safety for design.",
    example={
        'problem': 'Pin-pin column, 50x50 mm square, L=3 m, E=200 GPa. Find P_cr.',
        'steps': [
            'I = 50^4/12 mm^4 = 520,833 mm^4 = 5.208e-7 m^4',
            'Le = K*L = 1.0*3 = 3 m',
            'P_cr = pi^2*200e9*5.208e-7 / 3^2 = 1,027,600/9 = 114,178 N',
        ],
        'answer': 'P_cr = 114 kN'})


# ══════════════════════════════════════════════════════════════════════════════
# ME 105/205  MANUFACTURING TECHNOLOGY
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 105 / 205', 'Manufacturing Technology')
subtopic_hdr('Metal Cutting - Tool Life')

eq_block(74, "Taylor's Tool Life Equation",
    'V * T^n = C',
    [('V', 'Cutting speed', 'm/min'),
     ('T', 'Tool life', 'min'),
     ('n', "Taylor's exponent (HSS=0.1-0.15; carbide=0.2-0.3; ceramic=0.4-0.5)", 'dimensionless'),
     ('C', "Taylor constant (cutting speed for T=1 min)", 'm/min')],
    note='Rearranged: T = (C/V)^(1/n). Log form: log V + n*log T = log C (straight line on log-log plot).',
    definition="Empirical power-law relating cutting speed to tool life for a given tool-workpiece combination. Higher cutting speed always reduces tool life. The exponent n characterises sensitivity to speed: HSS n = 0.1-0.15 (very sensitive), carbide n = 0.2-0.3, ceramic n = 0.4-0.5. Calibrate C and n from two test points on a log-log plot. An extended form also includes feed and depth-of-cut terms.",
    example={
        'problem': 'At V=200 m/min, T=60 min; n=0.25. Find T at V=250 m/min.',
        'steps': [
            'C = V*T^n = 200*60^0.25 = 200*2.783 = 556.6 m/min',
            'T = (C/V)^(1/n) = (556.6/250)^(1/0.25) = (2.226)^4 = 24.5 min',
        ],
        'answer': 'T = 24.5 min at V=250 m/min (41% of original life)'})

subtopic_hdr('Metal Cutting - Cutting Geometry')

eq_block(75, 'Chip Thickness Ratio',
    'rc = t1 / t2',
    [('rc', 'Chip thickness ratio (cutting ratio), 0 < rc < 1', 'dimensionless'),
     ('t1', 'Uncut chip thickness (depth of cut)', 'mm'),
     ('t2', 'Actual (deformed) chip thickness', 'mm')],
    definition="Ratio of nominal uncut chip thickness to the actual measured chip thickness; always less than 1 because chips thicken and shorten due to plastic deformation. A higher rc (thin chip, rc closer to 1) indicates efficient cutting with lower shear strain and cutting force. Determined experimentally by collecting and measuring chips from a cutting test.",
    example={
        'problem': 'Uncut chip t1=0.25 mm, measured chip t2=0.75 mm. Find rc.',
        'steps': ['rc = t1/t2 = 0.25/0.75 = 0.333'],
        'answer': 'rc = 0.333'})

eq_block(76, 'Shear Plane Angle',
    'tan(phi) = rc*cos(alpha) / (1 - rc*sin(alpha))',
    [('phi',   'Shear plane angle', 'deg'),
     ('rc',    'Chip thickness ratio', 'dimensionless'),
     ('alpha', 'Tool rake angle (positive if tilted toward workpiece)', 'deg')],
    definition="Geometric relationship giving the shear plane angle from chip ratio and rake angle in orthogonal cutting. A larger phi means a thinner shear zone, lower cutting forces, and more efficient machining. Derived from the geometry of the chip formation zone assuming a thin shear plane. Compare with Merchant's prediction (eq. 77) to assess whether friction conditions are as expected.",
    example={
        'problem': 'rc=0.333, rake angle alpha=10 deg. Find shear plane angle phi.',
        'steps': [
            'tan(phi) = 0.333*cos(10)/(1 - 0.333*sin(10))',
            '         = 0.333*0.9848 / (1 - 0.333*0.1736)',
            '         = 0.3280 / 0.9422 = 0.3481',
            'phi = arctan(0.3481) = 19.2 deg',
        ],
        'answer': 'phi = 19.2 deg'})

eq_block(77, "Merchant's Minimum Energy Equation",
    'phi = 45 + alpha/2 - lambda/2',
    [('phi',    'Shear plane angle', 'deg'),
     ('alpha',  'Tool rake angle', 'deg'),
     ('lambda', 'Friction angle  lambda = arctan(mu)', 'deg')],
    note="Based on minimum energy principle. Predicts optimal shear angle for given rake and friction.",
    definition="Theoretical prediction of the shear plane angle that minimises total cutting energy (2*phi + lambda - alpha = 90 deg). Based on a sharp single shear-plane model and constant friction angle lambda = arctan(mu). Provides a first estimate benchmark; measured phi values often differ by 5-15 deg due to strain hardening and velocity-dependent friction. A larger positive rake angle alpha increases phi and reduces cutting force.",
    example={
        'problem': 'alpha=10 deg, friction coefficient mu=0.5. Find Merchant prediction for phi.',
        'steps': [
            'lambda = arctan(mu) = arctan(0.5) = 26.6 deg',
            'phi = 45 + 10/2 - 26.6/2 = 45 + 5 - 13.3 = 36.7 deg',
        ],
        'answer': "phi = 36.7 deg (Merchant's prediction)"})

subtopic_hdr('Material Removal - Turning')

eq_block(78, 'Cutting Speed',
    'Vc = pi*D*N / 1000',
    [('Vc', 'Peripheral cutting speed', 'm/min'),
     ('D',  'Workpiece diameter', 'mm'),
     ('N',  'Spindle speed', 'rpm')],
    definition="Peripheral velocity of the workpiece surface relative to the cutting tool in turning; the primary process variable controlling tool life (via eq. 74) and surface finish. Factor 1000 converts mm to m. To find the required spindle speed from a recommended cutting speed: N = 1000*Vc / (pi*D). As the workpiece diameter decreases during turning, N must increase to maintain constant Vc.",
    example={
        'problem': 'Turning 80 mm diameter steel bar at N=800 rpm. Find Vc.',
        'steps': ['Vc = pi*80*800/1000 = 201.1 m/min'],
        'answer': 'Vc = 201 m/min'})

eq_block(79, 'Material Removal Rate - Turning',
    'MRR = pi*D*N*f*d = Vc*f*d',
    [('MRR', 'Material removal rate', 'mm^3/min'),
     ('D',   'Workpiece diameter', 'mm'),
     ('N',   'Spindle speed', 'rpm'),
     ('f',   'Feed per revolution', 'mm/rev'),
     ('d',   'Depth of cut (radial)', 'mm'),
     ('Vc',  'Cutting speed', 'm/min (use mm/min = m/min * 1000)')],
    definition="Volume of material removed per unit time in turning; product of cutting speed, feed, and depth of cut. Used to estimate machining cycle time (volume to remove / MRR) and required machine power (P = u*MRR, eq. 82-83). Keep consistent units: if Vc is in m/min, convert to mm/min (multiply by 1000) before multiplying by f and d in mm.",
    example={
        'problem': 'D=80 mm, N=800 rpm, f=0.25 mm/rev, d=2 mm. Find MRR.',
        'steps': [
            'MRR = Vc*f*d = 201,100 mm/min * 0.25 * 2',
            '    = 100,550 mm^3/min',
        ],
        'answer': 'MRR = 100,550 mm^3/min (= 100.6 cm^3/min)'})

subtopic_hdr('Material Removal - Milling')

eq_block(80, 'Feed Rate - Milling',
    'vf = fz * z * N',
    [('vf', 'Table feed rate', 'mm/min'),
     ('fz', 'Feed per tooth', 'mm/tooth'),
     ('z',  'Number of cutter teeth', 'dimensionless'),
     ('N',  'Spindle speed', 'rpm')],
    definition="Table (workpiece) feed rate in milling: the rate at which the workpiece advances through the rotating cutter. The feed per tooth fz is the fundamental chip-load parameter that controls surface roughness and cutting force per tooth. Higher vf increases productivity but raises cutting force per tooth; if fz exceeds the recommended chip load, the tool may break or deflect excessively.",
    example={
        'problem': 'fz=0.08 mm/tooth, z=6 teeth, N=400 rpm. Find table feed rate.',
        'steps': ['vf = 0.08 * 6 * 400 = 192 mm/min'],
        'answer': 'vf = 192 mm/min'})

eq_block(81, 'Material Removal Rate - Face Milling',
    'MRR = vf * d * W',
    [('MRR', 'Material removal rate', 'mm^3/min'),
     ('vf',  'Feed rate', 'mm/min'),
     ('d',   'Axial depth of cut', 'mm'),
     ('W',   'Width of cut (radial engagement)', 'mm')],
    definition="Volume removed per unit time in face milling; product of table feed rate, axial depth, and radial width of engagement. Axial depth d is the step-down per pass; radial width W is the step-over (W <= cutter diameter). Used to estimate cycle time and check whether required power (P = u*MRR, eq. 82-83) is within the machine's capacity.",
    example={
        'problem': 'vf=192 mm/min, axial depth d=3 mm, width W=50 mm. Find MRR.',
        'steps': ['MRR = 192 * 3 * 50 = 28,800 mm^3/min'],
        'answer': 'MRR = 28,800 mm^3/min'})

subtopic_hdr('Machining Power and Energy')

eq_block(82, 'Cutting Power',
    'P = Fc * Vc',
    [('P',  'Cutting power', 'W'),
     ('Fc', 'Principal (tangential) cutting force', 'N'),
     ('Vc', 'Cutting speed', 'm/s')],
    definition="Power consumed at the tool-workpiece interface; product of the principal (tangential) cutting force and cutting velocity. Vc must be in m/s (not m/min) to obtain watts. Divide by machine drivetrain efficiency (typically 0.70-0.85) to find the required motor power. Fc can be estimated as: Fc = u * f * d, where u is the specific cutting energy (eq. 83).",
    example={
        'problem': 'Fc=800 N, Vc=201 m/min = 3.35 m/s. Find cutting power.',
        'steps': ['P = 800 * 3.35 = 2,680 W'],
        'answer': 'P = 2.68 kW'})

eq_block(83, 'Specific Cutting Energy (Unit Power)',
    'u = P / MRR = Fc / (f*d)',
    [('u',   'Specific cutting energy', 'J/mm^3'),
     ('P',   'Cutting power', 'W'),
     ('MRR', 'Material removal rate', 'mm^3/s'),
     ('Fc',  'Cutting force', 'N'),
     ('f',   'Feed', 'mm/rev'),
     ('d',   'Depth of cut', 'mm')],
    definition="Energy required to remove one cubic millimetre of material; a process-material constant at given cutting conditions. Tabulated reference values: steel ~2-4 J/mm^3, cast iron ~1-3 J/mm^3, aluminium ~0.5-1.5 J/mm^3. Used to predict cutting force from chip cross-section (Fc = u*f*d) and to size machine tools. u decreases slightly at larger chip cross-sections (size effect) — adjust for very large feeds or depths.",
    example={
        'problem': 'P=2680 W, MRR=100,550 mm^3/min = 1676 mm^3/s. Find specific cutting energy.',
        'steps': ['u = P/MRR = 2680/1676 = 1.60 J/mm^3'],
        'answer': 'u = 1.60 J/mm^3 (= 1.60 GJ/m^3)'})

subtopic_hdr('Rolling and Forming')

eq_block(84, 'Rolling - Draft',
    'd = t0 - tf',
    [('d',  'Draft (reduction in thickness)', 'mm'),
     ('t0', 'Initial sheet/slab thickness', 'mm'),
     ('tf', 'Final sheet/slab thickness', 'mm')],
    definition="Thickness reduction achieved in a single rolling pass. Maximum draft per pass is limited by the bite angle: d_max = mu^2 * R (mu = friction coefficient, R = roll radius). Hot rolling typically achieves 10–30% reduction per pass; cold rolling 5–15%. The rolling force and torque scale with draft — excessive draft causes edge cracking or roll slipping. The pass schedule (sequence of drafts) is designed to achieve the target thickness while maintaining surface quality.",
    example={
        'problem': 'Steel slab reduced from 25 mm to 18 mm in one pass. Find draft and % reduction.',
        'steps': [
            'd = 25 - 18 = 7 mm',
            '% reduction = 7/25 * 100 = 28%',
        ],
        'answer': 'd = 7 mm (28% reduction per pass)'})

eq_block(85, 'Deep Drawing - Approximate Blank Diameter',
    'Db = sqrt(Dp^2 + 4*Dp*h)',
    [('Db', 'Blank diameter required', 'mm'),
     ('Dp', 'Punch (cup) diameter', 'mm'),
     ('h',  'Cup drawing depth', 'mm')],
    note='Approximate formula based on surface area conservation.',
    definition="Determines the circular blank size needed to draw a cylindrical cup of given diameter and height, assuming volume/surface area conservation and negligible wall thinning. Drawing Ratio DR = Db/Dp: typical limit DR ≤ 2.0 for a single draw without annealing (higher DR risks wrinkling or tearing). For DR > 2, use multiple draws with intermediate anneals. Blank-holder force is critical — too low causes wrinkling, too high causes tearing. The formula gives the minimum blank; add a trimming allowance of 5–10 mm in practice.",
    example={
        'problem': 'Cup Dp=80 mm, drawing depth h=50 mm. Find required blank diameter.',
        'steps': [
            'Db = sqrt(80^2 + 4*80*50) = sqrt(6400 + 16000) = sqrt(22400) = 149.7 mm',
        ],
        'answer': 'Db = 150 mm'})


# ══════════════════════════════════════════════════════════════════════════════
# ME 106/206  THERMODYNAMICS AND HEAT TRANSFER
# ══════════════════════════════════════════════════════════════════════════════
topic_hdr('ME 106 / 206', 'Thermodynamics and Heat Transfer')
subtopic_hdr('Ideal Gas Laws')

eq_block(86, 'Ideal Gas Law - Molar Form',
    'P*V = n*Rbar*T',
    [('P',    'Absolute pressure', 'Pa'),
     ('V',    'Volume', 'm^3'),
     ('n',    'Amount of substance', 'mol'),
     ('Rbar', 'Universal gas constant = 8.314', 'J/(mol*K)'),
     ('T',    'Absolute temperature', 'K')],
    definition="Equation of state for an ideal (perfect) gas — assumes molecules occupy negligible volume and exert no intermolecular forces. Valid when the gas density is low and the temperature is well above the critical temperature (e.g., air, nitrogen, oxygen at typical engineering conditions). T must be in Kelvin (absolute). Not accurate near the critical point, at high pressures, or for steam — use steam tables or van der Waals equation instead. The molar form is preferred in chemistry; the mass form (eq 87) is preferred in engineering thermodynamics.",
    example={
        'problem': '2 mol of ideal gas at T=300 K, P=200 kPa. Find volume.',
        'steps': ['V = n*Rbar*T/P = 2*8.314*300 / 200,000 = 4988.4/200,000 = 0.0249 m^3'],
        'answer': 'V = 0.0249 m^3 = 24.9 L'})

eq_block(87, 'Ideal Gas Law - Mass Form',
    'P*V = m*R*T     or     P = rho*R*T',
    [('P',   'Absolute pressure', 'Pa'),
     ('V',   'Volume', 'm^3'),
     ('m',   'Mass of gas', 'kg'),
     ('R',   'Specific gas constant R=Rbar/M (air=287)', 'J/(kg*K)'),
     ('T',   'Absolute temperature', 'K'),
     ('rho', 'Gas density', 'kg/m^3')],
    definition="Mass-based form of the ideal gas law; the specific gas constant R = Rbar/M where M is the molar mass (kg/mol). Key values: air R = 287 J/(kg·K), oxygen R = 260, hydrogen R = 4124, CO2 R = 189 J/(kg·K). The density form P = rho*R*T is directly useful for compressible flow and HVAC calculations. Always convert temperature to Kelvin (T_K = T_degC + 273.15) before applying. For a closed process: P1*V1/T1 = P2*V2/T2 = m*R (combined gas law).",
    example={
        'problem': '1 kg of air at P=100 kPa, T=20 degC=293 K. Find volume.',
        'steps': ['V = m*R*T/P = 1*287*293/100,000 = 84,091/100,000 = 0.841 m^3'],
        'answer': 'V = 0.841 m^3'})

eq_block(88, 'Specific Heat Relationship - Ideal Gas (Mayer Relation)',
    'cp - cv = R',
    [('cp', 'Specific heat at constant pressure', 'J/(kg*K)'),
     ('cv', 'Specific heat at constant volume', 'J/(kg*K)'),
     ('R',  'Specific gas constant', 'J/(kg*K)')],
    definition="Mayer's relation: the difference between cp and cv equals the specific gas constant R for any ideal gas. The difference arises because at constant pressure, the gas must also do expansion work (P*dV) against the surroundings — so extra energy is needed compared to constant-volume heating. Applies only to ideal gases; for real gases or liquids, cp - cv must be computed from thermodynamic relations involving the compressibility. Key values for air: cp = 1005, cv = 718, R = 287 J/(kg·K); for steam as ideal gas: cp ≈ 1872, R = 462 J/(kg·K).",
    example={
        'problem': 'Air: cp=1005 J/(kg*K), R=287 J/(kg*K). Find cv.',
        'steps': ['cv = cp - R = 1005 - 287 = 718 J/(kg*K)'],
        'answer': 'cv = 718 J/(kg*K)'})

eq_block(89, 'Heat Capacity Ratio',
    'gamma = cp / cv',
    [('gamma', 'Heat capacity ratio (monatomic~1.67; diatomic/air~1.40)', 'dimensionless'),
     ('cp',    'Specific heat at constant pressure', 'J/(kg*K)'),
     ('cv',    'Specific heat at constant volume', 'J/(kg*K)')],
    definition="Ratio of specific heats; governs isentropic behaviour, speed of sound (c = sqrt(gamma*R*T)), and compressible flow. Key values: monatomic gases (He, Ar) gamma = 5/3 ≈ 1.67; diatomic gases and air gamma = 7/5 = 1.40; triatomic gases (CO2, H2O vapour) gamma ≈ 1.2–1.3. Higher gamma → steeper pressure-temperature rise in isentropic compression (harder to compress adiabatically). Gamma decreases slightly with temperature in real gases as vibrational modes become active. The isentropic relations (eq 95) and Carnot efficiency analysis both depend on gamma.",
    example={
        'problem': 'Air: cp=1005, cv=718 J/(kg*K). Find gamma.',
        'steps': ['gamma = cp/cv = 1005/718 = 1.40'],
        'answer': 'gamma = 1.40 (diatomic gas, as expected for air)'})

subtopic_hdr('First Law of Thermodynamics')

eq_block(90, 'First Law - Closed System',
    'DeltaU = Q - W',
    [('DeltaU', 'Change in internal energy of the system', 'J'),
     ('Q',      'Heat transferred into the system (Q>0 in; Q<0 out)', 'J'),
     ('W',      'Work done by the system (W>0 out; W<0 in)', 'J')],
    definition="Energy conservation for a fixed mass (no flow across the boundary). Sign convention: Q positive when heat enters the system; W positive when the system does work on surroundings (expansion). For ideal gas: DeltaU = m*cv*DeltaT (internal energy depends on temperature only). For a complete thermodynamic cycle: DeltaU = 0 → Q_net = W_net (net heat in = net work out). Boundary work: W = integral(P*dV). Apply this form to pistons, sealed vessels, and batch processes.",
    example={
        'problem': 'A gas absorbs Q=500 J heat and does W=200 J work. Find DeltaU.',
        'steps': ['DeltaU = Q - W = 500 - 200 = 300 J'],
        'answer': 'DeltaU = 300 J (internal energy increases)'})

eq_block(91, 'Steady Flow Energy Equation (SFEE) - Open System',
    'Qdot - Wdot = mdot * [(h2-h1) + (1/2)*(V2^2-V1^2) + g*(z2-z1)]',
    [('Qdot',    'Rate of heat transfer into system', 'W'),
     ('Wdot',    'Rate of work output (shaft work)', 'W'),
     ('mdot',    'Mass flow rate', 'kg/s'),
     ('h1, h2',  'Specific enthalpy at inlet and outlet', 'J/kg'),
     ('V1, V2',  'Velocities at inlet and outlet', 'm/s'),
     ('z1, z2',  'Elevations at inlet and outlet', 'm'),
     ('g',       'Gravitational acceleration', 'm/s^2')],
    definition="Energy balance for steady-state open systems with continuous mass flow (turbines, compressors, pumps, heat exchangers, nozzles). Enthalpy h = u + P*v naturally accounts for flow work, so it replaces internal energy in open-system analysis. Common simplifications: turbine/compressor — KE and PE negligible, Qdot ≈ 0 → Wdot = mdot*(h1-h2); nozzle — Wdot = 0, Qdot ≈ 0, z same → V2 = sqrt(V1^2 + 2*(h1-h2)); heat exchanger — Wdot = 0, KE/PE negligible → Qdot = mdot*(h2-h1). Use steam tables for h when working with steam.",
    example={
        'problem': 'Steam turbine: mdot=5 kg/s, h1=3200 kJ/kg, h2=2400 kJ/kg, Qdot=0, KE/PE negligible. Find Wdot.',
        'steps': [
            '0 - Wdot = 5*(2400e3 - 3200e3) = 5*(-800e3) = -4,000,000 W',
            'Wdot = +4,000,000 W',
        ],
        'answer': 'Wdot = 4.0 MW (turbine power output)'})

subtopic_hdr('Second Law and Cycles')

eq_block(92, 'Carnot Efficiency (Maximum Theoretical Efficiency)',
    'eta_c = 1 - TL / TH',
    [('eta_c', 'Carnot (maximum) thermal efficiency', 'dimensionless'),
     ('TL',    'Temperature of cold reservoir (condenser)', 'K (must be absolute)'),
     ('TH',    'Temperature of hot reservoir (boiler)', 'K')],
    note='eta_c = W_net/Q_H = 1 - Q_C/Q_H. Minimum fuel consumption implies Carnot efficiency.',
    definition="Theoretical upper bound on the thermal efficiency of any heat engine operating between two fixed temperature reservoirs; derived from the Second Law. T must be in Kelvin — using Celsius gives a wrong (and usually optimistic) answer. Real cycles are always less efficient due to irreversibilities (friction, heat transfer across finite temperature differences). Corollaries: Carnot COP (refrigerator) = TL/(TH-TL); Carnot COP (heat pump) = TH/(TH-TL). To improve efficiency: raise TH (higher boiler pressure/temperature) or lower TL (better condenser cooling) — raising TH is usually more impactful.",
    example={
        'problem': 'Steam cycle: TH=560 degC=833 K, TL=40 degC=313 K. Find Carnot efficiency.',
        'steps': ['eta_c = 1 - TL/TH = 1 - 313/833 = 1 - 0.376 = 0.624'],
        'answer': 'eta_c = 62.4%'})

eq_block(93, 'Entropy Change (Reversible Process)',
    'dS = dQ_rev / T     =>     DeltaS = integral(dQ/T)_rev',
    [('dS',      'Infinitesimal entropy change', 'J/K'),
     ('dQ_rev',  'Infinitesimal reversible heat transfer', 'J'),
     ('T',       'Absolute temperature at boundary', 'K'),
     ('DeltaS',  'Total entropy change', 'J/K')],
    definition="Entropy is a state function measuring the thermodynamic disorder or unavailability of energy. This definition applies only for reversible heat transfers; for irreversible processes the actual entropy change is greater: dS_actual > dQ_actual/T. For ideal gas processes, entropy changes can be computed from: DeltaS = m*cp*ln(T2/T1) - m*R*ln(P2/P1) or DeltaS = m*cv*ln(T2/T1) + m*R*ln(V2/V1). Isentropic process: DeltaS = 0 (reversible and adiabatic). Phase change (isothermal): DeltaS = Q/T = m*h_fg/T.",
    example={
        'problem': '2 kg of water vaporised reversibly at 100 degC (373 K), h_fg=2257 kJ/kg. Find DeltaS.',
        'steps': [
            'Q = 2 * 2257e3 = 4,514,000 J',
            'DeltaS = Q/T = 4,514,000/373 = 12,104 J/K',
        ],
        'answer': 'DeltaS = 12.1 kJ/K'})

eq_block(94, 'Clausius Inequality',
    'closed_integral( dQ / T ) <= 0',
    [('closed_integral dQ/T', 'Cyclic integral of heat/temperature over a complete cycle', 'J/K')],
    note='= 0 for reversible cycle; < 0 for irreversible (real) cycle. Expression of 2nd Law.',
    definition="A mathematical statement of the Second Law: the cyclic integral of dQ/T over a complete thermodynamic cycle is always less than or equal to zero. Equality holds only for a fully reversible (Carnot) cycle; any real cycle gives a negative result because irreversibilities generate entropy. A positive result would violate the Second Law and is physically impossible. Use this to verify whether a proposed heat engine cycle is thermodynamically feasible. It is the foundation from which the entropy state function is formally derived.",
    example={
        'problem': 'Heat engine: QH=500 J at 800 K, rejects QL=350 J at 300 K. Evaluate cyclic integral.',
        'steps': [
            'sum(dQ/T) = QH/TH - QL/TL = 500/800 - 350/300 = 0.625 - 1.167 = -0.542 J/K',
        ],
        'answer': 'Cyclic integral = -0.54 J/K < 0 => irreversible (real) engine'})

eq_block(95, 'Isentropic Process Relations - Ideal Gas',
    'T2/T1 = (P2/P1)^((gamma-1)/gamma) = (V1/V2)^(gamma-1)',
    [('T1, T2', 'Temperatures before and after isentropic process', 'K'),
     ('P1, P2', 'Pressures before and after', 'Pa'),
     ('V1, V2', 'Specific volumes before and after', 'm^3/kg'),
     ('gamma',  'Heat capacity ratio cp/cv', 'dimensionless')],
    note='Valid for reversible adiabatic (isentropic) processes only.',
    definition="Relates temperature, pressure, and volume for a reversible adiabatic (isentropic) process with an ideal gas. Applies to: ideal compressor/turbine stages, choked nozzle flow, and acoustic wave propagation. Isentropic efficiency compares real vs ideal: eta_s(turbine) = actual work / isentropic work; eta_s(compressor) = isentropic work / actual work. Not valid for polytropic processes (replace gamma with the polytropic index n). For steam, read isentropic temperatures from the h-s (Mollier) diagram rather than using this formula.",
    example={
        'problem': 'Air compressed isentropically: P1=100 kPa, T1=300 K to P2=600 kPa. Find T2.',
        'steps': [
            'T2 = T1 * (P2/P1)^((gamma-1)/gamma) = 300 * (6)^(0.4/1.4)',
            '   = 300 * 6^0.2857 = 300 * 1.669 = 500.7 K',
        ],
        'answer': 'T2 = 501 K = 228 degC'})

subtopic_hdr('Heat Transfer - Conduction')

eq_block(96, "Fourier's Law of Heat Conduction",
    'q = -k*A*(dT/dx)     or     q_flux = -k*(dT/dx)',
    [('q',       'Rate of heat transfer', 'W'),
     ('q_flux',  'Heat flux', 'W/m^2'),
     ('k',       'Thermal conductivity of material', 'W/(m*K)'),
     ('A',       'Cross-sectional area perpendicular to heat flow', 'm^2'),
     ('dT/dx',   'Temperature gradient in direction of heat flow', 'K/m')],
    note='Negative sign: heat flows in the direction of decreasing temperature.',
    definition="Fundamental law of heat conduction; the negative sign ensures heat flows from hot to cold (positive flux in the direction of decreasing T). Thermal conductivity k values: copper ~400, aluminium ~200, steel ~50, glass ~1, water ~0.6, insulation ~0.04, air ~0.026 W/(m·K). Assumes steady-state, one-dimensional, homogeneous isotropic material — for multilayer walls, use thermal resistance networks (eq 97). For transient conduction, the heat diffusion equation applies: d²T/dx² = (rho*cp/k)*(dT/dt), governed by thermal diffusivity alpha = k/(rho*cp).",
    example={
        'problem': 'Copper wall (k=400 W/(m*K)), A=0.5 m^2, thickness L=10 mm, DeltaT=50 degC. Find q.',
        'steps': ['q = k*A*(DeltaT/L) = 400 * 0.5 * (50/0.010) = 400*0.5*5000 = 1,000,000 W'],
        'answer': 'q = 1.0 MW'})

eq_block(97, 'Thermal Resistance - Conduction (Flat Slab)',
    'R_cond = L / (k*A)',
    [('R_cond', 'Conductive thermal resistance', 'K/W'),
     ('L',      'Slab thickness', 'm'),
     ('k',      'Thermal conductivity', 'W/(m*K)'),
     ('A',      'Area', 'm^2')],
    note='Analogous to Ohm law: Q_dot = DeltaT / R_cond',
    definition="Thermal resistance concept (Ohm's law analogy): Q = DeltaT/R. Layers in series (e.g., wall + insulation): R_total = R1 + R2 + ... Parallel heat paths: 1/R_total = 1/R1 + 1/R2 + ... For a cylindrical wall (pipe insulation): R_cyl = ln(r_o/r_i)/(2*pi*k*L). Contact resistance R_contact must be added at interfaces between surfaces — significant in electronics cooling and thin-film applications. The concept extends to convection (1/(h*A)) and radiation resistances, enabling complete thermal circuit analysis of composite walls, fins, and heat exchangers.",
    example={
        'problem': 'Copper wall from eq.96: L=10 mm, k=400, A=0.5 m^2. Find R_cond.',
        'steps': ['R_cond = L/(k*A) = 0.010 / (400*0.5) = 0.010/200 = 5.0e-5 K/W'],
        'answer': 'R_cond = 5.0e-5 K/W (very low - copper is excellent conductor)'})

subtopic_hdr('Heat Transfer - Convection')

eq_block(98, "Newton's Law of Cooling (Convection)",
    'q = h*A*(Ts - T_inf)',
    [('q',     'Convective heat transfer rate', 'W'),
     ('h',     'Convective heat transfer coefficient', 'W/(m^2*K)'),
     ('A',     'Surface area in contact with fluid', 'm^2'),
     ('Ts',    'Surface temperature', 'K (or degC)'),
     ('T_inf', 'Free-stream (ambient) fluid temperature', 'K (or degC)')],
    definition="The convective heat transfer coefficient h is not a material property — it depends on the fluid, geometry, flow regime, and temperature difference. Typical ranges: free convection in air h = 5–25 W/(m²·K); forced convection in air h = 25–250; forced convection in water h = 1,000–15,000; boiling/condensation h = 2,500–100,000 W/(m²·K). Temperature difference (Ts - T_inf) can be in degC or K since it is a difference. h is determined experimentally or from Nusselt number correlations (eqs 101–103): h = Nu*k_fluid/L_char.",
    example={
        'problem': 'Flat plate A=0.2 m^2, Ts=80 degC, air at T_inf=25 degC, h=35 W/(m^2*K). Find q.',
        'steps': ['q = 35 * 0.2 * (80 - 25) = 35 * 0.2 * 55 = 385 W'],
        'answer': 'q = 385 W'})

eq_block(99, 'Thermal Resistance - Convection',
    'R_conv = 1 / (h*A)',
    [('R_conv', 'Convective thermal resistance', 'K/W'),
     ('h',      'Convective heat transfer coefficient', 'W/(m^2*K)'),
     ('A',      'Surface area', 'm^2')],
    definition="Convective thermal resistance in the same Ohm's law analogy as R_cond (eq 97). Combine with R_cond for composite wall analysis: R_total = R_conv1 + R_cond + R_conv2 → q = (T_fluid1 - T_fluid2)/R_total. The dominant resistance determines the controlling mechanism: if R_conv >> R_cond, improving insulation has little benefit — focus on increasing h (e.g., add fins, increase flow velocity). Fins increase A, thereby reducing R_conv. This concept is central to thermal management of electronics and HVAC design.",
    example={
        'problem': 'Same plate: h=35 W/(m^2*K), A=0.2 m^2. Find R_conv.',
        'steps': ['R_conv = 1/(h*A) = 1/(35*0.2) = 1/7 = 0.143 K/W'],
        'answer': 'R_conv = 0.143 K/W'})

eq_block(100, 'Overall Heat Transfer Coefficient (Plane Wall)',
    '1/(U*A) = 1/(h1*A) + L/(k*A) + 1/(h2*A) = R_total',
    [('U',       'Overall heat transfer coefficient', 'W/(m^2*K)'),
     ('A',       'Heat transfer area', 'm^2'),
     ('h1, h2',  'Convection coefficients on each side of wall', 'W/(m^2*K)'),
     ('L',       'Wall thickness', 'm'),
     ('k',       'Thermal conductivity of wall', 'W/(m*K)'),
     ('R_total', 'Total thermal resistance', 'K/W')],
    definition="U is the combined measure of heat transfer through a wall including both surface convection layers and wall conduction. Used in heat exchanger design: Q = U*A*LMTD (eq 105). Fouling factors R_f are added in series for dirty surfaces: 1/U_fouled = 1/U_clean + R_f1 + R_f2. For multi-layer walls: add additional L/(k*A) terms for each layer. Typical U values: double-glazed window ~2–3, building wall ~0.3–1, water-to-water heat exchanger ~1000–2500 W/(m²·K). The smallest individual h dominates U — it is the bottleneck.",
    example={
        'problem': 'Wall: h1=50, L=20 mm, k=1.0, h2=20 W/(m^2*K), A=1 m^2. Find U.',
        'steps': [
            '1/U = 1/50 + 0.020/1.0 + 1/20 = 0.020 + 0.020 + 0.050 = 0.090 m^2*K/W',
            'U = 1/0.090 = 11.1 W/(m^2*K)',
        ],
        'answer': 'U = 11.1 W/(m^2*K)'})

eq_block(101, 'Nusselt Number',
    'Nu = h*L / k',
    [('Nu', 'Nusselt number (ratio of convective to conductive heat transfer)', 'dimensionless'),
     ('h',  'Convective heat transfer coefficient', 'W/(m^2*K)'),
     ('L',  'Characteristic length', 'm'),
     ('k',  'Fluid thermal conductivity', 'W/(m*K)')],
    definition="Dimensionless heat transfer coefficient; Nu = 1 represents pure conduction through a stagnant fluid, higher Nu means convection is increasingly dominant over conduction. Used to find h from empirical correlations (Nu = f(Re, Pr)): h = Nu*k_fluid/L_char. Characteristic length L: diameter D for pipes; length from leading edge for flat plates; diameter for spheres and cylinders in cross-flow. Laminar pipe flow: Nu = 3.66 (constant wall T) or 4.36 (constant heat flux). Turbulent pipe flow: Dittus-Boelter (eq 103). Derive h from Nu after computing Re and Pr.",
    example={
        'problem': 'Pipe D=20 mm, h=2000 W/(m^2*K). Air k_fluid=0.026 W/(m*K). Find Nu.',
        'steps': ['Nu = h*L/k = 2000*0.020/0.026 = 40/0.026 = 1538'],
        'answer': 'Nu = 1538 (turbulent forced convection in pipe)'})

eq_block(102, 'Prandtl Number',
    'Pr = mu*cp / k = nu / alpha_th',
    [('Pr',       'Prandtl number (ratio of momentum to thermal diffusivity)', 'dimensionless'),
     ('mu',       'Dynamic viscosity of fluid', 'Pa*s'),
     ('cp',       'Specific heat of fluid at constant pressure', 'J/(kg*K)'),
     ('k',        'Fluid thermal conductivity', 'W/(m*K)'),
     ('nu',       'Kinematic viscosity', 'm^2/s'),
     ('alpha_th', 'Thermal diffusivity alpha=k/(rho*cp)', 'm^2/s')],
    definition="Pr governs the relative thickness of the velocity vs thermal boundary layers. Pr << 1 (liquid metals: Pr ~ 0.003–0.03): thermal boundary layer is much thicker than the velocity boundary layer — heat diffuses rapidly. Pr ~ 1 (air: Pr ~ 0.71; gases): velocity and thermal boundary layers are of similar thickness. Pr >> 1 (viscous oils: Pr ~ 100–10,000): momentum diffuses faster than heat — thick velocity, thin thermal boundary layer. Prandtl number appears in all forced convection Nu correlations (eqs 101, 103) and determines which fluid properties dominate the heat transfer.",
    example={
        'problem': 'Engine oil: mu=0.08 Pa*s, cp=1900 J/(kg*K), k=0.14 W/(m*K). Find Pr.',
        'steps': ['Pr = mu*cp/k = 0.08*1900/0.14 = 152/0.14 = 1086'],
        'answer': 'Pr = 1086 (highly viscous - thick thermal boundary layer)'})

eq_block(103, 'Dittus-Boelter Correlation (Turbulent Flow in Pipes)',
    'Nu = 0.023 * Re^0.8 * Pr^n',
    [('Nu', 'Nusselt number', 'dimensionless'),
     ('Re', 'Reynolds number', 'dimensionless'),
     ('Pr', 'Prandtl number', 'dimensionless'),
     ('n',  'n=0.4 for fluid being heated;  n=0.3 for fluid being cooled', 'dimensionless')],
    note='Valid for: Re > 10,000; 0.6 < Pr < 160; L/D > 10 (thermally/hydrodynamically developed).',
    definition="Most widely used correlation for turbulent forced convection in smooth circular pipes; derived by Dittus and Boelter (1930). Validity criteria must be met: fully turbulent (Re > 10,000), most common fluids (0.6 < Pr < 160), and hydrodynamically/thermally developed flow (L/D > 10 from the entrance). Use n = 0.4 when the wall is hotter than the fluid (heating); n = 0.3 when the wall is cooler (cooling). For transitional flow (2300 < Re < 10,000), use the Gnielinski correlation: Nu = (f/8)*(Re-1000)*Pr / [1 + 12.7*(f/8)^0.5*(Pr^(2/3)-1)]. Evaluate all fluid properties at the bulk mean temperature.",
    example={
        'problem': 'Water in pipe: Re=50,000, Pr=7.0 (fluid being heated). Find Nu.',
        'steps': [
            'Nu = 0.023 * Re^0.8 * Pr^0.4',
            '   = 0.023 * (50000)^0.8 * 7.0^0.4',
            '   = 0.023 * 4521 * 2.055 = 213.4',
        ],
        'answer': 'Nu = 213'})

eq_block(104, 'Fin Efficiency',
    'eta_f = tanh(m*L) / (m*L)     where  m = sqrt(h*P / (k*Ac))',
    [('eta_f', 'Fin efficiency (actual / maximum possible heat transfer)', 'dimensionless'),
     ('m',     'Fin parameter', 'm^-1'),
     ('L',     'Fin length', 'm'),
     ('h',     'Convective heat transfer coefficient at fin surface', 'W/(m^2*K)'),
     ('P',     'Fin perimeter (cross-sectional)', 'm'),
     ('k',     'Fin material thermal conductivity', 'W/(m*K)'),
     ('Ac',    'Fin cross-sectional area', 'm^2')],
    note='Valid for insulated fin tip. Uniform cross-section.',
    definition="Fin efficiency compares actual fin heat transfer to the ideal case where the entire fin is at the base temperature. eta_f = 1 is ideal; practical well-designed fins achieve eta_f ~ 0.7–0.95. Total fin heat transfer: q_fin = eta_f * h * A_fin * (Tbase - T_inf). Even though eta_f < 1, fins are still beneficial because A_fin >> A_base. Design guideline: increasing fin conductivity (use Al or Cu) or decreasing fin length/height both raise eta_f. The fin parameter m = sqrt(hP/kAc) — a short, wide, high-k fin gives small m*L and high efficiency. Applies to rectangular fins; different formulae for triangular, annular, and pin fins.",
    example={
        'problem': 'Aluminium fin (k=200) rectangular 2mm wide x 20mm deep, L=40mm, h=80 W/(m^2*K). Find eta_f.',
        'steps': [
            'Ac = 0.002*0.020 = 4.0e-5 m^2;  P = 2*(0.002+0.020) = 0.044 m',
            'm = sqrt(80*0.044/(200*4.0e-5)) = sqrt(3.52/0.008) = sqrt(440) = 20.98 m^-1',
            'm*L = 20.98*0.040 = 0.839',
            'eta_f = tanh(0.839)/0.839 = 0.685/0.839 = 0.817',
        ],
        'answer': 'eta_f = 81.7%'})

eq_block(105, 'Log Mean Temperature Difference (LMTD) - Heat Exchangers',
    'LMTD = (DeltaT1 - DeltaT2) / ln(DeltaT1/DeltaT2)',
    [('LMTD',   'Log mean temperature difference', 'K (or degC)'),
     ('DeltaT1', 'Temperature difference at one end of heat exchanger', 'K'),
     ('DeltaT2', 'Temperature difference at other end of heat exchanger', 'K')],
    note='Used in Q_dot = U*A*LMTD. For counter-flow HX use counter-flow temperature pairings.',
    definition="The effective mean temperature driving force for heat transfer in a heat exchanger. Counter-flow arrangement always gives a higher LMTD than parallel-flow for the same terminal temperatures — counter-flow is always more thermally efficient. When DeltaT1 = DeltaT2 (uniform temperature difference), LMTD = DeltaT (apply L'Hopital's rule). For multi-pass shell-and-tube or cross-flow heat exchangers, apply a correction factor F: Q = U*A*F*LMTD (F < 1 from charts, depends on geometry and temperature ratio). Full design: Q = U*A*LMTD = mdot*cp*DeltaT_fluid (energy balance) — two equations for two unknowns (U and A, or the two outlet temperatures).",
    example={
        'problem': 'Counter-flow HX: hot fluid 90->50 degC, cold fluid 20->70 degC. Find LMTD.',
        'steps': [
            'DeltaT1 = 90-70 = 20 degC (hot inlet vs cold outlet)',
            'DeltaT2 = 50-20 = 30 degC (hot outlet vs cold inlet)',
            'LMTD = (20-30)/ln(20/30) = -10/(-0.405) = 24.7 degC',
        ],
        'answer': 'LMTD = 24.7 degC'})

subtopic_hdr('Heat Transfer - Radiation')

eq_block(106, 'Stefan-Boltzmann Law - Net Radiation Heat Transfer',
    'q = eps*sigma*A*(Ts^4 - T_surr^4)',
    [('q',      'Net radiation heat transfer rate', 'W'),
     ('eps',    'Emissivity (0 <= eps <= 1; blackbody eps=1)', 'dimensionless'),
     ('sigma',  'Stefan-Boltzmann constant = 5.670e-8', 'W/(m^2*K^4)'),
     ('A',      'Surface area', 'm^2'),
     ('Ts',     'Absolute surface temperature', 'K'),
     ('T_surr', 'Absolute temperature of surroundings', 'K')],
    definition="Radiation heat transfer requires no medium and follows T^4 dependence — it dominates over conduction and convection at high temperatures (furnaces, combustion) but is significant even at room temperature. T must be in Kelvin; using degC causes large errors due to the 4th-power relationship. Emissivity ranges: polished metals 0.02–0.10 (low emitters); oxidised metals 0.6–0.8; non-metals and paints 0.8–0.95. For two large parallel grey surfaces: q = sigma*(T1^4-T2^4) / (1/eps1 + 1/eps2 - 1). Radiation shields (low-eps foils) dramatically reduce radiative losses by inserting additional surface resistances in series.",
    example={
        'problem': 'Blackbody (eps=1) at 600 degC=873 K, A=0.1 m^2, surroundings at 20 degC=293 K. Find q.',
        'steps': [
            'q = 1 * 5.670e-8 * 0.1 * (873^4 - 293^4)',
            '  = 5.670e-9 * (5.819e11 - 7.370e9)',
            '  = 5.670e-9 * 5.745e11 = 3258 W',
        ],
        'answer': 'q = 3.26 kW'})

eq_block(107, "Wien's Displacement Law",
    'lambda_max * T = 2.898e-3  m*K',
    [('lambda_max',   'Wavelength of peak radiation intensity', 'm'),
     ('T',            'Absolute temperature of blackbody', 'K'),
     ('2.898e-3 m*K', "Wien's displacement constant", 'm*K')],
    definition="As temperature increases, the peak emission wavelength shifts to shorter (higher-energy) wavelengths — hence hot objects glow red, then orange, then white. At 300 K (room temperature): lambda_max ≈ 9.7 μm (long-wave infrared — detected by thermal cameras). At ~900 K: lambda_max ≈ 3.2 μm (near infrared, steel beginning to glow dull red, visible to the naked eye). At ~6000 K (Sun): lambda_max ≈ 0.5 μm (green-yellow visible light). Derived by differentiating the Planck blackbody distribution with respect to wavelength and setting the derivative to zero. Used to estimate star surface temperatures from spectral peak measurements.",
    example={
        'problem': 'The Sun has lambda_max = 500 nm. Find its surface temperature.',
        'steps': ['T = 2.898e-3 / lambda_max = 2.898e-3 / 500e-9 = 5796 K'],
        'answer': 'T = 5796 K (~5800 K, consistent with known solar surface temperature)'})

eq_block(108, 'View Factor Reciprocity Relation',
    'A1*F12 = A2*F21',
    [('A1, A2', 'Areas of surfaces 1 and 2', 'm^2'),
     ('F12',    'View factor from surface 1 to 2 (fraction of radiation from 1 reaching 2)', 'dimensionless'),
     ('F21',    'View factor from surface 2 to 1', 'dimensionless')],
    note='Summation rule: sum(Fij) = 1 for all j (all radiation from surface i must reach some surface).',
    definition="View factor F12 is the fraction of radiation leaving surface 1 that is intercepted by surface 2 — purely geometric, depends only on shape and relative position. Summation rule: sum(F_ij for all j) = 1 (radiation must go somewhere). Self-view factor: F_11 = 0 for convex/flat surfaces (cannot see themselves); F_11 > 0 for concave surfaces. For a two-surface enclosure: F12 = 1 if surface 1 is entirely enclosed by surface 2. Reciprocity reduces the number of independent view factors to compute. Used with the radiosity method to solve multi-surface enclosure radiation problems (furnaces, rooms, solar collectors).",
    example={
        'problem': 'Two parallel plates: A1=2 m^2, A2=3 m^2. F12=0.6. Find F21.',
        'steps': [
            'A1*F12 = A2*F21',
            'F21 = A1*F12/A2 = 2*0.6/3 = 0.40',
        ],
        'answer': 'F21 = 0.40'})


# ── footer ────────────────────────────────────────────────────────────────────
ap('-' * 105, mono=True, sz=7, color=C['note'], before=10, after=10)
ap('Total: 108 equations  |  Topics: ME 101-106 / ME 201-206  |  All equations web-verified  |  SI units',
   sz=8, color=C['note'], align='center', before=0, after=6)
ap('-- HyESys Agent', bold=True, sz=9, color=C['title'], align='center', before=0, after=10)

doc.save(DST)
print(f'Saved: {DST}')
print(f'Total equations: 108')
