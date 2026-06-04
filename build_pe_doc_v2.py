"""
Builds FEE2026_Mechanical_Solutions_v2.docx from scratch
with exact questions from the PDF (pages 46-52) and correct solutions.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DST = r'C:\Users\JasonOng\Desktop\local docs\personal\PE\FEE2026_Mechanical_Solutions_v2.docx'

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

SEP = '─' * 80

# ── low-level helpers ─────────────────────────────────────────────────────────

def _rpr(bold=False, mono=False, color=None, sz_hp=20):
    rPr = OxmlElement('w:rPr')
    fn = 'Courier New' if mono else 'Calibri'
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), fn); rf.set(qn('w:hAnsi'), fn)
    rPr.append(rf)
    for tag in ('w:sz', 'w:szCs'):
        e = OxmlElement(tag); e.set(qn('w:val'), str(sz_hp)); rPr.append(e)
    if bold: rPr.append(OxmlElement('w:b'))
    if color:
        cl = OxmlElement('w:color'); cl.set(qn('w:val'), color); rPr.append(cl)
    return rPr

def _shading(hex_fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    return shd

def ap(text='', bold=False, mono=False, color=None, shading=None,
       align=None, sz_pt=10, before=60, after=60, keep=False):
    """Add paragraph to doc."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before / 6)
    pf.space_after  = Pt(after  / 6)
    if keep:
        pf.keep_with_next = True

    if shading:
        pPr = p._p.get_or_add_pPr()
        pPr.append(_shading(shading))

    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if text:
        run = p.add_run(text)
        run.bold = bold
        if mono: run.font.name = 'Courier New'
        run.font.size = Pt(sz_pt)
        if color:
            r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
            run.font.color.rgb = RGBColor(r, g, b)
    return p


def sep():
    ap(SEP, mono=True, sz_pt=8, color='999999', before=20, after=20)

def eq_block(equations):
    """Blue-shaded equations block. equations = list of (name, formula)."""
    ap('▸  EQUATIONS USED', bold=True, sz_pt=10, color='1F3A8E',
       shading='DCE8F8', before=30, after=20)
    for name, formula in equations:
        ap(f'   •  {name}:', bold=True, sz_pt=9, color='333333',
           shading='EDF3FC', before=10, after=4)
        ap(f'        {formula}', mono=True, sz_pt=9, color='0D0D7A',
           shading='EDF3FC', before=0, after=12)
    ap('', shading='EDF3FC', before=4, after=4)

def diag_block(lines):
    """Green-shaded ASCII diagram block."""
    ap('▸  DIAGRAM / SCHEMATIC', bold=True, sz_pt=10, color='14531A',
       shading='DFF0DC', before=20, after=10)
    for line in lines:
        ap(line, mono=True, sz_pt=8.5, color='1A1A1A',
           shading='F2FAF2', before=0, after=0)
    ap('', shading='F2FAF2', before=4, after=16)

def qhdr(text):
    ap(text, bold=True, sz_pt=11, color='000000', before=40, after=20)

def answer(text):
    ap(f'▶  ANSWER:  {text}', bold=True, sz_pt=10, color='8B0000',
       before=30, after=20)

def note(text):
    ap(f'Note: {text}', sz_pt=9, color='555555', before=10, after=10)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT HEADER
# ══════════════════════════════════════════════════════════════════════════════
ap('PROFESSIONAL ENGINEERS EXAMINATION 2026', bold=True, sz_pt=14,
   align='center', before=0, after=20)
ap('Fundamentals of Engineering Examination (FEE)', bold=True, sz_pt=12,
   align='center', before=0, after=10)
ap('MECHANICAL DISCIPLINE — WORKED SOLUTIONS', bold=True, sz_pt=12,
   align='center', before=0, after=10)
ap('Part 1: Multiple Choice Questions (Q1–Q10)  |  Part 2: Long Questions (Q1–Q7)',
   sz_pt=10, align='center', before=0, after=10)
ap('Prepared by: Jason Ong Zong Yi', sz_pt=10, align='center', before=0, after=6)
ap('Date: May 2026', sz_pt=10, align='center', before=0, after=20)
ap('Note: All solutions are for the Mechanical discipline questions only. '
   'Questions from other disciplines (Civil, Electrical, etc.) are excluded.',
   sz_pt=9, color='555555', before=0, after=30)

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — MCQ
# ══════════════════════════════════════════════════════════════════════════════
ap('PART 1 — Multiple Choice Questions', bold=True, sz_pt=13, before=20, after=10)
ap('10 questions, 2.5 marks each. Answer all questions.', sz_pt=10, before=0, after=20)

# ── Q1 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Reynolds Number', 'Re = ρVD/μ = VD/ν'),
    ('Turbulent friction factor — Colebrook-White', '1/√f = −2 log₁₀(ε/(3.7D) + 2.51/(Re√f))'),
    ('Viscous dissipation (turbulence)', 'ε_dissip = ν (∂uᵢ/∂xⱼ + ∂uⱼ/∂xᵢ)² / 2  [kinetic energy → heat]'),
    ('Turbulent skin-friction coefficient (flat plate)', 'C_f ≈ 0.074 / Re_L^0.2  (higher than laminar 1.328/√Re_L)'),
])
qhdr('Question 1 — Fluid Mechanics: Turbulence')
ap('Turbulence is a very important concept in fluid mechanics. '
   'Which of the following statement(s) on turbulence is(are) correct?\n'
   '(i)  Turbulence dissipates useful kinetic energy of flow to heat.\n'
   '(ii)  Turbulence increases the frictional drag.\n'
   '(iii) Turbulence increases convective heat transfer.\n\n'
   '(a) (i) and (ii)    (b) (i) and (iii)    (c) (ii) and (iii)    (d) All of the above',
   sz_pt=10)
diag_block([
    '     TURBULENCE — KEY EFFECTS',
    '     ─────────────────────────────────────────────────────────────────',
    '     (i)  Energy cascade: turbulent eddies → smaller eddies → viscous dissipation → HEAT ✓',
    '          Large eddies (L) → inertial cascade → Kolmogorov scale η → thermal energy',
    '',
    '     (ii) Frictional drag:',
    '          Laminar:    f = 64/Re  (Hagen-Poiseuille)',
    '          Turbulent:  f >> 64/Re  (Moody chart, depends on ε/D)',
    '          Turbulent boundary layer exerts higher wall shear stress → MORE DRAG ✓',
    '',
    '     (iii) Convective heat transfer:',
    '          h ∝ Nu ∝ Re^0.8 × Pr^0.4  (Dittus-Boelter, turbulent)',
    '          Turbulent mixing → steep temp gradient at wall → MORE heat transfer ✓',
    '',
    '     All three statements are correct.',
])
ap('(i)  TRUE — Turbulent eddies undergo viscous dissipation at the Kolmogorov scale, '
   'converting kinetic energy to heat. This is why turbulence is "wasteful" energetically.', sz_pt=10)
ap('(ii) TRUE — The turbulent velocity profile has a steeper gradient near the wall than '
   'laminar flow, producing higher wall shear stress τ_w and thus greater frictional drag.', sz_pt=10)
ap('(iii) TRUE — Turbulent eddies transport heat across streamlines far more effectively '
   'than molecular diffusion alone; Nu (Dittus-Boelter) ∝ Re^0.8.', sz_pt=10)
ap('∴  All three statements are correct.', bold=True, sz_pt=10)
answer('(d) All of the above')

# ── Q2 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Total strain at any fibre (superposition)', 'ε(y) = ε_axial + ε_bending(y)'),
    ('Linear strain distribution (Euler-Bernoulli)', 'ε(y) = A + B·y  (y measured from top)'),
    ('Neutral axis location (ε = 0)', 'y_NA = −A / B  (from top surface)'),
    ('Axial strain — at centroid (y_c = h/2)', 'ε_axial = ε(y_c) = A + B·(h/2)'),
    ('Axial stress', 'σ_axial = E × ε_axial'),
    ('Max bending stress (at extreme fibre from centroid)', 'σ_bend = E × |ε_bending_max| = E × |ε_extremefibre − ε_axial|'),
])
qhdr('Question 2 — Mechanics of Materials: Combined Bending & Axial Load')
ap('A solid rectangular beam of 250 mm depth is loaded by combined loading of bending '
   'moment and axial tensile force such that strain measurements showed that the top surface '
   'contracted by 200 microstrain and the bottom surface elongated by 500 microstrain '
   'longitudinally. Determine the position of the neutral axis (from the top surface) and '
   'calculate the axial stress caused by the tensile force and bending stress caused by the '
   'bending moment independently. E = 200 GPa.\n\n'
   '(a) 71.43 mm, 30 MPa, 200 MPa\n'
   '(b) 125 mm, 150 MPa, 150 MPa\n'
   '(c) 125 mm, 30 MPa, 200 MPa\n'
   '(d) 178.57 mm, 30 MPa, 200 MPa',
   sz_pt=10)
diag_block([
    '     STRAIN DISTRIBUTION ACROSS 250 mm BEAM DEPTH',
    '     ─────────────────────────────────────────────────────────────────',
    '     Top surface (y=0):    ε_top = −200 μɛ  (COMPRESSION / contraction)',
    '     ┌────────────────────────────────────────────────────────────────┐',
    '     │  y=0                     ε = −200 μɛ                         │',
    '     │                                                                │',
    '     │  y=71.43mm ─ ─ ─ ─ NA ─ ─ ─ ─  ε = 0  (neutral axis)       │',
    '     │                                                                │',
    '     │  y=125mm  ─ ─ ─ ─ CG ─ ─ ─ ─   ε = +150 μɛ  (centroid)     │',
    '     │                                                                │',
    '     │  y=250mm                         ε = +500 μɛ                 │',
    '     └────────────────────────────────────────────────────────────────┘',
    '     Bottom surface (y=250): ε_bot = +500 μɛ  (TENSION / elongation)',
    '     ─────────────────────────────────────────────────────────────────',
    '     Linear fit: ε(y) = −200 + 2.8y  (μɛ, y in mm from top)',
    '     ε_axial = ε at centroid (y=125): = −200 + 2.8×125 = +150 μɛ → σ_axial = 30 MPa',
    '     σ_bend_max = E × |ε_bot − ε_axial| = 200,000 × |500−150| × 10⁻⁶ = 70 MPa',
])
ap('Given: h = 250 mm, ε_top = −200×10⁻⁶ (compression), ε_bot = +500×10⁻⁶ (tension), E = 200 GPa', sz_pt=10)
ap('Step 1 — Linear strain fit:  ε(y) = A + B·y  (y from top)', sz_pt=10)
ap('  At y = 0:   A = −200×10⁻⁶', sz_pt=10)
ap('  At y = 250: −200×10⁻⁶ + 250B = 500×10⁻⁶  →  B = 700×10⁻⁶/250 = 2.8×10⁻⁶ mm⁻¹', sz_pt=10)
ap('Step 2 — Neutral axis (ε = 0):', sz_pt=10)
ap('  0 = −200×10⁻⁶ + 2.8×10⁻⁶ × y_NA  →  y_NA = 200/2.8 = 71.43 mm from top surface', bold=True, sz_pt=10)
ap('Step 3 — Axial strain at centroid (y_c = 125 mm):', sz_pt=10)
ap('  ε_axial = −200×10⁻⁶ + 2.8×10⁻⁶ × 125 = −200×10⁻⁶ + 350×10⁻⁶ = +150×10⁻⁶', sz_pt=10)
ap('  σ_axial = E × ε_axial = 200,000 MPa × 150×10⁻⁶ = 30 MPa  (tensile ✓)', bold=True, sz_pt=10)
ap('Step 4 — Bending strain at extreme fibre (bottom, 125 mm from centroid):', sz_pt=10)
ap('  ε_bend_bot = ε_bot − ε_axial = 500×10⁻⁶ − 150×10⁻⁶ = 350×10⁻⁶', sz_pt=10)
ap('  σ_bend = E × ε_bend_bot = 200,000 × 350×10⁻⁶ = 70 MPa', bold=True, sz_pt=10)
note('The answer choices state 200 MPa for bending stress; the rigorous calculation yields '
     '70 MPa. The NA position of 71.43 mm (option a) and σ_axial = 30 MPa are correct. '
     'The bending stress value in the answer key appears to be a typographical error.')
answer('(a) 71.43 mm, 30 MPa — NA and axial stress confirmed; σ_bend computed = 70 MPa')

# ── Q3 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Polar second moment of area — hollow shaft (thick-wall formula)', 'J = π(d_o⁴ − d_i⁴) / 32'),
    ('Torsion formula (Coulomb)', 'τ_max = T·r_o / J  =  T·(d_o/2) / J'),
    ('Angle of twist', 'φ = T·L / (G·J)  [radians]  →  φ° = φ × 180/π'),
    ('Shear modulus', 'G = τ/γ  (modulus of rigidity)'),
])
qhdr('Question 3 — Torsion: Hollow Circular Shaft Twist Angle')
ap('A hollow circular section steel rod 300 mm long with internal diameter 10 mm and '
   'external diameter 16 mm is used as a torsional spring with one end fixed and the '
   'other end free to twist. Determine the torsional twist in degrees when a torque of '
   '100 N·m is applied about the rod axis. Material G = 80 GPa. '
   '(Hint: use thick wall formula)\n\n'
   '(a) 5.33°    (b) 3.94°    (c) 1.55°    (d) 7.43°',
   sz_pt=10)
diag_block([
    '     HOLLOW SHAFT TORSION — THICK WALL FORMULA',
    '     ─────────────────────────────────────────────────────────────────',
    '     T=100 N·m ↺              ↻ (fixed end)',
    '     |←──────────── L = 300 mm ─────────────→|',
    '     ══════════════════════════════════════════════',
    '                                                   ',
    '     Cross-section (end view):',
    '          d_o = 16 mm',
    '       ┌────────────────┐',
    '       │  d_i = 10 mm   │',
    '       │   ┌────────┐   │',
    '       │   │ hollow │   │',
    '       │   └────────┘   │',
    '       └────────────────┘',
    '     ─────────────────────────────────────────────────────────────────',
    '     J = π(16⁴ − 10⁴)/32 = π(65536 − 10000)/32 = π×55536/32 = 5455 mm⁴',
    '     φ = TL/GJ = (100×10³ × 300) / (80,000 × 5455) = 30×10⁶/436.4×10⁶',
    '       = 0.06876 rad × (180/π) = 3.94°',
])
ap('Given: d_o = 16 mm, d_i = 10 mm, L = 300 mm, T = 100 N·m = 100,000 N·mm, G = 80 GPa = 80,000 N/mm²', sz_pt=10)
ap('Step 1 — Polar second moment of area (thick-wall formula):', sz_pt=10)
ap('  J = π(d_o⁴ − d_i⁴)/32 = π(16⁴ − 10⁴)/32', sz_pt=10)
ap('  16⁴ = 65,536 mm⁴    10⁴ = 10,000 mm⁴', sz_pt=10)
ap('  J = π × 55,536 / 32 = 5,455 mm⁴', sz_pt=10)
ap('Step 2 — Angle of twist:', sz_pt=10)
ap('  φ = T·L / (G·J) = (100,000 × 300) / (80,000 × 5,455)', sz_pt=10)
ap('  φ = 30,000,000 / 436,400,000 = 0.06876 rad', sz_pt=10)
ap('  φ = 0.06876 × (180/π) = 3.94°', bold=True, sz_pt=10)
answer('(b) 3.94°')

# ── Q4 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Stress concentration at inclusion (Inglis/Eshelby)', 'σ_max = K_t × σ_nom  (K_t elevated near inclusion boundary)'),
    ('Inclusion detectability threshold (visual)', 'Minimum visible defect size ≈ 0.1–0.5 mm (naked eye)'),
    ('ESR/VAR process — inclusion reduction', 'Electro-Slag Remelting (ESR) or Vacuum Arc Remelting (VAR)'),
])
qhdr('Question 4 — Materials Science: Inclusions in Steel')
ap('Inclusions may sometimes be found in steels of different compositions. '
   'These inclusions:\n\n'
   '(a) are sites for potential initiation of cracks\n'
   '(b) are invisible to naked-eye inspections of the steel pieces\n'
   '(c) may be significantly reduced through costly special processes\n'
   '(d) All of the above',
   sz_pt=10)
diag_block([
    '     STEEL INCLUSIONS — MECHANISMS AND CHARACTERISTICS',
    '     ─────────────────────────────────────────────────────────────────',
    '     (a) Crack initiation:',
    '         Inclusion (e.g., MnS, Al₂O₃) has different elastic modulus than matrix',
    '         → stress concentration at inclusion/matrix interface',
    '         → fatigue cracks initiate under cyclic loading (K_t > 1)',
    '',
    '     (b) Visibility:',
    '         Most inclusions: 1–100 μm in size (NOT visible to naked eye)',
    '         Requires: optical/SEM microscopy, ultrasonic testing, or X-ray',
    '',
    '     (c) Reduction processes (costly):',
    '         • ESR (Electro-Slag Remelting) — refines slag, reduces oxide inclusions',
    '         • VAR (Vacuum Arc Remelting) — removes dissolved gases and inclusions',
    '         • Secondary metallurgy (ladle refining) — adds cost',
    '',
    '     All three statements are true → (d) All of the above',
])
ap('(a) TRUE — Inclusions act as stress raisers (Kt > 1) due to elastic mismatch with '
   'the steel matrix. Fatigue cracks preferentially nucleate at inclusion surfaces, '
   'especially MnS stringers and Al₂O₃ clusters.', sz_pt=10)
ap('(b) TRUE — Typical inclusions range from 1–100 μm, far below the naked-eye '
   'resolution limit (~100 μm). Detection requires microscopy, ultrasonic, or radiographic NDT.', sz_pt=10)
ap('(c) TRUE — Special remelting processes (ESR, VAR) significantly reduce inclusion '
   'content but add considerable cost to the manufacturing process.', sz_pt=10)
ap('∴  All three statements are correct.', bold=True, sz_pt=10)
answer('(d) All of the above')

# ── Q5 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Conservation of Linear Momentum — perfectly inelastic collision', 'm_b·v_b + m_B·0 = (m_b + m_B)·v_f'),
    ('Work-Energy Theorem — sliding phase', '½(m_total)·v_f² = μ_k·(m_total)·g·d'),
    ('Post-impact block velocity', 'v_f = √(2·μ_k·g·d)'),
    ('Kinetic friction force', 'F_k = μ_k × N = μ_k × (m_b + m_B)·g'),
    ("Newton's Second Law", 'ΣF = m·a'),
])
qhdr('Question 5 — Dynamics: Bullet-Block Impact')
ap('A 21 g bullet hits a 2 kg block that is initially at rest. After the collision, '
   'the bullet becomes embedded in the block, and they slide a distance of 0.31 m. '
   'If the coefficient of kinetic friction between the block and the ground is '
   'μk = 0.7, determine the pre-impact speed of the bullet.\n\n'
   '(a) 198.5 m/s    (b) 251.2 m/s    (c) 351.5 m/s    (d) 405.2 m/s',
   sz_pt=10)
diag_block([
    '     PHASE 1: IMPACT — Conservation of Linear Momentum',
    '     ─────────────────────────────────────────────────────────────────',
    '     ● ──────────────────► + [■]  →  [■●] ──────►  v_f',
    '     21g @ v_b m/s          2 kg at rest   2.021 kg',
    '     m_b·v_b = (m_b + m_B)·v_f',
    '',
    '     PHASE 2: SLIDING — Work-Energy Theorem',
    '     ─────────────────────────────────────────────────────────────────',
    '          [■●] ─────────────────────────► STOPS',
    '          v_f                            d = 0.31 m',
    '               ← ← F_k = μ_k·m_total·g = friction',
    '          ↑ N = m_total·g',
    '     ground',
    '     ─────────────────────────────────────────────────────────────────',
    '     v_f = √(2×0.7×9.81×0.31) = √4.2537 = 2.063 m/s',
    '     v_b = (2.021 × 2.063) / 0.021 = 4.169 / 0.021 = 198.5 m/s',
])
ap('m_bullet = 0.021 kg, m_block = 2.0 kg, d = 0.31 m, μk = 0.70, g = 9.81 m/s²', sz_pt=10)
ap('Step 1 — Post-impact velocity (Work-Energy on sliding phase):', sz_pt=10)
ap('  ½(m_total)v_f² = μk·(m_total)·g·d  →  v_f = √(2·μk·g·d)', sz_pt=10)
ap('  v_f = √(2 × 0.70 × 9.81 × 0.31) = √4.254 = 2.063 m/s', sz_pt=10)
ap('Step 2 — Pre-impact bullet speed (Conservation of Momentum):', sz_pt=10)
ap('  m_b·v_b = (m_b + m_B)·v_f', sz_pt=10)
ap('  v_b = (0.021 + 2.0) × 2.063 / 0.021 = 2.021 × 2.063 / 0.021 = 198.5 m/s', bold=True, sz_pt=10)
answer('(a) 198.5 m/s')

# ── Q6 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('SUVAT — velocity from rest', 'v = u + a·t  →  t_acc = v_top / a'),
    ('SUVAT — distance during acceleration', 's = v²/(2a)  (u = 0)'),
    ('Cruise distance', 'd_cruise = d_total − d_acc − d_dec'),
    ('Cruise time', 't_cruise = d_cruise / v_top'),
    ('Total time', 't_total = t_acc + t_cruise + t_dec  (by symmetry t_acc = t_dec)'),
])
qhdr('Question 6 — Kinematics: High-Speed Rail Travel Time')
ap('A high-speed rail transportation system has a top speed of 100 m/s. For the comfort '
   'of the passengers, the magnitude of the acceleration and deceleration is limited to '
   '2 m/s². Determine the minimum time required for a trip of 100 km.\n\n'
   '(a) 0.30 hr    (b) 0.25 hr    (c) 0.17 hr    (d) 0.25 hr',
   sz_pt=10)
diag_block([
    '     VELOCITY-TIME GRAPH — RAIL TRIP',
    '     ─────────────────────────────────────────────────────────────────',
    '     v (m/s)',
    '     100 ─────────────────────────────────────────────────────────',
    '          /                  CRUISE @ 100 m/s                     \\',
    '         /   a=2m/s²         d_cruise = 95,000 m                   \\ a=2m/s²',
    '        /    t_acc=50s                                    t_dec=50s  \\',
    '       0─────────┬──────────────────────────────────────────┬────────► t',
    '               50s                                        1050s',
    '     ─────────────────────────────────────────────────────────────────',
    '     t_acc = 100/2 = 50 s   d_acc = 100²/(2×2) = 2500 m',
    '     t_dec = 50 s           d_dec = 2500 m',
    '     d_cruise = 100,000 − 2500 − 2500 = 95,000 m',
    '     t_cruise = 95,000/100 = 950 s',
    '     t_total = 50 + 950 + 50 = 1050 s = 0.292 hr ≈ 0.30 hr',
])
ap('Given: v_top = 100 m/s,  a = dec = 2 m/s²,  d_total = 100 km = 100,000 m', sz_pt=10)
ap('Step 1 — Acceleration phase:  (from rest to v_top)', sz_pt=10)
ap('  t_acc = v_top / a = 100 / 2 = 50 s', sz_pt=10)
ap('  d_acc = v_top² / (2a) = 100² / (2×2) = 2,500 m', sz_pt=10)
ap('Step 2 — Deceleration phase:  (symmetric by symmetry of a = dec)', sz_pt=10)
ap('  t_dec = 50 s,   d_dec = 2,500 m', sz_pt=10)
ap('Step 3 — Cruise phase:', sz_pt=10)
ap('  d_cruise = 100,000 − 2,500 − 2,500 = 95,000 m', sz_pt=10)
ap('  t_cruise = 95,000 / 100 = 950 s', sz_pt=10)
ap('Step 4 — Total time:', sz_pt=10)
ap('  t_total = 50 + 950 + 50 = 1,050 s = 1,050/3,600 hr = 0.292 hr ≈ 0.30 hr', bold=True, sz_pt=10)
answer('(a) 0.30 hr')

# ── Q7 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Standard 2nd-order characteristic equation', 's² + 2ζω_n·s + ω_n² = 0'),
    ('Natural frequency', 'ω_n = √(constant term)'),
    ('Damping ratio', 'ζ = (s-coefficient) / (2·ω_n)'),
    ('Quadratic formula', 's = (−b ± √(b²−4ac)) / (2a)'),
    ('Classification', 'ζ < 1: underdamped  |  ζ = 1: critically damped  |  ζ > 1: overdamped'),
])
qhdr('Question 7 — Control Systems: Damping Classification')
ap('If the characteristic equation of the closed-loop system is s² + 3s + 2 = 0, '
   'then the system is:\n\n'
   '(a) Over damped    (b) Critically damped    (c) Under damped    (d) Unstable',
   sz_pt=10)
diag_block([
    '     s-PLANE POLE LOCATIONS',
    '     ─────────────────────────────────────────────────────────────────',
    '          jω',
    '           │',
    '     ──────┼────────── σ',
    '      ×    │    ×',
    '     s=-2  │   s=-1',
    '           │',
    '     Two distinct REAL NEGATIVE poles → OVERDAMPED (ζ > 1)',
    '     ─────────────────────────────────────────────────────────────────',
    '     ω_n = √2 = 1.414 rad/s',
    '     ζ = 3 / (2 × 1.414) = 1.061 > 1  → Overdamped ✓',
    '     Roots: s = (−3 ± √(9−8))/2 = (−3 ± 1)/2 → s₁=−1, s₂=−2',
])
ap('Standard form: s² + 2ζω_n·s + ω_n² = 0. Comparing with s² + 3s + 2 = 0:', sz_pt=10)
ap('  ω_n² = 2  →  ω_n = √2 = 1.414 rad/s', sz_pt=10)
ap('  2ζω_n = 3  →  ζ = 3/(2×1.414) = 1.061', sz_pt=10)
ap('  ζ = 1.061 > 1.0  →  OVERDAMPED', bold=True, sz_pt=10)
ap('Verification via roots: s = (−3 ± √(9−8))/2 = (−3 ± 1)/2  →  s₁=−1, s₂=−2  '
   '(two real, distinct, negative roots ✓)', sz_pt=10)
answer('(a) Over damped')

# ── Q8 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('First Law of Thermodynamics — closed system', 'ΔU = Q − W'),
    ('State function property', 'ΔU_cycle = 0  (U depends only on state, not path)'),
    ('Cyclic energy balance', 'Q_net = W_net  (net heat in = net work out)'),
])
qhdr('Question 8 — Thermodynamics: First Law for Cyclic Process')
ap('In a cyclic process, the net change of internal energy is:\n\n'
   '(a) Equal to zero    (b) Equal to one    '
   '(c) Greater than one    (d) Smaller than one',
   sz_pt=10)
diag_block([
    '     P-V DIAGRAM — CLOSED THERMODYNAMIC CYCLE',
    '     ─────────────────────────────────────────────────────────────────',
    '     P │',
    '       │   2──────────3',
    '       │  ╱  Q_in ↑   ╲',
    '       │ ╱   (heat in)  ╲',
    '       │1   W_net (area) 4',
    '       │ ╲              ╱',
    '       │  ╲  Q_out ↓  ╱',
    '       │   2──────────1',
    '       └──────────────────► V',
    '     System returns to state 1 after full cycle → ΔU = 0',
    '     First Law: ΔU = Q − W = 0  ∴  Q_net = W_net',
])
ap('Internal energy U is a STATE FUNCTION — its value depends only on the thermodynamic '
   'state (T, P, V), not on the path taken.', sz_pt=10)
ap('After a complete cycle, the system returns to its exact initial state. '
   'Therefore ALL state functions (U, H, S) return to their initial values.', sz_pt=10)
ap('  ΔU_cycle = 0  →  Q_net = W_net  (First Law for complete cycle)', bold=True, sz_pt=10)
answer('(a) Equal to zero')

# ── Q9 MCQ ────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Adiabatic condition (defining equation)', 'Q = 0  (no heat exchange with surroundings)'),
    ('First Law for adiabatic process', 'ΔU = −W  (work done BY system decreases internal energy)'),
    ('Reversible adiabatic — ideal gas (Poisson)', 'PV^γ = constant  (γ = c_p/c_v)'),
    ('Practical condition for near-adiabatic', 'Process occurs very rapidly → no time for heat transfer'),
])
qhdr('Question 9 — Thermodynamics: Adiabatic Process')
ap('In an adiabatic process, which of the above statements are correct?\n\n'
   '(i)   the temperature change is zero.\n'
   '(ii)  the change in internal energy is zero.\n'
   '(iii) the heat interaction between system and the surroundings is zero.\n'
   '(iv)  the process is carried out very rapidly.\n\n'
   '(a) All of the above    (b) (i), (ii) and (iii)    '
   '(c) (i), (ii) and (iv)    (d) (iii) and (iv)',
   sz_pt=10)
diag_block([
    '     P-V DIAGRAM — ADIABATIC vs ISOTHERMAL',
    '     ─────────────────────────────────────────────────────────────────',
    '     P │',
    '       │  × ← same initial state',
    '       │   \\\\',
    '       │    \\\\ ← ADIABATIC: PV^γ = const (steeper slope)',
    '       │     \\\\   Q=0; T DROPS during expansion',
    '       │      \\',
    '       │       ─ ─ ─ ─ ─ ─ ─  ← ISOTHERMAL: PV=const',
    '       │                          T=const; ΔU=0',
    '       └────────────────────────► V',
    '     (i)  FALSE — T changes (adiabatic compression ↑T, expansion ↓T)',
    '     (ii) FALSE — ΔU = −W ≠ 0 (unless W=0 also)',
    '     (iii) TRUE  — Q=0 is the DEFINITION of adiabatic',
    '     (iv) TRUE  — Very rapid processes give no time for heat transfer → ≈ adiabatic',
])
ap('(i)  FALSE — Temperature changes in an adiabatic process. '
   'Adiabatic compression raises T; expansion lowers T.', sz_pt=10)
ap('(ii) FALSE — ΔU = Q − W = 0 − W = −W. '
   'Internal energy changes unless no work is done.', sz_pt=10)
ap('(iii) TRUE — Q = 0 is the DEFINITION of an adiabatic process.', sz_pt=10)
ap('(iv) TRUE — Processes that occur very rapidly do not allow time for heat exchange '
   'with surroundings, approximating adiabatic conditions. '
   '(e.g., rapid compression in diesel engines, sound wave propagation)', sz_pt=10)
ap('∴  Only (iii) and (iv) are correct statements about adiabatic processes.', bold=True, sz_pt=10)
answer('(d) (iii) and (iv)')

# ── Q10 MCQ ───────────────────────────────────────────────────────────────────
sep()
eq_block([
    ("Taylor's Tool Life Equation", 'V × T^n = C'),
    ('Solved for tool life', 'T = (C/V)^(1/n)'),
    ('Log form (linear on log-log plot)', 'log V + n·log T = log C  (slope = −n)'),
])
qhdr("Question 10 — Manufacturing: Taylor's Tool Life Equation")
ap("Which of the following statements is correct?\n\n"
   "(a) Tool life increases with the increase of cutting speed\n"
   "(b) Tool life decreases with the increase of cutting speed\n"
   "(c) Cutting speed has no influence on tool life\n"
   "(d) None of the above",
   sz_pt=10)
diag_block([
    "     TAYLOR'S TOOL LIFE CURVE (log-log scale)",
    '     ─────────────────────────────────────────────────────────────────',
    '     log V │',
    '           │  ×  V₁',
    '           │    \\  slope = −n',
    '           │     ×  V₂ > V₁',
    '           │      \\',
    '           │       ×  V₃',
    '           │        \\',
    '           └───────────────────────────── log T',
    '                T₃      T₂   T₁    (T₁ > T₂ > T₃)',
    '     Higher cutting speed → shorter tool life  (V↑ → T↓)',
    '     V₁T₁^n = V₂T₂^n = C  (same constant for given tool-workpiece pair)',
])
ap("From VT^n = C:  T = (C/V)^(1/n).", sz_pt=10)
ap("As V increases, C/V decreases, so T^(1/n) decreases, so T decreases.", sz_pt=10)
ap("∴  Tool life DECREASES with increasing cutting speed.", bold=True, sz_pt=10)
answer('(b) Tool life decreases with the increase of cutting speed')

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — LONG QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
sep()
ap('PART 2 — Long Questions', bold=True, sz_pt=13, before=20, after=10)
ap('Answer ALL questions. Show all working clearly.', sz_pt=10, before=0, after=20)

# ── LQ1 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Darcy-Weisbach (Fanning form — as given)', 'h_f = 4f·(l/d)·V²/(2g)   [f=0.01 given]'),
    ('Entry loss (sharp entrance)', 'h_entry = 0.5 × V²/(2g)'),
    ('Exit loss (sharp exit to reservoir)', 'h_exit = 1.0 × V²/(2g)'),
    ('Sudden expansion loss (Borda-Carnot)', 'h_exp = [1 − A₁/A₂]² × V₁²/(2g)'),
    ('Continuity equation', 'A₁V₁ = A₂V₂  →  V₂ = V₁·(D₁/D₂)²'),
    ('Extended Bernoulli (open reservoirs)', 'Δz = Σh_losses  (V_surface ≈ 0 at both reservoirs)'),
    ('Volumetric flow rate', 'Q = A·V = (πD²/4)·V'),
    ('Total Energy Line vs Hydraulic Grade Line', 'TEL − HGL = V²/(2g)  (velocity head)'),
])
qhdr('Question 1 — Fluid Mechanics: Two-Reservoir Pipe Flow System  (20 marks)')
ap('Two reservoirs A and B are connected by a pipeline: 100 mm diameter for the first 10 m '
   '(Pipe 1), then 150 mm diameter for the remaining 20 m (Pipe 2). The entrance from and '
   'exit to the reservoir are sharp, and the expansion from 100 mm to 150 mm is sudden. '
   'The water surface of the upper reservoir A is 10 m above that of lower reservoir B.\n\n'
   '(a) Tabulate the losses of head that occur and calculate the discharge flow rate.  (10 marks)\n'
   '(b) Draw the hydraulic gradient and total energy gradient along the pipeline.  (5 marks)\n'
   '(c) What is the difference between total energy line and hydraulic gradient line?  (2 marks)\n'
   '(d) [3 marks]\n\n'
   'Given: f = 0.01 (Fanning), g = 9.81 m/s², ρ = 1000 kg/m³, Pa = 101 kPa, '
   'K_entry = 0.5, K_exit = 1.0. Expansion loss = [1−A₁/A₂]²·V₁²/(2g)',
   sz_pt=10)
diag_block([
    '     TWO-RESERVOIR PIPELINE SYSTEM',
    '     ─────────────────────────────────────────────────────────────────',
    '     ┌──────────────┐ ← Free surface A (datum reference: z_A = 10 m)',
    '     │ Reservoir A  │',
    '     └──────┬───────┘',
    '            │←── Pipe 1: D₁=100mm, L₁=10m ──→│←── Pipe 2: D₂=150mm, L₂=20m ──→│',
    '            │  K_entry=0.5                expansion     K_exit=1.0                │',
    '     ────── → ────────────────────────────────────────────────────────────────────→',
    '                                              ↓ sudden expansion                   │',
    '     Δz=10m                                                               ┌────────┘',
    '                                                                           │ Reservoir B',
    '                                                                           └─────────────',
    '                                                               Free surface B (z_B=0)',
    '     ─────────────────────────────────────────────────────────────────',
    '     Bernoulli: z_A = z_B + Σh_losses  →  10 = Σh_losses',
    '     All losses expressed as multiples of V₁²/(2g):',
])
ap('Part (a) — Head Loss Tabulation:', bold=True, sz_pt=10)
ap('Geometry:  A₁ = π/4×0.1² = 7.854×10⁻³ m²,  '
   'A₂ = π/4×0.15² = 1.767×10⁻² m²', sz_pt=10)
ap('Continuity: V₂ = V₁×(A₁/A₂) = V₁×(0.1/0.15)² = (4/9)V₁ = 0.4444V₁', sz_pt=10)
ap('Let k = V₁²/(2g)  for convenience.', sz_pt=10)
ap('', sz_pt=6)
ap('  h_entry    = 0.5 × V₁²/(2g)                           = 0.500 k', sz_pt=10, mono=True)
ap('  h_f1       = 4×0.01×(10/0.1) × V₁²/(2g) = 4.0 k      = 4.000 k', sz_pt=10, mono=True)
ap('  h_exp      = [1−4/9]² × V₁²/(2g) = (5/9)² k          = 0.309 k', sz_pt=10, mono=True)
ap('  h_f2       = 4×0.01×(20/0.15) × V₂²/(2g)', sz_pt=10, mono=True)
ap('             = 5.333 × (4/9)² × V₁²/(2g)               = 1.053 k', sz_pt=10, mono=True)
ap('  h_exit     = 1.0 × V₂²/(2g) = (4/9)² k              = 0.198 k', sz_pt=10, mono=True)
ap('  ─────────────────────────────────────────────────────────────────', sz_pt=10, mono=True)
ap('  TOTAL Σh   = (0.500+4.000+0.309+1.053+0.198) k       = 6.059 k', sz_pt=10, mono=True, bold=True)
ap('', sz_pt=6)
ap('Δz = 6.059 × V₁²/(2g)', sz_pt=10)
ap('  10 = 6.059 × V₁²/(2×9.81)', sz_pt=10)
ap('  V₁² = 10 × 19.62 / 6.059 = 32.38 m²/s²', sz_pt=10)
ap('  V₁ = 5.69 m/s;   V₂ = (4/9)×5.69 = 2.53 m/s', bold=True, sz_pt=10)
ap('  Q = A₁×V₁ = 7.854×10⁻³ × 5.69 = 0.0447 m³/s = 44.7 L/s', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('Individual head loss values (substituting V₁²/(2g) = 32.38/19.62 = 1.651 m):', sz_pt=10)
ap('  h_entry=0.826m | h_f1=6.604m | h_exp=0.510m | h_f2=1.738m | h_exit=0.327m | Total=10.00m ✓', sz_pt=9, mono=True)
ap('', sz_pt=6)
ap('Part (b) — Hydraulic Gradient (HGL) and Total Energy Line (TEL):', bold=True, sz_pt=10)
ap('• At reservoir A surface: TEL = HGL = 10 m (V≈0, so TEL=HGL)', sz_pt=10)
ap('• Entering pipe 1 (just inside): TEL drops by h_entry=0.826m; V₁²/2g=1.651m; HGL=TEL−V₁²/2g', sz_pt=10)
ap('• Along pipe 1: TEL drops linearly (friction); HGL parallel, offset below by V₁²/2g', sz_pt=10)
ap('• At expansion: TEL drops abruptly by h_exp; V drops → V₂²/2g=0.327m; HGL rises (slower V)', sz_pt=10)
ap('• Along pipe 2: TEL drops linearly (lower gradient — larger pipe, lower V)', sz_pt=10)
ap('• At exit: TEL drops by h_exit; HGL = reservoir B surface = 0 m', sz_pt=10)
ap('', sz_pt=6)
ap('Part (c) — Difference between TEL and HGL:', bold=True, sz_pt=10)
ap('  TEL (Total Energy Line) = P/(ρg) + V²/(2g) + z  [pressure + velocity + elevation head]', sz_pt=10)
ap('  HGL (Hydraulic Grade Line) = P/(ρg) + z  [pressure + elevation head only]', sz_pt=10)
ap('  Difference = V²/(2g)  (velocity head). '
   'At any cross-section, TEL is always above HGL by the velocity head.', bold=True, sz_pt=10)

# ── LQ2 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Power-Torque-Angular velocity', 'P = T·ω  →  T = P/ω'),
    ('Angular velocity', 'ω = 2πN/60  [rad/s]'),
    ('Torsion formula — solid shaft (Coulomb)', 'τ_max = T·(d/2)/J = 16T/(πd³)'),
    ('Polar second moment — solid', 'J = πd⁴/32'),
    ('Polar second moment — hollow', 'J_h = π(d_o⁴ − d_i⁴)/32'),
    ('Shaft design equation (solid)', 'd³ = 16T/(π·τ_allow)'),
    ('Shaft design equation (hollow, given d_o)', 'τ = T·(d_o/2) / J_h  →  solve for d_i'),
])
qhdr('Question 2 — Solid Mechanics: Stepped Shaft Design  (20 marks)')
ap('In the figure below, a motor drives a stepped shaft to which two pulleys A and B are '
   'attached and connected to machine loads. The torques Ta and Tb at the pulleys are equal '
   'and exerted in the same direction. The motor outputs 300 kW at 3000 rpm.\n\n'
   '(a) Calculate the torques at sections OA and AB.  (4 marks)\n'
   '(b) Assuming solid circular sections and maximum allowable shear stress = 110 N/mm², '
   'calculate the minimum diameters of each section.  (8 marks)\n'
   '(c) To reduce machining cost, hollow circular sections are used with outer diameter '
   'constant at 40 mm. Calculate minimum internal diameters for OA and AB, and '
   'recommend a common inner diameter.  (8 marks)',
   sz_pt=10)
diag_block([
    '     STEPPED SHAFT — LAYOUT AND TORQUE DIAGRAM',
    '     ─────────────────────────────────────────────────────────────────',
    '     MOTOR(300kW)    PULLEY A(-Ta)         PULLEY B(-Tb)',
    '     @3000rpm        Ta=Tb=T_motor/2      Tb=T_motor/2',
    '     O──────────────A──────────────B──────(end)',
    '     ├── seg OA ─────┤──── seg AB ──┤',
    '     T_OA=T_motor=955N·m   T_AB=477N·m',
    '',
    '     TORQUE DIAGRAM:',
    '     T(N·m)',
    '     955│────────────────┐',
    '        │                │477',
    '        │                │──────────┐ 0',
    '        └────────O───────A──────────B─► x',
])
ap('Part (a) — Motor and shaft torques:', bold=True, sz_pt=10)
ap('  ω = 2π × 3000/60 = 100π rad/s = 314.16 rad/s', sz_pt=10)
ap('  T_motor = P/ω = 300,000 / 314.16 = 954.93 N·m ≈ 955 N·m', sz_pt=10)
ap('  Since Ta = Tb and Ta + Tb = T_motor:', sz_pt=10)
ap('    Ta = Tb = 955/2 = 477.5 N·m', sz_pt=10)
ap('  Segment OA carries full motor torque:  T_OA = 955 N·m', bold=True, sz_pt=10)
ap('  After pulley A removes Ta = 477.5 N·m:', sz_pt=10)
ap('  Segment AB:  T_AB = 955 − 477.5 = 477.5 N·m', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('Part (b) — Minimum solid shaft diameters:', bold=True, sz_pt=10)
ap('  τ_allow = 110 N/mm²', sz_pt=10)
ap('  d³ = 16T/(π·τ_allow)', sz_pt=10)
ap('  Segment OA:', sz_pt=10)
ap('    d_OA³ = 16×954,930/(π×110) = 15,278,880/345.58 = 44,211 mm³', sz_pt=10)
ap('    d_OA = ∛44,211 = 35.4 mm  →  use 36 mm', bold=True, sz_pt=10)
ap('  Segment AB:', sz_pt=10)
ap('    d_AB³ = 16×477,500/(π×110) = 7,640,000/345.58 = 22,107 mm³', sz_pt=10)
ap('    d_AB = ∛22,107 = 28.1 mm  →  use 29 mm', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('Part (c) — Hollow sections with d_o = 40 mm constant:', bold=True, sz_pt=10)
ap('  τ = T·(d_o/2) / J_h  →  J_h = T·(d_o/2)/τ_allow', sz_pt=10)
ap('  J_h = π(d_o⁴ − d_i⁴)/32  →  d_i⁴ = d_o⁴ − 32·J_h/π', sz_pt=10)
ap('  Segment OA:', sz_pt=10)
ap('    J_h_min = 955,000×20/110 = 173,636 mm⁴', sz_pt=10)
ap('    d_i_OA⁴ = 40⁴ − 32×173,636/π = 2,560,000 − 1,768,518 = 791,482', sz_pt=10)
ap('    d_i_OA_max = 791,482^(0.25) = 29.82 mm', bold=True, sz_pt=10)
ap('  Segment AB:', sz_pt=10)
ap('    J_h_min = 477,500×20/110 = 86,818 mm⁴', sz_pt=10)
ap('    d_i_AB⁴ = 2,560,000 − 32×86,818/π = 2,560,000 − 884,262 = 1,675,738', sz_pt=10)
ap('    d_i_AB_max = 1,675,738^(0.25) = 35.98 mm', bold=True, sz_pt=10)
ap('  Governing section: OA (smaller d_i_max = 29.82 mm)', sz_pt=10)
ap('  Recommended common inner diameter: 29 mm  (conservative, satisfies both sections)', bold=True, sz_pt=10)
ap('  Verification at d_i=29mm: J=π(40⁴−29⁴)/32=181,952mm⁴; '
   'τ_OA=955000×20/181952=105MPa<110✓; τ_AB=477500×20/181952=52.5MPa<110✓', sz_pt=9)

# ── LQ3 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ("Newton's Second Law", 'ΣF = m·a  (applied along and perpendicular to incline)'),
    ('Weight component along incline (down-slope)', 'F∥ = m·g·sinθ'),
    ('Normal force (perpendicular to incline)', 'N = m·g·cosθ'),
    ('Kinetic friction force', 'F_k = μ_k × N  (opposes relative motion)'),
    ('Check for relative slip between A and B', 'Compare f_AB_required vs f_AB_max = μ_k × N_AB'),
])
qhdr('Question 3 — Dynamics: Two Blocks on Inclined Plane  (20 marks)')
ap('If mA = 10 kg, mB = 40 kg, and the coefficient of kinetic friction between '
   'all surfaces is μk = 0.11, what is the acceleration of B down the inclined surface?\n\n'
   '[Figure shows block A resting on top of block B, which rests on a 30° inclined plane. '
   'Friction acts between A-B and between B-incline.]',
   sz_pt=10)
diag_block([
    '     FREE BODY DIAGRAMS — BLOCK A ON BLOCK B ON 30° INCLINE',
    '     ─────────────────────────────────────────────────────────────────',
    '     FBD Block A (10 kg):             FBD Block B (40 kg):',
    '     ↑ N_AB = m_A g cos30°            ↑ N_ramp = (m_A+m_B)g cos30°',
    '     [  A  ]                          [      B      ]',
    '     ↓ m_A g sin30°  (down slope)     ↓ (m_A+m_B)g sin30°  (down slope)',
    '     → f_AB (friction from B on A)    ← f_ramp = μ_k × N_ramp (from incline on B)',
    '       opposes A sliding on B          → f_AB reaction (A pushes B)',
    '     ─────────────────────────────────────────────────────────────────',
    '     SYSTEM VIEW:                θ = 30°',
    '          ┌─────┐',
    '          │  A  │  10 kg',
    '          ├─────┤──────────────────────────── (down slope direction)',
    '          │  B  │  40 kg',
    '          └─────┘',
    '         ─────────────────────────── μ_k=0.11 (incline-B interface)',
    '                  θ = 30°',
    '     ─────────────────────────────────────────────────────────────────',
    '     Key: μ_k = 0.11 between ALL surfaces (A-B AND B-incline)',
])
ap('Given: mA=10 kg, mB=40 kg, μk=0.11 (ALL surfaces), θ=30° (from figure), g=9.81 m/s²', sz_pt=10)
ap('Step 1 — Assume A and B move together as one unit:', sz_pt=10)
ap('  System mass = mA + mB = 50 kg', sz_pt=10)
ap('  N_ramp = 50 × 9.81 × cos30° = 50 × 9.81 × 0.866 = 424.8 N', sz_pt=10)
ap('  f_ramp = μk × N_ramp = 0.11 × 424.8 = 46.73 N  (opposes downward motion)', sz_pt=10)
ap('  Net force along incline:', sz_pt=10)
ap('    ΣF = (mA+mB)g sin30° − f_ramp = 50×9.81×0.5 − 46.73 = 245.25 − 46.73 = 198.52 N', sz_pt=10)
ap('    a = 198.52/50 = 3.97 m/s²', bold=True, sz_pt=10)
ap('Step 2 — Verify A does not slip on B (check friction required at A-B interface):', sz_pt=10)
ap('  N_AB = mA × g × cos30° = 10 × 9.81 × 0.866 = 84.97 N', sz_pt=10)
ap('  f_AB_max = μk × N_AB = 0.11 × 84.97 = 9.35 N', sz_pt=10)
ap('  For A alone: mA×a = mA×g×sin30° − f_AB', sz_pt=10)
ap('  f_AB_required = mA(g sin30° − a) = 10×(4.905−3.97) = 10×0.935 = 9.35 N', sz_pt=10)
ap('  f_AB_required (9.35 N) = f_AB_max (9.35 N)  → A and B are just at the slip limit.', sz_pt=10)
ap('  A and B move together (or just on the verge of slipping).', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('∴  Acceleration of B = 3.97 m/s² (down the incline)', bold=True, sz_pt=10)
ap('   Friction force between A and B = 9.35 N', sz_pt=10)

# ── LQ4 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Galvanic cell EMF (Nernst)', 'E_cell = E°_cathode − E°_anode  (> 0 → spontaneous)'),
    ('Anode oxidation half-reaction (Al corrodes)', 'Al → Al³⁺ + 3e⁻'),
    ('Cathode reduction half-reaction (Fe protected)', 'O₂ + 2H₂O + 4e⁻ → 4OH⁻'),
    ("Faraday's Law (mass corroded)", 'm = (M·I·t)/(n·F),  F = 96,485 C/mol'),
    ('Standard electrode potentials', 'Al: −0.76 V (anode);  Fe/C steel: −0.44 V (cathode)'),
])
qhdr('Question 4 — Materials Science: Corrosion of Al Rivet in Steel Joint  (20 marks)')
ap('Two thin sheets of 0.4 wt% C plain carbon steel are held together by an aluminium '
   'alloy rivet. List the possible types of corrosion that might arise. Suggest how '
   'corrosion might be minimised in such a situation.',
   sz_pt=10)
diag_block([
    '     GALVANIC CELL — Al RIVET + CARBON STEEL SHEETS',
    '     ─────────────────────────────────────────────────────────────────',
    '          e⁻ flow  ← ← ← ← ← ← ← ←',
    '     ┌──────────────────┐           ┌──────────────────────┐',
    '     │  ANODE (Al rivet)│           │ CATHODE (Steel sheet) │',
    '     │  Al→Al³⁺ + 3e⁻  │           │ O₂+2H₂O+4e⁻→4OH⁻   │',
    '     │  CORRODES         │           │ PROTECTED              │',
    '     │  E° = −0.76 V    │           │ E° = −0.44 V          │',
    '     └──────────────────┘           └──────────────────────┘',
    '           |  crevice at interface         |',
    '           └──────── ELECTROLYTE ──────────┘',
    '                 (moisture, condensation)',
    '     ─────────────────────────────────────────────────────────────────',
    '     ΔE = −0.44 − (−0.76) = +0.32 V → spontaneous galvanic cell',
    '     Al rivet (small anode) vs large steel cathode → ACCELERATED corrosion',
])
ap('(a) Types of Corrosion:', bold=True, sz_pt=10)
ap('1. Galvanic Corrosion:', bold=True, sz_pt=10)
ap('   Al alloy (E° ≈ −0.76 V) and 0.4 wt% C steel (Fe/C, E° ≈ −0.44 V) have different '
   'electrochemical potentials. In the presence of moisture (electrolyte), a galvanic cell forms: '
   'the Al rivet is the anode and corrodes preferentially. The small anode/large cathode area ratio '
   'accelerates the corrosion rate of the rivet severely.', sz_pt=10)
ap('   Anode (Al rivet):   Al → Al³⁺ + 3e⁻  (oxidation — corrosion)', sz_pt=10)
ap('   Cathode (Steel):     O₂ + 2H₂O + 4e⁻ → 4OH⁻  (reduction — protected)', sz_pt=10)
ap('2. Crevice Corrosion:', bold=True, sz_pt=10)
ap('   Stagnant moisture trapped in the crevice between the rivet head/shank and the '
   'steel sheet depletes oxygen inside the gap. The differential aeration cell (low O₂ '
   'inside vs high O₂ outside) creates a local low-pH, high-Cl⁻ environment that '
   'aggressively attacks the Al rivet.', sz_pt=10)
ap('3. Pitting Corrosion:', bold=True, sz_pt=10)
ap('   Chloride ions (from ambient humidity/salt) locally break down the passive Al₂O₃ '
   'film on the aluminium rivet, initiating deep pits that can penetrate through the rivet.', sz_pt=10)
ap('4. Stress Corrosion Cracking (SCC):', bold=True, sz_pt=10)
ap('   The rivet is under sustained tensile stress (fastening preload). Combined with '
   'the corrosive environment, SCC can cause intergranular or transgranular cracking '
   'through the rivet, especially in high-strength Al alloys (7xxx series).', sz_pt=10)
ap('5. Fretting Corrosion:', bold=True, sz_pt=10)
ap('   Micro-vibrations at the joint produce relative sliding between rivet and steel '
   'sheet, continuously disrupting protective oxide films and generating wear debris.', sz_pt=10)
ap('', sz_pt=6)
ap('(b) Mitigation Measures:', bold=True, sz_pt=10)
ap('1. Eliminate the galvanic couple: Replace the Al rivet with a material closer in '
   'the galvanic series to the steel (e.g., 316L stainless steel or Monel rivets).', sz_pt=10)
ap('2. Insulating separation: Install nylon/PTFE/neoprene bushings and washers between '
   'the Al rivet and steel sheets to break the electrical circuit.', sz_pt=10)
ap('3. Protective coatings: Apply zinc-rich primer or epoxy coat to the steel sheets. '
   'Anodise the Al rivet to thicken the Al₂O₃ passive film.', sz_pt=10)
ap('4. Sealant at interfaces: Apply corrosion-inhibiting sealant (polysulfide or '
   'epoxy) at all rivet-sheet interfaces to exclude moisture.', sz_pt=10)
ap('5. Material selection: If Al must be used, choose alloy 5083 or 6061-T6 '
   '(better corrosion resistance than 2024 or 7075).', sz_pt=10)
ap('∴  Primary risk: galvanic corrosion accelerated by unfavourable area ratio (small anode). '
   'Mitigation: replace rivet material or insulate the dissimilar metal contact.', bold=True, sz_pt=10)

# ── LQ5 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Plant transfer function (force → velocity)', 'G(s) = V(s)/F(s) = 1/(m·s) = 1/(1000s)'),
    ('PI controller', 'C(s) = K_p + K_i/s = (K_p·s + K_i)/s'),
    ('Closed-loop characteristic equation', '1 + C(s)G(s) = 0  →  1000s² + K_p·s + K_i = 0'),
    ('Standard 2nd-order matching', 'K_p = 2ζω_n·m ;  K_i = ω_n²·m'),
    ('Steady-state driving force (constant speed on incline)', 'F_ss = m·g·sinθ = 1000×9.81×sin20° = 3,355 N'),
    ('Target speed conversion', 'v_set = 90 km/h = 90/3.6 = 25 m/s'),
])
qhdr('Question 5 — Control Systems: PI Controller for Vehicle Speed  (20 marks)')
ap('It is desired to have a car of mass 1,000 kg go up the inclined plane shown below '
   'at a constant speed of 90 km/h. The car is initially at rest. Design a controller '
   'that will output force F to push the car up the inclined plane and maintain the '
   'desired constant speed. (Incline angle = 20°)\n\n'
   'State any assumptions or requirements to ensure the controller works well.',
   sz_pt=10)
diag_block([
    '     FREE BODY DIAGRAM — CAR ON 20° INCLINE',
    '     ─────────────────────────────────────────────────────────────────',
    '                 ↑ N = mg cos20°',
    '              ┌──────┐',
    '              │ Car  │──────────────────► F (PI controller output)',
    '              │1000kg│',
    '              └──────┘',
    '            ↙ mg sin20° = 3,355 N (gravity component down-slope)',
    '     ────────────────────────────── θ = 20°',
    '     ─────────────────────────────────────────────────────────────────',
    '     CONTROL BLOCK DIAGRAM:',
    '                      ┌──────────┐    F(s)  ┌────────────┐',
    '     v_set(25m/s)→ Σ──┤ PI C(s)  ├──────────┤ 1/(1000s)  ├─► v(s)',
    '                   ↑  └──────────┘           └────────────┘',
    '                  −|                                  |',
    '                   └──────────────────────────────────┘',
    '     PI integrator → zero steady-state error for step speed input ✓',
])
ap('Assumptions: (1) Aerodynamic drag neglected. (2) Rolling resistance neglected. '
   '(3) Plant is linear (Newton 2nd Law — force to velocity integrator). '
   '(4) Sensor measures vehicle speed continuously.', sz_pt=10)
ap('', sz_pt=6)
ap('Step 1 — Plant model (Newton 2nd Law along incline):', sz_pt=10)
ap('  F − mg sinθ = m·a  →  F_net = m·a', sz_pt=10)
ap('  Taking Laplace: F_net(s) = m·s·V(s)', sz_pt=10)
ap('  G(s) = V(s)/F_net(s) = 1/(m·s) = 1/(1000s)', sz_pt=10)
ap('  [F_net = F_drive − F_gravity; F_gravity = mg sin20° = 3,355 N treated as disturbance]', sz_pt=10)
ap('', sz_pt=6)
ap('Step 2 — Steady-state force requirement:', sz_pt=10)
ap('  At constant speed (a=0): F_drive = m·g·sinθ = 1000×9.81×sin20° = 3,355 N', sz_pt=10)
ap('  Target: v_set = 90 km/h = 25 m/s', sz_pt=10)
ap('', sz_pt=6)
ap('Step 3 — PI controller design:', sz_pt=10)
ap('  C(s) = K_p + K_i/s = (K_p·s + K_i)/s', sz_pt=10)
ap('  Open-loop: C(s)G(s) = (K_p·s + K_i)/(1000s²)', sz_pt=10)
ap('  Closed-loop characteristic equation: 1000s² + K_p·s + K_i = 0', sz_pt=10)
ap('  Dividing: s² + (K_p/1000)s + (K_i/1000) = 0', sz_pt=10)
ap('  Match to s² + 2ζω_n·s + ω_n² = 0:', sz_pt=10)
ap('', sz_pt=6)
ap('Step 4 — Choose performance: ζ = 1 (critically damped), ω_n = 2 rad/s:', sz_pt=10)
ap('  K_p = 2ζω_n × m = 2×1×2×1000 = 4,000 N·s/m', bold=True, sz_pt=10)
ap('  K_i = ω_n² × m = 4×1000 = 4,000 N/m', bold=True, sz_pt=10)
ap('  Rise time ≈ 1.8/ω_n = 0.9 s  (fast enough for the 25 m/s setpoint)', sz_pt=10)
ap('', sz_pt=6)
ap('Step 5 — Steady-state performance:', sz_pt=10)
ap('  PI controller includes an integrator → Type 2 open-loop system → '
   'zero steady-state error for step (constant speed) input. '
   'The integral action accumulates error until F exactly equals 3,355 N at cruise.', sz_pt=10)
ap('∴  F_ss = 3,355 N  |  K_p = 4,000 N·s/m  |  K_i = 4,000 N/m  |  ζ=1 (critically damped)', bold=True, sz_pt=10)

# ── LQ6 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Carnot efficiency — minimum fuel implies Carnot', 'η_c = 1 − T_L/T_H  (T in Kelvin)'),
    ('Heat input to plant', 'Q = W_net/η_c  →  Q = P_net/η_c'),
    ('Mass flow rate of fuel', 'ṁ_fuel = Q/CV  (CV = calorific value, MJ/kg)'),
    ('Daily fuel consumption', 'm_daily = ṁ_fuel × 86,400 s/day'),
])
qhdr('Question 6 — Thermodynamics: Steam Power Plant Fuel Consumption  (20 marks)')
ap('A 750-MW steam power plant burns fuel of calorific value 44 MJ/kg and generates '
   'electricity. The boiler operates at 560°C and the condenser releases waste heat '
   'at 30°C. Calculate the minimum daily fuel consumption of the plant.\n\n'
   '(Minimum fuel consumption implies Carnot efficiency.)',
   sz_pt=10)
diag_block([
    '     CARNOT HEAT ENGINE — BLOCK DIAGRAM AND T-s DIAGRAM',
    '     ─────────────────────────────────────────────────────────────────',
    '     T_H = 560+273 = 833 K      T-s DIAGRAM:',
    '          ↓ Q_H                  T │  833K  1────────2',
    '     ┌────────────────┐              │      4│         │2',
    '     │  STEAM PLANT   ├── W=750MW     │      │  cycle  │',
    '     │  η_c = 63.6%  │              │  303K 3────────4',
    '     └────────┬───────┘              └─────────────────► s',
    '              ↓ Q_c (rejected)',
    '     T_L = 30+273 = 303 K',
    '     ─────────────────────────────────────────────────────────────────',
    '     η_c = 1 − 303/833 = 0.636   (Carnot — maximum efficiency)',
    '     Q_H = 750/0.636 = 1179 MW',
    '     ṁ_fuel = 1179/44 = 26.8 kg/s',
    '     m_daily = 26.8 × 86,400 = 2.315 × 10⁶ kg/day',
])
ap('Given: P_net = 750 MW, T_H = 560+273 = 833 K, T_L = 30+273 = 303 K, CV = 44 MJ/kg', sz_pt=10)
ap('Step 1 — Carnot efficiency:', sz_pt=10)
ap('  η_c = 1 − T_L/T_H = 1 − 303/833 = (833−303)/833 = 530/833 = 0.636  (63.6%)', sz_pt=10)
ap('Step 2 — Heat input rate Q_H:', sz_pt=10)
ap('  Q_H = W_net/η_c = 750 MW / 0.636 = 1,179 MW', sz_pt=10)
ap('Step 3 — Mass flow rate of fuel:', sz_pt=10)
ap('  ṁ_fuel = Q_H/CV = 1,179 MJ/s ÷ 44 MJ/kg = 26.80 kg/s', sz_pt=10)
ap('Step 4 — Daily fuel consumption:', sz_pt=10)
ap('  m_daily = 26.80 × 86,400 = 2,315,520 kg/day', bold=True, sz_pt=10)
ap('  m_daily ≈ 2.316 × 10⁶ kg/day', bold=True, sz_pt=10)
note('The PDF solution uses η_c = 0.636 and obtains Q = 1195 MW and m_daily = 2.347×10⁶ kg. '
     'The small difference arises from rounding η_c. The exact calculation gives 2.316×10⁶ kg/day.')

# ── LQ7 ───────────────────────────────────────────────────────────────────────
sep()
eq_block([
    ('Volume of material removed (cylindrical turning)', 'V = π/4 × (D_stock² × L_stock − D_shaft² × L_shaft)'),
    ('Available cutting power', 'P_cut = P_motor × η_machine = 10 × 0.80 = 8 kW'),
    ('Material Removal Rate from power', 'MRR_available = P_cut / p_s  [m³/s]'),
    ('MRR for lathe turning', 'MRR = V_c [m/s] × f [m/rev] × d_oc [m]'),
    ('Cutting speed', 'V_c = π·D·N/1000  [m/min]  or  /60000 [m/s]'),
    ('Net machining time constraint', 't_machining = t_total − t_load_unload = 15 − 2 = 13 min'),
    ('Minimum MRR to meet time', 'MRR_required = V_removed / t_machining'),
])
qhdr('Question 7 — Manufacturing: Lathe Turning — Cutting Conditions & Tool Selection  (20 marks)')
ap('Steel bar stock of 200 mm length and 90 mm diameter is to be machined to a shaft of '
   '160 mm length and 80 mm diameter using a lathe with a 10 kW motor running at '
   '80% efficiency. The total machining time (roughing + finishing) must not exceed 15 min. '
   'Specific cutting energy p_s = 2.73 GJ/m³. Load/unload time = 2 min.\n\n'
   'Select proper cutting conditions, tool materials, and rake angles for both '
   'roughing and finishing operations.',
   sz_pt=10)
diag_block([
    '     LATHE TURNING — MATERIAL REMOVAL SCHEMATIC',
    '     ─────────────────────────────────────────────────────────────────',
    '     Bar stock (before):          Final shaft (after):',
    '     ┌─────────────────────────┐  ┌──────────────────┐',
    '     │  D=90mm × L=200mm        │  │ D=80mm × L=160mm  │',
    '     └─────────────────────────┘  └──────────────────┘',
    '     Radial stock to remove: (90−80)/2 = 5 mm per side',
    '     Axial stock to remove: 200−160 = 40 mm (facing)',
    '     ─────────────────────────────────────────────────────────────────',
    '     Turning operation:',
    '     Tool → ───────────────────────► feed direction',
    '             /|  d_oc (depth of cut)',
    '     D=90mm / |   f (feed/rev)',
    '     ───────   D_final=80mm',
    '     N rpm (spindle)',
])
ap('Step 1 — Volume of material removed:', sz_pt=10)
ap('  V_stock = π/4 × 90² × 200 = π/4 × 8,100 × 200 = 1,272,345 mm³', sz_pt=10)
ap('  V_shaft = π/4 × 80² × 160 = π/4 × 6,400 × 160 = 803,541 mm³', sz_pt=10)
ap('  V_removed = 1,272,345 − 803,541 = 468,804 mm³ = 4.688×10⁻⁴ m³', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('Step 2 — Available machine power for cutting:', sz_pt=10)
ap('  P_cut = P_motor × η = 10 kW × 0.80 = 8 kW = 8,000 W', sz_pt=10)
ap('', sz_pt=6)
ap('Step 3 — MRR available from machine power:', sz_pt=10)
ap('  MRR_available = P_cut / p_s = 8,000 / (2.73×10⁹) = 2.930×10⁻⁶ m³/s = 2,930 mm³/s', bold=True, sz_pt=10)
ap('', sz_pt=6)
ap('Step 4 — MRR required to meet 13-minute cutting time:', sz_pt=10)
ap('  t_available = 15 − 2 = 13 min = 780 s', sz_pt=10)
ap('  MRR_required = V_removed / t_available = 468,804 / 780 = 601 mm³/s = 6.01×10⁻⁷ m³/s', sz_pt=10)
ap('  Since MRR_required (601 mm³/s) < MRR_available (2,930 mm³/s): machine is NOT the limiting constraint ✓', bold=True, sz_pt=10)
ap('  The job is feasible within 13 min even at partial machine power.', sz_pt=10)
ap('', sz_pt=6)
ap('Step 5 — Select cutting conditions:', bold=True, sz_pt=10)
ap('  Radial stock = (90−80)/2 = 5 mm. Strategy: roughing pass (4 mm) + finishing pass (1 mm).', sz_pt=10)
ap('', sz_pt=6)
ap('  ROUGHING (D_avg ≈ 85 mm):', bold=True, sz_pt=10)
ap('  Target MRR ≈ 2,500 mm³/s (within power limit)', sz_pt=10)
ap('  Choose: depth d = 4 mm, feed f = 0.30 mm/rev', sz_pt=10)
ap('  V_c = MRR / (f × d) = 2,500 / (0.30 × 4) = 2,083 mm²/s = 2.083×10⁻³ m²/s', sz_pt=10)
ap('  Wait — V_c = MRR/(f[mm]×d[mm]) mm/s = 2500/1.2 = 2,083 mm/s = 125 m/min', sz_pt=10)
ap('  N = V_c×1000/(π×D) = 125,000/(π×85) = 468 rpm', sz_pt=10)
ap('  Power check: P = p_s × MRR = 2.73×10⁹ × 2.5×10⁻⁶ = 6,825 W < 8,000 W ✓', sz_pt=10)
ap('', sz_pt=6)
ap('  FINISHING (D ≈ 80 mm):', bold=True, sz_pt=10)
ap('  Choose: depth d = 1 mm, feed f = 0.10 mm/rev, V_c = 200 m/min (higher speed for finish)', sz_pt=10)
ap('  N = 200,000/(π×80) = 796 rpm', sz_pt=10)
ap('  MRR = 200,000/60 × 0.1 × 1 = 333 mm³/s', sz_pt=10)
ap('  P = 2.73×10⁹ × 333×10⁻⁹ = 909 W ≪ 8,000 W ✓  (low power, high quality)', sz_pt=10)
ap('', sz_pt=6)
ap('Step 6 — Tool material and rake angle selection:', bold=True, sz_pt=10)
ap('  ROUGHING — Tool: Coated carbide insert (TiAlN or TiCN coating)', bold=True, sz_pt=10)
ap('  • Reason: High hardness, good wear resistance at elevated temperatures; '
   'handles interrupted cuts and high feed rates on carbon steel.', sz_pt=10)
ap('  • Rake angle: NEGATIVE rake (−5° to −7°)', sz_pt=10)
ap('  • Reason: Stronger cutting edge geometry; negative rake redistributes cutting '
   'forces into compression on the insert (prevents chipping at high chip loads).', sz_pt=10)
ap('  FINISHING — Tool: Coated carbide (TiN coating) or CBN insert', bold=True, sz_pt=10)
ap('  • Reason: Fine grain carbide or CBN gives superior surface finish; '
   'TiN coating reduces friction and built-up edge formation on mild steel.', sz_pt=10)
ap('  • Rake angle: POSITIVE rake (+5° to +10°)', sz_pt=10)
ap('  • Reason: Positive rake reduces cutting forces and heat generation, '
   'produces better surface finish (Ra ≤ 1.6 μm target).', sz_pt=10)
ap('  Coolant: Flood coolant (soluble oil emulsion) for both operations — '
   'controls thermal expansion, prevents BUE, extends tool life.', sz_pt=10)
ap('', sz_pt=6)
ap('∴  V_removed = 4.688×10⁻⁴ m³  |  MRR_avail = 2,930 mm³/s  |  Job feasible in 13 min', bold=True, sz_pt=10)
ap('   Roughing: d=4mm, f=0.30mm/rev, V_c=125m/min, TiAlN carbide, −6° rake', bold=True, sz_pt=10)
ap('   Finishing: d=1mm, f=0.10mm/rev, V_c=200m/min, TiN/CBN, +7° rake', bold=True, sz_pt=10)

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
sep()
ap('SUMMARY OF ANSWERS — QUICK REFERENCE', bold=True, sz_pt=12, before=20, after=20)
ap('Part 1 — Multiple Choice Questions', bold=True, sz_pt=11, before=10, after=10)
ap('Q1 (Turbulence):     (d) All of the above', sz_pt=10, mono=True)
ap('Q2 (Beam strains):   (a) NA=71.43mm, σ_axial=30MPa  [σ_bend=70MPa by calculation]', sz_pt=10, mono=True)
ap('Q3 (Torsion):        (b) 3.94°', sz_pt=10, mono=True)
ap('Q4 (Inclusions):     (d) All of the above', sz_pt=10, mono=True)
ap('Q5 (Bullet-block):   (a) 198.5 m/s', sz_pt=10, mono=True)
ap('Q6 (Rail kinematics):(a) 0.30 hr  [t=1050s, v_top=100 m/s, accel+cruise+decel]', sz_pt=10, mono=True)
ap('Q7 (Control):        (a) Overdamped  [ζ=1.061]', sz_pt=10, mono=True)
ap('Q8 (Cyclic process): (a) Equal to zero', sz_pt=10, mono=True)
ap('Q9 (Adiabatic):      (d) (iii) and (iv)', sz_pt=10, mono=True)
ap('Q10 (Taylor):        (b) Tool life decreases', sz_pt=10, mono=True)
ap('', sz_pt=6)
ap('Part 2 — Long Questions', bold=True, sz_pt=11, before=10, after=10)
ap('LQ1 (Pipe flow):     V₁=5.69m/s, Q=44.7L/s  [two-pipe, Fanning f=0.01]', sz_pt=10, mono=True)
ap('LQ2 (Stepped shaft): T_OA=955N·m, T_AB=478N·m; d_OA=36mm, d_AB=29mm; hollow d_i=29mm', sz_pt=10, mono=True)
ap('LQ3 (Dynamics):      a=3.97m/s², f_AB=9.35N  [mA=10kg, mB=40kg, θ=30°, μk=0.11]', sz_pt=10, mono=True)
ap('LQ4 (Corrosion):     Galvanic+crevice+pitting+SCC; fix: insulate+coat+compatible material', sz_pt=10, mono=True)
ap('LQ5 (PI control):    F_ss=3355N; K_p=4000N·s/m, K_i=4000N/m; v_set=90km/h=25m/s', sz_pt=10, mono=True)
ap('LQ6 (Thermodynamics):η_c=63.6%, Q_H=1179MW, ṁ=26.8kg/s, m_daily=2.316×10⁶kg', sz_pt=10, mono=True)
ap('LQ7 (Machining):     V_rem=4.69×10⁻⁴m³; MRR=2930mm³/s; feasible in 13min', sz_pt=10, mono=True)

sep()
ap('— HyESys Agent', bold=True, sz_pt=10, align='center', before=20, after=20)

doc.save(DST)
print(f'Saved: {DST}')
