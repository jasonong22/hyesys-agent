"""
Creates a clean summary Word document of the HyESys HySBatt Datasheet.
Output: HyESys_HySBatt_Datasheet_Summary.docx in the same folder as the source.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\2024 HyESys\Pre-hyesys Projects\Datasheet (HyESys, HySBatt)"
    r"\HyESys_HySBatt_Datasheet_Summary.docx"
)

FONT        = "Calibri"
CLR_DARK    = RGBColor(0x1F, 0x37, 0x63)   # dark navy — headings
CLR_ACCENT  = RGBColor(0x2E, 0x75, 0xB6)   # mid blue — sub-headings
CLR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CLR_BODY    = RGBColor(0x26, 0x26, 0x26)
CLR_TBL_HDR = RGBColor(0x2E, 0x75, 0xB6)
CLR_TBL_ALT = RGBColor(0xDF, 0xEB, 0xF7)   # light blue row stripe


# ── helpers ───────────────────────────────────────────────────────────────────

def set_run(run, size, bold=False, italic=False, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)


def add_para(doc, text, size=11, bold=False, italic=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size, bold=bold, italic=italic, color=color or CLR_BODY)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = add_para(doc, text, size=15, bold=True, color=CLR_DARK,
                     space_before=14, space_after=4)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        add_para(doc, text, size=12, bold=True, color=CLR_ACCENT,
                 space_before=8, space_after=2)


def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, size=10, bold=False, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_run(run, size, bold=bold, color=color or CLR_BODY)


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for ci, h in enumerate(headers):
        cell = hdr_row.cells[ci]
        shade_cell(cell, "2E75B6")
        set_cell_text(cell, h, size=10, bold=True, color=CLR_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        fill = "DFF0F7" if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            set_cell_text(cell, str(val), size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return tbl


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run(run, size, color=CLR_BODY)


# ── build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover / Title ────────────────────────────────────────────────────────
    add_para(doc, "HyESys™ & HySBatt", size=22, bold=True, color=CLR_DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    add_para(doc, "Product Datasheet Summary", size=14, bold=False,
             color=CLR_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Advancer Smart Technology Pte Ltd  |  Issue: May 2026  |  Version 2",
             size=9, italic=True, color=RGBColor(0x70, 0x70, 0x70),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    doc.add_paragraph()

    # ── 1. Product Overview ──────────────────────────────────────────────────
    add_heading(doc, "1.  Product Overview")
    add_para(doc,
        "HyESys™ is a modular hybrid energy system integrating ultra-safe HySBatt lithium-ion "
        "battery packs with a high-efficiency Power Conversion System (PCS). It simultaneously "
        "delivers reactive power compensation (kVAr injection), 3-phase load balancing, and "
        "energy storage / solar load shaving from a single unit.",
        size=10.5, space_after=4)
    add_para(doc,
        "The system is governed by a Smart Energy Management System (SEMS), liquid-based thermal "
        "management, passive fire suppression, and a patent-pending Gaseous Management Unit (GMU) "
        "(Singapore Patent No. 10202302370U).",
        size=10.5, space_after=6)

    add_heading(doc, "Key Compliance & Certifications", level=2)
    for item in [
        "NFPA 855 — energy storage system fire safety",
        "SCDF 2023 — Singapore Civil Defence Force regulations",
        "UL9540A — unit-level thermal runaway testing",
        "BS476-22 — 2-hour fire rating",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Installation Options", level=2)
    for item in [
        "Wall-mounted",
        "Stacked",
        "Floor-mounted",
        "Caster-mounted",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()

    # ── 2. Available Models ──────────────────────────────────────────────────
    add_heading(doc, "2.  Available Models")
    add_para(doc,
        "All models share the same AC grid interface. The number of HySBatt packs shown is the "
        "minimum required; final installation may need additional packs based on site load profile.",
        size=10.5, space_after=6)

    headers = [
        "Model", "Power\n(kVA)", "Max Current\n(A)", "VDC Operating\nRange",
        "HySBatt Packs\n(min)", "Usable Energy\n(kWh)",
        "DC Voltage\n(Min / Max)", "Weight\n(kg)", "Area\n(m²)"
    ]
    rows = [
        ["H30",  "30",  "43.5", "210 – 850 V",  "7",  "69.3",  "231 V / 269.5 V",  "1,400", "2.1"],
        ["H50",  "50",  "72.5", "350 – 850 V",  "11", "108.9", "363 V / 423.5 V",  "2,200", "3.2"],
        ["H60",  "60",  "87",   "420 – 850 V",  "14", "138.6", "462 V / 539 V",    "2,800", "4.2"],
        ["H100", "100", "145",  "680 – 900 V",  "22", "217.8", "726 V / 847 V",    "4,400", "6.3"],
        ["H125", "125", "181",  "680 – 900 V",  "22", "217.8", "726 V / 847 V",    "4,400", "6.3"],
    ]
    add_table(doc, headers, rows,
              col_widths=[1.5, 1.3, 1.5, 2.2, 1.6, 1.6, 2.2, 1.4, 1.2])

    add_para(doc,
        "Note: H100 and H125 share the same battery pack count (22) and usable energy (217.8 kWh). "
        "The H125 delivers higher kVA output from the same battery configuration.",
        size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55), space_after=10)

    # ── 3. HySBatt Pack Specifications ──────────────────────────────────────
    add_heading(doc, "3.  HySBatt Pack Specifications")
    add_para(doc,
        "Each HySBatt pack is an independent, modular lithium-ion battery unit. Multiple packs "
        "are assembled to meet the energy requirements of each HyESys model.",
        size=10.5, space_after=6)

    headers2 = ["Parameter", "Specification"]
    rows2 = [
        ["Dimensions (H × W × D)", "1,250 mm × 500 mm × 550 mm"],
        ["Weight",                  "< 200 kg"],
        ["Usable Energy",           "10 kWh per pack"],
        ["Nominal Voltage",         "35 V"],
        ["Charge Cutoff Current",   "120 A"],
        ["Ingress Protection",      "IP54"],
        ["Min. Installation Gap",   "50 mm between packs"],
        ["Installation Options",    "Raised floor / frame, wall-hung, rack assembly, caster"],
        ["Relative Humidity",       "0% – 90% (non-condensing)"],
        ["Operating Temp (Discharge)", "-20 °C to +50 °C"],
        ["Operating Temp (Charge)", "0 °C to +45 °C"],
        ["Storage Temperature",     "10 °C to 35 °C"],
    ]
    add_table(doc, headers2, rows2, col_widths=[6, 10])

    # ── 4. Communication & Control ───────────────────────────────────────────
    add_heading(doc, "4.  Communication & Control")

    headers3 = ["Interface", "Specification"]
    rows3 = [
        ["BMS Communication",  "CAN or RS485 (supports mainstream BMS protocols)"],
        ["EMS Communication",  "RS485 (Modbus RTU)"],
        ["Dry Contacts",       "2× DI (emergency stop, reserved)  |  2× DO (alarm, status)"],
        ["Parallel Operation", "Up to 15 units AC side; dedicated parallel ports with termination resistors"],
        ["Status LEDs",        "Run, Alarm, Battery, Grid"],
        ["Debug Port",         "Internal DEBUG port"],
    ]
    add_table(doc, headers3, rows3, col_widths=[5, 11])

    # ── 5. Safety & Thermal Management ──────────────────────────────────────
    add_heading(doc, "5.  Integrated Safety & Thermal Management")

    headers4 = ["Feature", "Description"]
    rows4 = [
        ["Thermal Management",
         "High-performance liquid cooling for industrial reliability."],
        ["Pressure Relief",
         "Automatic pressure release — prevents hazardous pressure build-up inside the battery enclosure."],
        ["Gaseous Management Unit (GMU)",
         "Patent-pending system that converts flammable vent gases into non-flammable water vapour and CO₂. "
         "(Singapore Patent No. 10202302370U; additional patents pending.)"],
        ["Fire Suppression",
         "Passive, self-actuating suppression triggered at 170 °C — requires no external power."],
    ]
    add_table(doc, headers4, rows4, col_widths=[5, 11])

    # ── 6. Compliance & Standards ────────────────────────────────────────────
    add_heading(doc, "6.  Compliance & Standards")

    headers5 = ["Component", "Standard / Certification"]
    rows5 = [
        ["HyESys System",  "NFPA 855, SCDF 2023"],
        ["HySBatt Pack",   "UL9540A (unit level), BS476-22 (2-hour fire rated)"],
        ["System Patents", "Singapore Patent No. 10202302370U; additional GMU patents pending"],
    ]
    add_table(doc, headers5, rows5, col_widths=[5, 11])

    # ── 7. Electrical Interface Notes ────────────────────────────────────────
    add_heading(doc, "7.  Electrical Interface Notes")
    for item in [
        "PCS operating VDC ranges are nominal; absolute limits are as shown in Section 2.",
        "All models support AC grid connection via standard low-voltage interface (specific grid codes available on request).",
        "System includes integral EMS controller for real-time power conversion and HySBatt management.",
        "Cycle life and warranty details available upon request.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()

    # ── Footer note ──────────────────────────────────────────────────────────
    add_para(doc,
        "This document is a summary of the HyESys™ HySBatt Datasheet (Version 2, May 2026). "
        "For full technical specifications, refer to the original datasheet. "
        "Specifications subject to change without notice.",
        size=8.5, italic=True, color=RGBColor(0x80, 0x80, 0x80),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    build()
