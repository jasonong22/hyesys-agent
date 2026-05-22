"""
Creates the HyESys Parameter & Equation Reference Sheet.
A single Word document covering every parameter and equation across
global config, per-site config, Agent 1, and Agent 2.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\2024 HyESys\Software\claude code"
    r"\HyESys_Parameter_Equation_Reference.docx"
)

FONT       = "Calibri"
FONT_MONO  = "Courier New"
C_NAVY     = RGBColor(0x1F, 0x37, 0x63)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_ORANGE   = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_PURPLE   = RGBColor(0x70, 0x30, 0xA0)
C_RED      = RGBColor(0xC0, 0x00, 0x00)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY     = RGBColor(0x1A, 0x1A, 0x1A)
C_GREY     = RGBColor(0x59, 0x59, 0x59)

# Table header colours per section
HDR_GLOBAL  = "1F3763"
HDR_SITE    = "375623"
HDR_A1      = "843C0C"
HDR_A2      = "2F5496"
HDR_SAVINGS = "7030A0"
ALT_ROW     = "F2F7FF"
ALT_ROW2    = "FFF4EF"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run(run, size, bold=False, italic=False, color=None, mono=False):
    run.font.name  = FONT_MONO if mono else FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    f = FONT_MONO if mono else FONT
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), f)


def para(doc, text="", size=10.5, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, mono=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run = p.add_run(text)
        set_run(run, size, bold=bold, italic=italic,
                color=color or C_BODY, mono=mono)
    return p


def heading1(doc, text):
    p = para(doc, text, size=13, bold=True, color=C_NAVY, sb=14, sa=3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading2(doc, text, color=None):
    para(doc, text, size=11, bold=True, color=color or C_BLUE, sb=8, sa=2)


def heading3(doc, text, color=None):
    para(doc, text, size=10.5, bold=True, italic=True,
         color=color or C_GREY, sb=5, sa=1)


def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def cell_text(cell, text, size=9.5, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_run(run, size, bold=bold, color=color or C_BODY, mono=mono)


def make_table(doc, headers, rows, col_widths, hdr_color,
               alt_color=ALT_ROW, hdr_size=9.5, row_size=9.5):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        shade_cell(c, hdr_color)
        cell_text(c, h, size=hdr_size, bold=True,
                  color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        fill = alt_color if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tbl.rows[ri + 1].cells[ci]
            shade_cell(c, fill)
            mono = ci >= 2 and ("=" in str(val) or "√" in str(val) or "×" in str(val))
            cell_text(c, str(val), size=row_size, mono=False)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Widths
    for row in tbl.rows:
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def equation_block(doc, label, eq, note=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{label}:  ")
    set_run(r1, 9.5, bold=True, color=C_NAVY)
    r2 = p.add_run(eq)
    set_run(r2, 9.5, mono=True, color=C_BODY)
    if note:
        r3 = p.add_run(f"    ← {note}")
        set_run(r3, 8.5, italic=True, color=C_GREY)


# ── Document ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    # ── COVER ─────────────────────────────────────────────────────────────────
    para(doc, "HyESys Agent System", size=20, bold=True, color=C_NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sb=10, sa=2)
    para(doc, "Parameter & Equation Reference Sheet", size=14, color=C_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, "All-Sites Config  ·  Per-Site Config  ·  Agent 1  ·  Agent 2  ·  Savings",
         size=9.5, italic=True, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para(doc, "Advancer Smart Technology Pte Ltd  |  HyESys Dept  |  May 2026  |  Confidential",
         size=8.5, italic=True, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=14)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. GLOBAL CONFIG
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "1.  Global Config  ( config/global.yaml )")
    para(doc, "Fixed across every HyESys site. Sourced from core/schema.py and agent logic.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    # 1a. Power factor & voltage targets
    heading2(doc, "1a.  Power Factor & Voltage Constants", color=C_NAVY)
    make_table(doc,
        headers=["Parameter", "Variable", "Value", "Unit", "Description"],
        rows=[
            ["PF Target",              "PF_TARGET",           "0.98",    "–",  "Correction target. Not 1.0 — law of convergence makes unity impractical"],
            ["SP Penalty Threshold",   "PF_PENALTY_THRESHOLD","0.85",    "–",  "SP surcharge triggered when |PF| falls below this value"],
            ["THD Assumption",         "THD_ASSUMPTION",      "0.15",    "–",  "15% THD baseline for mixed buildings. Used in loss estimates when THD not measured"],
            ["Nominal Voltage",        "VOLTAGE_NOMINAL",     "230.0",   "V",  "Singapore LV standard line-to-neutral"],
            ["Voltage Min",            "VOLTAGE_MIN",         "195.5",   "V",  "230 × 0.85  (−15% tolerance)"],
            ["Voltage Max",            "VOLTAGE_MAX",         "264.5",   "V",  "230 × 1.15  (+15% tolerance)"],
        ],
        col_widths=[3.8, 4.2, 1.6, 1.0, 7.2],
        hdr_color=HDR_GLOBAL
    )

    # 1b. Reward weights
    heading2(doc, "1b.  Reward Weights  ( Agent 2 SAR Loop )", color=C_NAVY)
    make_table(doc,
        headers=["Parameter", "Variable", "Value", "Description"],
        rows=[
            ["PF Reward Weight",   "W_PF",   "0.60", "60% weight on PF improvement — directly affects SP penalty"],
            ["Loss Reward Weight", "W_LOSS", "0.40", "40% weight on I²R loss fraction — quantifies kWh savings"],
        ],
        col_widths=[4.0, 3.0, 2.0, 8.8],
        hdr_color=HDR_GLOBAL
    )

    # 1c. Outcome thresholds
    heading2(doc, "1c.  SAR Outcome Classification Thresholds", color=C_NAVY)
    make_table(doc,
        headers=["Outcome", "Variable", "Condition", "Meaning"],
        rows=[
            ["POSITIVE", "REWARD_POSITIVE_PF_DELTA", "r_PF  ≥  +0.01", "Measurable PF improvement — action was beneficial"],
            ["NEUTRAL",  "—",                        "−0.01 < r_PF < +0.01", "Below detection threshold — no significant change"],
            ["NEGATIVE", "REWARD_NEGATIVE_PF_DELTA", "r_PF  ≤  −0.01", "PF degraded — action was counterproductive"],
        ],
        col_widths=[2.8, 5.0, 4.2, 5.8],
        hdr_color=HDR_GLOBAL
    )

    # 1d. Demand tariff
    heading2(doc, "1d.  Demand Tariff  ( Agent 2 Demand Risk )", color=C_NAVY)
    make_table(doc,
        headers=["Parameter", "Variable", "Value", "Unit", "Description"],
        rows=[
            ["Demand Tariff Rate", "DEMAND_TARIFF_SGD_PER_KW", "10.0", "SGD/kW/month",
             "Proxy for Singapore LV demand charge. Used to estimate cost at risk when current kW approaches peak"],
        ],
        col_widths=[3.8, 5.0, 1.6, 3.2, 4.2],
        hdr_color=HDR_GLOBAL
    )

    # 1e. Saving priority
    heading2(doc, "1e.  Saving Priority Order", color=C_NAVY)
    make_table(doc,
        headers=["Priority", "Function", "Variable", "Reason"],
        rows=[
            ["1 — Highest", "Solar Storage",       "solar_storage",       "Maximises ROI — free energy captured"],
            ["2",           "Reactive Correction", "reactive_correction", "Eliminates SP penalty and I²R losses"],
            ["3 — Lowest",  "Load Balancing",      "load_balancing",      "Reduces neutral I²R — smaller but real saving"],
        ],
        col_widths=[2.5, 3.8, 4.2, 7.3],
        hdr_color=HDR_GLOBAL
    )

    # 1f. HyESys Models
    heading2(doc, "1f.  HyESys Hardware Models  ( HYESYS_MODELS )", color=C_NAVY)
    make_table(doc,
        headers=["Model", "kVA", "Max I (A)", "VDC Range", "Packs (min)", "Energy (kWh)", "DC V (Min/Max)", "Weight (kg)", "Area (m²)", "Price (SGD)"],
        rows=[
            ["H30",  "30",  "43.5",  "210–850 V",  "7",  "69.3",  "231 / 269.5 V",  "1,400", "2.1", "$100,000"],
            ["H50",  "50",  "72.5",  "350–850 V",  "11", "108.9", "363 / 423.5 V",  "2,200", "3.2", "$120,000"],
            ["H60",  "60",  "87.0",  "420–850 V",  "14", "138.6", "462 / 539 V",    "2,800", "4.2", "TBD"],
            ["H100", "100", "145.0", "680–900 V",  "22", "217.8", "726 / 847 V",    "4,400", "6.3", "TBD"],
            ["H125", "125", "181.0", "680–900 V",  "22", "217.8", "726 / 847 V",    "4,400", "6.3", "$100,000"],
        ],
        col_widths=[1.5, 1.2, 1.8, 2.3, 2.0, 2.2, 2.8, 2.2, 1.6, 2.2],
        hdr_color=HDR_GLOBAL
    )

    # 1g. Model sizing rule
    heading2(doc, "1g.  Model Sizing Rule  ( recommend_model() )", color=C_NAVY)
    equation_block(doc, "Required kVA", "required_kVA  =  avg_kVAr  ×  1.2",
                   "20% headroom above average reactive load")
    para(doc, sa=2)
    make_table(doc,
        headers=["Condition", "Has Solar", "Recommended Model"],
        rows=[
            ["required_kVA ≤ 30",  "Yes or No", "H30"],
            ["required_kVA ≤ 50",  "Yes",        "H50"],
            ["required_kVA ≤ 60",  "Yes",        "H60"],
            ["required_kVA ≤ 100", "Yes",        "H100"],
            ["required_kVA > 100", "Yes",        "H125"],
            ["Any",                "No",         "Capped at H50  (SCDF / space constraint)"],
        ],
        col_widths=[4.5, 3.0, 6.0],
        hdr_color=HDR_GLOBAL
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. PER-SITE CONFIG
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "2.  Per-Site Config  ( config/{site_id}.yaml )")
    para(doc, "One YAML file per site. Values here override global defaults where both exist.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    heading2(doc, "2a.  Site Identity", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Description"],
        rows=[
            ["Site ID",       "site_id",      "Baoyuan-MSB1",     "Unique identifier — used as DB filename and SAR key"],
            ["Display Name",  "display_name", "Baoyuan Industrial — MSB1", "Human-readable name for reports"],
            ["Country",       "country",      "China",            "Determines default voltage standard and language"],
            ["Timezone",      "timezone",     "Asia/Shanghai",    "For timestamp normalisation and scheduler offsets"],
        ],
        col_widths=[3.2, 3.2, 5.0, 6.4],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2b.  Installed Hardware", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Options", "Description"],
        rows=[
            ["HyESys Model",   "hyesys_model", "H30 / H50 / H60 / H100 / H125", "Determines kVA clamp in ΔQ injection calculation"],
            ["Has Solar",      "has_solar",    "true / false",                   "Enables solar storage priority; affects model sizing cap"],
        ],
        col_widths=[3.2, 3.2, 5.2, 6.2],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2c.  Tariff & Currency", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Unit", "Description"],
        rows=[
            ["Currency",         "tariff.currency",       "CNY / SGD",  "–",          "Reporting currency for savings calculations"],
            ["Tariff Rate",      "tariff.rate_per_kwh",   "0.85",       "CNY or SGD / kWh", "Electricity rate for kWh savings → monetary value"],
            ["Hours per Month",  "tariff.hours_per_month","720",        "h",          "720 h = 30 days × 24 h. Adjust for actual operating hours"],
        ],
        col_widths=[3.2, 4.2, 2.2, 3.0, 5.2],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2d.  Agent 1 — Site-Specific Validation Thresholds", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Unit", "Description"],
        rows=[
            ["Voltage Min",             "agent1.voltage_min_v",         "360",  "V",  "Override global 195.5 V for non-SG sites (e.g. China 380 V −5%)"],
            ["Voltage Max",             "agent1.voltage_max_v",         "440",  "V",  "Override global 264.5 V for non-SG sites (e.g. China 380 V +15%)"],
            ["Max Gap Minutes",         "agent1.max_gap_minutes",       "30",   "min","Consecutive missing timestamps beyond this → SUSPECT gap flag"],
            ["Zero Current Threshold",  "agent1.zero_current_threshold_a","5",  "A",  "Current below this treated as plant-off, not fault"],
            ["Expected Phases",         "agent1.expected_phases",       "3",    "–",  "Phase count validation. Flags rows with missing phases"],
        ],
        col_widths=[4.0, 5.0, 1.8, 1.4, 5.6],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2e.  Agent 2 — Site-Specific Decision Parameters", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Unit", "Description"],
        rows=[
            ["kVAr Injection Step",      "agent2.kvar_injection_step",      "5.0",           "kVAr", "Step size per control interval. Larger = more aggressive correction"],
            ["Demand Risk Threshold",    "agent2.demand_risk_threshold_kw", "800",           "kW",   "Site-specific peak demand ceiling. Triggers REDUCE action when breached"],
            ["Savings Method",           "agent2.savings_method",           "i2r_fraction",  "–",    "i2r_fraction (default) or direct_kw (for sites with poor measurement quality)"],
        ],
        col_widths=[4.2, 5.0, 2.8, 1.4, 4.4],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2f.  Savings Baseline  ( measured at site )", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Unit", "Description"],
        rows=[
            ["Baseline Current",    "savings.baseline_current_a",  "3405",  "A",   "Measured I_rms at MSB incomer before HyESys activation"],
            ["Baseline Active Power","savings.baseline_kw",        "785",   "kW",  "Average kW before activation. Used as denominator in savings fraction"],
            ["THD Override",        "savings.baseline_thd_override","0.45", "–",   "Override global 15% if site THD is measured (e.g. Baoyuan Phase C = 142%)"],
        ],
        col_widths=[4.2, 5.0, 1.8, 1.2, 5.6],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    heading2(doc, "2g.  Retraining Schedule", color=C_GREEN)
    make_table(doc,
        headers=["Parameter", "Variable", "Example", "Unit", "Description"],
        rows=[
            ["Retrain Time Offset", "retrain.time_offset_minutes", "10", "min",
             "Minutes after midnight to run nightly retrain. Staggered per site to avoid CPU collision"],
        ],
        col_widths=[4.0, 4.8, 1.8, 1.4, 5.8],
        hdr_color=HDR_SITE,
        alt_color="F0F7EC"
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. AGENT 1
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "3.  Agent 1  —  Data Quality & Validation  ( agent1/validator.py )")
    para(doc, "Rule-based, LLM-free. Processes each row sequentially. Tags: CLEAN / SUSPECT / REJECTED.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    heading2(doc, "3a.  Input Fields Per Row", color=C_ORANGE)
    make_table(doc,
        headers=["Field", "Variable", "Type", "Description"],
        rows=[
            ["Site ID",    "site_id",   "string",  "Identifies which site this record belongs to"],
            ["Timestamp",  "timestamp", "string",  "Raw timestamp — normalised to ISO 8601 by parser"],
            ["Active Power","kW",       "float",   "Active (real) power at MSB incomer [kW]"],
            ["Reactive Power","kVAr",   "float",   "Reactive power [kVAr]. Negative = leading (capacitive) — valid"],
            ["Power Factor","PF",       "float",   "Signed power factor. Negative = leading PF"],
            ["Voltage",    "voltage_V", "float",   "Line-to-neutral RMS voltage [V]"],
        ],
        col_widths=[3.2, 3.0, 1.8, 9.8],
        hdr_color=HDR_A1,
        alt_color=ALT_ROW2
    )

    heading2(doc, "3b.  Validation Rules  ( applied in sequence )", color=C_ORANGE)
    make_table(doc,
        headers=["Rule #", "Check", "Condition → Tag", "Parameter Used", "Notes"],
        rows=[
            ["1", "Timestamp parseable",       "Fail → REJECTED",  "—",                   "Catches malformed / mixed-format timestamps"],
            ["2", "Duplicate timestamp",        "Duplicate → REJECTED", "—",              "Stateful — tracks all seen timestamps per site"],
            ["3", "Numeric field parse",        "TypeError → REJECTED", "—",              "Catches string values in numeric columns"],
            ["4", "Voltage physically valid",   "voltage ≤ 0 → REJECTED", "—",           "Negative or zero voltage is physically impossible"],
            ["5", "kW range",                   "|kW| > 10,000 → REJECTED", "—",         "Sanity ceiling. Negative kW valid for solar export"],
            ["6", "kVAr range",                 "|kVAr| > 5,000 → REJECTED", "—",        "Sanity ceiling. Negative kVAr = capacitive load"],
            ["7", "PF firmware bug",            "|PF| > 1.0 → REJECTED", "—",            "PF cannot exceed 1.0 — firmware error"],
            ["8", "All-zero row",               "kW=kVAr=PF=0 → SUSPECT", "—",           "Meter dropout / comms loss. Data retained but flagged"],
            ["9", "PF firmware saturation",     "|PF| = 1.0 exactly → SUSPECT", "—",     "Firmware saturates at ±1.0 — readings unreliable"],
            ["10","Voltage outside LV range",   "V < VOLTAGE_MIN or V > VOLTAGE_MAX → SUSPECT", "VOLTAGE_MIN=195.5 V\nVOLTAGE_MAX=264.5 V", "Singapore 230 V ±15%. Override per-site for China (380 V ±15%)"],
            ["11","PF below SP threshold",      "0 < |PF| < 0.85 → CLEAN (flagged)", "PF_PENALTY_THRESHOLD=0.85", "Not rejected — valid data but signals SP penalty risk"],
        ],
        col_widths=[1.2, 4.2, 4.0, 4.2, 4.2],
        hdr_color=HDR_A1,
        alt_color=ALT_ROW2
    )

    heading2(doc, "3c.  Output Quality Tags", color=C_ORANGE)
    make_table(doc,
        headers=["Tag", "Written to DB?", "Meaning"],
        rows=[
            ["CLEAN",    "Yes", "Row passed all checks. Used by Agent 2 for state building"],
            ["SUSPECT",  "Yes", "Row has anomaly but is retained. Agent 2 may use with reduced confidence"],
            ["REJECTED", "No",  "Row discarded — not written to site.db"],
        ],
        col_widths=[2.5, 3.5, 11.8],
        hdr_color=HDR_A1,
        alt_color=ALT_ROW2
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. AGENT 2 — STATE VARIABLES
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "4.  Agent 2  —  State Variables  ( agent2/state.py )")
    para(doc, "Each 15-minute interval produces one State snapshot. All derived quantities are computed from the four measured inputs.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    heading2(doc, "4a.  Measured Inputs", color=C_BLUE)
    make_table(doc,
        headers=["Symbol", "Variable", "Unit", "Description"],
        rows=[
            ["P",   "kW",        "kW",  "Active (real) power at MSB incomer"],
            ["Q",   "kVAr",      "kVAr","Reactive power. Negative = leading (capacitive)"],
            ["PF",  "PF",        "–",   "Measured power factor (signed). Negative = leading"],
            ["V_LN","voltage_V", "V",   "Line-to-neutral RMS voltage"],
        ],
        col_widths=[1.8, 3.0, 1.5, 11.5],
        hdr_color=HDR_A2
    )

    heading2(doc, "4b.  Derived State Properties & Equations", color=C_BLUE)
    make_table(doc,
        headers=["Property", "Variable", "Equation", "Unit", "Notes"],
        rows=[
            ["Apparent Power",    "kVA",                  "S  =  √( P²  +  Q² )",                         "kVA",  "Total electrical demand on the supply"],
            ["PF Angle",          "phi_rad",               "φ  =  arccos( |PF| )",                          "rad",  "0 = unity PF; π/2 = purely reactive"],
            ["PF Angle (degrees)","phi_deg",               "φ°  =  φ  ×  (180 / π)",                       "°",    "For human-readable reporting"],
            ["tan(φ)",            "tan_phi",               "tan(φ)  =  Q / P",                              "–",    "Used in Q correction formula"],
            ["Q Target",          "q_target",              "Q_tgt  =  P  ×  tan( arccos(0.98) )\n         =  P  ×  0.2031", "kVAr", "Reactive power at target PF = 0.98"],
            ["Q Correction",      "q_correction_needed",   "ΔQ  =  Q_current  −  Q_target\n     =  P × ( tan φ_current − tan φ_target )", "kVAr", "+ = lagging; inject capacitive  |  − = leading; inject inductive"],
            ["Line Current",      "current_A",             "I  =  ( S × 1000 )  /  ( 3 × V_LN )",          "A",    "3-phase, line-to-neutral convention (230 V nominal)"],
            ["Loss vs Unity PF",  "loss_fraction_vs_unity_pf", "f_unity  =  1  −  PF²",                    "–",    "Fraction of I²R recoverable if PF corrected to 1.0"],
            ["Recoverable Loss",  "recoverable_loss_fraction",  "S_tgt = P / PF_target\nf_rec = 1 − ( S_tgt / S )²", "–", "Fraction recoverable by correcting to PF_target = 0.98"],
        ],
        col_widths=[3.2, 4.4, 5.5, 1.2, 3.5],
        hdr_color=HDR_A2
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. AGENT 2 — PF CORRECTION
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "5.  Agent 2  —  PF Correction  ( agent2/tools.py → compute_pf_correction() )")
    para(doc, "Calculates required kVAr injection to drive site PF to PF_TARGET = 0.98. Executed every 15-minute interval.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    heading2(doc, "5a.  Pre-computed Constants", color=C_BLUE)
    make_table(doc,
        headers=["Constant", "Variable", "Value", "Derivation"],
        rows=[
            ["PF Target Angle",     "PHI_TARGET",     "0.1997 rad  (11.48°)", "arccos(0.98)"],
            ["tan(PF_TARGET angle)","TAN_PHI_TARGET",  "0.2031",              "tan( arccos(0.98) )"],
        ],
        col_widths=[4.0, 4.5, 4.0, 5.3],
        hdr_color=HDR_A2
    )

    heading2(doc, "5b.  Step-by-Step Equations", color=C_BLUE)
    steps = [
        ("Step 1 — Q Target",      "Q_target  =  P  ×  tan( arccos(0.98) )  =  P  ×  0.2031",             "Reactive power needed to reach PF = 0.98"),
        ("Step 2 — ΔQ Required",   "ΔQ_required  =  Q_current  −  Q_target",                               "+ = lagging (inject capacitive)   − = leading (inject inductive)"),
        ("Step 3 — Clamp to Model","ΔQ_injected  =  clamp( ΔQ_required,  −model_kVA,  +model_kVA )",       "Cannot inject more than the installed HyESys model capacity"),
        ("Step 4a — Q After",      "Q_after  =  Q_current  −  ΔQ_injected",                                "Predicted reactive power after injection"),
        ("Step 4b — S After",      "S_after  =  √( P²  +  Q_after² )",                                     "Predicted apparent power after injection"),
        ("Step 4c — PF After",     "PF_after  =  P  /  S_after",                                           "Predicted achievable power factor"),
        ("Step 5a — Savings Fraction","f  =  1  −  ( S_after / S_before )²",                               "Fraction of I²R distribution losses eliminated"),
        ("Step 5b — Loss Estimate","P_loss_base  =  P  ×  THD_ASSUMPTION  (0.15)\nkW_saved_est  =  f  ×  P_loss_base", "Estimated kW saving per 15-min interval"),
    ]
    for label, eq, note in steps:
        equation_block(doc, label, eq, note)
    para(doc, sa=4)

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. AGENT 2 — DEMAND RISK
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "6.  Agent 2  —  Demand Risk  ( agent2/tools.py → assess_demand_risk() )")
    para(doc, "Evaluates whether current kW approaches historical peak. Triggers REDUCE or battery storage action.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    heading2(doc, "6a.  Equations", color=C_BLUE)
    demand_eqs = [
        ("Peak Demand %",         "demand_pct  =  ( P_current / P_peak_historical )  ×  100",              "% of historical 15-min peak"),
        ("Headroom",              "P_headroom  =  P_peak  −  P_current",                                   "kW buffer before new peak is set"),
        ("Excess Demand",         "P_excess  =  max( P_current − P_peak,  0 )",                            "kW above historical peak"),
        ("Demand Charge at Risk", "charge_risk  =  P_excess  ×  DEMAND_TARIFF_SGD_PER_KW  (10.0)",        "SGD at risk if P_excess sustained for billing period"),
    ]
    for label, eq, note in demand_eqs:
        equation_block(doc, label, eq, note)
    para(doc, sa=4)

    heading2(doc, "6b.  Risk Levels", color=C_BLUE)
    make_table(doc,
        headers=["Risk Level", "Condition", "recommend_store", "Agent 2 Action"],
        rows=[
            ["CRITICAL", "demand_pct  ≥  95%",  "True",                "REDUCE output — prevent new peak from being set"],
            ["HIGH",     "demand_pct  ≥  85%",  "True",                "REDUCE and discharge battery storage"],
            ["MEDIUM",   "demand_pct  ≥  70%",  "True if has_solar",   "Monitor; discharge solar storage if available"],
            ["LOW",      "demand_pct  <  70%",  "False",               "Normal INJECT_KVAR operation continues"],
        ],
        col_widths=[2.8, 4.5, 3.5, 7.0],
        hdr_color=HDR_A2
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. AGENT 2 — REWARD
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "7.  Agent 2  —  Reward Computation  ( agent2/outcome.py → compute_reward() )")
    para(doc, "Closes the SAR loop. Computed from the before/after State pair following each injection decision.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    reward_eqs = [
        ("Step 1 — PF Reward",       "r_PF  =  PF_after  −  PF_before",                                               "> 0 = improvement;  < 0 = overcorrected or leading"),
        ("Step 2 — Loss Fraction",   "r_loss  =  1  −  ( S_after / S_before )²",                                       "Fraction of I²R losses eliminated this interval"),
        ("Step 3 — Combined Reward", "r_total  =  W_PF × r_PF  +  W_LOSS × r_loss\n         =  0.60 × r_PF  +  0.40 × r_loss", "Weighted sum. PF dominates to prioritise SP penalty avoidance"),
        ("Step 4 — THD Estimate",    "THD_est  =  √( ( PF_before / PF_target )²  /  ( 1 − r_loss )  −  1 )",          "Back-calculated site THD. Decimal fraction (0.15 = 15%)"),
        ("Step 5 — Classification",  "r_PF ≥ +0.01  →  POSITIVE\nr_PF ≤ −0.01  →  NEGATIVE\nelse          →  NEUTRAL", "Determines SAR outcome label stored in site.db"),
    ]
    for label, eq, note in reward_eqs:
        equation_block(doc, label, eq, note)
    para(doc, sa=6)

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. SAVINGS COMPUTATION
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "8.  Savings Computation  ( models/savings.py )")
    para(doc, "Primary validation observable: I_rms at MSB incomer. Cable resistance R cancels from all equations — no impedance modelling required.",
         size=9.5, italic=True, color=C_GREY, sa=6)

    savings_eqs = [
        ("kVA Before",          "kVA_before  =  √( kW²  +  kVAr_before² )",                                "Apparent power before HyESys injection"),
        ("kVA After",           "kVA_after   =  √( kW²  +  kVAr_after² )",                                 "Apparent power after injection"),
        ("Savings Fraction",    "f  =  1  −  ( kVA_after / kVA_before )²",                                  "R cancels because P_loss = 3 × I² × R, and I ∝ S/V"),
        ("Loss Component",      "P_loss_base  =  avg_kW  ×  THD_ASSUMPTION  (0.15)",                        "Approximate I²R loss component at 15% THD baseline"),
        ("kWh Savings",         "kWh_saved  =  f  ×  P_loss_base  ×  hours",                                "Monthly: hours = 720.  Annual: hours = 8,760"),
        ("Monetary Savings",    "savings_value  =  kWh_saved  ×  tariff_rate_per_kwh",                      "Currency (CNY / SGD) from per-site config"),
        ("kVAr Required",       "ΔQ  =  kW × ( tan(arccos(PF_current)) − tan(arccos(PF_target)) )",        "kVAr injection needed to reach PF_target = 0.98"),
        ("THD Back-Calc",       "THD  =  √( ( PF_before / PF_target )²  /  ( 1 − f )  −  1 )",             "Returns None if f ≥ 1.0 or PF inputs invalid"),
    ]
    for label, eq, note in savings_eqs:
        equation_block(doc, label, eq, note)
    para(doc, sa=6)

    heading2(doc, "8a.  Why R Cancels  ( Measurement Philosophy )", color=C_PURPLE)
    para(doc,
         "P_loss = 3 × I² × R  (3-phase cable losses)\n"
         "Savings = P_loss_before − P_loss_after = 3R × (I_before² − I_after²) = 3R × I_before² × f\n"
         "Dividing: f = (P_loss_before − P_loss_after) / P_loss_before = 1 − (I_after/I_before)²\n"
         "R, V, and cable length all cancel. Only the ratio of currents matters.",
         size=9.5, mono=False, color=C_BODY, sb=2, sa=4)

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. SAR LOOP SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "9.  SAR Loop Summary  ( State → Action → Reward )")
    make_table(doc,
        headers=["Component", "Content", "Source"],
        rows=[
            ["State (S)", "site_id, timestamp, kW, kVAr, PF, voltage_V, solar",         "agent2/state.py — built from CLEAN rows in site.db"],
            ["Action (A)", "INJECT_KVAR / HOLD / REDUCE  +  ΔQ_injected [kVAr]",         "agent2/agent.py — decision engine using site config params"],
            ["Reward (R)", "r_PF, r_loss, r_total, THD_est, outcome (POSITIVE/NEUTRAL/NEGATIVE)", "agent2/outcome.py — computed from before/after State pair"],
        ],
        col_widths=[2.8, 8.4, 6.6],
        hdr_color=HDR_A2
    )

    # Footer
    para(doc,
         "This document reflects the codebase as of May 2026. "
         "Parameters in config/global.yaml and config/{site_id}.yaml will govern runtime behaviour "
         "once the multi-site config architecture is implemented.",
         size=8.5, italic=True, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, sb=12)

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    build()
