"""
Generates HyESys Agent Rules Reference Word document.
Covers all Agent 1 and Agent 2 rules with colour-coded tables.
Saved to OneDrive Software/claude code folder.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DST = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\2024 HyESys\Software\claude code"
    r"\HyESys_Agent_Rules_Reference_2026-05-19.docx"
)

# ── Colour palette ────────────────────────────────────────────────────────────
C_NAVY      = RGBColor(0x1F, 0x37, 0x63)   # section headers
C_MIDBLUE   = RGBColor(0x2E, 0x75, 0xB6)   # sub-headers
C_ORANGE    = RGBColor(0xED, 0x7D, 0x31)   # suspect / warning
C_RED       = RGBColor(0xC0, 0x00, 0x00)   # rejected / critical
C_GREEN     = RGBColor(0x37, 0x56, 0x23)   # clean / positive
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK     = RGBColor(0x00, 0x00, 0x00)
C_DARKGREY  = RGBColor(0x40, 0x40, 0x40)

# Background fill hex strings (for shading XML)
BG_NAVY     = "1F3763"
BG_MIDBLUE  = "2E75B6"
BG_ORANGE   = "FFC000"
BG_RED      = "FF0000"
BG_LTRED    = "FFE0E0"
BG_LTORAN   = "FFF2CC"
BG_LTGREEN  = "E2EFDA"
BG_LTBLUE   = "DEEAF1"
BG_LTGREY   = "F2F2F2"
BG_WHITE    = "FFFFFF"
BG_DARKGREY = "595959"


# ── Helper: set cell shading ──────────────────────────────────────────────────
def shade_cell(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    val.get("sz",  "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


# ── Helper: paragraph styling ─────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=C_NAVY, size=14, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold     = True
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p


def add_subheading(doc, text, color=C_MIDBLUE, size=11, space_before=10, space_after=4):
    return add_heading(doc, text, color=color, size=size,
                       space_before=space_before, space_after=space_after)


def add_body(doc, text, size=10, color=C_DARKGREY, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p


# ── Helper: table builder ─────────────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths,
               hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
               row_bgs=None, row_fgs=None):
    """
    headers   : list of header strings
    rows      : list of row tuples
    col_widths: list of Inches() widths
    row_bgs   : optional list (same length as rows) of per-row bg hex strings
    row_fgs   : optional list of per-row font RGBColor
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style          = "Table Grid"
    table.alignment      = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit  = False

    # Header row
    hdr_row = table.rows[0]
    for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, headers, col_widths)):
        cell.width = w
        shade_cell(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold           = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = hdr_fg
        run.font.name      = "Calibri"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row     = table.rows[r_idx + 1]
        bg_hex  = (row_bgs[r_idx]  if row_bgs  else BG_WHITE)
        fg_col  = (row_fgs[r_idx]  if row_fgs  else C_BLACK)

        for c_idx, (cell, val, w) in enumerate(zip(row.cells, row_data, col_widths)):
            cell.width = w
            shade_cell(cell, bg_hex)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = fg_col
            run.font.name      = "Calibri"
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()   # spacing after table
    return table


def banner_cell(table, text, bg_hex, fg_col=C_WHITE, size=10):
    """Merge all cells in a row and use as a colour banner label."""
    row  = table.add_row()
    row.cells[0].merge(row.cells[-1])
    cell = row.cells[0]
    shade_cell(cell, bg_hex)
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(size)
    run.font.color.rgb = fg_col
    run.font.name      = "Calibri"
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Cover block ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("HyESys Agent Rules Reference")
run.bold           = True
run.font.size      = Pt(20)
run.font.color.rgb = C_NAVY
run.font.name      = "Calibri"

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("Complete Rule Set — Agent 1 (Data Validation) & Agent 2 (Decision & Analysis)")
run2.font.size      = Pt(11)
run2.font.color.rgb = C_MIDBLUE
run2.font.name      = "Calibri"
run2.italic         = True

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(14)
run3 = p3.add_run(
    "Advancer Smart Technology Pte Ltd  ·  HyESys Department  ·  May 2026  ·  Confidential"
)
run3.font.size      = Pt(9)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run3.font.name      = "Calibri"

# Horizontal rule (via border on an empty paragraph)
hr_p  = doc.add_paragraph()
hr_pf = hr_p.paragraph_format
hr_pf.space_before = Pt(0)
hr_pf.space_after  = Pt(14)
pPr   = hr_p._p.get_or_add_pPr()
pBdr  = OxmlElement("w:pBdr")
bot   = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "8")
bot.set(qn("w:space"), "1")
bot.set(qn("w:color"), "1F3763")
pBdr.append(bot)
pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════
# SECTION 1 — AGENT 1
# ═══════════════════════════════════════════════════════════════════

add_heading(doc, "1.  AGENT 1 — Data Quality & Ingestion", size=14,
            color=C_NAVY, space_before=0)
add_body(doc,
    "Agent 1 is fully rule-based and LLM-free. It validates every incoming row before "
    "writing to the database. Each row receives one of three quality tags: "
    "REJECTED (discarded), SUSPECT (stored but excluded from training), or "
    "CLEAN (used by Agent 2).",
    size=9.5)

# ── 1.1 AC Meter Data ─────────────────────────────────────────────
add_subheading(doc, "1.1  AC Meter Data Validation", space_before=12)
add_body(doc,
    "Source: agent1/validator.py and mqtt_ingest.py  |  "
    "Fields: site_id, timestamp, kW, kVAr, PF, voltage_V",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

W = [Inches(0.4), Inches(1.5), Inches(2.8), Inches(1.6)]

# REJECTED table
add_body(doc, "", space_before=2, space_after=0)
rejected_rows = [
    ("#1", "Unparseable timestamp",   "Cannot convert field to ISO 8601 format",          "Discard row"),
    ("#2", "Duplicate timestamp",     "Same site_id + timestamp already in database",     "Discard row"),
    ("#3", "Non-numeric field",       "kW / kVAr / PF / voltage_V cannot be cast to float","Discard row"),
    ("#4", "Impossible voltage",      "voltage_V ≤ 0",                                    "Discard row"),
    ("#5", "kW out of range",         "kW outside −10,000 … +10,000",                     "Discard row"),
    ("#6", "kVAr out of range",       "kVAr outside −5,000 … +5,000",                     "Discard row"),
    ("#7", "PF firmware bug",         "PF outside [−1.0, +1.0]",                          "Discard row"),
]
t = make_table(
    doc,
    headers=["No.", "Rule Name", "Trigger Condition", "Action"],
    rows=rejected_rows,
    col_widths=W,
    hdr_bg="C00000", hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED] * len(rejected_rows),
    row_fgs=[C_BLACK]  * len(rejected_rows),
)

# SUSPECT table
suspect_rows = [
    ("#8",  "Meter dropout",           "kW = 0 AND kVAr = 0 AND PF = 0 simultaneously",   "Store — excluded from training"),
    ("#9",  "PF firmware saturation",  "|PF| exactly equals 1.0",                          "Store — excluded from training"),
    ("#10", "Voltage out of LV range", "voltage_V outside 195.5 – 264.5 V  (230V ±15%)",  "Store — excluded from training"),
]
make_table(
    doc,
    headers=["No.", "Rule Name", "Trigger Condition", "Action"],
    rows=suspect_rows,
    col_widths=W,
    hdr_bg="C55A11", hdr_fg=C_WHITE,
    row_bgs=[BG_LTORAN] * len(suspect_rows),
    row_fgs=[C_BLACK]   * len(suspect_rows),
)

# CLEAN annotated
clean_rows = [
    ("#11", "SP penalty risk flag",
     "0 < |PF| < 0.85  (below SP penalty threshold)",
     "CLEAN — stored and used for Agent 2;\nreject_reason field populated with warning"),
]
make_table(
    doc,
    headers=["No.", "Rule Name", "Trigger Condition", "Action"],
    rows=clean_rows,
    col_widths=W,
    hdr_bg="375623", hdr_fg=C_WHITE,
    row_bgs=[BG_LTGREEN],
    row_fgs=[C_BLACK],
)

# Legend for 1.1
add_body(doc,
    "Tag legend:  ■ RED = REJECTED (discarded)   "
    "■ ORANGE = SUSPECT (stored, excluded from training)   "
    "■ GREEN = CLEAN (stored, used by Agent 2)",
    size=8.5, color=RGBColor(0x50, 0x50, 0x50))


# ── 1.2 BMS Data ──────────────────────────────────────────────────
add_subheading(doc, "1.2  BMS Data Validation  (Battery Management System)", space_before=14)
add_body(doc,
    "Source: mqtt_ingest.py  |  Topic: stsc/aems/message/+  |  "
    "Fields: type, cabinetId, reportTimeTs, soc, soh, tempMain  |  "
    "cabinetId is mapped to site_id via CABINET_TO_SITE dict",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WB = [Inches(0.4), Inches(1.6), Inches(2.8), Inches(1.5)]

bms_rej = [
    ("#1", "Missing required fields",  "Any of: type, cabinetId, reportTimeTs, soc, soh, tempMain absent", "Discard message"),
    ("#2", "Invalid timestamp",        "reportTimeTs cannot be converted from Unix milliseconds",           "Discard message"),
    ("#3", "SOC out of range",         "SOC outside 0 – 100 %",                                            "Discard message"),
    ("#4", "SOH out of range",         "SOH outside 0 – 100 %",                                            "Discard message"),
    ("#5", "Temperature out of range", "tempMain outside −20 … +80 °C",                                    "Discard message"),
]
make_table(
    doc,
    headers=["No.", "Rule Name", "Trigger Condition", "Action"],
    rows=bms_rej,
    col_widths=WB,
    hdr_bg="C00000", hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED] * len(bms_rej),
    row_fgs=[C_BLACK]  * len(bms_rej),
)

bms_sus = [
    ("#6", "Critically low SOC",    "SOC < 10 %",                                         "Store — excluded from training"),
    ("#7", "Thermal risk",          "tempMain > 50 °C",                                    "Store — excluded from training"),
    ("#8", "Degraded battery",      "SOH < 80 %",                                          "Store — excluded from training"),
    ("#9", "Cell voltage imbalance","Cell voltage spread (singleVoltageDiff) > 0.1 V",     "Store — excluded from training"),
]
make_table(
    doc,
    headers=["No.", "Rule Name", "Trigger Condition", "Action"],
    rows=bms_sus,
    col_widths=WB,
    hdr_bg="C55A11", hdr_fg=C_WHITE,
    row_bgs=[BG_LTORAN] * len(bms_sus),
    row_fgs=[C_BLACK]   * len(bms_sus),
)

add_body(doc,
    "Tag legend:  ■ RED = REJECTED (discarded)   "
    "■ ORANGE = SUSPECT (stored, excluded from training)",
    size=8.5, color=RGBColor(0x50, 0x50, 0x50))


# ═══════════════════════════════════════════════════════════════════
# SECTION 2 — AGENT 2
# ═══════════════════════════════════════════════════════════════════

doc.add_page_break()

add_heading(doc, "2.  AGENT 2 — Decision & Analysis Rules", size=14,
            color=C_NAVY, space_before=0)
add_body(doc,
    "Agent 2 is 100% local Python with no LLM at runtime. It reads 15-minute state "
    "snapshots from hyesys.db, runs a PI control law, detects events, issues kVAr "
    "injection commands, and logs every State → Action → Reward (SAR) triplet.",
    size=9.5)


# ── 2.1 Decision Engine ──────────────────────────────────────────
add_subheading(doc, "2.1  Decision Engine — Priority Order", space_before=12)
add_body(doc,
    "Source: agent2/agent.py  |  "
    "Rules are evaluated in priority order — the first match wins.",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WD2 = [Inches(0.5), Inches(1.5), Inches(3.2), Inches(1.1)]
dec_rows = [
    ("P1", "Solar + CRITICAL demand",
     "Site has solar AND demand_pct ≥ 95 %",
     "REDUCE — reserve hardware for peak shaving"),
    ("P2", "Within dead-band",
     "|e(t)| = |0.98 − PF_current| < 0.005",
     "HOLD — suppress micro-oscillations"),
    ("P3", "Below hardware resolution",
     "|ΔQ_clamped| < 0.5 kVAr after PI calculation",
     "HOLD"),
    ("P4", "Lagging PF",
     "ΔQ_commanded > 0  (load draws inductive reactive power)",
     "INJECT_KVAR — capacitive injection"),
    ("P5", "Leading PF or overcorrected",
     "ΔQ_commanded < 0  (load has excess capacitive reactive power)",
     "REDUCE — inductive injection"),
]
make_table(
    doc,
    headers=["Priority", "Condition Name", "Trigger", "Action Issued"],
    rows=dec_rows,
    col_widths=WD2,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED, BG_LTBLUE, BG_LTBLUE, BG_LTGREEN, BG_LTORAN],
    row_fgs=[C_BLACK] * 5,
)


# ── 2.2 PI Control Law Parameters ────────────────────────────────
add_subheading(doc, "2.2  PI Control Law Parameters", space_before=12)
add_body(doc,
    "Source: agent2/agent.py  |  "
    "Applied every 15-minute interval (Δt = 0.25 hr).",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WC = [Inches(1.5), Inches(0.8), Inches(2.0), Inches(2.0)]
ctrl_rows = [
    ("PF target (PF_TARGET)",      "0.98",  "Target power factor",
     "Not 1.0 — unity is impractical due to law of convergence"),
    ("Dead-band (ε)",              "0.005", "Minimum PF error to trigger action",
     "Prevents hunting / oscillation around target"),
    ("Proportional gain (K_P)",    "1.00",  "ΔQ_P = K_P × ΔQ_required",
     "Full analytical correction per 15-min step"),
    ("Integral gain (K_I)",        "0.50",  "ΔQ_I = K_I × I(t)",
     "Handles sustained PF deficit across multiple intervals"),
    ("Integrator anti-windup (I_MAX)", "±20", "Clamp on I(t)  [kVAr·hr]",
     "Prevents runaway ≈ 4 hrs × 5 kVAr sustained error"),
    ("Interval (Δt)",              "0.25 hr", "I(t) = I(t−1) + e(t) × Δt",
     "15-minute meter resolution"),
    ("Hardware resolution floor",  "0.5 kVAr", "If |ΔQ_clamped| < 0.5 → HOLD",
     "Below HyESys minimum controllable step"),
]
make_table(
    doc,
    headers=["Parameter", "Value", "Formula / Role", "Notes"],
    rows=ctrl_rows,
    col_widths=WC,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTBLUE if i % 2 == 0 else BG_WHITE for i in range(len(ctrl_rows))],
    row_fgs=[C_BLACK] * len(ctrl_rows),
)


# ── 2.3 Demand Risk Assessment ────────────────────────────────────
add_subheading(doc, "2.3  Demand Risk Assessment", space_before=12)
add_body(doc,
    "Source: agent2/tools.py  |  "
    "demand_pct = (current_kW / historical_peak_kW) × 100  |  "
    "Singapore tariff: demand charge ∝ max 15-min kW in billing period  |  "
    "Proxy tariff rate used: SGD 10 / kW / month.",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WR = [Inches(1.0), Inches(1.5), Inches(1.5), Inches(2.3)]
risk_rows = [
    ("CRITICAL", "demand_pct ≥ 95 %", "Yes",  "Near or at historical peak — immediate storage discharge recommended"),
    ("HIGH",     "demand_pct ≥ 85 %", "Yes",  "Likely to set new peak — activate storage"),
    ("MEDIUM",   "demand_pct ≥ 70 %", "Solar sites only", "Elevated — monitor; solar sites pre-charge storage"),
    ("LOW",      "demand_pct < 70 %", "No",   "Normal operating range"),
    ("UNKNOWN",  "No history yet",    "No",   "Treat as 100 % until baseline accumulates"),
]
make_table(
    doc,
    headers=["Risk Level", "Condition", "Recommend Store?", "Meaning"],
    rows=risk_rows,
    col_widths=WR,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED, "FFD7D7", BG_LTORAN, BG_LTGREEN, BG_LTGREY],
    row_fgs=[C_BLACK] * 5,
)


# ── 2.4 Event Detection ───────────────────────────────────────────
add_subheading(doc, "2.4  Event Detection", space_before=12)
add_body(doc,
    "Source: agent2/events.py  |  "
    "Four event types run every 15-minute interval and are sorted by severity (CRITICAL → WARNING → INFO).",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WE = [Inches(1.1), Inches(1.4), Inches(0.9), Inches(2.9)]

# Threshold
add_body(doc, "Threshold Events — fixed absolute limits", size=9.5,
         color=C_NAVY, space_before=6, space_after=2)
thr_rows = [
    ("PF_PENALTY_RISK",  "0 < PF < 0.85",           "CRITICAL", "Below SP penalty threshold — immediate injection required; kVAr correction calculated"),
    ("PF_LOW",           "0 < PF < 0.98",            "WARNING",  "Below PF target — injection needed; recoverable loss fraction computed"),
    ("PF_LEADING",       "Q < −5 kVAr",              "WARNING",  "Capacitive (leading) load — leading PF is as harmful as lagging; inductive injection needed"),
]
make_table(
    doc,
    headers=["Event Subtype", "Trigger", "Severity", "Description"],
    rows=thr_rows,
    col_widths=WE,
    hdr_bg=BG_DARKGREY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED, BG_LTORAN, BG_LTORAN],
    row_fgs=[C_BLACK] * 3,
)

# Statistical
add_body(doc,
    "Statistical Events — EMA z-score (α = 2/(8+1) ≈ 0.222, ~2-hour window)",
    size=9.5, color=C_NAVY, space_before=6, space_after=2)
add_body(doc,
    "z = (x_t − μ_t) / σ_t   where μ_t = EMA mean,  σ_t = EMA standard deviation",
    size=9, color=RGBColor(0x50, 0x50, 0x50), space_before=0, space_after=3)
stat_rows = [
    ("PF_ANOMALY_LOW",   "PF z-score < −3.0σ",       "CRITICAL", "PF dropped sharply vs EMA baseline — unusual event"),
    ("PF_TREND_DOWN",    "PF z-score < −2.0σ",        "WARNING",  "PF trending downward vs EMA baseline"),
    ("DEMAND_SPIKE",     "kW z-score > +3.0σ",        "CRITICAL", "Sudden demand spike well above typical load"),
    ("DEMAND_ELEVATED",  "kW z-score > +2.0σ",        "WARNING",  "Demand noticeably above typical — monitor"),
    ("REACTIVE_SURGE",   "|kVAr| z-score > +3.0σ",   "WARNING",  "Reactive load surge — large inductive equipment switched on"),
]
make_table(
    doc,
    headers=["Event Subtype", "Trigger", "Severity", "Description"],
    rows=stat_rows,
    col_widths=WE,
    hdr_bg=BG_DARKGREY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED, BG_LTORAN, BG_LTRED, BG_LTORAN, BG_LTORAN],
    row_fgs=[C_BLACK] * 5,
)

# Composite
add_body(doc, "Composite Events — two conditions must be simultaneously true",
         size=9.5, color=C_NAVY, space_before=6, space_after=2)
comp_rows = [
    ("LOW_PF_HIGH_DEMAND", "PF < 0.98  AND  kW > 90% of last 1-hr peak",   "CRITICAL",
     "Double fault — I²R losses compounded by both low PF and high current"),
    ("PF_VOLTAGE_SAG",     "PF < 0.98  AND  voltage_V < 220 V",             "WARNING",
     "High reactive current causing resistive voltage drop along distribution cable"),
]
make_table(
    doc,
    headers=["Event Subtype", "Trigger", "Severity", "Description"],
    rows=comp_rows,
    col_widths=WE,
    hdr_bg=BG_DARKGREY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTRED, BG_LTORAN],
    row_fgs=[C_BLACK] * 2,
)

# Scheduled
add_body(doc, "Scheduled Events — Singapore time-of-use tariff",
         size=9.5, color=C_NAVY, space_before=6, space_after=2)
sched_rows = [
    ("PEAK_PERIOD",    "Hour 08:00 – 21:59",  "INFO", "On-peak — demand charges active; prioritise storage discharge"),
    ("OFFPEAK_PERIOD", "Hour 22:00 – 07:59",  "INFO", "Off-peak — opportunity window to charge battery storage"),
]
make_table(
    doc,
    headers=["Event Subtype", "Trigger", "Severity", "Description"],
    rows=sched_rows,
    col_widths=WE,
    hdr_bg=BG_DARKGREY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTBLUE, BG_LTBLUE],
    row_fgs=[C_BLACK] * 2,
)


# ── 2.5 Reward Computation ────────────────────────────────────────
add_subheading(doc, "2.5  Reward Computation  (SAR Loop Feedback)", space_before=12)
add_body(doc,
    "Source: agent2/outcome.py  |  "
    "Computed after every action to close the State → Action → Reward loop. "
    "Reward drives nightly model retraining.",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WW = [Inches(1.5), Inches(2.0), Inches(2.8)]
rew_rows = [
    ("r_PF  (PF improvement)",
     "r_PF = PF_after − PF_before",
     "Positive = injection improved PF; negative = overcorrected"),
    ("r_loss  (Loss fraction)",
     "r_loss = 1 − (S_after / S_before)²\n= 1 − (P² + Q_after²) / (P² + Q_before²)",
     "Fraction of I²R cable losses eliminated. R cancels — no cable modelling needed."),
    ("r_total  (Combined reward)",
     "r_total = 0.60 × r_PF + 0.40 × r_loss",
     "PF weighted 60% (affects SP penalty directly); loss weighted 40% (quantifies kWh savings)"),
    ("THD back-calculation",
     "THD = √( (PF_before / PF_target)² / (1 − r_loss) − 1 )",
     "Estimates site harmonic distortion without a power analyser. "
     "Starting assumption: 15% THD for mixed buildings."),
]
make_table(
    doc,
    headers=["Component", "Formula", "Interpretation"],
    rows=rew_rows,
    col_widths=WW,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTBLUE if i % 2 == 0 else BG_WHITE for i in range(len(rew_rows))],
    row_fgs=[C_BLACK] * len(rew_rows),
)

add_body(doc, "Outcome classification:", size=9.5, color=C_NAVY,
         space_before=6, space_after=2)
WO = [Inches(1.1), Inches(1.5), Inches(3.7)]
out_rows = [
    ("POSITIVE", "r_PF ≥ +0.01",          "Measurable PF improvement — action was beneficial"),
    ("NEUTRAL",  "−0.01 < r_PF < +0.01",  "Below detection threshold — no significant change"),
    ("NEGATIVE", "r_PF ≤ −0.01",          "PF degraded — action was counterproductive (overcorrection or mode mismatch)"),
]
make_table(
    doc,
    headers=["Outcome", "Condition", "Meaning"],
    rows=out_rows,
    col_widths=WO,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTGREEN, BG_LTBLUE, BG_LTRED],
    row_fgs=[C_BLACK] * 3,
)


# ── 2.6 Model Sizing Rule ─────────────────────────────────────────
add_subheading(doc, "2.6  HyESys Model Sizing Recommendation", space_before=12)
add_body(doc,
    "Source: agent2/tools.py  |  "
    "Sizing headroom rule: required_kVA = avg_kVAr × 1.2  (20% headroom above average reactive load)  |  "
    "No-solar sites hard-capped at H50 (SCDF and space constraints).",
    size=9, color=RGBColor(0x60, 0x60, 0x60))

WM = [Inches(0.8), Inches(1.3), Inches(1.2), Inches(1.0), Inches(1.5), Inches(0.5)]
model_rows = [
    ("H30",  "≤ 30 kVA",  "30 kVA",  "43.5 A",  "Any site",           "$100,000"),
    ("H50",  "≤ 50 kVA",  "50 kVA",  "72.5 A",  "Any site (no-solar cap)", "$120,000"),
    ("H60",  "≤ 60 kVA",  "60 kVA",  "87 A",    "Solar sites only",   "TBD"),
    ("H100", "≤ 100 kVA", "100 kVA", "145 A",   "Solar sites only",   "TBD"),
    ("H125", "> 100 kVA", "125 kVA", "181 A",   "Solar sites only",   "$100,000"),
]
make_table(
    doc,
    headers=["Model", "Required kVA range", "Rated kVA", "Max Current", "Eligibility", "Price (SGD)"],
    rows=model_rows,
    col_widths=WM,
    hdr_bg=BG_NAVY, hdr_fg=C_WHITE,
    row_bgs=[BG_LTBLUE if i % 2 == 0 else BG_WHITE for i in range(len(model_rows))],
    row_fgs=[C_BLACK] * len(model_rows),
)
add_body(doc,
    "Note: H100 and H125 share the same 22-pack battery configuration (217.8 kWh). "
    "H125 delivers higher kVA output from the same hardware.",
    size=8.5, color=RGBColor(0x50, 0x50, 0x50))


# ── Footer rule ───────────────────────────────────────────────────
doc.add_paragraph()
hr2_p  = doc.add_paragraph()
hr2_pf = hr2_p.paragraph_format
hr2_pf.space_before = Pt(8)
hr2_pf.space_after  = Pt(4)
pPr2 = hr2_p._p.get_or_add_pPr()
pBdr2 = OxmlElement("w:pBdr")
top2  = OxmlElement("w:top")
top2.set(qn("w:val"),   "single")
top2.set(qn("w:sz"),    "4")
top2.set(qn("w:space"), "1")
top2.set(qn("w:color"), "AAAAAA")
pBdr2.append(top2)
pPr2.append(pBdr2)

ft = doc.add_paragraph()
ft.paragraph_format.space_after = Pt(0)
r_ft = ft.add_run(
    "Advancer Smart Technology Pte Ltd  ·  HyESys Department  ·  May 2026  ·  Confidential  "
    "·  Generated by HyESys Agent (Claude Code)"
)
r_ft.font.size      = Pt(7.5)
r_ft.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_ft.font.name      = "Calibri"
ft.alignment        = WD_ALIGN_PARAGRAPH.CENTER


# ── Save ──────────────────────────────────────────────────────────
doc.save(DST)
print(f"Saved: {DST}")
