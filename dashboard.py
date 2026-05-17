"""
HyESys Master Dashboard
Run:  streamlit run dashboard.py
Deps: pip install streamlit pandas python-docx
"""

import io
import math
from datetime import datetime
from pathlib import Path

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.store import (
    get_connection, init_db,
    get_sites, read_clean_records, read_sar,
    get_sar_summary, get_record_count,
)
from core.schema import (
    HYESYS_MODELS, SITE_CONFIG,
    PF_TARGET, PF_PENALTY_THRESHOLD, THD_ASSUMPTION,
)

# ── Page config ────────────────────────────────────────────────────
st.set_page_config(
    page_title="HyESys Dashboard",
    layout="wide",
    initial_sidebar_state="expanded",
)

BLUE   = "#2c5f8a"
ACCENT = "#4a9fd4"


# ── DB connection (cached for session) ────────────────────────────
@st.cache_resource
def _get_conn():
    conn = get_connection()
    init_db(conn)
    return conn

conn = _get_conn()


# ── Helper: aggregate SAR rows into per-site stats ────────────────

def _aggregate_sar(site_id: str) -> dict:
    """
    Reads SAR log for one site and returns aggregated savings metrics.

    kWh saved estimate per 15-min interval:
      P_loss_base = state_kW × THD_ASSUMPTION     [approx I²R at nominal]
      kWh_saved   = reward_fraction × P_loss_base × 0.25   [15-min → hours]
    Summed only over POSITIVE-outcome intervals.
    """
    rows = [dict(r) for r in read_sar(conn, site_id)]
    if not rows:
        return {}

    total       = len(rows)
    pos         = sum(1 for r in rows if r["outcome"] == "POSITIVE")
    neu         = sum(1 for r in rows if r["outcome"] == "NEUTRAL")
    neg         = sum(1 for r in rows if r["outcome"] == "NEGATIVE")

    pf_deltas   = [r["reward_pf_delta"]  for r in rows if r["reward_pf_delta"]  is not None]
    fractions   = [r["reward_fraction"]  for r in rows if r["reward_fraction"]   is not None]
    pf_values   = [abs(r["state_PF"])    for r in rows if r["state_PF"]          is not None]

    avg_pf_delta  = sum(pf_deltas) / len(pf_deltas)   if pf_deltas  else 0.0
    avg_fraction  = sum(fractions) / len(fractions)   if fractions  else 0.0
    avg_pf        = sum(pf_values) / len(pf_values)   if pf_values  else 0.0

    # kWh saved: POSITIVE rows only
    kwh_saved = sum(
        (r["reward_fraction"] or 0) * r["state_kW"] * THD_ASSUMPTION * 0.25
        for r in rows if r["outcome"] == "POSITIVE"
    )

    actions = {}
    for r in rows:
        actions[r["action"]] = actions.get(r["action"], 0) + 1

    timestamps = sorted(r["timestamp"] for r in rows)

    return {
        "total": total, "positive": pos, "neutral": neu, "negative": neg,
        "pos_pct": pos / total * 100 if total else 0,
        "avg_pf_delta": round(avg_pf_delta, 4),
        "avg_fraction": round(avg_fraction, 4),
        "avg_pf": round(avg_pf, 4),
        "kwh_saved": round(kwh_saved, 2),
        "actions": actions,
        "start_ts": timestamps[0]  if timestamps else "—",
        "end_ts":   timestamps[-1] if timestamps else "—",
        "rows": rows,
    }


def _get_meter_dates(site_id: str) -> tuple[str, str]:
    rows = read_clean_records(conn, site_id)
    if not rows:
        return "—", "—"
    ts = sorted(r["timestamp"] for r in rows)
    return ts[0], ts[-1]


# ── Sidebar ────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f"### HyESys")
    st.markdown("**Master Dashboard**")
    st.markdown("---")
    st.markdown("**Project:** HyESys Multi-Agent Pipeline")
    st.markdown("**Company:** Advancer Smart Technology Pte Ltd")
    st.markdown("**Engineer:** Jason Ong")
    st.markdown(f"**PF Target:** {PF_TARGET}")
    st.markdown(f"**SP Penalty Below:** {PF_PENALTY_THRESHOLD}")
    st.markdown("---")
    if st.button("Refresh Data", use_container_width=True):
        st.cache_resource.clear()
        st.rerun()
    st.caption(f"Updated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")


# ── Title ──────────────────────────────────────────────────────────
st.title("HyESys Multi-Agent Dashboard")
st.caption("Active Digital Power Compensator — Reactive Compensation | Demand Shaving | Solar Storage")
st.divider()


# ── Tabs ───────────────────────────────────────────────────────────
tab_overview, tab_models, tab_savings, tab_report = st.tabs([
    "Overview", "Model Specs", "Site Savings", "Download Report"
])


# ══════════════════════════════════════════════════════════════════
# TAB 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════
with tab_overview:
    st.subheader("Project Overview")
    st.markdown(
        "HyESys is an **active digital power compensator** that simultaneously delivers "
        "reactive compensation (kVAr injection), 3-phase load balancing, and solar/storage "
        "peak shaving. All three functions share the unit's rated output."
    )

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### Agent 1 — Data Quality & Ingestion")
        st.markdown("""
- Validates raw CSV meter data
- Tags each row **CLEAN / SUSPECT / REJECTED**
- Rules: duplicate timestamps, PF firmware bugs, zero rows, data gaps, mixed formats
- 100% rule-based — no LLM, edge-deployable
- Only CLEAN rows passed to Agent 2
        """)

    with col2:
        st.markdown("#### Agent 2 — Analysis & Recommendation")
        st.markdown("""
- Event-driven — reads 15-min state snapshots from `hyesys.db`
- Detects: Threshold / Statistical / Composite / Scheduled events
- Issues **INJECT_KVAR / HOLD / REDUCE** decisions
- PI controller with dead-band (ε = 0.005)
- Logs **STATE → ACTION → REWARD** (SAR) triplets
- Nightly model retraining from SAR history (<5 s)
        """)

    st.divider()
    st.subheader("Codebase Structure")

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("""
**Core pipeline**
```
agent1/validator.py     Data quality rules
agent1/simulator.py     Historical CSV runner
agent2/state.py         15-min State struct + equations
agent2/tools.py         PF correction, demand risk, loss reduction
agent2/events.py        Event detection (EMA z-score, thresholds)
agent2/outcome.py       Reward computation (SAR loop close)
agent2/agent.py         Decision engine + PI controller
```
        """)
    with col_b:
        st.markdown("""
**Support modules**
```
core/schema.py          Constants, model specs, site config
core/parser.py          CSV ingest + timestamp normalisation
core/store.py           SQLite read/write (meter_records, sar_log)
models/site_model.py    Per-site model: reactive load curve
models/savings.py       kWh savings + THD back-calculation
train.py                Nightly retraining script
main.py                 Pipeline entry point
dashboard.py            This dashboard
```
        """)

    st.divider()
    st.subheader("Key Equations")
    col_eq1, col_eq2 = st.columns(2)
    with col_eq1:
        st.markdown("""
**Power:**
```
S  = √(P² + Q²)          [kVA]
PF = P / S = cos(φ)
φ  = arccos(PF)
I  = S / (√3 × V_L)      [A]
```
**PF Correction:**
```
Q_target = P × tan(arccos(0.98))
ΔQ       = Q_current − Q_target
f        = 1 − (S_after / S_before)²
```
        """)
    with col_eq2:
        st.markdown("""
**Reward:**
```
r_PF    = PF_after − PF_before
r_loss  = 1 − (S_after / S_before)²
r_total = 0.6 × r_PF + 0.4 × r_loss
```
**THD Back-calculation:**
```
THD = √( (PF_before/PF_target)² / (1−f) − 1 )
```
        """)


# ══════════════════════════════════════════════════════════════════
# TAB 2 — MODEL SPECS
# ══════════════════════════════════════════════════════════════════
with tab_models:
    st.subheader("HyESys Available Models")
    st.caption("Source: HyESy.HySBatt Datasheet, Section 2 — May 2026, Version 2")

    MAX_CURRENT = {"H30": 43.5, "H50": 72.5, "H60": 87.0, "H100": 145.0, "H125": 181.0}

    model_rows = []
    for model, specs in HYESYS_MODELS.items():
        price = specs.get("price_sgd")
        model_rows.append({
            "Model":           model,
            "kVA":             specs["kVA"],
            "Max Current (A)": MAX_CURRENT.get(model, "—"),
            "Storage (kWh)":   specs["kWh"],
            "HySBatt Packs":   specs["packs"],
            "Price (SGD)":     f"${price:,}" if price else "TBD",
        })

    df_models = pd.DataFrame(model_rows).set_index("Model")
    st.dataframe(df_models, use_container_width=True)

    st.markdown("""
**Notes:**
- H100 and H125 share the same battery configuration (22 packs, 217.8 kWh) — H125 delivers higher kVA output
- No-solar sites capped at H50 (SCDF and space constraints)
- Sizing rule: model kVA ≥ avg_kVAr × 1.2 (20% headroom)
- Each HySBatt pack: 10 kWh usable, 35 V nominal, 1,250×500×550 mm, <200 kg, IP54
    """)

    st.divider()
    st.subheader("Known Deployment Sites")

    site_rows = []
    for site_id, cfg in SITE_CONFIG.items():
        site_rows.append({
            "Site ID":           site_id,
            "Solar":             "Yes" if cfg["solar"] else "No",
            "Recommended Model": cfg["recommended_model"],
            "Notes":             cfg.get("notes", ""),
        })

    df_sites = pd.DataFrame(site_rows).set_index("Site ID")
    st.dataframe(df_sites, use_container_width=True)


# ══════════════════════════════════════════════════════════════════
# TAB 3 — SITE SAVINGS
# ══════════════════════════════════════════════════════════════════
with tab_savings:
    st.subheader("Site Savings Summary")

    db_sites = get_sites(conn)

    if not db_sites:
        st.info("No site data in database yet. Run Agent 1 first:  `python main.py --csv data/yoursite.csv`")
    else:
        # ── All-sites summary table ──────────────────────────────
        summary_rows = []
        for sid in db_sites:
            counts    = get_record_count(conn, sid)
            s         = _aggregate_sar(sid)
            dt_start, dt_end = _get_meter_dates(sid)
            summary_rows.append({
                "Site":          sid,
                "Records":       sum(counts.values()),
                "CLEAN":         counts.get("CLEAN", 0),
                "Decisions":     s.get("total", 0),
                "Positive %":    f"{s.get('pos_pct', 0):.1f}%",
                "Avg PF":        f"{s.get('avg_pf', 0):.3f}",
                "Avg PF Delta":  f"{s.get('avg_pf_delta', 0):+.4f}",
                "Avg Loss Frac": f"{s.get('avg_fraction', 0):.4f}",
                "Est kWh Saved": f"{s.get('kwh_saved', 0):.1f}",
                "From":          dt_start[:10] if dt_start != "—" else "—",
                "To":            dt_end[:10]   if dt_end   != "—" else "—",
            })

        df_summary = pd.DataFrame(summary_rows).set_index("Site")
        st.dataframe(df_summary, use_container_width=True)

        st.divider()

        # ── Per-site drill-down ──────────────────────────────────
        selected = st.selectbox("Drill down into site:", db_sites)
        s = _aggregate_sar(selected)

        if not s:
            st.info(f"No SAR decisions logged for {selected} yet.")
        else:
            # KPI metrics
            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("Total Decisions",    s["total"])
            m2.metric("Positive Outcomes",  f"{s['positive']}  ({s['pos_pct']:.0f}%)")
            m3.metric("Avg PF",             f"{s['avg_pf']:.3f}")
            m4.metric("Avg Loss Fraction",  f"{s['avg_fraction']*100:.2f}%")
            m5.metric("Est. kWh Saved",     f"{s['kwh_saved']:.1f} kWh")

            col_left, col_right = st.columns(2)

            # Outcome distribution
            with col_left:
                st.markdown("**SAR Outcome Distribution**")
                df_outcomes = pd.DataFrame({
                    "Outcome": ["POSITIVE", "NEUTRAL", "NEGATIVE"],
                    "Count":   [s["positive"], s["neutral"], s["negative"]],
                }).set_index("Outcome")
                st.bar_chart(df_outcomes)

            # Action distribution
            with col_right:
                st.markdown("**Action Distribution**")
                if s["actions"]:
                    df_actions = pd.DataFrame(
                        list(s["actions"].items()), columns=["Action", "Count"]
                    ).set_index("Action")
                    st.bar_chart(df_actions)

            # PF trend
            st.markdown("**PF Trend Over Time**")
            rows = s["rows"]
            if len(rows) > 1:
                step = max(1, len(rows) // 200)   # downsample to 200 points max
                df_pf = pd.DataFrame([
                    {"timestamp": r["timestamp"][:16], "PF": abs(r["state_PF"])}
                    for r in rows[::step]
                ]).set_index("timestamp")
                st.line_chart(df_pf)
                st.caption(f"Data from {s['start_ts'][:10]} to {s['end_ts'][:10]}")

                # PF target reference line note
                st.caption(f"PF target = {PF_TARGET}  |  SP penalty threshold = {PF_PENALTY_THRESHOLD}")


# ══════════════════════════════════════════════════════════════════
# TAB 4 — DOWNLOAD REPORT
# ══════════════════════════════════════════════════════════════════
with tab_report:
    st.subheader("Download Savings Report")
    st.markdown("Generates a Word document summarising all sites, model specs, and savings since project start.")

    def _set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _header_cell(cell, text: str):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "2C5F8A")

    def _data_cell(cell, text: str, bold: bool = False):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.size = Pt(9)
        run.font.bold = bold

    def generate_word_report() -> bytes:
        doc = Document()

        # ── Page margins ──────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # ── Cover ─────────────────────────────────────────────────
        title = doc.add_heading("HyESys Savings Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Generated:     {datetime.now().strftime('%d %B %Y  %H:%M')}\n"
            f"Engineer:      Jason Ong  |  jason@advancer.sg\n"
            f"Company:       Advancer Smart Technology Pte Ltd\n"
            f"PF Target:     {PF_TARGET}    SP Penalty Threshold: {PF_PENALTY_THRESHOLD}\n"
            f"THD Assumption: {THD_ASSUMPTION*100:.0f}% (mixed building)"
        )
        doc.add_page_break()

        # ── Section 1: Agent Roles ─────────────────────────────────
        doc.add_heading("1.  Multi-Agent Architecture", level=1)

        doc.add_heading("Agent 1 — Data Quality & Ingestion", level=2)
        doc.add_paragraph(
            "Validates raw 15-minute CSV meter data and tags each row CLEAN / SUSPECT / REJECTED. "
            "Rules cover duplicate timestamps, PF firmware saturation, zero rows, voltage out of range, "
            "and mixed timestamp formats. Only CLEAN records are passed to Agent 2. "
            "Fully rule-based — no LLM dependency, suitable for edge deployment."
        )

        doc.add_heading("Agent 2 — Analysis & Recommendation", level=2)
        doc.add_paragraph(
            "Event-driven engine that reads validated 15-minute state snapshots from hyesys.db and "
            "issues injection decisions (INJECT_KVAR / HOLD / REDUCE). Detects four event classes: "
            "Threshold (fixed PF limits), Statistical (EMA z-score anomalies), Composite (multi-condition), "
            "and Scheduled (peak/off-peak periods). Uses a PI controller with dead-band (epsilon = 0.005) "
            "for reactive correction. Logs STATE-ACTION-REWARD (SAR) triplets for nightly model retraining."
        )

        # ── Section 2: Model Specs ─────────────────────────────────
        doc.add_heading("2.  HyESys Model Specifications", level=1)

        MAX_CURR = {"H30": 43.5, "H50": 72.5, "H60": 87.0, "H100": 145.0, "H125": 181.0}
        headers  = ["Model", "kVA", "Max Current (A)", "Storage (kWh)", "Packs", "Price (SGD)"]
        tbl      = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            _header_cell(tbl.rows[0].cells[i], h)
        for model, specs in HYESYS_MODELS.items():
            row   = tbl.add_row()
            price = specs.get("price_sgd")
            vals  = [
                model,
                str(specs["kVA"]),
                str(MAX_CURR.get(model, "—")),
                str(specs["kWh"]),
                str(specs["packs"]),
                f"${price:,}" if price else "TBD",
            ]
            for i, v in enumerate(vals):
                _data_cell(row.cells[i], v, bold=(i == 0))

        doc.add_paragraph(
            "\nSizing rule: recommended model kVA >= avg kVAr x 1.2 (20% headroom). "
            "No-solar sites capped at H50 due to SCDF and space constraints."
        )

        # ── Section 3: Known Sites ─────────────────────────────────
        doc.add_heading("3.  Known Deployment Sites", level=1)

        s_headers = ["Site ID", "Solar", "Recommended Model", "Notes"]
        s_tbl     = doc.add_table(rows=1, cols=len(s_headers))
        s_tbl.style = "Table Grid"
        for i, h in enumerate(s_headers):
            _header_cell(s_tbl.rows[0].cells[i], h)
        for site_id, cfg in SITE_CONFIG.items():
            row = s_tbl.add_row()
            vals = [site_id, "Yes" if cfg["solar"] else "No",
                    cfg["recommended_model"], cfg.get("notes", "")]
            for i, v in enumerate(vals):
                _data_cell(row.cells[i], v, bold=(i == 0))

        # ── Section 4: Savings Summary ─────────────────────────────
        doc.add_heading("4.  Site Savings Summary", level=1)
        doc.add_paragraph(
            "Savings computed from the SAR log in hyesys.db. "
            "Estimated kWh saved per 15-min interval = reward_fraction x state_kW x THD_assumption x 0.25 hr, "
            "summed over POSITIVE-outcome decisions only. "
            "Loss fraction = 1 - (S_after / S_before)^2 (cable R cancels — no impedance modelling required)."
        )

        db_sites_local = get_sites(conn)

        if not db_sites_local:
            doc.add_paragraph("No site data in database yet. Run Agent 1 to ingest CSV data first.")
        else:
            sav_headers = [
                "Site", "Records", "Decisions",
                "POSITIVE", "NEUTRAL", "NEGATIVE", "Positive %",
                "Avg PF", "Avg Loss Frac", "Est kWh Saved",
                "From", "To",
            ]
            sav_tbl = doc.add_table(rows=1, cols=len(sav_headers))
            sav_tbl.style = "Table Grid"
            for i, h in enumerate(sav_headers):
                _header_cell(sav_tbl.rows[0].cells[i], h)

            for sid in db_sites_local:
                counts        = get_record_count(conn, sid)
                s             = _aggregate_sar(sid)
                dt_start, dt_end = _get_meter_dates(sid)
                row = sav_tbl.add_row()
                vals = [
                    sid,
                    str(sum(counts.values())),
                    str(s.get("total", 0)),
                    str(s.get("positive", 0)),
                    str(s.get("neutral",  0)),
                    str(s.get("negative", 0)),
                    f"{s.get('pos_pct', 0):.1f}%",
                    f"{s.get('avg_pf', 0):.3f}",
                    f"{s.get('avg_fraction', 0):.4f}",
                    f"{s.get('kwh_saved', 0):.1f}",
                    dt_start[:10] if dt_start != "—" else "—",
                    dt_end[:10]   if dt_end   != "—" else "—",
                ]
                for i, v in enumerate(vals):
                    _data_cell(row.cells[i], v, bold=(i == 0))

        # ── Section 5: Savings Philosophy ─────────────────────────
        doc.add_heading("5.  Measurement & Savings Philosophy", level=1)
        doc.add_paragraph(
            "Primary validation observable: I_rms at the MSB incomer (not cable impedance modelling). "
            "Cable resistance R is age, temperature, and length dependent — impractical to model directly. "
            "R cancels in the savings fraction: f = 1 - (I_after / I_before)^2 = 1 - (S_after / S_before)^2. "
            "\n\n"
            "kW_meter = kW_loads + I2R_distribution_losses. "
            "Load kW is constant; distribution losses scale with I-squared. "
            "HyESys eliminates reactive and harmonic currents at the MSB — I drops — I2R losses drop — "
            "kW at meter drops. When presenting savings to customers, anchor to measured current reduction "
            "at the incomer."
            "\n\n"
            "THD back-calculation: THD = sqrt( (PF_before / PF_target)^2 / (1 - f) - 1 ). "
            "Starting assumption: 15% THD (mixed commercial building). "
            "A site THD > 20% indicates active harmonic filtering is also required."
        )

        # ── Footer note ────────────────────────────────────────────
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Report generated by HyESys Multi-Agent Dashboard  |  "
            f"Advancer Smart Technology Pte Ltd  |  {datetime.now().strftime('%d %b %Y')}"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── Generate + download ────────────────────────────────────────
    col_btn, col_info = st.columns([1, 3])

    with col_btn:
        if st.button("Generate Report", use_container_width=True, type="primary"):
            with st.spinner("Building Word document..."):
                report_bytes = generate_word_report()
            filename = f"HyESys_Savings_Report_{datetime.now().strftime('%Y%m%d')}.docx"
            st.download_button(
                label="Download .docx",
                data=report_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    with col_info:
        st.markdown("""
**Report contains:**
- Agent 1 & 2 roles and architecture summary
- HyESys model specifications table (H30 → H125)
- Known deployment sites with recommended models
- Per-site savings: decisions, outcomes, avg PF, loss fraction, est. kWh saved
- Measurement & savings philosophy (current reduction method, THD back-calc)
        """)
