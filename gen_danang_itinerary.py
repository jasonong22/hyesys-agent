# -*- coding: utf-8 -*-
"""
Danang + Hoi An 5D4N Travel Itinerary Generator
27-31 August 2026 | 2 Pax from Singapore
"""
import sys
import io as _io
# Force UTF-8 stdout so emojis don't crash on Windows cp1252 console
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
sys.stdout = _io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import os
import io
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.request
import tempfile

OUTPUT_PATH = r"C:\Users\JasonOng\Desktop\local docs\personal\viet\Danang_HoiAn_5D4N_Aug2026.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1A, 0x3A, 0x5C)   # deep navy – headings
GOLD     = RGBColor(0xC9, 0xA0, 0x2E)   # golden – accents / table headers
TEAL     = RGBColor(0x00, 0x7E, 0x8A)   # teal – sub-headings
CREAM    = RGBColor(0xFF, 0xF8, 0xF0)   # cream – table alt rows (XML only)
DARK     = RGBColor(0x22, 0x22, 0x22)   # near-black – body text
LIGHT_BG = "FFF3E0"                      # orange-cream for table header bg (hex str)
ALT_ROW  = "FFF8F0"                      # cream alt row
WHITE    = "FFFFFF"

# ── Image URLs ──────────────────────────────────────────────────────────────
IMAGES = {
    # Using Wikimedia REST API thumbnails — different endpoint, more permissive rate limits
    # Format: https://en.wikipedia.org/w/api.php?action=query&titles=File:XXX&prop=imageinfo&iiprop=url&iiurlwidth=640
    # We resolve these at runtime using the _resolve_wikimedia_url() helper
    "my_khe_beach": (
        "wikimedia:My_Khe_Beach_Da_Nang.jpg",
        "My Khe Beach, Da Nang"
    ),
    "golden_bridge": (
        "wikimedia:Golden_Bridge_above_the_clouds_Ba_Na_Hills_Da_Nang_Vietnam.jpg",
        "Golden Bridge, Ba Na Hills"
    ),
    "marble_mountains": (
        "wikimedia:Buddha_Statue%2C_Marble_Mountain.jpg",
        "Buddha Statue, Marble Mountains"
    ),
    "hoi_an_lanterns": (
        "wikimedia:Lanterns_in_Hoi_An_4.jpg",
        "Lanterns in Hoi An"
    ),
    "hoi_an_ancient": (
        "wikimedia:Hoi_An_Ancient_Town%2C_Vietnam_%287090653523%29.jpg",
        "Hoi An Ancient Town"
    ),
    "an_bang_beach": (
        "wikimedia:Playa_An_Bang%2C_Hoi_An%2C_Vietnam_%2839510560665%29.jpg",
        "An Bang Beach, Hoi An"
    ),
    "japanese_bridge": (
        "wikimedia:Hoi_An_Ancient_town.jpg",
        "Hoi An Ancient Town Street"
    ),
    "ba_na_hills": (
        "wikimedia:Da_Nang_Golden_Bridge%2C_Sun_World_Ba_Na_Hills.jpg",
        "Da Nang Golden Bridge, Sun World Ba Na Hills"
    ),
}

FALLBACK_IMAGES = {
    "my_khe_beach": (
        "https://live.staticflickr.com/65535/51234567890_abc123_b.jpg",
        "My Khe Beach, Da Nang"
    ),
    "golden_bridge": (
        "https://upload.wikimedia.org/wikipedia/commons/thumb/b/b5/"
        "Da_Nang_in_the_morning.jpg/1280px-Da_Nang_in_the_morning.jpg",
        "Da Nang"
    ),
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def _resolve_wikimedia_url(filename, width=640):
    """Use Wikimedia Commons REST API to get a thumbnail URL for a file."""
    import time
    import urllib.parse
    # filename may be URL-encoded already; decode first to normalise
    decoded = urllib.parse.unquote(filename)
    encoded = urllib.parse.quote(decoded, safe="")
    api_url = (
        "https://commons.wikimedia.org/w/api.php"
        f"?action=query&titles=File:{encoded}&prop=imageinfo"
        f"&iiprop=url&iiurlwidth={width}&format=json"
    )
    headers = {
        "User-Agent": (
            "DanangTravelItineraryBot/1.0 "
            "(jason@advancer.sg; personal travel document generator)"
        ),
        "Accept": "application/json",
    }
    try:
        r = requests.get(api_url, headers=headers, timeout=15)
        r.raise_for_status()
        data = r.json()
        pages = data.get("query", {}).get("pages", {})
        for page_id, page in pages.items():
            ii = page.get("imageinfo", [])
            if ii:
                thumb_url = ii[0].get("thumburl") or ii[0].get("url")
                if thumb_url:
                    return thumb_url
    except Exception as e:
        print(f"  API resolve failed for {filename}: {e}")
    return None


def download_image(url, label, retries=3, delay=4):
    """Download an image from URL (or wikimedia:filename), return (BytesIO, ext) or (None, None)."""
    import time
    # Resolve Wikimedia filenames via API
    if url.startswith("wikimedia:"):
        filename = url[len("wikimedia:"):]
        resolved = _resolve_wikimedia_url(filename, width=640)
        if resolved:
            url = resolved
            print(f"    Resolved: {url[:80]}...")
        else:
            print(f"  FAIL: Could not resolve Wikimedia URL for {label}")
            return None, None
    headers_list = [
        {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            ),
            "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
            "Referer": "https://www.google.com/",
        },
        {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/605.1.15 (KHTML, like Gecko) "
                "Version/17.0 Safari/605.1.15"
            ),
            "Accept": "image/webp,image/png,image/svg+xml,image/*;q=0.8,*/*;q=0.5",
            "Referer": "https://en.wikipedia.org/",
        },
        {
            "User-Agent": "Wget/1.21.4",
            "Accept": "*/*",
        },
    ]
    for attempt in range(retries):
        headers = headers_list[attempt % len(headers_list)]
        try:
            if attempt > 0:
                time.sleep(delay)
            r = requests.get(url, headers=headers, timeout=30, stream=True)
            r.raise_for_status()
            ct = r.headers.get("Content-Type", "")
            if "jpeg" in ct or "jpg" in ct:
                ext = "jpg"
            elif "png" in ct:
                ext = "png"
            elif "webp" in ct:
                ext = "webp"
            else:
                ext = url.split(".")[-1].split("?")[0].lower() or "jpg"
            buf = io.BytesIO(r.content)
            buf.seek(0)
            print(f"  OK Downloaded: {label}")
            return buf, ext
        except Exception as e:
            print(f"  FAIL (attempt {attempt+1}): {label} -- {e}")
    return None, None


def set_cell_bg(cell, hex_color):
    """Set cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, flag in [("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single" if flag else "none")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_style(para, font_name="Calibri", size=11, bold=False, italic=False,
               colour=None, align=None, space_before=0, space_after=6):
    if align:
        para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if colour:
            run.font.color.rgb = colour


def add_heading(doc, text, level=1, colour=NAVY, size=18, space_before=12,
                space_after=4, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    return p


def add_body(doc, text, size=11, colour=DARK, italic=False, space_after=4,
             space_before=0, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.font.italic = italic
    run.font.bold = bold
    return p


def add_bullet(doc, text, size=11, colour=DARK, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = "Calibri"
        r1.font.size = Pt(size)
        r1.font.bold = True
        r1.font.color.rgb = TEAL
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(size)
        r2.font.color.rgb = colour
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = colour
    return p


def add_horizontal_rule(doc, colour_hex="C9A02E"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _to_jpeg_buf(img_buf):
    """Convert any image format (incl. WebP) to JPEG BytesIO using Pillow."""
    try:
        from PIL import Image as PILImage
        img_buf.seek(0)
        pil_img = PILImage.open(img_buf)
        if pil_img.mode in ("RGBA", "P", "LA"):
            pil_img = pil_img.convert("RGB")
        out = io.BytesIO()
        pil_img.save(out, format="JPEG", quality=85)
        out.seek(0)
        return out
    except Exception as e:
        print(f"  Pillow convert fail: {e}")
        img_buf.seek(0)
        return img_buf


def add_image_to_doc(doc, img_buf, caption, width_inches=5.5):
    if img_buf is None:
        return
    try:
        # Convert to JPEG (handles WebP and other formats python-docx can't read)
        jpeg_buf = _to_jpeg_buf(img_buf)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(jpeg_buf, width=Inches(width_inches))
    except Exception as e:
        print(f"  EMBED FAIL ({caption}): {type(e).__name__}: {e}")
        return
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = cap.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(10)
    run = cap.runs[0]
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_session_block(doc, session_emoji, session_name, details):
    """Add a Morning/Afternoon/Evening block with coloured label."""
    colour_map = {
        "Morning":   GOLD,
        "Afternoon": TEAL,
        "Evening":   NAVY,
    }
    colour = colour_map.get(session_name, NAVY)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    r = p.add_run(f"{session_emoji}  {session_name.upper()}")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = colour

    for bullet_text in details:
        add_bullet(doc, bullet_text)


def add_info_row(doc, label, value):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(value)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    print("Building itinerary document...")

    # ── Download images ─────────────────────────────────────────────────────
    import time
    print("\nDownloading images...")
    imgs = {}
    for key, (url, label) in IMAGES.items():
        buf, ext = download_image(url, label)
        imgs[key] = (buf, label)
        if buf is not None:
            time.sleep(2)  # be polite between successful downloads

    # ── Create document ──────────────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DANANG & HOI AN")
    r.font.name = "Calibri"
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("5 Days  •  4 Nights  •  2 Pax")
    r2.font.name = "Calibri"
    r2.font.size = Pt(18)
    r2.font.bold = False
    r2.font.color.rgb = GOLD

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("27 – 31 August 2026")
    r3.font.name = "Calibri"
    r3.font.size = Pt(16)
    r3.font.color.rgb = TEAL

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Embed cover image — My Khe Beach
    buf, label = imgs.get("my_khe_beach", (None, "My Khe Beach"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.8)
    else:
        buf2, label2 = imgs.get("ba_na_hills", (None, "Ba Na Hills"))
        if buf2:
            add_image_to_doc(doc, buf2, label2, width_inches=5.8)

    # Trip Overview table
    doc.add_paragraph()
    add_heading(doc, "TRIP OVERVIEW", level=2, colour=NAVY, size=14,
                space_before=8, space_after=4)
    add_horizontal_rule(doc)

    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = "Table Grid"
    overview_rows = [
        ("Travellers",     "Jason Ong + 1 companion  (2 pax, ~30 years old)"),
        ("Departure",      "Singapore Changi Airport (SIN)"),
        ("Destination",    "Da Nang International Airport (DAD)"),
        ("Trip Dates",     "Wednesday 27 Aug – Sunday 31 Aug 2026  (5D4N)"),
        ("Base Cities",    "Days 1–2: Da Nang  |  Days 3–4: Hoi An  |  Day 5: Da Nang"),
        ("Flights",        "VietJet Air / Singapore Airlines  (direct, ~2h 50m)"),
        ("Currency",       "Vietnamese Dong (VND) — 1 SGD ≈ 17,000 VND"),
        ("Weather",        "28–33 °C, humid 80–90%, afternoon showers. Pack light + umbrella."),
    ]
    for i, (lbl, val) in enumerate(overview_rows):
        row = tbl.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, "1A3A5C")
        set_cell_bg(c1, ALT_ROW if i % 2 == 0 else WHITE)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(lbl)
        r0.font.name = "Calibri"; r0.font.size = Pt(10); r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = "Calibri"; r1.font.size = Pt(10)
        r1.font.color.rgb = DARK
        for cell in (c0, c1):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    # Set col widths
    for row in tbl.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.0)

    doc.add_paragraph()

    # ── FLIGHT INFO ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "FLIGHTS:  SIN → DAD → SIN", level=1, colour=NAVY,
                size=16, space_before=0, space_after=6)
    add_horizontal_rule(doc)

    add_heading(doc, "Airlines & Booking", level=2, colour=TEAL, size=12,
                space_before=6, space_after=2)
    add_bullet(doc, "VietJet Air (VZ) — Best budget option. Direct, ~2h 50m. "
               "One-way from SGD 90–160 (book 4–6 weeks ahead). Baggage add-on ~SGD 15–25.", )
    add_bullet(doc, "Singapore Airlines (SQ) — Premium comfort. Direct, ~2h 50m. "
               "One-way from SGD 230–350. Includes baggage.")
    add_bullet(doc, "Scoot (TR) — Budget LCC, similar to VietJet pricing, direct.")

    add_heading(doc, "Recommended Flight Times", level=2, colour=TEAL, size=12,
                space_before=6, space_after=2)
    add_bullet(doc, "Outbound (27 Aug Wed): Depart SIN ~07:00–09:00  →  Arrive DAD ~09:50–12:00. "
               "Gives full afternoon. SQ 954 departs SIN 07:05, arrives DAD 09:55.")
    add_bullet(doc, "Return (31 Aug Sun): Depart DAD ~17:00–20:00  →  Arrive SIN ~20:00–23:00. "
               "Maximises final morning in Da Nang.")

    add_heading(doc, "Budget Estimate (Flights, per person)", level=2, colour=TEAL, size=12,
                space_before=6, space_after=2)
    add_bullet(doc, "VietJet Air (return): SGD 180–320 per person (incl. 20 kg baggage add-on)")
    add_bullet(doc, "Singapore Airlines (return): SGD 460–700 per person (incl. 25 kg bag)")
    add_bullet(doc, "Tip: Book SQ outbound + VZ return for comfort + savings")
    add_bullet(doc, "Visa: Singapore passport holders — 45-day visa-free entry to Vietnam")

    # ── HOTELS ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "ACCOMMODATION", level=1, colour=NAVY, size=16,
                space_before=0, space_after=6)
    add_horizontal_rule(doc)

    add_heading(doc, "Da Nang Hotels  (Nights 1 & 2 — 27–29 Aug)", level=2,
                colour=NAVY, size=13, space_before=6, space_after=3)

    hotels_dn = [
        ("LUXURY — Hyatt Regency Danang Resort & Spa",
         "SGD 220–340/night", "My Khe Beach, Non Nuoc area",
         "5-star beachfront. Stunning pool. World-class spa. Highly recommended for couples. "
         "Avg rating 4.7/5 TripAdvisor. Book direct for best rate."),
        ("LUXURY — Fusion Maia Danang",
         "SGD 250–380/night", "Truong Sa Road, My Khe Beach",
         "All-inclusive spa resort. Every room has private pool. Unique adults-only atmosphere. "
         "Exceptional reviews for honeymoon/romantic trips."),
        ("MID-RANGE — Novotel Danang Premier Han River",
         "SGD 100–160/night", "Han River Waterfront, City Centre",
         "4-star city hotel with river views. Infinity rooftop pool. Walking distance to "
         "restaurants and Dragon Bridge. Great value."),
        ("MID-RANGE — Azura Da Nang Hotel",
         "SGD 80–130/night", "An Thuong tourist area, My Khe Beach",
         "4-star, 5-min walk to beach. Pool, gym, great breakfast included. "
         "Strong reviews, friendly staff. Good base for beach days."),
        ("BUDGET — A La Carte Da Nang Beach Hotel",
         "SGD 55–95/night", "My Khe Beach front",
         "3-star with excellent beach location. Rooftop pool and bar. "
         "Rooms simple but clean. Very popular with backpackers."),
    ]
    for name, price, loc, desc in hotels_dn:
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(1)
        r = p.add_run(f"  {name}")
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = NAVY
        add_info_row(doc, "    Price", price)
        add_info_row(doc, "    Location", loc)
        add_body(doc, f"    {desc}", size=10, colour=DARK, space_after=4)

    add_heading(doc, "Hoi An Hotels  (Nights 3 & 4 — 29–31 Aug)", level=2,
                colour=NAVY, size=13, space_before=10, space_after=3)

    hotels_ha = [
        ("LUXURY — La Siesta Hoi An Resort & Spa",
         "SGD 180–280/night", "Le Hong Phong, near Old Town",
         "TripAdvisor 2026 #1 in Vietnam, #2 Asia, #6 World. Spectacular saltwater pool, "
         "200m² infinity pool, world-class spa. 5-min tuk-tuk to Ancient Town."),
        ("LUXURY — Little Hoi An Boutique Hotel & Spa",
         "SGD 130–200/night", "Old Town edge, walking distance",
         "Charming boutique. Traditional lanterns, patterned tiles, warm golden tones. "
         "TripAdvisor Traveller's Choice. Excellent breakfast. Romantic atmosphere."),
        ("MID-RANGE — The Silk River Hotel & Spa",
         "SGD 70–120/night", "Thu Bon Riverfront, east of Old Town",
         "4-star, quiet riverside setting. Pool, spa. 10-min walk to Ancient Town. "
         "Excellent service reviews. Bicycles provided free."),
        ("MID-RANGE — RiverTown Hoi An Resort & Spa",
         "SGD 90–150/night", "An Hoi Island",
         "5-star on An Hoi Island, two pools (adult + family). "
         "Beautiful river views. Short walk to Old Town via pedestrian bridge."),
    ]
    for name, price, loc, desc in hotels_ha:
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(1)
        r = p.add_run(f"  {name}")
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = NAVY
        add_info_row(doc, "    Price", price)
        add_info_row(doc, "    Location", loc)
        add_body(doc, f"    {desc}", size=10, colour=DARK, space_after=4)

    add_body(doc, "Recommendation: Hyatt Regency Danang (nights 1–2) + La Siesta Hoi An (nights 3–4) "
             "for the best overall experience at mid-luxury budget.",
             size=10, italic=True, colour=TEAL, space_before=4, space_after=6)

    # ── DAY 1 ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "DAY 1  —  Wednesday, 27 August 2026", level=1,
                colour=NAVY, size=17, space_before=0, space_after=2)
    add_body(doc, "FLY SIN → DAN  |  Arrive & Settle In  |  Intensity: LOW",
             size=11, bold=True, colour=GOLD, space_after=4)
    add_horizontal_rule(doc)

    add_session_block(doc, "☀", "Morning", [
        "Depart Singapore Changi Airport, Terminal 1. Check in 2 hours early.",
        "Recommended flight: SQ 954 departs 07:05, arrives Da Nang 09:55 (direct, ~2h 50m). "
        "Alternatively VietJet Air morning departure ~08:30 arrival ~11:20.",
        "Tip: Grab a coffee and kaya toast at Toast Box T1 before departure.",
    ])

    add_session_block(doc, "☀", "Afternoon", [
        "Land at Da Nang International Airport (DAD). Clear immigration — Singapore passports "
        "visa-free 45 days. Fast process (~15 min).",
        "Transport to hotel: Open Grab app immediately at arrival hall. Grab Car to My Khe Beach "
        "hotels costs ~80,000–120,000 VND (~SGD 4.70–7.00). Journey ~15–20 min.",
        "Check in to Hyatt Regency Danang (or chosen hotel). Rooms typically ready from 14:00; "
        "drop luggage if early.",
        "Freshen up, then walk along My Khe Beach — 30 km of white sand. Dip your toes in. "
        "The water is calm and warm in August.",
        "Lunch near the hotel: Roly Poly Fresh Spring Rolls (40 Vo Nguyen Giap) — "
        "hand-rolled spring rolls with lemongrass beef, prawn & vegetables. ~SGD 8–12 per person.",
    ])

    # Embed My Khe Beach image
    buf, label = imgs.get("my_khe_beach", (None, "My Khe Beach"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "🌙", "Evening", [
        "Explore Da Nang's Han River Promenade — take a Grab (~SGD 3) to Dragon Bridge area.",
        "Dinner: Madame Lan Restaurant (4 Bach Dang, Han River) — beautifully designed courtyard "
        "mimicking a Vietnamese village. Signature dishes: Grilled Semi-Dried Squid, Banh Xeo "
        "(sizzling crepe). Budget: ~SGD 20–30 per person incl. drinks.",
        "After dinner: Stroll along Bach Dang riverside. The Dragon Bridge lights up at 21:00 "
        "on weekends — check if Sat/Sun show applies to your schedule.",
        "Watch (if Sat/Sun): Dragon Bridge breathes fire and water at 21:00 and 21:30. "
        "Free to watch from riverbank.",
        "Return to hotel via Grab by 22:30. Rest up for an early Day 2.",
    ])

    add_info_row(doc, "Day 1 Transport", "Grab from airport: ~SGD 5–7. Han River Grab: ~SGD 3 each way.")
    add_info_row(doc, "Day 1 Est. Cost", "~SGD 80–130 per person (incl. flight day share + meals + transport)")

    # ── DAY 2 ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "DAY 2  —  Thursday, 28 August 2026", level=1,
                colour=NAVY, size=17, space_before=0, space_after=2)
    add_body(doc, "BA NA HILLS + GOLDEN BRIDGE  |  Intensity: MEDIUM",
             size=11, bold=True, colour=GOLD, space_after=4)
    add_horizontal_rule(doc)

    # Ba Na Hills image
    buf, label = imgs.get("ba_na_hills", (None, "Ba Na Hills"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "☀", "Morning", [
        "Early start — leave hotel by 07:30. Ba Na Hills is best experienced before crowds "
        "arrive (~10:00).",
        "Pre-book Ba Na Hills tickets online via Klook or official Sun World app: "
        "Adult ticket ~950,000 VND (~SGD 56) — includes all cable cars, unlimited rides on "
        "Fantasy Park attractions, and entry to French Village.",
        "Transport: Book a Grab Car or hotel taxi to Ba Na Hills (25 km west of city centre). "
        "~350,000–450,000 VND (~SGD 20–26) one-way. Journey ~40 min.",
        "Cable Car: Ride the world-record-breaking Ba Na Hills cable car — 5.1 km gondola, "
        "gliding above waterfalls and misty forests. Spectacular views.",
        "Golden Bridge: The iconic 150 m golden bridge held by two giant stone hands. "
        "Best photos before 10:00. Clouds roll in by mid-morning for dramatic shots.",
    ])

    # Golden Bridge image
    buf, label = imgs.get("golden_bridge", (None, "Golden Bridge, Ba Na Hills"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "☀", "Afternoon", [
        "Explore the French Village on the mountain — charming neo-Gothic architecture, "
        "wine cellar, flower gardens. Temperature at 1,487m elevation is ~18–22°C "
        "(bring a light layer — it will feel cold vs Da Nang heat).",
        "Fantasy Park amusement rides (included in ticket): roller coaster, wax museum, "
        "4D cinema. Fun for 1–2 hours.",
        "Lunch at Ba Na Hills: multiple restaurants on site. Recommend the buffet lunch "
        "(included in combo ticket ~1,300,000 VND) or à la carte at Fantasy Park food court. "
        "~SGD 15–30 per person.",
        "Cable car descent ~15:00–16:00. Return to Da Nang via Grab.",
    ])

    add_session_block(doc, "🌙", "Evening", [
        "Freshen up at hotel. Head to My Khe Beach for a sunset swim/walk (~17:30–18:30). "
        "The beach faces east, so sunsets are reflected on water and city.",
        "Dinner: Mi Quang 1A (1 Hai Phong Street) — Da Nang's most beloved noodle dish. "
        "Turmeric-yellow noodles with pork, shrimp, roasted peanuts, and fresh herbs. "
        "Legendary local spot. ~40,000–60,000 VND (~SGD 2.50–3.50) per bowl. "
        "Arrive before 19:00 to avoid queues.",
        "Or upgrade to Nén Danang (Michelin Green Star, 27 Tran Thi Ly) for a refined "
        "tasting menu of hyper-local Central Vietnamese ingredients. ~SGD 40–60 per person.",
        "Night cap: Rooftop bar at your hotel or nearby Sky36 Bar (36th floor, Novotel) "
        "for panoramic city views. ~SGD 8–15 per cocktail.",
    ])

    add_info_row(doc, "Day 2 Transport", "Hotel→Ba Na Hills: ~SGD 20–26 one-way. "
                 "Return ~SGD 20–26. City Grab: ~SGD 2–4.")
    add_info_row(doc, "Day 2 Est. Cost", "~SGD 130–180 per person (Ba Na Hills ticket + meals + transport)")

    # ── DAY 3 ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "DAY 3  —  Friday, 29 August 2026", level=1,
                colour=NAVY, size=17, space_before=0, space_after=2)
    add_body(doc, "MARBLE MOUNTAINS  →  TRAVEL TO HOI AN  →  HOI AN ANCIENT TOWN",
             size=11, bold=True, colour=GOLD, space_after=2)
    add_body(doc, "Intensity: MEDIUM  |  Check out Da Nang, check in Hoi An",
             size=10, italic=True, colour=DARK, space_after=4)
    add_horizontal_rule(doc)

    add_session_block(doc, "☀", "Morning", [
        "Check out of Da Nang hotel. Store luggage at hotel while you visit Marble Mountains "
        "(usually OK until early afternoon — confirm with hotel).",
        "Grab taxi to Marble Mountains / Ngu Hanh Son (9 km south of Da Nang, en-route to Hoi An). "
        "~100,000–150,000 VND (~SGD 6–9).",
        "Marble Mountains (Ngu Hanh Son): Five limestone formations named after the five "
        "elements — Water, Wood, Fire, Metal, Earth. Explore Thuy Son (Water Mountain) — "
        "entrance ticket ~40,000 VND (~SGD 2.40). Elevator to top: 15,000 VND (~SGD 0.90).",
        "Cave temples: Huyen Khong Cave — natural skylights pierce the limestone ceiling, "
        "illuminating ancient Buddha shrines. Atmospheric and spiritual.",
        "Panoramic views from the summit over My Khe Beach and the South China Sea.",
        "Spend 1.5–2 hours exploring. Bring water — August heat is intense.",
    ])

    # Marble Mountains image
    buf, label = imgs.get("marble_mountains", (None, "Marble Mountains, Da Nang"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "☀", "Afternoon", [
        "Quick lunch near Marble Mountains or at a roadside stop: "
        "Bun Cha Ca (Da Nang-style fish cake noodle soup) at local eateries — ~30,000–50,000 VND.",
        "Collect luggage from Da Nang hotel.",
        "Travel Da Nang → Hoi An: ~30 km, 30–40 min. "
        "Options: (a) Grab Car — ~250,000–320,000 VND (~SGD 15–19), most convenient with luggage. "
        "(b) Shuttle bus from Da Nang to Hoi An — ~120,000–180,000 VND (~SGD 7–11) per person.",
        "Check in to La Siesta Hoi An Resort & Spa (or chosen Hoi An hotel). "
        "Rooms from 14:00. Pool time to cool down.",
        "Banh Mi Phuong (2B Phan Chau Trinh) — Anthony Bourdain's favourite banh mi in the world. "
        "~30,000–50,000 VND (~SGD 1.80–3.00). Queue is part of the experience. Go before 15:00.",
    ])

    add_session_block(doc, "🌙", "Evening", [
        "First night in Hoi An Ancient Town — buy your entrance ticket (~120,000 VND/~SGD 7) "
        "at the booth on Le Loi Street. Ticket grants entry to 5 heritage sites.",
        "Wander the Japanese Covered Bridge (Chùa Cầu) — the most photographed landmark in "
        "Hoi An. Built in the 1590s by Japanese traders. Beautiful at night with lantern glow.",
        "White Rose Restaurant (533 Hai Ba Trung) — the only authentic source of White Rose "
        "Dumplings (Banh Bao Vac). Watch women fold each dumpling by hand. "
        "~60,000–80,000 VND for a plate (~SGD 3.50–4.70). Must try.",
        "Dinner: Morning Glory Signature Restaurant (106 Nguyen Thai Hoc) — "
        "riverfront balcony seating. Central Vietnamese classics elevated to fine dining. "
        "Chef Vy's signature dishes: Cao Lau, Banh Xeo, White Rose. ~SGD 20–35 per person.",
        "After dinner: River lanterns. Buy silk lanterns on the riverside (~SGD 1–2 each) "
        "and float them on the Thu Bon River for good luck.",
        "Stroll Nguyen Hoang Night Market for handcrafted souvenirs, lanterns, and tailor samples.",
    ])

    # Japanese Bridge image
    buf, label = imgs.get("japanese_bridge", (None, "Japanese Covered Bridge, Hoi An"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_info_row(doc, "Day 3 Transport", "Marble Mtns Grab: ~SGD 6–9. Da Nang→Hoi An Grab: ~SGD 15–19.")
    add_info_row(doc, "Day 3 Est. Cost", "~SGD 80–120 per person (transport + Marble Mtns + meals + Old Town ticket)")

    # ── DAY 4 ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "DAY 4  —  Saturday, 30 August 2026", level=1,
                colour=NAVY, size=17, space_before=0, space_after=2)
    add_body(doc, "AN BANG BEACH  |  COOKING CLASS  |  TAILORS  |  LANTERN FESTIVAL NIGHT",
             size=11, bold=True, colour=GOLD, space_after=2)
    add_body(doc, "Intensity: LOW–MEDIUM  |  Best day of the trip",
             size=10, italic=True, colour=DARK, space_after=4)
    add_horizontal_rule(doc)

    add_session_block(doc, "☀", "Morning", [
        "Breakfast at hotel. Then bicycle ride to An Bang Beach — hire bicycles from hotel "
        "or street rental (~30,000–50,000 VND/~SGD 2–3 per hour). "
        "Scenic 3 km ride through rice paddies and vegetable gardens (~15 min).",
        "An Bang Beach — quieter and more laid-back than My Khe. Soft white sand, "
        "warm turquoise water, beach bars with sun loungers (~50,000–100,000 VND for a lounger). "
        "Swimming is excellent in August.",
        "Beach eats: Soul Kitchen (Beach road, An Bang) — Australian-run beach bar, "
        "excellent breakfasts, smoothies, and fresh seafood. ~SGD 10–18 per person.",
        "Or: La Plage Beach Bar for cocktails and light bites by the water.",
        "Spend 2–3 hours at the beach before the heat peaks (~11:30).",
    ])

    # An Bang Beach image
    buf, label = imgs.get("an_bang_beach", (None, "An Bang Beach, Hoi An"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "☀", "Afternoon", [
        "Return to Old Town. Head to tailor shops — Hoi An tailors are world-famous for "
        "24–48 hour turnaround. Visit on Day 4 (Sat) afternoon to collect on Day 5 Sunday morning.",
        "Recommended tailors: Yaly Couture (358 Nguyen Duy Hieu) — most respected, "
        "professional fittings, wide fabric selection. A Dong Silk (62 Tran Phu) — good variety, "
        "competitive pricing. B'lan Silk (2 Nguyen Hoang) — custom ao dai and western suits.",
        "Prices: Men's 3-piece suit SGD 90–200, women's dress SGD 25–60, custom shirt SGD 20–45. "
        "Bring reference photos. First fitting today; collect tomorrow.",
        "Vietnamese Cooking Class (13:00–17:00): Book in advance via Klook/GetYourGuide. "
        "Recommended: Red Bridge Cooking School (~SGD 45–55 per person, includes market tour + "
        "4 dishes + boat ride to school). Or Morning Glory Cooking Class (~SGD 35–45).",
        "Cooking class teaches: Cao Lau noodles, fresh spring rolls, Banh Xeo, White Rose dumplings. "
        "You eat everything you cook.",
        "Mid-afternoon snack: Thanh Cao Lau (26 Thai Phien) — tiny family noodle shop, "
        "best Cao Lau in town. Arrive before 17:00. ~40,000–60,000 VND (~SGD 2.40–3.50).",
    ])

    add_session_block(doc, "🌙", "Evening", [
        "Dinner at Cargo Club Restaurant (107–109 Nguyen Thai Hoc) — multi-level restaurant "
        "with river views. Vietnamese + Western fusion. Famous for its patisserie and cakes. "
        "~SGD 18–30 per person.",
        "Hoi An Full Moon / Lantern Festival Night: 29 August falls near the full moon "
        "(29th lunar month of July 2026). Check exact date — typically 13th–15th lunar day. "
        "If the festival is on 30 Aug, all electric lights go out and the Ancient Town "
        "glows entirely by lantern light. Absolutely spectacular.",
        "Lantern-Making Workshop (pre-book): 30–45 min sessions in Old Town. "
        "~100,000–150,000 VND (~SGD 6–9). You build your own silk lantern to take home.",
        "Float paper lanterns on the Thu Bon River from Cam Nam Bridge — "
        "beautiful and deeply atmospheric. Lanterns ~SGD 1–2 each.",
        "Riverside stroll: walk from Japanese Bridge to the Night Market, "
        "then along the riverfront. Old Town is at its most beautiful after dark.",
    ])

    # Hoi An Lanterns image
    buf, label = imgs.get("hoi_an_lanterns", (None, "Hoi An Lantern Festival, Thu Bon River"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_info_row(doc, "Day 4 Transport", "Bicycle An Bang: ~SGD 3–5. Old Town tuk-tuk/Grab: ~SGD 2–4.")
    add_info_row(doc, "Day 4 Est. Cost", "~SGD 120–180 per person (cooking class + tailors + meals + activities)")

    # ── DAY 5 ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "DAY 5  —  Sunday, 31 August 2026", level=1,
                colour=NAVY, size=17, space_before=0, space_after=2)
    add_body(doc, "HOI AN → DA NANG  |  LAST MEAL  |  FLY HOME",
             size=11, bold=True, colour=GOLD, space_after=2)
    add_body(doc, "Intensity: LOW  |  Fly home evening",
             size=10, italic=True, colour=DARK, space_after=4)
    add_horizontal_rule(doc)

    add_session_block(doc, "☀", "Morning", [
        "Leisurely breakfast at hotel. Check out by 11:00 (store luggage).",
        "Collect tailored garments from tailor shop (should be ready by 09:00–10:00).",
        "Final wander through the Ancient Town market — buy last-minute souvenirs: "
        "silk scarves, hand-painted lanterns, lacquerware, Vietnamese coffee, and pho spice mix.",
        "Com Linh Restaurant (Cam Pho area) — highly recommended for Hoi An Com Ga "
        "(chicken rice). Iconic late-morning meal. ~50,000–70,000 VND (~SGD 3–4).",
        "Visit Hoi An Market (Cho Hoi An, 46 Tran Phu) for fresh fruit, dried goods, "
        "and local snacks to bring back to Singapore.",
    ])

    # Hoi An Ancient Town image
    buf, label = imgs.get("hoi_an_ancient", (None, "Hoi An Ancient Town & Thu Bon River"))
    if buf:
        add_image_to_doc(doc, buf, label, width_inches=5.0)

    add_session_block(doc, "☀", "Afternoon", [
        "Travel Hoi An → Da Nang Airport: Grab Car (~250,000–320,000 VND/~SGD 15–19). "
        "Journey ~30–40 min. Allow 2.5 hours before departure.",
        "If time allows before leaving Hoi An: Banh Mi Phuong for one last sandwich.",
        "Arrive Da Nang Airport: Check in, clear security. International terminal is compact "
        "— 30 min is sufficient for security.",
        "Lunch/snacks at airport: noodle stalls and cafes airside.",
        "Recommended return flight: SQ 955 departs DAD 17:10, arrives SIN 20:10. "
        "Or VietJet Air evening flight ~19:00 arrives SIN ~22:00.",
    ])

    add_session_block(doc, "🌙", "Evening", [
        "Board flight back to Singapore. Arrive Changi Airport.",
        "Tip: Declare any durian products (prohibited) but all other Vietnamese food items "
        "(packed dried goods, coffee, cookies) are allowed into Singapore.",
        "Last Vietnamese meal in the air or at Changi: Jewel Changi has Pho Hoa if craving strikes.",
    ])

    add_info_row(doc, "Day 5 Transport", "Hoi An→Da Nang Airport Grab: ~SGD 15–20.")
    add_info_row(doc, "Day 5 Est. Cost", "~SGD 50–80 per person (transport + souvenirs + airport + last meals)")

    # ── FOOD GUIDE ────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "FOOD GUIDE — NAMED RESTAURANTS & STREET FOOD", level=1,
                colour=NAVY, size=16, space_before=0, space_after=6)
    add_horizontal_rule(doc)

    add_heading(doc, "Da Nang Eats", level=2, colour=TEAL, size=13,
                space_before=6, space_after=3)

    food_dn = [
        ("Mi Quang 1A", "1 Hai Phong St", "Da Nang's iconic turmeric noodle dish. "
         "Legendary local institution. Arrive early. ~SGD 2.50–3.50/bowl."),
        ("Banh Xeo Ba Duong", "K280/23 Hoang Dieu", "Michelin Bib Gourmand sizzling crepes "
         "(banh xeo) and grilled pork skewers (nem lui). Cash only. ~SGD 5–8/person."),
        ("Madame Lan", "4 Bach Dang St", "Vietnamese village–style courtyard. "
         "Grilled squid, banh xeo, central Vietnamese dishes. ~SGD 15–25/person."),
        ("Nén Danang (Michelin Green Star)", "27 Tran Thi Ly", "Vietnam's first Green Star. "
         "Hyper-local tasting menu with farm-to-table philosophy. ~SGD 40–65/person."),
        ("My Hanh Seafood", "Bai Bien My Khe area", "Freshest seafood in Da Nang. "
         "Whole grilled fish, oysters, prawns. ~SGD 20–35/person."),
        ("Roly Poly Spring Rolls", "40 Vo Nguyen Giap", "Made-to-order fresh spring rolls "
         "near My Khe Beach. Lemongrass beef, prawn options. ~SGD 8–12/person."),
        ("Jeremy's Kitchen (Café)", "An Thuong area", "American expat café. "
         "House-made bread sandwiches, cinnamon rolls, coffee. ~SGD 6–12/person."),
    ]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr = tbl2.rows[0].cells
    for i, h in enumerate(["Restaurant", "Address", "Notes"]):
        set_cell_bg(hdr[i], "1A3A5C")
        r = hdr[i].paragraphs[0].add_run(h)
        r.font.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, (name, addr, note) in enumerate(food_dn):
        row = tbl2.add_row().cells
        bg = ALT_ROW if idx % 2 == 0 else WHITE
        for c in row: set_cell_bg(c, bg)
        for c, txt in zip(row, [name, addr, note]):
            r = c.paragraphs[0].add_run(txt)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            r.font.color.rgb = DARK
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)

    add_heading(doc, "Hoi An Eats", level=2, colour=TEAL, size=13,
                space_before=10, space_after=3)

    food_ha = [
        ("White Rose Restaurant", "533 Hai Ba Trung", "The only authentic White Rose Dumplings "
         "(Banh Bao Vac). Watch live folding. ~SGD 3–5/plate."),
        ("Thanh Cao Lau", "26 Thai Phien St", "Family-run. Best Cao Lau in Hoi An — "
         "thick noodles, roasted pork, crispy croutons. ~SGD 2.50–3.50/bowl."),
        ("Banh Mi Phuong", "2B Phan Chau Trinh", "Anthony Bourdain's famous banh mi. "
         "Queue expected. Worth it. ~SGD 1.50–2.50."),
        ("Morning Glory Signature", "106 Nguyen Thai Hoc", "Fine dining with river views. "
         "Central Vietnamese classics elevated. ~SGD 20–35/person."),
        ("Cargo Club", "107–109 Nguyen Thai Hoc", "Multi-level riverside restaurant. "
         "Patisserie, Vietnamese + Western fusion. Great for sunset dining. ~SGD 18–28/person."),
        ("Com Linh", "Cam Pho area", "Local com ga (chicken rice). Busy and authentic. "
         "Excellent value. ~SGD 3–4/person."),
        ("Soul Kitchen", "An Bang Beach road", "Australian-run beach bar. Breakfasts, "
         "smoothies, fresh seafood. ~SGD 10–18/person."),
    ]
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    hdr3 = tbl3.rows[0].cells
    for i, h in enumerate(["Restaurant", "Address", "Notes"]):
        set_cell_bg(hdr3[i], "007E8A")
        r = hdr3[i].paragraphs[0].add_run(h)
        r.font.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, (name, addr, note) in enumerate(food_ha):
        row = tbl3.add_row().cells
        bg = ALT_ROW if idx % 2 == 0 else WHITE
        for c in row: set_cell_bg(c, bg)
        for c, txt in zip(row, [name, addr, note]):
            r = c.paragraphs[0].add_run(txt)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            r.font.color.rgb = DARK
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)

    # ── BUDGET SUMMARY ────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "BUDGET SUMMARY — PER PERSON (SGD)", level=1,
                colour=NAVY, size=16, space_before=0, space_after=6)
    add_horizontal_rule(doc)

    add_body(doc, "Exchange rate: 1 SGD ≈ 17,000 VND  (as of mid-2026). "
             "All estimates based on mid-range travel style.",
             size=10, italic=True, colour=DARK, space_after=6)

    budget_rows = [
        ("FLIGHTS", "VietJet Air return (incl. 20 kg bag)", "SGD 180", "SGD 320"),
        ("FLIGHTS", "Singapore Airlines return (incl. 25 kg bag)", "SGD 460", "SGD 700"),
        ("", "(Budget with VietJet recommended)", "", ""),
        ("HOTEL (2N Da Nang)", "Hyatt Regency per night (~SGD 280/night ÷ 2)", "SGD 280", "SGD 340"),
        ("HOTEL (2N Hoi An)", "La Siesta per night (~SGD 220/night ÷ 2)", "SGD 220", "SGD 280"),
        ("ACTIVITIES", "Ba Na Hills ticket (per person)", "SGD 56", "SGD 76"),
        ("ACTIVITIES", "Marble Mountains entrance", "SGD 3", "SGD 4"),
        ("ACTIVITIES", "Hoi An Old Town ticket", "SGD 7", "SGD 7"),
        ("ACTIVITIES", "Cooking Class (per person)", "SGD 35", "SGD 55"),
        ("ACTIVITIES", "Lantern workshop (per person)", "SGD 6", "SGD 9"),
        ("FOOD", "Da Nang (3 meals/day × 2 days)", "SGD 35", "SGD 70"),
        ("FOOD", "Hoi An (3 meals/day × 2 days)", "SGD 40", "SGD 80"),
        ("TRANSPORT", "Airport transfers (Grab, both ways)", "SGD 10", "SGD 15"),
        ("TRANSPORT", "Da Nang to Hoi An (Grab)", "SGD 15", "SGD 20"),
        ("TRANSPORT", "Marble Mountains Grab (return)", "SGD 12", "SGD 18"),
        ("TRANSPORT", "Ba Na Hills Grab (return)", "SGD 40", "SGD 52"),
        ("TRANSPORT", "Local Grab / tuk-tuk / bicycles", "SGD 20", "SGD 35"),
        ("SHOPPING", "Tailoring (suit or 2 garments)", "SGD 100", "SGD 250"),
        ("SHOPPING", "Souvenirs, lanterns, silk goods", "SGD 30", "SGD 80"),
        ("MISC", "Tips, entry top-ups, incidentals", "SGD 20", "SGD 40"),
    ]

    tbl_budget = doc.add_table(rows=1, cols=4)
    tbl_budget.style = "Table Grid"
    hdr_b = tbl_budget.rows[0].cells
    for i, h in enumerate(["Category", "Item", "Low (SGD)", "High (SGD)"]):
        set_cell_bg(hdr_b[i], "1A3A5C")
        r = hdr_b[i].paragraphs[0].add_run(h)
        r.font.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_b[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (cat, item, lo, hi) in enumerate(budget_rows):
        row = tbl_budget.add_row().cells
        if not cat and not lo:  # note row
            set_cell_bg(row[0], "E8F4F8")
            set_cell_bg(row[1], "E8F4F8")
            set_cell_bg(row[2], "E8F4F8")
            set_cell_bg(row[3], "E8F4F8")
            r = row[1].paragraphs[0].add_run(item)
            r.font.name = "Calibri"; r.font.size = Pt(9); r.font.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            continue
        bg = ALT_ROW if idx % 2 == 0 else WHITE
        for c in row: set_cell_bg(c, bg)
        for c, txt, bold in zip(row, [cat, item, lo, hi], [True, False, False, False]):
            r = c.paragraphs[0].add_run(txt)
            r.font.name = "Calibri"; r.font.size = Pt(9); r.font.bold = bold
            r.font.color.rgb = NAVY if bold else DARK
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Totals row
    row_t = tbl_budget.add_row().cells
    for c in row_t: set_cell_bg(c, "C9A02E")
    totals = ["TOTAL", "Per Person (VietJet flights + mid-range)", "SGD 929", "SGD 1,371"]
    for c, txt in zip(row_t, totals):
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)

    # Both pax total
    row_2 = tbl_budget.add_row().cells
    for c in row_2: set_cell_bg(c, "1A3A5C")
    totals2 = ["TOTAL (2 PAX)", "Both travellers combined", "SGD 1,858", "SGD 2,742"]
    for c, txt in zip(row_2, totals2):
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)

    # Column widths
    col_widths = [Cm(3.0), Cm(8.5), Cm(2.5), Cm(2.5)]
    for row in tbl_budget.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # ── PRACTICAL TIPS ────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "PRACTICAL TIPS & ESSENTIALS", level=1, colour=NAVY, size=16,
                space_before=0, space_after=6)
    add_horizontal_rule(doc)

    tips = [
        ("Weather & Packing",
         "Late August is hot (28–33°C) with humidity 80–90% and afternoon showers. Pack: "
         "2–3 light cotton/linen outfits, swimwear, reef-safe sunscreen SPF 50+, "
         "travel umbrella/poncho, portable fan, insect repellent. At Ba Na Hills (~1,500m) "
         "it will be cool (~18–22°C) — bring one light jacket."),
        ("Money & Payments",
         "Withdraw VND at airport ATM (Vietcombank/Techcombank for low fees). "
         "Budget ~2,000,000 VND/day for 2 pax street food + transport. "
         "Cards accepted at hotels and mid-range restaurants. "
         "Markets and street food = cash only. 1 SGD ≈ 17,000 VND."),
        ("Grab App",
         "Download Grab before you fly. Link a Singapore card. "
         "Use GrabCar for all intercity trips — fixed fare shown upfront, no haggling. "
         "GrabBike for solo short hops. Avoid unmetered taxis at touristy spots."),
        ("SIM Card",
         "Buy a local Vietnamese SIM at Da Nang Airport on arrival. Viettel or Vinaphone. "
         "10-day tourist SIM with 20 GB data: ~100,000–150,000 VND (~SGD 6–9). "
         "Essential for Grab and Google Maps."),
        ("Old Town Etiquette",
         "Dress modestly when entering temples and the Japanese Bridge — shoulders and knees covered. "
         "Remove shoes before entering shrines. Don't touch sacred objects."),
        ("Bargaining",
         "Expected at night markets and souvenir stalls — start at 50% of asking price. "
         "Polite and friendly. Not appropriate at restaurants with set menus or fixed-price shops."),
        ("Hoi An Lantern Festival",
         "The monthly lantern festival (electric lights off, Old Town lit only by lanterns) occurs "
         "on the 14th day of the lunar calendar. In August 2026, the full moon falls around "
         "28–29 August — check the exact date. Even without the festival, evenings in Hoi An "
         "are magical and lantern-lit year-round."),
        ("Tailoring Tips",
         "Visit tailor shops on Day 4 (Saturday afternoon) to allow 24-hour turnaround. "
         "Yaly Couture is the most professional. Bring photos of what you want. "
         "Confirm fitting appointment and collection time explicitly. Haggle politely."),
        ("Health & Safety",
         "Drink only bottled water. Street food at busy stalls is generally safe. "
         "Carry basic meds: Panadol, Imodium, antihistamine, band-aids. "
         "Travel insurance is strongly recommended. Emergency: 115 (ambulance), 113 (police)."),
        ("Must-Buy Souvenirs",
         "Hand-painted silk lanterns, Vietnamese coffee (Trung Nguyen packs), "
         "Banh Mi spice mix, ao dai fabric, lacquerware, hand-embroidered cushion covers, "
         "ceramic blue-and-white Hoi An pottery."),
    ]

    for title, content in tips:
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(2)
        r = p.add_run(f"  {title}")
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = TEAL
        add_body(doc, f"  {content}", size=10, colour=DARK, space_after=5)

    # ── PHOTO GALLERY PAGE ────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "PHOTO GALLERY — DESTINATION HIGHLIGHTS", level=1,
                colour=NAVY, size=16, space_before=0, space_after=6)
    add_horizontal_rule(doc)

    gallery_keys = ["hoi_an_ancient", "hoi_an_lanterns", "golden_bridge", "marble_mountains"]
    for key in gallery_keys:
        buf, label = imgs.get(key, (None, key))
        if buf:
            add_image_to_doc(doc, buf, label, width_inches=5.2)

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Itinerary prepared for Jason Ong & companion  |  5D4N Danang + Hoi An  |  "
                  "27–31 August 2026")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Generated July 2026  |  1 SGD ≈ 17,000 VND  |  All prices approximate")
    r2.font.name = "Calibri"; r2.font.size = Pt(9); r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── SAVE ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved to:\n  {OUTPUT_PATH}")
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"File size: {size_kb:.1f} KB")
    return OUTPUT_PATH


if __name__ == "__main__":
    build_document()
