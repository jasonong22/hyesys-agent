"""
Generate SCDF Material Submission Cover Sheet — 5 materials
Project: Boon Lay ST Engineering — Product X Deployment
Based on HiLT (Stanley) 8-point cover sheet guideline, 11 Aug 2026

Sources used:
  - PRV: 蒲微泄压阀产品规格书-PB04A4K-50G-7S-V2.0-PUW.pdf
  - Activated Carbon: DS-240 Inspection Report image + AST-SCDF 260326 v3.0.pptx slides 8-16
  - Fire Rated Board: Supalux Calcium Silicate Board.pdf
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd\AST BD"
    r"\HyESys Dept\7. Client Projects\STeng\SCDF"
    r"\HySBatt_SCDF_Material_Cover_Sheet.docx"
)

NAVY     = RGBColor(0x1F, 0x49, 0x7D)
RED      = RGBColor(0xCC, 0x00, 0x00)
GREY     = RGBColor(0x80, 0x80, 0x80)
HDR_FILL = "D9E1F2"
RED_FILL = "FFE0E0"
NOTE_FILL= "FFFDE7"


# ── helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _hdr_row(tbl, labels):
    # Use the first row for headers
    row = tbl.rows[0]
    for cell, label in zip(row.cells, labels):
        cell.text = label
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, HDR_FILL)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def add_field(doc, label, value="[To be filled]", placeholder=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(11)
    rv = p.add_run(value)
    rv.font.size = Pt(11)
    if placeholder:
        rv.font.color.rgb = GREY


def add_red_flag(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"★  RED CLOUD REQUIRED — {text}")
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(10)


def add_table(doc, headers, rows, red_rows=None):
    red_rows = red_rows or set()
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    _hdr_row(tbl, headers)
    for i, row_data in enumerate(rows):
        row = tbl.rows[1 + i]
        for j, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.text = val
            if i in red_rows:
                _set_cell_bg(cell, RED_FILL)
        if i in red_rows and row.cells[-1].paragraphs[0].runs:
            row.cells[-1].paragraphs[0].runs[0].font.color.rgb = RED
            row.cells[-1].paragraphs[0].runs[0].bold = True
    return tbl


def add_review_box(doc):
    add_h2(doc, "Review Comments")
    tbl = doc.add_table(rows=6, cols=1)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        _set_cell_bg(row.cells[0], NOTE_FILL)
    tbl.rows[0].cells[0].paragraphs[0].add_run(
        "[Space for consultant review comments]"
    ).font.color.rgb = GREY


def add_approval_box(doc):
    add_h2(doc, "Approval Sign-Off")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Prepared By (Contractor)", "Reviewed By (Consultant)", "Approved By (Consultant)"])
    for labels in [("Name:", "Name:", "Name:"),
                   ("Signature:", "Signature:", "Signature:"),
                   ("Date:", "Date:", "Date:")]:
        row = tbl.add_row()
        for cell, val in zip(row.cells, labels):
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True


def section_end(doc):
    doc.add_paragraph()
    add_review_box(doc)
    doc.add_paragraph()
    add_approval_box(doc)


# ── Material 1 — H50 ESS ─────────────────────────────────────────────────────

def build_h50_ess(doc):
    add_h1(doc, "Material 1: H50 Energy Storage System (HySBatt ESS)")

    add_h2(doc, "1. Material Description")
    add_field(doc, "Material Name", "HySBatt H50 Energy Storage System (ESS)")
    add_field(doc, "Material Type", "Battery Energy Storage System — Lithium Iron Phosphate (LFP)")
    add_field(doc, "Submission Purpose", "Material approval prior to installation — SCDF NOA")

    add_h2(doc, "2. Proposed Material — Dimensions, Sizes & Configuration")
    add_field(doc, "Model", "HySBatt H50")
    add_field(doc, "Number of Battery Packs", "11 packs")
    add_field(doc, "Pack Dimensions (each)", "1,250 × 500 × 550 mm")
    add_field(doc, "Min. Gap Between Packs", "50 mm")
    add_field(doc, "IP Rating", "IP54")
    add_field(doc, "Installation Footprint", "3.2 m²")
    add_field(doc, "Configuration / Arrangement", "[To be filled — describe physical layout in ESS room]", placeholder=True)

    add_h2(doc, "3. System Description")
    add_field(doc, "Proposed System", "Energy Storage System (ESS) / Active Power Compensation System")
    add_field(doc, "Location of Installation", "[To be filled — specify room / zone / level]", placeholder=True)
    add_field(doc, "Interface with Building / Fire Safety Systems",
              "MSB incomer; flammable gas detection; fire detection & alarm; mechanical ventilation (fire-rated ductwork)")

    add_h2(doc, "4. Technical Specifications & Product Data")
    add_field(doc, "Manufacturer", "Advancer Smart Technology Pte Ltd")
    add_field(doc, "Country of Origin", "[To be filled]", placeholder=True)
    add_field(doc, "Product Catalogue / Datasheet", "HySBatt Datasheet — May 2026, Version 2")

    add_h2(doc, "5. Applicable Product Standards & Codes")
    add_table(doc,
        ["Standard / Code", "Description", "Compliance"],
        [
            ("BS 476-20:1987", "Fire tests — Method for determination of fire resistance", "Compliant — 2HR certified"),
            ("UL 9540A",       "Test Method for Evaluating Thermal Runaway Fire Propagation in Battery ESS (Unit Level)", "Compliant — Unit Level certified"),
            ("[Additional]",   "[To be filled]", "[To be filled]"),
        ]
    )

    add_h2(doc, "6. Singapore Standards & Fire Code Compliance")
    add_field(doc, "Singapore Fire Code Clause(s)", "[To be filled — cite applicable SCDF clause(s) for ESS]", placeholder=True)
    add_field(doc, "SS / CP Reference", "[To be filled]", placeholder=True)
    doc.add_paragraph().add_run("Deviations:").bold = True
    add_table(doc,
        ["Clause / Requirement", "Description of Deviation", "Justification"],
        [("[None identified]", "—", "—")]
    )

    add_h2(doc, "7. Test Reports, Certificates & Third-Party Approvals")
    doc.add_paragraph(
        "For the ESS, the following are compiled into one document. "
        "The two starred certificates must be highlighted by RED CLOUD in the submission package."
    ).runs[0].font.size = Pt(10)
    add_table(doc,
        ["Certificate / Report", "Issuing Body", "Reference / Doc No.", "Highlight"],
        [
            ("BS 476-20 2HR Fire Test Certificate", "[Test Lab]", "[Doc No.]", "YES — RED CLOUD"),
            ("UL 9540A Unit Level Test Report",     "[Test Lab]", "[Doc No.]", "YES — RED CLOUD"),
            ("HySBatt Datasheet",                  "Advancer Smart Technology Pte Ltd", "May 2026 v2", "No"),
        ],
        red_rows={0, 1},
    )
    doc.add_paragraph()
    add_red_flag(doc, "BS 476-20 2HR Fire Test Certificate — box in red cloud in compiled PDF")
    add_red_flag(doc, "UL 9540A Unit Level Test Report — box in red cloud in compiled PDF")

    add_h2(doc, "8. Track Record of Local Projects")
    add_table(doc,
        ["Project Name", "Location", "Year", "Scope", "Reference Contact"],
        [
            ("[To be filled]", "[To be filled]", "[To be filled]", "ESS / HySBatt deployment", "[To be filled]"),
            ("[To be filled]", "[To be filled]", "[To be filled]", "ESS / HySBatt deployment", "[To be filled]"),
        ]
    )

    section_end(doc)


# ── Material 2 — Fire Rated Board (Supalux Calcium Silicate Board) ────────────

def build_fire_rated_board(doc):
    add_h1(doc, "Material 2: Fire Rated Board (Supalux® Calcium Silicate Board)")

    add_h2(doc, "1. Material Description")
    add_field(doc, "Material Name", "Supalux® Calcium Silicate Board")
    add_field(doc, "Material Type", "Fire-Protective Construction Board — Non-combustible Calcium Silicate")
    add_field(doc, "Submission Purpose", "Material approval for 2HR fire-rated ESS room enclosure")

    add_h2(doc, "2. Proposed Material — Dimensions, Sizes & Configuration")
    add_field(doc, "Composition", "Non-combustible calcium silicate matrix — reinforced fibres and fillers, 100% asbestos free")
    add_field(doc, "Standard Board Size", "2440 mm × 1220 mm")
    add_field(doc, "Thickness Selected", "[To be confirmed — available: 9 / 12 / 15 / 20 / 25 mm]", placeholder=True)
    add_field(doc, "Nominal Dry Density", "~975 kg/m³")
    add_field(doc, "Moisture Content", "~6.0%")
    add_field(doc, "Alkalinity", "pH 12")
    add_field(doc, "Configuration / Arrangement", "[To be filled — describe wall / ceiling lining arrangement]", placeholder=True)

    add_h2(doc, "3. System Description")
    add_field(doc, "Proposed System", "Passive fire protection — 2HR fire-rated enclosure / wall lining for ESS room")
    add_field(doc, "Location of Installation", "[To be filled — ESS room walls / ceiling / floor]", placeholder=True)
    add_field(doc, "Interface with Building / Fire Safety Systems",
              "Structural enclosure for HySBatt ESS; integrated with fire-rated ductwork and fire door assembly")

    add_h2(doc, "4. Technical Specifications & Product Data")
    add_field(doc, "Brand", "Supalux®")
    add_field(doc, "Manufacturer", "[To be confirmed — Promat / Etex Group]", placeholder=True)
    add_field(doc, "Country of Origin", "[To be filled]", placeholder=True)
    add_field(doc, "Product Catalogue / Datasheet", "Supalux Calcium Silicate Board Product Datasheet")

    add_h2(doc, "Mechanical Properties")
    add_table(doc,
        ["Property", "Longitudinal", "Transverse", "Unit"],
        [
            ("Modulus of Elasticity (E)", "4.1 kN/mm²", "4.0 kN/mm²", "kN/mm²"),
            ("Flexural Strength (F)",     "10 N/mm²",   "7 N/mm²",    "N/mm²"),
            ("Tensile Strength (T)",      "4.11 N/mm²", "2.15 N/mm²", "N/mm²"),
            ("Compressive Strength (⊥)", "9.3 N/mm²",  "—",          "N/mm²"),
        ]
    )
    doc.add_paragraph()
    add_field(doc, "Thermal Conductivity", "0.242 W/m²K")

    add_h2(doc, "5. Applicable Product Standards & Codes")
    add_table(doc,
        ["Standard / Code", "Description", "Compliance"],
        [
            ("EN 13501-1",        "Reaction to fire classification — A1 (Non-combustible)", "Compliant — A1 Classification"),
            ("BS 476: Part 4",    "Fire tests — Non-combustibility test",                  "Compliant — Non-combustible"),
            ("BS 476: Part 6 & 7","Surface spread of flame — Class O",                    "Compliant — Class O"),
            ("BS 5234",           "Performance of partitions — Heavy duty impact resistance","Compliant"),
            ("ISO 9001:2015",     "Quality management system",                             "Certified"),
            ("ISO 14001:2015",    "Environmental management",                              "Certified"),
            ("ISO 45001:2018",    "Occupational health & safety",                          "Certified"),
        ]
    )

    add_h2(doc, "6. Singapore Standards & Fire Code Compliance")
    add_field(doc, "Singapore Fire Code Clause(s)", "[To be filled — cite applicable SCDF clause(s) for fire-rated enclosure]", placeholder=True)
    add_field(doc, "SS / CP Reference", "[To be filled]", placeholder=True)
    doc.add_paragraph().add_run("Deviations:").bold = True
    add_table(doc,
        ["Clause / Requirement", "Description of Deviation", "Justification"],
        [("[None identified]", "—", "—")]
    )

    add_h2(doc, "7. Test Reports, Certificates & Third-Party Approvals")
    add_table(doc,
        ["Certificate / Report", "Issuing Body", "Reference / Doc No.", "Highlight"],
        [
            ("EN 13501-1 A1 Classification Certificate", "[Notified Body]",  "[Doc No.]", "No"),
            ("BS 476 Part 4 Non-Combustibility Report",  "[Test Lab]",        "[Doc No.]", "No"),
            ("BS 476 Part 6 & 7 Surface Burning Report", "[Test Lab]",        "[Doc No.]", "No"),
            ("ISO 9001:2015 Certificate",                "[Certifying Body]", "[Doc No.]", "No"),
            ("Supalux Calcium Silicate Board Datasheet", "Supalux / Promat",  "—",         "No"),
        ]
    )

    add_h2(doc, "8. Track Record of Local Projects")
    add_table(doc,
        ["Project Name", "Location", "Year", "Scope", "Reference Contact"],
        [
            ("[To be filled]", "[To be filled]", "[To be filled]", "Fire-rated board enclosure", "[To be filled]"),
            ("[To be filled]", "[To be filled]", "[To be filled]", "Fire-rated board enclosure", "[To be filled]"),
        ]
    )

    section_end(doc)


# ── Material 3 — Rockwool ─────────────────────────────────────────────────────

def build_rockwool(doc):
    add_h1(doc, "Material 3: Rockwool (Mineral Wool Insulation)")

    add_h2(doc, "1. Material Description")
    add_field(doc, "Material Name", "Rockwool / Mineral Wool Insulation")
    add_field(doc, "Material Type", "Passive Fire Protection — Mineral Wool (Stone Wool) Insulation")
    add_field(doc, "Submission Purpose", "Material approval for fire-rated ESS room insulation / passive fire protection")

    add_h2(doc, "2. Proposed Material — Dimensions, Sizes & Configuration")
    add_field(doc, "Product Model / Grade", "[To be filled]", placeholder=True)
    add_field(doc, "Thickness", "[To be filled]", placeholder=True)
    add_field(doc, "Standard Board / Slab Size", "[To be filled]", placeholder=True)
    add_field(doc, "Density", "[To be filled]", placeholder=True)
    add_field(doc, "Thermal Conductivity", "[To be filled]", placeholder=True)
    add_field(doc, "Melting Point", "[To be filled — typically >1000°C for stone wool]", placeholder=True)
    add_field(doc, "Configuration / Arrangement", "[To be filled]", placeholder=True)

    add_h2(doc, "3. System Description")
    add_field(doc, "Proposed System", "Passive fire protection — thermal and acoustic insulation within 2HR fire-rated ESS room walls")
    add_field(doc, "Location of Installation", "[To be filled]", placeholder=True)
    add_field(doc, "Interface with Building / Fire Safety Systems",
              "Infill between fire-rated board (Supalux) framework; contributes to overall 2HR fire rating of enclosure")

    add_h2(doc, "4. Technical Specifications & Product Data")
    add_field(doc, "Manufacturer", "[To be filled]", placeholder=True)
    add_field(doc, "Country of Origin", "[To be filled]", placeholder=True)
    add_field(doc, "Product Catalogue / Datasheet", "[To be filled]", placeholder=True)

    add_h2(doc, "5. Applicable Product Standards & Codes")
    add_table(doc,
        ["Standard / Code", "Description", "Compliance"],
        [
            ("[To be filled]", "Non-combustibility / Reaction to fire", "[To be filled]"),
            ("[To be filled]", "Thermal insulation performance",        "[To be filled]"),
            ("[To be filled]", "[To be filled]",                        "[To be filled]"),
        ]
    )

    add_h2(doc, "6. Singapore Standards & Fire Code Compliance")
    add_field(doc, "Singapore Fire Code Clause(s)", "[To be filled]", placeholder=True)
    add_field(doc, "SS / CP Reference", "[To be filled]", placeholder=True)
    doc.add_paragraph().add_run("Deviations:").bold = True
    add_table(doc,
        ["Clause / Requirement", "Description of Deviation", "Justification"],
        [("[To be filled]", "[To be filled]", "[To be filled]")]
    )

    add_h2(doc, "7. Test Reports, Certificates & Third-Party Approvals")
    add_table(doc,
        ["Certificate / Report", "Issuing Body", "Reference / Doc No.", "Highlight"],
        [
            ("[Fire Test / Reaction to Fire Certificate]", "[Test Lab]",  "[Doc No.]", "No"),
            ("[Thermal Performance Test Report]",          "[Test Lab]",  "[Doc No.]", "No"),
            ("[Product Datasheet]",                        "[Manufacturer]", "—",      "No"),
        ]
    )

    add_h2(doc, "8. Track Record of Local Projects")
    add_table(doc,
        ["Project Name", "Location", "Year", "Scope", "Reference Contact"],
        [
            ("[To be filled]", "[To be filled]", "[To be filled]", "Rockwool insulation", "[To be filled]"),
            ("[To be filled]", "[To be filled]", "[To be filled]", "Rockwool insulation", "[To be filled]"),
        ]
    )

    section_end(doc)


# ── Material 4 — Pressure Relief Valve (PUW PB04A4K-50G-7S) ─────────────────

def build_prv(doc):
    add_h1(doc, "Material 4: Pressure Relief Valve (PUW PB04A4K-50G-7S)")

    add_h2(doc, "1. Material Description")
    add_field(doc, "Material Name", "Waterproof Breathable Anti-Explosion Pressure Relief Valve (防水透气防爆阀)")
    add_field(doc, "Material Type", "Mechanical Safety Device — Pressure Relief / Anti-Explosion Valve")
    add_field(doc, "Submission Purpose", "Material approval for ESS battery room pressure management and thermal runaway gas venting")

    add_h2(doc, "2. Proposed Material — Dimensions, Sizes & Configuration")
    add_field(doc, "Model",                    "PB04A4K-50G-7S")
    add_field(doc, "Main Body Material",       "Aluminum alloy — precision CNC machined, hardened, corrosion-resistant surface treatment")
    add_field(doc, "Waterproof Breathable Membrane", "Imported E-PTFE composite (micropore diameter: 0.1–10 μm)")
    add_field(doc, "Sealing Ring",             "Silicone rubber — UL 94-V0 flame retardant, imported raw material")
    add_field(doc, "Thread Specification",     "M4 × 0.7, depth 5.5 mm")
    add_field(doc, "Recommended Torque",       "2–3 N.m")
    add_field(doc, "Min. Exhaust Clearance",   "≥ 20 mm above valve cover")
    add_field(doc, "Quantity per HySBatt Unit","[To be filled — confirm from installation drawing]", placeholder=True)

    add_h2(doc, "3. System Description")
    add_field(doc, "Proposed System", "ESS battery room pressure management; thermal runaway gas exhaust venting")
    add_field(doc, "Location of Installation", "Installed on HySBatt battery enclosure / ESS room panel — [To be confirmed]", placeholder=True)
    add_field(doc, "Interface with Building / Fire Safety Systems",
              "Vents overpressure gas from battery enclosure during thermal runaway; interfaces with flammable gas detection and mechanical ventilation system")
    add_field(doc, "Operating Mode",
              "Normal: breathable valve — equalises pressure via E-PTFE membrane. "
              "Thermal runaway: piston opens at blast pressure, rapidly vents gas; auto-resets when pressure drops (reusable)")

    add_h2(doc, "4. Technical Specifications & Product Data")
    add_field(doc, "Manufacturer",  "DONGGUAN PUW EPTFE MATERIAL Co., LTD (东莞蒲微防水透气膜材料有限公司)")
    add_field(doc, "Country of Origin", "China (Dongguan)")
    add_field(doc, "Product Catalogue / Datasheet", "蒲微泄压阀产品规格书 PB04A4K-50G-7S, Version V2.0")
    add_field(doc, "Patents",
              "ZL201621184072.7; ZL201821571934.0; ZL201920745289.8; ZL202030097744.6; CN202010319131.1")

    add_h2(doc, "Performance Specifications")
    add_table(doc,
        ["Parameter", "Specification"],
        [
            ("Protection Grade",              "IP67"),
            ("Breathing Rate (Normal)",       "≥ 1800 ml/min @ 2.5 kPa"),
            ("Valve Core Opening Pressure",   "4 ± 1 kPa"),
            ("Membrane Oil Resistance",       "Level 7"),
            ("Temperature Tolerance",         "−40 to 150 °C"),
            ("Operating Temperature",         "−40 to 120 °C"),
            ("Flame Retardant Grade (Sealing Ring)", "UL 94-V0"),
            ("Prohibited Substances",         "RoHS compliant; ELV compliant"),
        ]
    )

    add_h2(doc, "5. Applicable Product Standards & Codes")
    add_table(doc,
        ["Standard / Code", "Description", "Compliance"],
        [
            ("IP67 (IEC 60529)",  "Ingress protection — dust tight, water immersion up to 1 m", "Compliant"),
            ("UL 94-V0",         "Flammability standard — sealing ring material",                "Compliant"),
            ("RoHS Directive",   "Restriction of Hazardous Substances",                          "Compliant"),
            ("ELV Directive",    "End of Life Vehicles — prohibited substances",                  "Compliant"),
        ]
    )

    add_h2(doc, "6. Singapore Standards & Fire Code Compliance")
    add_field(doc, "Singapore Fire Code Clause(s)", "[To be filled — cite applicable SCDF clause(s) for pressure relief in ESS room]", placeholder=True)
    add_field(doc, "SS / CP Reference", "[To be filled]", placeholder=True)
    doc.add_paragraph().add_run("Deviations:").bold = True
    add_table(doc,
        ["Clause / Requirement", "Description of Deviation", "Justification"],
        [("[None identified]", "—", "—")]
    )

    add_h2(doc, "7. Test Reports, Certificates & Third-Party Approvals")
    add_table(doc,
        ["Certificate / Report", "Issuing Body", "Reference / Doc No.", "Highlight"],
        [
            ("UL 94-V0 Flame Retardancy Certificate (Sealing Ring)", "[UL / Authorised Lab]", "[Doc No.]", "No"),
            ("IP67 Ingress Protection Test Report",                  "[Test Lab]",             "[Doc No.]", "No"),
            ("RoHS Compliance Certificate",                          "[Certifying Body]",      "[Doc No.]", "No"),
            ("PUW Product Specification Sheet V2.0",                 "PUW EPTFE Material Co.", "PB04A4K-50G-7S V2.0", "No"),
        ]
    )

    add_h2(doc, "8. Track Record of Local Projects")
    add_table(doc,
        ["Project Name", "Location", "Year", "Scope", "Reference Contact"],
        [
            ("[To be filled]", "[To be filled]", "[To be filled]", "PRV for ESS / battery enclosure", "[To be filled]"),
            ("[To be filled]", "[To be filled]", "[To be filled]", "PRV for ESS / battery enclosure", "[To be filled]"),
        ]
    )

    section_end(doc)


# ── Material 5 — Activated Carbon Filter (DS-240) ────────────────────────────

def build_activated_carbon(doc):
    add_h1(doc, "Material 5: Activated Carbon Filter (DS-240)")

    add_h2(doc, "1. Material Description")
    add_field(doc, "Material Name", "DS-240 Granular Activated Carbon (DS-240型活性炭)")
    add_field(doc, "Material Type", "Air / Gas Filtration Media — Granular Activated Carbon")
    add_field(doc, "Submission Purpose",
              "Material approval for ESS battery room ventilation gas filtration — "
              "adsorption of THC / VOC from thermal runaway events")

    add_h2(doc, "2. Proposed Material — Dimensions, Sizes & Configuration")
    doc.add_paragraph().add_run("2a. Material Specifications (DS-240 Inspection Report — Batch 20240426-d, tested 30 Apr 2024)").bold = True
    add_table(doc,
        ["Parameter", "Standard / Requirement", "Tested Result"],
        [
            ("Abrasion Resistance",         "> 93 %",          "95.8 %"),
            ("Particle Size (diameter)",    "4 mm > 90 %",     "96.5 %"),
            ("Float Loss",                  "< 3 %",           "1.5 %"),
            ("pH",                          "> 7",             "7.8"),
            ("Iodine Adsorption Value",     "1150 ± 5% mg/g",  "1165 mg/g"),
            ("Bulk Density",                "400 ± 5% mg/cm³", "390 mg/cm³"),
            ("Ash Content",                 "< 15 %",          "9.2 %"),
            ("Ignition Point",              "> 400 °C",        "450 °C"),
            ("CCl₄ Desorption Rate",        "70 ± 5 %",        "73.3 %"),
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph().add_run("2b. Filter Sizing (per AST-SCDF 260326 v3.0 Calculation Report, Slides 8–16)").bold = True
    add_table(doc,
        ["Parameter", "Value", "Basis"],
        [
            ("Target EBCT",                   "360 seconds (0.1 hr)",  "Safety factor of 6 applied"),
            ("THC per LFP Cell",              "17.6 L = 0.0176 m³",   "UL 9540A Cell Test Report"),
            ("Number of Cells per HySBatt",   "11 cells",              "H50 configuration"),
            ("Total THC per HySBatt",         "193.6 L = 0.1936 m³",  "11 × 17.6 L"),
            ("Cell Vent to Thermal Runaway",  "16 minutes (0.268 hr)", "UL 9540A Cell Test Report"),
            ("Gas Flow Rate (Q)",             "0.72 m³/hr",            "0.1936 m³ ÷ 0.268 hr"),
            ("Volume of AC Required",         "0.072 m³",              "EBCT × Q = 0.1 hr × 0.72 m³/hr"),
            ("AC Unit Dimensions",            "100 × 100 × 400 mm",   "Single unit = 0.004 m³"),
            ("Number of AC Units Required",   "18 units",              "0.072 m³ ÷ 0.004 m³"),
            ("Iodine Value (Selected)",       "1165 mg/g",             "Confirmed from DS-240 Inspection Report"),
            ("Propane : AC Adsorption Ratio", "2.28 g : 1 g",         "Research study — Teknomekanik 2024"),
            ("Propane Density",               "0.5 kg/L",              "Reference value"),
            ("Total THC Mass (propane equiv.)","96.8 kg",              "193.6 L × 0.5 kg/L"),
            ("Theoretical AC Mass Required",  "42.4 kg",              "96.8 kg ÷ 2.28"),
            ("Total AC Mass Deployed",        "54 kg",                 "0.072 m³ × 750 kg/m³ — includes 20% redundancy, > 42.4 kg ✓"),
            ("VOC Removal Efficiency",        "≥ 97 %",               "At EBCT ≥ 60 s — reference Water Research, 2008"),
        ]
    )

    add_h2(doc, "3. System Description")
    add_field(doc, "Proposed System", "ESS battery room ventilation — activated carbon gas filtration system")
    add_field(doc, "Location of Installation",
              "18 AC filter units installed in dedicated housing within ESS room ventilation exhaust path — [confirm location]")
    add_field(doc, "Interface with Building / Fire Safety Systems",
              "Downstream of ESS room mechanical ventilation exhaust; filters THC and VOC gases generated during thermal runaway "
              "before discharge to atmosphere; interlocked with flammable gas detection system")

    add_h2(doc, "4. Technical Specifications & Product Data")
    add_field(doc, "Manufacturer", "溧阳市德胜活性炭厂 (Liyang Desheng Activated Carbon Factory)")
    add_field(doc, "Brand",        "漾竹 (Lizhu)")
    add_field(doc, "Country of Origin", "China (Liyang, Jiangsu)")
    add_field(doc, "Product",      "DS-240型活性炭 (DS-240 Activated Carbon)")
    add_field(doc, "Batch No.",    "20240426-d (tested 30 April 2024)")
    add_field(doc, "Sizing Reference", "AST-SCDF 260326 v3.0 — Activated Carbon Filter Sizing Calculation Report")

    add_h2(doc, "5. Applicable Product Standards & Codes")
    add_table(doc,
        ["Standard / Code", "Description", "Compliance"],
        [
            ("GB/T 7702-2023",  "Coal-based activated carbon test methods (China National Standard)", "Compliant — verified in DS-240 Inspection Report"),
            ("[To be filled]",  "Singapore / international standard for activated carbon filtration",  "[To be filled]"),
        ]
    )

    add_h2(doc, "6. Singapore Standards & Fire Code Compliance")
    add_field(doc, "Singapore Fire Code Clause(s)", "[To be filled — cite applicable SCDF clause(s) for ESS room gas filtration]", placeholder=True)
    add_field(doc, "SS / CP Reference", "[To be filled]", placeholder=True)
    doc.add_paragraph().add_run("Deviations:").bold = True
    add_table(doc,
        ["Clause / Requirement", "Description of Deviation", "Justification"],
        [("[None identified]", "—", "—")]
    )

    add_h2(doc, "7. Test Reports, Certificates & Third-Party Approvals")
    add_table(doc,
        ["Certificate / Report", "Issuing Body", "Reference / Doc No.", "Highlight"],
        [
            ("DS-240 Activated Carbon Inspection Report (GB/T7702-2023)", "溧阳市德胜活性炭厂", "Batch 20240426-d, 30 Apr 2024", "No"),
            ("AST-SCDF Activated Carbon Filter Sizing Calculation Report", "Advancer Smart Technology Pte Ltd", "AST-SCDF 260326 v3.0", "No"),
            ("[Third-party test / certificate if required]", "[To be filled]", "[To be filled]", "No"),
        ]
    )

    add_h2(doc, "8. Track Record of Local Projects")
    add_table(doc,
        ["Project Name", "Location", "Year", "Scope", "Reference Contact"],
        [
            ("[To be filled]", "[To be filled]", "[To be filled]", "Activated carbon filter for ESS gas filtration", "[To be filled]"),
            ("[To be filled]", "[To be filled]", "[To be filled]", "Activated carbon filter for ESS gas filtration", "[To be filled]"),
        ]
    )

    section_end(doc)


# ── main ─────────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cover page
    title = doc.add_heading("MATERIAL SUBMISSION COVER SHEET", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph("Boon Lay ST Engineering — Product X Deployment")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY

    doc.add_paragraph()

    info_tbl = doc.add_table(rows=1, cols=2)
    info_tbl.style = "Table Grid"
    for label, value in [
        ("Project",                  "Boon Lay ST Engineering — Product X Deployment"),
        ("Contractor / Installer",   "[To be filled]"),
        ("Consultant",               "HiLT Pte Ltd (A Jensen Hughes Company)"),
        ("Submission Date",          "11 August 2026"),
        ("Submission Reference No.", "[To be filled]"),
    ]:
        row = info_tbl.add_row()
        row.cells[0].text = label
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    tbl_el = info_tbl._tbl
    tbl_el.remove(tbl_el.tr_lst[0])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Materials Submitted:").bold = True
    for i, mat in enumerate([
        "H50 Energy Storage System (HySBatt ESS)",
        "Fire Rated Board — Supalux® Calcium Silicate Board",
        "Rockwool (Mineral Wool Insulation)",
        "Pressure Relief Valve — PUW PB04A4K-50G-7S",
        "Activated Carbon Filter — DS-240",
    ], 1):
        doc.add_paragraph(f"{i}.  {mat}", style="List Number")

    doc.add_page_break()

    build_h50_ess(doc)
    doc.add_page_break()

    build_fire_rated_board(doc)
    doc.add_page_break()

    build_rockwool(doc)
    doc.add_page_break()

    build_prv(doc)
    doc.add_page_break()

    build_activated_carbon(doc)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
