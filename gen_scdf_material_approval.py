"""
Generate SCDF Material Approval Forms — Boon Lay ST Engineering
Format mirrors the Breeching Inlet example (one form per page)
5 materials + Table of Contents page
Contractor: Advancer Smart Technology Pte Ltd
Date: 11 Aug 2026
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd\AST BD"
    r"\HyESys Dept\7. Client Projects\STeng\SCDF"
    r"\Material_Approval_Boon_Lay_STEng.docx"
)

PROJECT_TITLE    = (
    "Supply, Delivery, Installation, Testing and Commissioning of "
    "Energy Storage System (HySBatt H50) and Associated Fire Safety Works "
    "at Boon Lay, ST Engineering"
)
PROJECT_LOCATION = "Boon Lay, ST Engineering [Address To Be Confirmed]"
CONTRACTOR       = "Advancer Smart Technology Pte Ltd"
DATE             = "11/08/26"
COMPANY_LINE     = "Advancer Smart Technology Pte Ltd"

# ── low-level helpers ─────────────────────────────────────────────────────────

def _tcPr(cell):
    return cell._tc.get_or_add_tcPr()


def set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    _tcPr(cell).append(shd)


def set_cell_border(cell, edges=("top","left","bottom","right"), sz=8, color="000000"):
    tcBorders = OxmlElement("w:tcBorders")
    all_edges = ("top","left","bottom","right","insideH","insideV")
    for e in all_edges:
        el = OxmlElement(f"w:{e}")
        if e in edges:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    _tcPr(cell).append(tcBorders)


def set_row_height(row, height_cm, exact=True):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    twips = int(height_cm * 567)  # 1 cm ≈ 567 twips
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trHeight)


def set_col_widths(table, widths_cm):
    """Set column widths. widths_cm is a list per column."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))
        tblGrid.append(gc)
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"),    str(int(widths_cm[i] * 567)))
                tcW.set(qn("w:type"), "dxa")


def cell_para(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=0):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def cell_para_append(cell, text, bold=False, size=10):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def merge_row_cols(row, start, end):
    """Merge cells from index start to end (inclusive) in a row."""
    cells = row.cells
    merged = cells[start]
    for i in range(start + 1, end + 1):
        merged = merged.merge(cells[i])
    return merged


def all_borders(cell, sz=8):
    set_cell_border(cell, edges=("top","left","bottom","right"), sz=sz)


def no_borders(cell):
    tcBorders = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    _tcPr(cell).append(tcBorders)


# ── page-level helpers ────────────────────────────────────────────────────────

def add_page_header(doc, page_num):
    """
    Title area:  [horizontal rule]  MATERIAL APPROVAL  [page box]
    Matches the Breeching Inlet PDF header style.
    """
    # Outer table: 1 row, 2 cells — left=title area, right=page-number box
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col widths: title=14cm, box=3cm  (total 17cm)
    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    # Size right box
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(3 * 567)))
    tcW.set(qn("w:type"), "dxa")
    right_cell._tc.get_or_add_tcPr().append(tcW)

    # Left cell: top+bottom border, centered "MATERIAL APPROVAL"
    set_cell_border(left_cell, edges=("top","bottom"), sz=12)
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("MATERIAL APPROVAL")
    r.bold = True
    r.font.size = Pt(16)

    # Right cell: all borders, page number top-right
    all_borders(right_cell, sz=12)
    p2 = right_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    r2 = p2.add_run(str(page_num))
    r2.font.size = Pt(9)

    doc.add_paragraph()  # small gap


def add_company_footer(doc):
    p = doc.add_paragraph(COMPANY_LINE)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(8)


def add_approval_form(doc, page_num, item, brand_model, supplier_origin,
                      system, standards, description):
    """
    Build one Material Approval page matching the Breeching Inlet layout.
    Table columns: [label 3.5cm | value 9.5cm | secondary-label 2cm | secondary-value 2cm]
    """
    COL_W = [3.5, 9.5, 2.0, 2.0]   # total = 17 cm
    LABEL_SZ = 9
    VALUE_SZ = 10

    add_page_header(doc, page_num)

    # ── Main form table ───────────────────────────────────────────────────────
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = "Table Grid"

    def add_row():
        return tbl.add_row()

    # ── Row: PROJECT TITLE ────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 1.1, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "PROJECT TITLE",    bold=True, size=LABEL_SZ)
    cell_para(val, PROJECT_TITLE,                  size=VALUE_SZ)

    # ── Row: PROJECT LOCATION ─────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.7, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "PROJECT LOCATION", bold=True, size=LABEL_SZ)
    cell_para(val, PROJECT_LOCATION,               size=VALUE_SZ)

    # ── Row: CONTRACTOR + DATE ────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.7, exact=False)
    c0, c1, c2, c3 = r.cells[0], r.cells[1], r.cells[2], r.cells[3]
    cell_para(c0, "CONTRACTOR",  bold=True, size=LABEL_SZ)
    cell_para(c1, CONTRACTOR,              size=VALUE_SZ)
    cell_para(c2, "DATE",        bold=True, size=LABEL_SZ, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(c3, DATE,                    size=VALUE_SZ,  align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Row: DESCRIPTION OF ITEM header ──────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=True)
    hdr = merge_row_cols(r, 0, 3)
    set_cell_bg(hdr, "D9D9D9")
    cell_para(hdr, "DESCRIPTION OF ITEM", bold=True, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    hdr._tc.get_or_add_tcPr()  # ensure tcPr exists

    # ── Row: ITEM ─────────────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "ITEM",        bold=True, size=LABEL_SZ)
    cell_para(val, item,                    size=VALUE_SZ)

    # ── Row: BRAND / MODEL ────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "BRAND / MODEL",    bold=True, size=LABEL_SZ)
    cell_para(val, brand_model,                   size=VALUE_SZ)

    # ── Row: SUPPLIER / ORIGIN ────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "SUPPLIER / ORIGIN", bold=True, size=LABEL_SZ)
    cell_para(val, supplier_origin,               size=VALUE_SZ)

    # ── Row: SYSTEM ───────────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "SYSTEM",     bold=True, size=LABEL_SZ)
    cell_para(val, system,                 size=VALUE_SZ)

    # ── Row: STANDARDS / APPROVAL ─────────────────────────────────────────────
    r = add_row(); set_row_height(r, 1.2, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    lbl._tc.get_or_add_tcPr()
    cell_para(lbl, "STANDARDS /\nAPPROVAL", bold=True, size=LABEL_SZ)
    # Multi-line standards
    lines = [s.strip() for s in standards.split(";") if s.strip()]
    cell_para(val, lines[0] if lines else "", size=VALUE_SZ)
    for line in lines[1:]:
        cell_para_append(val, line, size=VALUE_SZ)

    # ── Row: DESCRIPTION ──────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 1.4, exact=False)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "DESCRIPTION", bold=True, size=LABEL_SZ)
    cell_para(val, description,              size=VALUE_SZ)

    # ── Row: Approval checkboxes ──────────────────────────────────────────────
    r = add_row(); set_row_height(r, 2.2, exact=True)
    chk = merge_row_cols(r, 0, 3)
    cell_para(chk, "☐  Approval with No Comment",    size=10, space_before=3)
    cell_para_append(chk, "☐  Approval with Comments",     size=10)
    cell_para_append(chk, "☐  Revise and Resubmit",        size=10)
    cell_para_append(chk, "☐  Reject",                     size=10)

    # ── Row: COMMENTS ─────────────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 2.5, exact=True)
    lbl = r.cells[0]; val = merge_row_cols(r, 1, 3)
    cell_para(lbl, "COMMENTS", bold=True, size=LABEL_SZ)
    cell_para(val, "",                    size=VALUE_SZ)     # empty for handwriting

    # ── Row: Sign-off header ──────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 0.65, exact=True)
    left_hdr  = merge_row_cols(r, 0, 1)
    right_hdr = merge_row_cols(r, 2, 3)
    set_cell_bg(left_hdr,  "D9D9D9")
    set_cell_bg(right_hdr, "D9D9D9")
    cell_para(left_hdr,  "SUBMITTED BY CONTRACTOR",   bold=True, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(right_hdr, "APPROVAL FROM CONSULTANT",  bold=True, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Row: Signature space ──────────────────────────────────────────────────
    r = add_row(); set_row_height(r, 2.0, exact=True)
    left_sig  = merge_row_cols(r, 0, 1)
    right_sig = merge_row_cols(r, 2, 3)
    # Signature line — underscores at the bottom
    cell_para(left_sig,  "\n\n________________________________", size=9,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_para_append(left_sig,  "NAME / SIGNATURE / DATE", size=8)
    cell_para(right_sig, "\n\n________________________________", size=9,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_para_append(right_sig, "NAME / SIGNATURE / DATE", size=8)

    # Apply column widths
    set_col_widths(tbl, COL_W)

    doc.add_paragraph()
    add_company_footer(doc)


# ── Table of Contents page ────────────────────────────────────────────────────

def add_toc_page(doc, page_num):
    add_page_header(doc, page_num)

    p = doc.add_paragraph("TABLE OF CONTENTS:")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for cell, text in zip(hdr.cells, ["No.", "Title", "Page", "Remarks"]):
        cell_para(cell, text, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9D9D9")

    rows = [
        ("1",  "Material Approval: H50 Energy Storage System (HySBatt ESS)",           "1",  ""),
        ("2",  "Material Approval: Fire Rated Board — Supalux Calcium Silicate Board",  "2",  ""),
        ("3",  "Material Approval: Rockwool (Mineral Wool Insulation)",                 "3",  ""),
        ("4",  "Material Approval: Pressure Relief Valve (PUW PB04A4K-50G-7S)",        "4",  ""),
        ("5",  "Material Approval: Activated Carbon Filter — DS-240",                   "5",  ""),
        ("6",  "HySBatt Datasheet (May 2026 v2)",                                       "Att",""),
        ("7",  "BS 476-20 2HR Fire Test Certificate",                                   "Att","Red Cloud"),
        ("8",  "UL 9540A Unit Level Test Report",                                       "Att","Red Cloud"),
        ("9",  "Supalux Calcium Silicate Board Datasheet",                              "Att",""),
        ("10", "PUW Pressure Relief Valve Specification Sheet (PB04A4K-50G-7S V2.0)",  "Att",""),
        ("11", "DS-240 Activated Carbon Inspection Report (GB/T 7702-2023, Batch 20240426-d)","Att",""),
        ("12", "Activated Carbon Filter Sizing Calculation Report (AST-SCDF 260326 v3.0)","Att",""),
    ]

    for no, title, page, remarks in rows:
        r = tbl.add_row()
        cell_para(r.cells[0], no,      size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(r.cells[1], title,   size=9)
        cell_para(r.cells[2], page,    size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        remark_run = cell_para(r.cells[3], remarks, size=9)
        if "Red Cloud" in remarks:
            remark_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            remark_run.bold = True

    set_col_widths(tbl, [1.0, 11.0, 1.5, 3.5])
    doc.add_paragraph()
    add_company_footer(doc)


# ── Material definitions ──────────────────────────────────────────────────────

MATERIALS = [
    dict(
        item            = "H50 Energy Storage System (ESS)",
        brand_model     = "HySBatt / H50",
        supplier_origin = "Advancer Smart Technology Pte Ltd / Singapore",
        system          = "Energy Storage System — Active Power Compensation and Battery Storage",
        standards       = (
            "BS 476-20:1987 — Fire Resistance Test (2HR);"
            "UL 9540A — Test Method for Evaluating Thermal Runaway Fire Propagation "
            "in Battery Energy Storage Systems (Unit Level)"
        ),
        description     = (
            "HySBatt H50 active digital power compensator with lithium iron phosphate (LFP) "
            "battery storage. 11 battery packs, IP54, 3.2 m² footprint. "
            "Simultaneously delivers reactive compensation, 3-phase load balancing, and energy storage."
        ),
    ),
    dict(
        item            = "Fire Rated Board — Supalux Calcium Silicate Board",
        brand_model     = "Supalux / [Specify thickness: 9 / 12 / 15 / 20 / 25 mm]",
        supplier_origin = "[Manufacturer — Promat / Etex Group] / [Country of Origin]",
        system          = "Passive Fire Protection — 2-Hour Fire-Rated ESS Room Enclosure",
        standards       = (
            "EN 13501-1 — Reaction to Fire Classification: A1 (Non-Combustible);"
            "BS 476: Part 4 — Non-Combustibility Test;"
            "BS 476: Part 6 and 7 — Class O Surface Burning;"
            "BS 5234 — Heavy Duty Impact Resistance;"
            "ISO 9001:2015 — Quality Management System"
        ),
        description     = (
            "Non-combustible calcium silicate board, 100% asbestos-free. "
            "Nominal dry density approx. 975 kg/m³. Standard size 2440 mm x 1220 mm. "
            "Thickness [to be confirmed]. Thermal conductivity 0.242 W/m²K. "
            "Flexural strength 10 N/mm² (longitudinal). Complies with EN 13501-1 Class A1 and BS 476 Class O."
        ),
    ),
    dict(
        item            = "Rockwool — Mineral Wool Insulation",
        brand_model     = "[Brand / Model — To Be Confirmed]",
        supplier_origin = "[Manufacturer] / [Country of Origin]",
        system          = "Passive Fire Protection — Thermal Insulation Infill within 2HR Fire-Rated ESS Room Walls",
        standards       = (
            "[Applicable fire resistance / reaction to fire standard — To Be Confirmed];"
            "[Applicable thermal performance standard — To Be Confirmed]"
        ),
        description     = (
            "Mineral wool (stone wool) insulation. Grade, thickness, density and dimensions "
            "to be confirmed. Non-combustible. Infill between Supalux fire-rated board framework "
            "to achieve overall 2HR fire-rated enclosure for ESS room."
        ),
    ),
    dict(
        item            = "Pressure Relief Valve — Waterproof Breathable Anti-Explosion Valve",
        brand_model     = "PUW EPTFE / PB04A4K-50G-7S",
        supplier_origin = "Dongguan PUW EPTFE Material Co., Ltd / China",
        system          = "ESS Battery Room Pressure Management — Thermal Runaway Gas Venting",
        standards       = (
            "IP67 (IEC 60529) — Ingress Protection: Dust Tight, Water Immersion to 1 m;"
            "UL 94-V0 — Flammability Standard (Sealing Ring Material);"
            "RoHS Directive — Restriction of Hazardous Substances;"
            "ELV Directive — End of Life Vehicles (Prohibited Substances)"
        ),
        description     = (
            "Waterproof breathable anti-explosion pressure relief valve. "
            "Aluminium alloy body (CNC machined, hardened). "
            "E-PTFE membrane (micropore 0.1 to 10 um). "
            "Silicone sealing ring, UL 94-V0 flame retardant. "
            "Opening pressure: 4 +/- 1 kPa. Operating temperature: -40 to 120 degrees C. "
            "Protection grade: IP67. Auto-resets after pressure relief (reusable)."
        ),
    ),
    dict(
        item            = "Activated Carbon Filter — DS-240 Granular Activated Carbon",
        brand_model     = "Desheng (Lizhu) / DS-240",
        supplier_origin = "Liyang Desheng Activated Carbon Factory / China",
        system          = "ESS Battery Room Ventilation — Gas Filtration (THC / VOC Adsorption from Thermal Runaway Events)",
        standards       = (
            "GB/T 7702-2023 — Activated Carbon Test Methods (China National Standard, "
            "verified in DS-240 Inspection Report, Batch 20240426-d, 30 Apr 2024)"
        ),
        description     = (
            "DS-240 granular activated carbon. Particle size 4 mm (> 90%), "
            "iodine adsorption value 1165 mg/g, bulk density 390 mg/cm³, "
            "ignition point 450 degrees C, abrasion resistance 95.8%, CCl4 desorption rate 73.3%. "
            "18 filter units (100 x 100 x 400 mm each), 54 kg total deployed. "
            "Empty bed contact time (EBCT) 360 seconds. VOC removal efficiency >= 97%."
        ),
    ),
]


# ── build document ────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    for page_num, mat in enumerate(MATERIALS, start=1):
        add_approval_form(doc, page_num, **mat)
        doc.add_page_break()

    # Table of Contents (last page)
    add_toc_page(doc, len(MATERIALS) + 1)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
