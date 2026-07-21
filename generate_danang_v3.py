from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x23, 0x5E, 0x9F)

def h3(doc, text, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = color or RGBColor(0x34, 0x7A, 0x47)

def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run('─' * 62)
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def info(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f'{label}: '); r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(value); r2.font.size = Pt(10.5)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    r = p.add_run(text); r.font.size = Pt(10.5)

def food_header(doc):
    h3(doc, '🍽  Good Food Options (Online Reviewed)', color=RGBColor(0xC0, 0x53, 0x00))

# ─── TITLE ────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Da Nang + Hoi An — 4D1M Itinerary')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('August 2026  •  v3  •  4 Nights / 5 Days')
r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r2.italic = True
doc.add_paragraph()

# ─── TRIP OVERVIEW ────────────────────────────────────────────────────
h1(doc, '✈  Trip Overview')
divider(doc)
info(doc, 'Route', 'Singapore → Da Nang → Ba Na Hills (Night 1) → Hoi An (Night 2) → Da Nang (Nights 3 & 4) → Singapore')
info(doc, 'Duration', '4 nights / 5 days')
info(doc, 'Month', 'August 2026')
doc.add_paragraph()

h2(doc, 'Flights')
bullet(doc, 'Outbound: Scoot TR510  SIN → DAD  |  Depart 08:10  |  Ref: NBT4NW')
bullet(doc, 'Return:   VietJet VJ889  DAD → SIN  |  Depart 14:15  |  Ref: NBT4NW')
doc.add_paragraph()

h2(doc, 'Transport Summary')
legs = [
    ('Airport (DAD) → Ba Na Hills cable car base', 'Grab', '~30 min', 'SGD 20–26'),
    ('Ba Na Hills → Hoi An', 'Grab', '~60–75 min', 'SGD 20–26'),
    ('Hoi An → Da Nang Wellness Resort', 'Grab', '~35–45 min', 'SGD 12–15'),
    ('Da Nang Resort → Da Nang Airport', 'Grab', '~15–20 min', 'SGD 4–6'),
]
for leg, mode, time, cost in legs:
    bullet(doc, f'{leg}  |  {mode}  |  {time}  |  {cost}')
doc.add_paragraph()

# ─── ACCOMMODATION ────────────────────────────────────────────────────
h1(doc, '🏨  Accommodation Overview')
divider(doc)
hotels = [
    ('Night 1', 'Ba Na Hills', 'Mercure Ba Na Hills French Village', 'SGD 120–180/night', 'Within complex; cable car included; summit at 1,487 m'),
    ('Night 2', 'Hoi An', 'Hoi An Trails Resort / La Siesta Hoi An', 'SGD 60–100/night', 'Pool + spa; 3 km from Ancient Town'),
    ('Nights 3 & 4', 'Da Nang', 'TIA Wellness Resort (recommended)', 'SGD 120–180/night', 'My Khe Beach; daily spa, yoga, pilates, pools'),
]
for nights, city, name, price, notes in hotels:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f'{nights} — {city}:  '); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(name); r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x23, 0x5E, 0x9F)
    bullet(doc, f'{notes}  |  Est. {price}', level=1)
doc.add_paragraph()

# ─── DA NANG WELLNESS RESORTS ─────────────────────────────────────────
h1(doc, '🧘  Da Nang Wellness Resort Options (Nights 3 & 4)')
divider(doc)
info(doc, 'Area', 'My Khe Beach strip — 5–10 min Grab to Han River, Dragon Bridge, Sky 36 Bar, Night Market')
doc.add_paragraph()

resorts = [
    ('1.  TIA Wellness Resort  ★ Recommended', [
        'Price: SGD 120–180/night',
        'Daily spa credit included, yoga studio, pilates, 2 pools, gym, direct beach access',
        'Water sports: jet ski, kayak, SUP on-site at My Khe Beach (~100 m)',
        'Best for: all-inclusive wellness at mid-range pricing; no surcharge for daily treatments',
    ]),
    ('2.  Furama Resort Da Nang', [
        'Price: SGD 100–160/night',
        'Full spa (Furama Spa), yoga, 3 pools, watersports centre, beach club',
        'On-site parasailing, banana boat, jet ski, windsurfing',
        'Best for: more variety in water sports; large resort grounds',
    ]),
    ('3.  Pullman Da Nang Beach Resort', [
        'Price: SGD 150–200/night (upper end)',
        'Spa, yoga, pilates studio, infinity pool, gym',
        'Non Nuoc Beach frontage; quieter; 25 min to city centre',
        'Best for: couples seeking quiet luxury',
    ]),
    ('4.  Alacarte Da Nang Beach Hotel  (Budget pick)', [
        'Price: SGD 80–120/night',
        'Rooftop pool, spa, gym; My Khe Beach 200 m; Han River 5 min walk',
        'Best for: lowest cost + most central city access',
    ]),
]
for title_text, bullets in resorts:
    p = doc.add_paragraph(); r = p.add_run(title_text); r.bold = True; r.font.size = Pt(11)
    for b in bullets:
        bullet(doc, b)
    doc.add_paragraph()

# ─── DAY-BY-DAY ───────────────────────────────────────────────────────
h1(doc, '📅  Day-by-Day Itinerary')
divider(doc)

# ══ DAY 1 ══
h2(doc, 'Day 1 (Arrival) — Singapore → Ba Na Hills')
info(doc, 'Night', 'Mercure Ba Na Hills French Village')
info(doc, 'Travel', 'Airport → Ba Na Hills  |  Grab ~30 min  |  SGD 20–26')
doc.add_paragraph()

h3(doc, 'Schedule')
d1 = [
    ('08:10', 'Depart Changi on TR510 (Scoot)'),
    ('~10:30', 'Arrive Da Nang Airport (DAD)  |  Collect bags, get VND cash at ATM, buy local SIM'),
    ('11:00', 'Grab to Ba Na Hills cable car base (~30 min, SGD 20–26)'),
    ('12:00', 'Cable car up (20 min; stunning valley + jungle panorama)'),
    ('12:30', 'Check in: Mercure Ba Na Hills French Village  |  Drop bags; freshen up'),
    ('13:00', 'Lunch at L\'Indochine (inside Ba Na Hills) — Vietnamese-French fusion'),
    ('14:30', 'Golden Bridge — the two giant stone hands holding a golden walkway; iconic sunset photos'),
    ('15:30', 'Explore French Village cobblestone streets, flower gardens, fountains'),
    ('17:00', 'Fantasy Park amusement area — arcade, rides, wax museum (included in cable car ticket)'),
    ('18:30', 'Sunset watch from Ba Na Hills summit (1,487 m) — clouds below, golden light above'),
    ('19:30', 'Dinner at Bistro de la Tour or hotel buffet restaurant'),
    ('21:00', 'Evening stroll through illuminated French Village; rest early for Day 2 sunrise'),
]
for t, a in d1:
    bullet(doc, f'{t}  —  {a}')
doc.add_paragraph()

h3(doc, 'Ba Na Hills Tips')
bullet(doc, 'Cable car ticket (~SGD 25–35 pp) covers all rides + Fantasy Park entry; Mercure guests get priority access')
bullet(doc, 'Summit temperature ~15–20°C cooler than coast — bring a light jacket for evening and early morning')
bullet(doc, 'Set alarm for 04:45 on Day 2 — sunrise at summit is one of the most spectacular experiences in Vietnam')
bullet(doc, 'August mornings are usually clearest; cloud cover builds from midday onwards')
doc.add_paragraph()

food_header(doc)
bullet(doc, 'L\'Indochine (inside Ba Na Hills) — Vietnamese-French buffet; popular for lunch; ★★★★ TripAdvisor')
bullet(doc, 'Bistro de la Tour (inside Ba Na Hills) — European-style dining, good steak and pasta')
bullet(doc, 'Le Jardin Restaurant (Ba Na Hills) — garden terrace; local & Western; relaxed atmosphere')
bullet(doc, 'Note: Dining inside Ba Na Hills is resort-priced; eat a proper lunch as options are limited at dinner')
doc.add_paragraph()

# ══ DAY 2 ══
h2(doc, 'Day 2 — Ba Na Hills Sunrise → Hoi An')
info(doc, 'Night', 'Hoi An Trails Resort / La Siesta Hoi An')
info(doc, 'Travel', 'Ba Na Hills → Hoi An  |  Grab ~60–75 min  |  SGD 20–26')
doc.add_paragraph()

h3(doc, 'Schedule')
d2 = [
    ('04:45', 'Wake up — wrap up warm; head to Ba Na Hills summit viewpoint'),
    ('05:15', 'Sunrise watch at Ba Na Hills (sunrise ~05:30–05:50 in August)  |  Sea of clouds below the ridge, first light over the valley — one of Vietnam\'s most photographed natural moments'),
    ('06:30', 'Golden Bridge at dawn — near-empty; best photography window before day-trippers arrive'),
    ('07:15', 'Breakfast at Mercure restaurant — enjoy the cool crisp mountain morning'),
    ('08:30', 'Last walk through French Village; cable car down'),
    ('09:30', 'Check out; Grab to Hoi An (~60–75 min, SGD 20–26)'),
    ('11:00', 'Arrive Hoi An; check in or store bags at hotel'),
    ('11:30', 'Hoi An Ancient Town on foot — UNESCO heritage streets, lantern-lit alleyways'),
    ('12:30', 'Lunch at Morning Glory or Banh Mi Phuong (see food section below)'),
    ('14:00', 'Japanese Covered Bridge (Chùa Cầu) + Phuc Kien Assembly Hall + Museum of Trading Ceramics'),
    ('16:00', 'Cycling to An Bang Beach (~4 km) — hire bicycle SGD 2–4; quieter and more local than Cua Dai'),
    ('17:30', 'Sunset at An Bang Beach; beachside bars (Soul Kitchen, An Bang Seaside Village)'),
    ('19:00', 'Return to hotel; freshen up'),
    ('19:30', 'Hoi An Night Market on the river — lanterns, silk, local street food'),
    ('20:30', 'Release a lantern on Thu Bon River (~SGD 1–2)'),
    ('21:30', 'Post-market drinks at White Marble Wine Bar or riverside cafe'),
    ('22:30', 'Return to hotel; rest'),
]
for t, a in d2:
    bullet(doc, f'{t}  —  {a}')
doc.add_paragraph()

h3(doc, 'Hoi An Tips')
bullet(doc, 'Ancient Town entry ticket ~SGD 3 pp; covers 5 heritage sites including Japanese Bridge + Assembly Halls')
bullet(doc, 'Early morning is best for Ancient Town — quiet and photogenic before tour groups arrive (~09:00+)')
bullet(doc, 'An Bang Beach has calmer energy than Cua Dai; good spot for swimming in August')
bullet(doc, 'Lantern Festival is every full moon — if timing aligns, the river is covered in floating lanterns')
doc.add_paragraph()

food_header(doc)
bullet(doc, 'Morning Glory Restaurant (Hội An) — iconic; cao lau, white rose dumplings, banh xeo; ★★★★★ TripAdvisor; book ahead')
bullet(doc, 'Banh Mi Phuong — world-famous after Anthony Bourdain\'s visit; long queues but worth it; SGD 1.50 per bánh mì')
bullet(doc, 'White Rose Restaurant (Quán Hoa Trắng) — only place that makes authentic white rose dumplings; family-run; ★★★★')
bullet(doc, 'Cơm Gà A Hai — Hoi An chicken rice, must-try local dish; simple shop, incredible flavour; SGD 2–3')
bullet(doc, 'The Cargo Club — riverside; good for evening drinks + desserts + Western options; ★★★★')
bullet(doc, 'Nu Eatery — modern Vietnamese small plates; excellent online reviews from expats and travellers')
doc.add_paragraph()

# ══ DAY 3 ══
h2(doc, 'Day 3 — Hoi An → Da Nang (Wellness & Beach)')
info(doc, 'Night', 'TIA Wellness Resort (or chosen Da Nang wellness resort)')
info(doc, 'Travel', 'Hoi An → Da Nang resort  |  Grab ~35–45 min  |  SGD 12–15')
doc.add_paragraph()

h3(doc, 'Schedule')
d3 = [
    ('07:30', 'Breakfast at Hoi An hotel'),
    ('08:30', 'Morning walk through Ancient Town lanes before the crowds — best light for photos'),
    ('09:30', 'Tra Que Vegetable Village (~3 km from town) — herb and vegetable farm; short guided walking tour + cooking experience; SGD 5–8 pp'),
    ('11:00', 'Check out; Grab to Da Nang wellness resort (~35–45 min, SGD 12–15)'),
    ('12:00', 'Check in: TIA Wellness Resort; drop bags'),
    ('12:30', 'Lunch at resort beach club or nearby My Khe seafood restaurant'),
    ('14:00', 'My Khe Beach — swim, sunbathe, beach relaxation'),
    ('15:00', 'Water sports at My Khe Beach: jet ski, banana boat, SUP, parasailing  |  SGD 15–40/activity'),
    ('17:00', 'Resort spa or yoga session (included in TIA package; book slot at check-in)'),
    ('18:30', 'Freshen up; Grab to Han River waterfront (~15 min, SGD 5–8)'),
    ('19:00', 'Stroll Han River promenade; view Dragon Bridge lit up at night'),
    ('20:00', 'Dinner on Bach Dang Street — Madame Lan or local Da Nang speciality restaurant'),
    ('21:30', 'Return to resort  |  SGD 5–8'),
    ('22:00', 'Evening dip in resort pool'),
]
for t, a in d3:
    bullet(doc, f'{t}  —  {a}')
doc.add_paragraph()

h3(doc, 'My Khe Beach Water Sports Costs')
water = [
    ('Jet ski', '15 min', 'SGD 15–20'),
    ('Banana boat', 'Per ride', 'SGD 8–12 pp'),
    ('SUP (Stand-Up Paddleboard)', '1 hour', 'SGD 10–15'),
    ('Parasailing', 'Per flight', 'SGD 25–40 pp'),
    ('Kayak rental', '1 hour', 'SGD 8–12'),
]
for sport, dur, cost in water:
    bullet(doc, f'{sport}  |  {dur}  |  {cost}')
doc.add_paragraph()

food_header(doc)
bullet(doc, 'Madame Lan (Bach Dang St) — upscale Vietnamese classics; riverside views; ★★★★ TripAdvisor; popular with tourists & locals')
bullet(doc, 'Waterfront Restaurant (Bach Dang St) — international + Vietnamese; excellent seafood; ★★★★')
bullet(doc, 'Bún Chả Cá Hòa Lộc — famous fish cake noodle soup (Da Nang specialty); busy local spot; SGD 2–3')
bullet(doc, 'Mì Quảng Ông Hải — best mi quang in Da Nang per locals; turmeric noodles with shrimp + pork; SGD 2')
bullet(doc, 'Pizza 4P\'s Da Nang — popular Japanese-Italian fusion; great for group dinners; book ahead; ★★★★★')
bullet(doc, 'Fatfish Restaurant — casual Vietnamese + grilled seafood; consistently good reviews online')
doc.add_paragraph()

# ══ DAY 4 ══
h2(doc, 'Day 4 — Da Nang (City, Sky 36 & Night Market)')
info(doc, 'Night', 'TIA Wellness Resort (Night 4)')
doc.add_paragraph()

h3(doc, 'Schedule')
d4 = [
    ('06:30', 'Early morning yoga class at resort (typically 07:00 start; confirm with resort) or sunrise jog on My Khe Beach'),
    ('07:30', 'Breakfast at resort'),
    ('09:00', 'Grab to Son Tra Peninsula — Linh Ung Pagoda + Lady Buddha (67 m tall; visible from city)  |  ~20 min, SGD 7–10'),
    ('10:00', 'Son Tra Peninsula coastal drive (hire motorbike or private car ~SGD 15–25 for 2 hrs)  |  Monkey Mountain viewpoints, dramatic coastal cliffs, coral reef lookouts'),
    ('12:00', 'Back to city; lunch at local seafood restaurant near My Khe  |  Fresh prawns, squid, clam hotpot  |  SGD 15–25 pp'),
    ('13:30', 'Han Market (Chợ Hàn) — Da Nang\'s central wet market; dried mango, coconut candy, local produce, cheap souvenirs  |  SGD 5–15 for shopping'),
    ('15:00', 'Da Nang Museum of Cham Sculpture — largest collection of Cham artefacts in the world; free or SGD 2 entry'),
    ('16:30', 'Return to resort; afternoon spa treatment or pilates class (book in advance at resort)'),
    ('18:00', 'Grab to Sky 36 Bar, Novotel Da Nang Premier Han River (~15 min, SGD 5–8)'),
    ('18:30', 'Sky 36 Bar (36th floor) — panoramic 360° view over Han River and Da Nang coastline; cocktails SGD 7–15; best at golden hour 18:30–19:30'),
    ('19:30', 'Dinner — Waterfront Restaurant (Bach Dang St) or Madame Lan; riverside setting'),
    ('21:00', 'Bach Dang Night Market — grilled corn, bánh mì, fresh sugarcane juice, street food, local handicrafts'),
    ('22:30', 'Grab back to resort  |  SGD 5–8'),
]
for t, a in d4:
    bullet(doc, f'{t}  —  {a}')
doc.add_paragraph()

h3(doc, 'Day 4 Highlights')
bullet(doc, 'Sky 36 Bar: No cover charge; VND 120,000+ per cocktail (~SGD 7). Golden hour 18:30–19:30 is peak; can get crowded on weekends')
bullet(doc, 'Linh Ung Pagoda: Free entry; 17 km from city centre; combine with Son Tra drive for a half-day trip')
bullet(doc, 'Cham Museum: Often overlooked but very worthwhile — 2,000+ Cham stone artefacts dating to 7th–15th century')
bullet(doc, 'Dragon Bridge fire show: Every Saturday and Sunday 21:00 — free; stand on the bank for best view')
bullet(doc, 'Bach Dang Night Market: Open ~17:00–23:00; food stalls most active after 20:00')
doc.add_paragraph()

food_header(doc)
bullet(doc, 'Sky 36 Bar (Novotel, 36F) — cocktails + finger food; rooftop with best view in Da Nang; ★★★★ TripAdvisor')
bullet(doc, 'Madame Lan — reliable Da Nang upscale dining; Vietnamese herbs + seafood; great for group dinners')
bullet(doc, 'GoGi House Da Nang — Korean BBQ chain; great value; very popular with locals; ★★★★')
bullet(doc, 'Cơm Niêu Saigon — clay pot rice; crispy rice base; authentic Southern Vietnamese; ★★★★')
bullet(doc, 'My Khe Seafood (various stalls near beach) — order by weight; best for fresh squid, crab, prawns grilled at the table')
bullet(doc, 'Bach Dang Night Market stalls — try bột chiên (fried rice cake), bắp nướng (grilled corn), chè (sweet dessert soups)')
doc.add_paragraph()

# ══ DAY 5 ══
h2(doc, 'Day 5 (Departure) — Da Nang → Singapore')
info(doc, 'Flight', 'VietJet VJ889  DAD → SIN  |  Departs 14:15')
info(doc, 'Travel', 'Resort → Airport  |  Grab ~15–20 min  |  SGD 4–6')
doc.add_paragraph()

h3(doc, 'Schedule')
d5 = [
    ('07:30', 'Leisurely breakfast at resort'),
    ('08:30', 'Morning swim or final spa session'),
    ('09:30', 'Check out; store bags at resort reception'),
    ('10:00', 'Final Da Nang walk: Helio Night Market area (daytime) or Han River bridge views'),
    ('10:30', 'My Khe Beach one last dip or beach coffee'),
    ('11:30', 'Collect bags from resort; Grab to Da Nang Airport (~15–20 min, SGD 4–6)'),
    ('12:00', 'Airport arrival; check-in + security (VietJet recommends 2 hrs before)'),
    ('14:15', 'Board VJ889 — depart Da Nang'),
    ('~17:00', 'Arrive Changi Airport, Singapore'),
]
for t, a in d5:
    bullet(doc, f'{t}  —  {a}')
doc.add_paragraph()

food_header(doc)
bullet(doc, 'Da Nang Airport food stalls: Grab a bánh mì or phở before boarding — better and cheaper than in-flight')
bullet(doc, 'Highlands Coffee (airport level) — popular Vietnamese coffee chain; good iced coffee for the flight')
doc.add_paragraph()

# ─── BUDGET ───────────────────────────────────────────────────────────
h1(doc, '💰  Estimated Budget (Per Person, SGD)')
divider(doc)
budget = [
    ('Flights (roundtrip SIN–DAD)', 'SGD 150–300'),
    ('Accommodation (4 nights)', 'SGD 400–560'),
    ('Transport (all Grabs + cable car)', 'SGD 80–110'),
    ('Meals (5 days, mix local + resort)', 'SGD 150–250'),
    ('Water sports + entry tickets', 'SGD 60–100'),
    ('Sky 36 Bar + nightlife', 'SGD 30–50'),
    ('Shopping + souvenirs', 'SGD 50–100'),
    ('Miscellaneous (SIM, tips, extras)', 'SGD 30–50'),
]
for cat, est in budget:
    info(doc, cat, est)
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Total estimate per person: '); r.bold = True; r.font.size = Pt(11)
r2 = p.add_run('SGD 950 – 1,520')
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x34, 0x7A, 0x47)
doc.add_paragraph()

# ─── PRACTICAL TIPS ───────────────────────────────────────────────────
h1(doc, '🎒  Practical Tips')
divider(doc)
tips = [
    'Currency: Vietnamese Dong (VND). SGD 1 ≈ VND 17,500. Withdraw at airport ATM on arrival.',
    'SIM: Buy Viettel / Vietnamobile SIM at airport; unlimited data ~SGD 5 for 7 days.',
    'Grab: Works throughout Da Nang and Hoi An. Download the app and top up before leaving Singapore.',
    'Weather: August is peak rainy season. Mornings clear; rain builds from afternoon. Pack a foldable raincoat.',
    'Ba Na Hills sunrise: Alarm at 04:45. Summit is cold and cloud-free at dawn. Bring the hotel blanket or jacket.',
    'Dragon Bridge fire show: Every Sat & Sun 21:00 on Han River — free, no ticket required.',
    'Tipping: Not mandatory; SGD 1–2 at restaurants, SGD 3–5 for spa. Greatly appreciated.',
    'Health: Drink bottled water only. Most mid-range hotels provide filtered water in rooms.',
    'Power: Vietnam uses Type A/C outlets (same as Singapore for Type C). No adapter needed.',
    'Hoi An Ancient Town: Entry ticket ~SGD 3 pp covers 5 heritage site visits. Buy at the gate.',
]
for tip in tips:
    bullet(doc, tip)
doc.add_paragraph()

# ─── FOOTER ───────────────────────────────────────────────────────────
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Da Nang + Hoi An Itinerary v3  |  August 2026  |  Generated by HyESys Agent')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); r.italic = True

output = r'C:\Users\JasonOng\Desktop\local docs\personal\viet\Danang_HoiAn_4D1M_Aug2026_v3.docx'
doc.save(output)
print(f'Saved: {output}')
