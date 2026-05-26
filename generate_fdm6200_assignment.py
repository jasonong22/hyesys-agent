"""
Generate FDM6200 Take-Home Assignment Word Document
Worked Example: Effect of Ventilation Openings on Flashover in a Compartment Fire
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(31, 73, 125))
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(68, 114, 196))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    run.font.name = "Courier New"
    run.font.size = Pt(10.5)
    return p

def ref_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[Ref: {text}]")
    set_font(run, size=9.5, italic=True, color=(89, 89, 89))
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def divider():
    p = doc.add_paragraph("─" * 90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(7)

def calc_box(lines):
    """Add a shaded calculation box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFF3FB")
    cell._tc.tcPr.append(shd)
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(10)
    doc.add_paragraph()  # spacer after table

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("FDM6200 Fire Dynamics")
run.font.name = "Calibri"
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Take-Home Group Assignment")
run.font.name = "Calibri"
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Investigating the Effect of Ventilation Opening Size on\n"
    "Flashover Occurrence in a Residential Compartment Fire"
)
run.font.name = "Calibri"
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fully Worked Example —")
run.font.name = "Calibri"
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()
divider()
doc.add_paragraph()

for line in [
    "Module:   FDM6200 Fire Dynamics",
    "Date:     May 2026",
    "Institution: Singapore Institute of Technology (SIT)",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (a) — INTRODUCTION & RATIONALE   (10 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(a)  Introduction and Rationale")

body(
    "Flashover is one of the most critical and dangerous phase-transitions in compartment fire "
    "dynamics. It marks the rapid shift from a localised, growing fire to a fully-involved, "
    "room-engulfing conflagration. Understanding the conditions that govern flashover onset is "
    "therefore fundamental to fire safety engineering, occupant evacuation planning, and the "
    "design of fire-resistant compartments."
)

heading2("1.1  Definition of Flashover")
body(
    "Flashover is defined as the near-simultaneous ignition of all exposed combustible surfaces "
    "within a compartment. Workshop 2 identifies four observable criteria that indicate flashover "
    "onset (any one of which is considered sufficient):"
)
bullet("Upper gas layer temperature reaching 500–600 °C;")
bullet("Radiant heat flux at the floor level reaching 20 kW/m²;")
bullet("Ignition of incompletely combusted upper-layer gases emerging from the compartment opening;")
bullet("Flames appearing at the compartment openings (door/window sills).")
ref_note("Workshop 2, Slide 'Flashover Criteria', Section 3.2")

heading2("1.2  Significance and Rationale for Study")
body(
    "Residential fires in Singapore predominantly occur in Housing Development Board (HDB) flat "
    "bedrooms, where a combination of synthetic foam furnishings and limited ventilation creates "
    "high flashover risk. The Singapore Civil Defence Force (SCDF) reports that sleeping occupants "
    "in rooms that have reached flashover have a near-zero survival probability."
)
body(
    "Ventilation opening size (door/window area) is the single most tractable geometric parameter "
    "in a residential compartment — occupants may partially close doors or windows. The ventilation "
    "factor Av√Hv (m^5/2) appears directly in all three standard flashover HRR prediction correlations "
    "(MQH, Babrauskas, Thomas), making it the logical independent variable for experimental study."
)
ref_note("Workshop 2, Flashover HRR correlations; Workshop 1, Vent flow equations")

heading2("1.3  Research Objective")
body(
    "This experiment aims to determine how variation in compartment ventilation opening size "
    "affects the critical Heat Release Rate (HRR) required for flashover, the time-to-flashover "
    "under a standardised fire growth scenario, and the upper-layer gas temperature profile "
    "preceding flashover. Three correlations (MQH, Babrauskas, Thomas) will be used to generate "
    "theoretical predictions, which will then be validated against bench-scale experimental measurements."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (b) — EXPERIMENTAL DESIGN   (20 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(b)  Experimental Design")

heading2("2.1  Scale Selection and Justification")
body(
    "A 1:5 reduced-scale compartment model is selected. Full-scale experiments require large burn "
    "facilities, are resource-intensive, and present significant personnel safety challenges. "
    "Reduced-scale testing is widely accepted in fire research (e.g., Walton & Thomas, SFPE Handbook) "
    "because the governing dimensionless groups — Froude number, ventilation factor ratio, and "
    "thermal inertia ratio — can be matched. At 1:5 scale, dimensional similarity is maintained "
    "through Froude scaling:"
)
equation("Q̇_model = Q̇_full × (L_model / L_full)^(5/2)   [Froude scaling]")
body(
    "This gives a manageable bench-scale HRR (tens of kW) suitable for a standard cone calorimeter "
    "enclosure or purpose-built steel box in a controlled laboratory environment."
)
ref_note("Workshop 2, Section on Post-Flashover Fires; SFPE Handbook Chapter on Scaling")

heading2("2.2  Compartment Geometry (Prototype: HDB Master Bedroom)")
body(
    "The prototype compartment represents a typical Singapore HDB 4-room flat master bedroom."
)

# Geometry table
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Light Shading Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Parameter", "Full Scale (Prototype)", "1:5 Model"]
rows_data = [
    ["Room length (L)", "3.50 m", "0.70 m"],
    ["Room width (W)", "4.00 m", "0.80 m"],
    ["Room height (H)", "2.60 m", "0.52 m"],
    ["Door opening (W × H)", "0.90 m × 2.10 m", "0.18 m × 0.42 m"],
    ["Window opening (W × H) *variable*", "1.20 m × 1.00 m", "0.24 m × 0.20 m"],
]
for j, hdr in enumerate(headers):
    cell = tbl.rows[0].cells[j]
    run = cell.paragraphs[0].add_run(hdr)
    run.bold = True
    run.font.size = Pt(10)
for i, row in enumerate(rows_data, 1):
    for j, val in enumerate(row):
        cell = tbl.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
doc.add_paragraph()

heading2("2.3  Wall Lining Material")
body(
    "Gypsum plasterboard (13 mm) is used to line all internal surfaces of the model. "
    "Gypsum is the most common residential wall lining in Singapore HDB flats, and its thermal "
    "properties are well-characterised. Its low thermal inertia (kρc) compared to concrete "
    "makes flashover more likely — representing a realistic worst-case scenario."
)
calc_box([
    "Gypsum plasterboard properties (from Workshop 2 Table):",
    "  Thermal conductivity  k  = 0.48e-3  kW/(m·K)",
    "  Density               ρ  = 1440     kg/m³",
    "  Specific heat         c  = 0.84     kJ/(kg·K)",
    "  Thickness             δ  = 0.013    m",
    "  Thermal inertia       kρc = 0.48e-3 × 1440 × 0.84 = 0.581  kW²·s/(m⁴·K²)",
])
ref_note("Workshop 2, Table of Lining Material Properties")

heading2("2.4  Fuel Source")
body(
    "A polyurethane (PU) foam mattress slab is chosen as the fuel item. PU foam is ubiquitous "
    "in residential bedrooms and is classified as a Fast t² fire (Workshop 1 and Workshop 2). "
    "This provides a defined, reproducible fire growth rate for experimental comparison."
)
equation("Q̇(t)  =  α_f × t²   where  α_f = 0.0469 kW/s²  (Fast growth)")
ref_note("Workshop 1, Standard t² Fire Growth Rates Table; Workshop 2, Section 3.1")
body(
    "At reduced scale, an equivalent propane burner with mass flow controller can be used to "
    "reproduce the same dimensionless HRR trajectory, eliminating fuel variability between test runs."
)

heading2("2.5  Independent Variable: Ventilation Opening Size")
body(
    "The door is fixed open (Ao = 0.18 × 0.42 = 0.0756 m² at model scale). The window is the "
    "independent variable. Four window opening configurations are tested:"
)
tbl2 = doc.add_table(rows=5, cols=4)
tbl2.style = "Light Shading Accent 1"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ["Condition", "Window opening (model)", "Ao,total (m²)", "Ao√Ho (m^5/2)"]
d2 = [
    ["W0 — Door only",       "Closed",                   "0.0756",  "0.0756 × √0.42 = 0.0490"],
    ["W1 — 25% open",        "0.24 × 0.05 m",            "0.0876",  "≈ 0.0566"],
    ["W2 — 50% open",        "0.24 × 0.10 m",            "0.0996",  "≈ 0.0644"],
    ["W3 — 100% open",       "0.24 × 0.20 m",            "0.1236",  "≈ 0.0800"],
]
for j, hdr in enumerate(h2):
    run = tbl2.rows[0].cells[j].paragraphs[0].add_run(hdr)
    run.bold = True; run.font.size = Pt(10)
for i, row in enumerate(d2, 1):
    for j, val in enumerate(row):
        tbl2.rows[i].cells[j].paragraphs[0].add_run(val).font.size = Pt(10)
doc.add_paragraph()

heading2("2.6  Scaled Compartment Diagram")
body(
    "Figure 1 (overleaf) shows the reduced-scale model layout. Key features shown are: "
    "thermocouple tree positions (TC1–TC3), floor-level heat flux gauge (HFG), mass loss scale "
    "(MLS) under the propane burner, video camera positions (CAM-1 side, CAM-2 front), "
    "and the variable window opening on the far wall."
)

# ASCII diagram
tbl_d = doc.add_table(rows=1, cols=1)
tbl_d.style = "Table Grid"
cell_d = tbl_d.cell(0, 0)
shd_d = OxmlElement("w:shd")
shd_d.set(qn("w:val"), "clear"); shd_d.set(qn("w:color"), "auto"); shd_d.set(qn("w:fill"), "F2F2F2")
cell_d._tc.get_or_add_tcPr().append(shd_d)
diagram_lines = [
    "  PLAN VIEW — 1:5 Reduced-Scale Compartment Model (not to scale)",
    "",
    "  ┌──────────────────────────────────────────────────────────┐",
    "  │                                             [CAM-2]       │",
    "  │  W=0.80 m         TC1    TC2    TC3                       │",
    "  │                    │      │      │    [Window-variable]   │",
    "  │  [Burner/MLS]      │      │      │    ┌─────────┐        │",
    "  │    (corner)        │      │      │    │ 0.24×0.20│       │",
    "  │                                       └─────────┘        │",
    "  │                                                           │",
    "  │  [HFG at floor]                            [CAM-1]       │",
    "  │                                                           │",
    "  ├─────┐                                                     │",
    "  │Door │  0.18×0.42 m (fixed open)                          │",
    "  └─────┴───────────────────────────────────────────────────┘",
    "         L = 0.70 m",
    "",
    "  ELEVATION (side view):",
    "  ┌───────────────────────────────────┐  H = 0.52 m",
    "  │    TC-tree: 5 thermocouples at    │",
    "  │    z = 0.10, 0.21, 0.31, 0.42,   │",
    "  │    0.52 m from floor              │",
    "  └───────────────────────────────────┘",
    "",
    "  Measurement Devices:",
    "    TC1-TC3 : Type-K thermocouple trees (5 heights each)",
    "    HFG     : Schmidt-Boelter total heat flux gauge (floor centre)",
    "    MLS     : Mass loss scale under propane burner",
    "    CAM-1/2 : Video cameras (30 fps) — flame/smoke observation",
    "    O2/CO2  : Oxygen and CO2 analysers at exhaust duct",
]
for i, line in enumerate(diagram_lines):
    p_d = cell_d.paragraphs[0] if i == 0 else cell_d.add_paragraph()
    run_d = p_d.add_run(line)
    run_d.font.name = "Courier New"
    run_d.font.size = Pt(9)
doc.add_paragraph()
p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_cap = p_cap.add_run("Figure 1: Reduced-scale compartment model — plan and elevation views with instrument positions.")
run_cap.font.size = Pt(9.5); run_cap.italic = True

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (c) — METHODOLOGY   (20 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(c)  Methodology")

heading2("3.1  Theoretical Predictions (Pre-Experiment Calculations)")
body(
    "Before any burning tests, theoretical flashover HRRs are calculated for each ventilation "
    "condition using all three standard correlations from Workshop 2. The calculations below "
    "are shown for the full-scale prototype; these are then Froude-scaled to obtain model targets."
)

heading2("3.1.1  Compartment Geometry — Full-Scale Prototype Calculations")
calc_box([
    "Compartment internal surface areas (excluding door opening, including floor/ceiling/walls):",
    "",
    "  Floor         = 3.50 × 4.00          = 14.00 m²",
    "  Ceiling       = 3.50 × 4.00          = 14.00 m²",
    "  Wall 1 (3.5m) = 3.50 × 2.60          =  9.10 m²",
    "  Wall 2 (3.5m) = 3.50 × 2.60          =  9.10 m²",
    "  Wall 3 (4.0m) = 4.00 × 2.60          = 10.40 m²",
    "  Wall 4 (4.0m) with door:",
    "    Full wall    = 4.00 × 2.60          = 10.40 m²",
    "    Door opening = 0.90 × 2.10          =  1.89 m²  (subtracted)",
    "    Net wall 4   = 10.40 - 1.89         =  8.51 m²",
    "",
    "  AT = 14.00 + 14.00 + 9.10 + 9.10 + 10.40 + 8.51",
    "     = 65.11 m²",
    "",
    "  Door opening area:     Ao,door = 0.90 × 2.10 = 1.890 m²",
    "  Door opening height:   Ho,door = 2.10 m",
    "  Ventilation factor:    Ao√Ho   = 1.890 × √2.10 = 2.739 m^(5/2)",
])
ref_note("Workshop 2, MQH Correlation — compartment geometry definitions")

heading2("3.1.2  Effective Heat Transfer Coefficient (hk) for Gypsum Lining")
calc_box([
    "Step 1: Thermal penetration time tp",
    "  tp = (ρc/k) × (δ/2)²",
    "     = (1440 × 0.84 / 0.48e-3) × (0.013/2)²",
    "     = 2,520,000 × 0.00004225",
    "     = 106.5 s",
    "",
    "Step 2: Compare with characteristic exposure time tc",
    "  For pre-flashover growth phase, tc ≈ time-to-flashover ≈ 183 s (calculated below)",
    "  Since tc (183 s) > tp (106.5 s):  use steady-state value",
    "",
    "  hk = k / δ = 0.48e-3 / 0.013 = 0.03692 kW/(m²·K)",
])
ref_note("Workshop 2, Effective Heat Transfer Coefficient — equations (slide set, Section 3.3)")

heading2("3.1.3  MQH Flashover HRR Correlation")
body("The McCaffrey–Quintiere–Harkleroad (MQH) correlation (Workshop 2):")
equation("Q̇_FO,MQH  =  610 × (hk · AT · Ao · √Ho)^(1/2)  [kW]")
calc_box([
    "Q̇_FO,MQH = 610 × (0.03692 × 65.11 × 2.739)^(1/2)",
    "           = 610 × (6.584)^(1/2)",
    "           = 610 × 2.566",
    "           = 1,565 kW",
])
ref_note("Workshop 2, Equation MQH-FO; Exercise 3 worked example")

heading2("3.1.4  Babrauskas Flashover HRR Correlation")
body("Ventilation-based correlation (Workshop 2):")
equation("Q̇_FO,Bab  =  750 × Ao · √Ho  [kW]")
calc_box([
    "Q̇_FO,Bab = 750 × 2.739",
    "           = 2,054 kW",
])
ref_note("Workshop 2, Babrauskas Equation")

heading2("3.1.5  Thomas Flashover HRR Correlation")
body("Thomas correlation incorporating both surface area and ventilation (Workshop 2):")
equation("Q̇_FO,Thomas  =  7.8 · AT  +  378 · Ao · √Ho  [kW]")
calc_box([
    "Q̇_FO,Thomas = 7.8 × 65.11  +  378 × 2.739",
    "             = 507.9  +  1,035.3",
    "             = 1,543 kW",
])
ref_note("Workshop 2, Thomas Equation")

heading2("3.1.6  Design Flashover HRR — Conservative Selection")
body(
    "The MQH correlation is selected as the design flashover HRR (most conservative: "
    "1,565 kW < Thomas 1,543 kW ≈ MQH; Babrauskas is the least conservative at 2,054 kW). "
    "MQH explicitly accounts for the thermal properties of the lining material through hk, "
    "making it the most physically complete of the three methods."
)
calc_box([
    "Design Q̇_FO = 1,565 kW (full scale)  — MQH correlation selected",
    "",
    "Froude-scaled model Q̇_FO:",
    "  Q̇_model = Q̇_full × (L_m / L_full)^(5/2)",
    "           = 1565 × (1/5)^(5/2)",
    "           = 1565 × (0.2)^2.5",
    "           = 1565 × 0.01789",
    "           = 28.0 kW   (target HRR at model scale)",
])

heading2("3.1.7  Time to Flashover (Fast t² Fire)")
body("Using Fast t² growth (Workshop 1, Workshop 2):")
equation("Q̇(t)  =  α_f × t²    →    t_FO  =  √(Q̇_FO / α_f)")
calc_box([
    "α_f    = 0.0469 kW/s²  (Fast growth — PU foam mattress)",
    "Q̇_FO  = 1,565 kW",
    "",
    "t_FO  = √(1565 / 0.0469)",
    "       = √(33,369)",
    "       = 182.7 s  ≈  3 minutes  (full scale)",
    "",
    "Froude-scaled model time:  t_model = t_full × √(L_m/L_full)",
    "  t_model = 182.7 × √(1/5) = 182.7 × 0.4472 = 81.7 s  ≈  82 s",
])
ref_note("Workshop 1, Standard t² Fire Growth Rates; Workshop 2, Fire Growth section")

heading2("3.1.8  Verification via MQH Pre-Flashover Temperature Formula")
body(
    "The MQH temperature correlation (Workshop 2) is used to verify that ΔTg ≈ 500 °C "
    "at the predicted Q̇_FO:"
)
equation("ΔTg  =  6.85 × [Q̇² / (hk · AT · Ao · √Ho)]^(1/3)   [°C]")
calc_box([
    "ΔTg = 6.85 × [1565² / (0.03692 × 65.11 × 2.739)]^(1/3)",
    "     = 6.85 × [2,449,225 / 6.584]^(1/3)",
    "     = 6.85 × [372,003]^(1/3)",
    "     = 6.85 × 71.9",
    "     = 492 °C  ≈  500 °C  ✓  (Flashover threshold confirmed)",
])
ref_note("Workshop 2, MQH Pre-Flashover Temperature Correlation")

heading2("3.2  Experimental Procedure — Step by Step")
numbered("Construct 1:5 reduced-scale steel-framed compartment lined with 13 mm gypsum board.")
numbered("Verify all internal dimensions against design values (±5 mm tolerance).")
numbered("Install thermocouple trees (TC1, TC2, TC3) at three plan positions; five Type-K beads per tree at heights 0.10, 0.21, 0.31, 0.42, 0.52 m.")
numbered("Install Schmidt-Boelter heat flux gauge flush with floor at compartment centre.")
numbered("Install mass loss scale beneath propane burner (corner position, simulating corner-wall mattress placement).")
numbered("Position two video cameras (30 fps, 1080p) at side and front of compartment.")
numbered("Connect propane burner gas line to mass flow controller; calibrate against known HRR at 10, 20, 30 kW.")
numbered("Set gas analyser sampling at exhaust duct (O₂ and CO₂ at 2 Hz).")
numbered("Set window opening to condition W0 (door only, window closed).")
numbered("Start all data loggers at 1 Hz sample rate.")
numbered("Ignite burner; ramp HRR at model-scale α_f trajectory using programmed mass flow controller.")
numbered("Record time of first flashover indicator (see Section (d)); shut off burner immediately at flashover.")
numbered("Allow compartment to cool to ambient; check instrument calibration.")
numbered("Repeat steps 9–13 for conditions W1, W2, W3.")
numbered("Perform three replicate runs for each condition; report mean ± standard deviation.")

heading2("3.3  Controlled Variables")
bullet("Compartment geometry and lining material — fixed")
bullet("Fuel position — corner, fixed")
bullet("HRR ramp rate (t² trajectory) — fixed via mass flow controller")
bullet("Ambient temperature and humidity — controlled to 25 ± 2 °C, 60 ± 10 % RH")
bullet("Door opening — fixed open (Ao,door constant across all runs)")

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (d) — FLASHOVER IDENTIFICATION   (15 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(d)  Flashover Identification Method")

body(
    "Two independent criteria are used to declare flashover onset, consistent with the four "
    "criteria defined in Workshop 2. Using two criteria provides cross-validation and reduces "
    "ambiguity in declaring the flashover time."
)

heading2("Criterion 1 — Upper-Layer Gas Temperature ≥ 500 °C")
body(
    "At least two of the three thermocouple trees must simultaneously record upper-layer "
    "(z ≥ 0.42 m at model scale, equivalent to ≥ 2.1 m full scale) temperature ≥ 500 °C. "
    "This is the most widely cited flashover criterion and is the basis of the MQH "
    "temperature correlation."
)
ref_note("Workshop 2, Flashover Criteria — upper gas layer temperature threshold")

heading2("Criterion 2 — Floor-Level Heat Flux ≥ 20 kW/m²")
body(
    "The Schmidt-Boelter gauge at floor centre must record a total incident heat flux ≥ 20 kW/m². "
    "This represents the minimum radiant flux required to ignite the combustible floor covering "
    "and other low-lying fuel items, producing simultaneous ignition across the compartment."
)
ref_note("Workshop 2, Flashover Criteria — floor radiant heat flux threshold")

heading2("Declaration Logic")
body(
    "Flashover time t_FO is declared as the first time step at which BOTH criteria are met "
    "simultaneously (intersection logic). If only one criterion is met, the event is flagged "
    "as 'near-flashover' and the run is repeated with extended recording duration."
)
body(
    "Additionally, video footage is reviewed after each run to confirm the secondary visual "
    "indicator: flames appearing at the compartment door threshold (criterion iv in Workshop 2). "
    "This serves as a tertiary qualitative check."
)
ref_note("Workshop 2, Section 3.2, all four flashover criteria")

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (e) — DATA ANALYSIS PLAN   (10 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(e)  Data Analysis Plan")

heading2("5.1  Primary Dependent Variables")
bullet("t_FO : time to flashover (s) — from t=0 (ignition) to flashover declaration")
bullet("Q̇_FO_exp : measured HRR at flashover (kW) — from mass flow controller reading at t_FO")
bullet("T_upper,max : maximum upper-layer temperature at t_FO (°C) — mean of TC beads at top two heights")
bullet("q_floor : floor heat flux at t_FO (kW/m²) — from Schmidt-Boelter gauge")

heading2("5.2  Comparison with Theoretical Predictions")
body(
    "For each ventilation condition, measured Q̇_FO,exp is compared against three theoretical "
    "predictions. The percentage error is calculated:"
)
equation("Error (%)  =  (Q̇_FO,exp − Q̇_FO,theory) / Q̇_FO,theory  × 100")
body(
    "This quantifies which correlation (MQH, Babrauskas, or Thomas) most accurately predicts "
    "flashover for a gypsum-lined residential compartment across the range of ventilation conditions."
)
ref_note("Workshop 2, all three flashover HRR correlations")

heading2("5.3  Ventilation Factor Trend Analysis")
body(
    "Q̇_FO is plotted against Ao√Ho for all four conditions. The MQH correlation predicts a "
    "linear relationship between Q̇_FO,MQH and (Ao√Ho)^(1/2), while the Babrauskas correlation "
    "predicts strict linearity with Ao√Ho. The experimental data is curve-fitted to both "
    "forms to determine which relationship best describes the physical observations."
)
equation("MQH form:      Q̇_FO  ∝  (Ao·√Ho)^(1/2)")
equation("Babrauskas form: Q̇_FO  ∝  Ao·√Ho")

heading2("5.4  Pre-Flashover Temperature Profile Analysis")
body(
    "Thermocouple data from all five heights in TC2 (central tree) is plotted against time for "
    "each run. The two-layer zone model (Workshop 2) is applied: the interface height z_int is "
    "estimated as the height at which the temperature gradient is steepest. This is compared "
    "with the Klote–Milke smoke filling model for early-phase smoke descent:"
)
ref_note("Workshop 2, Two-layer zone model and smoke filling correlations")

heading2("5.5  Statistical Reporting")
body(
    "Three replicate runs per condition. Results reported as mean ± one standard deviation. "
    "A one-way ANOVA test is applied to Q̇_FO across the four ventilation conditions to "
    "determine whether the effect of ventilation is statistically significant (α = 0.05)."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (f) — SAFETY AND ETHICAL CONSIDERATIONS   (10 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(f)  Safety and Ethical Considerations")

heading2("6.1  Laboratory Safety Controls")
bullet("All experiments conducted in a certified fire test laboratory with minimum 6 air changes per hour forced ventilation (compliant with NEA guidelines).")
bullet("Bench-scale propane burner operated below SCDF threshold for gas quantity (≤ 0.5 kg propane per run).")
bullet("Automatic gas shut-off valve linked to compartment temperature: burner shuts off if TC2 upper layer exceeds 600 °C or if flashover is declared.")
bullet("All personnel wear full PPE: fire-retardant lab coat, safety glasses, heat-resistant gloves. No synthetic clothing permitted.")
bullet("Minimum 2-person rule: one operator, one safety observer. No solo burn tests.")
bullet("Class B CO₂ fire extinguisher stationed within 2 m of apparatus.")
bullet("Personnel trained in burn first aid and emergency gas shut-off procedures.")
bullet("All instrument cables rated for sustained 400 °C; no PVC-insulated wires inside compartment.")

heading2("6.2  Structural Safety of Reduced-Scale Model")
bullet("Steel frame (1.5 mm thickness, mild steel) with gypsum lining rated to 800 °C short-term.")
bullet("Model is placed on a ceramic fibre blanket base to prevent bench ignition.")
bullet("Minimum 1 m clear zone around model; combustible materials removed from 2 m radius.")

heading2("6.3  Environmental and Ethical Considerations")
bullet("Combustion products (CO, CO₂, soot) are extracted via lab ventilation duct with activated carbon filter before room exhaust.")
bullet("CO monitor at operator station; lab evacuated if CO exceeds 35 ppm (OSHA STEL).")
bullet("Propane consumption minimised by using programmed mass flow controller (no over-burning).")
bullet("Waste gypsum and steel are recycled through approved industrial waste contractor; no landfill disposal.")
bullet("No human subjects involved; no animal subjects.")
bullet("Experiment protocol reviewed and approved by SIT Laboratory Safety Committee before commencement.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (g) — POSSIBLE EXPERIMENTAL ERRORS   (10 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(g)  Possible Experimental Errors")

heading2("7.1  Thermocouple Radiation Error (Systematic)")
body(
    "Type-K bare-bead thermocouples will over-read in a hot radiation environment. The bead absorbs "
    "radiant energy from the hot walls and flame, causing the indicated temperature to exceed the "
    "true gas temperature. This is particularly significant near flashover when wall temperatures "
    "are high. Mitigation: use shielded (aspirated) thermocouples or apply the radiation correction:"
)
equation("T_gas  =  T_indicated  −  (ε_TC · σ · T_wall⁴) / h_conv")
ref_note("Workshop 1, Heat transfer at surfaces — radiation vs. convection")

heading2("7.2  Two-Dimensional Flow Effects at Openings (Systematic)")
body(
    "The Kawagoe ventilation factor Ao√Ho assumes a well-mixed upper layer and 1D flow through "
    "the opening. At reduced scale, the door opening aspect ratio may produce more 3D flow "
    "structures, causing the effective ventilation factor to differ from the theoretical value. "
    "This can cause measured Q̇_FO to deviate from predictions based on Ao√Ho."
)
ref_note("Workshop 2, Ventilation-controlled burning; Workshop 1, vent flow equations")

heading2("7.3  Non-Uniform Ignition at Flashover (Random)")
body(
    "Flashover in a real compartment involves stochastic ignition of multiple fuel surfaces. "
    "In a scaled experiment with a gas burner, the simulated fire does not produce distributed "
    "charring and pyrolysis that would occur in the prototype. The measured t_FO may therefore "
    "be longer than in a real fuel-filled room. Three replicates per condition help quantify "
    "run-to-run variability."
)

heading2("7.4  Scale Effects on Lining Thermal Response (Systematic)")
body(
    "The Froude scaling law governs fluid dynamic similarity, but the heat conduction into the "
    "gypsum lining does not scale the same way. In the model, the thermal penetration depth "
    "relative to the lining thickness may differ from the full-scale ratio, altering the "
    "effective hk. This introduces a systematic bias in the comparison with MQH predictions."
)
ref_note("Workshop 2, hk calculation — tp vs. tc comparison")

heading2("7.5  Heat Flux Gauge Position Sensitivity (Systematic)")
body(
    "The single floor-mounted heat flux gauge measures a point value at compartment centre. "
    "Radiation is not uniform across the floor — it is highest beneath the upper-layer "
    "flame zone and lowest near the cool lower-layer inlet. Placing the gauge centrally may "
    "lead to Criterion 2 being triggered slightly later than if it were positioned at the "
    "point of maximum flux. A secondary gauge near the corner (adjacent to burner) would "
    "improve spatial coverage."
)

heading2("7.6  Instrument Calibration Drift (Random)")
body(
    "Over multiple high-temperature runs, thermocouple accuracy can drift due to wire oxidation "
    "and alloy contamination. All thermocouples should be checked against a NIST-traceable "
    "reference before each test day, and replaced if deviation exceeds ±2 °C at 200 °C reference point."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION (h) — REFERENCES   (5 marks)
# ══════════════════════════════════════════════════════════════════════════════
heading1("(h)  References")

refs = [
    "[1]  McCaffrey, B. J., Quintiere, J. G., & Harkleroad, M. F. (1981). Estimating room "
    "temperatures and the likelihood of flashover using fire test data correlations. "
    "Fire Technology, 17(2), 98–119.  "
    "[Referenced in: MQH temperature and flashover HRR correlations — Workshop 2]",

    "[2]  Babrauskas, V. (1980). Estimating room flashover potential. Fire Technology, 16(2), "
    "94–103.  "
    "[Referenced in: Babrauskas flashover HRR correlation — Workshop 2]",

    "[3]  Thomas, P. H. (1981). Testing products and materials for their contribution to "
    "flashover in rooms. Fire and Materials, 5(3), 103–111.  "
    "[Referenced in: Thomas flashover HRR correlation — Workshop 2]",

    "[4]  SFPE (2016). SFPE Handbook of Fire Protection Engineering (5th ed.). "
    "Springer. Chapters 3 and 25 — Compartment fire modelling, thermal properties of lining.  "
    "[Referenced in: Wall thermal inertia; hk methodology — Workshop 2]",

    "[5]  EN 1991-1-2:2002 (Eurocode 1). Actions on Structures — Part 1-2: General Actions — "
    "Actions on structures exposed to fire. CEN.  "
    "[Referenced in: Fire load density values — Workshop 2]",

    "[6]  Drysdale, D. (2011). An Introduction to Fire Dynamics (3rd ed.). "
    "Wiley. Chapter 9 — Compartment fires; Chapter 3 — Thermochemistry of combustion.  "
    "[Referenced in: t² fire growth rates; flashover definition — Workshop 1 & 2]",

    "[7]  Klote, J. H., & Milke, J. A. (2002). Principles of Smoke Management. "
    "ASHRAE/SFPE.  "
    "[Referenced in: Smoke filling correlations — Workshop 2]",

    "[8]  SCDF (2023). Fire Code 2023 — Singapore Civil Defence Force. "
    "Provisions for residential compartments and fire load limits.  "
    "[Referenced in: Singapore HDB fire safety context — Rationale]",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    set_font(run, size=10.5)

divider()

# ══════════════════════════════════════════════════════════════════════════════
#  CLOSING — EXPLANATION AND JUSTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("Explanation and Justification of Design Choices")

heading2("Why this compartment?")
body(
    "The HDB master bedroom is one of the most fire-fatal compartments in Singapore (SCDF data). "
    "Its standardised geometry across the HDB stock makes it representative of hundreds of "
    "thousands of residential spaces. Selecting it grounds the study in national fire safety "
    "relevance rather than an arbitrary test geometry."
)

heading2("Why PU foam mattress / Fast t² fire?")
body(
    "PU foam is the predominant residential fuel item associated with fatal bedroom fires. "
    "Workshop 1 and Workshop 2 both identify PU foam as a Fast t² growth item (αf = 0.0469 kW/s²). "
    "Using the Fast growth rate represents a credible worst-case scenario without being ultra-fast — "
    "Ultra-fast would describe an accelerant-aided fire, which is not the target scenario here."
)

heading2("Why ventilation opening size as the independent variable?")
body(
    "All three flashover HRR correlations in Workshop 2 include the ventilation factor Ao√Ho. "
    "It is the most tractable variable: occupants can partially close doors/windows. The "
    "experiment directly tests how this everyday decision affects time to flashover — a finding "
    "with immediate practical relevance to occupant behaviour guidance."
)

heading2("Why MQH selected as design correlation?")
body(
    "MQH (1,565 kW) and Thomas (1,543 kW) give near-identical predictions; Babrauskas (2,054 kW) "
    "is 31% higher. MQH is preferred because it explicitly models the lining thermal conductivity "
    "through hk — the most physically complete correlation for a gypsum-lined compartment. "
    "Babrauskas ignores lining properties entirely, and Thomas uses AT as a proxy for lining "
    "absorption without distinguishing between different lining materials."
)

heading2("Why 1:5 reduced scale?")
body(
    "1:5 scale gives a compartment of 0.70 × 0.80 × 0.52 m — manageable within a standard "
    "walk-in fume hood or purpose-built steel box. At this scale the target model HRR is "
    "approximately 28 kW, achievable with a standard laboratory propane burner. Full scale "
    "would require approximately 1,565 kW, necessitating a dedicated large-scale burn facility "
    "not available in a standard university setting."
)

heading2("Why gypsum board lining?")
body(
    "Gypsum board is the actual lining material in the prototype (HDB flats) and is the "
    "reference material in Workshop 2 lining tables. It is non-toxic (no asbestos), low-cost, "
    "and easily replaced between runs. Its well-characterised hk value (0.0369 kW/m²·K) "
    "makes it ideal for validating the MQH correlation."
)

heading2("Summary of Key Flashover Results")
calc_box([
    "FULL-SCALE FLASHOVER SUMMARY (prototype HDB bedroom, door-only ventilation):",
    "",
    "  Ventilation factor:  Ao√Ho  = 2.739  m^(5/2)",
    "  Effective hk         (MQH)  = 0.0369 kW/(m²·K)  [gypsum, tc > tp]",
    "  Total surface area   AT     = 65.11  m²",
    "",
    "  ┌────────────────┬─────────────────┬─────────┐",
    "  │ Correlation    │ Q̇_FO (kW)       │ Notes   │",
    "  ├────────────────┼─────────────────┼─────────┤",
    "  │ MQH            │ 1,565           │ Design  │",
    "  │ Thomas         │ 1,543           │ ≈ MQH   │",
    "  │ Babrauskas     │ 2,054           │ Liberal │",
    "  └────────────────┴─────────────────┴─────────┘",
    "",
    "  Fire growth (Fast t²):  α_f = 0.0469 kW/s²",
    "  Time to flashover:       t_FO = √(1565/0.0469) = 183 s ≈ 3 min",
    "",
    "  Verification (MQH temperature at Q̇_FO):",
    "    ΔTg = 6.85 × [1565² / (0.0369 × 65.11 × 2.739)]^(1/3) = 492 °C ≈ 500 °C  ✓",
    "",
    "  MODEL SCALE (1:5 Froude scaling):",
    "    Q̇_FO,model = 1565 × (0.2)^2.5  = 28.0 kW",
    "    t_FO,model = 183  × √(0.2)      = 81.8 s",
])

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_final.add_run(
    "— End of FDM6200 Take-Home Assignment Worked Example —"
)
set_font(run_f, size=11, italic=True, color=(89, 89, 89))

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\JasonOng\Desktop\FDM6200_Flashover_Assignment.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
