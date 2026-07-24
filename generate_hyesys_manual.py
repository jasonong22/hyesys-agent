"""
HyESys Client User Manual generator
Theme: Navy #0E2841 + Green #1B5E20 + Gold #FFC000  (matches HyESys_DesignSample_v2)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\HyESys Dept\6. Documentation\HyESys_Client_User_Manual.docx"
)

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0E, 0x28, 0x41)   # corporate header
STEEL_BLUE  = RGBColor(0x15, 0x60, 0x82)   # info / sub-bars
DK_GREEN    = RGBColor(0x1B, 0x5E, 0x20)   # HyESys identity
MID_GREEN   = RGBColor(0x2E, 0x7D, 0x32)   # secondary green
GOLD        = RGBColor(0xFF, 0xC0, 0x00)   # brand / KPI values
LT_GREEN    = RGBColor(0xE8, 0xF5, 0xE9)   # solution card bg
PAGE_BG     = RGBColor(0xF5, 0xFA, 0xF5)
PROBLEM_RED = RGBColor(0xB7, 0x1C, 0x1C)
LT_RED      = RGBColor(0xFF, 0xEB, 0xEE)
BODY_TEXT   = RGBColor(0x21, 0x21, 0x21)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

HEX = dict(
    navy     ="0E2841",
    green    ="1B5E20",
    gold     ="FFC000",
    lt_green ="E8F5E9",
    lt_red   ="FFEBEE",
    lt_blue  ="EFF6FF",
    steel    ="156082",
    label    ="D9EAD3",   # softer green for label cells
    alt_row  ="F0F7F0",   # very light green for alternate rows
    white    ="FFFFFF",
)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def set_table_borders(table, colour='C8E6C9', sz=4):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   'single')
                b.set(qn('w:sz'),    str(sz))
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), colour)
                tcBorders.append(b)
            tcPr.append(tcBorders)

def set_row_height(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trPr.append(trHeight)

def add_para_border_bottom(para, colour='1B5E20', sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), colour)
    pBdr.append(bot)
    pPr.append(pBdr)

# ── Style helpers ─────────────────────────────────────────────────────────────
def heading1(doc, text):
    """Navy text + green bottom rule — matches .section-title from HTML."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text.upper())
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    run.font.name  = 'Aptos'
    add_para_border_bottom(p, colour='1B5E20', sz=8)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = DK_GREEN
    run.font.name  = 'Aptos'
    return p

def body(doc, text, size=9.5, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for r in p.runs:
        r.font.size  = Pt(size)
        r.font.color.rgb = BODY_TEXT
        r.font.name  = 'Aptos'
    return p

def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size  = Pt(size)
        r.font.color.rgb = BODY_TEXT
        r.font.name  = 'Aptos'
    return p

def numbered(doc, text, size=9.5):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size  = Pt(size)
        r.font.color.rgb = BODY_TEXT
        r.font.name  = 'Aptos'
    return p

def spacer(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)

# ── Table cell writers ────────────────────────────────────────────────────────
def hdr_cell(cell, text, size=9, centre=True):
    """Navy bg, white bold text — table column header."""
    set_cell_bg(cell, HEX['navy'])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = True
    r.font.color.rgb = WHITE
    r.font.name  = 'Aptos'

def lbl_cell(cell, text, size=9):
    """Light green bg, dark green bold — label column."""
    set_cell_bg(cell, HEX['label'])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = True
    r.font.color.rgb = DK_GREEN
    r.font.name  = 'Aptos'

def data_cell(cell, text, size=9, bold=False, centre=False, colour=None, bg=None):
    if bg:
        set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.color.rgb = colour if colour else BODY_TEXT
    r.font.name  = 'Aptos'

def sub_hdr_row(table, row_idx, text, ncols, colour=HEX['green']):
    """Green sub-header spanning all columns — matches tr.sub-head in HTML."""
    row = table.rows[row_idx]
    merged = row.cells[0].merge(row.cells[ncols - 1])
    set_cell_bg(merged, colour)
    p = merged.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size  = Pt(9)
    r.font.bold  = True
    r.font.color.rgb = WHITE
    r.font.name  = 'Aptos'

def alt_row(table, row_idx):
    """Light green alternate row background."""
    for cell in table.rows[row_idx].cells:
        set_cell_bg(cell, HEX['alt_row'])

# ── Callout boxes ─────────────────────────────────────────────────────────────
def _make_box(doc, label, text, label_colour, text_colour, bg_hex, border_hex):
    """Simulate left-border callout box using a single-row borderless table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # left stripe column (narrow, coloured)
    stripe = tbl.rows[0].cells[0]
    set_cell_bg(stripe, border_hex)
    tbl.columns[0].width = Cm(0.25)
    stripe.paragraphs[0].text = ''

    # content column
    content = tbl.rows[0].cells[1]
    set_cell_bg(content, bg_hex)
    p = content.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label + '  ')
    r1.font.bold  = True
    r1.font.size  = Pt(9)
    r1.font.color.rgb = label_colour
    r1.font.name  = 'Aptos'
    r2 = p.add_run(text)
    r2.font.size  = Pt(9)
    r2.font.color.rgb = text_colour
    r2.font.name  = 'Aptos'

    # remove all borders
    for cell in [stripe, content]:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'none')
            b.set(qn('w:sz'),    '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    spacer(doc, 4)

def note_box(doc, label, text):
    _make_box(doc, label, text,
              label_colour=DK_GREEN, text_colour=BODY_TEXT,
              bg_hex='E8F5E9', border_hex='1B5E20')

def warn_box(doc, label, text):
    _make_box(doc, label, text,
              label_colour=PROBLEM_RED, text_colour=RGBColor(0x50,0x0A,0x0A),
              bg_hex='FFEBEE', border_hex='B71C1C')

def gold_box(doc, label, text):
    _make_box(doc, label, text,
              label_colour=GOLD, text_colour=WHITE,
              bg_hex='0E2841', border_hex='FFC000')

def info_box(doc, label, text):
    _make_box(doc, label, text,
              label_colour=STEEL_BLUE, text_colour=NAVY,
              bg_hex='EFF6FF', border_hex='156082')

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# ── Navy header block ─────────────────────────────────────────────────────────
cover_hdr = doc.add_table(rows=2, cols=1)
cover_hdr.style = 'Table Grid'

# Row 0 — navy bg with brand name
r0 = cover_hdr.rows[0]
set_cell_bg(r0.cells[0], HEX['navy'])
set_row_height(r0, 60)
ph = r0.cells[0].paragraphs[0]
ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph.paragraph_format.space_before = Pt(16)
ph.paragraph_format.space_after  = Pt(4)
rh1 = ph.add_run('HyE')
rh1.font.size = Pt(32); rh1.font.bold = True
rh1.font.color.rgb = WHITE; rh1.font.name = 'Aptos'
rh2 = ph.add_run('Sys')
rh2.font.size = Pt(32); rh2.font.bold = True
rh2.font.color.rgb = GOLD; rh2.font.name = 'Aptos'
ph2 = r0.cells[0].add_paragraph()
ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph2.paragraph_format.space_after = Pt(10)
rh3 = ph2.add_run('Active Digital Power Compensator')
rh3.font.size = Pt(13); rh3.font.color.rgb = RGBColor(0xCC,0xDD,0xEE)
rh3.font.name = 'Aptos'

# Row 1 — green stripe
r1 = cover_hdr.rows[1]
set_cell_bg(r1.cells[0], HEX['green'])
set_row_height(r1, 6)
r1.cells[0].paragraphs[0].text = ''

# remove all table borders
for row in cover_hdr.rows:
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
            tcBorders.append(b)
        tcPr.append(tcBorders)

spacer(doc, 8)

# ── Sub-title ─────────────────────────────────────────────────────────────────
ps = doc.add_paragraph()
ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
ps.paragraph_format.space_after = Pt(20)
rs = ps.add_run('CLIENT USER MANUAL')
rs.font.size = Pt(16); rs.font.bold = True
rs.font.color.rgb = NAVY; rs.font.name = 'Aptos'

# ── Cover info table ──────────────────────────────────────────────────────────
cov = doc.add_table(rows=7, cols=2)
cov.style = 'Table Grid'
cov_data = [
    ('Project Name',       '[Project Name]'),
    ('Installation Site',  '[Site Address]'),
    ('HyESys Model',       '[H30 / H50 / H125]'),
    ('Serial Number',      '[Serial Number]'),
    ('Document Revision',  'Rev 1.0'),
    ('Issue Date',         '[DD MMM YYYY]'),
    ('Prepared By',        'Advancer Smart Technology Pte Ltd'),
]
for i, (lbl, val) in enumerate(cov_data):
    if i % 2 == 1:
        alt_row(cov, i)
    lbl_cell(cov.rows[i].cells[0], lbl)
    data_cell(cov.rows[i].cells[1], val)
set_table_borders(cov, colour='C8E6C9')

spacer(doc, 14)
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run('Advancer Smart Technology Pte Ltd  |  www.advancer.sg')
rf.font.size = Pt(8.5); rf.font.color.rgb = STEEL_BLUE; rf.font.name = 'Aptos'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — DOCUMENT CONTROL
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '1.  Document Control')
body(doc, 'This document is issued by Advancer Smart Technology Pte Ltd (AST) for the client\'s reference and ongoing operation of the HyESys unit at the above site.')

dc = doc.add_table(rows=9, cols=2)
dc.style = 'Table Grid'
dc_data = [
    ('Project Name',       '[Project Name]'),
    ('Installation Site',  '[Full site address]'),
    ('HyESys Model',       '[H30 / H50 / H125]'),
    ('Unit Serial Number', '[SN-XXXXXX]'),
    ('MSB / Panel',        '[Panel designation]'),
    ('Commissioning Date', '[DD MMM YYYY]'),
    ('Document Revision',  'Rev 1.0'),
    ('Issue Date',         '[DD MMM YYYY]'),
    ('Approved By',        '[Name, Designation]'),
]
for i, (lbl, val) in enumerate(dc_data):
    if i % 2 == 1:
        alt_row(dc, i)
    lbl_cell(dc.rows[i].cells[0], lbl)
    data_cell(dc.rows[i].cells[1], val)
set_table_borders(dc, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, 'Revision History')
rh = doc.add_table(rows=2, cols=4)
rh.style = 'Table Grid'
for j, txt in enumerate(['Rev', 'Date', 'Description', 'Approved By']):
    hdr_cell(rh.rows[0].cells[j], txt)
for j, val in enumerate(['1.0', '[DD MMM YYYY]', 'Initial issue', '[Name]']):
    data_cell(rh.rows[1].cells[j], val, centre=(j == 0))
set_table_borders(rh, colour='C8E6C9')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '2.  System Overview')

heading2(doc, '2.1  Purpose')
body(doc, 'HyESys is an active digital power compensator installed at the Main Switchboard (MSB) incomer. It delivers three simultaneous functions, sharing the unit\'s rated kVA output:')
bullet(doc, 'Reactive power compensation — injects leading or lagging kVAr to correct power factor, reducing reactive current drawn from the utility.')
bullet(doc, 'Three-phase load balancing — redistributes current across phases to eliminate neutral I²R losses.')
bullet(doc, 'Energy storage / solar load shaving — stores excess solar generation and releases it to reduce peak demand charges.')
note_box(doc, 'Target PF:', 'HyESys targets a power factor of 0.98. Unity (1.0) is not the design target. The SP penalty threshold is PF < 0.85.')

heading2(doc, '2.2  Operating Principle')
body(doc, 'HyESys measures the MSB incomer current waveform at high speed, calculates the reactive, imbalance and harmonic components, and injects an exact compensating current through a dedicated circuit breaker. The result is a reduced, balanced current at the utility meter — lowering I²R cable losses and improving power factor.')

heading2(doc, '2.3  System Capacity')
cap = doc.add_table(rows=6, cols=5)
cap.style = 'Table Grid'
for j, txt in enumerate(['Model', 'Rated Output (kVA)', 'Max Current (A)', 'Usable Energy (kWh)', 'Weight (kg)']):
    hdr_cell(cap.rows[0].cells[j], txt)
cap_data = [
    ('H30',  '30',  '43.5', '69.3',  '1,400'),
    ('H50',  '50',  '72.5', '108.9', '2,200'),
    ('H60',  '60',  '87',   '138.6', '2,800'),
    ('H100', '100', '145',  '217.8', '4,400'),
    ('H125', '125', '181',  '217.8', '4,400'),
]
for i, row in enumerate(cap_data):
    if i % 2 == 1:
        alt_row(cap, i + 1)
    for j, val in enumerate(row):
        bold = (j == 0)
        clr  = GOLD if j == 0 else BODY_TEXT
        data_cell(cap.rows[i+1].cells[j], val, centre=True, bold=bold, colour=clr)
set_table_borders(cap, colour='C8E6C9')

spacer(doc, 4)
heading2(doc, '2.4  Functional Block Diagram')
info_box(doc, 'Diagram:', 'Insert functional block diagram — MSB busbars → dedicated MCCB → HyESys unit → DC battery string → monitoring gateway → data logger → router → cloud/SCADA.')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — COMPONENTS HANDED OVER
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '3.  Components Handed Over to Client')
body(doc, 'Verify all items below during the handover inspection. Any discrepancies must be recorded in Section 10.')

comp = doc.add_table(rows=9, cols=6)
comp.style = 'Table Grid'
for j, txt in enumerate(['#', 'Description', 'Qty', 'Serial / Ref No.', 'Location', 'Warranty']):
    hdr_cell(comp.rows[0].cells[j], txt)
comp_data = [
    ('1', 'HyESys Main Unit',           '1',     '[SN-XXXXXX]', '[MSB Room / Panel]',  '2 years'),
    ('2', 'HySBatt Battery Pack(s)',     '[qty]', '[SN-XXXXXX]', '[MSB Room]',           '2 years'),
    ('3', 'Monitoring Gateway',          '1',     '[SN-XXXXXX]', '[Panel / DIN Rail]',   '1 year'),
    ('4', 'Data Logger (Splitter)',      '1',     '[SN-XXXXXX]', '[Panel / DIN Rail]',   '1 year'),
    ('5', 'Router / SIM Module',         '1',     '[SN-XXXXXX]', '[Panel / Cabinet]',    '1 year'),
    ('6', 'Communication & Power Cables','1 set', 'N/A',          'As installed',         'N/A'),
    ('7', 'Keys & Access Items',         '[qty]', '[Key No.]',   '[Refer to KT]',        'N/A'),
    ('8', 'Manufacturer Docs & Certs',   '1 set', 'N/A',          'With this manual',     'N/A'),
]
for i, row in enumerate(comp_data):
    if i % 2 == 1:
        alt_row(comp, i + 1)
    for j, val in enumerate(row):
        data_cell(comp.rows[i+1].cells[j], val, centre=(j in [0, 2]))
set_table_borders(comp, colour='C8E6C9')
note_box(doc, 'Note:', 'Key handover details to be completed by KT. Manufacturer datasheets, FAT certificates and warranty cards are in Appendix F.')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — KEY WIRING REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '4.  Key Wiring Requirements')
body(doc, 'The following is a generic overview only. All final connections must follow the approved Site Single Line Diagram (SLD) and installation drawings in Appendix A.')

heading2(doc, '4.1  Generic Wiring Overview')
wt = doc.add_table(rows=7, cols=2)
wt.style = 'Table Grid'
hdr_cell(wt.rows[0].cells[0], 'Connection')
hdr_cell(wt.rows[0].cells[1], 'Requirement', centre=False)
wt_data = [
    ('Three-Phase Power (L1, L2, L3)',
     'Connected to HyESys AC terminals via the dedicated circuit breaker. Phasing must be correct and terminations torqued to specification.'),
    ('Neutral (N)',
     'Connected to the HyESys neutral terminal. Do not share this neutral with other downstream loads.'),
    ('Protective Earth (PE)',
     'Dedicated earth conductor to MSB earth busbar. Must not be daisy-chained with other equipment.'),
    ('Dedicated Circuit Breaker',
     'Correctly rated MCCB or MCB as specified in the approved SLD. Final rating to be confirmed by Jeff.'),
    ('Isolation Switch',
     'Lockable isolation upstream of the dedicated breaker for safe maintenance isolation. Final arrangement to be confirmed by Jeff.'),
    ('Communication Wiring',
     'CT sensor cables, RS-485/Modbus and Ethernet/SIM are routed per installation drawings. Do not re-route without AST approval.'),
]
for i, (lbl, val) in enumerate(wt_data):
    if i % 2 == 1:
        alt_row(wt, i + 1)
    lbl_cell(wt.rows[i+1].cells[0], lbl)
    data_cell(wt.rows[i+1].cells[1], val)
set_table_borders(wt, colour='C8E6C9')

spacer(doc, 4)
heading2(doc, '4.2  Cable and Breaker Ratings')
info_box(doc, 'To be confirmed:', 'Final cable and breaker ratings are confirmed by AST/Jeff per the approved SLD, based on cable length, ambient temperature and local authority requirements.')
warn_box(doc, 'WARNING:', 'No wiring modifications are permitted without AST written approval. Incorrect wiring voids warranty and creates electrical hazards.')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — POWERING UP SEQUENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '5.  Powering Up Sequence')
note_box(doc, 'Routine operation:', 'HyESys starts and stops automatically. No daily client action is required. This section covers full power-up from a de-energised state.')

heading2(doc, '5.1  Power-Up from Cold')
for step in [
    'Confirm all maintenance work is complete and personnel are clear.',
    'Verify the dedicated circuit breaker is in the OFF position.',
    'Close the upstream isolation switch.',
    'Switch on the auxiliary supply (if a separate auxiliary breaker is provided).',
    'Confirm the HyESys display panel illuminates and shows the startup screen.',
    'Close the dedicated circuit breaker.',
    'Confirm the unit enters STANDBY or AUTO mode on the display.',
    'Verify PF and current readings appear normal on the monitoring interface.',
]:
    numbered(doc, step)

spacer(doc, 6)
heading2(doc, '5.2  Safe Isolation for Maintenance (De-Energise)')
for step in [
    'Select MANUAL HOLD or disable automatic operation on the HyESys control panel.',
    'Wait for the unit to ramp down — confirm output current reaches zero on the display.',
    'Open (turn OFF) the dedicated circuit breaker.',
    'Switch off the auxiliary supply.',
    'Apply lockout/tagout on the isolation switch per your site safety procedure.',
    'Wait at least 5 minutes for DC bus capacitors to discharge.',
    'Verify zero voltage with an approved voltage tester before opening the enclosure.',
]:
    numbered(doc, step)

warn_box(doc, 'WARNING:', 'Never open the HyESys enclosure without first confirming zero voltage. Lethal DC bus voltages may remain present after AC isolation.')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — NORMAL OPERATION
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '6.  Normal Operation')

heading2(doc, '6.1  Operating Modes')
modes = doc.add_table(rows=5, cols=3)
modes.style = 'Table Grid'
for j, txt in enumerate(['Mode', 'Display Indicator', 'Description']):
    hdr_cell(modes.rows[0].cells[j], txt)
modes_data = [
    ('AUTO',    'Green — AUTO',    'Unit operating normally. All active compensation functions running based on site demand.'),
    ('STANDBY', 'Amber — STANDBY','Unit energised and ready but not actively injecting. Normal during low-load or overnight periods.'),
    ('FAULT',   'Red — FAULT',    'Fault detected. Unit has safely tripped. Check alarm log and contact AST before taking any action.'),
    ('MANUAL',  'Blue — MANUAL',  'Under manual control for commissioning or maintenance only. Not a normal client operating mode.'),
]
indicator_colours = [HEX['green'], 'FFC000', 'B71C1C', '156082']
for i, (mode, ind, desc) in enumerate(modes_data):
    if i % 2 == 1:
        alt_row(modes, i + 1)
    data_cell(modes.rows[i+1].cells[0], mode, bold=True, centre=True,
              colour=DK_GREEN if i == 0 else (GOLD if i == 1 else (PROBLEM_RED if i == 2 else STEEL_BLUE)))
    data_cell(modes.rows[i+1].cells[1], ind, bold=True)
    data_cell(modes.rows[i+1].cells[2], desc)
set_table_borders(modes, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '6.2  Expected Readings (Normal Conditions)')
rdgs = doc.add_table(rows=5, cols=3)
rdgs.style = 'Table Grid'
for j, txt in enumerate(['Parameter', 'Normal Range', 'Note']):
    hdr_cell(rdgs.rows[0].cells[j], txt)
rdgs_data = [
    ('Power Factor (PF)',     '>= 0.95 lagging',     'Target 0.98. Values above 0.85 avoid SP penalty tariff.'),
    ('Phase Current Balance', '< 5% imbalance',      'Significant reduction from pre-installation baseline expected.'),
    ('Output Current',        '<= Rated unit current','Must not exceed nameplate rating continuously.'),
    ('DC Bus Voltage',        'Within model range',   'H30: 210–850 V  |  H50: 350–850 V  |  H125: 680–900 V'),
]
for i, row in enumerate(rdgs_data):
    if i % 2 == 1:
        alt_row(rdgs, i + 1)
    lbl_cell(rdgs.rows[i+1].cells[0], row[0])
    data_cell(rdgs.rows[i+1].cells[1], row[1], bold=True, colour=DK_GREEN, centre=True)
    data_cell(rdgs.rows[i+1].cells[2], row[2])
set_table_borders(rdgs, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '6.3  Confirming Correct Operation')
bullet(doc, 'The monitoring dashboard (browser/app) should show real-time PF trending >= 0.95.')
bullet(doc, 'Monthly energy savings reports are generated automatically by the AST cloud platform.')
bullet(doc, 'If the display is blank or shows an unexpected state, note the error code and contact AST before taking any action.')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ALARMS AND BASIC TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '7.  Alarms and Basic Troubleshooting')
note_box(doc, 'Scope:', 'The client is not expected to repair the unit. The actions below are limited to safe observation and basic resets only.')

alm = doc.add_table(rows=9, cols=4)
alm.style = 'Table Grid'
for j, txt in enumerate(['Alarm / Symptom', 'Possible Cause', 'Client Action', 'Contact AST?']):
    hdr_cell(alm.rows[0].cells[j], txt)
alm_data = [
    ('FAULT — Overcurrent',         'Load surge or short circuit downstream',      'Check for obvious faults. Do not reset without AST guidance.',                                       'YES'),
    ('FAULT — Overvoltage',         'Grid voltage spike or improper connection',    'Record fault code and time. Do not reset. Contact AST.',                                             'YES'),
    ('FAULT — Overtemperature',     'Blocked ventilation or high ambient temp',     'Check ventilation openings are clear. Allow unit to cool. Attempt one reset.',                      'Yes if fault persists'),
    ('FAULT — Comms Loss',          'Router/SIM issue or cable fault',              'Check router power and SIM signal. Restart router.',                                                 'Yes if unresolved'),
    ('STANDBY > 24 hrs',            'Low load or unit held in manual mode',         'Check display — confirm AUTO mode is selected.',                                                     'Yes if AUTO shown but inactive'),
    ('Display blank',               'Auxiliary supply off or internal fault',       'Check auxiliary circuit breaker. Reset once if tripped.',                                            'Yes if blank persists'),
    ('PF not improving (dashboard)','Unit in STANDBY or PF already at target',     'Verify AUTO mode. Review dashboard trend — PF may already be 0.95–0.98.',                          'Yes if PF consistently < 0.90'),
    ('Unusual noise or smell',      'Mechanical or electrical fault',               'Immediately open the dedicated circuit breaker. Do not re-energise. Contact AST.',                  'YES — immediately'),
]
for i, row in enumerate(alm_data):
    if i % 2 == 1:
        alt_row(alm, i + 1)
    lbl_cell(alm.rows[i+1].cells[0], row[0])
    data_cell(alm.rows[i+1].cells[1], row[1])
    data_cell(alm.rows[i+1].cells[2], row[2])
    is_yes = row[3].startswith('YES')
    data_cell(alm.rows[i+1].cells[3], row[3],
              bold=is_yes,
              colour=PROBLEM_RED if is_yes else DK_GREEN,
              centre=True)
set_table_borders(alm, colour='C8E6C9')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — DOS AND DON'TS
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "8.  Dos and Don'ts")

ddt = doc.add_table(rows=2, cols=2)
ddt.style = 'Table Grid'
# Green header / Red header
hdr_cell(ddt.rows[0].cells[0], 'DO')
set_cell_bg(ddt.rows[0].cells[0], HEX['green'])
hdr_cell(ddt.rows[0].cells[1], 'DO NOT')
set_cell_bg(ddt.rows[0].cells[1], 'B71C1C')

dos = [
    'Keep ventilation openings unobstructed at all times.',
    'Conduct a visual inspection monthly — check for dust, water ingress and loose labels.',
    'Verify the monitoring dashboard is accessible and data is updating weekly.',
    'Report any unusual alarms, smells or sounds to AST immediately.',
    'Keep the area around the unit clear of combustible materials.',
    'Ensure the dedicated circuit breaker is accessible and labelled.',
    'Follow the isolation procedure before any maintenance near the unit.',
    'Retain all documentation in a safe and accessible location.',
]
donts = [
    'Open the enclosure unless authorised and trained to do so.',
    'Modify or re-route any wiring without AST written approval.',
    'Stack items against or on top of the unit.',
    'Attempt to reset the unit more than once without AST guidance.',
    'Allow unauthorised personnel to operate the controls.',
    'Apply water or liquid cleaners near the unit.',
    'Connect additional loads to the dedicated circuit breaker.',
    'Disable alarms or override fault conditions independently.',
]
do_cell   = ddt.rows[1].cells[0]
dont_cell = ddt.rows[1].cells[1]
set_cell_bg(do_cell,   HEX['lt_green'])
set_cell_bg(dont_cell, HEX['lt_red'])
for item in dos:
    p = do_cell.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9); r.font.color.rgb = DK_GREEN; r.font.name = 'Aptos'
for item in donts:
    p = dont_cell.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9); r.font.color.rgb = PROBLEM_RED; r.font.name = 'Aptos'
set_table_borders(ddt, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '8.1  Inspection and Maintenance Schedule')
maint = doc.add_table(rows=5, cols=3)
maint.style = 'Table Grid'
for j, txt in enumerate(['Frequency', 'Activity', 'Responsible']):
    hdr_cell(maint.rows[0].cells[j], txt)
maint_data = [
    ('Monthly',     'Visual inspection — ventilation, cleanliness, labels, indicator status.',              'Client Facility Manager'),
    ('Quarterly',   'Visual check of cable terminations for security and damage (external only).',          'Client Facility Manager'),
    ('Bi-annually', 'Full inspection and functional test.',                                                 'AST Service Engineer'),
    ('Annually',    'Preventive maintenance, firmware check, battery health assessment.',                   'AST Service Engineer'),
]
for i, row in enumerate(maint_data):
    if i % 2 == 1:
        alt_row(maint, i + 1)
    lbl_cell(maint.rows[i+1].cells[0], row[0])
    data_cell(maint.rows[i+1].cells[1], row[1])
    data_cell(maint.rows[i+1].cells[2], row[2])
set_table_borders(maint, colour='C8E6C9')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — SAFETY AND EMERGENCY RESPONSE
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '9.  Safety and Emergency Response')

heading2(doc, '9.1  Electrical Hazards')
bullet(doc, 'HyESys operates at 400 V AC (three-phase) and high DC bus voltages up to 900 V DC. Contact with live terminals is potentially fatal.')
bullet(doc, 'DC bus capacitors retain charge for up to 5 minutes after AC isolation. Always allow discharge time and verify with a live voltage indicator.')
bullet(doc, 'Do not assume the unit is safe because the display is off — internal voltages may still be present.')

spacer(doc, 4)
heading2(doc, '9.2  Required PPE')
bullet(doc, 'Insulated safety gloves — minimum Cat III / 1000 V rated.')
bullet(doc, 'Arc flash-rated face shield and flame-resistant clothing (appropriate to MSB arc flash level).')
bullet(doc, 'Non-conductive safety footwear.')
bullet(doc, 'Voltage detection / live line indicator.')

spacer(doc, 4)
heading2(doc, '9.3  Emergency Shutdown')
for step in [
    'On detecting fire, smoke, abnormal sound or burning smell — immediately open the dedicated circuit breaker.',
    'If safe to do so, open the upstream isolation switch and apply lockout/tagout.',
    'Evacuate the area and follow your site emergency response plan.',
    'Contact AST and your site emergency coordinator immediately.',
    'Do not re-energise the unit until AST has inspected and issued written clearance.',
]:
    numbered(doc, step)
warn_box(doc, 'FIRE:', 'Use CO₂ or dry powder extinguisher only on electrical equipment. Never use water.')

spacer(doc, 4)
heading2(doc, '9.4  Emergency Contacts')
ec = doc.add_table(rows=4, cols=3)
ec.style = 'Table Grid'
for j, txt in enumerate(['Contact', 'Organisation / Name', 'Number']):
    hdr_cell(ec.rows[0].cells[j], txt)
ec_data = [
    ('AST 24-hr Support',      'Advancer Smart Technology Pte Ltd', '[+65 XXXX XXXX]'),
    ('Site Facility Manager',  '[Name]',                             '[Contact Number]'),
    ('SP PowerGrid Emergency', 'SP PowerGrid Ltd',                   '1800 778 8888'),
]
for i, row in enumerate(ec_data):
    if i % 2 == 1:
        alt_row(ec, i + 1)
    lbl_cell(ec.rows[i+1].cells[0], row[0])
    data_cell(ec.rows[i+1].cells[1], row[1])
    data_cell(ec.rows[i+1].cells[2], row[2], bold=True)
set_table_borders(ec, colour='C8E6C9')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — CLIENT HANDOVER AND TRAINING
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '10.  Client Handover and Training')

heading2(doc, '10.1  Training Attendance')
tr = doc.add_table(rows=5, cols=4)
tr.style = 'Table Grid'
for j, txt in enumerate(['Full Name', 'Designation', 'Organisation', 'Signature']):
    hdr_cell(tr.rows[0].cells[j], txt)
for i in range(1, 5):
    if i % 2 == 1:
        alt_row(tr, i)
    for j in range(4):
        data_cell(tr.rows[i].cells[j], '')
set_table_borders(tr, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '10.2  Handover Checklist')
hc = doc.add_table(rows=10, cols=3)
hc.style = 'Table Grid'
for j, txt in enumerate(['Item', 'Status', 'Remarks']):
    hdr_cell(hc.rows[0].cells[j], txt)
hc_data = [
    'Site commissioning completed and signed off.',
    'Equipment verified against handover list (Section 3).',
    'Keys and access items handed over (KT).',
    'Client trained on normal operation and monitoring dashboard.',
    'Client trained on isolation and emergency shutdown procedure.',
    'Monitoring dashboard login credentials provided.',
    'All manufacturer documents and certificates handed over.',
    'Approved SLD and wiring drawings handed over.',
    'Outstanding items noted with agreed resolution dates.',
]
for i, item in enumerate(hc_data):
    if i % 2 == 1:
        alt_row(hc, i + 1)
    data_cell(hc.rows[i+1].cells[0], item)
    data_cell(hc.rows[i+1].cells[1], 'Complete  /  Pending', centre=True)
    data_cell(hc.rows[i+1].cells[2], '')
set_table_borders(hc, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '10.3  Outstanding Items')
oi = doc.add_table(rows=4, cols=4)
oi.style = 'Table Grid'
for j, txt in enumerate(['No.', 'Description', 'Responsible Party', 'Target Date']):
    hdr_cell(oi.rows[0].cells[j], txt)
for i in range(1, 4):
    if i % 2 == 1:
        alt_row(oi, i)
    data_cell(oi.rows[i].cells[0], str(i), centre=True)
    for j in range(1, 4):
        data_cell(oi.rows[i].cells[j], '')
set_table_borders(oi, colour='C8E6C9')

spacer(doc, 6)
heading2(doc, '10.4  Client Acknowledgement')
body(doc, 'We confirm that the HyESys system has been commissioned and handed over, and that we have received sufficient training to operate the system in accordance with this manual.')
spacer(doc, 4)
ack = doc.add_table(rows=4, cols=3)
ack.style = 'Table Grid'
hdr_cell(ack.rows[0].cells[0], '')
hdr_cell(ack.rows[0].cells[1], 'AST Representative')
hdr_cell(ack.rows[0].cells[2], 'Client Representative')
ack_rows = [('Name & Designation', '', ''), ('Signature', '', ''), ('Date', '', '')]
for i, row in enumerate(ack_rows):
    if i % 2 == 1:
        alt_row(ack, i + 1)
    lbl_cell(ack.rows[i+1].cells[0], row[0])
    data_cell(ack.rows[i+1].cells[1], row[1])
    data_cell(ack.rows[i+1].cells[2], row[2])
set_table_borders(ack, colour='C8E6C9')
spacer(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — SITE SPECIFIC APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '11.  Site-Specific Appendices')
body(doc, 'The following site-specific documents are attached. All must be updated if any changes are made to the installation.')

app = doc.add_table(rows=9, cols=3)
app.style = 'Table Grid'
for j, txt in enumerate(['Appendix', 'Document', 'Status']):
    hdr_cell(app.rows[0].cells[j], txt)
app_data = [
    ('A', 'Approved Single Line Diagram (SLD)'),
    ('B', 'Wiring and Termination Drawings'),
    ('C', 'Equipment Layout Drawing'),
    ('D', 'Cable Route Drawing'),
    ('E', 'Commissioning Records and Sign-Off Sheet'),
    ('F', 'Test Reports (FAT, SAT, insulation, CT accuracy)'),
    ('G', 'Equipment Datasheets (HyESys, HySBatt, Gateway, Logger)'),
    ('H', 'Warranty Cards and Certificates'),
]
for i, (app_id, desc) in enumerate(app_data):
    if i % 2 == 1:
        alt_row(app, i + 1)
    data_cell(app.rows[i+1].cells[0], app_id, bold=True, colour=GOLD, centre=True)
    data_cell(app.rows[i+1].cells[1], desc)
    data_cell(app.rows[i+1].cells[2], 'Attached  /  Pending', centre=True)
set_table_borders(app, colour='C8E6C9')

spacer(doc, 8)
gold_box(doc, 'Reminder:', 'This manual should be reviewed and updated after any change to the installation, firmware upgrade or change in site conditions. Contact AST for the latest version.')

# ── Footer strip ──────────────────────────────────────────────────────────────
spacer(doc, 10)
ft = doc.add_table(rows=2, cols=1)
ft.style = 'Table Grid'
set_cell_bg(ft.rows[0].cells[0], HEX['green'])
set_row_height(ft.rows[0], 4)
ft.rows[0].cells[0].paragraphs[0].text = ''
set_cell_bg(ft.rows[1].cells[0], HEX['navy'])
fp = ft.rows[1].cells[0].paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(5)
fp.paragraph_format.space_after  = Pt(5)
r_hy = fp.add_run('HyESys')
r_hy.font.size = Pt(9); r_hy.font.bold = True; r_hy.font.color.rgb = GOLD; r_hy.font.name = 'Aptos'
r_rest = fp.add_run('  ·  Advancer Smart Technology Pte Ltd  ·  www.advancer.sg  ·  Confidential')
r_rest.font.size = Pt(8.5); r_rest.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE); r_rest.font.name = 'Aptos'
for cell in [ft.rows[0].cells[0], ft.rows[1].cells[0]]:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
