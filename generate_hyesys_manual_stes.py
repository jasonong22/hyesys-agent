"""
HyESys Client User Manual — ST Engineering Land Systems Ltd. (STES)
Model: H50  |  Site: 249 Jln Boon Lay, Singapore 619523
Theme: Navy #0E2841 + Green #1B5E20 + Gold #FFC000
No costings, no pricing anywhere.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\HyESys Dept\6. Documentation"
    r"\HyESys_Client_User_Manual_STES_H50.docx"
)

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0E, 0x28, 0x41)
STEEL_BLUE  = RGBColor(0x15, 0x60, 0x82)
DK_GREEN    = RGBColor(0x1B, 0x5E, 0x20)
GOLD        = RGBColor(0xFF, 0xC0, 0x00)
PROBLEM_RED = RGBColor(0xB7, 0x1C, 0x1C)
BODY_TEXT   = RGBColor(0x21, 0x21, 0x21)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

HEX = dict(
    navy    ="0E2841",
    green   ="1B5E20",
    gold    ="FFC000",
    lt_green="E8F5E9",
    lt_red  ="FFEBEE",
    lt_blue ="EFF6FF",
    steel   ="156082",
    label   ="D9EAD3",
    alt_row ="F0F7F0",
    white   ="FFFFFF",
    red     ="B71C1C",
)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex6); tcPr.append(shd)

def set_table_borders(table, colour='C8E6C9', sz=4):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
            tcB = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
                b.set(qn('w:space'),'0'); b.set(qn('w:color'),colour)
                tcB.append(b)
            tcPr.append(tcB)

def no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        tcB.append(b)
    tcPr.append(tcB)

def set_row_height(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_pt * 20)))
    trPr.append(trH)

def green_rule(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8')
    b.set(qn('w:space'),'4'); b.set(qn('w:color'),'1B5E20')
    pBdr.append(b); pPr.append(pBdr)

# ── Style helpers ─────────────────────────────────────────────────────────────
def H1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text.upper())
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = NAVY; r.font.name = 'Aptos'
    green_rule(p); return p

def H2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = DK_GREEN; r.font.name = 'Aptos'; return p

def body(doc, text, size=9.5, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for r in p.runs:
        r.font.size = Pt(size); r.font.color.rgb = BODY_TEXT; r.font.name = 'Aptos'
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9.5); r.font.color.rgb = BODY_TEXT; r.font.name = 'Aptos'

def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9.5); r.font.color.rgb = BODY_TEXT; r.font.name = 'Aptos'

def sp(doc, pt=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pt)

# ── Table cell helpers ────────────────────────────────────────────────────────
def hdr(cell, text, size=9, centre=True):
    set_cell_bg(cell, HEX['navy'])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True
    r.font.color.rgb = WHITE; r.font.name = 'Aptos'

def lbl(cell, text, size=9):
    set_cell_bg(cell, HEX['label'])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True
    r.font.color.rgb = DK_GREEN; r.font.name = 'Aptos'

def dat(cell, text, size=9, bold=False, centre=False, colour=None, bg=None):
    if bg: set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = colour if colour else BODY_TEXT; r.font.name = 'Aptos'

def alt(table, i):
    for c in table.rows[i].cells: set_cell_bg(c, HEX['alt_row'])

# ── Callout boxes ─────────────────────────────────────────────────────────────
def _box(doc, label, text, lc, tc, bg, border):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    stripe = tbl.rows[0].cells[0]
    set_cell_bg(stripe, border)
    tbl.columns[0].width = Cm(0.25)
    stripe.paragraphs[0].text = ''
    content = tbl.rows[0].cells[1]
    set_cell_bg(content, bg)
    p = content.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + '  ')
    r1.font.bold = True; r1.font.size = Pt(9)
    r1.font.color.rgb = lc; r1.font.name = 'Aptos'
    r2 = p.add_run(text)
    r2.font.size = Pt(9); r2.font.color.rgb = tc; r2.font.name = 'Aptos'
    for cell in [stripe, content]: no_border(cell)
    sp(doc, 4)

def note(doc, lbl_text, text):
    _box(doc, lbl_text, text, DK_GREEN, BODY_TEXT, 'E8F5E9', '1B5E20')

def warn(doc, lbl_text, text):
    _box(doc, lbl_text, text, PROBLEM_RED, RGBColor(0x50,0x0A,0x0A), 'FFEBEE', 'B71C1C')

def info(doc, lbl_text, text):
    _box(doc, lbl_text, text, STEEL_BLUE, NAVY, 'EFF6FF', '156082')

def gold_box(doc, lbl_text, text):
    _box(doc, lbl_text, text, GOLD, WHITE, '0E2841', 'FFC000')

def make_header_block(doc, title_white, title_gold, subtitle, tag_line):
    """Navy cover header with green stripe."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0]
    set_cell_bg(r0.cells[0], HEX['navy'])
    set_row_height(r0, 70)
    # Title
    ph = r0.cells[0].paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.space_before = Pt(18); ph.paragraph_format.space_after = Pt(4)
    rw = ph.add_run(title_white); rw.font.size = Pt(32); rw.font.bold = True
    rw.font.color.rgb = WHITE; rw.font.name = 'Aptos'
    rg = ph.add_run(title_gold); rg.font.size = Pt(32); rg.font.bold = True
    rg.font.color.rgb = GOLD; rg.font.name = 'Aptos'
    # Subtitle
    ph2 = r0.cells[0].add_paragraph()
    ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph2.paragraph_format.space_after = Pt(6)
    rs = ph2.add_run(subtitle)
    rs.font.size = Pt(12); rs.font.color.rgb = RGBColor(0xCC,0xDD,0xEE)
    rs.font.name = 'Aptos'
    # Tag line
    ph3 = r0.cells[0].add_paragraph()
    ph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph3.paragraph_format.space_after = Pt(14)
    rt = ph3.add_run(tag_line)
    rt.font.size = Pt(10); rt.font.color.rgb = RGBColor(0x88,0xAA,0xCC)
    rt.font.name = 'Aptos'
    # Green stripe row
    r1 = tbl.rows[1]
    set_cell_bg(r1.cells[0], HEX['green'])
    set_row_height(r1, 5)
    r1.cells[0].paragraphs[0].text = ''
    for row in tbl.rows:
        for cell in row.cells: no_border(cell)

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
make_header_block(doc,
    title_white='HyE', title_gold='Sys',
    subtitle='Active Digital Power Compensator  ·  Model H50',
    tag_line='249 Jln Boon Lay, Singapore 619523  ·  ST Engineering Land Systems Ltd.')

sp(doc, 8)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(20)
rs = p_sub.add_run('CLIENT USER MANUAL')
rs.font.size = Pt(16); rs.font.bold = True
rs.font.color.rgb = NAVY; rs.font.name = 'Aptos'

cov = doc.add_table(rows=7, cols=2)
cov.style = 'Table Grid'
cov_data = [
    ('Project Name',       'HyESys H50 — Sub-MSB Power Quality Improvement'),
    ('Client',             'ST Engineering Land Systems Ltd. (STES)'),
    ('Installation Site',  '249 Jln Boon Lay, Singapore 619523'),
    ('HyESys Model',       'H50  (50 kVAr / 110 kWh)'),
    ('MSB / Panel',        'Sub-MSB Meter  —  IGS Contestable'),
    ('Document Revision',  'Rev 1.0'),
    ('Prepared By',        'Advancer Smart Technology Pte Ltd'),
]
for i, (l, v) in enumerate(cov_data):
    if i % 2 == 1: alt(cov, i)
    lbl(cov.rows[i].cells[0], l)
    dat(cov.rows[i].cells[1], v)
set_table_borders(cov, 'C8E6C9')

sp(doc, 14)
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run('Advancer Smart Technology Pte Ltd  |  www.advancer.sg')
rf.font.size = Pt(8.5); rf.font.color.rgb = STEEL_BLUE; rf.font.name = 'Aptos'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — DOCUMENT CONTROL
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '1.  Document Control')
body(doc, 'This document is issued by Advancer Smart Technology Pte Ltd (AST) for ST Engineering Land Systems Ltd. (STES) as the operating reference for the HyESys H50 unit installed at the Sub-MSB at 249 Jln Boon Lay, Singapore 619523.')

dc = doc.add_table(rows=9, cols=2)
dc.style = 'Table Grid'
dc_data = [
    ('Client',              'ST Engineering Land Systems Ltd. (STES)'),
    ('Installation Site',   '249 Jln Boon Lay, Singapore 619523'),
    ('MSB / Panel',         'Sub-MSB Meter  —  IGS Contestable'),
    ('HyESys Model',        'H50  (50 kVAr / 110 kWh)'),
    ('Unit Serial Number',  '[To be confirmed at installation]'),
    ('Commissioning Date',  '[To be confirmed]'),
    ('Document Revision',   'Rev 1.0'),
    ('Issue Date',          '24 Jul 2026'),
    ('Approved By',         '[Name, Designation]'),
]
for i, (l, v) in enumerate(dc_data):
    if i % 2 == 1: alt(dc, i)
    lbl(dc.rows[i].cells[0], l)
    dat(dc.rows[i].cells[1], v)
set_table_borders(dc, 'C8E6C9')

sp(doc, 6)
H2(doc, 'Revision History')
rh_t = doc.add_table(rows=2, cols=4)
rh_t.style = 'Table Grid'
for j, t in enumerate(['Rev', 'Date', 'Description', 'Approved By']):
    hdr(rh_t.rows[0].cells[j], t)
for j, v in enumerate(['1.0', '24 Jul 2026', 'Initial issue', '[Name]']):
    dat(rh_t.rows[1].cells[j], v, centre=(j == 0))
set_table_borders(rh_t, 'C8E6C9')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '2.  System Overview')

H2(doc, '2.1  Purpose')
body(doc, 'The HyESys H50 is installed at the Sub-MSB meter at 249 Jln Boon Lay to improve power quality and reduce distribution losses. It delivers three simultaneous functions:')
bullet(doc, 'Reactive power compensation — corrects power factor by injecting up to 50 kVAr, reducing reactive current drawn from the SP Group network.')
bullet(doc, 'Three-phase load balancing — redistributes current across phases to eliminate neutral I²R losses arising from the site\'s measured phase imbalance.')
bullet(doc, 'Solar export capture — stores weekend solar spillover (which would otherwise be exported at negligible credit under the IGS Contestable tariff) and releases it during on-site consumption periods.')
note(doc, 'Site context:', 'Rooftop solar is confirmed on site. Weekend export is concentrated on Fri/Sat/Sun when daytime load is lower. The H50 battery (110 kWh) is sized to capture this spillover within a single day\'s export cycle.')

H2(doc, '2.2  Operating Principle')
body(doc, 'HyESys measures the Sub-MSB incomer current waveform continuously. It calculates the reactive, imbalance and harmonic components and injects an exact compensating current through a dedicated circuit breaker connected to the Sub-MSB busbars. The result is a reduced, balanced current at the Sub-MSB incomer — lowering I²R distribution losses and improving power factor toward the 0.98 target.')

H2(doc, '2.3  H50 System Specification')
cap = doc.add_table(rows=2, cols=5)
cap.style = 'Table Grid'
for j, t in enumerate(['Model', 'Rated Output (kVAr)', 'Battery Capacity (kWh)', 'Max Current (A)', 'Supply Voltage']):
    hdr(cap.rows[0].cells[j], t)
for j, v in enumerate(['H50', '50', '110', '72.5', '350 – 850 V DC bus\n(400 V AC, 3-phase)']):
    dat(cap.rows[1].cells[j], v, centre=True,
        bold=(j == 0), colour=GOLD if j == 0 else BODY_TEXT)
set_table_borders(cap, 'C8E6C9')

sp(doc, 4)
H2(doc, '2.4  Functional Block Diagram')
info(doc, 'Diagram:', 'Insert functional block diagram — Sub-MSB busbars → dedicated MCCB → HyESys H50 unit → DC battery string (110 kWh) → monitoring gateway → data logger → router → AST cloud platform.')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — COMPONENTS HANDED OVER
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '3.  Components Handed Over to Client')
body(doc, 'Verify all items below during the handover inspection. Discrepancies must be recorded in Section 10.')

comp = doc.add_table(rows=9, cols=6)
comp.style = 'Table Grid'
for j, t in enumerate(['#', 'Description', 'Qty', 'Serial / Ref No.', 'Location', 'Warranty']):
    hdr(comp.rows[0].cells[j], t)
comp_data = [
    ('1', 'HyESys H50 Main Unit',         '1',      '[SN-XXXXXX]',   'Sub-MSB Room / Panel',     '2 years'),
    ('2', 'HySBatt Battery Pack(s)',       '11',     '[SN-XXXXXX]',   'Sub-MSB Room',              '2 years'),
    ('3', 'Monitoring Gateway',            '1',      '[SN-XXXXXX]',   'Panel / DIN Rail',          '1 year'),
    ('4', 'Data Logger (Splitter)',        '1',      '[SN-XXXXXX]',   'Panel / DIN Rail',          '1 year'),
    ('5', 'Router / SIM Module',           '1',      '[SN-XXXXXX]',   'Panel / Cabinet',           '1 year'),
    ('6', 'Communication & Power Cables', '1 set',  'N/A',            'As installed',              'N/A'),
    ('7', 'Keys & Access Items',           '[qty]',  '[Key No.]',      '[Refer to KT]',            'N/A'),
    ('8', 'Manufacturer Docs & Certs',    '1 set',  'N/A',            'With this manual',          'N/A'),
]
for i, row in enumerate(comp_data):
    if i % 2 == 1: alt(comp, i + 1)
    for j, v in enumerate(row):
        dat(comp.rows[i+1].cells[j], v, centre=(j in [0, 2]))
set_table_borders(comp, 'C8E6C9')
note(doc, 'Note:', 'Key handover details to be confirmed by KT. 11 HySBatt packs are supplied for the H50 configuration. Manufacturer datasheets, FAT certificates and warranty cards are in Appendix F.')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — KEY WIRING REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '4.  Key Wiring Requirements')
body(doc, 'The following is a generic overview only. All final connections must follow the approved Site Single Line Diagram (SLD) and installation drawings in Appendix A.')

H2(doc, '4.1  Generic Wiring Overview')
wt = doc.add_table(rows=7, cols=2)
wt.style = 'Table Grid'
hdr(wt.rows[0].cells[0], 'Connection')
hdr(wt.rows[0].cells[1], 'Requirement', centre=False)
wt_data = [
    ('Three-Phase Power (L1, L2, L3)',
     'Connected to H50 AC terminals via the dedicated circuit breaker at the Sub-MSB. Phasing must be correct and terminations torqued to specification.'),
    ('Neutral (N)',
     'Connected to the H50 neutral terminal. Do not share this neutral with other downstream loads on the Sub-MSB.'),
    ('Protective Earth (PE)',
     'Dedicated earth conductor to the Sub-MSB earth busbar. Must not be daisy-chained with other equipment.'),
    ('Dedicated Circuit Breaker',
     'Correctly rated MCCB at the Sub-MSB as specified in the approved SLD. Final rating to be confirmed by Jeff.'),
    ('Isolation Switch',
     'Lockable isolation upstream of the dedicated breaker for safe maintenance isolation. Final arrangement to be confirmed by Jeff.'),
    ('Communication Wiring',
     'CT sensor cables (Sub-MSB incomer), RS-485/Modbus and Ethernet/SIM are routed per installation drawings. Do not re-route without AST approval.'),
]
for i, (l, v) in enumerate(wt_data):
    if i % 2 == 1: alt(wt, i + 1)
    lbl(wt.rows[i+1].cells[0], l)
    dat(wt.rows[i+1].cells[1], v)
set_table_borders(wt, 'C8E6C9')

sp(doc, 4)
H2(doc, '4.2  Cable and Breaker Ratings')
info(doc, 'To be confirmed:', 'Final cable and breaker ratings are confirmed by AST/Jeff per the approved SLD, based on cable length, ambient temperature and local authority requirements.')
warn(doc, 'WARNING:', 'No wiring modifications are permitted without AST written approval. Incorrect wiring voids warranty and creates electrical hazards.')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — POWERING UP SEQUENCE
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '5.  Powering Up Sequence')
note(doc, 'Routine operation:', 'HyESys starts and stops automatically. No daily client action is required. This section covers full power-up from a de-energised state only.')

H2(doc, '5.1  Power-Up from Cold')
for step in [
    'Confirm all maintenance work is complete and all personnel are clear of the Sub-MSB area.',
    'Verify the H50 dedicated circuit breaker is in the OFF position.',
    'Close the upstream isolation switch at the Sub-MSB.',
    'Switch on the auxiliary supply (if a separate auxiliary breaker is provided).',
    'Confirm the H50 display panel illuminates and shows the startup screen.',
    'Close the dedicated circuit breaker.',
    'Confirm the unit enters STANDBY or AUTO mode on the display.',
    'Verify that power factor and current readings appear normal on the monitoring interface.',
]:
    numbered(doc, step)

sp(doc, 6)
H2(doc, '5.2  Safe Isolation for Maintenance (De-Energise)')
for step in [
    'Select MANUAL HOLD or disable automatic operation on the H50 control panel.',
    'Wait for the unit to ramp down — confirm output current reaches zero on the display.',
    'Open (turn OFF) the dedicated circuit breaker at the Sub-MSB.',
    'Switch off the auxiliary supply.',
    'Apply lockout/tagout on the isolation switch per STES site safety procedure.',
    'Wait at least 5 minutes for DC bus capacitors to discharge.',
    'Verify zero voltage with an approved voltage tester before opening the enclosure.',
]:
    numbered(doc, step)

warn(doc, 'WARNING:', 'Never open the H50 enclosure without first confirming zero voltage. The DC bus operates at 350–850 V and retains charge after AC isolation.')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — NORMAL OPERATION
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '6.  Normal Operation')

H2(doc, '6.1  Operating Modes')
modes = doc.add_table(rows=5, cols=3)
modes.style = 'Table Grid'
for j, t in enumerate(['Mode', 'Display Indicator', 'Description']):
    hdr(modes.rows[0].cells[j], t)
modes_data = [
    ('AUTO',    'Green — AUTO',    'Unit operating normally. Reactive compensation, phase balancing and solar capture active based on Sub-MSB demand.'),
    ('STANDBY', 'Amber — STANDBY','Unit energised and ready but not actively injecting. Normal during very low-load or overnight periods.'),
    ('FAULT',   'Red — FAULT',    'Fault detected. Unit has safely tripped. Note the error code and contact AST before any further action.'),
    ('MANUAL',  'Blue — MANUAL',  'Under manual control for commissioning or maintenance only. Not a normal STES operating mode.'),
]
for i, (mode, ind, desc) in enumerate(modes_data):
    if i % 2 == 1: alt(modes, i + 1)
    colours_mode = [DK_GREEN, GOLD, PROBLEM_RED, STEEL_BLUE]
    dat(modes.rows[i+1].cells[0], mode, bold=True, centre=True, colour=colours_mode[i])
    dat(modes.rows[i+1].cells[1], ind, bold=True)
    dat(modes.rows[i+1].cells[2], desc)
set_table_borders(modes, 'C8E6C9')

sp(doc, 6)
H2(doc, '6.2  Expected Readings — Sub-MSB (Normal Conditions)')
rdgs = doc.add_table(rows=6, cols=3)
rdgs.style = 'Table Grid'
for j, t in enumerate(['Parameter', 'Normal Range / Expectation', 'Note']):
    hdr(rdgs.rows[0].cells[j], t)
rdgs_data = [
    ('Power Factor (PF)',      '>= 0.95 lagging',            'Target is 0.98. Pre-installation mean PF at this sub-MSB was 0.8189.'),
    ('Phase Current Balance',  '< 5% imbalance',             'Pre-installation imbalance was 29.9%. Active balancing reduces this from Day 1.'),
    ('Reactive Current (kVAr)','< 20 kVAr (within H50 range)','Pre-installation mean was 61.7 kVAr. H50 provides full correction below 50 kVAr.'),
    ('Output Current',         '<= 72.5 A (H50 rated)',      'Must not exceed H50 nameplate current continuously.'),
    ('DC Bus Voltage',         '350 – 850 V DC',             'H50 operating range. Normal variation reflects battery state of charge and reactive demand.'),
]
for i, row in enumerate(rdgs_data):
    if i % 2 == 1: alt(rdgs, i + 1)
    lbl(rdgs.rows[i+1].cells[0], row[0])
    dat(rdgs.rows[i+1].cells[1], row[1], bold=True, colour=DK_GREEN, centre=True)
    dat(rdgs.rows[i+1].cells[2], row[2])
set_table_borders(rdgs, 'C8E6C9')

sp(doc, 6)
H2(doc, '6.3  Confirming Correct Operation')
bullet(doc, 'The AST monitoring dashboard (browser/app) should show real-time PF trending >= 0.95 at the Sub-MSB incomer.')
bullet(doc, 'Monthly performance reports are generated by the AST cloud platform — compare Sub-MSB kVArh month-on-month against the Dec 2025 SP Group billing baseline.')
bullet(doc, 'If the display is blank or shows an unexpected mode, note the error code and contact AST before any action.')
note(doc, 'Validation benchmark:', 'Post-activation validation is via CT measurement at the Sub-MSB incomer and month-on-month comparison of SP Group billing kVArh and reactive energy readings.')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ALARMS AND BASIC TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '7.  Alarms and Basic Troubleshooting')
note(doc, 'Scope:', 'STES personnel are not expected to repair the unit. The actions below are limited to safe observation and basic resets only. Contact AST for all fault investigations.')

alm = doc.add_table(rows=9, cols=4)
alm.style = 'Table Grid'
for j, t in enumerate(['Alarm / Symptom', 'Possible Cause', 'Client Action', 'Contact AST?']):
    hdr(alm.rows[0].cells[j], t)
alm_data = [
    ('FAULT — Overcurrent',        'Load surge or short circuit at Sub-MSB',    'Check for obvious faults. Do not reset without AST guidance.',                                    'YES'),
    ('FAULT — Overvoltage',        'Grid voltage spike or wiring issue',         'Record fault code and time. Do not reset. Contact AST.',                                          'YES'),
    ('FAULT — Overtemperature',    'Blocked ventilation or high ambient temp',   'Check ventilation openings are clear. Allow unit to cool. Attempt one reset.',                   'Yes if fault persists'),
    ('FAULT — Comms Loss',         'Router/SIM issue or cable fault',            'Check router power and SIM signal. Restart router. Contact AST if unresolved.',                  'Yes if unresolved'),
    ('STANDBY > 24 hrs',           'Low reactive load or unit in manual mode',   'Check display — confirm AUTO mode is selected. Low-load periods (nights/weekends) are normal.',  'Yes if AUTO shown but inactive'),
    ('Display blank',              'Auxiliary supply off or internal fault',      'Check auxiliary circuit breaker. Reset once if tripped. Contact AST if blank persists.',        'Yes if persists'),
    ('PF not improving (dashboard)','Unit in STANDBY or reactive load at target','Verify AUTO mode. Review dashboard trend — PF may already be >= 0.95.',                         'Yes if PF < 0.90 consistently'),
    ('Unusual noise or smell',     'Mechanical or electrical fault',              'Immediately open the dedicated circuit breaker. Do not re-energise. Contact AST urgently.',      'YES — immediately'),
]
for i, row in enumerate(alm_data):
    if i % 2 == 1: alt(alm, i + 1)
    lbl(alm.rows[i+1].cells[0], row[0])
    dat(alm.rows[i+1].cells[1], row[1])
    dat(alm.rows[i+1].cells[2], row[2])
    is_yes = row[3].startswith('YES')
    dat(alm.rows[i+1].cells[3], row[3], bold=is_yes,
        colour=PROBLEM_RED if is_yes else DK_GREEN, centre=True)
set_table_borders(alm, 'C8E6C9')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — DOS AND DON'TS
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, "8.  Dos and Don'ts")

ddt = doc.add_table(rows=2, cols=2)
ddt.style = 'Table Grid'
hdr(ddt.rows[0].cells[0], 'DO'); set_cell_bg(ddt.rows[0].cells[0], HEX['green'])
hdr(ddt.rows[0].cells[1], 'DO NOT'); set_cell_bg(ddt.rows[0].cells[1], HEX['red'])

dos = [
    'Keep ventilation openings unobstructed at all times.',
    'Conduct a visual inspection monthly — check for dust, water ingress, loose labels and indicator status.',
    'Verify the monitoring dashboard is accessible and data is updating weekly.',
    'Report any unusual alarms, smells or sounds to AST immediately.',
    'Keep the Sub-MSB area around the unit clear of combustible materials.',
    'Ensure the dedicated circuit breaker is accessible, labelled and not obstructed.',
    'Follow the isolation procedure before any maintenance work near the unit.',
    'Retain all documentation in a safe and accessible location at the site.',
]
donts = [
    'Open the enclosure unless authorised and trained to do so.',
    'Modify or re-route any wiring without AST written approval.',
    'Stack items against or on top of the H50 unit or battery packs.',
    'Attempt to reset the unit more than once without AST guidance.',
    'Allow unauthorised personnel to operate the H50 controls.',
    'Apply water or liquid cleaners to or near the unit.',
    'Connect additional loads to the H50 dedicated circuit breaker.',
    'Disable alarms or override fault conditions independently.',
]
do_c = ddt.rows[1].cells[0]; dont_c = ddt.rows[1].cells[1]
set_cell_bg(do_c, HEX['lt_green']); set_cell_bg(dont_c, HEX['lt_red'])
for item in dos:
    p = do_c.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(9); r.font.color.rgb = DK_GREEN; r.font.name = 'Aptos'
for item in donts:
    p = dont_c.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(9); r.font.color.rgb = PROBLEM_RED; r.font.name = 'Aptos'
set_table_borders(ddt, 'C8E6C9')

sp(doc, 6)
H2(doc, '8.1  Inspection and Maintenance Schedule')
maint = doc.add_table(rows=5, cols=3)
maint.style = 'Table Grid'
for j, t in enumerate(['Frequency', 'Activity', 'Responsible']):
    hdr(maint.rows[0].cells[j], t)
maint_data = [
    ('Monthly',     'Visual inspection — ventilation, cleanliness, labels, HySBatt pack condition, indicator status.',    'STES Facility Manager'),
    ('Quarterly',   'Visual check of cable terminations at the Sub-MSB for security and damage (external only).',          'STES Facility Manager'),
    ('Bi-annually', 'Full inspection and functional test at Sub-MSB.',                                                     'AST Service Engineer'),
    ('Annually',    'Preventive maintenance, firmware update, battery health assessment, CT calibration check.',           'AST Service Engineer'),
]
for i, row in enumerate(maint_data):
    if i % 2 == 1: alt(maint, i + 1)
    lbl(maint.rows[i+1].cells[0], row[0])
    dat(maint.rows[i+1].cells[1], row[1])
    dat(maint.rows[i+1].cells[2], row[2])
set_table_borders(maint, 'C8E6C9')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — SAFETY AND EMERGENCY RESPONSE
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '9.  Safety and Emergency Response')

H2(doc, '9.1  Electrical Hazards')
bullet(doc, 'The H50 operates at 400 V AC (three-phase) and DC bus voltages of 350–850 V. Contact with live terminals is potentially fatal.')
bullet(doc, 'DC bus capacitors retain charge for up to 5 minutes after AC isolation. Always allow discharge time and verify with a live voltage indicator before opening the enclosure.')
bullet(doc, 'Do not assume the unit is safe because the display is off — internal voltages may still be present.')

sp(doc, 4)
H2(doc, '9.2  Required PPE')
bullet(doc, 'Insulated safety gloves — minimum Cat III / 1000 V rated.')
bullet(doc, 'Arc flash-rated face shield and flame-resistant clothing appropriate to the Sub-MSB arc flash level.')
bullet(doc, 'Non-conductive safety footwear.')
bullet(doc, 'Approved voltage detection / live line indicator.')

sp(doc, 4)
H2(doc, '9.3  Emergency Shutdown')
for step in [
    'On detecting fire, smoke, abnormal sound or burning smell — immediately open the H50 dedicated circuit breaker at the Sub-MSB.',
    'If safe to do so, open the upstream isolation switch and apply STES lockout/tagout.',
    'Evacuate the Sub-MSB area and follow the STES site emergency response plan.',
    'Contact AST and the STES site emergency coordinator immediately.',
    'Do not re-energise the H50 until AST has inspected the unit and issued written clearance.',
]:
    numbered(doc, step)
warn(doc, 'FIRE:', 'Use CO₂ or dry powder extinguisher only on electrical equipment. Never use water on the H50 or Sub-MSB equipment.')

sp(doc, 4)
H2(doc, '9.4  Emergency Contacts')
ec = doc.add_table(rows=4, cols=3)
ec.style = 'Table Grid'
for j, t in enumerate(['Contact', 'Organisation / Name', 'Number']):
    hdr(ec.rows[0].cells[j], t)
ec_data = [
    ('AST 24-hr Support',      'Advancer Smart Technology Pte Ltd',  '[+65 XXXX XXXX]'),
    ('STES Facility Manager',  '[Name]',                              '[Contact Number]'),
    ('SP PowerGrid Emergency', 'SP PowerGrid Ltd',                    '1800 778 8888'),
]
for i, row in enumerate(ec_data):
    if i % 2 == 1: alt(ec, i + 1)
    lbl(ec.rows[i+1].cells[0], row[0])
    dat(ec.rows[i+1].cells[1], row[1])
    dat(ec.rows[i+1].cells[2], row[2], bold=True)
set_table_borders(ec, 'C8E6C9')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — CLIENT HANDOVER AND TRAINING
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '10.  Client Handover and Training')

H2(doc, '10.1  Training Attendance')
tr = doc.add_table(rows=5, cols=4)
tr.style = 'Table Grid'
for j, t in enumerate(['Full Name', 'Designation', 'Organisation', 'Signature']):
    hdr(tr.rows[0].cells[j], t)
for i in range(1, 5):
    if i % 2 == 1: alt(tr, i)
    for j in range(4): dat(tr.rows[i].cells[j], '')
set_table_borders(tr, 'C8E6C9')

sp(doc, 6)
H2(doc, '10.2  Handover Checklist')
hc = doc.add_table(rows=10, cols=3)
hc.style = 'Table Grid'
for j, t in enumerate(['Item', 'Status', 'Remarks']):
    hdr(hc.rows[0].cells[j], t)
hc_items = [
    'Site commissioning completed and signed off at Sub-MSB.',
    'Equipment verified against handover list (Section 3), including 11 HySBatt packs.',
    'Keys and access items handed over (KT).',
    'STES personnel trained on normal operation and monitoring dashboard.',
    'STES personnel trained on isolation and emergency shutdown procedure.',
    'Monitoring dashboard login credentials provided to STES.',
    'All manufacturer documents and certificates handed over.',
    'Approved SLD and wiring drawings for Sub-MSB handed over.',
    'Outstanding items noted with agreed resolution dates.',
]
for i, item in enumerate(hc_items):
    if i % 2 == 1: alt(hc, i + 1)
    dat(hc.rows[i+1].cells[0], item)
    dat(hc.rows[i+1].cells[1], 'Complete  /  Pending', centre=True)
    dat(hc.rows[i+1].cells[2], '')
set_table_borders(hc, 'C8E6C9')

sp(doc, 6)
H2(doc, '10.3  Outstanding Items')
oi = doc.add_table(rows=4, cols=4)
oi.style = 'Table Grid'
for j, t in enumerate(['No.', 'Description', 'Responsible Party', 'Target Date']):
    hdr(oi.rows[0].cells[j], t)
for i in range(1, 4):
    if i % 2 == 1: alt(oi, i)
    dat(oi.rows[i].cells[0], str(i), centre=True)
    for j in range(1, 4): dat(oi.rows[i].cells[j], '')
set_table_borders(oi, 'C8E6C9')

sp(doc, 6)
H2(doc, '10.4  Client Acknowledgement')
body(doc, 'We confirm that the HyESys H50 system at the Sub-MSB has been commissioned and handed over, and that we have received sufficient training to operate the system in accordance with this manual.')
sp(doc, 4)
ack = doc.add_table(rows=4, cols=3)
ack.style = 'Table Grid'
hdr(ack.rows[0].cells[0], '')
hdr(ack.rows[0].cells[1], 'AST Representative')
hdr(ack.rows[0].cells[2], 'STES Representative')
for i, label_text in enumerate(['Name & Designation', 'Signature', 'Date']):
    if i % 2 == 1: alt(ack, i + 1)
    lbl(ack.rows[i+1].cells[0], label_text)
    dat(ack.rows[i+1].cells[1], '')
    dat(ack.rows[i+1].cells[2], '')
set_table_borders(ack, 'C8E6C9')
sp(doc, 4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — SITE SPECIFIC APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
H1(doc, '11.  Site-Specific Appendices')
body(doc, 'The following site-specific documents for 249 Jln Boon Lay are attached. All must be updated if any changes are made to the installation.')

app = doc.add_table(rows=9, cols=3)
app.style = 'Table Grid'
for j, t in enumerate(['Appendix', 'Document', 'Status']):
    hdr(app.rows[0].cells[j], t)
app_data = [
    ('A', 'Approved Single Line Diagram (SLD) — Sub-MSB, 249 Jln Boon Lay'),
    ('B', 'Wiring and Termination Drawings — H50 at Sub-MSB'),
    ('C', 'Equipment Layout Drawing — Sub-MSB Room'),
    ('D', 'Cable Route Drawing'),
    ('E', 'Commissioning Records and Sign-Off Sheet'),
    ('F', 'Test Reports (FAT, SAT, insulation, CT accuracy at Sub-MSB incomer)'),
    ('G', 'Equipment Datasheets (H50 unit, HySBatt packs ×11, Gateway, Logger)'),
    ('H', 'Warranty Cards and Certificates'),
]
for i, (app_id, desc) in enumerate(app_data):
    if i % 2 == 1: alt(app, i + 1)
    dat(app.rows[i+1].cells[0], app_id, bold=True, colour=GOLD, centre=True)
    dat(app.rows[i+1].cells[1], desc)
    dat(app.rows[i+1].cells[2], 'Attached  /  Pending', centre=True)
set_table_borders(app, 'C8E6C9')

sp(doc, 8)
gold_box(doc, 'Reminder:', 'This manual should be reviewed and updated after any change to the Sub-MSB installation, firmware upgrade or change in site operating conditions. Contact AST for the latest version.')

# ── Footer strip ──────────────────────────────────────────────────────────────
sp(doc, 10)
ft = doc.add_table(rows=2, cols=1)
ft.style = 'Table Grid'
set_cell_bg(ft.rows[0].cells[0], HEX['green'])
set_row_height(ft.rows[0], 4)
ft.rows[0].cells[0].paragraphs[0].text = ''
set_cell_bg(ft.rows[1].cells[0], HEX['navy'])
fp = ft.rows[1].cells[0].paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(5); fp.paragraph_format.space_after = Pt(5)
rg = fp.add_run('HyESys H50')
rg.font.size = Pt(9); rg.font.bold = True; rg.font.color.rgb = GOLD; rg.font.name = 'Aptos'
rr = fp.add_run('  ·  ST Engineering Land Systems Ltd.  ·  249 Jln Boon Lay, S619523  ·  Advancer Smart Technology Pte Ltd  ·  Confidential')
rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor(0xCC,0xDD,0xEE); rr.font.name = 'Aptos'
for cell in [ft.rows[0].cells[0], ft.rows[1].cells[0]]: no_border(cell)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
