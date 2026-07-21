from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd\AST BD\HyESys Dept\7. Client Projects\MAS\tender\AST_Tender_Proposal_CSD_T_26_HES_KCH_005.docx"

# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def add_run(para, text, bold=False, size=11, colour=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*bytes.fromhex(colour))
    return run

def heading(doc, text, level, colour_hex=None):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    run.font.size = Pt(16 - (level * 2))
    if colour_hex:
        run.font.color.rgb = RGBColor(*bytes.fromhex(colour_hex))
    return p

def body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def page_break(doc):
    doc.add_page_break()

def two_col_table(doc, rows, col_widths=(2.0, 4.5), header=None, header_bg='1F4E79'):
    tbl = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    offset = 0
    if header:
        r = tbl.rows[0]
        for ci, text in enumerate(header):
            c = r.cells[ci]
            c.width = Inches(col_widths[ci])
            set_cell_bg(c, header_bg)
            p = c.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        offset = 1
    for ri, (k, v) in enumerate(rows):
        r = tbl.rows[ri + offset]
        c0 = r.cells[0]
        c0.width = Inches(col_widths[0])
        c0.paragraphs[0].add_run(k).font.size = Pt(11)
        c0.paragraphs[0].runs[0].bold = True
        c1 = r.cells[1]
        c1.width = Inches(col_widths[1])
        c1.paragraphs[0].add_run(v).font.size = Pt(11)
    return tbl

# ─────────────────────────────────────────────────────────────
# Document
# ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
add_run(p, 'TENDER PROPOSAL', bold=True, size=22, colour='1F4E79')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p,
        'Supply, Installation, Testing, Commissioning and Integration of a\n'
        'Hybrid Energy System (HES)\nat MAS Building, 10 Shenton Way, Singapore 079117',
        size=14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Tender Reference: CSD/T/26/HES/KCH/005', bold=True, size=13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Submitted by:', bold=True, size=12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Advancer Smart Technology Pte. Ltd.', bold=True, size=13, colour='1F4E79')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '55 Toh Guan Rd E, #06-02, Singapore 608601', size=11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Submitted to:', bold=True, size=12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Monetary Authority of Singapore', bold=True, size=12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Contact Person:', bold=True, size=11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Dr Lim Chee Chong\nDirector\nAdvancer Smart Technology Pte. Ltd.\nTel: _______________', size=11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Date: 17 June 2026', bold=True, size=11)

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════
heading(doc, '1.  EXECUTIVE SUMMARY', 1, '1F4E79')

body(doc,
     'Advancer Smart Technology Pte. Ltd. (AST) submits this proposal in response to ITT Reference '
     'CSD/T/26/HES/KCH/005. AST proposes the HySBatt Hybrid Energy System comprising 11 battery packs '
     'with a 50 kW Power Conversion System (PCS), delivering 110 kWh of usable energy storage with '
     'integrated reactive compensation, three-phase load balancing, and peak demand management at '
     'MASB Level 6 Carpark. SCDF Conditional Approval (Ref: WVR/00109/26, dated 17 April 2026) '
     'has been obtained for this site.')

two_col_table(doc, [
    ('Tender Reference',           'CSD/T/26/HES/KCH/005'),
    ('Client',                     'Monetary Authority of Singapore'),
    ('Site',                       'MAS Building, 10 Shenton Way, Singapore 079117'),
    ('Deployment Location',        'Level 6 Carpark'),
    ('System',                     '11 × HySBatt + 50 kW PCS'),
    ('Total Usable Energy',        '110 kWh'),
    ('Footprint',                  '3 m (L) × 2 m (W) × 2.5 m (H)'),
    ('Fire Rating',                '2HR, BS 476-20 / UL9540A Unit Level Certified'),
    ('SCDF Approval',              'WVR/00109/26 — Conditionally Approved, 17 Apr 2026'),
    ('Min. Energy Savings Target', '≥ 4,000 kWh/month'),
    ('Power Factor Target',        '≥ 0.98 (min. 15% improvement)'),
    ('Harmonic Reduction Target',  '≥ 25% THD reduction'),
    ('System Uptime',              '≥ 99%'),
], header=['Parameter', 'Detail'])

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 2 — COMPANY CREDENTIALS
# ════════════════════════════════════════════════════════════
heading(doc, '2.  COMPANY CREDENTIALS AND TRIAL CAPABILITY', 1, '1F4E79')

heading(doc, '2.1  Company Profile', 2)
body(doc,
     'AST is a Singapore-based smart building IoT and AI company specialising in energy management '
     'and power quality optimisation. Products include HyESys and HySBatt, designed and engineered in '
     'Singapore for commercial and industrial buildings.')

two_col_table(doc, [
    ('Company Name',       'Advancer Smart Technology Pte. Ltd.'),
    ('Address',            '55 Toh Guan Rd E, #06-02, Singapore 608601'),
    ('Contact Person',     'Dr Lim Chee Chong, Director'),
    ('Telephone',          '_______________'),
    ('Email',              '_______________'),
    ('Key Product',        'HyESys / HySBatt Hybrid Energy System'),
], header=['Detail', 'Information'])

heading(doc, '2.2  Technical Team', 2)
bullet(doc, 'Power systems engineering: reactive compensation, load balancing, harmonic mitigation')
bullet(doc, 'Battery energy storage systems: cell-level design, BMS integration, thermal management')
bullet(doc, 'AI and data analytics: multi-agent monitoring, predictive maintenance, SAR-based control')
bullet(doc, 'Regulatory expertise: SCDF fire safety submissions, SP PowerGrid electrical submissions')
bullet(doc, 'Project management: phased deployments, method of statement, commissioning and handover')

heading(doc, '2.3  Financial Capability', 2)
body(doc,
     'AST bears all costs for the Initial Period — equipment, installation, testing, commissioning, '
     'support, and reporting — at no cost to MAS. Financial statements are available upon request.')

heading(doc, '2.4  Insurance Coverage', 2)
bullet(doc, 'Public liability insurance covering all on-site works')
bullet(doc, 'Professional indemnity insurance for engineering services')
bullet(doc, 'Equipment insurance covering all HySBatt and PCS hardware')
bullet(doc, 'Workers\' compensation insurance for all personnel')
body(doc, 'Insurance certificates will be provided prior to commencement of works.', italic=True)

heading(doc, '2.5  Reference Project', 2)
two_col_table(doc, [
    ('Project',       'Energy Storage System Installation — Galen Building'),
    ('Site',          '61 Science Park Road, Singapore 117525 (Level 1 Outdoor)'),
    ('Applicant',     'Advancer Smart Technology Pte. Ltd.'),
    ('SCDF Approval', 'RBP/A02034/24 — Full Notice of Approval, 11 March 2024'),
    ('Basis',         'Section 55(5), Fire Safety Act 1993'),
    ('QP',            'SYT Consultants Pte Ltd'),
    ('Outcome',       'Full approval granted; system installed and operational'),
], header=['Item', 'Detail'])

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 3 — TECHNICAL PROPOSAL
# ════════════════════════════════════════════════════════════
heading(doc, '3.  TECHNICAL PROPOSAL', 1, '1F4E79')

heading(doc, '3.1  System Architecture', 2)
body(doc,
     'The proposed HySBatt Hybrid Energy System is an active digital power compensator simultaneously '
     'delivering reactive compensation, 3-phase load balancing, and energy storage from a single unit.')

two_col_table(doc, [
    ('Power Conversion System (PCS)',   '50 kW active digital inverter — reactive compensation, load balancing, harmonic filtering'),
    ('HySBatt Battery Packs',           '11 units in series; 10 kWh usable per pack; 35 V nominal; 110 kWh total'),
    ('Battery Management System (BMS)', 'Cell-level voltage monitoring, SoH tracking, thermal management, fault isolation'),
    ('IoT Sensor Network',              'Real-time monitoring: kW, kVAr, PF, current per phase, voltage, THD, temperature'),
    ('AI Control Engine',               'Multi-agent neural network; self-learning optimisation; real-time anomaly detection'),
    ('Web Dashboard',                   'Real-time analytics, remote access, customisable reporting'),
    ('Fire-Rated Enclosure',            '2HR per BS 476-20; 2HR motorised damper (SS333 CoC); Rockwool/gypsum ducting'),
    ('Gas Detection',                   'H₂ sensor at filter inlet + smoke sensor at ceiling; auto-triggers carpark purge'),
    ('Automated Transfer Switch (ATS)', 'Seamless switching between grid and HySBatt modes; emergency isolation for business continuity'),
], header=['Component', 'Specification'], col_widths=(2.3, 4.2))

heading(doc, '3.2  Deployment Configuration — Level 6 Carpark', 2)
two_col_table(doc, [
    ('HySBatt Packs',         '11 units'),
    ('PCS Rating',            '50 kW'),
    ('Usable Energy',         '110 kWh'),
    ('Arrangement',           'Vertical stacking with spreader plate'),
    ('Pack Spacing',          '50 mm (UL9540A specification)'),
    ('Footprint',             '3 m (L) × 2 m (W) × 2.5 m (H)'),
    ('Fire Rating',           '2HR, BS 476-20 / UL9540A Unit Level Certified'),
    ('Structural Compliance', 'Floor rated 7.5 kN/m²; load within safe limits with spreader plate'),
    ('SCDF Approval',         'WVR/00109/26 — Conditionally Approved, 17 Apr 2026'),
], header=['Parameter', 'Detail'])

heading(doc, '3.3  SCDF Compliance', 2)
body(doc, 'All four SCDF waiver conditions (WVR/00109/26) are fully addressed by the system design:')

tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
scdf_widths = [Inches(0.4), Inches(2.7), Inches(3.4)]
for ci, h in enumerate(['No.', 'SCDF Condition', 'AST Compliance']):
    c = tbl.rows[0].cells[ci]
    c.width = scdf_widths[ci]
    set_cell_bg(c, '1F4E79')
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

conditions = [
    ('1', 'UL9540A test at unit level',
          'All HySBatt units carry UL9540A unit-level certification'),
    ('2', 'Flammable gas + smoke detectors to activate carpark purge',
          'H₂ sensor at filter inlet + smoke sensor at ceiling; wired to carpark purge'),
    ('3', 'Fencing at 1.2 m clearance, 1.8 m height',
          'Safety fencing at 1.2 m standoff, 1.8 m height'),
    ('4', 'ESS enclosure per NFPA 69:2019 Ch.13 deflagration control',
          'Enclosure validated by QP against NFPA 69:2019 Ch.13'),
]
for ri, (no, req, comp) in enumerate(conditions):
    r = tbl.rows[ri + 1]
    for ci, val in enumerate([no, req, comp]):
        c = r.cells[ci]
        c.width = scdf_widths[ci]
        c.paragraphs[0].add_run(val).font.size = Pt(10)

heading(doc, '3.4  Performance Projections', 2)
two_col_table(doc, [
    ('Energy Savings',            '≥ 4,000 kWh/month for 6 consecutive months'),
    ('Power Factor',              '≥ 0.98 (min. 15% improvement from baseline)'),
    ('Harmonic Reduction',        '≥ 25% THD reduction'),
    ('Neutral Current Reduction', '≥ 80%'),
    ('System Availability',       '≥ 99% uptime'),
    ('Peak Savings (Optimal)',    'Up to 30% energy savings under optimal conditions'),
], header=['Metric', 'Target'])

body(doc,
     'Savings are measured using the current-reduction fraction method anchored to RMS current at '
     'the MSB incomer. Baseline is established over a minimum 30-day pre-installation period '
     'coordinated with MAS\'s BMS contractor.')

heading(doc, '3.5  Multi-Floor Scalability', 2)
body(doc,
     'The system is pre-wired for 3-floor operation during the initial 6-week installation phase — '
     'all cable trays, conduits, and communication backbones for additional floors are installed and '
     'tested, with no further civil works required for expansion.')
bullet(doc, 'Phase 1: Level 6 active; 3-floor infrastructure pre-installed and tested')
bullet(doc, 'Expansion: each additional floor activated within 5 working days upon MAS written request')
bullet(doc, 'No disruption to Level 6 operations during expansion')
bullet(doc, 'AI control parameters auto-adapt to expanded load profile')

heading(doc, '3.6  Integration Methodology', 2)
bullet(doc, 'Installation between revenue meter and MSB — no changes to downstream distribution boards')
bullet(doc, 'BMS integration via Modbus, BACnet, or equivalent communication protocols')
bullet(doc, 'Compatible with existing SP PowerGrid revenue-grade digital meters')
bullet(doc, 'Zero changes to existing equipment, user behaviour, or operations')
bullet(doc, 'Web dashboard accessible to MAS facility managers with remote control capability')

heading(doc, '3.7  Reference Installation: Galen Building', 3)
two_col_table(doc, [
    ('Site',          'The Galen Building, 61 Science Park Road, Singapore 117525'),
    ('SCDF Ref.',     'RBP/A02034/24'),
    ('Approval Date', '11 March 2024'),
    ('Approval Type', 'Full Notice of Approval under Section 55(5), Fire Safety Act 1993'),
    ('QP Firm',       'SYT Consultants Pte Ltd'),
    ('Outcome',       'Full SCDF approval; system installed and commissioned'),
], header=['Item', 'Detail'])

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 4 — INITIAL PERIOD MANAGEMENT PLAN
# ════════════════════════════════════════════════════════════
heading(doc, '4.  INITIAL PERIOD MANAGEMENT PLAN', 1, '1F4E79')

heading(doc, '4.1  Project Timeline', 2)
tbl = doc.add_table(rows=9, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for ci, h in enumerate(['Phase', 'Duration', 'Activity']):
    c = tbl.rows[0].cells[ci]
    set_cell_bg(c, '1F4E79')
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

timeline = [
    ('Pre-Installation',       'Day 1–30',  'Baseline data collection; BMS contractor coordination; site survey; method of statement submission'),
    ('Site Preparation',       'Week 1–2',  'Unistrut frame; cable trays; wall anchoring; 3-floor infrastructure pre-wiring'),
    ('Equipment Installation', 'Week 3–4',  'HySBatt and PCS installation; BMS and IoT sensors; safety fencing'),
    ('Safety Systems',         'Week 5',    'Gas/smoke detectors + carpark purge integration; fire damper; enclosure sealing'),
    ('Commissioning',          'Week 6',    'Full system testing; integration with existing infrastructure; AI calibration'),
    ('Phase 1 Trial',          'Month 1–2', 'Level 6 operations; continuous monitoring; weekly reports; optimisation'),
    ('Trial (Active)',         'Month 3–5', 'Performance stabilisation; monthly reviews; expansion readiness testing'),
    ('Trial Completion',       'Month 6',   'Final verification; 6-month summary report; Option to Purchase evaluation'),
]
for ri, (phase, dur, act) in enumerate(timeline):
    r = tbl.rows[ri + 1]
    r.cells[0].paragraphs[0].add_run(phase).font.size = Pt(10)
    r.cells[1].paragraphs[0].add_run(dur).font.size = Pt(10)
    r.cells[2].paragraphs[0].add_run(act).font.size = Pt(10)

heading(doc, '4.2  Performance Monitoring and Reporting', 2)
bullet(doc, 'Continuous AI-driven monitoring every 15 minutes: kW, kVAr, PF, current, THD, voltage')
bullet(doc, 'Live web dashboard accessible to MAS facility managers')
bullet(doc, 'Weekly summaries; monthly energy savings verification; quarterly comprehensive reviews')
bullet(doc, 'All savings measured against a mutually agreed 30-day pre-installation baseline')
bullet(doc, 'Advance notice (min. 48 hours) for all scheduled maintenance windows')
bullet(doc, '24/7 technical support with a dedicated AST project manager as single point of contact')

heading(doc, '4.3  Expansion Protocol', 2)
numbered(doc, 'Submit written expansion method of statement to MAS ≥ 5 working days in advance')
numbered(doc, 'Verify Phase 1 performance is stable prior to expansion works')
numbered(doc, 'Conduct load simulation on pre-installed floor circuits before energisation')
numbered(doc, 'Energise additional floor and validate all performance metrics within 24 hours of activation')

heading(doc, '4.4  Risk Management', 2)
two_col_table(doc, [
    ('Performance Shortfall',    'Root cause investigation and corrective action. If 4,000 kWh/month cannot be sustained, AST removes the system at no cost and reinstates the original electrical configuration.'),
    ('System Failure',           '≤ 4-hour response (business hours); ≤ 2-hour response (safety-critical). Backup procedures ensure immediate restoration of original configuration.'),
    ('Expansion Risk',           'All risks borne by AST. Issues resolved at no cost to MAS with no disruption to Level 6 operations.'),
    ('Safety Incident',          'Comprehensive insurance in place. Immediate remediation protocol activated. Zero-tolerance policy for incidents attributable to HySBatt installation.'),
    ('Technology Obsolescence',  'Firmware, software, and algorithm updates provided throughout the Initial Period at no additional cost.'),
], header=['Risk', 'Mitigation'], col_widths=(1.8, 4.7))

heading(doc, '4.5  Stakeholder Communication', 2)
bullet(doc, 'Designated AST project manager as single point of contact for MAS facility management')
bullet(doc, '24/7 technical support hotline for all queries and emergency response')
bullet(doc, 'Monthly performance review meetings with MAS representatives')
bullet(doc, 'Advance notice (min. 48 hours) for all scheduled maintenance windows')
bullet(doc, 'Immediate notification for any safety-critical events or system anomalies')

heading(doc, '4.6  Success Criteria', 2)
two_col_table(doc, [
    ('Energy Savings',      '≥ 4,000 kWh/month sustained for 6 consecutive months'),
    ('Power Factor',        '≥ 0.98 (min. 15% improvement from baseline)'),
    ('Harmonic Reduction',  '≥ 25% THD reduction sustained throughout trial'),
    ('System Availability', '≥ 99% uptime over the 6-month trial'),
    ('Load Balancing',      'Neutral current reduction ≥ 80% at all times'),
    ('Scalability',         'Expansion to additional floors within 5 working days upon request'),
    ('Safety',              'Zero incidents or electrical failures attributable to HySBatt installation'),
    ('User Impact',         'No disruption to normal MASB operations at any stage'),
], header=['Criterion', 'Target'])

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 5 — HES MAINTENANCE PLAN
# ════════════════════════════════════════════════════════════
heading(doc, '5.  HES MAINTENANCE PLAN', 1, '1F4E79')

body(doc,
     'All maintenance is provided at no cost throughout the Initial Period, with a maximum of '
     '8 scheduled maintenance hours per month and a minimum 48-hour advance notice to MAS.')

heading(doc, '5.1  Physical Safety Inspection (Monthly)', 2)
bullet(doc, 'Structural integrity of battery rack, PCS mounting, and fencing')
bullet(doc, 'Cable condition, piping, and ventilation for wear, overheating, or damage')
bullet(doc, 'Fire-rated walls, dampers, and structural components')
bullet(doc, 'General cleanliness and safety signage')
body(doc, 'Equipment: flashlight/headlamp, inspection mirror, cleaning cloth, soft brush, non-conductive vacuum, digital thermometer, cable tie gun/flush cutters.', italic=True)

heading(doc, '5.2  Connectivity Check (Monthly)', 2)
bullet(doc, 'Network communication between HES and Sub Alarm Panel — ping and packet loss verification')
bullet(doc, 'BMS data feed to MAS BMS and AST remote monitoring platform')
bullet(doc, 'IoT sensor data streams — verify all within expected ranges')
body(doc, 'Equipment: laptop with terminal emulator, CAN scanner via RS-485 cables.', italic=True)

heading(doc, '5.3  Battery Health Check (Quarterly)', 2)
bullet(doc, 'State of Health (SoH) for all 11 packs')
bullet(doc, 'Cell voltage balance and internal resistance per pack')
bullet(doc, 'Temperature uniformity across the pack array')
bullet(doc, 'BMS alarm log review; insulation resistance per pack')
body(doc, 'Equipment: BMS interface software, digital multimeter with test leads, battery internal resistance tester, infrared thermometer.', italic=True)

heading(doc, '5.4  Condition Assessment (Quarterly)', 2)
bullet(doc, 'Thermal imaging of battery packs and PCS connections to identify hot spots')
bullet(doc, 'Gas sensor calibration — verify H₂ and smoke sensor response thresholds')
bullet(doc, 'Re-torque all battery terminal connections to manufacturer specifications')
bullet(doc, 'Check damper clearance')
body(doc, 'Equipment: gas calibration kit for 4-in-1 sensor, thermal imaging camera, torque wrench set, feeler gauge.', italic=True)

heading(doc, '5.5  Consumable Replacement (As Required)', 2)
bullet(doc, 'Activated carbon filter in gas filtration system')
bullet(doc, 'Any consumables identified during scheduled inspections (cable ties, gaskets, signage)')
bullet(doc, 'All replacement parts supplied by AST at no additional cost')

heading(doc, '5.6  Maintenance Schedule Summary', 2)
two_col_table(doc, [
    ('Physical Safety Inspection', 'Monthly'),
    ('Connectivity Check',         'Monthly'),
    ('Battery Health Check',       'Quarterly (or on BMS alert)'),
    ('Condition Assessment',       'Quarterly'),
    ('Consumable Replacement',     'As required during scheduled inspection'),
    ('Remote Monitoring',          'Continuous — automated alerts to AST operations centre'),
    ('Emergency Response',         '≤ 4 hours (business hours) / ≤ 2 hours (safety-critical)'),
    ('Max. Scheduled Downtime',    '≤ 8 hours/month; 48-hour advance notice to MAS'),
], header=['Activity', 'Frequency / SLA'])

page_break(doc)

# ════════════════════════════════════════════════════════════
# SECTION 6 — BACKGROUND INTELLECTUAL PROPERTY (15.2e)
# ════════════════════════════════════════════════════════════
heading(doc, '6.  BACKGROUND INTELLECTUAL PROPERTY', 1, '1F4E79')

body(doc,
     'The following Background Intellectual Property (Background IP) owned by AST will be used '
     'under the Contract:')

two_col_table(doc, [
    ('HySBatt System',           'Hardware design, firmware, and Battery Management System (BMS) software for the HySBatt Hybrid Energy System'),
    ('HyESys AI Platform',       'Multi-agent AI control algorithms, energy optimisation engine, and autonomous decision-making framework'),
    ('IoT Monitoring Dashboard', 'Web-based monitoring platform, real-time data analytics engine, and remote access control system'),
    ('Power Quality Algorithms', 'Reactive compensation, 3-phase load balancing, and harmonic filtering algorithms embedded in PCS firmware'),
], header=['Background IP', 'Description'])

body(doc,
     'All Background IP listed above remains the property of AST. MAS is granted a licence to use '
     'the Background IP solely for the purposes of the Contract. Upon exercise of the Option to '
     'Purchase, MAS shall be granted a perpetual non-exclusive licence to use the Background IP '
     'for ongoing operation of the HES.', italic=True)

# ─── Closing ───────────────────────────────────────────────
doc.add_paragraph()
body(doc, 'Submitted by Advancer Smart Technology Pte. Ltd.', bold=True)
body(doc, 'For and on behalf of:', italic=True)
doc.add_paragraph()
body(doc, 'Dr Lim Chee Chong', bold=True)
body(doc, 'Director')
body(doc, 'Advancer Smart Technology Pte. Ltd.')
body(doc, '55 Toh Guan Rd E, #06-02, Singapore 608601')
body(doc, 'Tel: _______________')
doc.add_paragraph()
body(doc, 'Date: 17 June 2026')

# Save
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print("Saved:", OUTPUT)
