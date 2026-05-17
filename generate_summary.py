from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Styles ---
def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def heading3(doc, text):
    doc.add_heading(text, level=3)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def equation_box(doc, label, eq):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    runner = p.add_run(f"{label}:  {eq}")
    runner.bold = True
    runner.font.name = 'Courier New'
    runner.font.size = Pt(10)
    runner.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def note(doc, text):
    p = doc.add_paragraph()
    runner = p.add_run(f"Note: {text}")
    runner.italic = True
    runner.font.size = Pt(9)
    runner.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# =========================================================
# TITLE PAGE
# =========================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FDM6200 Fire Dynamics Workshop 2")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Compartment Fires and Fire Safety\nKey Concepts & Equations Summary")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph("Presenter: DC Teong How Hwa, Deputy Commissioner (Strategy & Corporate Services)")
doc.add_paragraph("Reference: Drysdale (2011), SFPE Handbook 5th Ed., Buchanan (2002), BS EN 1991-1-2:2002")
doc.add_paragraph()

# =========================================================
# SECTION 1: FIRE GROWTH STAGES
# =========================================================
heading1(doc, "1. Compartment Fire Growth Stages")

body(doc, "A compartment fire progresses through four distinct stages:")

stages = [
    ("Ignition", "The period during which the fire begins."),
    ("Growth (Pre-Flashover)", "Fire grows independent of compartment (fuel-controlled). Temperature rises are localised; escape is relatively easy. The later part of the growth stage is the last opportunity for escape."),
    ("Fully-Developed (Post-Flashover)", "All combustibles are involved. HRR is at maximum. Fire becomes ventilation-controlled — produces large CO quantities; untenable for humans. Structural damage is likely."),
    ("Decay", "Fuel is consumed; temperatures fall. Fire may change from ventilation-controlled to fuel-controlled. Decay commences when average temperature falls to 80% of peak value."),
]

for name, desc in stages:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

heading2(doc, "1.1 Flashover")
body(doc, "Flashover is the rapid transition from the Growth to the Fully-Developed stage, characterised by sudden total ignition of all combustibles in the compartment.")
body(doc, "Criteria for onset of Flashover:")
bullet(doc, "Upper gas layer temperature reaches 500–600 °C")
bullet(doc, "Radiant heat flux at floor level reaches 20 kW/m²")
bullet(doc, "Ignition of unburned gases in the upper layer")
bullet(doc, "Appearance of flames from openings")

heading2(doc, "1.2 Flashover vs Backdraught")
body(doc, "Flashover: occurs when a fire already has adequate ventilation and the upper layer reaches critical temperature.")
body(doc, "Backdraught: occurs under poor ventilation; initiated by sudden influx of air (door opening, window breaking).")

# =========================================================
# SECTION 2: PRE-FLASHOVER FIRES
# =========================================================
heading1(doc, "2. Pre-Flashover Fires")

heading2(doc, "2.1 Rate of Burning — Single Fuel Item")
body(doc, "The rate of burning (mass loss per unit area) is determined by the heat flux balance at the fuel surface:")
equation_box(doc, "Free burning", "ṁ'' = (Q̇''_F - Q̇''_L) / L_v")
equation_box(doc, "With enclosure radiation feedback", "ṁ'' = (Q̇''_F - Q̇''_L) / L_v  +  Q̇''_E / L_v")

body(doc, "Where:")
bullet(doc, "ṁ'' = rate of burning (kg/m²·s)")
bullet(doc, "Q̇''_F = heat flux from flame to surface (kW/m²)")
bullet(doc, "Q̇''_L = rate of heat loss per unit area (kW/m²)")
bullet(doc, "Q̇''_E = additional radiant flux from enclosure (kW/m²)")
bullet(doc, "L_v = latent heat of volatilization (kJ/kg)")
note(doc, "Enclosure confinement can increase burning rate ~3x and achieve maximum in ~1/3 the time vs open burning (PMMA experiments, Friedman 1975).")

heading2(doc, "2.2 Heat Release Rate (HRR)")
body(doc, "The rate of energy released by the fire:")
equation_box(doc, "HRR", "Q̇ = ṁ · χ · ΔH_c = ṁ'' · A_f · χ · ΔH_c  (kW)")

body(doc, "Where:")
bullet(doc, "ṁ = rate of burning of fuel (kg/s)")
bullet(doc, "ṁ'' = rate of burning per unit area (kg/m²·s)")
bullet(doc, "A_f = area of fuel undergoing burning (m²)")
bullet(doc, "χ = combustion efficiency (0 ≤ χ ≤ 1)")
bullet(doc, "ΔH_c = heat of combustion of fuel (kJ/kg)")

heading2(doc, "2.3 Parabolic (t²) Fire Growth")
body(doc, "Fire growth in the pre-flashover phase is often represented as a parabolic function:")
equation_box(doc, "Parabolic growth", "Q̇ = α_f · (t - t_o)²  (kW)")

body(doc, "Where:")
bullet(doc, "Q̇ = heat release rate (kW)")
bullet(doc, "α_f = fire growth coefficient (kW/s²)")
bullet(doc, "t = time (s)")
bullet(doc, "t_o = effective ignition time (s) — the x-intercept of the parabolic extrapolation")

heading2(doc, "2.4 Standard t² Fires")
body(doc, "Four standard fire growth rates (all attain ~1050 kW at their characteristic time):")

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Category"
hdr[1].text = "Equation"
hdr[2].text = "Time to 1050 kW"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ("Slow",      "Q̇ = 0.00293 t²", "600 s"),
    ("Medium",    "Q̇ = 0.01172 t²", "300 s"),
    ("Fast",      "Q̇ = 0.0469 t²",  "150 s"),
    ("Ultra-Fast","Q̇ = 0.1876 t²",   "75 s"),
]
for i, (cat, eq, time) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = cat
    row[1].text = eq
    row[2].text = time

doc.add_paragraph()

heading2(doc, "2.5 Predicting Pre-Flashover Upper Layer Temperature (MQH Method)")
body(doc, "Based on a two-zone heat balance on the upper gas layer (McCaffrey, Quintiere & Harkleroad):")
equation_box(doc, "Temperature rise", "ΔT_g = 6.85 · ( Q̇² / (h_k · A_T · A_o · √H_o) )¹ᐟ³  (K)")

body(doc, "Where:")
bullet(doc, "ΔT_g = upper gas layer temperature rise above ambient (K)")
bullet(doc, "Q̇ = heat release rate of fire (kW)")
bullet(doc, "h_k = effective heat transfer coefficient of enclosure boundaries (kW/m²·K)")
bullet(doc, "A_T = total internal surface area of enclosure (m²)")
bullet(doc, "A_o = area of ventilation opening (m²)")
bullet(doc, "H_o = height of ventilation opening (m)")
note(doc, "Constants C=1.63, N=2/3, M=-1/3 derived from 100+ experimental fires. Valid for upper layer temperatures up to ~600°C.")

heading2(doc, "2.6 Effective Heat Transfer Coefficient (h_k)")
body(doc, "Two cases depending on the characteristic burning time (t_c) vs thermal penetration time (t_p):")

body(doc, "Case 1 — Long fires (t_c > t_p): steady-state conduction through boundary:")
equation_box(doc, "h_k (steady-state)", "h_k = k / δ  (kW/m²·K)")

body(doc, "Case 2 — Short fires (t_c ≤ t_p): boundary stores heat; little lost to exterior:")
equation_box(doc, "h_k (transient)", "h_k = √(kρc / t_c)  (kW/m²·K)")

body(doc, "Where:")
bullet(doc, "k = thermal conductivity of boundary material (kW/m·K)")
bullet(doc, "δ = thickness of boundary (m)")
bullet(doc, "ρ = density of boundary material (kg/m³)")
bullet(doc, "c = specific heat of boundary material (kJ/kg·K)")
bullet(doc, "kρc = thermal inertia (kJ²/m⁴·K²·s)")
bullet(doc, "α = k/(ρc) = thermal diffusivity (m²/s)")

heading2(doc, "2.7 Thermal Penetration Time")
equation_box(doc, "t_p", "t_p = (δ/2)² · (ρc/k)  (s)")

body(doc, "For multiple boundary materials, use area-weighted average:")
equation_box(doc, "(h_k)_EFF", "(h_k)_EFF = (1/A_T) · Σ h_k,i · A_i")

heading2(doc, "2.8 Conservation of Mass")
body(doc, "Mass balance across the compartment:")
equation_box(doc, "Mass balance", "ṁ_g = ṁ_a + ṁ_f  (kg/s)")

body(doc, "Mass flow rate of air into a door/window opening (Kawagoe):")
equation_box(doc, "Air inflow", "ṁ_a = 0.5 · A_o · √H_o  (kg/s)")

note(doc, "A_o√H_o is the Ventilation Factor. Many aspects of compartment fire behaviour (pre and post flashover) vary with this term.")

# =========================================================
# SECTION 3: FLASHOVER PREDICTION
# =========================================================
heading1(doc, "3. Flashover — Minimum HRR Required")

body(doc, "Three established methods to estimate the minimum HRR required to achieve flashover in a given compartment:")

heading2(doc, "3.1 McCaffrey, Quintiere & Harkleroad (MQH) Method")
equation_box(doc, "Q̇_FO (general)", "Q̇_FO = 610 · (h_k · A_T · A_o · √H_o)¹ᐟ²  (kW)")
equation_box(doc, "Q̇_FO (simplified, t_c > t_p)", "Q̇_FO = 610 · √(k/δ · A_T · A_o · √H_o)  (kW)")
note(doc, "Based on ΔT_FO = 500°C as onset criterion. Assumes near-cubical compartment, free ventilation, well-mixed upper layer, fire at centre.")

heading2(doc, "3.2 Babrauskas Method")
equation_box(doc, "Q̇_FO", "Q̇_FO = 750 · A_o · √H_o  (kW)")
note(doc, "Derived from correlation with 33 room fires (wood & polyurethane). Finding: Q̇_FO ≈ 0.5 × Q̇_max where Q̇_max = 1500 A_o√H_o.")

heading2(doc, "3.3 Thomas Method")
equation_box(doc, "Q̇_FO", "Q̇_FO = 7.8 · A_T + 378 · A_o · √H_o  (kW)")
note(doc, "Based on ΔT_g = 600 K at flashover and c_p = 1.26 kJ/kg·K.")

heading2(doc, "3.4 Stoichiometric Maximum HRR")
equation_box(doc, "Q̇_max", "Q̇_max = 1500 · A_o · √H_o  (kW)")

body(doc, "Note: Flashover HRR is reduced when the fire is located near a wall or in a corner (enhanced flame spread to upper layer).")

# =========================================================
# SECTION 4: POST-FLASHOVER FIRES
# =========================================================
heading1(doc, "4. Post-Flashover Fires")

heading2(doc, "4.1 Ventilation-Controlled Burning")
body(doc, "In typical rooms, post-flashover fires are ventilation-controlled. The mass burning rate of wood cribs:")
equation_box(doc, "Burning rate (ventilation-controlled)", "ṁ = 0.09 · A_o · √H_o  (kg/s)  =  5.5 · A_o · √H_o  (kg/min)")

heading2(doc, "4.2 Fuel vs Ventilation Control — Distinction")
body(doc, "Using fuel area A_f and ventilation factor A_o√H_o:")
equation_box(doc, "Fuel-controlled condition", "A_T / (A_o√H_o) > 8–10 m⁻¹ᐟ²")
equation_box(doc, "Ventilation-controlled condition", "A_T / (A_o√H_o) < 8–10 m⁻¹ᐟ²")

heading2(doc, "4.3 Post-Flashover Fire Temperature Model (Buchanan)")
body(doc, "Temperature of fire gases as a function of time:")
equation_box(doc, "Fire gas temperature", "T_f(t) - T_f(0) = β · t¹ᐟ⁶")

equation_box(doc, "β constant", "β = 3 · T_f(0) · (A_o√H_o / (A_T√(kρc)))¹ᐟ³   (K·s⁻¹ᐟ⁶)")

body(doc, "Duration of post-flashover fire:")
equation_box(doc, "Fire duration", "t_D = (L · A_F) / (0.09 · A_o · √H_o)  (s)")

body(doc, "Where:")
bullet(doc, "L = fuel load (kg/m²)")
bullet(doc, "A_F = floor area (m²)")

body(doc, "Heat absorbed per unit area of fire separations/structural members:")
equation_box(doc, "Heat absorbed", "q̇''(t_D) = (3β/2) · √(kρc) · t_D²ᐟ³  (kJ/m²)")
note(doc, "ISO 834: β = 230 K·s⁻¹ᐟ⁶;  ASTM E119: β = 229 K·s⁻¹ᐟ⁶")

heading2(doc, "4.4 Equivalent Fire Resistance Duration (Harmathy & Mehaffey)")
body(doc, "Duration of a real fire equivalent to an ISO 834 standard fire test:")
equation_box(doc, "Equivalent duration", "t_eq = (β / 230)³ᐟ² · t_D")

heading2(doc, "4.5 Ventilation Factor (Opening Factor)")
equation_box(doc, "F_v", "F_v = A_v · √H_v / A_t   (m¹ᐟ²)")
body(doc, "Where A_t = total internal area of bounding surfaces (including openings).")
note(doc, "Swedish Fire Curves (Magnusson & Thelandersson, 1970) use this factor with fuel load density to produce time-temperature curves for structural design.")

heading2(doc, "4.6 Fire Spread from an Enclosure — External Flame Dimensions")
body(doc, "For flames projecting from ventilation openings:")
equation_box(doc, "Flame height above soffit", "z + H = 12.8 · (ṁ/B)²ᐟ³  (m)")
equation_box(doc, "Horizontal flame reach", "x / H = 0.454 / n⁰ᐟ⁵³")

body(doc, "Where:")
bullet(doc, "H, B = height and width of ventilation opening (m)")
bullet(doc, "z = height of flame tip above window soffit (m)")
bullet(doc, "ṁ = rate of burning (kg/s)")
bullet(doc, "x = horizontal reach of flame from building face (m)")
bullet(doc, "n = 2B/H (shape factor)")

# =========================================================
# SECTION 5: SMOKE FILLING
# =========================================================
heading1(doc, "5. Smoke Filling")

heading2(doc, "5.1 Interface Height — Empirical Correlations (Klote & Milke / ASHRAE)")
body(doc, "For a constant HRR fire (Q̇ = Q̇_o):")
equation_box(doc, "Z_i/H (constant HRR)", "Z_i/H = 1.11 - 0.28 · ln[t · Q̇_o¹ᐟ³ · H⁻⁴ᐟ³ · (S/H²)⁻¹]")

body(doc, "For a t² growth fire (Q̇ = 1000(t/t_ig)² kW):")
equation_box(doc, "Z_i/H (t² fire)", "Z_i/H = 0.91 · [t · t_ig⁻²ᐟ⁵ · H⁻⁴ᐟ⁵ · (S/H²)⁻³ᐟ⁵]⁻¹ᐟ⁴⁵")
equation_box(doc, "t_ig", "t_ig = (1000/α)¹ᐟ²  (s)")

body(doc, "Both equations valid for 0.2 ≤ Z_i/H ≤ 1.0 and 0.9 ≤ S/H² ≤ 14.0 (unobstructed plume)")

body(doc, "Where:")
bullet(doc, "Z_i = smoke layer height from floor (m)")
bullet(doc, "H = ceiling height (m)")
bullet(doc, "S = floor area (m²)")
bullet(doc, "t_ig = characteristic t² fire growth time (s)")
bullet(doc, "α = fire growth coefficient (kW/s²)")

heading2(doc, "5.2 Upper Layer Temperature (Klote & Milke)")
body(doc, "Energy balance for the upper layer (no heat loss to boundaries — upper limit):")
equation_box(doc, "Upper layer temperature", "1 - T_a/T_h = ∫ Q̇_conv dt / (ρ_a c_p T_a S(H - z_i))")
equation_box(doc, "Convective HRR", "Q̇_conv = 0.7 · Q̇")

body(doc, "For constant Q̇:    ∫ Q̇_conv dt = 0.7 · Q̇_o · t")
body(doc, "For t² fire:         ∫ Q̇_conv dt = 0.7 · (1/3) · α · t³")

heading2(doc, "5.3 Dimensionless Smoke Filling Method")
equation_box(doc, "Dimensionless ODE", "dy/dτ + Q̇* + 0.21(Q̇*)¹ᐟ³ · y⁵ᐟ³ = 0")
equation_box(doc, "Dimensionless HRR", "Q̇* = Q̇ / (ρ_a c_p T_a √g H⁵ᐟ²)")
equation_box(doc, "Dimensionless time", "τ = t · √(g/H) · H²/S")

body(doc, "Upper layer density and temperature:")
equation_box(doc, "Hot layer density", "ρ_h = ρ_a · (1 - Q̇*τ/(1-y))")
equation_box(doc, "Hot layer temperature", "T_h = 353 / ρ_h  (K)")

# =========================================================
# SECTION 6: FIRE LOAD DENSITY
# =========================================================
heading1(doc, "6. Fire Load Density")

heading2(doc, "6.1 Key Definitions and Equations")
body(doc, "Total fire load energy:")
equation_box(doc, "Total energy", "E = m · ΔH_c  (MJ)")

body(doc, "Average heat release rate:")
equation_box(doc, "Average HRR", "Q = E / t  (MW)")

body(doc, "Fire load density based on floor area (most common):")
equation_box(doc, "e_f (floor area basis)", "e_f = E / A_f  (MJ/m²)")

body(doc, "Fire load density based on total bounding surface (European standard):")
equation_box(doc, "e_t (surface area basis)", "e_t = E / A_t  (MJ/m²)")

heading2(doc, "6.2 Hazard Classification")
table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = "Fire Load Density (MJ/m²)"
hdr2[1].text = "Hazard Level"
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True
table2.rows[1].cells[0].text = "< 200"
table2.rows[1].cells[1].text = "Low Hazard"
table2.rows[2].cells[0].text = "200–400"
table2.rows[2].cells[1].text = "Moderate Hazard"
table2.rows[3].cells[0].text = "> 400"
table2.rows[3].cells[1].text = "High Hazard"
doc.add_paragraph()

heading2(doc, "6.3 Fuel Load Statistics (Buchanan)")
table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = "Occupancy"
hdr3[1].text = "Mean L (kg/m²)"
hdr3[2].text = "σ_L (kg/m²)"
for cell in hdr3:
    cell.paragraphs[0].runs[0].bold = True
occ_data = [
    ("Dwelling", "30.1", "4.4"),
    ("Office", "24.8", "8.6"),
    ("School", "17.5", "5.1"),
    ("Hospital", "25.1", "7.8"),
    ("Hotel", "14.6", "4.2"),
]
for i, (o, m, s) in enumerate(occ_data):
    row = table3.rows[i+1].cells
    row[0].text = o; row[1].text = m; row[2].text = s
doc.add_paragraph()

heading2(doc, "6.4 EN 1991-1-2 Characteristic Fire Load Densities")
table4 = doc.add_table(rows=10, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = "Occupancy"
hdr4[1].text = "Mean (MJ/m²)"
hdr4[2].text = "Std Dev (MJ/m²)"
for cell in hdr4:
    cell.paragraphs[0].runs[0].bold = True
en_data = [
    ("Dwelling", "780", "234"),
    ("Hospital (room)", "230", "69"),
    ("Hotel (room)", "310", "93"),
    ("Library", "1500", "450"),
    ("Office", "420", "126"),
    ("Classroom", "285", "85.5"),
    ("Shopping Centre", "600", "180"),
    ("Theatre/Cinema", "300", "90"),
    ("Transport (public)", "100", "30"),
]
for i, row_data in enumerate(en_data):
    row = table4.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
doc.add_paragraph()

# =========================================================
# SECTION 7: HUMAN BEHAVIOUR IN FIRES
# =========================================================
heading1(doc, "7. Human Behaviour in Fires")

heading2(doc, "7.1 Why It Matters")
body(doc, "Majority of fire deaths are attributed to inappropriate decision-making. Understanding human behaviour helps design safer buildings, better evacuation tools, and more effective fire safety management.")

heading2(doc, "7.2 Response Timeline")
body(doc, "Human response is broadly divided into two periods:")
bullet(doc, "Pre-Evacuation Period: From ignition to commencement of purposive evacuation. Includes Pre-Alarm Phase, Evacuation Decision-Making Phase, and Protective Action Phase.")
bullet(doc, "Movement Period: Time during which occupants physically move to safety.")

heading2(doc, "7.3 Discarded Myths")
body(doc, "Three myths are NOT typical of actual fire behaviour:")
bullet(doc, "Panic Behaviour: People generally behave rationally and cooperatively, not irrationally (Fahy, 2012; WTC evidence).")
bullet(doc, "Disaster Shock: Victims are more likely to show personal initiative and assist others rather than being dazed.")
bullet(doc, "Group Mind: Individuals do not lose their individual decision-making to group 'mob' behaviour.")

heading2(doc, "7.4 Five Factors Influencing Behaviour")

factors = [
    ("1. Social Influence",
     "People’s actions (including exit choice) are influenced by others. Bystander effect: in presence of passive others, only 10% report an emergency vs 75% when alone (Darley & Latane).",
     "Strategically place fire wardens throughout the building to prompt evacuation."),
    ("2. Stress",
     "Emergencies create uncertainty, information overload and time pressure. Stress narrows the field of perception — people process fewer cues.",
     "Design more noticeable signage and use luminous materials for better visibility under stress."),
    ("3. Built Environment (Movement to Familiar)",
     "People evacuate via familiar routes. Summerland 1973: 93% staff used familiar exit; 61% visitors used familiar entrance.",
     "Conduct regular fire drills so all exits become familiar."),
    ("4. Leadership",
     "Pre-event leaders influence group actions. When leadership fails, new leadership emerges.",
     "Train managers in fire safety; ensure they provide accurate instructions during emergencies."),
    ("5. Demographics (Gender)",
     "Males more involved in firefighting activities; females more likely to call emergency services and evacuate first (Bryan 1977, Wood 1972).",
     "Design evacuation procedures that account for individual and demographic variation."),
]

for factor, evidence, implication in factors:
    heading3(doc, f"Factor {factor}")
    body(doc, f"Evidence: {evidence}")
    p = doc.add_paragraph()
    r = p.add_run(f"Engineering Implication: {implication}")
    r.italic = True

heading2(doc, "7.5 11 Behavioural Facts (Kuligowski & Gwynne)")
facts = [
    "People’s first instinct is to feel (sometimes inappropriately) safe — NOT to panic.",
    "Providing information does not guarantee appropriate occupant response. Perception, attention, and comprehension of information is critical.",
    "Occupants must perceive a credible threat and personalise the risk before taking protective action.",
    "People will engage in information-seeking when cues are ambiguous or inconsistent.",
    "People are likely to engage in preparation activities (e.g. collecting belongings) before evacuation — this delays response.",
    "Generally, people act rationally and altruistically during building fires.",
    "The surrounding population influences an individual’s decision-making process.",
    "Stress narrows a person’s field of perception — individuals may miss or ignore certain cues.",
    "People move to the familiar. Prior relationships with the building and its occupants influence response.",
    "People do not instantaneously switch roles in a fire. Pre-event roles carry into the event.",
    "People are heterogeneous — individual differences in demographics can significantly influence behaviour.",
]
for i, fact in enumerate(facts):
    bullet(doc, f"Fact #{i+1}: {fact}")

# =========================================================
# SECTION 8: KEY SYMBOLS SUMMARY
# =========================================================
heading1(doc, "8. Quick Reference — Key Symbols & Units")

symbols = [
    ("Q̇", "Heat Release Rate", "kW"),
    ("ΔH_c", "Heat of Combustion", "kJ/kg"),
    ("ṁ", "Mass burning rate", "kg/s"),
    ("ṁ''", "Mass burning rate per unit area", "kg/m²·s"),
    ("χ", "Combustion efficiency", "—"),
    ("A_f", "Fuel area", "m²"),
    ("A_o / A_v", "Ventilation opening area", "m²"),
    ("H_o / H_v", "Ventilation opening height", "m"),
    ("A_T / A_t", "Total internal surface area", "m²"),
    ("A_F", "Floor area", "m²"),
    ("h_k", "Effective heat transfer coefficient", "kW/m²·K"),
    ("k", "Thermal conductivity", "kW/m·K"),
    ("ρ", "Density", "kg/m³"),
    ("c", "Specific heat", "kJ/kg·K"),
    ("kρc", "Thermal inertia", "kJ²/m⁴·K²·s"),
    ("α = k/(ρc)", "Thermal diffusivity", "m²/s"),
    ("δ", "Boundary thickness", "m"),
    ("t_p", "Thermal penetration time", "s"),
    ("t_c", "Characteristic burning time", "s"),
    ("t_D", "Duration of post-flashover fire", "s"),
    ("t_eq", "Equivalent ISO 834 fire duration", "s"),
    ("β", "Temperature-time constant", "K·s⁻¹ᐟ⁶"),
    ("L", "Specific fuel load", "kg/m²"),
    ("e_f", "Fire load density (floor area)", "MJ/m²"),
    ("F_v", "Ventilation / Opening factor", "m¹ᐟ²"),
    ("Z_i", "Smoke layer height from floor", "m"),
    ("Q̇*", "Dimensionless HRR", "—"),
    ("τ", "Dimensionless time", "—"),
    ("ΔT_g", "Upper gas layer temperature rise", "K"),
    ("ΔT_FO", "Temperature rise at flashover onset", "K (typically 500–600 K)"),
]

table5 = doc.add_table(rows=len(symbols)+1, cols=3)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
hdr5[0].text = "Symbol"
hdr5[1].text = "Meaning"
hdr5[2].text = "Units"
for cell in hdr5:
    cell.paragraphs[0].runs[0].bold = True
for i, (sym, meaning, unit) in enumerate(symbols):
    row = table5.rows[i+1].cells
    row[0].text = sym
    row[1].text = meaning
    row[2].text = unit

doc.add_paragraph()
body(doc, "End of Summary — FDM6200 Fire Dynamics Workshop 2")

# Save
out_path = r"C:\Users\JasonOng\Desktop\local docs\personal\SIT\Fire Dynamics Workshop 2 - Summary.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
