from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DIR = r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd\AST BD\HyESys Dept\3. Hardware (PCS.BATT)\v2.2 - data center\UPS"

ORANGE = RGBColor(0xE8, 0x6B, 0x1A)
DARK_GREY = RGBColor(0x3C, 0x3C, 0x3C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = color or ORANGE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or ORANGE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or DARK_GREY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        rest = p.add_run(text)
        rest.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_para(doc, text, bold=False, italic=False, size=10, color=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'E86B1A')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT 1 — ATP Series
# ─────────────────────────────────────────────────────────────────────────────
def build_atp():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title block
    add_heading(doc, 'ATP Series Intelligent Online UPS (10–200 kVA)', level=1)
    add_para(doc, 'AVIC-TECH (TOBTAK) | www.avic-tech.com.cn', italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))
    add_para(doc, 'Configuration: 3-Phase In / 3-Phase Out, Tower-Type', bold=True, size=10)
    add_divider(doc)

    # Application
    add_heading(doc, 'Application', level=2)
    add_para(doc,
             'Uninterrupted power supply for IT equipment, network communications devices, intelligent '
             'equipment, precision instruments, and automatic control systems in small-to-medium data centres.',
             size=10)

    # Key Features
    add_heading(doc, 'Key Features', level=2)
    features = [
        'DSP digital control — more stable and reliable control system',
        'Ultra-wide input voltage range — adapts to harsh grid environments',
        '30–50 battery cells selectable — intelligent charge management',
        'Rich communication interfaces and optional accessories',
        'Superior load-carrying capacity — strong load adaptability, compatible with generator operation',
        'Supports 4-unit parallel connection with shared battery banks',
        '4.3-inch colour touchscreen — simple and intuitive operation',
    ]
    for f in features:
        add_bullet(doc, f)

    add_divider(doc)

    # Specifications table
    add_heading(doc, 'Technical Specifications', level=2)
    add_para(doc, 'ATP Series Intelligent Online UPS — 3-Phase In / 3-Phase Out, Tower-Type', bold=True, size=10)

    # Build table
    models = ['ATP-10K/C', 'ATP-20K/C', 'ATP-30K/C', 'ATP-40K/C', 'ATP-60K/C',
              'ATP-80K/C', 'ATP-100K/C', 'ATP-120K/C', 'ATP-160K/C', 'ATP-200K/C']

    spec_rows = [
        # (Category, Parameter, Value)  — category='' means same as above
        ('', 'Model', ' | '.join(models)),
        ('', 'Rated Output Power (kVA)', '10 | 20 | 30 | 40 | 60 | 80 | 100 | 120 | 160 | 200'),
        ('INPUT', 'Rated Voltage', '380/400/415 VAC (3-phase 5-wire)'),
        ('', 'Frequency Range', '40–70 Hz'),
        ('', 'Power Factor', '≥ 0.99'),
        ('OUTPUT', 'Voltage', '380/400/415 VAC'),
        ('', 'Voltage Regulation', '±1%'),
        ('', 'Frequency', 'Mains mode: ±1%/±2%/±4%/±5%/±10% selectable\nBattery mode: (50/60 ± 0.1%) Hz'),
        ('', 'Power Factor', '0.9'),
        ('', 'Waveform Distortion (THDv)', '≤2% (100% linear load)   ≤5% (100% non-linear load)'),
        ('', 'Transfer Time', '0 ms'),
        ('', 'Waveform', 'Sinusoidal'),
        ('', 'Overload', '110%: 60 min   125%: 10 min   150%: 1 min   >150%: 400 ms'),
        ('BYPASS', 'Rated Voltage', '380/400/415 VAC'),
        ('', 'Sync Frequency Range', '46–54 Hz or 56–64 Hz'),
        ('BATTERY', 'Battery Voltage', '±180 ~ ±300 VDC (30–50 cells, even numbers adjustable)'),
        ('CONTROL', 'Communication (Standard)', 'RS232, RS485'),
        ('', 'Communication (Optional)', 'SNMP Card'),
        ('', 'Display', 'LCD — UPS status, load level, battery capacity, input/output voltage, discharge time, fault indicator'),
        ('ENVIRONMENT', 'Operating Temperature', '0–40°C'),
        ('', 'Relative Humidity', '0–95% (non-condensing)'),
        ('OVERALL', 'Efficiency', 'Up to 96%'),
        ('', 'Noise (1 m from front)', '<55 dB (10–40 kVA)   <58 dB (60 kVA)   <60 dB (80–100 kVA)   <65 dB (120–200 kVA)'),
        ('', 'Dimensions W×D×H (mm)',
         '10–40 kVA: 250×720×535\n60 kVA: 250×868×862\n80–200 kVA: 440×885×1200'),
        ('', 'Net Weight (kg)', '10kVA:32 | 20kVA:33 | 30kVA:35 | 40kVA:38 | 60kVA:70.5 | 80kVA:150 | 100kVA:160 | 120kVA:162 | 160kVA:196 | 200kVA:200'),
        ('', 'Compliance Standard', 'YD/T 1095-2018'),
    ]

    table = doc.add_table(rows=len(spec_rows) + 1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    col_widths = [Cm(2.8), Cm(5.5), Cm(9.5)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_labels = ['Category', 'Parameter', 'Value / Specification']
    for i, (cell, label) in enumerate(zip(hdr_cells, hdr_labels)):
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, 'E86B1A')

    # Data rows
    prev_cat = ''
    for row_idx, (cat, param, val) in enumerate(spec_rows):
        row = table.rows[row_idx + 1]
        c0, c1, c2 = row.cells

        display_cat = cat if cat != prev_cat else ''
        if cat:
            prev_cat = cat

        c0.text = display_cat
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9)
        c0.paragraphs[0].runs[0].font.color.rgb = ORANGE if display_cat else DARK_GREY
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c1.text = param
        c1.paragraphs[0].runs[0].font.size = Pt(9)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c2.text = val
        c2.paragraphs[0].runs[0].font.size = Pt(9)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bg = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)

    add_divider(doc)

    # Company info
    add_heading(doc, 'Manufacturer', level=2)
    info_lines = [
        'AVIC-TECH (XIAMEN) ELECTRIC POWER TECHNOLOGY CO., LTD.',
        'Website: www.avic-tech.com.cn',
        'Business Hotline: 18559028961',
        'Service Hotline: 400-800-1592',
        'Service Email: sale@avic-tech.com.cn',
    ]
    for line in info_lines:
        add_para(doc, line, size=9)

    # Disclaimer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(
        'Disclaimer: The company reserves the right to modify product design and specifications. '
        'All information has been carefully verified for accuracy. The company accepts no liability for '
        'printing errors or translation inaccuracies. Product specifications and appearance are subject '
        'to the actual product. For detailed specifications and operating instructions, refer to the product manual.'
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out_path = OUTPUT_DIR + r'\ATP Series Intelligent Online UPS (10-200kVA) — English.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT 2 — ARP Series
# ─────────────────────────────────────────────────────────────────────────────
def build_arp():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title
    add_heading(doc, 'ARP Series Modular UPS (30–1200 kVA)', level=1)
    add_para(doc, 'AVIC-TECH (TOBTAK) | www.avic-tech.com.cn', italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))
    add_para(doc, 'Configuration: 3-Phase In / 3-Phase Out, Modular', bold=True, size=10)
    add_divider(doc)

    # About
    add_heading(doc, 'About AVIC-TECH', level=2)
    add_para(doc,
             'AVIC-TECH (Xiamen) Electric Power Technology Co., Ltd. was established in 2006. '
             'Specialising in R&D and manufacturing of new digital power energy systems — including lithium battery '
             'integration, energy storage systems, industrial power supplies, and network energy products. '
             'Recognised as a National High-Tech Enterprise and National-Level Specialised & Innovative "Little Giant" enterprise. '
             'Operations cover all provinces in China and overseas markets.',
             size=10)

    stats = [
        ('20+ years', 'Experience in the power energy sector'),
        ('120+', 'Intellectual properties (invention patents, software copyrights)'),
        ('48,000 m²', 'Total building area — AVIC-TECH Smart Industrial Park'),
        ('100+', 'Honorary titles'),
    ]
    table_s = doc.add_table(rows=1, cols=4)
    table_s.style = 'Table Grid'
    for i, (num, label) in enumerate(stats):
        cell = table_s.rows[0].cells[i]
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(num + '\n')
        r1.bold = True
        r1.font.color.rgb = ORANGE
        r1.font.size = Pt(11)
        r2 = p1.add_run(label)
        r2.font.size = Pt(8)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, 'FFF5EE')
    doc.add_paragraph()

    add_divider(doc)

    # Product intro
    add_heading(doc, 'Product Overview', level=2)
    add_para(doc,
             'The new-generation ARP Modular UPS is a smart power solution designed for mission-critical applications. '
             'It integrates native lithium battery management, 100% unbalanced load capability, Super ECO and reserve-integrated '
             'mode, and optionally uses 3rd-generation SiC semiconductor components — delivering ultimate efficiency and reliability. '
             'Full-redundancy design, key-component lifespan warning, and intelligent O&M ensure high availability in demanding '
             'environments such as data centres, finance, and government applications.',
             size=10)

    # Four core features
    add_heading(doc, 'Four Core Features', level=2)
    core_features = [
        ('1. Native Lithium Battery',
         'Independent charge/discharge circuit, naturally neutral-free, perfectly adapted for lithium batteries — '
         'dual improvement in lifespan and efficiency.'),
        ('2. Unlimited Load Capacity',
         '100% unbalanced load capability — seamlessly adapts to AI loads with no application restrictions.'),
        ('3. SiC Materials',
         'Optional 3rd-generation silicon carbide semiconductor components deliver higher power density and ultra-low losses.'),
        ('4. Integrated Intelligence',
         'Combines Super ECO mode with reserve functions to significantly reduce operating costs and maximise economic efficiency.'),
    ]
    for title, desc in core_features:
        p = doc.add_paragraph()
        r1 = p.add_run(title + '  ')
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = ORANGE
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    add_divider(doc)

    # Additional highlights
    add_heading(doc, 'Additional Smart Highlights', level=2)

    highlights = [
        ('1. Ultimate Energy Efficiency — Green & Low-Carbon', [
            'High efficiency: System efficiency up to 96.5%, significantly reducing energy losses.',
            'Intelligent sleep: Modules wake on demand, further optimising light-load efficiency.',
        ]),
        ('2. Full-Chain Reliability — Intelligent Early Warning', [
            'Full-redundancy design: No single point of failure in key components, ensuring continuous power.',
            'Lifespan warning: Precise prediction of key component lifespan, preventing failures before they occur.',
        ]),
        ('3. Smart O&M — Easy & Convenient', [
            'Online maintenance: Modules support hot-swap, enabling true 7×24-hour maintenance without power interruption.',
            'Fault diagnosis: One-click fault recording with second-level issue localisation.',
            'Self-aging test: System testing without renting dummy loads — saves time and cost.',
        ]),
    ]
    for h_title, bullets in highlights:
        add_heading(doc, h_title, level=3)
        for b in bullets:
            add_bullet(doc, b)

    add_divider(doc)

    # Feature showcase
    add_heading(doc, 'Feature Showcase', level=2)

    showcases = [
        ('1. Standard Neutral-Free / Optional Neutral-Wire', [
            'Flexible configuration: Within 30–50 cell range, supports individual cell independent adjustment.',
            'New projects: Flexible battery cell configuration — reduces one power cable and one 1P circuit breaker.',
            'Retrofit projects: Fully utilises existing batteries with no additional modification required.',
            'Battery fault: Quickly removes faulty batteries to ensure reliable system operation.',
        ]),
        ('2. Native Lithium Battery Adaptation', [
            'Direct lithium connection with fine-grained management: Supports lithium-specific settings with BMS communication '
            '— smarter and safer battery management.',
        ]),
        ('3. Reserve-Integrated Function', [
            'High-speed charging: Single module 20A/30A high-current charging, meeting rapid energy storage recharge demands.',
            'Value: Uninterrupted supply (emergency backup), peak shaving / valley filling (cost optimisation), easy capacity expansion.',
            'Planned curve mode: Set schedules based on electricity tariffs — suitable for simple energy dispatch.',
            'EMS dispatch mode: External EMS commands for battery management — suitable for complex energy dispatch.',
        ]),
        ('4. Super ECO Function', [
            'Bypass main supply with harmonic management: Supports bypass main-supply mode; inverter provides APF/SVG active '
            'harmonic current compensation — ensures power quality while reducing energy losses.',
        ]),
        ('5. AI Load Adaptation', [
            'Purpose-built for AI: Supports 100% extreme load imbalance across racks and servers.',
            'Uses three-phase four-arm (balanced half-bridge) circuit — tailored for AI computing environments.',
        ]),
        ('6. SiC Component Application', [
            'Higher efficiency, better density: 3rd-generation SiC semiconductor components deliver higher conversion efficiency '
            'and power density — dual improvement in efficiency and single-unit power rating.',
        ]),
    ]
    for s_title, bullets in showcases:
        add_heading(doc, s_title, level=3)
        for b in bullets:
            add_bullet(doc, b)

    add_divider(doc)

    # ── SPEC TABLE 1: 30–150 kVA ──────────────────────────────────────────────
    add_heading(doc, 'Technical Specifications — ARP Series 30–150 kVA', level=2)

    spec_30_150 = [
        # (Section, Parameter, Values...)  — 4 models
        ('INPUT', 'Rated Input Voltage', '380/400/415 Vac (3-phase 4-wire + PE)', '', '', ''),
        ('', 'Input Voltage Range', '138–485 Vac (138–305 Vac: linear derating; 305–485 Vac: no derating)', '', '', ''),
        ('', 'Input Frequency Range', '40–70 Hz', '', '', ''),
        ('', 'Input Power Factor', '≥ 0.99', '', '', ''),
        ('', 'Input THDi', '≤3% (100% linear load)', '', '', ''),
        ('OUTPUT', 'Rated Output Voltage', '380/400/415 Vac ±1% (3-phase 4-wire + PE)', '', '', ''),
        ('', 'Output Frequency', 'Mains: tracks bypass   Battery: 50/60 Hz ±0.1%', '', '', ''),
        ('', 'Current Crest Ratio', '3:1', '', '', ''),
        ('', 'Overload Capacity',
         '110%: bypass after 60 min   125%: bypass after 10 min   150%: bypass after 60 s', '', '', ''),
        ('', 'Output THDv', '≤1% (linear load)   ≤3% (non-linear load)', '', '', ''),
        ('BATTERY', 'Rated Battery Voltage',
         '300–600 Vdc (30–50 cells, default 32 cells)   409.6 Vdc (lithium)', '', '', ''),
        ('', 'Max Module Charge Current', '20 A', '', '', ''),
        ('', 'Battery Type', 'LiFePO₄ / Lead-Acid', '', '', ''),
        ('', 'Shared Battery Bank', 'Supported', '', '', ''),
        ('COMMS', 'Standard', 'RS485 / Dry Contact', '', '', ''),
        ('', 'Optional', 'SNMP Card', '', '', ''),
        ('SYSTEM', 'Output Power Factor', '1', '', '', ''),
        ('', 'System Efficiency', 'Up to 96.5%', '', '', ''),
        ('', 'Parallel Capacity', '4 units', '', '', ''),
        ('', 'Alarm Functions', 'Input abnormality, battery low voltage, overload, fault, etc.', '', '', ''),
        ('', 'Protection Functions',
         'Output short circuit, over/under voltage, overload, over-temperature, battery under-voltage', '', '', ''),
        ('ENVIRONMENT', 'Operating Temperature', '0–40°C', '', '', ''),
        ('', 'Relative Humidity', '0–95% (non-condensing)', '', '', ''),
        ('', 'Altitude', '1000 m; above 1000 m derate per GB/T3859.2-2013', '', '', ''),
        ('OVERALL', 'Noise', '<70 dB', '', '', ''),
        ('', 'Switch Configuration', 'Standard: Main, Bypass, Output, Maintenance (4 switches)', '', '', ''),
        ('', 'Module Dimensions W×D×H', '440×640×86 mm (2U)', '', '', ''),
        ('', 'Cabinet Dimensions W×D×H', '600×850×1200 mm', '', '', ''),
        ('', 'Module Net Weight', '21 kg', '', '', ''),
        ('', 'Cabinet Net Weight', '200 kg (ARP-60/90/120) | 250 kg (ARP-150)', '', '', ''),
        ('', 'Protection Class', 'IP20', '', '', ''),
        ('', 'Cable Entry/Exit', 'Bottom-in, Bottom-out', '', '', ''),
        ('', 'Standard', 'YD/T 2165-2017', '', '', ''),
    ]

    models_30 = ['ARP-60-30/2\n(60 kVA)', 'ARP-90-30/3\n(90 kVA)', 'ARP-120-30/4\n(120 kVA)', 'ARP-150-30/5\n(150 kVA)']

    _build_spec_table(doc, models_30, spec_30_150, col_widths=[Cm(2.5), Cm(5.0), Cm(10.3)])

    doc.add_paragraph()

    # ── SPEC TABLE 2: 50–600 kVA ──────────────────────────────────────────────
    add_heading(doc, 'Technical Specifications — ARP Series 50–600 kVA', level=2)

    spec_50_600 = [
        ('INPUT', 'Rated Input Voltage', '380/400/415 Vac (3-phase 4-wire + PE)'),
        ('', 'Input Voltage Range', '138–485 Vac (138–305 Vac: linear derating; 305–485 Vac: no derating)'),
        ('', 'Input Frequency Range', '40–70 Hz'),
        ('', 'Input Power Factor', '≥ 0.99'),
        ('', 'Input THDi', '≤3% (100% linear load)'),
        ('OUTPUT', 'Rated Output Voltage', '380/400/415 Vac ±1% (3-phase 4-wire + PE)'),
        ('', 'Output Frequency', 'Mains: tracks bypass   Battery: 50/60 Hz ±0.1%'),
        ('', 'Current Crest Ratio', '3:1'),
        ('', 'Overload Capacity', '110%: bypass after 60 min   125%: after 10 min   150%: after 60 s'),
        ('', 'Output THDv', '≤1% (linear load)   ≤3% (non-linear load)'),
        ('BATTERY', 'Rated Battery Voltage', '300–600 Vdc (30–50 cells, default 40 cells)   512 Vdc (lithium)'),
        ('', 'Max Module Charge Current', '30 A'),
        ('', 'Battery Type', 'LiFePO₄ / Lead-Acid'),
        ('', 'Shared Battery Bank', 'Supported'),
        ('COMMS', 'Standard', 'RS485 / Dry Contact'),
        ('', 'Optional', 'SNMP Card'),
        ('SYSTEM', 'Output Power Factor', '1'),
        ('', 'System Efficiency', 'Up to 96.5%'),
        ('', 'Parallel Capacity', '4 units'),
        ('', 'Alarm Functions', 'Input abnormality, battery low voltage, overload, fault, etc.'),
        ('', 'Protection Functions', 'Output short circuit, over/under voltage, overload, over-temperature, battery under-voltage'),
        ('ENVIRONMENT', 'Operating Temperature', '0–40°C'),
        ('', 'Relative Humidity', '0–95% (non-condensing)'),
        ('', 'Altitude', '1000 m; above 1000 m derate per GB/T3859.2-2013'),
        ('OVERALL', 'Noise', '<70 dB'),
        ('', 'Switch Configuration', 'Standard: Main, Bypass, Output, Maintenance (4 switches)'),
        ('', 'Module Dimensions W×D×H', '440×640×130 mm (3U)'),
        ('', 'Cabinet Dimensions W×D×H',
         '200–300 kVA: 600×850×2000 mm\n400–600 kVA: 1200×850×2000 mm'),
        ('', 'Module Net Weight', '32 kg'),
        ('', 'Cabinet Net Weight', '200–300 kVA: 400 kg   400–600 kVA: 650 kg'),
        ('', 'Protection Class', 'IP20'),
        ('', 'Cable Entry/Exit', 'Bottom-in/bottom-out or Top-in/top-out (selectable); top/bottom compatible'),
        ('', 'Standard', 'YD/T 2165-2017'),
    ]

    models_50 = ['ARP-200-50/4\n(200 kVA)', 'ARP-300-50/6\n(300 kVA)', 'ARP-400-50/8\n(400 kVA)',
                 'ARP-500-50/10\n(500 kVA)', 'ARP-600-50/12\n(600 kVA)']

    _build_spec_table_simple(doc, spec_50_600)

    doc.add_paragraph()

    # ── SPEC TABLE 3: 100–1200 kVA ────────────────────────────────────────────
    add_heading(doc, 'Technical Specifications — ARP Series 100–1200 kVA', level=2)

    spec_100_1200 = [
        ('RATED CAPACITY', 'Module Capacity', '100 kVA', '100 kVA'),
        ('', 'Max Modules', '6', '12'),
        ('MAIN INPUT', 'Rated Input Voltage', '380/400/415 Vac', '380/400/415 Vac'),
        ('', 'Input Voltage Range', '138–485 Vac', '138–485 Vac'),
        ('', 'Wiring', '3-phase 5-wire', '3-phase 5-wire'),
        ('', 'Frequency Range', '40–70 Hz', '40–70 Hz'),
        ('', 'Power Factor', '≥ 0.99', '≥ 0.99'),
        ('', 'THDi', '≤3% (100% linear load)', '≤3% (100% linear load)'),
        ('BYPASS INPUT', 'Rated Voltage', '380/400/415 Vac', '380/400/415 Vac'),
        ('', 'Wiring', '3-phase 5-wire', '3-phase 5-wire'),
        ('', 'Sync Tracking Range', '±10%', '±10%'),
        ('OUTPUT', 'Voltage', '380/400/415 Vac ±1%', '380/400/415 Vac ±1%'),
        ('', 'Power Factor', '1', '1'),
        ('', 'Frequency (Mains mode)', '±1%/±2%/±4%/±5%/±10% selectable', '±1%/±2%/±4%/±5%/±10% selectable'),
        ('', 'Frequency (Battery mode)', '50/60 ±0.1%', '50/60 ±0.1%'),
        ('', 'Waveform', 'Sinusoidal', 'Sinusoidal'),
        ('', 'Current Crest Ratio', '3:1', '3:1'),
        ('', 'THDv', '≤1% (linear)   ≤3% (non-linear)', '≤1% (linear)   ≤3% (non-linear)'),
        ('', 'Transfer Time', '0 ms', '0 ms'),
        ('', 'System Efficiency', '97%', '97%'),
        ('', 'Overload', '110%: 60 min   125%: 10 min   150%: 1 min (then bypass)',
         '110%: 60 min   125%: 10 min   150%: 1 min (then bypass)'),
        ('BATTERY', 'Max Charge Current', '30 A', '30 A'),
        ('', 'Battery Voltage', '360–600 Vdc (30–50 cells, neutral-free)', '360–600 Vdc (30–50 cells, neutral-free)'),
        ('ENVIRONMENT', 'Operating Temperature', '0°C–40°C', '0°C–40°C'),
        ('', 'Storage Temperature', '-25°C–55°C (excluding batteries)', '-25°C–55°C (excluding batteries)'),
        ('', 'Humidity', '0–95% (non-condensing)', '0–95% (non-condensing)'),
        ('', 'Operating Altitude', '<1500 m; above 1500 m derate per GB/T3859.2', '<1500 m; above 1500 m derate per GB/T3859.2'),
        ('FUNCTIONS', 'Alarms', 'Overload, mains abnormality, UPS fault, battery under-voltage, etc.',
         'Overload, mains abnormality, UPS fault, battery under-voltage, etc.'),
        ('', 'Protection',
         'Short circuit, overload, over-temp, battery under-voltage, output over/under voltage, fan fault, lightning protection, bypass backfeed',
         'Short circuit, overload, over-temp, battery under-voltage, output over/under voltage, fan fault, lightning protection, bypass backfeed'),
        ('', 'Communication',
         'CAN, RS485, network interface, dry contact, parallel interface, LBS interface, smart slot, temperature sensor interface',
         'CAN, RS485, network interface, dry contact, parallel interface, LBS interface, smart slot, temperature sensor interface'),
        ('DIMENSIONS', 'Cabinet (W×D×H)', '800×1000×2000 mm', '1600×1000×2000 mm'),
        ('', 'Module (W×D×H)', '750×440×130 mm', '750×440×130 mm'),
        ('WEIGHT', 'Cabinet Net Weight', '410 kg', '840 kg'),
        ('', 'Module Net Weight', '50 kg', '50 kg'),
        ('', 'Standard', 'YD/T 2165-2017', 'YD/T 2165-2017'),
    ]

    tbl = doc.add_table(rows=len(spec_100_1200) + 1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cw = [Cm(2.8), Cm(4.8), Cm(5.5), Cm(5.5)]
    for i, col in enumerate(tbl.columns):
        for cell in col.cells:
            cell.width = cw[i]

    hdr = tbl.rows[0].cells
    for i, label in enumerate(['Category', 'Parameter', 'ARP-600-100/6 (600 kVA)', 'ARP-1200-100/12 (1200 kVA)']):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], 'E86B1A')

    prev_cat = ''
    for row_idx, row_data in enumerate(spec_100_1200):
        cat, param, v1, v2 = row_data
        row = tbl.rows[row_idx + 1]
        c0, c1, c2, c3 = row.cells

        display_cat = cat if cat != prev_cat else ''
        if cat:
            prev_cat = cat

        c0.text = display_cat
        if c0.paragraphs[0].runs:
            c0.paragraphs[0].runs[0].bold = True
            c0.paragraphs[0].runs[0].font.size = Pt(9)
            c0.paragraphs[0].runs[0].font.color.rgb = ORANGE if display_cat else DARK_GREY
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for cell, val in [(c1, param), (c2, v1), (c3, v2)]:
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bg = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'
        for cell in [c0, c1, c2, c3]:
            set_cell_bg(cell, bg)

    add_divider(doc)

    # Applications
    add_heading(doc, 'Application Sectors', level=2)
    sectors = ['Data Centres', 'Government', 'Healthcare', 'Education', 'Enterprise', 'Large Internet Data Centres']
    for s in sectors:
        add_bullet(doc, s)

    add_divider(doc)

    # Company info
    add_heading(doc, 'Manufacturer', level=2)
    info_lines = [
        'AVIC-TECH (XIAMEN) ELECTRIC POWER TECHNOLOGY CO., LTD.',
        'Industrial Park: No. 26 Xinle Road, Xinyang Industrial Zone, Haicang District, Xiamen, Fujian, China',
        'Business Hotline: 18559028961',
        'Service Hotline: 400-800-1592',
        'Website: www.avic-tech.com.cn',
        'Service Email: sale@avic-tech.com.cn',
    ]
    for line in info_lines:
        add_para(doc, line, size=9)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(
        'Disclaimer: The company reserves the right to modify product design and specifications. '
        'All information has been carefully verified for accuracy. The company accepts no liability for '
        'printing errors or translation inaccuracies. Product specifications and appearance are subject '
        'to the actual product. For detailed specifications and operating instructions, refer to the product manual. '
        'Version: TOBTAK-ARP-CN-20260109'
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out_path = OUTPUT_DIR + r'\ARP Series Modular UPS (30-1200kVA) — English.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


def _build_spec_table(doc, models, spec_rows, col_widths):
    """Single-value spec table (all models share same value)."""
    tbl = doc.add_table(rows=len(spec_rows) + 1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, col in enumerate(tbl.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    hdr = tbl.rows[0].cells
    for i, label in enumerate(['Category', 'Parameter', 'Specification (all models)']):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], 'E86B1A')

    prev_cat = ''
    for row_idx, row_data in enumerate(spec_rows):
        cat, param = row_data[0], row_data[1]
        val = row_data[2]
        row = tbl.rows[row_idx + 1]
        c0, c1, c2 = row.cells

        display_cat = cat if cat != prev_cat else ''
        if cat:
            prev_cat = cat

        c0.text = display_cat
        if c0.paragraphs[0].runs:
            c0.paragraphs[0].runs[0].bold = True
            c0.paragraphs[0].runs[0].font.size = Pt(9)
            c0.paragraphs[0].runs[0].font.color.rgb = ORANGE if display_cat else DARK_GREY
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c1.text = param
        if c1.paragraphs[0].runs:
            c1.paragraphs[0].runs[0].font.size = Pt(9)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c2.text = val
        if c2.paragraphs[0].runs:
            c2.paragraphs[0].runs[0].font.size = Pt(9)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bg = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'
        for cell in [c0, c1, c2]:
            set_cell_bg(cell, bg)


def _build_spec_table_simple(doc, spec_rows):
    """Simple 3-col table."""
    _build_spec_table(doc, [], spec_rows, col_widths=[Cm(2.5), Cm(5.0), Cm(10.3)])


if __name__ == '__main__':
    atp_path = build_atp()
    arp_path = build_arp()
    print('\nDone. Both documents saved.')
