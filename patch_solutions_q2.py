"""
Patch FEE2026_Mechanical_Solutions_v2.docx — Question 2 clarifications.

Changes:
  1. After Step 1 calculation (para 87): insert note explaining A = y-intercept, B = slope.
  2. Fix typo in para 91: remove stray trailing '2' from +150×10⁻⁶2.
  3. After Step 3 result (para 92): insert note explaining WHY centroid isolates ε_axial.
  4. Reword para 96 Note to be more definitive that 70 MPa is correct.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import copy

SRC = r"C:\Users\JasonOng\Desktop\local docs\personal\PE\FEE2026_Mechanical_Solutions_v2.docx"
DST = SRC  # overwrite in place

doc = Document(SRC)


# ── helper: insert a styled note paragraph after a reference paragraph ──────
def insert_note_after(ref_para, label, body):
    """
    Insert a new paragraph immediately after ref_para.
    label  — short bold prefix (e.g. '  ↳ Why A vs B:')
    body   — italic explanation text
    """
    new_p_xml = OxmlElement('w:p')
    ref_para._p.addnext(new_p_xml)

    # Copy paragraph properties (indentation, spacing) from ref_para
    if ref_para._p.pPr is not None:
        pPr = copy.deepcopy(ref_para._p.pPr)
        new_p_xml.insert(0, pPr)

    # Bold label run
    r_lbl = OxmlElement('w:r')
    rPr_lbl = OxmlElement('w:rPr')
    b = OxmlElement('w:b'); rPr_lbl.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr_lbl.append(sz)  # 9 pt
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1A1A60'); rPr_lbl.append(color)
    r_lbl.append(rPr_lbl)
    t_lbl = OxmlElement('w:t')
    t_lbl.text = label
    t_lbl.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_lbl.append(t_lbl)
    new_p_xml.append(r_lbl)

    # Italic body run
    r_body = OxmlElement('w:r')
    rPr_body = OxmlElement('w:rPr')
    i = OxmlElement('w:i'); rPr_body.append(i)
    sz2 = OxmlElement('w:sz'); sz2.set(qn('w:val'), '18'); rPr_body.append(sz2)
    color2 = OxmlElement('w:color'); color2.set(qn('w:val'), '1A1A60'); rPr_body.append(color2)
    r_body.append(rPr_body)
    t_body = OxmlElement('w:t')
    t_body.text = body
    t_body.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_body.append(t_body)
    new_p_xml.append(r_body)


# ── 1. Insert Step 1 clarification after para 87 ───────────────────────────
para87 = doc.paragraphs[87]
assert '250B' in para87.text or 'B = 700' in para87.text, f"Unexpected para 87: {para87.text!r}"

insert_note_after(
    para87,
    label='  ↳ Why A needs no multiplication, but B does:  ',
    body=(
        'ε(y) = A + B·y follows Euler-Bernoulli theory — plane sections remain plane, '
        'so strain must vary linearly with depth. '
        'A is the y-intercept: the strain value when y = 0 (top surface). '
        'Substituting y = 0 makes the B·y term vanish, so A = ε_top directly. '
        'B is the slope (strain gradient per mm). '
        'Substituting y = 250 mm gives −200 + 250B = 500, from which B is solved. '
        'The coordinate origin (y = 0 at top) was chosen deliberately to make Step 1 clean.'
    )
)

# ── 2. Fix typo in para 91: trailing stray '2' on +150×10⁻⁶2 ──────────────
# After the insertion above, para indices from 88 onward shift by 1.
# para 91 (original) is now at index 92.
para91_new_idx = 92
para91 = doc.paragraphs[para91_new_idx]
assert 'ε_axial' in para91.text, f"Unexpected para at new idx {para91_new_idx}: {para91.text!r}"

for run in para91.runs:
    if '10⁻⁶' + '2' in run.text:   # ×10⁻⁶ followed by stray 2
        run.text = run.text.replace('10⁻⁶' + '2', '10⁻⁶')
    elif run.text.endswith('2') and '10' in run.text:
        run.text = run.text[:-1]

# ── 3. Insert Step 3 clarification after para 92 (original), now at 93 ─────
para92_new_idx = 93
para92 = doc.paragraphs[para92_new_idx]
assert 'σ_axial' in para92.text and '30 MPa' in para92.text, \
    f"Unexpected para at new idx {para92_new_idx}: {para92.text!r}"

insert_note_after(
    para92,
    label='  ↳ Why evaluate at the centroid?  ',
    body=(
        'By superposition: ε_total(y) = ε_axial + ε_bending(y). '
        'The bending strain formula is ε_bending(y) = −M·(y − y_c)/(E·I). '
        'At y = y_c (centroid): (y − y_c) = 0, so ε_bending = 0 exactly. '
        'Therefore ε_total at the centroid = ε_axial + 0 = ε_axial — '
        'the centroid is the only point where the bending component drops out, '
        'leaving ε_axial directly readable from the linear fit. '
        'σ_axial = 30 MPa is uniform across the entire cross-section, not just at the centroid.'
    )
)

# ── 4. Reword the Note paragraph (originally para 96, now at 98) ────────────
para96_new_idx = 98
para96 = doc.paragraphs[para96_new_idx]
assert 'typographical' in para96.text or '200 MPa' in para96.text, \
    f"Unexpected para at new idx {para96_new_idx}: {para96.text!r}"

# Clear all runs and rewrite
for run in para96.runs:
    run.text = ''

# Rebuild as two runs: bold label + normal body
para96.runs[0].text = 'Correction — Bending Stress:  '
para96.runs[0].bold = True
para96.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# Add a second run for the body if needed
if len(para96.runs) > 1:
    para96.runs[1].text = (
        '70 MPa is the correct answer (verified above). '
        'The answer choices listing 200 MPa for bending stress is a typographical error in the question paper. '
        'Correct answer: (a) — NA at 71.43 mm from top, σ_axial = 30 MPa (tensile), σ_bend = 70 MPa.'
    )
    para96.runs[1].bold = False
    para96.runs[1].italic = False
else:
    from docx.shared import RGBColor as RC
    new_run = para96.add_run(
        '70 MPa is the correct answer (verified above). '
        'The answer choices listing 200 MPa for bending stress is a typographical error in the question paper. '
        'Correct answer: (a) — NA at 71.43 mm from top, σ_axial = 30 MPa (tensile), σ_bend = 70 MPa.'
    )
    new_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(DST)
print("Saved:", DST)

# Verify — print the patched region
print("\nVerification — paragraphs around Steps 1, 3, 4:")
doc2 = Document(DST)
for i in range(85, 102):
    print(f"  [{i}] {doc2.paragraphs[i].text[:120]!r}")
