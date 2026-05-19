"""
Translate Baoyuan HyESys Savings Report — English to Simplified Chinese.
All translations are hardcoded (no API key required).
Applies Microsoft YaHei font throughout for clean, professional Chinese output.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\2024 HyESys\Projects\China\葆元 baoyuan\report"
    r"\设备报告Baoyuan_HyESys_Savings_Report_2026-05-12.docx"
)
DST = (
    r"C:\Users\JasonOng\OneDrive - Advancer Global Ltd"
    r"\AST BD\2024 HyESys\Projects\China\葆元 baoyuan\report"
    r"\葆元_HyESys节能分析报告_2026-05-12_CN.docx"
)

FONT = "Microsoft YaHei"

# ── Translation map ────────────────────────────────────────────────────────────
# Key: stripped original English text (or substring that uniquely identifies it)
# Value: full Chinese replacement text
PARA_MAP = {
    "HyESys Energy Savings Analysis Report":
        "HyESys 节能分析报告",

    "Baoyuan Industrial (诸暨市葆元实业有限公司)":
        "葆元实业（诸暨市葆元实业有限公司）",

    "Prepared by: AST — Advancer Smart Technology Pte Ltd\nDate: 12 May 2026\nAnalysis Period: 20 March – 6 May 2026\nHyESys Unit: H125 (125 kVAr / 200 kWh)":
        "编制单位：AST — Advancer Smart Technology Pte Ltd\n日期：2026年5月12日\n分析周期：2026年3月20日 – 5月6日\nHyESys 设备型号：H125（125 kVAr / 200 kWh）",

    "1.  Executive Summary":
        "1.  执行摘要",

    "This report presents the quantified energy savings delivered by the HyESys H125 active digital power compensator installed at the Baoyuan MSB1 main switchboard. The analysis is based on 47 days of 15-minute interval data from the MSB1 incomer meter (20 March – 6 May 2026), cross-validated against the utility grid meter (April – May 2026), using the known HyESys activation schedule provided by the site.":
        "本报告呈现了安装于葆元 MSB1 总配电柜的 HyESys H125 有源数字功率补偿器所实现的量化节能成果。分析基于 MSB1 进线电表 47 天的 15 分钟间隔数据（2026年3月20日 – 5月6日），并与电网公用电表（2026年4月 – 5月）进行交叉验证，采用现场提供的 HyESys 已知启动时间表。",

    "A transition-event methodology was applied — comparing electrical parameters in the ±2-hour window around each HyESys ON event — to isolate the HyESys effect from background production-schedule variation.":
        "采用过渡事件分析法——对每次 HyESys 开启事件前后 ±2 小时窗口内的电气参数进行对比——以将 HyESys 效果与生产计划波动的背景变化相隔离。",

    "2.  Site Load Profile":
        "2.  现场负荷特性",

    "The MSB1 incomer meter records 15-minute interval data covering all loads within MSB1, including the HyESys unit itself, but excluding a separate capacitor bank which compensates reactive power at the high-voltage (10 kV) side. The baseline statistics below are drawn from 1–13 April 2026, when HyESys was confirmed OFF. ":
        "MSB1 进线电表记录 15 分钟间隔数据，涵盖 MSB1 内所有负载（含 HyESys 设备本身），但不含另设的电容器组（该电容器组在高压侧 10 kV 进行无功补偿）。以下基线统计数据取自 2026年4月1日至13日 HyESys 确认关闭期间。",

    "Note: The large gap between displacement PF (0.825) and true apparent PF (0.57–0.60) on phases A and C indicates very high harmonic current content — primarily from non-linear loads such as VFDs. HyESys, as an active digital compensator, addresses both reactive and harmonic currents simultaneously. This is a key reason why the observed current reduction exceeds what pure reactive compensation alone would predict.":
        "注：A 相和 C 相的位移功率因数（0.825）与真实视在功率因数（0.57–0.60）之间存在较大差距，表明谐波电流含量极高——主要来自变频器（VFD）等非线性负载。HyESys 作为有源数字补偿器，可同时处理无功电流和谐波电流。这也是实测电流降幅超出纯无功补偿理论预测值的主要原因。",

    "Figure 1 — Per-phase displacement PF vs true apparent PF. The gap between bars represents harmonic current. THD of ~142% on Phase C is abnormally high.":
        "图1 — 各相位移功率因数与真实视在功率因数对比。柱形之间的差距代表谐波电流。C 相 THD 约 142%，异常偏高。",

    "3.  Transition Event Results":
        "3.  过渡事件结果",

    "Three OFF→ON activation events were identified within the analysis period. For each event, electrical parameters were averaged over the 2-hour window immediately before and after activation.":
        "在分析周期内共识别出三次 OFF→ON 启动事件。针对每次事件，对启动前后 2 小时窗口内的电气参数取均值进行对比。",

    "Figure 2 — I_rms_3ph in the ±3-hour window around each activation. The green dashed line marks the HyESys ON moment. Current drops immediately and consistently across all three events.":
        "图2 — 每次启动事件前后 ±3 小时窗口内的 I_rms_3ph 变化。绿色虚线标记 HyESys 开启时刻。三次事件中电流均即时稳定下降。",

    "4  Grid Meter limitation":
        "4.  电网电表局限性",

    "The utility grid meter (covering MSB1 + MSB2 combined, at 10 kV) was cross-checked at the same transition events. Because MSB1 and MSB2 load is not separately metered, its variation dilutes the data, but both available events show consistent directional improvement:":
        "电网公用电表（覆盖 MSB1 + MSB2 合并，10 kV 侧）在相同过渡事件处进行了交叉验证。由于 MSB1 和 MSB2 负荷未分开计量，其变化对数据产生稀释效应，但两次可用事件均显示出一致的改善方向：",

    "5.  Electrical Parameter Overview":
        "5.  电气参数总览",

    "The chart below shows I_rms, reactive power, and active power at the MSB1 incomer over the full analysis period. Green shaded bands mark HyESys ON periods. While the ON/OFF effect is not visible to the naked eye at this scale (due to production load variation), the transition-event analysis in Section 3 isolates the effect precisely.":
        "下图展示了整个分析周期内 MSB1 进线处的 I_rms、无功功率和有功功率变化。绿色阴影区域标记 HyESys 开启时段。由于生产负荷波动，开/关效果在此比例下肉眼不可见，但第3节的过渡事件分析能精确量化该效果。",

    "Figure 3 — MSB1 incomer parameters from 1 April to 6 May 2026. Green bands = HyESys ON. Grey areas = HyESys OFF.":
        "图3 — 2026年4月1日至5月6日 MSB1 进线参数变化。绿色区域 = HyESys 开启，灰色区域 = HyESys 关闭。",

    "6.  Savings Quantification":
        "6.  节能量化",

    "Two savings estimates are provided. The conservative figure (direct observed kW delta) uses what was directly measured at the MSB1 meter during transition events. The upper bound (I²R fraction method) applies the theoretical loss reduction fraction to the full baseline kW, which may include cable losses not fully captured at the incomer measurement point.":
        "本报告提供两种节能估算值。保守值（直接观测 kW 变化量）采用过渡事件期间 MSB1 电表的直接测量数据。上限值（I²R 比例法）将理论损耗削减比例应用于全部基线 kW，可能涵盖进线测量点未能完全捕获的电缆损耗。",

    "Figure 4 — Left: I²R loss reduction fraction per activation event. Right: Estimated annual savings range (CNY).":
        "图4 — 左：每次启动事件的 I²R 损耗削减比例。右：预计年节省范围（CNY）。",

    "★  Recommended figure for customer reporting: CNY 257,000/year (conservative). This is directly observable from the meter data and is fully defensible.":
        "★  建议向客户报告的数据：CNY 257,000 /年（保守估算）。该数值可直接从电表数据观测得出，具有充分的可靠依据。",

    "7.  Issues and Recommendations":
        "7.  问题与建议",

    "1.  Meter limitation at MSB1 and MSB2 level.":
        "1.  MSB1 和 MSB2 级别的电表局限性",

    "As HyESys is located at feeder 2 of Cabinet A, the scale of effect is greatly diluted. Due to limitations of not being able to read data from MSB1 & 2 level, the only way to compute the effect is at the Main Grid level, which is diluted by the presence of MSB2 and the neighboring MSB1 Cabinets. Recommendation to install meter at the MSB 1 & 2 level to read data for a clearer computation. ":
        "由于 HyESys 安装于 A 柜馈线2，效果被大幅稀释。因无法读取 MSB1 和 MSB2 级别数据，只能在总电网级别进行计算，而该级别因 MSB2 及邻近 MSB1 柜的存在而被进一步稀释。建议在 MSB1 和 MSB2 级别安装专用电表，以获取更清晰的计算数据。",

    "2.  Cotton trapped in the system  ":
        "2.  系统内棉絮堵塞",

    "We suspected Cotton has filled the internal parts of the system. Clogging the system electrical components. When system was running, the system stopped shortly due to overheating, this may be due to cotton being trapped at the ventilation fan. Recommendation to clean the system and clear of any cotton that has been trapped in the internal components by means of maybe compressed air.":
        "我们怀疑棉絮已填充至系统内部，堵塞电气元件。系统运行时因过热短暂停机，可能是棉絮堵塞通风扇所致。建议对系统进行清洁，采用压缩空气等方式清除内部元件中滞留的棉絮。",

    "2.  Longer period (3-5months) of operation.":
        "3.  延长运行观察周期（3–5 个月）",

    "As changes are minute and diluted due to the scale of HyESys machine to the electrical Main Grid. The transitional effect at section 3 might not be captured in detailed. A longer period of operation (3-5months) is recommended to observe a normalized changes. ":
        "由于 HyESys 设备规模相对于电气总网较小，变化微小且被稀释，第3节的过渡效果可能未能详细捕获。建议延长运行周期至 3–5 个月，以观察归一化后的变化规律。",

    "Prepared by AST — Advancer Smart Technology Pte Ltd  |  12 May 2026  |  Confidential":
        "编制单位：AST — Advancer Smart Technology Pte Ltd  |  2026年5月12日  |  机密",
}

# Table cell translations — key: original cell text (stripped), value: Chinese
CELL_MAP = {
    # Table 0 — Key Findings (each bullet is a separate paragraph in the cell)
    "Key Findings":
        "主要发现",
    "■  MSB1 incomer current drops 93–130 A (3-phase RMS) each time HyESys activates":
        "■  每次 HyESys 激活，MSB1 进线电流下降 93–130 A（三相均方根值）",
    "■  I²R loss reduction fraction: 5.2–7.5%, mean 6.7%":
        "■  I²R 损耗削减比例：5.2–7.5%，均值 6.7%",
    "■  Direct observed kW reduction at MSB1 meter: 35 kW per activation":
        "■  MSB1 电表直接观测有功功率降低：每次激活 35 kW",
    "■  Grid utility meter improves: −16 kW, −20 kVAr, PF +0.009 ":
        "■  电网电表改善：−16 kW、−20 kVAr，PF +0.009",
    "■  Conservative annual savings: CNY 257,000 (~SGD 54,000)":
        "■  保守年节省：CNY 257,000（约 SGD 54,000）",
    "■  Upper-bound annual savings: CNY 390,000 (I²R fraction method)":
        "■  年节省上限：CNY 390,000（I²R 比例法）",
    "■  Extreme harmonic content (THD ~142% Phase C) — HyESys compensates both reactive AND harmonic currents":
        "■  谐波含量极高（C 相 THD ~142%）—— HyESys 同时补偿无功电流和谐波电流",

    # Table 1 — Baseline stats (label column)
    "Active power (mean)":        "有功功率（均值）",
    "Reactive power (mean)":      "无功功率（均值）",
    "Per-phase current A / B / C":"各相电流 A / B / C",
    "3-Phase I_rms combined":     "三相 I_rms 合计",
    "Displacement PF  (kW/√(kW²+kVAr²))":
                                  "位移功率因数（kW/√(kW²+kVAr²)）",
    "True (apparent) PF  (kW/V×I)  A/B/C":
                                  "真实（视在）功率因数（kW/V×I）A/B/C",
    "Implied THD  A / B / C":     "推算谐波畸变率 A / B / C",

    # Table 2 — Transition events (header row)
    "Event":        "事件",
    "I_before (A)": "补偿前电流 (A)",
    "I_after (A)":  "补偿后电流 (A)",
    "ΔI (A)":       "ΔI (A)",
    "Fraction":     "削减比例",
    "ΔkW (direct)": "ΔkW（直接）",
    "Mean":         "均值",

    # Table 3 — Savings (header + method column)
    "Method":   "方法",
    "kW Saved": "节省 kW",
    "kWh/month\n(720 h)":          "kWh/月\n（720 h）",
    "CNY/month\n(¥0.85/kWh)":     "CNY/月\n（¥0.85/kWh）",
    "CNY/year": "CNY/年",
    "Conservative\n(Direct observed ΔkW)":
                                   "保守估算\n（直接观测 ΔkW）",
    "Upper bound\n(I²R fraction × kW_baseline)":
                                   "上限估算\n（I²R 比例 × kW 基线）",
}


# ── Font helpers ───────────────────────────────────────────────────────────────

def apply_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)


def set_para_text(para, text, size_pt, bold=False, italic=False, color=None):
    """Replace all runs in *para* with a single run containing *text*."""
    # Preserve paragraph-level XML (spacing, borders, shading) but clear runs
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    run = para.add_run(text)
    apply_font(run, size_pt, bold=bold, italic=italic, color=color)


# ── Size selection ─────────────────────────────────────────────────────────────

def size_for_style(style_name: str, text: str) -> tuple:
    """Return (size_pt, bold) based on paragraph style."""
    s = style_name
    if "Title" in s:
        return 20, True
    if "Heading 1" in s:
        return 14, True
    if "Heading 2" in s:
        return 13, True
    # Footer-like lines
    if "Confidential" in text or ("Prepared by" in text and "|" in text):
        return 9, False
    # Figure captions / notes
    if text.startswith("Figure") or text.startswith("图") or text.startswith("Note:") or text.startswith("注："):
        return 10, False
    return 11, False


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    print("Loading document …")
    doc = Document(SRC)

    # ── Paragraphs ──────────────────────────────────────────────────────────
    print("Translating paragraphs …")
    for para in doc.paragraphs:
        original = para.text  # keep original spacing/newlines for map lookup
        stripped = original.strip()
        if not stripped:
            continue

        cn_text = PARA_MAP.get(stripped) or PARA_MAP.get(original)
        if cn_text is None:
            # Try partial match for lines that may have trailing whitespace variants
            for k, v in PARA_MAP.items():
                if stripped == k.strip():
                    cn_text = v
                    break

        if cn_text:
            size_pt, bold = size_for_style(para.style.name, cn_text)
            set_para_text(para, cn_text, size_pt, bold=bold)
            print(f"  ✓  {cn_text[:55]}{'…' if len(cn_text) > 55 else ''}")
        else:
            # Paragraph not in map — just reformat font without changing text
            size_pt, bold = size_for_style(para.style.name, stripped)
            for run in para.runs:
                if run.text.strip():
                    apply_font(run, size_pt, bold=bold)
            print(f"  –  (kept as-is) {stripped[:55]}")

    # ── Tables ───────────────────────────────────────────────────────────────
    print("\nTranslating tables …")
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    stripped = original.strip()
                    if not stripped:
                        continue
                    cn_text = CELL_MAP.get(stripped) or CELL_MAP.get(original.strip())
                    if cn_text is None:
                        for k, v in CELL_MAP.items():
                            if stripped == k.strip():
                                cn_text = v
                                break
                    if cn_text:
                        set_para_text(para, cn_text, 10)
                        print(f"  [T{ti}] ✓  {cn_text[:50]}")
                    else:
                        for run in para.runs:
                            if run.text.strip():
                                apply_font(run, 10)

    print(f"\nSaving → {DST}")
    doc.save(DST)
    print("Done ✓")


if __name__ == "__main__":
    main()
