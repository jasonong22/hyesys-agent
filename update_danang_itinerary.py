"""
Generates updated Danang + Hoi An itinerary Word doc.
Changes from v1:
1. Night 2 moved to Mercure Ba Na Hills (overnight) instead of Da Nang hotel
2. Day 3 morning: sunrise at Golden Bridge (exclusive hotel guest access)
3. Danang-first order confirmed as optimal (no change to order)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\JasonOng\Desktop\local docs\personal\viet\Danang_HoiAn_4D1M_Aug2026_v2.docx"

# ── helpers ────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def title_block(doc, title, subtitle, dates, img_caption=""):
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(24)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    d = doc.add_paragraph(dates)
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in d.runs:
        r.font.size = Pt(11)
        r.bold = True

def h1(doc, text, color=None):
    p = doc.add_heading(text, level=1)
    c = color or RGBColor(0x1F, 0x49, 0x7D)
    for r in p.runs:
        r.font.color.rgb = c
        r.font.size = Pt(14)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        r.font.size = Pt(12)
    return p

def body(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r1 = p.add_run(bold_prefix + "  ")
        r1.bold = True
        r1.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p

def day_header(doc, day_num, date, title, subtitle=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    shading_elm = OxmlElement('w:pPr')
    r1 = p.add_run(f"DAY {day_num}  —  {date}")
    r1.font.size = Pt(13)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Add background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F497D')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.1)

    p2 = doc.add_paragraph(title)
    for r in p2.runs:
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    if subtitle:
        p3 = doc.add_paragraph(subtitle)
        for r in p3.runs:
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def period(doc, icon, label):
    p = doc.add_paragraph()
    r = p.add_run(f"{icon}  {label}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC5, 0x5A, 0x11)
    p.paragraph_format.space_before = Pt(6)

def night_box(doc, night_num, hotel, note=""):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'DEEAF1')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    r1 = p.add_run(f"  TONIGHT — NIGHT {night_num}  |  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(hotel)
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x00, 0x50, 0x00)
    if note:
        r3 = p.add_run(f"\n  {note}")
        r3.font.size = Pt(9.5)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def cost_line(doc, label, cost):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(cost)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x20, 0x60, 0x20)

def tip(doc, text):
    p = doc.add_paragraph()
    r1 = p.add_run("Tip: ")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x7F, 0x6B, 0x00)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x7F, 0x6B, 0x00)

def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDD7EE')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Build document ─────────────────────────────────────────────────────

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── Cover ──────────────────────────────────────────────────────────────
title_block(doc,
    "DANANG & HOI AN",
    "4 Full Days + 1 Morning  •  4 Nights  •  2 Pax\n1 Night Da Nang  •  1 Night Ba Na Hills  •  2 Nights Hoi An",
    "27 – 31 August 2026  |  v2 — Ba Na Hills Overnight Edition\nScoot TR510 (SIN→DAD)  •  VietJet VJ889 (DAD→SIN)  •  Booking Ref: NBT4NW")
doc.add_paragraph()

# ── Planning Notes ────────────────────────────────────────────────────
h1(doc, "PLANNING NOTES — KEY CHANGES FROM v1")
h2(doc, "Change 1: Ba Na Hills — Overnight Stay Added")
body(doc, "The original plan did Ba Na Hills as a day trip and returned to Da Nang hotel. The revised plan checks in to the Mercure Danang French Village Bana Hills (the only on-mountain hotel) on Night 2, allowing:")
bullet(doc, "Exclusive early access to the Golden Bridge at sunrise (~05:45–06:30) before the first public cable car arrives at 06:00. Day-trippers begin arriving after 10:00 — you have the bridge entirely to yourselves.")
bullet(doc, "Ba Na Hills Laser Light Show at 19:00 and 21:00 (only accessible to overnight guests after the park closes to day visitors).")
bullet(doc, "20–25% discount on Ba Na Hills entry tickets for hotel guests.")
bullet(doc, "No wasted round-trip transport on Day 2 — the hotel is on the mountain.")
bullet(doc, "Cooler, mist-covered Golden Bridge at dawn — the most photographed experience at Ba Na Hills.")
body(doc, "Hotel: Mercure Danang French Village Bana Hills — 4-star, 1,487m elevation, ~SGD 115–175/night. World's Leading Themed Resort, World Travel Awards 2025.", size=10)

doc.add_paragraph()
h2(doc, "Change 2: Danang First vs Hoi An First — Verdict")
body(doc, "Danang first (the current order) is confirmed as optimal. Reasons:")
bullet(doc, "Da Nang International Airport is in Da Nang — zero extra travel on arrival day. Going to Hoi An first would add 30 km on Day 1 when you are already tired from flying.")
bullet(doc, "Dragon Bridge fire show is Friday at 21:00 — Day 3 of this itinerary. Friday is the least-crowded show night (busyness 68/100 vs Saturday 100/100). Perfect alignment.")
bullet(doc, "Ba Na Hills is 25 km northwest of Da Nang — a Da Nang base makes this excursion natural. From Hoi An it would be 55+ km each way.")
bullet(doc, "Geographic flow is logical: Da Nang (arrive) → Ba Na Hills (northwest) → Marble Mountains (south of Da Nang) → Hoi An (further south) → Da Nang Airport (depart). No backtracking.")
bullet(doc, "Hoi An's craft villages and cooking school work best saved for Day 4 — the full day in town after An Bang Beach.")
body(doc, "Conclusion: No change to the Danang-first order. It is the correct and most logical sequence.", bold=True, color=RGBColor(0x20, 0x60, 0x20))

rule(doc)

# ── Flights ────────────────────────────────────────────────────────────
h1(doc, "FLIGHTS  —  CONFIRMED BOOKINGS")
body(doc, "Booking Reference: NBT4NW  (both passengers — Jason Ong & Rachel Teo)", bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
doc.add_paragraph()

h2(doc, "Outbound  —  Scoot TR510  |  27 Aug 2026")
t_out = doc.add_table(rows=4, cols=2)
t_out.style = 'Table Grid'
for ri, (lbl, val) in enumerate([
    ("Flight", "Scoot TR510"),
    ("Departure", "08:10  —  Singapore Changi Airport Terminal 1"),
    ("Arrival", "09:55  —  Da Nang International Airport Terminal 2"),
    ("Duration", "~1h 45m direct"),
]):
    t_out.cell(ri, 0).text = lbl
    t_out.cell(ri, 1).text = val
    set_cell_bg(t_out.cell(ri, 0), 'BDD7EE')
    if t_out.cell(ri, 0).paragraphs[0].runs:
        t_out.cell(ri, 0).paragraphs[0].runs[0].bold = True
        t_out.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(10)
    if t_out.cell(ri, 1).paragraphs[0].runs:
        t_out.cell(ri, 1).paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()
bullet(doc, "Arrive at Changi T1 by 06:00 (check-in closes ~07:40). Leave home no later than 05:30.")
bullet(doc, "Baggage: NO free checked baggage — carry-on only. 1 cabin bag (max 54×38×23 cm) + 1 personal item (max 40×30×10 cm). Total weight ≤ 10 kg per person.")

h2(doc, "Return  —  VietJet Air VJ889  |  31 Aug 2026")
t_ret = doc.add_table(rows=4, cols=2)
t_ret.style = 'Table Grid'
for ri, (lbl, val) in enumerate([
    ("Flight", "VietJet Air VJ889"),
    ("Departure", "14:15  —  Da Nang International Airport Terminal 2"),
    ("Arrival", "18:10  —  Singapore Changi Airport Terminal 4"),
    ("Duration", "~1h 55m direct"),
]):
    t_ret.cell(ri, 0).text = lbl
    t_ret.cell(ri, 1).text = val
    set_cell_bg(t_ret.cell(ri, 0), 'BDD7EE')
    if t_ret.cell(ri, 0).paragraphs[0].runs:
        t_ret.cell(ri, 0).paragraphs[0].runs[0].bold = True
        t_ret.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(10)
    if t_ret.cell(ri, 1).paragraphs[0].runs:
        t_ret.cell(ri, 1).paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()
bullet(doc, "Arrive at Da Nang Airport by 11:15 (3h before departure recommended). Leave Hoi An ~10:30.")
bullet(doc, "Baggage: NO free checked baggage — carry-on only. 1 cabin bag (max 56×36×23 cm) + 1 personal item. Total weight ≤ 7 kg per person.")

p_warn = doc.add_paragraph()
pPr = p_warn._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'FCE4D6')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
r_warn = p_warn.add_run("  ⚠  BAGGAGE: Both flights have ZERO free checked baggage. Scoot carry-on limit = 10 kg; VietJet return limit = 7 kg per person. For a 5-day trip with market shopping, add checked baggage now — airport add-ons are 2–3× more expensive. Scoot add-on ~SGD 30–50/person; VietJet add-on ~SGD 25–40/person. Add via Trip.com (Ref: NBT4NW) or each airline's app.")
r_warn.bold = True
r_warn.font.size = Pt(10)
r_warn.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()
body(doc, "Visa: Singapore passport holders — 45-day visa-free entry to Vietnam. No action needed.")

rule(doc)

# ── Accommodation ──────────────────────────────────────────────────────
h1(doc, "ACCOMMODATION")
h2(doc, "Night 1  —  Da Nang  (27 Aug)")
bullet(doc, "LUXURY — Hyatt Regency Danang Resort & Spa  |  SGD 220–340/night  |  My Khe Beach, 5-star beachfront, pool, spa. Rating 4.7/5.")
bullet(doc, "LUXURY — Fusion Maia Danang  |  SGD 250–380/night  |  All-inclusive spa, private pool in every room.")
bullet(doc, "MID-RANGE — Novotel Danang Premier Han River  |  SGD 100–160/night  |  Han River waterfront, infinity rooftop pool.")
bullet(doc, "MID-RANGE — Azura Da Nang Hotel  |  SGD 80–130/night  |  My Khe Beach, pool, gym, breakfast included.")
bullet(doc, "BUDGET — HAIAN Beach Hotel & Spa  |  SGD 45–65/night  |  Directly across My Khe Beach. Rated 9.8/10 Booking.com (Exceptional). Rooftop pool, 24h front desk, luggage storage. One of Da Nang's highest-rated budget properties. Breakfast available add-on.")
bullet(doc, "BUDGET — Sala Danang Beach Hotel  |  SGD 50–70/night  |  My Khe Beach strip. Highly rated on TripAdvisor and Booking.com. Pool, modern rooms, luggage storage confirmed. Short walk to beach.")
bullet(doc, "BUDGET — Monarque Hotel Danang  |  SGD 50–65/night  |  My Khe Beach area. Rated 8.9/10 on Booking.com. Clean, well-maintained, 24h reception, luggage storage. Good value for a 1-night stay.")
bullet(doc, "BUDGET — Muong Thanh Luxury Da Nang Hotel  |  SGD 45–60/night  |  Central Da Nang. Reputable Vietnamese 4-star chain with strong local reviews. Pool, gym, luggage storage, breakfast included.")
tip(doc, "Only 1 night here — no need to overspend. HAIAN Beach Hotel is the standout budget pick: exceptional reviews, beachfront, and full luggage storage. Book early for August peak season.")

h2(doc, "Night 2  —  Ba Na Hills  (28 Aug)  ★ NEW")
bullet(doc, "★ Mercure Danang French Village Bana Hills  |  ~SGD 115–175/night  |  THE only hotel on Ba Na Hills mountain at 1,487m elevation.")
bullet(doc, "4-star. Nestled within the French Village of Sun World. 14-min walk to Golden Bridge.")
bullet(doc, "World's Leading Themed Resort — World Travel Awards 2025.")
bullet(doc, "Guests get: exclusive early access to Golden Bridge before cable car opens, Laser Light Show at 19:00 & 21:00, 20–25% discount on park entry tickets.")
bullet(doc, "Book direct at mercure-danang-banahills-french-village.com or via Booking.com / Agoda.")
tip(doc, "Book this hotel BEFORE buying Ba Na Hills tickets — the hotel discount covers the cost difference.")
p_bana = doc.add_paragraph()
pPr_bana = p_bana._p.get_or_add_pPr()
shd_bana = OxmlElement('w:shd')
shd_bana.set(qn('w:fill'), 'FFF2CC')
shd_bana.set(qn('w:val'), 'clear')
pPr_bana.append(shd_bana)
r_bana = p_bana.add_run("  ★  Ba Na Hills accommodation note: Mercure is the ONLY hotel on the mountain. All other cheaper options (Ebisu Onsen Resort, La Retreat, etc.) are at the base of Ba Na Hills — they do not have cable car access or Golden Bridge sunrise privileges. If budget is the constraint, the alternative is to do Ba Na Hills as a day trip and skip the overnight. However, you will lose the sunrise access which is the highlight of this itinerary.")
r_bana.font.size = Pt(10)
r_bana.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
doc.add_paragraph()
bullet(doc, "BASE-OF-MOUNTAIN ALTERNATIVE (no sunrise access) — Ebisu Onsen Resort  |  SGD 75–100/night  |  3.5-star, natural hot spring onsen, water park, outdoor pools. 10–15 min Grab from Ba Na Hills cable car base. Good reviews. Opt for this only if budget is a hard constraint and sunrise access is not a priority.")

h2(doc, "Nights 3 & 4  —  Hoi An  (29–31 Aug)")
bullet(doc, "LUXURY — La Siesta Hoi An Resort & Spa  |  SGD 180–280/night  |  TripAdvisor 2026: #1 Vietnam, #2 Asia, #6 World. 200m² infinity pool.")
bullet(doc, "LUXURY — Little Hoi An Boutique Hotel & Spa  |  SGD 130–200/night  |  Old Town edge, TripAdvisor Traveller's Choice.")
bullet(doc, "MID-RANGE — RiverTown Hoi An Resort & Spa  |  SGD 90–150/night  |  An Hoi Island, river views, short walk to Old Town.")
bullet(doc, "MID-RANGE — The Silk River Hotel & Spa  |  SGD 70–120/night  |  Riverfront, bicycles provided free.")
bullet(doc, "BUDGET — Cozy Savvy Boutique Hotel  |  SGD 40–60/night  |  Rooftop pool, free bikes, near Old Town. Highly praised on Booking.com and TripAdvisor for service, cleanliness, and location. Luggage storage at 24h front desk. Excellent value for 2 nights.")
bullet(doc, "BUDGET — Riverside White House Boutique Hotel  |  SGD 35–55/night  |  Rated 9.6/10 on Booking.com. Pool, terrace, 24h airport shuttle, free toiletries. Guests consistently cite safe, welcoming environment and helpful luggage storage. River-view rooms available.")
bullet(doc, "BUDGET — Hai Au Boutique Hotel & Spa  |  SGD 50–70/night  |  Rated 8.8/10. Outdoor pool, on-site spa, good location. Luggage storage available. Well-reviewed for cleanliness and safety.")
bullet(doc, "BUDGET — Cozy An Boutique Hotel  |  SGD 45–65/night  |  4-star, pool, quiet area with Thu Bon river views, 10-min walk to Old Town. Free bikes. Strong guest reviews for safety and staff responsiveness. Good for couples.")

body(doc, "Luxury combo: Hyatt Regency Da Nang + Mercure Ba Na Hills + La Siesta Hoi An.", bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
body(doc, "Budget combo: HAIAN Beach Hotel Da Nang + Mercure Ba Na Hills + Cozy Savvy / Riverside White House Hoi An.", bold=True, color=RGBColor(0x20, 0x60, 0x20))
body(doc, "Note: Mercure Ba Na Hills has no budget alternative on-mountain — keeping it is necessary for the sunrise experience that anchors Day 3.", italic=True, size=10, color=RGBColor(0x60, 0x60, 0x60))

rule(doc)

# ── DAY 1 ──────────────────────────────────────────────────────────────
day_header(doc, 1, "Wednesday, 27 August 2026",
           "FLY SIN → DAD  |  Afternoon in Da Nang",
           "Intensity: LOW  |  Arrival day — settle in, explore the waterfront")

period(doc, "☀", "MORNING — FLIGHT")
bullet(doc, "Leave home by 05:30. Arrive Changi Airport Terminal 1 by 06:00.", bold_prefix="05:30")
bullet(doc, "Scoot TR510 departs 08:10 (Booking Ref: NBT4NW). Check in online 48h before to secure seats together.", bold_prefix="08:10")
bullet(doc, "Flight duration ~1h 45m direct. Arrive Da Nang Airport Terminal 2 at 09:55.")
tip(doc, "Terminal 1 has Toast Box and Ya Kun open from 05:30 — grab breakfast airside after clearing security.")

period(doc, "☀", "LATE MORNING — ARRIVAL")
bullet(doc, "Land at Da Nang International Airport T2 (09:55). Singapore passports — visa-free, immigration ~15 min.", bold_prefix="09:55")
bullet(doc, "Grab Car to hotel: ~80,000–120,000 VND (~SGD 5–7), ~15–20 min. Drop bags — rooms ready from 14:00.", bold_prefix="10:30")
bullet(doc, "Lunch: Roly Poly Fresh Spring Rolls (40 Vo Nguyen Giap) — lemongrass beef, prawn and fresh rolls, Da Nang style. ~SGD 8–12/person. 5-min walk from My Khe Beach.", bold_prefix="11:30")

period(doc, "☀", "AFTERNOON")
bullet(doc, "Museum of Cham Sculpture (2 Thang 9 St, 5 min Grab from hotel) — the only museum in the world dedicated entirely to Cham civilisation art. Over 2,000 artefacts: intricate stone carvings, altars, apsara dancers. Entry ~40,000 VND. Allow 1–1.5h. Genuinely unique — nothing like this anywhere else in Vietnam.", bold_prefix="13:00")
bullet(doc, "Da Nang Cathedral (Nha Tho Con Ga / Pink Church) — a quick 10-min stop en route back. Built 1923, iconic pink facade and rooster weathervane. Good photo stop.")
bullet(doc, "Hotel room ready by ~15:00. Check in, freshen up, change into swimwear.", bold_prefix="15:00")
bullet(doc, "My Khe Beach — one of Asia's finest urban beaches. 30 km of white sand, calm warm water in August. Hire a sunbed (~50,000–80,000 VND), swim, decompress from travel. Best 15:30–17:30 when the heat softens.", bold_prefix="15:30")

period(doc, "🌙", "EVENING")
bullet(doc, "Grab to Dragon Bridge area (~SGD 3) — walk the Bach Dang riverside promenade.")
bullet(doc, "Dinner: Madame Lan Restaurant (4 Bach Dang, Han River) — Grilled Squid, Banh Xeo, courtyard ambiance. ~SGD 20–30/person.")
bullet(doc, "Admire Dragon Bridge illuminated — beautifully lit every evening. No fire show tonight (Wednesday). The Friday fire show is saved for Day 3.")
bullet(doc, "Return to hotel by 22:30. Early night — Ba Na Hills cable car first thing tomorrow.")
cost_line(doc, "Day 1 Transport:", "Airport Grab ~SGD 5–7  |  Han River Grab ~SGD 3 each way")
cost_line(doc, "Day 1 Est. Cost:", "~SGD 80–130/person (Scoot TR510 flight share + meals + transport)")
night_box(doc, "1", "Da Nang Hotel  (Hyatt Regency / Novotel / Azura)",
          "Check-in from 14:00. Last night at this hotel — check out tomorrow morning for Ba Na Hills.")

rule(doc)

# ── DAY 2 ──────────────────────────────────────────────────────────────
day_header(doc, 2, "Thursday, 28 August 2026",
           "SON TRA PENINSULA (MORNING)  →  BA NA HILLS OVERNIGHT  ★  LASER LIGHT SHOW",
           "Intensity: MEDIUM  |  Da Nang morning → check out → Ba Na Hills → Mercure overnight")

period(doc, "☀", "EARLY MORNING — SON TRA PENINSULA & LINH UNG PAGODA")
bullet(doc, "Breakfast at hotel (07:00). Pack your overnight bag for Ba Na Hills — main luggage can be stored with the hotel front desk until you check out later.", bold_prefix="07:00")
bullet(doc, "Grab Car to Son Tra Peninsula (~15 min, ~80,000–100,000 VND). Son Tra is a lush protected jungle headland jutting into the sea — the lungs of Da Nang.", bold_prefix="07:45")
bullet(doc, "Linh Ung Pagoda & Lady Buddha: Vietnam's tallest Lady Buddha statue at 67 m — visible from across Da Nang. The pagoda sits on the Son Tra cliff edge with sweeping views of the coastline, My Khe Beach, and the city. Free entry. Allow 1–1.5h.")
bullet(doc, "Son Tra viewpoints: Scan for the rare Red-shanked Douc langur monkeys in the forest canopy along the road. Often spotted early morning.")
bullet(doc, "Grab back to hotel (~10:00).", bold_prefix="10:00")
tip(doc, "Son Tra is at its best before 09:30 — quiet, misty forest views, no tour buses. This pairs perfectly with the early hotel breakfast.")

period(doc, "☀", "MID-MORNING — CHECK OUT & HEAD TO BA NA HILLS")
bullet(doc, "Check out of Da Nang hotel. Pack for 1 night on the mountain.", bold_prefix="10:30")
bullet(doc, "Grab Car to Ba Na Hills cable car base station (25 km northwest). ~350,000–450,000 VND (~SGD 20–26), ~40 min.", bold_prefix="11:00")
bullet(doc, "Cable car up: world-record-breaking 5.1 km gondola, gliding above waterfalls and misty forest canopy. Arrive summit ~12:00.", bold_prefix="11:45")
bullet(doc, "Check in to Mercure Danang French Village Bana Hills. Drop luggage, freshen up. Rooms available from 14:00 — leave bags and explore.")
tip(doc, "As a hotel guest you get 20–25% off Ba Na Hills entry tickets — buy at the Mercure front desk, not via Klook.")

period(doc, "☀", "AFTERNOON")
bullet(doc, "Golden Bridge — the iconic 150m bridge held by two giant stone hands. Best photos before crowds build (~10:00). Dramatic mist and cloud from early afternoon.")
bullet(doc, "French Village — neo-Gothic architecture, wine cellar, manicured flower gardens. Temperature ~18–22°C at 1,487m. Bring a light layer.")
bullet(doc, "Fantasy Park (included in ticket): roller coaster, wax museum, 4D cinema. Allow 1–2 hours.")
bullet(doc, "Lunch: Buffet at Ba Na Hills (combo ticket includes buffet ~1,300,000 VND) or à la carte at Fantasy Park food court. ~SGD 15–30/person.")

period(doc, "🌙", "EVENING — EXCLUSIVE TO HOTEL GUESTS")
bullet(doc, "After day-trippers leave (~17:00–18:00), the park becomes quiet. Stroll the French Village and Golden Bridge without crowds.")
bullet(doc, "★  Ba Na Hills Laser Light Show — 19:00 (first show). Spectacular light-and-sound performance. Only available to overnight guests after park closes.")
bullet(doc, "★  Second show at 21:00. Watch from the French Village amphitheatre.")
bullet(doc, "Dinner at Mercure hotel restaurant — French-Vietnamese cuisine. ~SGD 25–40/person.")
bullet(doc, "Early rest — wake-up call at 05:15 tomorrow for sunrise.")
tip(doc, "Set an alarm for 05:15. The sunrise at the Golden Bridge is the highlight of this entire trip.")
cost_line(doc, "Day 2 Transport:", "Son Tra Grab return ~SGD 9–12  |  Da Nang → Ba Na Hills ~SGD 20–26 (no return — sleeping on mountain)")
cost_line(doc, "Day 2 Est. Cost:", "~SGD 200–280/person (Mercure hotel + Ba Na Hills ticket + meals + transport)")
night_box(doc, "2", "Mercure Danang French Village Bana Hills  ★",
          "1,487m elevation. Laser Light Show at 19:00 & 21:00. Set alarm 05:15 for Golden Bridge sunrise.")

rule(doc)

# ── DAY 3 ──────────────────────────────────────────────────────────────
day_header(doc, 3, "Friday, 29 August 2026",
           "GOLDEN BRIDGE SUNRISE  →  MARBLE MOUNTAINS  →  DRAGON BRIDGE FIRE SHOW  →  HOI AN",
           "Intensity: MEDIUM-HIGH  |  Long day — worth every minute  |  Check out Ba Na Hills → arrive Hoi An late")

period(doc, "★", "EARLY MORNING — GOLDEN BRIDGE SUNRISE (05:30–07:30)")
bullet(doc, "Wake up 05:15. Be at Golden Bridge by 05:45.", bold_prefix="05:15")
bullet(doc, "As hotel guests you have EXCLUSIVE access to the bridge before the first public cable car arrives at 06:00. The bridge is yours alone — no crowds, no tour buses.", bold_prefix="05:45")
bullet(doc, "Watch sunrise over the Truong Son mountains and Da Nang coastline from the golden arc. Mist rolls through the giant stone hands in the early morning light.")
bullet(doc, "Best photography window: 05:45–07:00. Soft golden light, wispy mountain fog, dramatic clouds.")
bullet(doc, "Return to Mercure for breakfast (07:30). Hot buffet breakfast included with room.")
tip(doc, "This is the single best photo opportunity in all of Central Vietnam. Bring your camera. The bridge is 150m long — walk to the far end for the full view.")

period(doc, "☀", "MORNING — CHECK OUT & LAST LOOK")
bullet(doc, "After breakfast, explore Ba Na Hills one final time — the French Village flower gardens are beautiful in morning light without the crowds.")
bullet(doc, "Check out of Mercure ~09:30–10:00. Leave luggage with hotel concierge if needed for a final walk.")
bullet(doc, "Cable car down ~10:30. Grab Car waiting at base station.")
bullet(doc, "Luggage storage for the day: Two options — (1) call your Night 1 Da Nang hotel in advance; most 4-star hotels hold luggage for 24–48h for prior guests — collect when passing through Da Nang this afternoon; or (2) WhaleLO luggage storage: 20+ drop-off locations in Da Nang, ~SGD 3–5/bag/day, 24/7, bookable via Klook. Book in advance for peace of mind.")

period(doc, "☀", "LATE MORNING / AFTERNOON — MARBLE MOUNTAINS")
bullet(doc, "Grab to Marble Mountains / Ngu Hanh Son — ~45–50 min from Ba Na Hills base (Ba Na Hills is 25–35 km west of Da Nang; Marble Mountains is 9 km south of Da Nang — ~34 km total, driving east through Da Nang then south). ~350,000–450,000 VND (~SGD 20–26).")
bullet(doc, "Five limestone formations named after the five elements. Explore Thuy Son (Water Mountain) — entry ~40,000 VND, elevator ~15,000 VND.")
bullet(doc, "Huyen Khong Cave — natural skylights pierce limestone, illuminating ancient Buddha shrines. Atmospheric and spiritual.")
bullet(doc, "Panoramic views from summit over My Khe Beach and the South China Sea. Allow 1.5–2 hours.")
bullet(doc, "Quick lunch nearby: Bun Cha Ca (Da Nang fish cake noodle soup) at local eateries. ~30,000–50,000 VND.")

period(doc, "☀", "AFTERNOON — DA NANG CITY")
bullet(doc, "Return to Da Nang city (~13:30–14:00). Visit Han Market (Cho Han, Tran Phu St) — silk scarves, lacquerware, Vietnamese coffee, dried goods. Cash only. Allow 1 hour.")
bullet(doc, "Stroll Bach Dang promenade, Han Riverfront. Coffee at riverside café (~15:00–17:00).")
bullet(doc, "No hotel room in Da Nang tonight — freshen up at a riverside café or restaurant. Most mid-range establishments near Bach Dang / Han River will let you use their facilities. Plan a 30–45 min rest stop before the Dragon Bridge show.")

period(doc, "🌙", "EVENING — DRAGON BRIDGE FIRE SHOW")
bullet(doc, "Dinner at Madame Lan (4 Bach Dang) or Bach Dang riverside restaurant (~18:30–20:00). Position yourself near Dragon Bridge viewing area.")
bullet(doc, "★  DRAGON BRIDGE FIRE & WATER SHOW — 21:00 sharp (Fri/Sat/Sun only). The 150m dragon breathes real fire then sprays water. Show runs to ~21:30. Free. Today is Friday — least-crowded show night (busyness 68/100 vs Saturday's 100/100). Perfect timing.")
bullet(doc, "After show (~21:30): Grab Car directly to Hoi An — ~30–40 min, ~250,000–320,000 VND (~SGD 15–19). Arrive Hoi An ~22:15–22:30.")
bullet(doc, "Late check-in at La Siesta Hoi An — pre-advise hotel of late arrival. Front desk 24h. Drop bags, rest. Full Hoi An day tomorrow.")
cost_line(doc, "Day 3 Transport:", "Ba Na Hills base → Marble Mtns ~SGD 20–26  |  Da Nang city → Hoi An Grab (late) ~SGD 15–19")
cost_line(doc, "Day 3 Est. Cost:", "~SGD 80–120/person (transport + Marble Mtns + meals + Dragon Bridge)")
night_box(doc, "3  (First Hoi An Night)", "La Siesta Hoi An Resort & Spa  (or chosen Hoi An hotel)",
          "Late arrival ~22:30. Inform hotel at booking. Full Hoi An day starts tomorrow morning.")

rule(doc)

# ── DAY 4 ──────────────────────────────────────────────────────────────
day_header(doc, 4, "Saturday, 30 August 2026",
           "AN BANG BEACH  |  COCONUT BASKET BOAT  |  CAM KIM ISLAND  |  COOKING CLASS  |  OLD TOWN EVENING",
           "Intensity: MEDIUM  |  Best day of the trip")

period(doc, "☀", "MORNING")
bullet(doc, "Breakfast at hotel. Bicycle to An Bang Beach — hire from hotel (~30,000–50,000 VND/hr). Scenic 3 km through rice paddies (~15 min).")
bullet(doc, "An Bang Beach — quieter than My Khe. Soft white sand, turquoise water, beach bar sun loungers (~50,000–100,000 VND). Excellent swimming in August.")
bullet(doc, "Beach eats: Soul Kitchen (Beach road, An Bang) — Australian-run bar, breakfasts, smoothies, fresh seafood. ~SGD 10–18/person.")
bullet(doc, "Spend 2–3 hours before heat peaks (~11:30).")

period(doc, "☀", "AFTERNOON")
bullet(doc, "Return to Old Town ~12:00. Banh Mi Phuong (2B Phan Chau Trinh) — Anthony Bourdain's famous banh mi. ~30,000–50,000 VND. Queue expected.", bold_prefix="12:00")
bullet(doc, "Cam Thanh Coconut Basket Boat — 3 km from Old Town, ~15 min Grab (~30,000–40,000 VND). Round woven bamboo basket boats on the river weaving through a dense coconut palm water forest. Local boatmen spin, dip, and dance the baskets on the water — genuinely fun and photogenic. ~30–40 min on the water. ~SGD 10–15/person. Book via Klook or walk-up at Cam Thanh village.", bold_prefix="13:00")
bullet(doc, "Cam Kim Island — from An Hoi Islet (5 min walk from Old Town), take a bumboat across (~20,000–30,000 VND, ~10 min). A quiet island of traditional craftsmen: woodcarving workshops, mat weaving looms, boat builders at work. Completely unaffected by tourism — locals still living the old Hoi An trades. Cycle or walk the island (~45–60 min). Back to town by ~15:00.", bold_prefix="14:00")
bullet(doc, "Vietnamese Cooking Class (15:00–18:30): Red Bridge Cooking School via Klook (~SGD 45–55 — includes market tour, cook 4 dishes, scenic boat ride to the school). Cao Lau, fresh spring rolls, Banh Xeo, White Rose dumplings.", bold_prefix="15:00")

period(doc, "🌙", "EVENING — HOI AN OLD TOWN")
bullet(doc, "Ancient Town entrance ticket ~120,000 VND (~SGD 7) at Le Loi Street booth — grants entry to 5 heritage sites.")
bullet(doc, "Japanese Covered Bridge (Chùa Cầu) — built 1590s by Japanese traders. Glows golden under lantern light after dark.")
bullet(doc, "White Rose Restaurant (533 Hai Ba Trung) — authentic White Rose Dumplings (Banh Bao Vac). Watch hand-folding. ~60,000–80,000 VND/plate.")
bullet(doc, "Dinner: Morning Glory Signature Restaurant (106 Nguyen Thai Hoc) — riverfront balcony, Cao Lau, Banh Xeo. ~SGD 20–35/person. Or Cargo Club for multi-level river views.")
bullet(doc, "Lantern-Making Workshop (pre-book): 30–45 min, ~SGD 6–9. Build your own silk lantern to take home.")
bullet(doc, "Float paper lanterns on the Thu Bon River from Cam Nam Bridge — ~SGD 1–2 each.")
cost_line(doc, "Day 4 Transport:", "Bicycle An Bang ~SGD 3–5  |  Cam Thanh Grab ~SGD 2  |  Cam Kim bumboat ~SGD 2–4  |  Old Town Grab ~SGD 2")
cost_line(doc, "Day 4 Est. Cost:", "~SGD 100–160/person (cooking class + coconut boat + meals + Old Town activities)")
night_box(doc, "4  (Final Night)", "La Siesta Hoi An Resort & Spa",
          "Check-out by 10:00. Grab to Da Nang Airport at 10:30 tomorrow — VietJet VJ889 departs 14:15.")

rule(doc)

# ── DAY 5 ──────────────────────────────────────────────────────────────
day_header(doc, 5, "Sunday, 31 August 2026",
           "RELAXED MORNING IN HOI AN  |  FLY HOME 14:15  |  ARRIVE SIN 18:10",
           "Intensity: LOW  |  VietJet VJ889 14:15 departure — more morning time than expected")

period(doc, "☀", "MORNING")
bullet(doc, "Breakfast at hotel (07:00). No rush — flight departs 14:15.", bold_prefix="07:00")
bullet(doc, "Old Town morning stroll — the Ancient Town at 08:00 is serene before tour groups arrive. Walk the Japanese Covered Bridge, Phuc Kien Assembly Hall (grand Fujian ancestral hall with incense smoke and ornate altars), and the quiet Thu Bon riverside. Free. The best light of the whole trip.", bold_prefix="08:00")
bullet(doc, "Com Linh Restaurant (Cam Pho area) — best Com Ga Hoi An (turmeric chicken rice). ~SGD 3–4. Go early before they sell out.", bold_prefix="09:00")
bullet(doc, "Hoi An Market (46 Tran Phu) — last chance for Vietnamese drip coffee packets, pho spice kits, dried jackfruit, lanterns. Budget ~SGD 20–40 for souvenirs.", bold_prefix="09:30–10:00")
bullet(doc, "Return to hotel, final pack. Check out by 10:00–10:15 (request late checkout if needed — most hotels allow to 11:00 free).", bold_prefix="10:00")
bullet(doc, "Grab Car Hoi An → Da Nang Airport T2 — ~250,000–320,000 VND (~SGD 15–19), ~30–40 min. Book in advance via Grab.", bold_prefix="10:30")
bullet(doc, "Arrive Da Nang Airport T2 by 11:10–11:15. Check in (3h before 14:15 as recommended). Drop checked bags if you purchased add-on.", bold_prefix="11:10")
bullet(doc, "Da Nang Airport is compact — security 20–30 min. Airside: small duty-free, coffee shop, basic food. Browse and relax.")
bullet(doc, "VietJet Air VJ889 departs 14:15. Arrive Singapore Changi Terminal 4 at 18:10.", bold_prefix="14:15")
tip(doc, "All Vietnamese dried goods, coffee packets, and cookies clear Singapore customs. Declare fresh fruit / durian if any.")
cost_line(doc, "Day 5 Transport:", "Hoi An → Da Nang Airport Grab ~SGD 15–19")
cost_line(doc, "Day 5 Est. Cost:", "~SGD 50–80/person (transport + last meals + airport + souvenirs)")

rule(doc)

# ── Budget Summary ─────────────────────────────────────────────────────
h1(doc, "BUDGET SUMMARY — PER PERSON (SGD)")
body(doc, "Exchange rate: 1 SGD ≈ 17,000 VND  |  Mid-range travel style", italic=True, size=10)

t = doc.add_table(rows=8, cols=3)
t.style = 'Table Grid'
headers = ["Category", "Budget", "Mid-Luxury"]
rows_data = [
    ("Flights (return, incl. baggage)", "SGD 180–320", "SGD 460–700"),
    ("Night 1 — Da Nang (HAIAN Beach / Sala / Novotel / Hyatt)", "SGD 45–80", "SGD 220–340"),
    ("Night 2 — Mercure Ba Na Hills (only on-mountain option)", "SGD 115–175", "SGD 115–175"),
    ("Nights 3–4 — Hoi An ×2 (Cozy Savvy / RiverTown / La Siesta)", "SGD 80–240", "SGD 360–560"),
    ("Activities (Ba Na Hills, Cham Museum, Old Town, cooking class, coconut boat)", "SGD 110–150", "SGD 150–190"),
    ("Food & drinks (5 days)", "SGD 100–150", "SGD 200–300"),
    ("Transport (all Grabs, 5 days)", "SGD 60–80", "SGD 80–100"),
]
for ci, h in enumerate(headers):
    cell = t.cell(0, ci)
    cell.text = h
    set_cell_bg(cell, '1F497D')
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
for ri, row in enumerate(rows_data, start=1):
    for ci, val in enumerate(row):
        cell = t.cell(ri, ci)
        cell.text = val
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        if ri % 2 == 0:
            set_cell_bg(cell, 'DEEAF1')

doc.add_paragraph()

# ── Practical Tips ─────────────────────────────────────────────────────
h1(doc, "PRACTICAL TIPS & ESSENTIALS")

h2(doc, "Weather & Packing")
bullet(doc, "Late August: 28–33°C, humidity 80–90%, afternoon showers. Pack light cotton/linen outfits, swimwear, SPF 50+ sunscreen, travel umbrella, portable fan, insect repellent.")
bullet(doc, "Ba Na Hills (~1,500m): 18–22°C — bring one light jacket or layer. This is significantly cooler than Da Nang sea level.")
bullet(doc, "Night 2 at Mercure: Pack a small overnight bag for Ba Na Hills only — leave main luggage at Da Nang hotel if preferred, or take it all up the mountain.")

h2(doc, "Checked Baggage Add-On — Action Required")
bullet(doc, "Both flights have ZERO free checked baggage. You need to add checked baggage before the trip — do not wait until check-in (airport add-on is 2–3x more expensive).")
bullet(doc, "Outbound Scoot TR510 (Booking Ref: NBT4NW): Add 20 kg checked bag via Trip.com or the Scoot app (scoot.com). ~SGD 30–50/person.")
bullet(doc, "Return VietJet VJ889 (Booking Ref: NBT4NW): Add 20 kg checked bag via Trip.com or VietJet app (vietjetair.com). ~SGD 25–40/person.")
bullet(doc, "Current carry-on limits: Outbound Scoot = 10 kg total; Return VietJet = 7 kg total. Vietnamese market souvenirs + 5 days of clothing will push this limit.")

h2(doc, "Money & Payments")
bullet(doc, "Withdraw VND at airport ATM (Vietcombank/Techcombank — low fees). Budget ~2,000,000 VND/day for 2 pax street food + transport.")
bullet(doc, "Cards at hotels and mid-range restaurants. Markets and street food = cash only. 1 SGD ≈ 17,000 VND.")

h2(doc, "Grab App")
bullet(doc, "Download before flying. Link a Singapore card. GrabCar for all intercity trips — fixed fare, no haggling.")
bullet(doc, "From Hoi An, Dragon Bridge is 45–55 min (~SGD 10–15). No public bus after 19:00 — private transfer (Grab) is only practical option for the fire show.")

h2(doc, "Ba Na Hills — Overnight Tips")
bullet(doc, "Book Mercure hotel well in advance (August = peak season). Request a room facing the French Village for best views.")
bullet(doc, "Golden Bridge sunrise: Be at the bridge by 05:45 — first public cable car arrives 06:00. Bring a jacket (cold and misty at dawn).")
bullet(doc, "Laser Light Show: 19:00 and 21:00. Dress warmly for the evening shows — temperature drops after sunset.")
bullet(doc, "Hotel discount on tickets: Present your Mercure room key at the ticket counter for 20–25% off.")

h2(doc, "Dragon Bridge Fire Show")
bullet(doc, "Runs every Friday, Saturday, Sunday at 21:00. Free from Bach Dang promenade.")
bullet(doc, "Friday is least crowded (busyness score 68/100 vs Saturday 100/100). This itinerary places you there on Friday — optimal.")
bullet(doc, "Show: dragon breathes fire for 2 min, then water spray for 3 min, cycling for ~30 min total until 21:30.")

h2(doc, "SIM Card")
bullet(doc, "Buy Viettel or Vinaphone tourist SIM at Da Nang Airport. 10-day, 20 GB: ~100,000–150,000 VND (~SGD 6–9). Essential for Grab and Google Maps.")

h2(doc, "Health & Safety")
bullet(doc, "Drink bottled water only. Carry: Panadol, Imodium, antihistamine, band-aids. Travel insurance strongly recommended.")
bullet(doc, "Emergency: 115 (ambulance), 113 (police).")

h2(doc, "Old Town Etiquette")
bullet(doc, "Dress modestly at temples and Japanese Bridge — shoulders and knees covered. Remove shoes before shrines.")

rule(doc)

# ── Food Guide ─────────────────────────────────────────────────────────
h1(doc, "FOOD GUIDE — DA NANG & HOI AN")
body(doc, "Quick-reference restaurant tables. All prices per person excluding drinks.", italic=True, size=10)

h2(doc, "Da Nang — Must-Eat Restaurants")
da_nang_headers = ["Restaurant", "Must-Order", "Price / Person", "Location"]
da_nang_rows = [
    ("Madame Lan", "Grilled squid, Banh Xeo, Com Hen (clam rice)", "SGD 20–30", "4 Bach Dang — Han River waterfront"),
    ("Roly Poly Fresh Spring Rolls", "Lemongrass beef rolls, prawn spring rolls", "SGD 8–12", "40 Vo Nguyen Giap — beachfront strip"),
    ("Mi Quang Ba Mua", "Mi Quang — Da Nang's signature peanut-broth noodle", "SGD 3–5", "19 Tran Binh Trong"),
    ("Bun Cha Ca street stalls", "Bun Cha Ca — fish cake noodle soup", "SGD 2–4", "Various — ask hotel for nearest stall"),
    ("Hai Hai Restaurant", "Mam Tom marinated BBQ (pungent shrimp paste)", "SGD 15–25", "100 Nguyen Chi Thanh"),
]
t1 = doc.add_table(rows=len(da_nang_rows)+1, cols=4)
t1.style = 'Table Grid'
for ci, h in enumerate(da_nang_headers):
    cell = t1.cell(0, ci)
    cell.text = h
    set_cell_bg(cell, '1F497D')
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
for ri, row in enumerate(da_nang_rows, start=1):
    for ci, val in enumerate(row):
        cell = t1.cell(ri, ci)
        cell.text = val
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        if ri % 2 == 0:
            set_cell_bg(cell, 'DEEAF1')
doc.add_paragraph()

h2(doc, "Hoi An — Must-Eat Restaurants")
hoi_an_headers = ["Restaurant", "Must-Order", "Price / Person", "Location"]
hoi_an_rows = [
    ("Banh Mi Phuong", "Signature Banh Mi (pork belly + pâté)", "SGD 2–3", "2B Phan Chau Trinh — Anthony Bourdain's pick"),
    ("White Rose Restaurant", "White Rose Dumplings (Banh Bao Vac) — watch hand-folding", "SGD 5–8", "533 Hai Ba Trung"),
    ("Morning Glory Signature", "Cao Lau (Hoi An noodles), Banh Xeo, fried wontons", "SGD 20–35", "106 Nguyen Thai Hoc — riverfront balcony"),
    ("Cargo Club", "French toast, river-view terrace (day); multi-level river dining (night)", "SGD 15–25", "107 Nguyen Thai Hoc"),
    ("Com Linh", "Com Ga Hoi An — Hoi An-style turmeric chicken rice", "SGD 3–4", "Cam Pho area — best before 09:00"),
    ("Soul Kitchen", "Full Australian-style breakfasts, fresh seafood, fruit smoothies", "SGD 10–18", "An Bang Beach road — beachfront"),
]
t2 = doc.add_table(rows=len(hoi_an_rows)+1, cols=4)
t2.style = 'Table Grid'
for ci, h in enumerate(hoi_an_headers):
    cell = t2.cell(0, ci)
    cell.text = h
    set_cell_bg(cell, '1F497D')
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
for ri, row in enumerate(hoi_an_rows, start=1):
    for ci, val in enumerate(row):
        cell = t2.cell(ri, ci)
        cell.text = val
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        if ri % 2 == 0:
            set_cell_bg(cell, 'DEEAF1')
doc.add_paragraph()

body(doc, "Street Food to Hunt Down", bold=True)
bullet(doc, "Banh Canh Cua (crab udon noodle soup) — Da Nang morning street stalls, ~SGD 2–3.")
bullet(doc, "Che (Vietnamese sweet soup dessert) — sold from push carts in Hoi An Old Town evenings, ~SGD 1.")
bullet(doc, "Ca Phe Trung (egg coffee) — Hoi An cafes, ~SGD 2–3. Rich custard egg yolk foam on dark Vietnamese coffee.")
bullet(doc, "Banh It La Gai (black sticky rice cake) — Hoi An local specialty, ~SGD 0.50 each at market stalls.")

rule(doc)

# ── Footer ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Itinerary for Jason Ong & Rachel Teo  |  4 Full Days + 1 Morning  |  Danang + Hoi An  |  27–31 August 2026\n"
    "Scoot TR510 (08:10 SIN→DAD)  •  VietJet VJ889 (14:15 DAD→SIN)  •  Booking Ref: NBT4NW\n"
    "v2 — Ba Na Hills Overnight Edition  |  Updated July 2026  |  1 SGD ≈ 17,000 VND  |  All prices approximate"
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save(OUT)
print(f"Saved: {OUT}")
